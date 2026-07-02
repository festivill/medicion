"""Informes de cargo/denuncia (PDF y Word) y tolerancias.

Extraído de app.py sin modificaciones (refactor modular v1.5).
"""
import hashlib

try:
    _real_md5 = hashlib.md5
    def patched_md5(data=b'', **kwargs):
        # Eliminamos el argumento 'usedforsecurity' si existe, porque Python 3.8 no lo entiende
        if 'usedforsecurity' in kwargs:
            del kwargs['usedforsecurity']
        return _real_md5(data, **kwargs)
    hashlib.md5 = patched_md5
except:
    pass


import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import json
import os
from datetime import datetime
import math
import base64
from io import BytesIO
import ctypes
import re
import platform
import subprocess
import traceback
import sqlite3, pathlib

from db import (
    _get_db_path, _init_db,
    db_buscar_funcionarios, db_guardar_funcionario,
    db_todos_funcionarios, db_eliminar_funcionario,
    _get_aduana_db_path, _init_aduana_db,
    db_get_aduanas, db_guardar_aduana, db_eliminar_aduana,
    db_get_lugares_operativos, db_guardar_lugar_operativo,
    db_eliminar_lugar_operativo,
    _init_funciones_db,
    db_get_funciones, db_guardar_funcion, db_eliminar_funcion,
)
from assets.icons import ICON_WINDOW_B64, ICON_REPORT_B64
from applog import app_dir as _app_dir, instalar_hooks as _instalar_hooks


class CargosMixin:
    def _generar_cargo_docx(self, ddt_obj, dif_kv, target_k,
                            ctrl_doc=True, ctrl_sal=False, ctrl_lab=False,
                            modo_comp_forzado=None, tipo_operacion="importacion",
                            tipo_operacion_info=None, output_folder=None):
        """Genera el informe de cargo/denuncia como documento Word (.docx)."""
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_ALIGN_VERTICAL
        except ImportError:
            messagebox.showerror("Error", "python-docx no está instalado.\nInstale con: pip install python-docx")
            return None

        if target_k == 0: return None
        diff_k = dif_kv - target_k
        permil_k = (diff_k / target_k * 1000) if target_k != 0 else 0
        if abs(permil_k) <= 6.0: return None
        diff_abs_k = abs(diff_k)
        pct_k = (diff_abs_k / target_k * 100) if target_k != 0 else 0
        es_faltante = diff_k < 0
        tipo_diff = "FALTANTE" if es_faltante else "SOBRANTE"
        es_denuncia = pct_k >= 2.0
        modo_comp = modo_comp_forzado or ("laboratorio" if ctrl_lab else ("salida" if ctrl_sal else "documento"))

        # Info operación
        if tipo_operacion_info:
            art_principal  = tipo_operacion_info.get("art_principal", "Art. 954 del Código Aduanero")
            art_infraccion = tipo_operacion_info.get("art_inc", "Art. 954 inc. c) C.A.")
            desc_operacion = tipo_operacion_info.get("descripcion", "OPERACIÓN ADUANERA")
            cod_subregimen = tipo_operacion_info.get("codigo", "")
        elif tipo_operacion in ("exportacion", "remo_carga"):
            art_principal = "Art. 959 del Código Aduanero"; art_infraccion = "Art. 959 inc. c) C.A."
            desc_operacion = "EXPORTACIÓN"; cod_subregimen = ""
        else:
            art_principal = "Art. 954 del Código Aduanero"; art_infraccion = "Art. 954 inc. c) C.A."
            desc_operacion = "IMPORTACIÓN"; cod_subregimen = ""

        aduana    = self.aduana_nombre()
        buque     = self.get_var('car_buque').get()
        operador  = self._ddt_actor(ddt_obj, "impexp")
        dest_num  = ddt_obj["numero"].get()
        producto  = ddt_obj["producto"].get()
        doc_k     = self.parse_float(ddt_obj["kilos"].get())
        fecha_op  = self.get_var('car_fecha').get()
        fecha_hoy = datetime.now().strftime("%d/%m/%Y")
        tipo_doc_str = "DENUNCIA" if es_denuncia else "CARGO TRIBUTARIO"
        modo_labels   = {"documento": "Documento declarado", "salida": "Salida de Zona Primaria", "laboratorio": "Análisis de Laboratorio"}
        _vl_var = ddt_obj.get("valor_litro")
        valor_kg = self.parse_float(_vl_var.get() if isinstance(_vl_var, tk.StringVar) else str(_vl_var or "0"))
        divisa   = ddt_obj.get("divisa", tk.StringVar(value="ARS"))
        divisa   = divisa.get() if isinstance(divisa, tk.StringVar) else str(divisa)
        divisa_sym = {"ARS": "$", "USD": "U$S", "EUR": "€", "BRL": "R$", "GBP": "£"}.get(divisa, "$")
        _tc_var = ddt_obj.get("tipo_cambio")
        tc_ddt  = self.parse_float(_tc_var.get() if isinstance(_tc_var, tk.StringVar) else str(_tc_var or "0"))
        # T/C global del documento (carátula) para la línea de equivalencia en USD
        tc_global = self.parse_float(self.get_var("car_tipo_cambio", "").get())
        # Regla: para divisa!=ARS usamos el T/C del DDT (convierte precio a ARS)
        #        para la línea "Equiv. en USD" usamos SIEMPRE el T/C global del documento
        tipo_cambio = tc_ddt if (divisa != "ARS" and tc_ddt > 0) else tc_global

        doc = Document()

        # ── Márgenes ──
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        def add_heading(text, level=1, color=None):
            p = doc.add_heading(text, level=level)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
            if color:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(*color)
            return p

        def add_para(text, bold=False, italic=False, size=10, color=None, align=None):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = bold; run.italic = italic; run.font.size = Pt(size)
            if color: run.font.color.rgb = RGBColor(*color)
            if align: p.alignment = align
            return p

        def add_table_row(table, cells_data):
            row = table.add_row()
            for i, (text, bold, bg) in enumerate(cells_data):
                cell = row.cells[i]
                cell.text = text
                if bold:
                    for run in cell.paragraphs[0].runs: run.bold = True
                cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            return row

        # ── ENCABEZADO ──
        color_warn = (192, 57, 43) if es_denuncia else (230, 126, 34)
        add_heading(f"INFORME DE {tipo_doc_str}", level=1, color=color_warn)
        add_para(f"Generado: {fecha_hoy}", italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

        # ── DATOS ──
        add_heading("I. DATOS DE LA OPERACIÓN", level=2)
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        tbl.columns[0].width = Cm(6); tbl.columns[1].width = Cm(11)
        rows_data = [
            ("Aduana", aduana), ("Buque / Nave", buque), ("Operador / Imp-Exp", operador),
            ("Documento", dest_num), ("Subregimen", f"{desc_operacion} {cod_subregimen}".strip()),
            ("Producto", producto), ("Fecha operación", fecha_op),
            ("Base de comparación", modo_labels.get(modo_comp, modo_comp)),
        ]
        for lbl, val in rows_data:
            row = tbl.add_row()
            row.cells[0].text = lbl; row.cells[1].text = str(val)
            for run in row.cells[0].paragraphs[0].runs: run.bold = True
        doc.add_paragraph()

        # ── DIFERENCIA ──
        add_heading("II. DIFERENCIA DETECTADA", level=2)
        tbl2 = doc.add_table(rows=0, cols=2); tbl2.style = "Table Grid"
        tbl2.columns[0].width = Cm(8); tbl2.columns[1].width = Cm(9)
        diff_rows = [
            ("Tipo de diferencia", tipo_diff),
            ("Diferencia en Kilos", f"{diff_abs_k:,.0f} kg"),
            ("Por mil (‰)", f"{abs(permil_k):.2f}‰"),
            ("Porcentaje (%)", f"{pct_k:.2f}%"),
            ("Total declarado (kg)", f"{doc_k:,.0f} kg"),
            ("Clasificación", "DENUNCIA (>=2%)" if es_denuncia else "CARGO TRIBUTARIO (entre 6 o/oo y 2%)"),
        ]
        for lbl, val in diff_rows:
            row = tbl2.add_row()
            row.cells[0].text = lbl; row.cells[1].text = val
            for run in row.cells[0].paragraphs[0].runs: run.bold = True
        doc.add_paragraph()

        # ── ENCUADRE TÉCNICO ──
        add_heading("III. ENCUADRE TÉCNICO Y NORMATIVO", level=2)
        if es_denuncia:
            add_para(
                f"La diferencia detectada ({diff_abs_k:,.0f} kg) representa el {pct_k:.2f}% del total declarado, "
                f"SUPERANDO el umbral del 2%, lo que configura la presunción de infracción al {art_principal} "
                f"por declaración inexacta de cantidad de mercadería.", size=10)
        else:
            add_para(
                f"La diferencia detectada ({diff_abs_k:,.0f} kg) representa el {pct_k:.2f}% del total declarado, "
                f"superando la franquicia técnica del 6‰ pero resultando INFERIOR al 2%, "
                f"por lo cual es de aplicación la excusa absolutoria del {art_infraccion}. "
                f"No obstante, en materia tributaria existe un hecho imponible que genera el cargo correspondiente.", size=10)
        doc.add_paragraph()
        add_para("Marco Normativo:", bold=True, size=10)
        normas = [
            f"• {art_principal}: Infracción por declaración inexacta de cantidad.",
            "• Arts. 637/638 C.A.: Determinación y nacimiento de la obligación tributaria.",
            "• Arts. 790/791 C.A.: Liquidación de tributos y actuación aduanera.",
            "• Res. ex-ANA 2220/90: Tolerancia técnica 0,6% para graneles líquidos.",
        ]
        for n in normas:
            doc.add_paragraph(n, style="List Bullet")
        doc.add_paragraph()

        # ── LIQUIDACIÓN ──
        # ── Conversión de divisa ──────────────────────────────────────────────
        # Si el precio está en moneda extranjera, convertir a ARS usando el
        # tipo de cambio del documento para la base imponible final.
        usa_tc = (divisa != "ARS") and (tipo_cambio > 0)
        if usa_tc:
            # valor_kg_ars: precio por kg en pesos (para la liquidación ARS)
            valor_kg_ars = valor_kg * tipo_cambio
            tc_nota = (f"T/C: {divisa_sym} 1 = $ {tipo_cambio:,.2f}  →  "
                       f"Precio en ARS: {divisa_sym} {valor_kg:,.4f} × {tipo_cambio:,.2f} = $ {valor_kg_ars:,.4f}/Kg")
        else:
            valor_kg_ars = valor_kg   # ya en ARS, sin conversión
            tc_nota = ""

        if es_denuncia:
            add_heading("IV. BASE IMPONIBLE ESTIMADA — DENUNCIA", level=2)
            base_k = doc_k if doc_k > 0 else diff_abs_k
            add_para(f"La base imponible correspondería a la totalidad de la carga: {base_k:,.0f} kg", size=10)
            if valor_kg > 0:
                base_imp_orig = base_k * valor_kg          # en divisa original
                base_imp      = base_k * valor_kg_ars      # en ARS (= base_imp_orig si ya es ARS)
                if usa_tc:
                    add_para(tc_nota, size=9, italic=True)
                    add_para(
                        f"Total: {base_k:,.0f} Kg × {divisa_sym} {valor_kg:,.4f}/Kg"
                        f" = {divisa_sym} {base_imp_orig:,.2f}"
                        f"  →  $ {base_imp:,.2f} (tipo de cambio $ {tipo_cambio:,.2f})",
                        bold=True, size=10)
                else:
                    add_para(f"Total: {base_k:,.0f} Kg × $ {valor_kg:,.4f}/Kg = $ {base_imp:,.2f}", bold=True, size=10)
                tbl3 = doc.add_table(rows=0, cols=4); tbl3.style = "Table Grid"
                tbl3.add_row().cells[0].text = "CONCEPTO"
                tbl3.rows[0].cells[1].text = "ALÍC."; tbl3.rows[0].cells[2].text = "BASE (ARS)"; tbl3.rows[0].cells[3].text = "IMPORTE EST. (ARS)"
                for run in tbl3.rows[0].cells[0].paragraphs[0].runs: run.bold = True
                total_t = 0
                for conc, alic in (self.get_tributos_activos() or [("Derechos de Imp/Exp (ref.)", 8.0)]):
                    imp = base_imp * alic / 100; total_t += imp
                    r = tbl3.add_row()
                    r.cells[0].text = conc; r.cells[1].text = f"{alic}%"
                    r.cells[2].text = f"$ {base_imp:,.2f}"; r.cells[3].text = f"$ {imp:,.2f}"
                r = tbl3.add_row()
                r.cells[0].text = "TOTAL ESTIMADO"; r.cells[3].text = f"$ {total_t:,.2f}"
                for run in r.cells[0].paragraphs[0].runs: run.bold = True
        else:
            add_heading("IV. LIQUIDACIÓN ESTIMADA DE CARGO", level=2)
            add_para(f"El cargo tributario corresponde al {tipo_diff.lower()} de {diff_abs_k:,.0f} kg.", size=10)
            if valor_kg > 0:
                base_imp_orig = diff_abs_k * valor_kg      # en divisa original
                base_imp      = diff_abs_k * valor_kg_ars  # en ARS
                if usa_tc:
                    add_para(tc_nota, size=9, italic=True)
                    add_para(
                        f"{tipo_diff}: {diff_abs_k:,.0f} Kg × {divisa_sym} {valor_kg:,.4f}/Kg"
                        f" = {divisa_sym} {base_imp_orig:,.2f}"
                        f"  →  $ {base_imp:,.2f} (tipo de cambio $ {tipo_cambio:,.2f})",
                        bold=True, size=10)
                else:
                    add_para(f"{tipo_diff}: {diff_abs_k:,.0f} Kg × $ {valor_kg:,.4f}/Kg = $ {base_imp:,.2f}", bold=True, size=10)
                tbl3 = doc.add_table(rows=0, cols=4); tbl3.style = "Table Grid"
                r0 = tbl3.add_row()
                r0.cells[0].text = "CONCEPTO"; r0.cells[1].text = "ALÍC."
                r0.cells[2].text = "BASE (ARS)"; r0.cells[3].text = "IMPORTE EST. (ARS)"
                for run in r0.cells[0].paragraphs[0].runs: run.bold = True
                total_t = 0
                for conc, alic in (self.get_tributos_activos() or [("Derechos de Imp/Exp (ref.)", 8.0)]):
                    imp = base_imp * alic / 100; total_t += imp
                    r = tbl3.add_row()
                    r.cells[0].text = conc; r.cells[1].text = f"{alic}%"
                    r.cells[2].text = f"$ {base_imp:,.2f}"; r.cells[3].text = f"$ {imp:,.2f}"
                r = tbl3.add_row()
                r.cells[0].text = f"TOTAL ESTIMADO {tipo_diff}"; r.cells[3].text = f"$ {total_t:,.2f}"
                for run in r.cells[0].paragraphs[0].runs: run.bold = True
        doc.add_paragraph()

        # ── CONCLUSIÓN ──
        add_heading("V. CONCLUSIÓN", level=2)
        if es_denuncia:
            add_para(
                f"Se elevarían las presentes actuaciones a fin de que la superioridad evaluara la posible "
                f"instrucción de sumario al amparo del {art_principal}, en virtud del {tipo_diff.lower()} "
                f"detectado de {diff_abs_k:,.0f} kilogramos ({pct_k:.2f}% del total declarado), "
                f"sin perjuicio del cargo tributario que eventualmente correspondiere. Es todo cuanto se informaría.", size=10)
        else:
            add_para(
                f"Se elevarían las presentes actuaciones a fin de que la superioridad evaluara la procedencia "
                f"del cargo tributario en los términos del {art_principal}, en virtud del {tipo_diff.lower()} "
                f"detectado de {diff_abs_k:,.0f} kilogramos ({pct_k:.2f}% del total declarado). Es todo cuanto se informaría.", size=10)
        doc.add_paragraph()
        add_para(f"Fecha: {fecha_hoy}", italic=True, size=9)

        # ── GUARDAR ──
        clean_name  = self.clean_filename(buque)
        fecha_arch  = datetime.now().strftime("%Y%m%d")
        tipo_arch   = "Denuncia" if es_denuncia else "Cargo"
        docx_file   = f"{tipo_arch}_Tributario_{clean_name}_{self.clean_filename(dest_num)}_{fecha_arch}.docx"
        docx_path   = os.path.join(output_folder or ".", docx_file)
        doc.save(docx_path)
        return docx_path

    def detectar_tolerancia_y_cargo(self, ddt_obj, dif_kv, target_k, output_folder,
                                   ctrl_doc=True, ctrl_sal=False, ctrl_lab=False,
                                   modo_comp_forzado=None, shared_canvas=None,
                                   tipo_operacion_override=None, tipo_operacion_info=None):
        """Detecta si excede tolerancia (en KILOS) y genera cargo o denuncia."""
        if target_k == 0: return
        diff_k = dif_kv - target_k
        permil_k = (diff_k / target_k * 1000) if target_k != 0 else 0
        if abs(permil_k) <= 6.0: return
        diff_abs_k = abs(diff_k)
        pct_k = (diff_abs_k / target_k * 100) if target_k != 0 else 0
        es_faltante = diff_k < 0
        tipo_diff = "FALTANTE" if es_faltante else "SOBRANTE"
        es_denuncia = pct_k >= 2.0
        modo_comp = modo_comp_forzado or ("laboratorio" if ctrl_lab else ("salida" if ctrl_sal else "documento"))

        if tipo_operacion_override:
            return self._generar_cargo_pdf(ddt_obj, diff_abs_k, permil_k,
                                           pct_k, tipo_diff, es_faltante, output_folder, modo_comp,
                                           es_denuncia=es_denuncia, tipo_operacion=tipo_operacion_override,
                                           shared_canvas=shared_canvas, tipo_operacion_info=tipo_operacion_info)

        # Diálogo: solo importación/exportación
        dlg = tk.Toplevel(self.root)
        dlg.title("Tolerancia Excedida - Confirmar Cargo / Denuncia")
        dlg.transient(self.root)
        dlg.grab_set()
        dlg.update_idletasks()
        sw, sh = dlg.winfo_screenwidth(), dlg.winfo_screenheight()
        w_dlg, h_dlg = 560, 280
        dlg.geometry(f"{w_dlg}x{h_dlg}+{(sw-w_dlg)//2}+{(sh-h_dlg)//2}")
        color_warn = "#C0392B" if es_denuncia else "#E67E22"
        titulo = "DENUNCIA - SUPERA 2% EN KILOS" if es_denuncia else "CARGO TRIBUTARIO - Entre 6 o/oo y 2% en Kilos"
        tk.Label(dlg, text=titulo, font=("Arial", 8, "bold"), fg=color_warn, bg="#FDF2F2").pack(fill="x", ipady=8)
        comp_labels = {"documento": "Documento declarado", "salida": "Salida de Zona Primaria", "laboratorio": "Análisis de Laboratorio"}
        dest_num_val = ddt_obj["numero"].get()
        lbl_info  = f"Destino: {dest_num_val}\n"
        lbl_info += f"Diferencia: {diff_abs_k:,.0f} kg {tipo_diff}  ({abs(permil_k):.2f} o/oo  |  {pct_k:.2f}%)\n"
        lbl_info += f"Comparación: {comp_labels.get(modo_comp, modo_comp)}"
        tk.Label(dlg, text=lbl_info, font=("Arial", 10), justify="left").pack(padx=20, pady=8)
        f_op = ttk.LabelFrame(dlg, text="Tipo de operación")
        f_op.pack(fill="x", padx=20, pady=4)
        var_op = tk.StringVar(value="importacion")
        tk.Radiobutton(f_op, text="Importación  (Art. 954 C.A.)", variable=var_op, value="importacion", font=("Arial", 10)).pack(side="left", padx=20)
        tk.Radiobutton(f_op, text="Exportación  (Art. 959 C.A.)", variable=var_op, value="exportacion", font=("Arial", 10)).pack(side="left", padx=20)
        resultado = {"generar": False, "operacion": "importacion"}
        def aceptar():
            resultado["generar"] = True; resultado["operacion"] = var_op.get(); dlg.destroy()
        def cancelar(): dlg.destroy()
        lbl_btn = "GENERAR DENUNCIA" if es_denuncia else "GENERAR CARGO"
        f_bot = tk.Frame(dlg); f_bot.pack(pady=10)
        tk.Button(f_bot, text=lbl_btn, bg="#C0392B", fg="white", font=("Arial", 8, "bold"), command=aceptar).pack(side="left", padx=10)
        tk.Button(f_bot, text="Cancelar", bg="gray", fg="white", font=("Arial", 10), command=cancelar).pack(side="left", padx=10)
        dlg.wait_window()
        if not resultado["generar"]: return
        return self._generar_cargo_pdf(ddt_obj, diff_abs_k, permil_k,
                                pct_k, tipo_diff, es_faltante, output_folder, modo_comp,
                                es_denuncia=es_denuncia, tipo_operacion=resultado["operacion"],
                                shared_canvas=shared_canvas)

    def _generar_cargo_en_canvas(self, c, ddt_obj, dif_kv, target_k,
                                  ctrl_doc=True, ctrl_sal=False, ctrl_lab=False,
                                  modo_comp_forzado=None, tipo_operacion="importacion", tipo_operacion_info=None):
        """Calcula tolerancia en KILOS y si supera dibuja cargo/denuncia en el canvas c."""
        if target_k == 0: return
        diff_k = dif_kv - target_k
        permil_k = (diff_k / target_k * 1000) if target_k != 0 else 0
        if abs(permil_k) <= 6.0: return
        diff_abs_k = abs(diff_k)
        pct_k = (diff_abs_k / target_k * 100) if target_k != 0 else 0
        es_faltante = diff_k < 0
        tipo_diff = "FALTANTE" if es_faltante else "SOBRANTE"
        es_denuncia = pct_k >= 2.0
        modo_comp = modo_comp_forzado or ("laboratorio" if ctrl_lab else ("salida" if ctrl_sal else "documento"))
        self._generar_cargo_pdf(
            ddt_obj, diff_abs_k, permil_k,
            pct_k, tipo_diff, es_faltante, output_folder=None, modo_comp=modo_comp,
            es_denuncia=es_denuncia, tipo_operacion=tipo_operacion,
            shared_canvas=c, tipo_operacion_info=tipo_operacion_info
        )

    def _generar_cargo_pdf(self, ddt_obj, diff_abs_k, permil_k,
                           pct_max, tipo_diff, es_faltante, output_folder, modo_comp,
                           es_denuncia=False, tipo_operacion="importacion", shared_canvas=None,
                           tipo_operacion_info=None, _buf_override=None):
        """Genera el informe de cargo/denuncia — todo en KILOS, sin litros.
        Si shared_canvas: dibuja en él y retorna None.
        Si output_folder: guarda a archivo y retorna la ruta.
        """
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.utils import simpleSplit
        import io
        
        clean_name = self.clean_filename(self.get_var('car_buque').get())
        dest_num = ddt_obj["numero"].get()
        tipo_doc_str = "Denuncia" if es_denuncia else "Cargo"
        filename = f"{tipo_doc_str}_Tributario_{clean_name}_{self.clean_filename(dest_num)}_{datetime.now().strftime('%Y%m%d')}.pdf"
        full_path = os.path.join(output_folder, filename) if output_folder else None
        
        # Artículos y descripción de operación
        if tipo_operacion_info:
            art_principal  = tipo_operacion_info.get("art_principal", "Art. 954 del Código Aduanero")
            art_infraccion = tipo_operacion_info.get("art_inc",        "Art. 954 inc. c) C.A.")
            art_denuncia   = tipo_operacion_info.get("art_principal",  "Art. 954 del Código Aduanero")
            desc_operacion = tipo_operacion_info.get("descripcion",    "OPERACIÓN ADUANERA")
            cod_subregimen = tipo_operacion_info.get("codigo",         "")
        elif tipo_operacion in ("exportacion", "remo_carga"):
            art_principal  = "Art. 959 del Código Aduanero"
            art_infraccion = "Art. 959 inc. c) C.A."
            art_denuncia   = "Art. 959 del Código Aduanero"
            desc_operacion = "EXPORTACIÓN"
            cod_subregimen = ""
        else:
            art_principal  = "Art. 954 del Código Aduanero"
            art_infraccion = "Art. 954 inc. c) C.A."
            art_denuncia   = "Art. 954 del Código Aduanero"
            desc_operacion = "IMPORTACIÓN"
            cod_subregimen = ""
        
        standalone = shared_canvas is None
        if shared_canvas is not None:
            # Dibujar directamente en el canvas compartido (ya en portrait A4)
            c = shared_canvas
            c.setPageSize(A4)  # asegurar portrait
        else:
            _buf = _buf_override if _buf_override is not None else io.BytesIO()
            try:
                c = canvas.Canvas(_buf, pagesize=A4)
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo crear buffer de cargo: {e}")
                return
        
        w, h = A4  # 595 x 842
        margin = 50
        usable_w = w - 2 * margin
        
        aduana = self.aduana_nombre()
        buque = self.get_var('car_buque').get() or self.get_var('car_patente').get() or "S/I"
        operador = self._ddt_actor(ddt_obj, "impexp")
        cuit_op = self._ddt_actor(ddt_obj, "cuit_impexp")
        ata = self._ddt_actor(ddt_obj, "ata")
        cuit_ata = self._ddt_actor(ddt_obj, "cuit_ata")
        fecha_med = self.get_var('final_Fecha').get() or self.get_var('inicial_Fecha').get()
        
        _vl_var2 = ddt_obj.get("valor_litro")
        valor_litro = self.parse_float(_vl_var2.get() if isinstance(_vl_var2, tk.StringVar) else str(_vl_var2 or "0"))
        divisa = ddt_obj.get("divisa", tk.StringVar(value="ARS")).get() if isinstance(ddt_obj.get("divisa"), tk.StringVar) else ddt_obj.get("divisa", "ARS")
        divisa_desc = ddt_obj.get("divisa_desc", tk.StringVar()).get() if isinstance(ddt_obj.get("divisa_desc"), tk.StringVar) else ddt_obj.get("divisa_desc", "")
        if divisa == "Otra" and divisa_desc: divisa = divisa_desc
        tipo_cambio_ddt = self.parse_float(ddt_obj.get("tipo_cambio", tk.StringVar()).get() if isinstance(ddt_obj.get("tipo_cambio"), tk.StringVar) else ddt_obj.get("tipo_cambio", "0"))
        tc_global_pdf   = self.parse_float(self.get_var("car_tipo_cambio", "").get())
        # Para divisa!=ARS: T/C del DDT convierte precio a ARS; para equiv USD: T/C global
        tipo_cambio = tipo_cambio_ddt if (divisa != "ARS" and tipo_cambio_ddt > 0) else tc_global_pdf
        
        doc_l = self.parse_float(ddt_obj['litros'].get())
        doc_k = self.parse_float(ddt_obj['kilos'].get())
        dens_doc = self.parse_float(ddt_obj['densidad'].get())
        producto_desc = ddt_obj.get('producto', tk.StringVar()).get() if isinstance(ddt_obj.get('producto'), tk.StringVar) else str(ddt_obj.get('producto', ''))
        pos_arancel = ddt_obj.get('pos_arancel', tk.StringVar()).get() if isinstance(ddt_obj.get('pos_arancel'), tk.StringVar) else str(ddt_obj.get('pos_arancel', ''))
        # Salidas declaradas en la destinación
        salidas_list = []
        for s in (ddt_obj.get('salidas') or []):
            s_num = s['numero'].get() if isinstance(s.get('numero'), tk.StringVar) else str(s.get('numero',''))
            s_k   = self.parse_float(s['kilos'].get()  if isinstance(s.get('kilos'),   tk.StringVar) else s.get('kilos',  '0'))
            s_l   = self.parse_float(s['litros'].get() if isinstance(s.get('litros'),  tk.StringVar) else s.get('litros', '0'))
            s_d   = self.parse_float(s['densidad'].get() if isinstance(s.get('densidad'), tk.StringVar) else s.get('densidad','0'))
            if s_num or s_k or s_l:
                salidas_list.append({'numero': s_num, 'kilos': s_k, 'litros': s_l, 'densidad': s_d})
        sal_k_total = sum(s['kilos']  for s in salidas_list)
        sal_l_total = sum(s['litros'] for s in salidas_list)
        
        def draw_text(txt, y_pos, font_name="Helvetica", font_size=10, bold=False, centered=False):
            fn = f"{font_name}-Bold" if bold else font_name
            c.setFont(fn, font_size)
            if centered:
                c.drawCentredString(w/2, y_pos, txt)
            else:
                c.drawString(margin, y_pos, txt)
            return y_pos - font_size - 4
        
        def draw_wrapped(txt, y_pos, font_name="Helvetica", font_size=10, indent=0):
            c.setFont(font_name, font_size)
            lines = simpleSplit(txt, font_name, font_size, usable_w - indent)
            for line in lines:
                c.drawString(margin + indent, y_pos, line)
                y_pos -= font_size + 3
            return y_pos
        
        def new_page_if_needed(y_pos, needed=80):
            if y_pos < needed:
                c.showPage()
                return h - 50
            return y_pos
        
        # === I. TÍTULO Y ENCABEZADO ===
        # ======================================================
        # === I. ENCABEZADO
        # ======================================================
        y = h - 50
        # Datos de lugar operativo
        lop_cod  = self.get_var("car_lop_codigo").get().strip()
        lop_desc = self.get_var("car_lop_desc").get().strip()
        lugar_str = aduana
        if lop_cod or lop_desc:
            lugar_str = f"{aduana} — Lugar Operativo: {lop_cod} {lop_desc}".strip(" —")

        tipo_doc_titulo = "INFORME DE DENUNCIA POR POSIBLE INFRACCIÓN" if es_denuncia else "INFORME DE CARGO TRIBUTARIO"
        aduana_full = self.aduana_completa()  # "067 - USHUAIA"
        aduana_cod_h = self.aduana_codigo()   # "067"
        aduana_nom_h = self.aduana_nombre()   # "USHUAIA"
        # Primer renglón: solo aduana (código + nombre)
        y = draw_text(f"ADUANA {aduana_cod_h} - {aduana_nom_h.upper()} — DIRECCIÓN GENERAL DE ADUANAS — ARCA", y, font_size=11, bold=True, centered=True)
        # Segundo renglón: Lugar Operativo (si existe), separado
        if lop_cod or lop_desc:
            lop_str = f"Lugar Operativo: {lop_cod}  {lop_desc}".strip()
            y = draw_text(lop_str, y, font_size=9, centered=True)
        y = draw_text(f"{aduana_nom_h}, {fecha_med}", y, font_size=9, centered=True)
        y -= 8
        c.setFont("Helvetica-Bold", 13)
        c.drawCentredString(w/2, y, tipo_doc_titulo)
        y -= 16
        # Tipo de operación resaltado
        _desc_font = 11 if len(desc_operacion) <= 60 else (9 if len(desc_operacion) > 80 else 10)
        c.setFont("Helvetica-Bold", _desc_font)
        c.setFillColor(colors.HexColor("#1B3A5C"))
        c.drawCentredString(w/2, y, f"[ {desc_operacion} ]")
        c.setFillColor(colors.black)
        y -= (_desc_font + 7)
        asunto = (f"Posible infracción al {art_principal} — {tipo_diff.capitalize()} detectado en destinación."
                  if es_denuncia else
                  f"Posible procedencia de Cargo Tributario por {tipo_diff.lower()} en destinación.")
        y = draw_text(f"Ref.: {asunto}", y, font_size=10, bold=True)
        y = draw_text(f"Destinación N°: {dest_num}", y, font_size=10)
        y = draw_text(f"Operador: {operador}  —  CUIT: {cuit_op}", y, font_size=10)
        y = draw_text(f"Agente de Transporte Aduanero: {ata}  —  CUIT: {cuit_ata}", y, font_size=10)
        tipo_nave = self.label_unidad()
        tm_cargo = self.get_tipo_medio()
        lugar_body = aduana_nom_h
        if lop_cod or lop_desc:
            lugar_body = f"{aduana_nom_h} / Lugar Op.: {lop_cod} {lop_desc}".strip(" /")
        y = draw_text(f"Lugar: {lugar_body}  -  {tipo_nave}: {buque}", y, font_size=10)
        y -= 12

        # I.a FUNCIONARIOS INTERVINIENTES
        if self.funcionarios_data:
            y = new_page_if_needed(y, 120)
            y = draw_text("I. FUNCIONARIOS INTERVINIENTES", y, font_size=11, bold=True)
            y -= 4
            y = draw_wrapped(
                "Los abajo firmantes, en calidad de funcionarios de la Dirección General de Aduanas "
                "— ARCA, certificarían haber intervenido en la operación de medición de referencia "
                "ejerciendo las funciones que se detallan:",
                y, font_size=10)
            y -= 6
            for f_obj in self.funcionarios_data:
                apell = f_obj.get("apellido", tk.StringVar()).get() if isinstance(f_obj.get("apellido"), tk.StringVar) else ""
                nom = f_obj.get("nombre", tk.StringVar()).get() if isinstance(f_obj.get("nombre"), tk.StringVar) else ""
                nombre_completo = f"{apell.upper()}, {nom}".strip(", ")
                legajo = f_obj["legajo"].get() if isinstance(f_obj.get("legajo"), tk.StringVar) else str(f_obj.get("legajo", ""))
                cuil = f_obj["cuil"].get() if isinstance(f_obj.get("cuil"), tk.StringVar) else str(f_obj.get("cuil", ""))
                funcion = f_obj["funcion"].get() if isinstance(f_obj.get("funcion"), tk.StringVar) else str(f_obj.get("funcion", ""))
                linea = f"• {nombre_completo}  —  Leg. {legajo}  —  CUIL {cuil}  —  Función: {funcion}"
                y = draw_wrapped(linea, y, font_size=9, indent=10)
                y -= 2
            y -= 8

        # ======================================================
        # === II. RELACIÓN DE LOS HECHOS
        # ======================================================
        y = new_page_if_needed(y, 150)
        sec_num = "II" if self.funcionarios_data else "I"
        y = draw_text(f"{sec_num}. RELACIÓN DE LOS HECHOS Y CONSTANCIAS TÉCNICAS", y, font_size=11, bold=True)
        y -= 5
        comp_label = {"documento": "Documento declarado", "salida": "Salida de Zona Primaria", "laboratorio": "Análisis de Laboratorio"}
        comp_str = comp_label.get(modo_comp, modo_comp)

        # Párrafo 1: introducción — adaptado por tipo de medio
        _tm_legal = self.get_tipo_medio()
        if _tm_legal in ("OLEODUCTO","POLIDUCTO","GASODUCTO"):
            intro_p1 = (
                f"1. El presente informe se elevaría a los fines de que la superioridad determine el "
                f"tratamiento tributario{'e infraccional' if es_denuncia else ''} que pudiere corresponder "
                f"en relación a las diferencias detectadas en el volumen (y/o masa) de "
                f"{'gas natural / hidrocarburos gaseosos' if 'GASODUCTO' in _tm_legal else 'hidrocarburo líquido'} "
                f"transportado por el {tipo_nave.lower()} de referencia, entre lo declarado y lo medido "
                f"según las constancias instrumentales del período."
            )
        elif _tm_legal == "MEDICION ELECTRICA":
            intro_p1 = (
                f"1. El presente informe se elevaría a los fines de que la superioridad determine el "
                f"tratamiento tributario{'e infraccional' if es_denuncia else ''} que pudiere corresponder "
                f"en relación a las diferencias constatadas en la medición de energía eléctrica "
                f"en la instalación de referencia, entre lo declarado y lo registrado en el medidor "
                f"durante el período de control."
            )
        elif "CAMION GAS" in _tm_legal:
            intro_p1 = (
                f"1. El presente informe se elevaría a los fines de que la superioridad determine el "
                f"tratamiento tributario{'e infraccional' if es_denuncia else ''} que pudiere corresponder "
                f"en relación a las diferencias de masa detectadas en la carga de gas licuado / GLP "
                f"transportada en el vehículo cisterna de referencia."
            )
        else:
            intro_p1 = (
                f"1. El presente informe se elevaría a los fines de que la superioridad determine el "
                f"tratamiento tributario{'e infraccional' if es_denuncia else ''} que pudiere corresponder, "
                f"a efectos de salvaguardar la renta fiscal en relación a las diferencias de peso constatadas "
                f"a la finalización de la operación en el {tipo_nave.lower()} de referencia."
            )
        y = draw_wrapped(intro_p1, y, font_size=10, indent=15)
        y -= 5

        # Párrafo 2: producto
        prod_str = producto_desc if producto_desc else "sin especificar"
        pos_str  = f" (Pos. Arancel: {pos_arancel})" if pos_arancel else ""
        y = draw_wrapped(
            f"2. La mercadería objeto de la destinación {dest_num} sería: {prod_str}{pos_str}. "
            f"La documentación declaró {doc_l:,.0f} Litros y {doc_k:,.0f} Kilogramos"
            f"{f', con densidad {dens_doc:.4f} Kg/L' if dens_doc > 0 else ''}.",
            y, font_size=10, indent=15)
        y -= 5

        # Párrafo 3: salidas (si aplica modo salida o laboratorio y hay salidas)
        nro_item = 3
        if modo_comp in ("salida", "laboratorio") and salidas_list:
            y = new_page_if_needed(y, 80)
            y = draw_wrapped(
                f"{nro_item}. Las salidas de Zona Primaria declaradas en la destinación son las siguientes:",
                y, font_size=10, indent=15)
            y -= 3
            # Tabla de salidas
            c.setFont("Helvetica-Bold", 9)
            sx = margin + 25
            c.drawString(sx,       y, "N° Salida")
            c.drawString(sx+90,    y, "Kilos")
            c.drawString(sx+170,   y, "Litros")
            c.drawString(sx+250,   y, "Densidad (Kg/L)")
            y -= 3
            c.line(sx, y, sx+350, y)
            y -= 10
            c.setFont("Helvetica", 9)
            for s in salidas_list:
                c.drawString(sx,    y, str(s['numero']) if s['numero'] else "-")
                c.drawString(sx+90, y, f"{s['kilos']:,.0f}" if s['kilos'] else "-")
                c.drawString(sx+170,y, f"{s['litros']:,.0f}" if s['litros'] else "-")
                c.drawString(sx+250,y, f"{s['densidad']:.4f}" if s['densidad'] else "-")
                y -= 11
            y -= 4
            c.line(sx, y, sx+350, y)
            y -= 13          # bajar suficiente para que el texto quede bajo la línea
            # Total salidas
            c.setFont("Helvetica-Bold", 9)
            c.drawString(sx,     y, "TOTAL")
            c.drawString(sx+90,  y, f"{sal_k_total:,.0f}")
            c.drawString(sx+170, y, f"{sal_l_total:,.0f}")
            y -= 14
            nro_item += 1
            y -= 4

        # Párrafo: sondaje / medición (adaptado por tipo de medio)
        y = new_page_if_needed(y, 60)
        _tm_sond = self.get_tipo_medio()
        if _tm_sond in ("OLEODUCTO","POLIDUCTO","GASODUCTO"):
            sond_txt = (
                f"{nro_item}. Se habría procedido a realizar las lecturas de los contadores de caudal (inicio/fin del período), "
                f"la verificación de los medidores Coriolis (masa directa), el registro de condiciones de línea (T/P), "
                f"el cálculo del factor de compresibilidad Z y la corrección a condiciones base. "
                f"{'La cromatografía de gas (composición % mol) fue registrada según lo indicado en planilla adjunta.' if _tm_sond=='GASODUCTO' else ''} "
                f"Se adjuntan las actas de medición debidamente rubricadas."
            )
        elif _tm_sond == "MEDICION ELECTRICA":
            sond_txt = (
                f"{nro_item}. Se habría procedido a realizar la lectura inicial y final del medidor de energía, "
                f"el registro de tensión, corriente, factor de potencia y demanda máxima, aplicando la constante "
                f"de transformación correspondiente. Se adjuntan las actas de medición debidamente rubricadas."
            )
        else:
            sond_txt = (
                f"{nro_item}. Se habría procedido a realizar el correspondiente sondaje de tanques del "
                f"{tipo_nave.lower()} y la toma de muestras, adjuntándose las planillas de medición debidamente rubricadas."
            )
        y = draw_wrapped(sond_txt, y, font_size=10, indent=15)
        y -= 5
        nro_item += 1

        # Párrafo: explicación modo de control y diferencia
        y = new_page_if_needed(y, 100)
        if modo_comp == "laboratorio":
            if salidas_list:
                sal_desc = f"{sal_k_total:,.0f} Kg ({sal_l_total:,.0f} Litros)" if sal_l_total else f"{sal_k_total:,.0f} Kg"
                txt_modo = (
                    f"{nro_item}. El control se habría efectuado mediante análisis de laboratorio. "
                    f"Las salidas de la Declaración Detallada acumulan {sal_desc}. "
                    f"Al comparar las mediciones realizadas a bordo con la densidad determinada por análisis de laboratorio, "
                    f"se habría constatado una diferencia de {diff_abs_k:,.0f} Kg ({tipo_diff.lower()}) "
                    f"equivalente al {abs(permil_k):.2f}‰ ({pct_max:.2f}%) respecto de las salidas declaradas."
                )
            else:
                txt_modo = (
                    f"{nro_item}. El control se habría efectuado mediante análisis de laboratorio. "
                    f"Al aplicar la densidad determinada en laboratorio sobre las mediciones realizadas a bordo, "
                    f"se habría constatado una diferencia de {diff_abs_k:,.0f} Kg ({tipo_diff.lower()}) "
                    f"equivalente al {abs(permil_k):.2f}‰ ({pct_max:.2f}%) respecto de la carga declarada."
                )
        elif modo_comp == "salida":
            if salidas_list:
                sal_desc = f"{sal_k_total:,.0f} Kg ({sal_l_total:,.0f} Litros)" if sal_l_total else f"{sal_k_total:,.0f} Kg"
                txt_modo = (
                    f"{nro_item}. El control se habría efectuado cotejando las mediciones con las Salidas de Zona Primaria "
                    f"que acumulan {sal_desc}. La diferencia entre las salidas declaradas y las mediciones "
                    f"realizadas a bordo resultaría de {diff_abs_k:,.0f} Kg ({tipo_diff.lower()}), "
                    f"equivalente al {abs(permil_k):.2f}‰ ({pct_max:.2f}%)."
                )
            else:
                txt_modo = (
                    f"{nro_item}. El control se habría efectuado cotejando las mediciones con las Salidas de Zona Primaria. "
                    f"La diferencia entre las salidas declaradas y las mediciones realizadas a bordo resultaría de "
                    f"{diff_abs_k:,.0f} Kg ({tipo_diff.lower()}), equivalente al {abs(permil_k):.2f}‰ ({pct_max:.2f}%)."
                )
        else:  # documento
            txt_modo = (
                f"{nro_item}. El control se habría efectuado confrontando las mediciones realizadas a bordo con la "
                f"cantidad declarada en la documentación ({doc_k:,.0f} Kg / {doc_l:,.0f} Litros). "
                f"La diferencia entre la documentación y las mediciones resultaría de {diff_abs_k:,.0f} Kg ({tipo_diff.lower()}), "
                f"equivalente al {abs(permil_k):.2f}‰ ({pct_max:.2f}%) del total declarado."
            )
        y = draw_wrapped(txt_modo, y, font_size=10, indent=15)
        y -= 5
        nro_item += 1

        # Diferencia resumida (bold)
        y = new_page_if_needed(y, 60)
        y = draw_wrapped(f"{nro_item}. La medición habría arrojado la siguiente diferencia:", y, font_size=10, indent=15)
        c.setFont("Helvetica-Bold", 10)
        y -= 2
        c.drawString(margin + 25, y, f"• Diferencia en Kilos: {diff_abs_k:,.0f} Kg {tipo_diff}S.")
        y -= 13
        c.setFont("Helvetica", 10)
        c.drawString(margin + 25, y, f"• Por mil: {abs(permil_k):.2f} o/oo  |  Porcentaje: {pct_max:.2f}%")
        y -= 18

        y = new_page_if_needed(y, 200)
        y = draw_text("III. ENCUADRE TÉCNICO Y NORMATIVO", y, font_size=11, bold=True)
        y -= 5
        
        if es_denuncia:
            y = draw_wrapped(
                f"La diferencia detectada ({diff_abs_k:,.0f} kg) representa el {pct_max:.2f}% del total de la carga, "
                f"SUPERANDO el umbral del 2%, lo que configura la presunción de infracción al {art_principal}, "
                f"por declaración inexacta de cantidad de mercadería. La diferencia excede la franquicia técnica "
                f"de medición del 6‰ (0,6%) y el umbral de excusa absolutoria del 2%, habilitando la instrucción "
                f"de sumario por infracción.",
                y, font_size=10)
            y -= 8
            y = new_page_if_needed(y, 120)
            c.setFont("Helvetica-Bold", 10)
            c.drawString(margin, y, f"A. Instrucción de Sumario ({art_principal}):")
            y -= 14
            y = draw_wrapped(
                f"Dado que la diferencia supera el 2%, corresponde la instrucción de sumario por declaración inexacta "
                f"al amparo del {art_principal}, sin perjuicio del cargo tributario que proceda por el tributo omitido. "
                f"La excusa absolutoria del 2% prevista en el {art_infraccion} es inaplicable al superarse dicho umbral.",
                y, font_size=10, indent=15)
            y -= 8
        else:
            y = draw_wrapped(
                f"La diferencia detectada ({diff_abs_k:,.0f} kg) representa el {pct_max:.2f}% del total de la carga, "
                f"superando la franquicia técnica del 6‰ (0,6%) pero resultando INFERIOR al 2%, "
                f"por lo cual es de aplicación la excusa absolutoria del {art_infraccion} (no corresponde multa). "
                f"No obstante, en materia TRIBUTARIA rige el principio de integridad: existe un hecho imponible "
                f"que genera el cargo tributario correspondiente.",
                y, font_size=10)
            y -= 8
            y = new_page_if_needed(y, 120)
            c.setFont("Helvetica-Bold", 10)
            c.drawString(margin, y, f"A. Improcedencia de Multa — Procedencia de Cargo ({art_infraccion}):")
            y -= 14
            y = draw_wrapped(
                f"Al ser la diferencia inferior al 2%, resulta aplicable la excusa absolutoria del {art_infraccion}, "
                f"no correspondiendo sanción por infracción de declaración inexacta. Sin embargo, la diferencia "
                f"supera la tolerancia técnica del 0,6%, por lo que existe mercadería que físicamente falta "
                f"y no puede atribuirse a error instrumental, generando el hecho imponible.",
                y, font_size=10, indent=15)
            y -= 8
        
        y = new_page_if_needed(y, 150)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(margin, y, "B. Procedencia del Cargo Tributario (Art. 790 y 791 C.A.):")
        y -= 14
        y = draw_wrapped(
            f"El {tipo_diff.lower()} supera la tolerancia técnica del 0,6% (seis por mil) admitida para mermas "
            f"en transporte acuático de graneles líquidos (Res. ex-ANA 2220/90 y concordantes). "
            f"Dicha mercadería no arribó físicamente. La tolerancia del 2% es una ficción jurídica penal "
            f"(Art. {art_infraccion}). La tolerancia del 0,6% es técnica. Todo lo que supere el 0,6% y no llegue "
            f"al 2% es mercadería que físicamente falta y no puede atribuirse a error de instrumento. "
            f"Por ende, existe un Hecho Imponible (Arts. 637/638 C.A.).",
            y, font_size=10, indent=15)
        y -= 8
        
        y = new_page_if_needed(y, 100)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(margin, y, "C. Marco Normativo Aplicable:")
        y -= 14
        normas = [
            f"• {art_principal}: Infracción por declaración inexacta de cantidad.",
            "• Arts. 637/638 C.A.: Determinación y nacimiento de la obligación tributaria.",
            "• Arts. 790/791 C.A.: Liquidación de tributos y actuación aduanera.",
            "• Arts. 164/390 C.A.: Presunción de destino de mercadería faltante.",
            "• Res. ex-ANA 2220/90: Tolerancia técnica 0,6% para graneles líquidos.",
        ]
        for n in normas:
            y = draw_wrapped(n, y, font_size=9, indent=15)
            y -= 2
        y -= 6
        
        # === IV. JURISPRUDENCIA ===
        y = new_page_if_needed(y, 200)
        y = draw_text("IV. JURISPRUDENCIA Y ANTECEDENTES", y, font_size=11, bold=True)
        y -= 5
        
        juris = [
            ('CASO "ESSO": Autonomía del Tributo respecto de la Multa.',
             'Autos: "Esso S.A.P.A. c/ DGA s/ Recurso de Apelación" — Tribunal Fiscal de la Nación, Sala E.',
             'El TFN confirmó que las franquicias del Art. 959 del C.A. (2%) operan exclusivamente en el ámbito '
             'infraccional, pero no liberan del pago de tributos por mercadería faltante que exceda la merma técnica.'),
            ('CASO "SHELL": La Inexistencia de Tolerancia Fiscal.',
             'Autos: "Shell C.A.P.S.A. c/ DGA" — Tribunal Fiscal de la Nación, Sala F.',
             'Se estableció que la tolerancia del 2% del Art. 959 es una excusa absolutoria penal. '
             'En materia tributaria rige el principio de integridad. Si la merma supera lo técnicamente '
             'aceptable (0,6%), la diferencia está sujeta a tributación.'),
            ('FALLO CORTE SUPREMA (CSJN): "I.B.M. World Trade Corp."',
             'Referencia: Fallos 303:1363.',
             'La obligación tributaria nace del hecho imponible objetivo y es independiente de la intención '
             'dolosa o culposa del operador.')
        ]
        for titulo, ref, doctrina in juris:
            y = new_page_if_needed(y, 80)
            c.setFont("Helvetica-Bold", 9)
            c.drawString(margin + 10, y, titulo)
            y -= 12
            c.setFont("Helvetica-Oblique", 8)
            c.drawString(margin + 20, y, ref)
            y -= 12
            y = draw_wrapped(doctrina, y, font_size=9, indent=20)
            y -= 8
        
        # === V. LIQUIDACIÓN ESTIMADA ===
        y = new_page_if_needed(y, 200)
        if es_denuncia:
            sec_title = "V. BASE IMPONIBLE ESTIMADA — DENUNCIA (total de la carga)"
        else:
            sec_title = "V. LIQUIDACIÓN ESTIMADA DE CARGO (sobre el faltante/sobrante)"
        y = draw_text(sec_title, y, font_size=11, bold=True)
        y -= 5

        moneda_sym = {"ARS": "$", "USD": "U$S", "EUR": "€", "BRL": "R$", "GBP": "£"}.get(divisa, "$")

        lbl_tipo = "FALTANTE" if es_faltante else "SOBRANTE"
        # valor_litro en el DDT se usa como $/kg para el cargo (se trabaja todo en kilos)
        _vl_var = ddt_obj.get("valor_litro")
        valor_kg = self.parse_float(_vl_var.get() if isinstance(_vl_var, tk.StringVar) else str(_vl_var or "0"))
        divisa_sym = {"ARS": "$", "USD": "U$S", "EUR": "€", "BRL": "R$", "GBP": "£"}.get(divisa, "$")

        if es_denuncia:
            base_kilos = doc_k if doc_k > 0 else diff_abs_k
            y = draw_wrapped(
                f"En el presente caso se habría constatado un {lbl_tipo} de {diff_abs_k:,.0f} kilogramos, "
                f"representando el {pct_max:.2f}% del total declarado ({base_kilos:,.0f} kg). "
                f"Por tratarse de una denuncia, la base imponible estimada correspondería a la TOTALIDAD "
                f"de la mercadería objeto de la destinación ({base_kilos:,.0f} kg), "
                f"sin perjuicio de la liquidación que practicara el área de valoración.",
                y, font_size=10)
            y -= 8
            if valor_kg > 0:
                usa_tc_pdf = (divisa != "ARS") and (tipo_cambio > 0)
                base_imp_orig = base_kilos * valor_kg          # en divisa original
                base_imp_ars  = base_imp_orig * tipo_cambio if usa_tc_pdf else base_imp_orig
                if usa_tc_pdf:
                    y = draw_text(f"T/C: {divisa_sym} 1 = $ {tipo_cambio:,.2f}  →  precio ARS: {divisa_sym} {valor_kg:,.4f} × {tipo_cambio:,.2f} = $ {base_imp_orig * tipo_cambio / base_kilos:,.4f}/Kg", y, font_size=9)
                    y -= 4
                    y = draw_text(f"Total declarado: {base_kilos:,.0f} Kg × {divisa_sym} {valor_kg:,.4f}/Kg = {divisa_sym} {base_imp_orig:,.2f}  →  $ {base_imp_ars:,.2f}", y, font_size=10, bold=True)
                else:
                    y = draw_text(f"Total declarado: {base_kilos:,.0f} Kg  ×  $ {valor_kg:,.4f}/Kg = $ {base_imp_ars:,.2f}", y, font_size=10, bold=True)
                y -= 6
                y = draw_wrapped("Los valores a continuación serían estimativos. La liquidación definitiva la practicaría el área de valoración.", y, font_size=9)
                y -= 8
                tributos = self.get_tributos_activos() or [("Derechos de Importación/Exportación (ref.)", 8.0), ("Anticipo Ganancias (ref.)", 0.5)]
                c.setFont("Helvetica-Bold", 9)
                c.drawString(margin, y, "CONCEPTO"); c.drawString(270, y, "ALÍC."); c.drawString(345, y, "BASE (ARS)"); c.drawString(475, y, "IMPORTE EST. (ARS)")
                y -= 3; c.line(margin, y, w - margin, y); y -= 14
                c.setFont("Helvetica", 9)
                total_trib = 0
                for conc, alic in tributos:
                    importe = base_imp_ars * alic / 100; total_trib += importe
                    c.drawString(margin, y, conc); c.drawString(270, y, f"{alic}%")
                    c.drawString(345, y, f"$ {base_imp_ars:,.2f}"); c.drawString(475, y, f"$ {importe:,.2f}")
                    y -= 14
                c.setFont("Helvetica-Bold", 10)
                c.drawString(margin, y, "TOTAL ESTIMADO SOBRE TOTAL CARGA (referencia):")
                c.drawString(475, y, f"$ {total_trib:,.2f}")
                y -= 20; c.line(margin, y, w - margin, y); y -= 10
                # Equivalencia en divisa extranjera (solo si hay T/C cargado)
                if tipo_cambio > 0 and divisa != "ARS":
                    # ya mostramos en ARS arriba; agregamos nota en divisa original
                    c.setFont("Helvetica", 9)
                    c.drawString(margin, y, f"Monto en {divisa}: {divisa_sym} {total_trib / tipo_cambio:,.2f}  (T/C: 1 {divisa} = $ {tipo_cambio:,.2f})")
                    y -= 14
                elif tipo_cambio > 0 and divisa == "ARS":
                    # Precio en ARS, mostramos equiv en USD usando el T/C del documento
                    c.setFont("Helvetica", 9)
                    c.drawString(margin, y, f"Equiv. est. en USD: U$S {total_trib / tipo_cambio:,.2f}  (T/C ref.: 1 U$S = $ {tipo_cambio:,.2f})")
                    y -= 14
            else:
                y = draw_text("(Sin valor unitario — liquidación pendiente de valoración oficial)", y, font_size=10)
        else:
            # CARGO: base = solo el faltante/sobrante en KILOS
            y = draw_wrapped(
                f"En el presente caso se habría constatado un {lbl_tipo} de {diff_abs_k:,.0f} kilogramos "
                f"({pct_max:.2f}% del total declarado). El cargo tributario correspondería ÚNICAMENTE "
                f"a la diferencia detectada de {diff_abs_k:,.0f} kg, conforme Arts. 790/791 C.A.",
                y, font_size=10)
            y -= 8
            if valor_kg > 0:
                usa_tc_pdf = (divisa != "ARS") and (tipo_cambio > 0)
                base_imp_orig = diff_abs_k * valor_kg
                base_imp_ars  = base_imp_orig * tipo_cambio if usa_tc_pdf else base_imp_orig
                if usa_tc_pdf:
                    y = draw_text(f"T/C: {divisa_sym} 1 = $ {tipo_cambio:,.2f}  →  precio ARS: $ {valor_kg * tipo_cambio:,.4f}/Kg", y, font_size=9)
                    y -= 4
                    y = draw_text(f"{lbl_tipo}: {diff_abs_k:,.0f} Kg × {divisa_sym} {valor_kg:,.4f}/Kg = {divisa_sym} {base_imp_orig:,.2f}  →  $ {base_imp_ars:,.2f}", y, font_size=10, bold=True)
                else:
                    y = draw_text(f"{lbl_tipo}: {diff_abs_k:,.0f} Kg  ×  $ {valor_kg:,.4f}/Kg = $ {base_imp_ars:,.2f}", y, font_size=10, bold=True)
                y -= 6
                y = draw_wrapped("Los valores a continuación serían estimativos. La liquidación definitiva la practicaría el área de valoración.", y, font_size=9)
                y -= 8
                tributos = self.get_tributos_activos() or [("Derechos de Importación/Exportación (ref.)", 8.0), ("Anticipo Ganancias (ref.)", 0.5)]
                c.setFont("Helvetica-Bold", 9)
                c.drawString(margin, y, "CONCEPTO"); c.drawString(270, y, "ALÍC."); c.drawString(345, y, "BASE (ARS)"); c.drawString(475, y, "IMPORTE EST. (ARS)")
                y -= 3; c.line(margin, y, w - margin, y); y -= 14
                c.setFont("Helvetica", 9)
                total_trib = 0
                for conc, alic in tributos:
                    importe = base_imp_ars * alic / 100; total_trib += importe
                    c.drawString(margin, y, conc); c.drawString(270, y, f"{alic}%")
                    c.drawString(345, y, f"$ {base_imp_ars:,.2f}"); c.drawString(475, y, f"$ {importe:,.2f}")
                    y -= 14
                c.setFont("Helvetica-Bold", 10)
                c.drawString(margin, y, f"TOTAL ESTIMADO CARGO SOBRE {lbl_tipo} (referencia):")
                c.drawString(475, y, f"$ {total_trib:,.2f}")
                y -= 20; c.line(margin, y, w - margin, y); y -= 10
                if tipo_cambio > 0 and divisa != "ARS":
                    c.setFont("Helvetica", 9)
                    c.drawString(margin, y, f"Monto en {divisa}: {divisa_sym} {total_trib / tipo_cambio:,.2f}  (T/C: 1 {divisa} = $ {tipo_cambio:,.2f})")
                    y -= 14
                elif tipo_cambio > 0 and divisa == "ARS":
                    c.setFont("Helvetica", 9)
                    c.drawString(margin, y, f"Equiv. est. en USD: U$S {total_trib / tipo_cambio:,.2f}  (T/C ref.: 1 U$S = $ {tipo_cambio:,.2f})")
                    y -= 14
            else:
                y = draw_text("(Sin valor unitario — liquidación pendiente de valoración oficial)", y, font_size=10)

        # === VI. CONCLUSIÓN ===
        y = new_page_if_needed(y, 120)
        y -= 10
        y = draw_text("VI. CONCLUSIÓN", y, font_size=11, bold=True)
        y -= 5
        lbl_tipo_final = "FALTANTE" if es_faltante else "SOBRANTE"
        if es_denuncia:
            y = draw_wrapped(
                f"En mérito a lo expuesto, se elevarían las presentes actuaciones a fin de que la superioridad "
                f"evaluara la posible instrucción de sumario al amparo del {art_principal}, en virtud del "
                f"{lbl_tipo_final} detectado de {diff_abs_k:,.0f} kilogramos ({pct_max:.2f}% del total declarado), "
                f"sin perjuicio del cargo tributario que eventualmente correspondiere. "
                f"Es todo cuanto se informaría.",
                y, font_size=10)
        else:
            y = draw_wrapped(
                f"En mérito a lo expuesto, se elevarían las presentes actuaciones a fin de que la superioridad "
                f"evaluara la procedencia del cargo tributario en los términos del {art_principal}, "
                f"en virtud del {lbl_tipo_final} detectado de {diff_abs_k:,.0f} kilogramos "
                f"({pct_max:.2f}% del total declarado). Es todo cuanto se informaría.",
                y, font_size=10)
        y -= 15

        try:
            c.showPage()
            if shared_canvas is None:
                c.save()
        except Exception as _pdf_err:
            messagebox.showerror("Error PDF Cargo", f"Error al finalizar PDF de cargo:\n{_pdf_err}")
            return None
        # Si se usó canvas compartido, las páginas ya están en él, no guardar archivo
        if shared_canvas is not None:
            return None
        # Si es un buffer override (embed en PDF unificado), no guardar archivo
        if _buf_override is not None:
            return None
        # Escribir archivo del cargo (standalone)
        _buf.seek(0)
        _data = _buf.read()
        if full_path:
            try:
                with open(full_path, "wb") as f_out:
                    f_out.write(_data)
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo guardar {filename}: {e}")
                return None
            return full_path
        return None

