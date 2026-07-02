import hashlib

try:
    _real_md5 = hashlib.md5
    def patched_md5(data=b'', **kwargs):
        # Eliminamos el argumento 'usedforsecurity' si existe, porque Python 3.8 no lo entiende
        if 'usedforsecurity' in kwargs:
            del kwargs['usedforsecurity']
        return _real_md5(data, **kwargs)
    hashlib.md5 = patched_md5
except:
    pass


import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import json
import os
from datetime import datetime
import math
import base64
from io import BytesIO
import ctypes
import re
import platform
import subprocess
import traceback
import sqlite3, pathlib

from db import (
    _get_db_path, _init_db,
    db_buscar_funcionarios, db_guardar_funcionario,
    db_todos_funcionarios, db_eliminar_funcionario,
    _get_aduana_db_path, _init_aduana_db,
    db_get_aduanas, db_guardar_aduana, db_eliminar_aduana,
    db_get_lugares_operativos, db_guardar_lugar_operativo,
    db_eliminar_lugar_operativo,
    _init_funciones_db,
    db_get_funciones, db_guardar_funcion, db_eliminar_funcion,
)
from assets.icons import ICON_WINDOW_B64, ICON_REPORT_B64

class PlanillaFinalApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Sistema de Medición - V120 (VCF gases: GLP K0/rho2 | GNL alpha=0.00468 | NH3 alpha=0.00226)")
        self.root.geometry("1600x900") 
        
        self.setup_icon()
        
        self.ui_font_size = 9 
        self.tabs_built = {1: False, 2: False, 3: False}
        self.is_loading_data = False

        self.vars = {}
        self.norma_astm = tk.StringVar(value="1980")   # "1980" | "2004"
        self.ddt_data = [] 
        self.ddt_counter = 0
        self.combos_ddt = [] 
        self.funcionarios_data = []
        self.func_counter = 0

        self.lista_tanques = []
        for i in range(1, 9): 
            self.lista_tanques.append(f"TK {i} BABOR")
            self.lista_tanques.append(f"TK {i} ESTRIBOR")
        self.lista_tanques.append("SLOP BABOR")
        self.lista_tanques.append("SLOP ESTRIBOR")
        self.lista_carbonera = ["CARBONERA 1"]

        # --- MENU ---
        menubar = tk.Menu(root)
        filemenu = tk.Menu(menubar, tearoff=0)
        filemenu.add_command(label="Nueva Medición", command=self.nueva_medicion)
        filemenu.add_separator()
        filemenu.add_command(label="Cargar Datos (.meg)", command=self.cargar_datos)
        filemenu.add_command(label="Guardar Datos (.meg)...", command=self.guardar_datos_manual)
        filemenu.add_separator()
        filemenu.add_command(label="Generar Reportes PDF...", command=self.generar_con_seleccion)
        filemenu.add_separator()
        filemenu.add_command(label="Salir", command=root.quit)
        menubar.add_cascade(label="Archivo", menu=filemenu)
        root.config(menu=menubar)

        # --- NOTEBOOK ---
        self.notebook = ttk.Notebook(root)
        self.tab_caratula = ttk.Frame(self.notebook)
        self.tab_ini_prod = ttk.Frame(self.notebook)
        self.tab_fin_prod = ttk.Frame(self.notebook)
        self.tab_preview = ttk.Frame(self.notebook)
        
        self.notebook.add(self.tab_caratula, text=' 1. CARÁTULA (Documentos) ')
        self.notebook.add(self.tab_ini_prod, text=' 2. INICIAL (Completo) ')
        self.notebook.add(self.tab_fin_prod, text=' 3. FINAL (Completo) ')
        self.notebook.add(self.tab_preview, text=' 4. VISTA PREVIA / REPORTES ')
        
        self.notebook.pack(expand=True, fill="both")
        self.notebook.bind("<<NotebookTabChanged>>", self.on_tab_changed)

        self.construir_caratula()
        if not self.ddt_data: self.agregar_ddt_row(def_prod="GASOIL")
        self.lista_carbonera = ["CARBONERA 1"]

        self.status_bar = tk.Label(root, text="V120: Norma VCF seleccionable — 1980 (tablas impresas) | 2004 (API MPMS 11.1)", bd=1, relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        self._update_norma_status()
        self.auto_save_loop()
        
        # --- RESPONSIVE FONTS ---
        self._last_responsive_w = 0
        self._setup_responsive_fonts()
        root.bind("<Configure>", self._on_root_configure)

    def _update_norma_status(self):
        """Actualiza la barra de estado con la norma activa."""
        n = self.norma_astm.get()
        if n == "1980":
            info = "Norma activa: ASTM D1250-1980 (tablas impresas — 4 zonas densidad en 54B)"
        else:
            info = "Norma activa: API MPMS 11.1-2004 (digital — K0=346.4228 K1=0.4033 en 54B)"
        self.status_bar.config(text=info)

    def _on_norma_changed(self):
        """Llamado al cambiar la norma — recalcula todas las celdas visibles."""
        self._update_norma_status()
        # Forzar recálculo de todos los tanques (pestaña activa)
        try:
            for etapa in ("inicial", "final"):
                for tk_name in self.lista_tanques + self.lista_carbonera:
                    self.calc_volumen_prod_ui(etapa, tk_name)
        except Exception:
            pass

    def _vcf_k0k1_1980_54B(self, rho):
        """
        Coeficientes K0/K1 de la tabla 54B según ASTM D1250-1980 (tablas impresas).
        Cuatro zonas de densidad — reproduce exactamente las tablas físicas.

        Zona 1  rho ≤ 770          K0=346.42278  K1=0.43884
        Zona 2  770 < rho < 778    Transición: alpha = A + B/rho²
        Zona 3  778 ≤ rho < 839    K0=594.5418   K1=0
        Zona 4  rho ≥ 839          K0=186.9696   K1=0.48618
        """
        if rho <= 770.0:
            return 346.42278, 0.43884, None        # (k0, k1, alpha_override)
        elif rho < 778.0:
            # Zona de transición — alpha calculado directamente, no por k0/k1
            A, B = -0.0033612, 2680.32
            alpha_override = A + B / (rho ** 2)
            return None, None, alpha_override
        elif rho < 839.0:
            return 594.5418, 0.0, None
        else:
            return 186.9696, 0.48618, None

    def _calc_vcf_exponencial(self, rho, t, k0, k1, alpha_override=None):
        """Fórmula exponencial ASTM: VCF = exp(-α·ΔT·(1 + 0.8·α·ΔT))"""
        if alpha_override is not None:
            alpha = alpha_override
        else:
            alpha = (k0 / (rho ** 2)) + (k1 / rho)
        dt  = t - 15.0
        vcf = math.exp(-alpha * dt * (1.0 + 0.8 * alpha * dt))
        return round(vcf, 5)

    def _setup_responsive_fonts(self):
        """Calcula el tamaño de fuente base según el ancho de la ventana y aplica estilos ttk."""
        w = self.root.winfo_width() or 1600
        h = self.root.winfo_height() or 900
        # Escala: 1600px ancho = tamaño 9. Cap en 11 para evitar X11 BadLength con RENDER
        base = max(7, min(10, int(w / 178)))
        small = max(6, base - 1)
        big   = base + 2
        title = base + 4
        self.ui_font_size = base

        style = ttk.Style()
        style.configure(".", font=("Arial", base))
        style.configure("TLabel",    font=("Arial", base))
        style.configure("TEntry",    font=("Arial", base))
        style.configure("TButton",   font=("Arial", base))
        style.configure("TCombobox", font=("Arial", base))
        style.configure("TNotebook.Tab", font=("Arial", min(8, base), "bold"), padding=[8, 4])
        style.configure("TLabelframe.Label", font=("Arial", min(8, base), "bold"))
        style.configure("Treeview",  font=("Arial", small), rowheight=max(18, base + 8))
        style.configure("Treeview.Heading", font=("Arial", min(8, small), "bold"))
        # Tk native widgets need to be updated too - we store the fonts for reuse
        self._font_base  = ("Arial", base)
        self._font_small = ("Arial", small)
        self._font_big   = ("Arial", min(8, big), "bold")
        self._font_title = ("Arial", min(8, title), "bold")

    def _on_root_configure(self, event):
        """Actualiza fuentes cuando cambia el tamaño de la ventana principal."""
        if event.widget != self.root: return
        w = event.width
        # Solo actualizar si el ancho cambió más de 30px (evitar loops)
        if abs(w - self._last_responsive_w) < 30: return
        self._last_responsive_w = w
        self._setup_responsive_fonts()

    def setup_icon(self):
        try:
            myappid = 'arca.medicion.v108'
            ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
            raw = ICON_WINDOW_B64.strip()
            if "," in raw: raw = raw.split(",")[1]
            icon_data = base64.b64decode(raw)
            icon_image = tk.PhotoImage(data=icon_data)
            self.root.iconphoto(True, icon_image)
            self._icon_image_ref = icon_image  # evitar garbage collection
        except Exception as _ico_err:
            try:
                from PIL import Image, ImageTk
                import io
                raw2 = ICON_WINDOW_B64.strip()
                _img = Image.open(io.BytesIO(base64.b64decode(raw2)))
                _img = _img.resize((32, 32), Image.LANCZOS)
                _photo = ImageTk.PhotoImage(_img)
                self.root.iconphoto(True, _photo)
                self._icon_image_ref = _photo
            except: pass

    def on_tab_changed(self, event):
        current_tab_index = self.notebook.index(self.notebook.select())
        if current_tab_index in self.tabs_built and not self.tabs_built[current_tab_index]:
            self.root.after(50, lambda: self.build_specific_tab(current_tab_index))

    def _update_tab_titles(self):
        """Actualiza los títulos de las solapas con el tipo de medición activo."""
        try:
            _tm = self.get_tipo_medio()
            _tipo_abrev = {
                "BUQUE":              "⚓ BUQ",
                "BARCAZA":            "⚓ BAR",
                "BUQUE GASERO/GLP":   "⚙ GLP",
                "BUQUE QUIMIQUERO":   "⚗ QUI",
                "BUQUE METANERO/GNL": "❄ GNL",
                "DRAFT SURVEY":       "⚓ DFT",
                "TANQUE FIJO":        "🛢 TF",
                "TANQUE FLOTANTE":    "🛢 TTF",
                "ESFERA DE GAS":      "⚙ ESF",
                "CAMION CISTERNA":    "🚛 CAM",
                "CAMION GAS/GLP":     "🚛 CGP",
                "OLEODUCTO":          "≋ OLE",
                "GASODUCTO":          "≋ GAS",
                "POLIDUCTO":          "≋ POL",
                "ELECTRICO":          "⚡ ELC",
            }.get(_tm, _tm[:6])
            self.notebook.tab(1, text=f" 2. INICIAL [{_tipo_abrev}] ")
            self.notebook.tab(2, text=f" 3. FINAL   [{_tipo_abrev}] ")
        except:
            pass

    def build_specific_tab(self, index):
        if index == 1:
            self.construir_pantalla_unificada("inicial", self.tab_ini_prod)
        elif index == 2:
            self.construir_pantalla_unificada("final", self.tab_fin_prod)
        elif index == 3:
            self.construir_preview_tab()
        self.tabs_built[index] = True
        self.actualizar_combos_ddt()
        self._update_tab_titles()

    def _limpiar_vars_tipo_anterior(self):
        """Limpia variables de medición que no corresponden al tipo actual, evitando datos fantasma."""
        tm = self.get_tipo_medio()
        # Keys del tipo actual
        current_tanks = set(self.lista_tanques + self.lista_carbonera)
        # Prefijos de variables de tanque para los stages
        ETAPAS = ["ini", "fin"]
        # Sufijos que son exclusivos de tipos específicos — se limpian al cambiar de tipo
        SUFIJOS_MARITIMO  = ["_alt_uti","_num_uti","_alt_ref","_s_corr","_prod_s1","_prod_l1","_prod_s2","_prod_l2","_agua_s1","_agua_l1","_agua_s2","_agua_l2","_agua_lectura","_agua_desc","_agua_s_real","_vol_nat_agua"]
        SUFIJOS_TIERRA    = ["_s_tierra","_tf_offset","_vol_bruto"]
        SUFIJOS_CAMION    = ["_ticket_l","_ticket_k"]
        SUFIJOS_GAS       = ["_presion","_temp_liq","_factor_z","_dens_vapor","_fase","_vol_liq","_vol_vap","_masa_liq","_masa_vap","_pres_gnl","_temp_vap"]
        SUFIJOS_ELECTRICO = ["_el_ini_act","_el_fin_act","_el_kwh_act","_el_ini_rea","_el_fin_rea","_el_kwh_rea","_el_const","_el_fp","_el_dem","_el_V","_el_A","_el_VA","_el_fp_med","_el_fases"]
        SUFIJOS_DUCTO     = ["_cont_ini","_cont_fin","_vol_linea","_caudal_mh","_coriolis_kgh","_masa_coriolis","_P_lin","_T_lin","_Z","_vol_base","_vol_base_km3","_pig_diam","_pig_largo","_pig_vol","_pig_vel","_pig_obs"]
        SUFIJOS_ESFERA    = ["_esf_pres","_esf_temp","_esf_dens","_esf_vol_gas","_esf_masa","_esf_fase"]
        SUFIJOS_CAMION_GAS= ["_cg_pres","_cg_temp","_cg_dens","_cg_masa","_cg_vol","_cg_vol_gas"]
        SUFIJOS_GNL_MOL   = ["_gc_CH4","_gc_C2H6","_gc_C3H8","_gc_C4H10","_gc_N2","_gc_CO2","_gc_iC4","_gc_nC5","_gc_H2S","_gc_sum","_gc_PM"]

        # Sufijos comunes de medición — se limpian cuando el TIPO CAMBIA entre familias
        # (evita que datos de buque aparezcan en camión o viceversa)
        SUFIJOS_COMUNES_MEDICION = [
            "_dens_lab","_dens_doc","_dens_salida","_v15_lab","_kv_lab","_v15_doc","_kv_doc",
            "_v15_sal","_kv_sal","_vol_nat_prod","_vol_nat_agua","_prod_name","_ddt_assign",
            "_tabla_vcf","_temp","_num_uti","_desc_tubo","_alt_uti","_alt_ref"
        ]

        _mar  = self.es_maritimo()
        _tie  = self.es_tierra()
        _cam  = self.es_camion()
        _gas  = self.es_gasero() and not ("METANERO" in tm or "GNL" in tm)
        _met  = "METANERO" in tm or "GNL" in tm
        _esf  = self.es_esfera()
        _duc  = self.es_ducto()
        _el   = self.es_electrico()
        _cgb  = self.es_camion_gas()

        # Decide qué sufijos limpiar según el tipo ACTUAL
        sufijos_a_limpiar = []
        if not _mar:                  sufijos_a_limpiar += SUFIJOS_MARITIMO
        if not _tie:                  sufijos_a_limpiar += SUFIJOS_TIERRA
        if not _cam:                  sufijos_a_limpiar += SUFIJOS_CAMION
        if not (_gas or _met or _esf or _cgb): sufijos_a_limpiar += SUFIJOS_GAS
        if not _el:                   sufijos_a_limpiar += SUFIJOS_ELECTRICO
        if not _duc:                  sufijos_a_limpiar += SUFIJOS_DUCTO
        if not _esf:                  sufijos_a_limpiar += SUFIJOS_ESFERA
        if not _cgb:                  sufijos_a_limpiar += SUFIJOS_CAMION_GAS
        if not _met:                  sufijos_a_limpiar += SUFIJOS_GNL_MOL

        if not sufijos_a_limpiar: return
        sufijos_a_limpiar_set = set(sufijos_a_limpiar)

        # Pre-compute prefixes and suffix tuples for fast matching
        prefixes = {f"{et}_": len(f"{et}_") for et in ETAPAS}
        sufijos_tuple = tuple(sufijos_a_limpiar_set)

        # Limpiar variables: buscar claves que coincidan con etapa_tanque_sufijo
        # IMPORTANTE: se limpian TODOS los tanques (no solo los que ya no existen)
        # para evitar datos fantasma cuando el tipo cambia
        keys_to_clear = set()
        all_keys = list(self.vars.keys())
        for k in all_keys:
            if not k.endswith(sufijos_tuple):
                continue
            for prefix, plen in prefixes.items():
                if k.startswith(prefix):
                    for suf in sufijos_a_limpiar_set:
                        if k.endswith(suf) and len(k) > plen + len(suf):
                            keys_to_clear.add(k)
                            break
                    break
        for k in keys_to_clear:
            if k in self.vars:
                self.vars[k].set("")

        # Limpiar también los tanques que ya no existen — incluye sufijos comunes
        sufijos_no_existentes = sufijos_a_limpiar_set | set(SUFIJOS_COMUNES_MEDICION)
        sufijos_ne_tuple = tuple(sufijos_no_existentes)
        keys_to_clear2 = set()
        for k in all_keys:
            if k in keys_to_clear:
                continue
            if not k.endswith(sufijos_ne_tuple):
                continue
            for prefix, plen in prefixes.items():
                if k.startswith(prefix):
                    for suf in sufijos_no_existentes:
                        if k.endswith(suf) and len(k) > plen + len(suf):
                            tank_part = k[plen:-len(suf)]
                            if tank_part and tank_part not in current_tanks:
                                keys_to_clear2.add(k)
                                break
                    break
        for k in keys_to_clear2:
            if k in self.vars:
                self.vars[k].set("")

        # Limpiar variables de cabecera de etapa que no aplican al tipo actual
        vars_header_maritimo = [
            "ini_Calados Proa","ini_Calados Popa","ini_Calados Babor","ini_Calados Estribor",
            "ini_Trimación","ini_Lista",
            "fin_Calados Proa","fin_Calados Popa","fin_Calados Babor","fin_Calados Estribor",
            "fin_Trimación","fin_Lista",
        ]
        vars_header_tierra = [
            "ini_Temp_Amb","fin_Temp_Amb","ini_Modo_Sondaje","fin_Modo_Sondaje",
        ]
        vars_header_ducto = [
            "ini_P_linea","ini_T_linea","ini_Cond_base","fin_P_linea","fin_T_linea","fin_Cond_base",
        ]
        if not _mar:
            for vk in vars_header_maritimo:
                if vk in self.vars: self.vars[vk].set("")
        if not (_tie or _cam):
            for vk in vars_header_tierra:
                if vk in self.vars: self.vars[vk].set("")
        if not _duc:
            for vk in vars_header_ducto:
                if vk in self.vars: self.vars[vk].set("")

    def rebuild_all_tabs(self):
        # Limpiar variables que no corresponden al tipo actual
        self._limpiar_vars_tipo_anterior()
        for idx in [1, 2, 3]:
            self.tabs_built[idx] = False
            if idx == 1: tab_frame = self.tab_ini_prod
            elif idx == 2: tab_frame = self.tab_fin_prod
            else: tab_frame = self.tab_preview
            for w in tab_frame.winfo_children(): w.destroy()
        self.combos_ddt = []
        current = self.notebook.index(self.notebook.select())
        if current in self.tabs_built:
            self.build_specific_tab(current)
        self._update_tab_titles()

    def get_var(self, key, default=""):
        if key not in self.vars: self.vars[key] = tk.StringVar(value=default)
        return self.vars[key]

    # Actores propios de cada declaración detallada (con su CUIT)
    DDT_ACTOR_KEYS = ("despachante", "cuit_desp", "impexp", "cuit_impexp", "ata", "cuit_ata")
    # Mapeo actor del DDT → campo global de la carátula (fallback)
    DDT_ACTOR_GLOBAL = {
        "despachante": "car_despachante", "cuit_desp": "car_cuit_desp",
        "impexp": "car_impexp", "cuit_impexp": "car_cuit_impexp",
        "ata": "car_ata", "cuit_ata": "car_cuit_ata",
    }

    def _ddt_actor(self, ddt_obj, key):
        """Actor (despachante/imex/ATA o CUIT) del documento; si el documento
        no lo tiene cargado, cae a la carátula global."""
        v = ""
        try:
            sv = ddt_obj.get(key) if ddt_obj else None
            v = sv.get().strip() if isinstance(sv, tk.StringVar) else str(sv or "").strip()
        except: pass
        return v or self.get_var(self.DDT_ACTOR_GLOBAL[key]).get()

    def _actores_pdf(self, ddt_obj=None):
        """Actores a imprimir en un PDF. Con ddt_obj: los de ese documento
        (fallback carátula). Sin ddt_obj (reporte general): los valores
        distintos entre todos los documentos, unidos con ' / '."""
        res = {}
        for k in self.DDT_ACTOR_KEYS:
            if ddt_obj is not None:
                res[k] = self._ddt_actor(ddt_obj, k)
            else:
                vistos = []
                for d in self.ddt_data:
                    sv = d.get(k)
                    val = sv.get().strip() if isinstance(sv, tk.StringVar) else ""
                    if val and val not in vistos: vistos.append(val)
                res[k] = " / ".join(vistos) if vistos else self.get_var(self.DDT_ACTOR_GLOBAL[k]).get()
        return res

    def parse_float(self, value):
        if not value: return 0.0
        if isinstance(value, (int, float)): return float(value)
        val_str = str(value).strip()
        if not val_str: return 0.0
        if ',' in val_str and '.' in val_str:
            if val_str.rfind(',') > val_str.rfind('.'): val_str = val_str.replace('.', '').replace(',', '.')
            else: val_str = val_str.replace(',', '')
        elif ',' in val_str: val_str = val_str.replace(',', '.')
        val_str = re.sub(r'[^\d.-]', '', val_str)
        try: return float(val_str)
        except: return 0.0

    # Lookup de códigos de aduana → nombre
    _ADUANA_LOOKUP = {
        "001":"BUENOS AIRES","004":"IGUAZU REMOTO","006":"BAHIA BLANCA",
        "008":"BARRANQUERAS","009":"BELEN","010":"BUENOS AIRES NORTE",
        "011":"CLORINDA","012":"COLON (ER)","013":"COLON (BA)",
        "014":"COMODORO RIVADAVIA","015":"CONCEPCION DEL URUGUAY","016":"CONCORDIA",
        "017":"CORDOBA","018":"CORRIENTES","019":"PUERTO DESEADO",
        "020":"DIAMANTE","023":"ESQUEL","024":"FORMOSA","025":"GOYA",
        "026":"GUALEGUAYCHU","029":"IGUAZU","031":"JUJUY","033":"LA PLATA",
        "034":"LA QUIACA","037":"MAR DEL PLATA","038":"MENDOZA","040":"NECOCHEA",
        "041":"PARANA","042":"PASO DE LOS LIBRES","045":"POCITOS","046":"POSADAS",
        "047":"PUERTO MADRYN","048":"RIO GALLEGOS","049":"RIO GRANDE","052":"ROSARIO",
        "053":"SALTA","054":"SAN JAVIER","055":"SAN JUAN","057":"SAN LORENZO",
        "058":"SAN MARTIN DE LOS ANDES","059":"SAN NICOLAS","060":"SAN PEDRO",
        "061":"SANTA CRUZ","062":"SANTA FE","066":"TINOGASTA","067":"USHUAIA",
        "069":"VILLA CONSTITUCION","073":"EZEIZA","074":"TUCUMAN","075":"NEUQUEN",
        "076":"ORAN","078":"SAN RAFAEL","079":"LA RIOJA","080":"SAN ANTONIO OESTE",
        "082":"BERNARDO DE YRIGOYEN","083":"SAN LUIS","084":"SANTO TOME",
        "085":"VILLA REGINA","086":"OBERA","087":"CALETA OLIVIA",
        "088":"GENERAL DEHEZA","089":"SANTIAGO DEL ESTERO","090":"GENERAL PICO",
        "091":"BS.AS. NORTE","092":"BS.AS. SUR","093":"RAFAELA","099":"MULTIADUANA",
    }


    def aduana_nombre(self):
        """Devuelve solo el nombre de la aduana (consulta DB + dict estático)."""
        raw = self.get_var("car_lugar").get().strip()
        if " - " in raw:
            return raw.split(" - ", 1)[1].strip()
        cod = self.aduana_codigo()
        if cod:
            # Primero intenta en la DB dinámica
            try:
                aduanas_db = db_get_aduanas()
                for a in aduanas_db:
                    if a["codigo"] == cod:
                        return a["nombre"]
            except: pass
            # Fallback al lookup estático
            if cod in self._ADUANA_LOOKUP:
                return self._ADUANA_LOOKUP[cod]
        if raw and not raw[:3].isdigit():
            return raw
        return raw or "USHUAIA"

    def aduana_completa(self):
        """Devuelve la cadena completa '067 - USHUAIA'."""
        raw = self.get_var("car_lugar").get().strip()
        if " - " in raw:
            return raw
        cod = self.aduana_codigo()
        nom = self.aduana_nombre()
        if cod and nom:
            return f"{cod} - {nom}"
        return raw or "USHUAIA"

    def aduana_codigo(self):
        """Devuelve el código numérico de la aduana con ceros a la izquierda (3 dígitos).
        Ej: '67 - USHUAIA' → '067', '1 - BS.AS.' → '001'.
        Fallback: extrae el código del MANI (posición 2-5 del número).
        """
        raw = self.get_var("car_lugar").get().strip()
        if " - " in raw:
            cod = raw.split(" - ", 1)[0].strip()
            return cod.zfill(3) if cod.isdigit() else cod
        # Fallback: extraer del MANI (formato "26067MANI000618G" → "067")
        mani = self.get_var("car_mani").get().strip()
        if len(mani) >= 5 and mani[:2].isdigit() and mani[2:5].isdigit():
            return mani[2:5]
        return ""

    def inferir_tipo_operacion(self, ddt_obj):
        """Infiere tipo de operación a partir del subrégimen del documento.
        Para REMO: determina automáticamente carga/descarga comparando
        la aduana del documento con la aduana actual (car_lugar).
          - aduana_doc == aduana_actual → CARGA (Art. 959)
          - aduana_doc != aduana_actual → DESCARGA (Art. 954)
        Solo pide confirmación si no se puede determinar del número.
        """
        import re as _re
        num = ddt_obj["numero"].get().upper().strip()
        
        # Extraer subrégimen del número
        subr = ""
        for sr in ["ER01","ER02","ER03","ER05","ER06","ER07","ER08",
                   "EC01","EC03","EC06","IC04","IC05","IC06","IC07","IC65",
                   "IR01","IR02","IS01","IS04","IT01","IT04","IT14",
                   "RE01","REMO","TRAN","ZF01","ZF06"]:
            if sr in num:
                subr = sr
                break
        if not subr:
            td = ddt_obj.get("tipo_doc", None)
            if td:
                tdv = td.get() if hasattr(td, "get") else str(td)
                if "Rancho" in tdv: subr = "ER01"
                elif "Expediente" in tdv: subr = "RE01"

        # IC/IR/IS/IT → Importación (Art. 954)
        if subr.startswith("IC") or subr.startswith("IR") or subr.startswith("IS") or subr.startswith("IT"):
            return {"tipo":"importacion","art_principal":"Art. 954 del Código Aduanero",
                    "art_inc":"Art. 954 inc. c) C.A.","descripcion":f"IMPORTACIÓN ({subr})",
                    "codigo":subr,"necesita_pregunta":False}
        
        # EC* → Exportación (Art. 959)
        if subr.startswith("EC"):
            return {"tipo":"exportacion","art_principal":"Art. 959 del Código Aduanero",
                    "art_inc":"Art. 959 inc. c) C.A.","descripcion":f"EXPORTACIÓN ({subr})",
                    "codigo":subr,"necesita_pregunta":False}
        
        # Rancho bandera extranjera
        if subr in ("ER01","ER05"):
            return {"tipo":"rancho_ext","art_principal":"Art. 959 del Código Aduanero",
                    "art_inc":"Art. 959 inc. c) C.A.",
                    "descripcion":f"RANCHO DE COMBUSTIBLE - BANDERA EXTRANJERA ({subr})",
                    "codigo":subr,"necesita_pregunta":False}
        
        # Rancho bandera argentina
        if subr in ("ER02","ER06"):
            return {"tipo":"rancho_arg","art_principal":"Art. 954 del Código Aduanero",
                    "art_inc":"Art. 954 inc. c) C.A.",
                    "descripcion":f"RANCHO DE COMBUSTIBLE - BANDERA ARGENTINA ({subr})",
                    "codigo":subr,"necesita_pregunta":False}
        
        # REMO → determinar por comparación de aduanas (NUNCA preguntar para Detallada)
        if subr == "REMO":
            tipo_doc_val = ""
            td = ddt_obj.get("tipo_doc", None)
            if td: tipo_doc_val = td.get() if hasattr(td, "get") else str(td)
            es_detallada = "detallada" in tipo_doc_val.lower() or tipo_doc_val == ""
            aduana_actual = self.aduana_codigo()  # ej: "067"
            # El número de REMO tiene el código de aduana origen embebido
            # Formato típico: "AANNNNREMONNNNN" donde NNNN = cod aduana (4 dígitos o 3)
            # Intentar extraer los dígitos antes de "REMO"
            # Formato: YY(2) + ADUANA(3) + "REMO" + ...
            # Extrae los 3 dígitos en posición [2:5] del número
            remo_pos = num.find("REMO")
            if remo_pos >= 5:
                aduana_doc = num[remo_pos-3:remo_pos]          # mantener ceros: "033"
            else:
                m = _re.search(r"(\d{2})(\d{3})REMO", num)
                aduana_doc = m.group(2) if m else ""            # ya son 3 dígitos
            aduana_doc_padded = aduana_doc.zfill(3)            # asegurar 3 dígitos
            aduana_act_stripped = aduana_actual.lstrip("0")
            
            if aduana_doc and aduana_actual:
                if aduana_doc.lstrip("0") == aduana_act_stripped:
                    # Misma aduana → CARGA
                    return {"tipo":"remo_carga","art_principal":"Art. 959 del Código Aduanero",
                            "art_inc":"Art. 959 inc. c) C.A.",
                            "descripcion":"TRANSFERENCIA DE COMBUSTIBLE A LA CARGA (REMO) - MISMA ADUANA",
                            "codigo":"REMO","necesita_pregunta":False}
                else:
                    # Distinta aduana → DESCARGA
                    return {"tipo":"remo_descarga","art_principal":"Art. 954 del Código Aduanero",
                            "art_inc":"Art. 954 inc. c) C.A.",
                            "descripcion":f"REMO - DESCARGA - ADUANA ORIGEN {aduana_doc_padded} - {self._ADUANA_LOOKUP.get(aduana_doc_padded, aduana_doc_padded)}",
                            "codigo":"REMO","necesita_pregunta":False}
            # No se pudo determinar del número
            if es_detallada:
                # Detallada: nunca preguntar → asumir descarga (viene de otra aduana)
                return {"tipo":"remo_descarga","art_principal":"Art. 954 del Código Aduanero",
                        "art_inc":"Art. 954 inc. c) C.A.",
                        "descripcion":"TRANSFERENCIA DE COMBUSTIBLE A LA DESCARGA (REMO)",
                        "codigo":"REMO","necesita_pregunta":False}
            # No es Detallada → pedir confirmación
            return {"tipo":"remo","art_principal":"Art. 954 del Código Aduanero",
                    "art_inc":"Art. 954 inc. c) C.A.",
                    "descripcion":"TRANSFERENCIA DE COMBUSTIBLE (REMO)",
                    "codigo":"REMO","necesita_pregunta":True}
        
        # Default (RE01, TRAN, ZF*, etc)
        return {"tipo":"otro","art_principal":"Art. 954 del Código Aduanero",
                "art_inc":"Art. 954 inc. c) C.A.",
                "descripcion":f"OPERACIÓN ADUANERA ({subr or 'DETALLADA'})",
                "codigo":subr or "","necesita_pregunta":False}

    def clean_filename(self, text):
        return re.sub(r'[\\/*?:"<>|]', "", str(text)).strip()

    def validar_fecha(self, new_value):
        """Valida formato DD/MM/AAAA permitiendo escritura progresiva"""
        if new_value == "" or new_value == "DD/MM/AAAA":
            return True
        # Permitir solo digitos y /
        cleaned = new_value.replace("/", "")
        if not cleaned.isdigit():
            return False
        if len(new_value) > 10:
            return False
        # Auto-insertar / despues de DD y MM
        # Validar parcialmente segun lo que se lleva escrito
        parts = new_value.split("/")
        if len(parts) > 3:
            return False
        # Validar dia
        if parts[0]:
            if len(parts[0]) > 2: return False
            if len(parts[0]) == 2:
                d = int(parts[0])
                if d < 1 or d > 31: return False
        # Validar mes
        if len(parts) > 1 and parts[1]:
            if len(parts[1]) > 2: return False
            if len(parts[1]) == 2:
                m = int(parts[1])
                if m < 1 or m > 12: return False
        # Validar año
        if len(parts) > 2 and parts[2]:
            if len(parts[2]) > 4: return False
        return True

    def validar_hora(self, new_value):
        """Valida formato HH:MM permitiendo escritura progresiva"""
        if new_value == "" or new_value == "HH:MM":
            return True
        cleaned = new_value.replace(":", "")
        if not cleaned.isdigit():
            return False
        if len(new_value) > 5:
            return False
        parts = new_value.split(":")
        if len(parts) > 2:
            return False
        if parts[0]:
            if len(parts[0]) > 2: return False
            if len(parts[0]) == 2:
                h = int(parts[0])
                if h > 23: return False
        if len(parts) > 1 and parts[1]:
            if len(parts[1]) > 2: return False
            if len(parts[1]) == 2:
                m = int(parts[1])
                if m > 59: return False
        return True

    def on_fecha_focus_in(self, event, var):
        if var.get() == "DD/MM/AAAA":
            var.set("")

    def on_fecha_focus_out(self, event, var):
        if var.get() == "":
            var.set("DD/MM/AAAA")

    def on_hora_focus_in(self, event, var):
        if var.get() == "00:00" or var.get() == "HH:MM":
            var.set("")

    def on_hora_focus_out(self, event, var):
        if var.get() == "":
            var.set("00:00")

    def validar_cuit(self, cuit_str):
        """Valida CUIT/CUIL argentino. Retorna True/False y mensaje."""
        if not cuit_str or cuit_str.strip() == "":
            return True, ""
        clean = cuit_str.replace("-", "").replace(" ", "")
        if not clean.isdigit():
            return False, "Solo números y guiones"
        if len(clean) != 11:
            return False, f"Debe tener 11 dígitos (tiene {len(clean)})"
        # Verificar dígito verificador
        mult = [5, 4, 3, 2, 7, 6, 5, 4, 3, 2]
        total = sum(int(clean[i]) * mult[i] for i in range(10))
        resto = 11 - (total % 11)
        if resto == 11: dv = 0
        elif resto == 10: dv = 9
        else: dv = resto
        if int(clean[10]) != dv:
            return False, "Dígito verificador inválido"
        return True, "OK"

    def on_cuit_validate(self, widget, var):
        """Valida CUIT al salir del campo - solo feedback visual, sin popup."""
        val = var.get().strip()
        if not val:
            widget.configure(background="white")
            return
        clean = val.replace("-", "").replace(" ", "")
        if len(clean) < 11:
            # Aún incompleto, no marcar error
            widget.configure(background="#fff9c4")  # amarillo suave = en progreso
            return
        ok, msg = self.validar_cuit(val)
        if ok:
            widget.configure(background="#e8f5e9")  # verde claro
        else:
            widget.configure(background="#ffcdd2")  # rojo claro

    def _init_tanks_for_tipo(self, tipo_medio):
        """Inicializa lista de tanques apropiada según tipo_medio."""
        tm = tipo_medio
        self.lista_tanques = []
        self.lista_carbonera = []
        if tm in ("BUQUE", "BARCAZA"):
            for i in range(1, 9):
                self.lista_tanques.append(f"TK {i} BABOR")
                self.lista_tanques.append(f"TK {i} ESTRIBOR")
            self.lista_tanques += ["SLOP BABOR", "SLOP ESTRIBOR"]
            self.lista_carbonera = ["CARBONERA 1"]
        elif tm in ("BUQUE GASERO/GLP",):
            # Buque gasero tipo MOSS o similar: 4 tanques esféricos/prismáticos
            for i in range(1, 5):
                self.lista_tanques.append(f"TK {i}")
            self.lista_carbonera = ["CARBONERA 1"]
        elif tm in ("BUQUE QUIMIQUERO",):
            for i in range(1, 7):
                self.lista_tanques.append(f"TK {i} BABOR")
                self.lista_tanques.append(f"TK {i} ESTRIBOR")
            self.lista_carbonera = ["CARBONERA 1"]
        elif tm in ("BUQUE METANERO/GNL",):
            for i in range(1, 5):
                self.lista_tanques.append(f"TANQUE {i}")
            self.lista_carbonera = ["CARBONERA 1"]
        elif tm in ("TANQUE FIJO", "TANQUE FLOTANTE"):
            self.lista_tanques.append("TANQUE 1")
        elif tm == "ESFERA DE GAS":
            self.lista_tanques.append("ESFERA 1")
        elif "CAMION" in tm:
            n_comps = 1 if "GAS" in tm else 6
            for i in range(1, n_comps + 1):
                self.lista_tanques.append(f"COMPARTIMENTO {i}")
        elif tm == "DRAFT SURVEY":
            for i in range(1, 9):
                self.lista_tanques.append(f"TK {i} BABOR")
                self.lista_tanques.append(f"TK {i} ESTRIBOR")
            self.lista_tanques += ["SLOP BABOR", "SLOP ESTRIBOR"]
            self.lista_carbonera = ["CARBONERA 1"]
        elif self.es_ducto():
            self.lista_tanques = ["TRAMO 1"]
        elif tm == "MEDICION ELECTRICA":
            self.lista_tanques = ["MEDIDOR 1"]
        else:
            for i in range(1, 5): self.lista_tanques.append(f"TK {i}")

    def wizard_nueva_medicion(self):
        """Diálogo de inicio: pregunta tipo de medio y operación antes de empezar."""
        wizard = tk.Toplevel(self.root)
        wizard.title("Nueva Medición — Configuración Inicial")
        wizard.grab_set()
        wizard.transient(self.root)
        wizard.resizable(True, True)
        wizard.update_idletasks()
        sw, sh = wizard.winfo_screenwidth(), wizard.winfo_screenheight()
        w_dlg = min(720, sw - 40)
        h_dlg = min(640, sh - 80)
        wizard.geometry(f"{w_dlg}x{h_dlg}+{(sw-w_dlg)//2}+{(sh-h_dlg)//2}")
        wizard.minsize(560, 480)

        # ── Header (fijo arriba) ──────────────────────────────────────────────
        fh = tk.Frame(wizard, bg="#1B3A5C")
        fh.pack(fill="x", side="top")
        tk.Label(fh, text="NUEVA MEDICIÓN", bg="#1B3A5C", fg="white",
                 font=("Arial", 13, "bold")).pack(pady=10)
        tk.Label(fh, text="Seleccioná el tipo de medición y la operación",
                 bg="#1B3A5C", fg="#AED6F1", font=("Arial", 9)).pack(pady=(0,8))

        # ── Footer con botones (fijo abajo — DEBE IR ANTES del área scrollable) ──
        fbot = tk.Frame(wizard, bg="#1B3A5C")
        fbot.pack(fill="x", side="bottom")

        resultado = {"ok": False}
        v_tipo = tk.StringVar(value="BUQUE")
        v_op   = tk.StringVar(value="importacion")

        def _confirmar():
            resultado["ok"] = True
            resultado["tipo_medio"] = v_tipo.get()
            resultado["operacion"]  = v_op.get()
            wizard.destroy()

        tk.Button(fbot, text="  COMENZAR MEDICION  >>", bg="#27AE60", fg="white",
                  font=("Arial", 10, "bold"), command=_confirmar,
                  cursor="hand2").pack(side="right", padx=16, pady=10, ipadx=8, ipady=4)
        tk.Button(fbot, text="Cancelar", bg="#5D6D7E", fg="white",
                  font=("Arial", 9), command=wizard.destroy
                  ).pack(side="right", padx=4, pady=10)

        # ── Área scrollable central ───────────────────────────────────────────
        cv_scroll = tk.Canvas(wizard, bg="#F8F9FA", highlightthickness=0)
        vsb = ttk.Scrollbar(wizard, orient="vertical", command=cv_scroll.yview)
        cv_scroll.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        cv_scroll.pack(side="left", fill="both", expand=True)

        fmain = tk.Frame(cv_scroll, bg="#F8F9FA")
        cv_win = cv_scroll.create_window((0, 0), window=fmain, anchor="nw")

        def _on_frame_cfg(e):
            cv_scroll.configure(scrollregion=cv_scroll.bbox("all"))
        def _on_canvas_cfg(e):
            cv_scroll.itemconfig(cv_win, width=e.width)
        fmain.bind("<Configure>", _on_frame_cfg)
        cv_scroll.bind("<Configure>", _on_canvas_cfg)
        # Scroll con rueda del mouse
        def _scroll(e):
            cv_scroll.yview_scroll(int(-1*(e.delta/120)) if e.delta else (1 if e.num==5 else -1), "units")
        cv_scroll.bind("<MouseWheel>", _scroll)
        cv_scroll.bind("<Button-4>", _scroll)
        cv_scroll.bind("<Button-5>", _scroll)

        # ── Contenido: tipos de medio ─────────────────────────────────────────
        tk.Label(fmain, text="QUE VAS A MEDIR?", font=("Arial", 10, "bold"),
                 bg="#F8F9FA", fg="#1B3A5C").pack(anchor="w", padx=20, pady=(14,4))

        CATEGORIAS = [
            ("MARITIMO / FLUVIAL",   "#1B3A5C", ["BUQUE", "BARCAZA", "BUQUE GASERO/GLP", "BUQUE QUIMIQUERO", "BUQUE METANERO/GNL", "DRAFT SURVEY"]),
            ("TIERRA / PLANTA",      "#1D6A39", ["TANQUE FIJO", "TANQUE FLOTANTE", "ESFERA DE GAS"]),
            ("TRANSPORTE TERRESTRE", "#784212", ["CAMION CISTERNA", "CAMION GAS/GLP"]),
            ("DUCTOS / CANERIAS",    "#5D4037", ["OLEODUCTO", "POLIDUCTO", "GASODUCTO"]),
            ("ENERGIA ELECTRICA",    "#6A1B9A", ["MEDICION ELECTRICA"]),
        ]
        ICONS = {
            "BUQUE": "[B]", "BARCAZA": "[Ba]", "BUQUE GASERO/GLP": "[G]",
            "BUQUE QUIMIQUERO": "[Q]", "BUQUE METANERO/GNL": "[M]",
            "TANQUE FIJO": "[TF]", "TANQUE FLOTANTE": "[TFl]", "ESFERA DE GAS": "[E]",
            "CAMION CISTERNA": "[CC]", "CAMION GAS/GLP": "[CG]",
            "OLEODUCTO": "[O]", "POLIDUCTO": "[P]", "GASODUCTO": "[GD]",
            "MEDICION ELECTRICA": "[kWh]", "DRAFT SURVEY": "[DS]",
        }

        # Guardar todos los radiobuttons para poder resaltarlos entre categorías
        all_rbs = []

        def _select(rb_clicked, val, all_rbs_ref):
            v_tipo.set(val)
            for rb_item in all_rbs_ref:
                rb_item.config(relief="flat", bg="#F8F9FA")
            rb_clicked.config(relief="raised", bg="#D6EAF8")

        f_tipos = tk.Frame(fmain, bg="#F8F9FA")
        f_tipos.pack(fill="x", padx=16, pady=2)

        for cat_name, cat_color, cat_tipos in CATEGORIAS:
            fc = tk.LabelFrame(f_tipos, text=f"  {cat_name}  ", bg="#F8F9FA",
                               font=("Arial", 8, "bold"), fg=cat_color, relief="groove", bd=2)
            fc.pack(fill="x", padx=4, pady=3)
            fr = tk.Frame(fc, bg="#F8F9FA")
            fr.pack(fill="x", padx=6, pady=5)
            for t in cat_tipos:
                ico = ICONS.get(t, "[•]")
                rb = tk.Radiobutton(fr, text=f"{ico}  {t}", variable=v_tipo, value=t,
                                    bg="#F8F9FA", font=("Arial", 9),
                                    activebackground="#E8F0FE", selectcolor="#D6EAF8",
                                    indicatoron=0, relief="flat", padx=8, pady=5,
                                    cursor="hand2")
                rb.pack(side="left", padx=3, pady=2)
                all_rbs.append(rb)
                rb.config(command=lambda r=rb, val=t: _select(r, val, all_rbs))

        # Resaltar el default
        if all_rbs:
            all_rbs[0].config(relief="raised", bg="#D6EAF8")

        # ── Tipo de operación ─────────────────────────────────────────────────
        tk.Frame(fmain, bg="#BDC3C7", height=1).pack(fill="x", padx=16, pady=8)
        tk.Label(fmain, text="TIPO DE OPERACION", font=("Arial", 10, "bold"),
                 bg="#F8F9FA", fg="#1B3A5C").pack(anchor="w", padx=20, pady=(4,4))

        f_ops = tk.Frame(fmain, bg="#F8F9FA")
        f_ops.pack(fill="x", padx=20, pady=4)
        ops = [
            ("importacion",     "[+] IMPORTACION  (Art. 954 C.A.)"),
            ("exportacion",     "[-] EXPORTACION  (Art. 959 C.A.)"),
            ("remo_descarga",   "[R] REMO / CABOTAJE - Descarga"),
            ("trafico_interno", "[I] TRAFICO INTERNO / Control Planta"),
        ]
        for val, txt in ops:
            tk.Radiobutton(f_ops, text=txt, variable=v_op, value=val,
                           bg="#F8F9FA", font=("Arial", 9),
                           activebackground="#E8F5E9").pack(anchor="w", padx=10, pady=3)

        tk.Frame(fmain, bg="#F8F9FA", height=12).pack()  # espacio al pie

        wizard.wait_window()
        return resultado

    def nueva_medicion(self):
        if not messagebox.askyesno("Nueva Medición", "¿Está seguro? Se perderán todos los datos no guardados."):
            return
        # Show wizard
        resultado = self.wizard_nueva_medicion()
        if not resultado["ok"]: return

        tipo_medio  = resultado["tipo_medio"]
        operacion   = resultado["operacion"]

        for d in self.ddt_data[:]: d["main_frame"].destroy()
        self.ddt_data = []
        self.ddt_counter = 0
        keys_to_clear = list(self.vars.keys())
        for key in keys_to_clear: self.vars[key].set("")
        self.vars.clear()
        self._archivo_tipo_medio = tipo_medio   # tipo original del wizard
        self._archivo_bloqueado  = True          # bloquear categoría desde el inicio

        # Init tanks based on tipo
        self._init_tanks_for_tipo(tipo_medio)
        self.funcionarios_data = []
        self.func_counter = 0

        for w in self.tab_caratula.winfo_children(): w.destroy()
        self.combos_ddt = []
        self.construir_caratula()

        # Pre-set tipo_medio and operacion
        self.get_var("car_tipo_medio").set(tipo_medio)
        self.get_var("car_tipo_nave").set(tipo_medio)
        if operacion in ("importacion", "exportacion"):
            self.get_var("car_operacion_default").set(operacion)

        self.agregar_ddt_row(def_prod="GASOIL")
        self.rebuild_all_tabs()

    def construir_caratula(self):
        canvas = tk.Canvas(self.tab_caratula)
        sb = ttk.Scrollbar(self.tab_caratula, orient="vertical", command=canvas.yview)
        sf = ttk.Frame(canvas)
        sf.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=sf, anchor="nw")
        
        def on_canvas_configure(event):
             canvas.itemconfig(self.canvas_window_id, width=event.width)
        canvas.bind("<Configure>", on_canvas_configure)
        self.canvas_window_id = canvas.create_window((0, 0), window=sf, anchor="nw")

        canvas.configure(yscrollcommand=sb.set)
        canvas.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        
        def _on_mousewheel(event): canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        def _on_scroll_up(event): canvas.yview_scroll(-3, "units")
        def _on_scroll_down(event): canvas.yview_scroll(3, "units")
        def _bind_mw(e):
            canvas.bind_all("<MouseWheel>", _on_mousewheel)
            canvas.bind_all("<Button-4>", _on_scroll_up)
            canvas.bind_all("<Button-5>", _on_scroll_down)
        def _unbind_mw(e):
            canvas.unbind_all("<MouseWheel>")
            canvas.unbind_all("<Button-4>")
            canvas.unbind_all("<Button-5>")
        canvas.bind("<Enter>", _bind_mw)
        canvas.bind("<Leave>", _unbind_mw)

        f_btns = ttk.Frame(sf)
        f_btns.pack(fill="x", padx=20, pady=10)
        def _lbl_gestor():
            tm = self.get_tipo_medio()
            if "CAMION" in tm: return "GESTIONAR COMPARTIMENTOS"
            if "TANQUE" in tm: return "GESTIONAR TANQUES"
            return "GESTIONAR / EDITAR TANQUES"
        btn_tanks = tk.Button(f_btns, text=_lbl_gestor(), bg="#2196F3", fg="white", font=("Arial", 8, "bold"), command=self.abrir_gestor_tanques)
        btn_tanks.pack(side="left")
        self.get_var("car_tipo_medio").trace_add("write",
            lambda *a: btn_tanks.config(text=_lbl_gestor()))
        MARITIMOS_CON_DRAFT = ("BUQUE","BARCAZA","BUQUE GASERO/GLP","BUQUE QUIMIQUERO","BUQUE METANERO/GNL")
        btn_draft = tk.Button(f_btns, text="DRAFT SURVEY", bg="#5D4037", fg="white",
                              font=("Arial", 8, "bold"), command=self.abrir_draft_survey)
        def _update_draft_btn(*a):
            tm = self.get_tipo_medio()
            if tm in MARITIMOS_CON_DRAFT:
                btn_draft.pack(side="left", padx=6)
            else:
                btn_draft.pack_forget()
        self.get_var("car_tipo_medio").trace_add("write", _update_draft_btn)
        _update_draft_btn()  # aplicar al arranque

        frame_top = ttk.LabelFrame(sf, text="Datos Generales")
        frame_top.pack(padx=20, pady=10, fill="x")
        # Columnas impares (entries) se estiran con la ventana
        for _ci in [1, 3]: frame_top.columnconfigure(_ci, weight=1)

        # --- SELECCION TIPO NAVE ---
        _is_locked = getattr(self, "_archivo_bloqueado", False)
        _lock_icon = "[BLQ] " if _is_locked else ""
        ttk.Label(frame_top, text=f"{_lock_icon}Tipo de Medición:",
                  font=("Arial", min(8, max(7, self.ui_font_size)), "bold"),
                  foreground="#8B0000" if _is_locked else "").grid(row=0, column=0, sticky="e", padx=5, pady=6)
        _tm_var = self.get_var("car_tipo_medio", self.get_var("car_tipo_nave","BUQUE").get() or "BUQUE")
        # Restringir valores del combo a la categoría actual si está bloqueado
        _allowed_tipos = self.CATEGORIA_TIPOS.get(
            getattr(self, "_archivo_tipo_medio", ""),
            self.TIPO_MEDIOS
        ) if _is_locked else self.TIPO_MEDIOS
        cb_tipo = ttk.Combobox(frame_top, textvariable=_tm_var, state="readonly",
                               values=_allowed_tipos, width=27,
                               font=("Arial", max(10, self.ui_font_size+1)))
        cb_tipo.grid(row=0, column=1, sticky="w", padx=5)
        # Indicador visual si está bloqueado
        if _is_locked:
            _locked_cat = getattr(self, "_archivo_tipo_medio", "")
            _cat_label  = {"BUQUE":"MARÍTIMO","BARCAZA":"MARÍTIMO","BUQUE GASERO/GLP":"MARÍTIMO",
                           "BUQUE QUIMIQUERO":"MARÍTIMO","BUQUE METANERO/GNL":"MARÍTIMO",
                           "TANQUE FIJO":"TERRESTRE","TANQUE FLOTANTE":"TERRESTRE","ESFERA DE GAS":"TERRESTRE",
                           "CAMION CISTERNA":"AUTOMOTOR","CAMION GAS/GLP":"AUTOMOTOR",
                           "OLEODUCTO":"DUCTO","POLIDUCTO":"DUCTO","GASODUCTO":"DUCTO",
                           "ELECTRICO":"ELÉCTRICO"}.get(_locked_cat, _locked_cat)
            # Color del banner según categoría
            _lock_colors = {
                "MARÍTIMO":   ("#0A2A4A", "#5DADE2"),
                "TERRESTRE":  ("#0A3A0A", "#58D68D"),
                "AUTOMOTOR":  ("#3A2010", "#E59866"),
                "DUCTO":      ("#1A1A3A", "#85C1E9"),
                "ELÉCTRICO":  ("#2A2A00", "#F9E79F"),
            }
            _lbg, _lfg = _lock_colors.get(_cat_label, ("#8B0000", "#FFFFFF"))
            lbl_lock = tk.Label(frame_top,
                                text=f"  🔒 {_cat_label}  |  {_locked_cat}  —  No se puede cambiar categoría  ",
                                font=("Arial", max(7, self.ui_font_size-1), "bold"),
                                fg=_lfg, bg=_lbg, padx=6, pady=3, relief="flat")
            lbl_lock.grid(row=0, column=2, columnspan=2, sticky="w", padx=8)
        # sync car_tipo_nave for legacy compatibility
        _prev_tipo = [self.get_tipo_medio()]
        def _sync_tipo(*a):
            new_tm = _tm_var.get()
            # Bloqueo de categoría: si está bloqueado, solo se permite cambiar
            # entre tipos de la misma categoría (ej: BUQUE ↔ BARCAZA, no BUQUE → TANQUE)
            if getattr(self, "_archivo_bloqueado", False) and new_tm:
                locked = getattr(self, "_archivo_tipo_medio", "")
                allowed = self.CATEGORIA_TIPOS.get(locked, [locked])
                if new_tm not in allowed:
                    _tm_var.set(_prev_tipo[0])  # revertir
                    import tkinter.messagebox as mb
                    mb.showwarning("Tipo bloqueado",
                        f"Esta medición es de tipo «{locked}».\n"
                        f"No se puede cambiar a «{new_tm}».\n\n"
                        f"Para otro tipo, use: Archivo → Nueva Medición")
                    return
            self.get_var("car_tipo_nave").set(new_tm)
            _toggle_fields()
            if new_tm and new_tm != _prev_tipo[0]:
                _prev_tipo[0] = new_tm
                # Reinicializar tanques y tabs cuando el tipo cambia
                self._init_tanks_for_tipo(new_tm)
                self.rebuild_all_tabs()
        _tm_var.trace_add("write", _sync_tipo)


        _car_fnt = max(10, self.ui_font_size + 1)  # font mínimo 10 en Datos Generales
        _fnt_b = ("Arial", min(8, _car_fnt), "bold")
        _fnt_n = ("Arial", _car_fnt)

        # ── Fila 1: Nombre unidad + IMO/ID ──────────────────────────────────
        lbl_nombre = ttk.Label(frame_top, text="Nombre Buque/Barcaza:", font=_fnt_b)
        lbl_nombre.grid(row=1, column=0, sticky="e", padx=5, pady=6)
        entry_nombre = tk.Entry(frame_top, textvariable=self.get_var("car_buque"), font=_fnt_n)
        entry_nombre.grid(row=1, column=1, sticky="ew", padx=5, ipady=2)
        lbl_imo = ttk.Label(frame_top, text="N° IMO:", font=_fnt_b)
        lbl_imo.grid(row=1, column=2, sticky="e", padx=5, pady=6)
        entry_imo = tk.Entry(frame_top, textvariable=self.get_var("car_imo"), font=_fnt_n)
        entry_imo.grid(row=1, column=3, sticky="ew", padx=5, ipady=2)

        # ── Filas fijas comunes ──────────────────────────────────────────────
        campos_fijos = [
            ("Despachante de Aduana:", "car_despachante", 2, 0), ("CUIT Despachante:", "car_cuit_desp", 2, 2),
            ("Importador / Exportador:", "car_impexp", 3, 0), ("CUIT Imp/Exp:", "car_cuit_impexp", 3, 2),
            ("Agencia Marítima (ATA):", "car_ata", 4, 0), ("CUIT ATA:", "car_cuit_ata", 4, 2),
            ("N° MANI:", "car_mani", 5, 0), ("N° Viaje:", "car_conocimientos", 5, 2),
        ]
        cuit_fields = ["car_cuit_desp", "car_cuit_impexp", "car_cuit_ata"]
        lbl_ata = None; lbl_cuit_ata = None; e_ata = None; e_cuit_ata = None
        lbl_mani = None; e_mani = None
        for item in campos_fijos:
            lbl_txt, key, r, c_col = item
            lbl_w = ttk.Label(frame_top, text=lbl_txt, font=_fnt_b)
            lbl_w.grid(row=r, column=c_col, sticky="e", padx=5, pady=6)
            e_w = tk.Entry(frame_top, textvariable=self.get_var(key), font=_fnt_n)
            e_w.grid(row=r, column=c_col+1, sticky="ew", padx=5, ipady=2)
            if key in cuit_fields:
                var_ref = self.get_var(key)
                e_w.bind("<FocusOut>", lambda ev, wdg=e_w, v=var_ref: self.on_cuit_validate(wdg, v))
            # Save refs for dynamic hiding
            if key == "car_ata": lbl_ata = lbl_w; e_ata = e_w
            if key == "car_cuit_ata": lbl_cuit_ata = lbl_w; e_cuit_ata = e_w
            if key == "car_mani": lbl_mani = lbl_w; e_mani = e_w

        # ── Extra fields: tanques de tierra (techo flotante offset + diámetro) ──
        lbl_radio = ttk.Label(frame_top, text="Radio Interno (m):", font=_fnt_b)
        lbl_radio.grid(row=1, column=2, sticky="e", padx=5, pady=6)
        e_radio = tk.Entry(frame_top, textvariable=self.get_var("car_radio_m",""), font=_fnt_n, width=10)
        e_radio.grid(row=1, column=3, sticky="w", padx=5, ipady=2)

        lbl_tf_off = ttk.Label(frame_top, text="Offset Techo Flotante (mm):", font=_fnt_b)
        lbl_tf_off.grid(row=1, column=2, sticky="e", padx=5, pady=6)  # same pos as radio; toggle later
        e_tf_off = tk.Entry(frame_top, textvariable=self.get_var("car_tf_offset","0"), font=_fnt_n, width=10)
        e_tf_off.grid(row=1, column=3, sticky="w", padx=5, ipady=2)

        # ── Extra fields: camion ─────────────────────────────────────────────
        lbl_patente = ttk.Label(frame_top, text="Dominio / Patente:", font=_fnt_b)
        lbl_patente.grid(row=1, column=0, sticky="e", padx=5, pady=6)
        e_patente = tk.Entry(frame_top, textvariable=self.get_var("car_patente",""), font=_fnt_n)
        e_patente.grid(row=1, column=1, sticky="ew", padx=5, ipady=2)
        lbl_precinto = ttk.Label(frame_top, text="N° Precinto:", font=_fnt_b)
        lbl_precinto.grid(row=1, column=2, sticky="e", padx=5, pady=6)
        e_precinto = tk.Entry(frame_top, textvariable=self.get_var("car_precinto",""), font=_fnt_n, width=10)
        e_precinto.grid(row=1, column=3, sticky="w", padx=5, ipady=2)

        lbl_radio_cam = ttk.Label(frame_top, text="Radio Cisterna (m):", font=_fnt_b)
        lbl_radio_cam.grid(row=4, column=2, sticky="e", padx=5, pady=6)
        e_radio_cam = tk.Entry(frame_top, textvariable=self.get_var("car_radio_camion",""), font=_fnt_n, width=10)
        e_radio_cam.grid(row=4, column=3, sticky="w", padx=5, ipady=2)
        lbl_largo_cam = ttk.Label(frame_top, text="Largo Cisterna (m):", font=_fnt_b)
        lbl_largo_cam.grid(row=4, column=0, sticky="e", padx=5, pady=6)
        e_largo_cam = tk.Entry(frame_top, textvariable=self.get_var("car_largo_camion",""), font=_fnt_n, width=10)
        e_largo_cam.grid(row=4, column=1, sticky="w", padx=5, ipady=2)

        # ── Extra fields: gasero ─────────────────────────────────────────────
        lbl_presion = ttk.Label(frame_top, text="Presión operativa (kPa):", font=_fnt_b)
        lbl_presion.grid(row=4, column=0, sticky="e", padx=5, pady=6)
        e_presion = tk.Entry(frame_top, textvariable=self.get_var("car_presion_op",""), font=_fnt_n, width=10)
        e_presion.grid(row=4, column=1, sticky="w", padx=5, ipady=2)
        lbl_temp_op = ttk.Label(frame_top, text="Temp. operativa (°C):", font=_fnt_b)
        lbl_temp_op.grid(row=4, column=2, sticky="e", padx=5, pady=6)
        e_temp_op = tk.Entry(frame_top, textvariable=self.get_var("car_temp_op",""), font=_fnt_n, width=10)
        e_temp_op.grid(row=4, column=3, sticky="w", padx=5, ipady=2)

        def _toggle_fields():
            tm = self.get_tipo_medio()
            mar = self.es_maritimo()
            gas = self.es_gasero()
            tie = self.es_tierra()
            cam = self.es_camion()
            flot = "FLOTANTE" in tm

            # Fila 1 label switch
            for w in (lbl_nombre, entry_nombre, lbl_imo, entry_imo,
                      lbl_radio, e_radio, lbl_tf_off, e_tf_off,
                      lbl_patente, e_patente, lbl_precinto, e_precinto):
                try: w.grid_remove()
                except: pass

            if mar:
                lbl_nombre.config(text="Nombre Buque/Barcaza:" if "BARCAZA" not in tm else "Nombre Barcaza:")
                lbl_nombre.grid(); entry_nombre.grid()
                lbl_imo.grid(); entry_imo.grid()
            elif tie:
                lbl_nombre.config(text="Nombre del Tanque:")
                lbl_nombre.grid(); entry_nombre.grid()
                lbl_radio.grid(); e_radio.grid()
                if flot: lbl_tf_off.grid(); e_tf_off.grid()
            elif cam:
                lbl_patente.grid(); e_patente.grid()
                lbl_precinto.grid(); e_precinto.grid()

            elif self.es_esfera():
                lbl_nombre.config(text="Nombre Instalacion:")
                lbl_nombre.grid(); entry_nombre.grid()
                lbl_imo.config(text="N Expediente / Contrato:")
                lbl_imo.grid(); entry_imo.grid()

            # ATA: solo marino
            for w in (lbl_ata, e_ata, lbl_cuit_ata, e_cuit_ata):
                if w:
                    try:
                        if mar: w.grid()
                        else: w.grid_remove()
                    except: pass

            # MANI label
            if lbl_mani:
                lbl_mani.config(text="N° MANI:" if mar else "N° Expediente:")

            # Fila 4 extras
            for w in (lbl_radio_cam, e_radio_cam, lbl_largo_cam, e_largo_cam,
                      lbl_presion, e_presion, lbl_temp_op, e_temp_op):
                try: w.grid_remove()
                except: pass
            duc = self.es_ducto()
            cam_gas = self.es_camion_gas()
            elec = self.es_electrico()

            if cam or cam_gas:
                lbl_radio_cam.grid(); e_radio_cam.grid()
                lbl_largo_cam.grid(); e_largo_cam.grid()
            _esf_car = self.es_esfera()
            if gas or cam_gas or _esf_car:
                if _esf_car:
                    lbl_presion.config(text="Presion operativa (kPa):")
                    lbl_temp_op.config(text="Temp. operativa (C):")
                lbl_presion.grid(); e_presion.grid()
                lbl_temp_op.grid(); e_temp_op.grid()
            if duc:
                lbl_presion.config(text="Presion linea (kPa):"); lbl_presion.grid(); e_presion.grid()
                lbl_temp_op.config(text="Temp. linea (C):"); lbl_temp_op.grid(); e_temp_op.grid()
                lbl_nombre.config(text="Nombre del Ducto:"); lbl_nombre.grid(); entry_nombre.grid()
                lbl_imo.config(text="Diametro (pulg):"); lbl_imo.grid(); entry_imo.grid()
            if elec:
                lbl_nombre.config(text="Instalacion / Punto de Medicion:"); lbl_nombre.grid(); entry_nombre.grid()
                lbl_imo.config(text="N Contrato / Suministro:"); lbl_imo.grid(); entry_imo.grid()

        _toggle_fields()

        # --- SELECTOR ADUANA (códigos oficiales ARCA) ---
        ADUANAS_ARG = [
            "001 - BS.AS. (CAPITAL)",       "003 - BAHIA BLANCA",
            "004 - BARILOCHE",              "008 - CAMPANA",
            "010 - BARRANQUERAS",           "012 - CLORINDA",
            "013 - COLON",                  "014 - COMODORO RIVADAVIA",
            "015 - CONCEPCION DEL URUGUAY", "016 - CONCORDIA",
            "017 - CORDOBA",                "018 - CORRIENTES",
            "019 - PUERTO DESEADO",         "020 - DIAMANTE",
            "023 - ESQUEL",                 "024 - FORMOSA",
            "025 - GOYA",                   "026 - GUALEGUAYCHU",
            "029 - IGUAZU",                 "031 - JUJUY",
            "033 - LA PLATA",               "034 - LA QUIACA",
            "037 - MAR DEL PLATA",          "038 - MENDOZA",
            "040 - NECOCHEA",               "041 - PARANA",
            "042 - PASO DE LOS LIBRES",     "045 - POCITOS",
            "046 - POSADAS",                "047 - PUERTO MADRYN",
            "048 - RIO GALLEGOS",           "049 - RIO GRANDE",
            "052 - ROSARIO",                "053 - SALTA",
            "054 - SAN JAVIER",             "055 - SAN JUAN",
            "057 - SAN LORENZO",            "058 - S. MARTIN DE LOS ANDES",
            "059 - SAN NICOLAS",            "060 - SAN PEDRO",
            "061 - SANTA CRUZ",             "062 - SANTA FE",
            "066 - TINOGASTA",              "067 - USHUAIA",
            "069 - VILLA CONSTITUCION",     "073 - EZEIZA",
            "074 - TUCUMAN",                "075 - NEUQUEN",
            "076 - ORAN",                   "078 - SAN RAFAEL",
            "079 - LA RIOJA",               "080 - SAN ANTONIO OESTE",
            "082 - BERNARDO DE YRIGOYEN",   "083 - SAN LUIS",
            "084 - SANTO TOME",             "085 - VILLA REGINA",
            "086 - OBERA",                  "087 - CALETA OLIVIA",
            "088 - GENERAL DEHEZA",         "089 - SANTIAGO DEL ESTERO",
            "090 - GENERAL PICO",           "091 - BS.AS. NORTE",
            "092 - BS.AS. SUR",             "093 - RAFAELA",
            "099 - MULTIADUANA",
            "258 - Z.F GENERAL PICO",       "266 - Z.F CORONEL ROSALES",
            "267 - Z.F CONCEP.DEL.URUG.",   "268 - Z.F. V. CONSTITUCION",
            "269 - Z.F. PUERTO GALVAN",
        ]
        ttk.Label(frame_top, text="Aduana:", font=("Arial", min(8, max(7, self.ui_font_size)), "bold")).grid(row=6, column=0, sticky="e", padx=5, pady=6)
        self.get_var("car_lugar", "067 - USHUAIA")
        cb_aduana = ttk.Combobox(frame_top, textvariable=self.get_var("car_lugar"),
                                 values=ADUANAS_ARG, state="normal", width=28, font=("Arial", max(10, self.ui_font_size+1)))
        cb_aduana.grid(row=6, column=1, sticky="w", padx=5)
        # T/C global del documento (ARS por 1 U$S)
        ttk.Label(frame_top, text="T/C U$S:", font=("Arial", min(8, max(7, self.ui_font_size)), "bold"),
                  foreground="#6A1B9A").grid(row=6, column=2, sticky="e", padx=5, pady=6)
        ttk.Entry(frame_top, textvariable=self.get_var("car_tipo_cambio", ""),
                  width=12, font=("Arial", max(10, self.ui_font_size+1))).grid(row=6, column=3, sticky="w", padx=5)

        # --- LUGAR OPERATIVO: selector desde DB, filtrado por aduana ---
        f_lop = ttk.Frame(frame_top)
        f_lop.grid(row=7, column=0, columnspan=4, sticky="ew", padx=5, pady=3)
        ttk.Label(f_lop, text="Lugar Operativo (LOT):",
                  font=("Arial", min(8, max(7, self.ui_font_size)), "bold")).pack(side="left")
        _lop_display = tk.StringVar()
        _lop_fnt = ("Arial", max(10, self.ui_font_size + 1))
        cb_lop = ttk.Combobox(f_lop, textvariable=_lop_display, width=55,
                               font=_lop_fnt, state="normal")
        cb_lop.pack(side="left", padx=(6, 0))
        tk.Button(f_lop, text="Gestionar LOTs", bg="#1A5276", fg="white",
                  font=("Arial", 8), cursor="hand2",
                  command=self.abrir_gestor_aduanas_lop).pack(side="left", padx=6)
        lbl_lop_hint = tk.Label(f_lop, text="", font=("Arial", 7), fg="gray")
        lbl_lop_hint.pack(side="left", padx=2)

        def _lop_refresh(*_):
            adu_raw = self.get_var("car_lugar").get().strip()
            adu_cod = adu_raw.split(" - ")[0].strip() if " - " in adu_raw else adu_raw[:3]
            lots = db_get_lugares_operativos(adu_cod) if adu_cod else []
            if not lots:
                lots = db_get_lugares_operativos()
            opciones = [f"{r['codigo']} — {r['descripcion']}" for r in lots]
            cb_lop["values"] = opciones
            n = len(opciones)
            lbl_lop_hint.config(text=f"{n} LOT(s)" if n else "(sin LOTs — use Gestionar LOTs)")

        def _lop_sync(*_):
            val = _lop_display.get().strip()
            if " — " in val:
                cod, desc = val.split(" — ", 1)
            else:
                parts = val.split(None, 1)
                cod  = parts[0] if parts else val
                desc = parts[1] if len(parts) > 1 else ""
            self.get_var("car_lop_codigo").set(cod.strip())
            self.get_var("car_lop_desc").set(desc.strip())

        def _lop_preload():
            cod  = self.get_var("car_lop_codigo", "").get().strip()
            desc = self.get_var("car_lop_desc",   "").get().strip()
            if cod or desc:
                sep = " — " if desc else ""
                _lop_display.set(f"{cod}{sep}{desc}")

        cb_lop.bind("<<ComboboxSelected>>", _lop_sync)
        cb_lop.bind("<FocusOut>",           _lop_sync)
        self.get_var("car_lugar").trace_add("write", _lop_refresh)
        _lop_refresh()
        _lop_preload()

        # === SELECTOR NORMA VCF ===
        frame_norma = ttk.LabelFrame(sf, text="Norma de Cálculo VCF (Tablas 54)")
        frame_norma.pack(padx=20, pady=(4, 10), fill="x")
        _fn = ("Arial", max(9, self.ui_font_size), "bold")
        _fn2 = ("Arial", max(8, self.ui_font_size - 1))

        tk.Radiobutton(
            frame_norma, text="ASTM D1250-1980  (tablas impresas — 4 zonas en 54B)",
            variable=self.norma_astm, value="1980",
            font=_fn, fg="#1A5276", activeforeground="#1A5276",
            command=self._on_norma_changed
        ).grid(row=0, column=0, sticky="w", padx=15, pady=(8, 2))
        tk.Label(frame_norma,
                 text="  JP-1 rho=798 t=23°C → 0.99252 | Diesel rho=850 t=25°C → 0.99159  (4 zonas de densidad)",
                 font=_fn2, fg="#555").grid(row=1, column=0, sticky="w", padx=30, pady=(0, 6))

        tk.Radiobutton(
            frame_norma, text="API MPMS 11.1 / ASTM D1250-2004  (norma digital — fórmula única en 54B)",
            variable=self.norma_astm, value="2004",
            font=_fn, fg="#4A235A", activeforeground="#4A235A",
            command=self._on_norma_changed
        ).grid(row=0, column=1, sticky="w", padx=15, pady=(8, 2))
        tk.Label(frame_norma,
                 text="  JP-1 rho=798 t=23°C → 0.99096 | Diesel rho=850 t=25°C → 0.98996  (K0=346.4228 K1=0.4033)",
                 font=_fn2, fg="#555").grid(row=1, column=1, sticky="w", padx=30, pady=(0, 6))

        # === FUNCIONARIOS INTERVINIENTES ===
        self.frame_func_master = ttk.LabelFrame(sf, text="Funcionarios Intervinientes")
        self.frame_func_master.pack(padx=20, pady=10, fill="x")
        _fb = ttk.Frame(self.frame_func_master)
        _fb.pack(fill="x", padx=10, pady=5)
        tk.Button(_fb, text="+ Fila en blanco", bg="#6A1B9A", fg="white",
                  font=("Arial", 8, "bold"), command=self.agregar_funcionario_row).pack(side="left")
        tk.Button(_fb, text="Buscar en DB", bg="#27AE60", fg="white",
                  font=("Arial", 8, "bold"), command=self.buscar_y_agregar_funcionario_db).pack(side="left", padx=6)
        tk.Button(_fb, text="Administrar DB", bg="#37474F", fg="white",
                  font=("Arial", 8, "bold"), command=self.abrir_gestor_funcionarios_db).pack(side="left", padx=4)
        tk.Button(_fb, text="Gestionar Funciones", bg="#0D47A1", fg="white",
                  font=("Arial", 8, "bold"), command=self.abrir_gestor_funciones).pack(side="left", padx=4)
        self.func_stack = ttk.Frame(self.frame_func_master)
        self.func_stack.pack(fill="x", padx=5)
        # Cabecera con grid
        f_func_hdr = ttk.Frame(self.func_stack)
        f_func_hdr.pack(fill="x")
        for ci, (col_txt, col_w) in enumerate([("CUIL", 14), ("Legajo", 10), ("Apellido", 22), ("Nombre", 22), ("Función", 26), ("", 3)]):
            tk.Label(f_func_hdr, text=col_txt, font=("Arial", 8, "bold"), fg="gray", width=col_w, anchor="w").grid(row=0, column=ci, padx=2)

        self.frame_ddt_master = ttk.LabelFrame(sf, text="Documentos y Salidas")
        self.frame_ddt_master.pack(padx=20, pady=10, fill="both", expand=True)
        tk.Button(self.frame_ddt_master, text="+ AGREGAR DOCUMENTO", bg="#2E7D32", fg="white", font=("Arial", 8, "bold"), command=self.popup_tipo_documento).pack(pady=10, anchor="w", padx=10)
        self.ddt_stack = ttk.Frame(self.frame_ddt_master)
        self.ddt_stack.pack(fill="both", expand=True, padx=5)
        
    def abrir_gestor_tanques(self):
        try:
            toplevel = tk.Toplevel(self.root)
            toplevel.title("Gestión de Tanques y Carbonera")
            # Maximizar ventana
            try:
                toplevel.state("zoomed")       # Windows
            except:
                try:
                    toplevel.attributes("-zoomed", True)  # Linux
                except:
                    sw2, sh2 = toplevel.winfo_screenwidth(), toplevel.winfo_screenheight()
                    toplevel.geometry(f"{sw2}x{sh2}+0+0")
            
            # --- BOTONERA INFERIOR FIJA (fuera del PanedWindow) ---
            f_bot = tk.Frame(toplevel, bg="#2C3E50", bd=0, height=60)
            f_bot.pack(side="bottom", fill="x")
            f_bot.pack_propagate(False)  # Mantener altura fija
            
            # Frame interior centrado
            f_bot_inner = tk.Frame(f_bot, bg="#2C3E50")
            f_bot_inner.pack(expand=True)

            # PanedWindow para dividir lista (izq) y dibujo (der)
            pw = ttk.PanedWindow(toplevel, orient=tk.HORIZONTAL)
            pw.pack(fill="both", expand=True)

            left_frame = ttk.Frame(pw)
            right_frame = ttk.Frame(pw)
            pw.add(left_frame, weight=2)
            pw.add(right_frame, weight=3)

            # --- CANVAS SCROLL ---
            canvas_list = tk.Canvas(left_frame, bg="white")
            sb = ttk.Scrollbar(left_frame, orient="vertical", command=canvas_list.yview)
            scroll_f = ttk.Frame(canvas_list)
            
            scroll_f.bind("<Configure>", lambda e: canvas_list.configure(scrollregion=canvas_list.bbox("all")))
            window_id = canvas_list.create_window((0, 0), window=scroll_f, anchor="nw")

            def on_canvas_configure(event):
                canvas_list.itemconfig(window_id, width=event.width)
            canvas_list.bind("<Configure>", on_canvas_configure)

            canvas_list.configure(yscrollcommand=sb.set)
            sb.pack(side="right", fill="y")
            canvas_list.pack(fill="both", expand=True)
            # Mousewheel scoped to this canvas only
            def _cl_on_mw(event): canvas_list.yview_scroll(int(-1*(event.delta/120)), "units")
            def _cl_up(event): canvas_list.yview_scroll(-3, "units")
            def _cl_down(event): canvas_list.yview_scroll(3, "units")
            def _cl_bind(e):
                canvas_list.bind_all("<MouseWheel>", _cl_on_mw)
                canvas_list.bind_all("<Button-4>", _cl_up)
                canvas_list.bind_all("<Button-5>", _cl_down)
            def _cl_unbind(e):
                canvas_list.unbind_all("<MouseWheel>")
                canvas_list.unbind_all("<Button-4>")
                canvas_list.unbind_all("<Button-5>")
            canvas_list.bind("<Enter>", _cl_bind)
            canvas_list.bind("<Leave>", _cl_unbind)

            def _parse_side(name):
                """Extrae el lado del nombre"""
                if "BABOR" in name.upper(): return "BABOR"
                if "ESTRIBOR" in name.upper(): return "ESTRIBOR"
                return "AMBOS"

            def _strip_side(name):
                """Quita BABOR/ESTRIBOR del nombre"""
                n = name.strip()
                for s in [" BABOR", " ESTRIBOR", " babor", " estribor"]:
                    if n.endswith(s): n = n[:-len(s)].strip()
                return n

            # Datos temporales: (nombre_sin_lado, lado)
            self.temp_tanks = []
            for t in self.lista_tanques:
                self.temp_tanks.append({
                    "name": tk.StringVar(value=_strip_side(t)),
                    "side": tk.StringVar(value=_parse_side(t))
                })

            self.temp_carb = []
            for t in self.lista_carbonera:
                self.temp_carb.append({
                    "name": tk.StringVar(value=_strip_side(t)),
                    "side": tk.StringVar(value=_parse_side(t))
                })

            # Guía Visual (Canvas Derecho)
            gui_canvas = tk.Canvas(right_frame, bg="#f0f0f0")
            gui_canvas.pack(fill="both", expand=True)

            def dibujar_guia():
                gui_canvas.delete("all")
                w_cv = gui_canvas.winfo_width()
                h_cv = gui_canvas.winfo_height()
                if w_cv < 100 or h_cv < 100: return

                # Construir listas temporales con nombres completos
                temp_tank_full = []
                for item in self.temp_tanks:
                    name = item["name"].get().strip()
                    side = item["side"].get()
                    if not name: continue
                    if side == "AMBOS": temp_tank_full.append(name)
                    else: temp_tank_full.append(f"{name} {side}")

                temp_carb_full = []
                for item in self.temp_carb:
                    name = item["name"].get().strip()
                    side = item["side"].get()
                    if not name: continue
                    if side == "AMBOS": temp_carb_full.append(name)
                    else: temp_carb_full.append(f"{name} {side}")

                if self.es_maritimo():
                    half_h = (h_cv - 20) / 2
                    # Buques: 2 vistas
                    self.dibujar_unidad_tk(gui_canvas, 5, 5, w_cv - 10, half_h - 5, "BABOR", None, temp_tank_full, temp_carb_full)
                    self.dibujar_unidad_tk(gui_canvas, 5, 5 + half_h, w_cv - 10, half_h - 5, "ESTRIBOR", None, temp_tank_full, temp_carb_full)
                else:
                    # Tierra/esfera/camion/ducto: canvas completo
                    self.dibujar_unidad_tk(gui_canvas, 5, 5, w_cv - 10, h_cv - 10, "AMBOS", None, temp_tank_full, temp_carb_full)

            _guia_pending = [None]
            def _dibujar_guia_debounced(event=None):
                if _guia_pending[0] is not None:
                    try: toplevel.after_cancel(_guia_pending[0])
                    except: pass
                _guia_pending[0] = toplevel.after(50, dibujar_guia)
            gui_canvas.bind("<Configure>", _dibujar_guia_debounced)

            _render_pending = [None]
            def render_all():
                if _render_pending[0] is not None:
                    try: toplevel.after_cancel(_render_pending[0])
                    except: pass
                _render_pending[0] = toplevel.after(50, _do_render)

            def _do_render():
                for w_child in scroll_f.winfo_children(): w_child.destroy()
                
                # --- TANQUES ---
                _tm_now = self.get_tipo_medio()
                _lbl_tk_sec = "TANQUES DE CARGA"
                if "CAMION" in _tm_now: _lbl_tk_sec = "COMPARTIMENTOS DE CISTERNA"
                elif "TANQUE" in _tm_now: _lbl_tk_sec = "TANQUES DE ALMACENAMIENTO"
                elif "ESFERA" in _tm_now: _lbl_tk_sec = "ESFERAS DE GAS"
                elif "DUCTO" in _tm_now: _lbl_tk_sec = "TRAMOS DE DUCTO"
                elif "ELECTRICA" in _tm_now: _lbl_tk_sec = "MEDIDORES"
                tk.Label(scroll_f, text=_lbl_tk_sec, font=("Arial", 8, "bold"), bg="#D6EAF8").pack(fill="x", pady=5)
                
                for i, item in enumerate(self.temp_tanks):
                    row = ttk.Frame(scroll_f)
                    row.pack(fill="x", pady=1, padx=2)
                    
                    tk.Button(row, text="^", width=2, command=lambda idx=i: move_item(self.temp_tanks, idx, -1)).pack(side="left")
                    tk.Button(row, text="v", width=2, command=lambda idx=i: move_item(self.temp_tanks, idx, 1)).pack(side="left")
                    
                    ttk.Entry(row, textvariable=item["name"], width=20).pack(side="left", padx=3)
                    
                    _is_mar = self.es_maritimo()
                    if _is_mar:
                        cb = ttk.Combobox(row, textvariable=item["side"], values=["BABOR", "ESTRIBOR", "AMBOS"], state="readonly", width=10)
                        cb.pack(side="left", padx=3)
                        cb.bind("<<ComboboxSelected>>", lambda e: dibujar_guia())
                    
                    tk.Button(row, text="X", bg="#E74C3C", fg="white", width=2, command=lambda idx=i: remove_item(self.temp_tanks, idx)).pack(side="right")

                ttk.Separator(scroll_f, orient="horizontal").pack(fill="x", pady=10)
                
                # --- CARBONERAS / CONSUMO (solo buques maritimos clásicos) ---
                _show_carb = self.get_tipo_medio() in ("BUQUE","BARCAZA","BUQUE GASERO/GLP","BUQUE QUIMIQUERO","BUQUE METANERO/GNL")
                if _show_carb:
                    tk.Label(scroll_f, text="CARBONERAS / CONSUMO", font=("Arial", 8, "bold"), bg="#FEF9E7", fg="#6E4B00").pack(fill="x", pady=5)
                    
                    for i, item in enumerate(self.temp_carb):
                        row = ttk.Frame(scroll_f)
                        row.pack(fill="x", pady=1, padx=2)
                        tk.Button(row, text="^", width=2, command=lambda idx=i: move_item(self.temp_carb, idx, -1)).pack(side="left")
                        tk.Button(row, text="v", width=2, command=lambda idx=i: move_item(self.temp_carb, idx, 1)).pack(side="left")
                        ttk.Entry(row, textvariable=item["name"], width=20).pack(side="left", padx=3)
                        cb = ttk.Combobox(row, textvariable=item["side"], values=["BABOR", "ESTRIBOR", "AMBOS"], state="readonly", width=10)
                        cb.pack(side="left", padx=3)
                        cb.bind("<<ComboboxSelected>>", lambda e: dibujar_guia())
                        tk.Button(row, text="X", bg="#E74C3C", fg="white", width=2, command=lambda idx=i: remove_item(self.temp_carb, idx)).pack(side="right")
                
                dibujar_guia()

            def move_item(lista, index, direction):
                new_idx = index + direction
                if 0 <= new_idx < len(lista):
                    lista[index], lista[new_idx] = lista[new_idx], lista[index]
                    render_all()

            def remove_item(lista, index):
                del lista[index]
                render_all()

            def add_t():
                _side_def = "BABOR" if self.es_maritimo() else "AMBOS"
                _name_def = "NUEVO TK" if not self.es_maritimo() else "NUEVO TK"
                self.temp_tanks.append({"name": tk.StringVar(value=_name_def), "side": tk.StringVar(value=_side_def)})
                render_all()
            def add_c():
                self.temp_carb.append({"name": tk.StringVar(value="CARBONERA"), "side": tk.StringVar(value="AMBOS")})
                render_all()

            def save():
                # Construir nombre final: nombre + lado (excepto AMBOS)
                new_tanks = []
                for item in self.temp_tanks:
                    name = item["name"].get().strip()
                    if not name: continue
                    if self.es_maritimo():
                        side = item["side"].get()
                        if side == "AMBOS": new_tanks.append(name)
                        else: new_tanks.append(f"{name} {side}")
                    else:
                        new_tanks.append(name)
                
                new_carb = []
                for item in self.temp_carb:
                    name = item["name"].get().strip()
                    if not name: continue
                    if self.es_maritimo():
                        side = item["side"].get()
                        if side == "AMBOS": new_carb.append(name)
                        else: new_carb.append(f"{name} {side}")
                    else:
                        new_carb.append(name)
                
                self.lista_tanques = new_tanks
                self.lista_carbonera = new_carb
                self.rebuild_all_tabs()
                toplevel.destroy()
                messagebox.showinfo("Listo", "Configuración de tanques guardada.")

            # --- BOTONES EN f_bot ---
            _tm_btn = self.get_tipo_medio()
            _tk_lbl = "+ Compartimento" if ("CAMION" in _tm_btn or "TANQUE" in _tm_btn) else "+ TK Carga"
            tk.Button(f_bot_inner, text=f"  {_tk_lbl}  ", bg="#2980B9", fg="white", font=("Arial", 8, "bold"), command=add_t).pack(side="left", padx=10, pady=12)
            if self.get_tipo_medio() in ("BUQUE","BARCAZA","BUQUE GASERO/GLP","BUQUE QUIMIQUERO","BUQUE METANERO/GNL"):
                tk.Button(f_bot_inner, text="  + Carbonera  ", bg="#D4AC0D", fg="white", font=("Arial", 8, "bold"), command=add_c).pack(side="left", padx=10, pady=12)
            tk.Button(f_bot_inner, text="  GUARDAR CAMBIOS", bg="#27AE60", fg="white", font=("Arial", 8, "bold"), command=save, cursor="hand2").pack(side="right", padx=20, pady=12, ipadx=10, ipady=4)
            
            render_all()

        except Exception as e:
            traceback.print_exc()
            messagebox.showerror("Error", f"Error al abrir el gestor de tanques:\n{e}")

    
    def auto_calc_densidad(self, sv_lits, sv_kgs, sv_dens):
        try:
            l = self.parse_float(sv_lits.get())
            k = self.parse_float(sv_kgs.get())
            if l > 0: sv_dens.set(f"{k/l:.5f}")
        except: pass

    # --- SUBREGÍMENES DE DECLARACIÓN DETALLADA ---

    # ── Tipos de medio disponibles ───────────────────────────────
    TIPO_MEDIOS = [
        "BUQUE", "BARCAZA",
        "BUQUE GASERO/GLP", "BUQUE QUIMIQUERO", "BUQUE METANERO/GNL",
        "TANQUE FIJO", "TANQUE FLOTANTE",
        "ESFERA DE GAS",
        "CAMION CISTERNA", "CAMION GAS/GLP",
        "OLEODUCTO", "POLIDUCTO", "GASODUCTO",
        "MEDICION ELECTRICA",
        "DRAFT SURVEY",
    ]
    # Tipos permitidos en el mismo proyecto por categoría operativa
    CATEGORIA_TIPOS = {
        "BUQUE":              ["BUQUE","BARCAZA","BUQUE GASERO/GLP","BUQUE QUIMIQUERO","BUQUE METANERO/GNL","DRAFT SURVEY"],
        "BARCAZA":            ["BUQUE","BARCAZA","BUQUE GASERO/GLP","BUQUE QUIMIQUERO","BUQUE METANERO/GNL","DRAFT SURVEY"],
        "BUQUE GASERO/GLP":   ["BUQUE","BARCAZA","BUQUE GASERO/GLP","BUQUE QUIMIQUERO","BUQUE METANERO/GNL","DRAFT SURVEY"],
        "BUQUE QUIMIQUERO":   ["BUQUE","BARCAZA","BUQUE GASERO/GLP","BUQUE QUIMIQUERO","BUQUE METANERO/GNL","DRAFT SURVEY"],
        "BUQUE METANERO/GNL": ["BUQUE","BARCAZA","BUQUE GASERO/GLP","BUQUE QUIMIQUERO","BUQUE METANERO/GNL","DRAFT SURVEY"],
        "DRAFT SURVEY":       ["BUQUE","BARCAZA","BUQUE GASERO/GLP","BUQUE QUIMIQUERO","BUQUE METANERO/GNL","DRAFT SURVEY"],
        "TANQUE FIJO":        ["TANQUE FIJO","TANQUE FLOTANTE","ESFERA DE GAS"],
        "TANQUE FLOTANTE":    ["TANQUE FIJO","TANQUE FLOTANTE","ESFERA DE GAS"],
        "ESFERA DE GAS":      ["TANQUE FIJO","TANQUE FLOTANTE","ESFERA DE GAS"],
        "CAMION CISTERNA":    ["CAMION CISTERNA","CAMION GAS/GLP"],
        "CAMION GAS/GLP":     ["CAMION CISTERNA","CAMION GAS/GLP"],
        "OLEODUCTO":          ["OLEODUCTO","POLIDUCTO","GASODUCTO"],
        "POLIDUCTO":          ["OLEODUCTO","POLIDUCTO","GASODUCTO"],
        "GASODUCTO":          ["OLEODUCTO","POLIDUCTO","GASODUCTO"],
        "MEDICION ELECTRICA": ["MEDICION ELECTRICA"],
    }

    # ── Color palette por tipo de producto ──────────────────────────────────
    # AGUA siempre es #3498DB. Producto NUNCA azul.
    PROD_COLORS = {
        "default":     ("#F0B429", "#D4880A"),   # amber (gasoil genérico)
        "gasoil":      ("#F0B429", "#D4880A"),   # amber/dorado
        "fuel":        ("#3D1C0A", "#5C3317"),   # marrón oscuro (fuel oil)
        "crudo":       ("#1C0A00", "#3D1A00"),   # negro petróleo
        "nafta":       ("#FFD700", "#C9A800"),   # dorado brillante
        "glp":         ("#FF6B35", "#CC4A1A"),   # naranja GLP/propano
        "gnl":         ("#E8E8F0", "#A0A0C0"),   # blanco-plata (GNL criogénico, NUNCA azul)
        "quimico":     ("#8B6914", "#6B4F10"),   # bronce/marrón químico
        "lubricante":  ("#556B2F", "#3D4F22"),   # verde oscuro lubricante
        "agua_carga":  ("#85C1E9", "#5B9FC0"),   # azul claro (agua de carga)
        "slop":        ("#8D6E63", "#6D4C41"),   # marrón sucio slop
        "gas_comp":    ("#FF8C00", "#CC6600"),   # naranja gas comprimido
        "metanol":     ("#9B59B6", "#7D3C98"),   # violeta
        "acido":       ("#E74C3C", "#C0392B"),   # rojo ácido
    }

    @staticmethod
    def contrast_text(hex_color):
        """Devuelve '#FFFFFF' o '#000000' según la luminancia del color de fondo.
        Garantiza contraste legible sobre cualquier producto."""
        try:
            r = int(hex_color[1:3], 16)
            g = int(hex_color[3:5], 16)
            b = int(hex_color[5:7], 16)
            # Luminancia relativa WCAG
            lum = 0.2126 * r + 0.7152 * g + 0.0722 * b
            return "#FFFFFF" if lum < 140 else "#000000"
        except:
            return "#FFFFFF"

    def get_prod_color(self, tk_name, etapa_key=None):
        """Devuelve (fill, outline) para el producto del tanque.
        NUNCA retorna azul (#3498DB) que está reservado para AGUA."""
        prod = ""
        if etapa_key and tk_name:
            prod = self.get_var(f"{etapa_key}_{tk_name}_prod_name").get().lower()
        prod_low = prod.lower()
        if any(x in prod_low for x in ["gasoil","diesel","go "]): return self.PROD_COLORS["gasoil"]
        if any(x in prod_low for x in ["fuel","fo ","furnace"]): return self.PROD_COLORS["fuel"]
        if any(x in prod_low for x in ["crudo","crude","pf"]): return self.PROD_COLORS["crudo"]
        if any(x in prod_low for x in ["nafta","gasolin","naphtha"]): return self.PROD_COLORS["nafta"]
        if any(x in prod_low for x in ["glp","propano","butano","gas licuado"]): return self.PROD_COLORS["glp"]
        if any(x in prod_low for x in ["gnl","lng","metano","natural"]): return self.PROD_COLORS["gnl"]
        if any(x in prod_low for x in ["quim","acid","solvente","benceno","tolueno"]): return self.PROD_COLORS["quimico"]
        if any(x in prod_low for x in ["lubri","aceite"]): return self.PROD_COLORS["lubricante"]
        if any(x in prod_low for x in ["metanol","methanol","alcohol"]): return self.PROD_COLORS["metanol"]
        if any(x in prod_low for x in ["slop","residuo"]): return self.PROD_COLORS["slop"]
        return self.PROD_COLORS["default"]

    def get_tipo_medio(self):
        v = self.get_var("car_tipo_medio", "BUQUE").get().strip()
        if not v: v = self.get_var("car_tipo_nave", "BUQUE").get().strip()
        return v or "BUQUE"

    def es_maritimo(self):
        return self.get_tipo_medio() in ("BUQUE","BARCAZA",
            "BUQUE GASERO/GLP","BUQUE QUIMIQUERO","BUQUE METANERO/GNL","DRAFT SURVEY")
    def es_tierra(self):
        return self.get_tipo_medio() in ("TANQUE FIJO","TANQUE FLOTANTE")

    def es_camion(self):
        tm = self.get_tipo_medio()
        return "CAMION" in tm

    def es_camion_gas(self):
        tm = self.get_tipo_medio()
        return "CAMION GAS" in tm

    def es_ducto(self):
        tm = self.get_tipo_medio()
        return tm in ("OLEODUCTO","POLIDUCTO","GASODUCTO")

    def es_gasoducto(self):
        return self.get_tipo_medio() == "GASODUCTO"

    def es_electrico(self):
        return self.get_tipo_medio() == "MEDICION ELECTRICA"

    def es_esfera(self):
        return self.get_tipo_medio() == "ESFERA DE GAS"

    def es_no_tradicional(self):
        """Ductos + electricidad: no tienen tanques en el sentido habitual."""
        return self.es_ducto() or self.es_electrico() or self.es_esfera()

    def es_gasero(self):
        """Buques que transportan gas a presión/criogénico (NO incluye quimiquero que es líquido)."""
        tm = self.get_tipo_medio()
        # QUIMIQUERO es tanquero LÍQUIDO, no gasero — se excluye explícitamente
        return ("GASERO" in tm or "METANERO" in tm or "GNL" in tm or
                ("GLP" in tm and "CAMION" not in tm and "QUIMIQUERO" not in tm))

    def label_unidad(self):
        """Nombre genérico de la unidad operativa para labels y PDF."""
        tm = self.get_tipo_medio()
        if "BARCAZA" in tm:         return "BARCAZA"
        if "GASERO" in tm:          return "BUQUE GASERO"
        if "QUIMIQUERO" in tm:      return "BUQUE QUIMIQUERO"
        if "METANERO" in tm or "GNL" in tm: return "BUQUE METANERO/GNL"
        if "TANQUE FIJO" in tm:     return "PLANTA - TANQUE FIJO"
        if "TANQUE FLOTANTE" in tm: return "PLANTA - TANQUE FLOTANTE"
        if "CAMION GAS" in tm:      return "CAMION GAS/GLP"
        if "CAMION" in tm:          return "CAMION CISTERNA"
        if tm == "OLEODUCTO":       return "OLEODUCTO"
        if tm == "POLIDUCTO":       return "POLIDUCTO"
        if tm == "GASODUCTO":       return "GASODUCTO"
        if tm == "MEDICION ELECTRICA": return "INSTALACION ELECTRICA"
        if tm == "ESFERA DE GAS":       return "ESFERA DE GAS"
        return "BUQUE"

    def label_contenedor(self):
        """Nombre del identificador principal."""
        tm = self.get_tipo_medio()
        if self.es_tierra():          return "ID Tanque / N° Planta:"
        if "CAMION" in tm:            return "Dominio / Patente:"
        if self.es_ducto():           return "Nombre del Ducto:"
        if self.es_electrico():       return "Instalación / Punto de Medición:"
        return "Nombre Buque/Barcaza:"

    def label_doc_principal(self):
        """Etiqueta para MANI / documento de transporte."""
        if self.es_ducto():           return "N° Acta / Expediente:"
        if self.es_electrico():       return "N° Expediente / Contrato:"
        if self.es_ducto():                       return "N° Acta / Expediente:"
        if self.es_electrico():                   return "N° Expediente / Contrato:"
        if self.es_tierra() or self.es_camion():  return "N° Expediente:"
        return "N° MANI:"

    def calc_volumen_geometrico_tierra(self, sondaje_mm, radio_interno_m, modo="INNAGE"):
        """V en litros. radio en metros, sondaje en mm.
        Techo fijo/flotante: cilindro vertical π*r²*h."""
        try:
            h = float(sondaje_mm) / 1000.0   # mm → m
            r = float(radio_interno_m)
            vol_m3 = math.pi * r**2 * h
            return vol_m3 * 1000.0            # m³ → litros
        except: return 0.0

    def calc_volumen_cilindro_horizontal(self, varilla_mm, radio_interno_m, largo_m):
        """V en litros para cisterna horizontal (camión).
        varilla = nivel desde el fondo, en mm."""
        try:
            h = float(varilla_mm) / 1000.0
            r = float(radio_interno_m)
            L = float(largo_m)
            if h < 0: h = 0
            if h > 2*r: h = 2*r
            ratio = (r - h) / r
            ratio = max(-1.0, min(1.0, ratio))
            alpha = math.acos(ratio)
            vol_m3 = L * r**2 * (alpha - math.sin(alpha)*math.cos(alpha))
            return vol_m3 * 1000.0
        except: return 0.0


    SUBREGIMENES = [
        "REMO", "ER01", "ER02", "ER03", "ER05", "ER06",
        "EC01", "EC03", "ES01", "ES02", "ES06",
        "IC04", "IC05", "IC06", "IC07", "IC65",
        "IR01", "IR02", "IS01", "IS04",
        "IT01", "IT04", "IT14",
        "RE01", "TRAN", "ZF01", "ZF06"
    ]

    def calcular_letra_dest(self, num_str):
        """Intenta calcular la letra verificadora por módulo 23 (TRWAGMYFPDXBNJZSQVHLCKE).
        Nota: el algoritmo exacto del SIM/MALVINA no es público, esta es una aproximación."""
        tabla = "TRWAGMYFPDXBNJZSQVHLCKE"
        try:
            # Extraer solo dígitos
            digits = "".join(c for c in num_str if c.isdigit())
            if not digits: return "?"
            n = int(digits)
            return tabla[n % 23]
        except:
            return "?"

    def popup_tipo_documento(self, edit_obj=None):
        """Popup para seleccionar/editar tipo de documento."""
        dlg = tk.Toplevel(self.root)
        dlg.title("Tipo de Documento" if not edit_obj else "Modificar Documento")
        dlg.transient(self.root)
        dlg.grab_set()
        dlg.update_idletasks()
        # Centrar en pantalla
        sw, sh = dlg.winfo_screenwidth(), dlg.winfo_screenheight()
        w_dlg, h_dlg = 480, 420
        dlg.geometry(f"{w_dlg}x{h_dlg}+{(sw - w_dlg)//2}+{(sh - h_dlg)//2}")

        tk.Label(dlg, text="Seleccione el tipo de documento:", font=("Arial", 8, "bold")).pack(pady=10)

        # Pre-fill if editing
        _existing_num = edit_obj["numero"].get() if edit_obj else ""
        _existing_tipo = edit_obj["tipo_doc"].get() if edit_obj else "Detallada"

        tipo_var = tk.StringVar(value="DECLARACION_DETALLADA")
        if edit_obj:
            t = _existing_tipo
            if t in ("Particular",): tipo_var.set("SOLICITUD_PARTICULAR")
            elif t in ("Expediente",): tipo_var.set("EXPEDIENTE")
            else: tipo_var.set("DECLARACION_DETALLADA")
        for val, txt in [("DECLARACION_DETALLADA", "Declaración Detallada (SIM/MALVINA)"), 
                         ("SOLICITUD_PARTICULAR", "Solicitud Particular"),
                         ("EXPEDIENTE", "Expediente")]:
            tk.Radiobutton(dlg, text=txt, variable=tipo_var, value=val, font=("Arial", 10)).pack(anchor="w", padx=30)

        # Frame para Declaración Detallada
        f_dd = ttk.LabelFrame(dlg, text="Datos Declaración Detallada")
        f_dd.pack(padx=20, pady=10, fill="x")

        tk.Label(f_dd, text="Año (2 díg):", font=("Arial", 9)).grid(row=0, column=0, padx=3, pady=3)
        var_anio = tk.StringVar()
        ttk.Entry(f_dd, textvariable=var_anio, width=4).grid(row=0, column=1, padx=3)

        tk.Label(f_dd, text="Aduana (3 díg):", font=("Arial", 9)).grid(row=0, column=2, padx=3)
        var_aduana = tk.StringVar()
        ttk.Entry(f_dd, textvariable=var_aduana, width=5).grid(row=0, column=3, padx=3)

        tk.Label(f_dd, text="Subrégimen:", font=("Arial", 9)).grid(row=0, column=4, padx=3)
        var_subreg = tk.StringVar(value="REMO")
        ttk.Combobox(f_dd, textvariable=var_subreg, values=self.SUBREGIMENES, state="readonly", width=7).grid(row=0, column=5, padx=3)

        tk.Label(f_dd, text="Número (hasta 6 díg):", font=("Arial", 9)).grid(row=1, column=0, columnspan=2, padx=3, pady=3)
        var_num = tk.StringVar()
        ttk.Entry(f_dd, textvariable=var_num, width=8).grid(row=1, column=2, padx=3)

        tk.Label(f_dd, text="Letra:", font=("Arial", 9)).grid(row=1, column=3, padx=3)
        var_letra = tk.StringVar()
        e_letra = ttk.Entry(f_dd, textvariable=var_letra, width=3)
        e_letra.grid(row=1, column=4, padx=3)

        # Vista previa
        var_preview = tk.StringVar(value="")
        tk.Label(f_dd, textvariable=var_preview, font=("Arial", 8, "bold"), fg="blue").grid(row=2, column=0, columnspan=6, pady=5)

        def actualizar_preview(*args):
            a = var_anio.get().zfill(2)[-2:]
            ad = var_aduana.get().zfill(3)[-3:]
            sr = var_subreg.get()
            n = var_num.get().zfill(6)[-6:]
            code = f"{a}{ad}{sr}{n}"
            letra = var_letra.get().upper().strip()
            if not letra:
                letra = self.calcular_letra_dest(f"{a}{ad}{n}")
                var_letra.set(letra)
            var_preview.set(f"{code}{letra}")

        for v in [var_anio, var_aduana, var_subreg, var_num]:
            v.trace_add("write", actualizar_preview)

        # Frame Solicitud/Expediente
        f_sp = ttk.LabelFrame(dlg, text="Número de Solicitud/Expediente")
        f_sp.pack(padx=20, pady=5, fill="x")
        var_libre = tk.StringVar()
        ttk.Entry(f_sp, textvariable=var_libre, width=40).pack(padx=10, pady=5)

        def toggle_frames(*args):
            is_dd = tipo_var.get() == "DECLARACION_DETALLADA"
            for child in f_dd.winfo_children():
                try: child.configure(state="normal" if is_dd else "disabled")
                except: pass
            for child in f_sp.winfo_children():
                try: child.configure(state="disabled" if is_dd else "normal")
                except: pass
        tipo_var.trace_add("write", toggle_frames)
        toggle_frames()

        def aceptar():
            tipo = tipo_var.get()
            if tipo == "DECLARACION_DETALLADA":
                actualizar_preview()
                desc = var_preview.get()
                tipo_doc_val = "Detallada"
            elif tipo == "SOLICITUD_PARTICULAR":
                desc = var_libre.get() or "Solicitud Particular"
                tipo_doc_val = "Particular"
            else:
                desc = var_libre.get() or "Expediente"
                tipo_doc_val = "Expediente"
            dlg.destroy()
            if edit_obj:
                # Editar en lugar de agregar
                edit_obj["numero"].set(desc)
                edit_obj["tipo_doc"].set(tipo_doc_val)
            else:
                self.agregar_ddt_row(data={"numero": desc, "tipo_doc": tipo_doc_val, "producto": "", 
                                           "pos_arancel": "", "litros": "", "kilos": "", "densidad": "",
                                           "num_planilla": "", "valor_litro": ""})

        f_bot = tk.Frame(dlg)
        f_bot.pack(pady=10)
        tk.Button(f_bot, text="ACEPTAR", bg="#27AE60", fg="white", font=("Arial", 8, "bold"), command=aceptar).pack(side="left", padx=10)
        tk.Button(f_bot, text="Cancelar", bg="#E74C3C", fg="white", font=("Arial", 10), command=dlg.destroy).pack(side="left", padx=10)

    def agregar_ddt_row(self, def_prod="", data=None):
        self.ddt_counter += 1
        did = str(self.ddt_counter)
        obj = {
            "id": did,
            "numero": tk.StringVar(value=data["numero"] if data else ""),
            "num_planilla": tk.StringVar(value=data["num_planilla"] if data and "num_planilla" in data else ""),
            "tipo_doc": tk.StringVar(value=data["tipo_doc"] if data and "tipo_doc" in data else "Detallada"),
            "producto": tk.StringVar(value=data["producto"] if data else def_prod),
            "pos_arancel": tk.StringVar(value=data["pos_arancel"] if data else ""),
            "litros": tk.StringVar(value=data["litros"] if data else ""),
            "kilos": tk.StringVar(value=data["kilos"] if data else ""),
            "densidad": tk.StringVar(value=data["densidad"] if data else ""),
            "valor_litro": tk.StringVar(value=data["valor_litro"] if data and "valor_litro" in data else ""),
            "divisa": tk.StringVar(value=data["divisa"] if data and "divisa" in data else "ARS"),
            "tipo_cambio": tk.StringVar(value=data["tipo_cambio"] if data and "tipo_cambio" in data else ""),
            "salidas": [], "main_frame": None, "salidas_frame": None
        }
        # ── Actores del documento (cada detallada puede tener su propio
        #    despachante / importador-exportador / ATA). Al agregar un doc
        #    nuevo se precargan desde el primer documento que los tenga. ──
        def _actor_ini(k):
            if data is not None:
                return data.get(k, "")
            for d_prev in self.ddt_data:
                v_prev = d_prev.get(k)
                if isinstance(v_prev, tk.StringVar) and v_prev.get().strip():
                    return v_prev.get()
            return ""
        for _ak in self.DDT_ACTOR_KEYS:
            obj[_ak] = tk.StringVar(value=_actor_ini(_ak))
        obj["litros"].trace_add("write", lambda *args: self.auto_calc_densidad(obj["litros"], obj["kilos"], obj["densidad"]))
        obj["kilos"].trace_add("write", lambda *args: self.auto_calc_densidad(obj["litros"], obj["kilos"], obj["densidad"]))

        frame = ttk.LabelFrame(self.ddt_stack, text=f"Item Documento #{self.ddt_counter}")
        frame.pack(fill="x", pady=10, padx=5)
        obj["main_frame"] = frame
        # Actualizar título del LabelFrame cuando cambia la descripción
        def update_frame_title(*args, f=frame, o=obj, n=self.ddt_counter):
            desc = o["numero"].get().strip()
            if desc:
                f.configure(text=f"Item #{n}: {desc}")
            else:
                f.configure(text=f"Item Documento #{n}")
        obj["numero"].trace_add("write", update_frame_title)
        update_frame_title()  # Actualizar de inmediato si ya tiene datos
        f_header = ttk.Frame(frame)
        f_header.pack(fill="x", padx=5, pady=5)
        # ROW 0: Descripcion | entry | Modificar | N°Planilla | entry | Producto | entry | Pos.Aranc | entry | [ELIMINAR]
        tk.Label(f_header, text="Doc:", font=("Arial", 8, "bold")).grid(row=0, column=0, sticky="e", padx=2)
        e_num = ttk.Entry(f_header, textvariable=obj["numero"], width=28, state="readonly")
        e_num.grid(row=0, column=1, padx=3, sticky="ew")
        tk.Button(f_header, text="Modificar", bg="#1976D2", fg="white", font=("Arial", 8, "bold"),
                  command=lambda o=obj: self.popup_tipo_documento(edit_obj=o)).grid(row=0, column=2, padx=3)
        tk.Label(f_header, text="Producto:", font=("Arial", 8, "bold")).grid(row=0, column=3, sticky="e", padx=2)
        ttk.Entry(f_header, textvariable=obj["producto"], width=18).grid(row=0, column=4, padx=3, sticky="ew")
        tk.Label(f_header, text="Pos.Aranc:", font=("Arial", 8, "bold")).grid(row=0, column=5, sticky="e", padx=2)
        ttk.Entry(f_header, textvariable=obj["pos_arancel"], width=22).grid(row=0, column=6, padx=3, sticky="ew")
        # ROW 1: Litros | entry | Kilos | entry | Densidad | entry | $/Litro | entry | Divisa | combo | T/C | entry
        tk.Label(f_header, text="Litros:", font=("Arial", 9)).grid(row=1, column=0, sticky="e", padx=2, pady=2)
        ttk.Entry(f_header, textvariable=obj["litros"], width=12).grid(row=1, column=1, sticky="ew")
        tk.Label(f_header, text="Kilos:", font=("Arial", 9)).grid(row=1, column=2, sticky="e", padx=2)
        ttk.Entry(f_header, textvariable=obj["kilos"], width=12).grid(row=1, column=3)
        tk.Label(f_header, text="Densidad:", font=("Arial", 8, "bold"), fg="blue").grid(row=1, column=4, sticky="e", padx=2)
        ttk.Entry(f_header, textvariable=obj["densidad"], width=12, state="readonly").grid(row=1, column=5, sticky="ew")
        tk.Label(f_header, text="$/Litro:", font=("Arial", 8, "bold"), fg="#6A1B9A").grid(row=1, column=6, sticky="e", padx=2)
        ttk.Entry(f_header, textvariable=obj["valor_litro"], width=12).grid(row=1, column=7)
        tk.Label(f_header, text="Divisa:", font=("Arial", 8, "bold"), fg="#6A1B9A").grid(row=1, column=8, sticky="e", padx=2)
        cb_divisa = ttk.Combobox(f_header, textvariable=obj["divisa"], values=["ARS", "USD", "EUR", "BRL", "GBP", "Otra"], state="readonly", width=6)
        cb_divisa.grid(row=1, column=9, padx=2)
        tk.Label(f_header, text="T/C:", font=("Arial", 9), fg="#6A1B9A").grid(row=1, column=10, sticky="e", padx=2)
        ttk.Entry(f_header, textvariable=obj["tipo_cambio"], width=10).grid(row=1, column=11, padx=2)
        # ROW 2: Descripción divisa (si Otra) + ELIMINAR al final
        tk.Button(f_header, text="  ELIMINAR  ", bg="red", fg="white", font=("Arial", 8, "bold"),
                  command=lambda: self.borrar_ddt(obj)).grid(row=2, column=9, columnspan=3, padx=8, pady=3, sticky="e")
        # Campo descripción de divisa (visible solo si "Otra")
        obj["divisa_desc"] = tk.StringVar(value=data["divisa_desc"] if data and "divisa_desc" in data else "")
        lbl_divisa_desc = tk.Label(f_header, text="Desc. Divisa:", font=("Arial", 9), fg="#6A1B9A")
        e_divisa_desc = ttk.Entry(f_header, textvariable=obj["divisa_desc"], width=12)
        def toggle_divisa_desc(*args):
            if obj["divisa"].get() == "Otra":
                lbl_divisa_desc.grid(row=2, column=0, padx=2, sticky="e", columnspan=2)
                e_divisa_desc.grid(row=2, column=2, padx=2, columnspan=5, sticky="ew")
            else:
                lbl_divisa_desc.grid_remove()
                e_divisa_desc.grid_remove()
        obj["divisa"].trace_add("write", toggle_divisa_desc)
        toggle_divisa_desc()
        # ROW 3-4: Actores del documento (despachante / imp-exp / ATA + CUITs)
        _fga = ("Arial", 8, "bold")
        tk.Label(f_header, text="Despachante:", font=_fga, fg="#1B4F72").grid(row=3, column=0, sticky="e", padx=2, pady=(8, 0))
        ttk.Entry(f_header, textvariable=obj["despachante"], width=28).grid(row=3, column=1, sticky="ew", padx=3, pady=(8, 0))
        tk.Label(f_header, text="CUIT:", font=_fga, fg="#1B4F72").grid(row=3, column=2, sticky="e", padx=2, pady=(8, 0))
        e_cd = tk.Entry(f_header, textvariable=obj["cuit_desp"], width=14)
        e_cd.grid(row=3, column=3, sticky="w", padx=3, pady=(8, 0))
        tk.Label(f_header, text="Imp/Exp:", font=_fga, fg="#1B4F72").grid(row=3, column=4, sticky="e", padx=2, pady=(8, 0))
        ttk.Entry(f_header, textvariable=obj["impexp"], width=24).grid(row=3, column=5, columnspan=2, sticky="ew", padx=3, pady=(8, 0))
        tk.Label(f_header, text="CUIT:", font=_fga, fg="#1B4F72").grid(row=3, column=7, sticky="e", padx=2, pady=(8, 0))
        e_ci = tk.Entry(f_header, textvariable=obj["cuit_impexp"], width=14)
        e_ci.grid(row=3, column=8, columnspan=2, sticky="w", padx=3, pady=(8, 0))
        tk.Label(f_header, text="ATA:", font=_fga, fg="#1B4F72").grid(row=4, column=0, sticky="e", padx=2, pady=(2, 4))
        ttk.Entry(f_header, textvariable=obj["ata"], width=28).grid(row=4, column=1, sticky="ew", padx=3, pady=(2, 4))
        tk.Label(f_header, text="CUIT:", font=_fga, fg="#1B4F72").grid(row=4, column=2, sticky="e", padx=2, pady=(2, 4))
        e_ca = tk.Entry(f_header, textvariable=obj["cuit_ata"], width=14)
        e_ca.grid(row=4, column=3, sticky="w", padx=3, pady=(2, 4))
        for _ew, _ev in ((e_cd, obj["cuit_desp"]), (e_ci, obj["cuit_impexp"]), (e_ca, obj["cuit_ata"])):
            _ew.bind("<FocusOut>", lambda ev, wdg=_ew, v=_ev: self.on_cuit_validate(wdg, v))
        def _copiar_actores(o=obj):
            src = next((d for d in self.ddt_data if d is not o and any(
                isinstance(d.get(k), tk.StringVar) and d.get(k).get().strip()
                for k in self.DDT_ACTOR_KEYS)), None)
            if not src: return
            for k in self.DDT_ACTOR_KEYS:
                sv = src.get(k)
                if isinstance(sv, tk.StringVar): o[k].set(sv.get())
        tk.Button(f_header, text="⧉ Copiar del 1er doc", bg="#5D6D7E", fg="white", font=("Arial", 8),
                  command=_copiar_actores).grid(row=4, column=4, columnspan=3, sticky="w", padx=3, pady=(2, 4))
        # T/C now in row=1 cols 10-11 above
        f_sub = ttk.Frame(frame)
        f_sub.pack(fill="x", padx=20, pady=5)
        tk.Button(f_sub, text="+ Agregar Salida", bg="#1976D2", fg="white", font=("Arial", 8), command=lambda: self.agregar_salida(obj)).pack(anchor="w")
        obj["salidas_frame"] = ttk.Frame(f_sub)
        obj["salidas_frame"].pack(fill="x", pady=5)
        f_titles = ttk.Frame(obj["salidas_frame"])
        f_titles.pack(fill="x")
        tk.Label(f_titles, text="N° Salida", width=20, fg="gray", font=("Arial", 7)).pack(side="left", padx=2)
        tk.Label(f_titles, text="Litros", width=12, fg="gray", font=("Arial", 7)).pack(side="left", padx=2)
        tk.Label(f_titles, text="Kilos", width=12, fg="gray", font=("Arial", 7)).pack(side="left", padx=2)
        tk.Label(f_titles, text="Densidad", width=12, fg="blue", font=("Arial", 7)).pack(side="left", padx=2)
        if data and "salidas" in data:
            for s_data in data["salidas"]: self.agregar_salida(obj, s_data)
        self.ddt_data.append(obj)
        self.actualizar_combos_ddt()

    def agregar_funcionario_row(self, data=None):
        self.func_counter += 1
        funciones = db_get_funciones()
        _apellido_default = data.get("apellido", "") if data else ""
        _nombre_default   = data.get("nombre",   "") if data else ""
        if data and not _apellido_default and _nombre_default:
            parts = _nombre_default.strip().split(" ", 1)
            _apellido_default = parts[0]
            _nombre_default   = parts[1] if len(parts) > 1 else ""
        obj = {
            "cuil":     tk.StringVar(value=data["cuil"]    if data else ""),
            "legajo":   tk.StringVar(value=data["legajo"]  if data else ""),
            "apellido": tk.StringVar(value=_apellido_default),
            "nombre":   tk.StringVar(value=_nombre_default),
            "funcion":  tk.StringVar(value=data["funcion"] if data else funciones[0]),
            "frame":    None,
        }
        row = ttk.Frame(self.func_stack)
        row.pack(fill="x", pady=2)
        obj["frame"] = row

        def _autocomplete_from_db(*args):
            """Al salir del campo CUIL o APELLIDO, busca en DB y completa."""
            cuil_v = obj["cuil"].get().strip()
            apell_v = obj["apellido"].get().strip()
            found = []
            if cuil_v and len(cuil_v) >= 4:
                found = db_buscar_funcionarios(cuil_v, "cuil")
            if not found and apell_v and len(apell_v) >= 3:
                found = db_buscar_funcionarios(apell_v, "apellido")
            if found:
                r = found[0]
                obj["cuil"].set(r["cuil"]); obj["legajo"].set(r["legajo"])
                obj["apellido"].set(r["apellido"]); obj["nombre"].set(r["nombre"])
                obj["funcion"].set(r["funcion"] if r["funcion"] in funciones else funciones[0])

        def _save_to_db(*args):
            """Al salir del campo LEGAJO, guarda o actualiza en DB si hay datos suficientes."""
            cuil = obj["cuil"].get().strip()
            legajo = obj["legajo"].get().strip()
            apellido = obj["apellido"].get().strip()
            nombre = obj["nombre"].get().strip()
            funcion = obj["funcion"].get()
            aduana = self.get_var("car_lugar", "").get()
            if cuil and apellido:
                db_guardar_funcionario(cuil, legajo, apellido, nombre, funcion, aduana)

        e_cuil = tk.Entry(row, textvariable=obj["cuil"], width=16, font=("Arial", 9))
        e_cuil.grid(row=0, column=0, padx=2, sticky="w")
        e_cuil.bind("<FocusOut>", lambda ev, w=e_cuil, v=obj["cuil"]: (self.on_cuit_validate(w, v), _autocomplete_from_db()))
        e_cuil.bind("<Return>", lambda ev: _autocomplete_from_db())

        e_leg = ttk.Entry(row, textvariable=obj["legajo"], width=13)
        e_leg.grid(row=0, column=1, padx=2, sticky="w")
        e_leg.bind("<FocusOut>", lambda ev: _save_to_db())

        e_apell = ttk.Entry(row, textvariable=obj["apellido"], width=22)
        e_apell.grid(row=0, column=2, padx=2, sticky="w")
        e_apell.bind("<FocusOut>", lambda ev: _autocomplete_from_db())
        e_apell.bind("<Return>", lambda ev: _autocomplete_from_db())

        ttk.Entry(row, textvariable=obj["nombre"], width=22).grid(row=0, column=3, padx=2, sticky="w")
        cb_fun = ttk.Combobox(row, textvariable=obj["funcion"], values=funciones, state="readonly", width=28)
        cb_fun.grid(row=0, column=4, padx=2, sticky="w")
        cb_fun.bind("<<ComboboxSelected>>", lambda ev: _save_to_db())

        tk.Button(row, text="X", bg="#E74C3C", fg="white",
                  command=lambda o=obj: (self.func_stack.pack_propagate(True),
                                        o["frame"].destroy(),
                                        self.funcionarios_data.remove(o))
                  ).grid(row=0, column=5, padx=2)

        self.funcionarios_data.append(obj)

    def buscar_y_agregar_funcionario_db(self):
        """Popup busqueda rapida en DB de funcionarios para agregar a caratula."""
        try:
            top = tk.Toplevel(self.root)
            top.title("Buscar Funcionario en Base de Datos")
            top.geometry("920x520")
            top.resizable(True, True)
            top.grab_set()

            fh = tk.Frame(top, bg="#1D6A39")
            fh.pack(fill="x")
            tk.Label(fh, text="BUSCAR Y AGREGAR FUNCIONARIOS A LA CARATULA",
                     bg="#1D6A39", fg="white", font=("Arial", 10, "bold")).pack(side="left", padx=14, pady=8)
            tk.Label(fh, text="Seleccion multiple con Ctrl+Click  |  Doble click = agregar y cerrar",
                     bg="#1D6A39", fg="#A9DFBF", font=("Arial", 8)).pack(side="left")

            f_srch = tk.Frame(top, bg="#EBF5FB")
            f_srch.pack(fill="x")
            tk.Label(f_srch, text="Buscar:", bg="#EBF5FB",
                     font=("Arial", 9, "bold")).pack(side="left", padx=8, pady=6)
            v_q = tk.StringVar()
            e_q = tk.Entry(f_srch, textvariable=v_q, width=34, font=("Arial", 10))
            e_q.pack(side="left", padx=4, pady=6)
            tk.Label(f_srch, text="(apellido, CUIL o funcion — vacio = todos)",
                     bg="#EBF5FB", fg="gray", font=("Arial", 8)).pack(side="left", padx=6)

            cols   = ("cuil", "legajo", "apellido", "nombre", "funcion", "aduana")
            hdrs   = ("CUIL",  "Legajo",  "Apellido",  "Nombre",   "Funcion",   "Aduana")
            widths = (120,      80,        150,          150,         140,          120)
            f_tree = tk.Frame(top)
            f_tree.pack(fill="both", expand=True, padx=10, pady=6)
            tree = ttk.Treeview(f_tree, columns=cols, show="headings",
                                selectmode="extended", height=14)
            for c, h, w in zip(cols, hdrs, widths):
                tree.heading(c, text=h)
                tree.column(c, width=w, minwidth=40)
            vsb = ttk.Scrollbar(f_tree, orient="vertical", command=tree.yview)
            tree.configure(yscrollcommand=vsb.set)
            vsb.pack(side="right", fill="y")
            tree.pack(fill="both", expand=True)
            tree.tag_configure("odd",  background="#EAF4FB")
            tree.tag_configure("even", background="white")

            def _load(q=""):
                for row in tree.get_children():
                    tree.delete(row)
                if q:
                    rows = (db_buscar_funcionarios(q, "apellido") +
                            db_buscar_funcionarios(q, "cuil") +
                            db_buscar_funcionarios(q, "funcion"))
                    # Dedup por registro completo: un agente con varias
                    # funciones aparece una vez POR FUNCIÓN
                    seen = set(); unique = []
                    for r in rows:
                        k = (r.get("cuil",""), r.get("legajo",""), r.get("funcion",""))
                        if k not in seen:
                            seen.add(k); unique.append(r)
                    rows = unique
                else:
                    rows = db_todos_funcionarios()
                for i, r in enumerate(rows):
                    tree.insert("", "end", tags=("odd" if i%2 else "even",), values=(
                        r.get("cuil",""), r.get("legajo",""),
                        r.get("apellido",""), r.get("nombre",""),
                        r.get("funcion",""), r.get("aduana","")
                    ))

            def _on_search(*_):  _load(v_q.get().strip())
            v_q.trace_add("write", _on_search)
            e_q.bind("<Return>", _on_search)

            def _agregar():
                sels = tree.selection()
                if not sels:
                    messagebox.showwarning("Sin seleccion",
                        "Seleccione al menos un funcionario.", parent=top)
                    return
                for s in sels:
                    v = tree.item(s, "values")
                    self.agregar_funcionario_row({
                        "cuil": v[0], "legajo": v[1],
                        "apellido": v[2], "nombre": v[3], "funcion": v[4]
                    })
                top.destroy()

            tree.bind("<Double-1>", lambda e: _agregar())

            f_bot = tk.Frame(top, bg="#2C3E50")
            f_bot.pack(fill="x", side="bottom")
            tk.Button(f_bot, text="Agregar seleccionados a caratula",
                      bg="#27AE60", fg="white", font=("Arial", 9, "bold"),
                      command=_agregar, cursor="hand2"
                      ).pack(side="left", padx=12, pady=8, ipadx=10, ipady=3)
            tk.Button(f_bot, text="Cancelar", bg="#7F8C8D", fg="white",
                      font=("Arial", 8), command=top.destroy
                      ).pack(side="right", padx=12, pady=8)
            _load()
            e_q.focus_set()
        except Exception as e:
            import traceback; traceback.print_exc()
            messagebox.showerror("Error", str(e))


    def abrir_gestor_funcionarios_db(self):
        """Ventana completa de gestión de la base de datos de funcionarios: listar, agregar, editar, eliminar, importar CSV."""
        try:
            top = tk.Toplevel(self.root)
            top.title("Base de Datos de Funcionarios")
            top.geometry("1000x620")
            top.resizable(True, True)

            # ── Header ──────────────────────────────────────────────────────
            fh = tk.Frame(top, bg="#1B3A5C")
            fh.pack(fill="x")
            tk.Label(fh, text="BASE DE DATOS DE FUNCIONARIOS", bg="#1B3A5C", fg="white",
                     font=("Arial", 10, "bold")).pack(side="left", padx=14, pady=8)
            tk.Label(fh, text="Agregue, edite o elimine funcionarios. Importe desde CSV.",
                     bg="#1B3A5C", fg="#AED6F1", font=("Arial", 8)).pack(side="left")

            # ── Barra de búsqueda + botones de acción ────────────────────────
            f_act = tk.Frame(top, bg="#EBF5FB", bd=1, relief="flat")
            f_act.pack(fill="x", padx=0, pady=0)
            tk.Label(f_act, text="Buscar:", bg="#EBF5FB", font=("Arial", 8, "bold")).pack(side="left", padx=8, pady=6)
            v_search = tk.StringVar()
            e_search = tk.Entry(f_act, textvariable=v_search, width=26, font=("Arial", 9))
            e_search.pack(side="left", padx=4, pady=6)
            tk.Button(f_act, text="Agregar nuevo", bg="#27AE60", fg="white",
                      font=("Arial", 8, "bold"), command=lambda: _abrir_form()).pack(side="left", padx=8, pady=6, ipadx=6, ipady=2)
            tk.Button(f_act, text="Editar seleccionado", bg="#2980B9", fg="white",
                      font=("Arial", 8, "bold"), command=lambda: _editar()).pack(side="left", padx=4, pady=6, ipadx=6, ipady=2)
            tk.Button(f_act, text="Usar en operacion", bg="#8E44AD", fg="white",
                      font=("Arial", 8, "bold"), command=lambda: _usar()).pack(side="left", padx=4, pady=6, ipadx=6, ipady=2)
            tk.Button(f_act, text="Importar CSV", bg="#D35400", fg="white",
                      font=("Arial", 8, "bold"), command=lambda: _importar_csv()).pack(side="left", padx=4, pady=6, ipadx=6, ipady=2)
            tk.Button(f_act, text="Eliminar", bg="#E74C3C", fg="white",
                      font=("Arial", 8, "bold"), command=lambda: _delete()).pack(side="left", padx=4, pady=6, ipadx=6, ipady=2)
            lbl_count = tk.Label(f_act, text="", bg="#EBF5FB", font=("Arial", 8), fg="#1B3A5C")
            lbl_count.pack(side="right", padx=12)

            # ── Treeview ─────────────────────────────────────────────────────
            f_tree = tk.Frame(top)
            f_tree.pack(fill="both", expand=True, padx=8, pady=4)
            cols = ("CUIL", "Legajo", "Apellido", "Nombre", "Funcion", "Aduana", "Lugar Op.")
            tree = ttk.Treeview(f_tree, columns=cols, show="headings", height=18, selectmode="extended")
            tree.heading("CUIL",     text="CUIL",        command=lambda: _sort("CUIL"))
            tree.heading("Legajo",   text="Legajo",      command=lambda: _sort("Legajo"))
            tree.heading("Apellido", text="Apellido",    command=lambda: _sort("Apellido"))
            tree.heading("Nombre",   text="Nombre",      command=lambda: _sort("Nombre"))
            tree.heading("Funcion",  text="Funcion",     command=lambda: _sort("Funcion"))
            tree.heading("Aduana",   text="Aduana",      command=lambda: _sort("Aduana"))
            tree.heading("Lugar Op.",text="Lugar Operativo", command=lambda: _sort("Lugar Op."))
            tree.column("CUIL",     width=120); tree.column("Legajo",   width=70)
            tree.column("Apellido", width=150); tree.column("Nombre",   width=130)
            tree.column("Funcion",  width=170); tree.column("Aduana",   width=150)
            tree.column("Lugar Op.",width=180)
            vsb = ttk.Scrollbar(f_tree, orient="vertical",   command=tree.yview)
            hsb = ttk.Scrollbar(f_tree, orient="horizontal", command=tree.xview)
            tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
            tree.grid(row=0, column=0, sticky="nsew")
            vsb.grid(row=0, column=1, sticky="ns")
            hsb.grid(row=1, column=0, sticky="ew")
            f_tree.grid_rowconfigure(0, weight=1); f_tree.grid_columnconfigure(0, weight=1)
            tree.bind("<Double-1>", lambda e: _editar())

            # ── Hint CSV ─────────────────────────────────────────────────────
            f_hint = tk.Frame(top, bg="#FEF9E7")
            f_hint.pack(fill="x", padx=8, pady=2)
            tk.Label(f_hint, text="CSV para importar: cuil,legajo,apellido,nombre,funcion,aduana,lugar_operativo  (una fila por funcionario, puede tener encabezado)",
                     bg="#FEF9E7", fg="#6E4B00", font=("Arial", 7), anchor="w").pack(fill="x", padx=8, pady=3)

            # ── Footer ───────────────────────────────────────────────────────
            f_bot = tk.Frame(top, bg="#2C3E50")
            f_bot.pack(fill="x", side="bottom")
            tk.Button(f_bot, text="Cerrar", bg="#7F8C8D", fg="white",
                      font=("Arial", 8, "bold"), command=top.destroy).pack(side="right", padx=10, pady=6, ipadx=8, ipady=2)

            # ── Sort state ───────────────────────────────────────────────────
            _sort_col = [None]; _sort_rev = [False]

            def _load(query=""):
                tree.delete(*tree.get_children())
                if query:
                    rows = db_buscar_funcionarios(query, "apellido") + db_buscar_funcionarios(query, "cuil") + db_buscar_funcionarios(query, "funcion")
                    seen = set(); unique = []
                    for r in rows:
                        k = (r["cuil"], r["legajo"], r["funcion"])
                        if k not in seen: seen.add(k); unique.append(r)
                else:
                    unique = db_todos_funcionarios()
                for r in unique:
                    tree.insert("", "end", values=(r["cuil"],r["legajo"],r["apellido"],r["nombre"],r["funcion"],r["aduana"],r.get("lugar_operativo","")))
                lbl_count.config(text=f"{len(unique)} registro(s)")

            def _sort(col):
                items = [(tree.set(iid, col), iid) for iid in tree.get_children("")]
                rev = _sort_col[0] == col and not _sort_rev[0]
                items.sort(reverse=rev)
                for i, (_, iid) in enumerate(items): tree.move(iid, "", i)
                _sort_col[0] = col; _sort_rev[0] = rev

            def _on_search(*a): _load(v_search.get().strip())
            v_search.trace_add("write", _on_search)
            e_search.bind("<Return>", _on_search)

            def _abrir_form(existing_vals=None):
                """Formulario para agregar o editar un funcionario."""
                frm = tk.Toplevel(top)
                frm.title("Nuevo Funcionario" if not existing_vals else "Editar Funcionario")
                frm.geometry("480x400")
                frm.resizable(False, False)
                frm.grab_set()
                fh2 = tk.Frame(frm, bg="#1B3A5C"); fh2.pack(fill="x")
                tk.Label(fh2, text="DATOS DEL FUNCIONARIO", bg="#1B3A5C", fg="white",
                         font=("Arial", 9, "bold")).pack(padx=12, pady=8)
                ff = tk.Frame(frm, pady=8); ff.pack(fill="both", expand=True, padx=16)
                # Opciones de aduana desde la DB
                _adu_list = db_get_aduanas()
                _adu_opts = [f"{a['codigo']} - {a['nombre']}" for a in _adu_list]
                funcion_opciones = db_get_funciones()
                vs = {}
                simple_fields = [("CUIL:", "cuil"), ("Legajo:", "legajo"),
                                  ("Apellido:", "apellido"), ("Nombre:", "nombre")]
                for i, (lbl, key) in enumerate(simple_fields):
                    tk.Label(ff, text=lbl, font=("Arial", 8, "bold"), width=14, anchor="e").grid(row=i, column=0, sticky="e", pady=3, padx=4)
                    v = tk.StringVar(value=existing_vals.get(key, "") if existing_vals else "")
                    vs[key] = v
                    tk.Entry(ff, textvariable=v, width=30, font=("Arial", 9)).grid(row=i, column=1, sticky="w", pady=3, padx=4)
                # Funcion combobox
                v_fun = tk.StringVar(value=existing_vals.get("funcion","") if existing_vals else (funcion_opciones[0] if funcion_opciones else ""))
                vs["funcion"] = v_fun
                tk.Label(ff, text="Funcion:", font=("Arial", 8, "bold"), width=14, anchor="e").grid(row=4, column=0, sticky="e", pady=3, padx=4)
                ttk.Combobox(ff, textvariable=v_fun, values=funcion_opciones, width=28, state="normal").grid(row=4, column=1, sticky="w", pady=3, padx=4)
                # Aduana combobox desde DB
                v_adu = tk.StringVar()
                vs["aduana"] = v_adu
                if existing_vals and existing_vals.get("aduana"):
                    adu_raw = existing_vals["aduana"]
                    match = [o for o in _adu_opts if adu_raw.split(" - ")[0].zfill(3) in o or adu_raw in o]
                    v_adu.set(match[0] if match else adu_raw)
                else:
                    cur_adu = self.get_var("car_lugar").get()
                    match = [o for o in _adu_opts if cur_adu.split(" - ")[0].zfill(3) in o]
                    v_adu.set(match[0] if match else (_adu_opts[0] if _adu_opts else ""))
                tk.Label(ff, text="Aduana:", font=("Arial", 8, "bold"), width=14, anchor="e").grid(row=5, column=0, sticky="e", pady=3, padx=4)
                cb_adu = ttk.Combobox(ff, textvariable=v_adu, values=_adu_opts, state="normal", width=28)
                cb_adu.grid(row=5, column=1, sticky="w", pady=3, padx=4)
                tk.Label(ff, text="(Solo aduanas registradas en ABM de Aduanas)",
                         font=("Arial", 7), fg="#888", anchor="w").grid(row=6, column=1, sticky="w", padx=4)

                # ── Lugar Operativo (filtrado por aduana seleccionada) ─────────────
                v_lop = tk.StringVar(value=existing_vals.get("lugar_operativo","") if existing_vals else "")
                vs["lugar_operativo"] = v_lop
                tk.Label(ff, text="Lugar Operativo:", font=("Arial", 8, "bold"), width=14, anchor="e").grid(row=7, column=0, sticky="e", pady=3, padx=4)
                cb_lop = ttk.Combobox(ff, textvariable=v_lop, values=[], state="normal", width=28)
                cb_lop.grid(row=7, column=1, sticky="w", pady=3, padx=4)
                tk.Label(ff, text="(Lugares operativos de la aduana seleccionada)",
                         font=("Arial", 7), fg="#888", anchor="w").grid(row=8, column=1, sticky="w", padx=4)

                def _update_lop(*a):
                    """Filtra lugares operativos según la aduana elegida."""
                    adu_sel = v_adu.get().strip()
                    adu_cod = adu_sel.split(" - ")[0].strip().zfill(3) if " - " in adu_sel else adu_sel.strip().zfill(3)
                    lops = db_get_lugares_operativos(adu_cod)
                    lop_opts = [f"{l['codigo']} - {l['descripcion']}" for l in lops]
                    cb_lop["values"] = lop_opts
                    # Mantener el valor si aún es válido, sino limpiar
                    cur_lop = v_lop.get()
                    if cur_lop and cur_lop not in lop_opts:
                        # Intentar encontrar por código
                        match_l = [o for o in lop_opts if cur_lop.split(" - ")[0] in o]
                        v_lop.set(match_l[0] if match_l else "")

                v_adu.trace_add("write", _update_lop)
                _update_lop()  # cargar al abrir

                # Validación de CUIL
                def _validar_cuil(val):
                    cuil = val.replace("-","").replace(".","").replace(" ","")
                    if len(cuil) == 11 and cuil.isdigit(): return True, cuil
                    return False, cuil
                def _guardar_form():
                    ok_cuil, cuil_clean = _validar_cuil(vs["cuil"].get())
                    if not ok_cuil:
                        messagebox.showwarning("CUIL inválido", "El CUIL debe tener 11 dígitos.", parent=frm); return
                    apellido = vs["apellido"].get().strip()
                    if not apellido:
                        messagebox.showwarning("Dato faltante", "El Apellido es obligatorio.", parent=frm); return
                    adu_val = vs["aduana"].get().strip()
                    lop_val = vs["lugar_operativo"].get().strip()
                    legajo_new = vs["legajo"].get().strip()
                    # If editing and cuil/legajo changed, delete old record first
                    if existing_vals:
                        old_cuil = existing_vals.get("cuil", "")
                        old_legajo = existing_vals.get("legajo", "")
                        old_funcion = existing_vals.get("funcion", "")
                        if (old_cuil != cuil_clean or old_legajo != legajo_new
                                or old_funcion != vs["funcion"].get().strip()):
                            db_eliminar_funcionario(old_cuil, old_legajo, old_funcion)
                    db_guardar_funcionario(cuil_clean, legajo_new,
                                          apellido, vs["nombre"].get().strip(),
                                          vs["funcion"].get().strip(), adu_val, lop_val)
                    _load(v_search.get()); frm.destroy()
                fb2 = tk.Frame(frm, bg="#2C3E50"); fb2.pack(fill="x", side="bottom")
                tk.Button(fb2, text="Guardar", bg="#27AE60", fg="white", font=("Arial", 8, "bold"),
                          command=_guardar_form).pack(side="left", padx=10, pady=6, ipadx=10, ipady=2)
                tk.Button(fb2, text="Cancelar", bg="#7F8C8D", fg="white", font=("Arial", 8),
                          command=frm.destroy).pack(side="right", padx=10, pady=6)

            def _editar():
                sel = tree.selection()
                if not sel: return
                vals = tree.item(sel[0], "values")
                existing = {"cuil":vals[0],"legajo":vals[1],"apellido":vals[2],"nombre":vals[3],"funcion":vals[4],"aduana":vals[5],"lugar_operativo":vals[6] if len(vals)>6 else ""}
                _abrir_form(existing)

            def _usar():
                sel = tree.selection()
                if not sel: return
                cur_adu_cod = ""
                cur_lugar = self.get_var("car_lugar", "").get()
                if cur_lugar and " - " in cur_lugar:
                    cur_adu_cod = cur_lugar.split(" - ")[0].strip().zfill(3)
                elif cur_lugar:
                    cur_adu_cod = cur_lugar.strip().zfill(3)
                advertencias = []
                for item in sel:
                    vals = tree.item(item, "values")
                    data = {"cuil":vals[0],"legajo":vals[1],"apellido":vals[2],"nombre":vals[3],"funcion":vals[4],"aduana":vals[5],"lugar_operativo":vals[6] if len(vals)>6 else ""}
                    # Validación aduanero: debe pertenecer a la aduana de la medición
                    if data["funcion"] == "ADUANERO" and cur_adu_cod:
                        func_adu_cod = data["aduana"].split(" - ")[0].strip().zfill(3) if " - " in data["aduana"] else data["aduana"].strip().zfill(3)
                        if func_adu_cod and func_adu_cod != cur_adu_cod:
                            advertencias.append(f"{data['apellido']}, {data['nombre']} (Aduana {data['aduana']})")
                if advertencias:
                    msg = ("⚠️ ADVERTENCIA: Los siguientes ADUANEROS pertenecen a una aduana diferente a la de la medición actual "
                           f"(Aduana {cur_adu_cod}):\n\n" + "\n".join(advertencias) +
                           "\n\n¿Desea agregarlos de todas formas?")
                    if not messagebox.askyesno("Conflicto de Aduana", msg, parent=top):
                        return
                for item in sel:
                    vals = tree.item(item, "values")
                    data = {"cuil":vals[0],"legajo":vals[1],"apellido":vals[2],"nombre":vals[3],"funcion":vals[4],"aduana":vals[5],"lugar_operativo":vals[6] if len(vals)>6 else ""}
                    self.agregar_funcionario_row(data=data)
                messagebox.showinfo("OK", f"{len(sel)} funcionario(s) agregado(s) a la operación.", parent=top)

            def _delete():
                sel = tree.selection()
                if not sel: return
                if not messagebox.askyesno("Confirmar", f"Eliminar {len(sel)} registro(s) seleccionado(s) de la base de datos?", parent=top): return
                for item in sel:
                    vals = tree.item(item, "values")
                    db_eliminar_funcionario(vals[0], vals[1], vals[4] if len(vals) > 4 else None)
                _load(v_search.get())

            def _importar_csv():
                """Importar funcionarios desde CSV."""
                f_path = filedialog.askopenfilename(
                    title="Importar Funcionarios desde CSV",
                    filetypes=[("CSV","*.csv"),("Texto","*.txt"),("Todos","*.*")],
                    parent=top)
                if not f_path: return
                try:
                    import csv as csv_mod
                    importados = 0; errores = 0
                    with open(f_path, "r", encoding="utf-8-sig") as f_csv:
                        # Detectar si tiene encabezado
                        sample = f_csv.read(512); f_csv.seek(0)
                        has_header = any(kw in sample.lower() for kw in ["cuil","apellido","legajo","nombre","funcion"])
                        reader = csv_mod.reader(f_csv)
                        if has_header:
                            header_row = next(reader, None)
                            # Map columns by name if header exists
                            if header_row:
                                hdr = [h.strip().lower().replace("ó","o").replace("ú","u").replace("í","i") for h in header_row]
                                def _idx(names):
                                    for nm in names:
                                        if nm in hdr: return hdr.index(nm)
                                    return -1
                                ci = {"cuil":_idx(["cuil"]),"legajo":_idx(["legajo","leg"]),"apellido":_idx(["apellido","apell"]),"nombre":_idx(["nombre","nom"]),"funcion":_idx(["funcion","func"]),"aduana":_idx(["aduana"]),"lugar_operativo":_idx(["lugar_operativo","lugar","lop"])}
                            else:
                                ci = {"cuil":0,"legajo":1,"apellido":2,"nombre":3,"funcion":4,"aduana":5,"lugar_operativo":6}
                        else:
                            ci = {"cuil":0,"legajo":1,"apellido":2,"nombre":3,"funcion":4,"aduana":5,"lugar_operativo":6}
                        for row in reader:
                            if not row or not any(row): continue
                            try:
                                def _g(key, default=""):
                                    idx = ci.get(key, -1)
                                    if idx < 0 or idx >= len(row): return default
                                    return row[idx].strip()
                                cuil = _g("cuil","").replace("-","").replace(".","").replace(" ","")
                                legajo = _g("legajo",""); apellido = _g("apellido",""); nombre = _g("nombre","")
                                funcion = _g("funcion","Verificador"); aduana = _g("aduana",""); lugar_operativo = _g("lugar_operativo","")
                                if len(cuil) == 11 and cuil.isdigit() and apellido:
                                    db_guardar_funcionario(cuil, legajo, apellido, nombre, funcion, aduana, lugar_operativo)
                                    importados += 1
                                else:
                                    errores += 1
                            except: errores += 1
                    _load(v_search.get())
                    messagebox.showinfo("Importacion CSV", f"Importados: {importados} funcionarios.\nFilas con error / saltadas: {errores}.", parent=top)
                except Exception as ex:
                    messagebox.showerror("Error", f"No se pudo leer el archivo:\n{ex}", parent=top)

            _load()
            e_search.focus_set()
        except Exception as e:
            traceback.print_exc()
            messagebox.showerror("Error", str(e))




    def abrir_gestor_funciones(self):
        """ABM simple para gestionar las funciones de funcionarios."""
        top = tk.Toplevel(self.root)
        top.title("Gestionar Funciones")
        top.geometry("400x420")
        top.resizable(False, False)
        top.grab_set()

        fh = tk.Frame(top, bg="#0D47A1"); fh.pack(fill="x")
        tk.Label(fh, text="FUNCIONES DE FUNCIONARIOS", bg="#0D47A1", fg="white",
                 font=("Arial", 10, "bold")).pack(padx=12, pady=8)

        frame_list = tk.Frame(top)
        frame_list.pack(fill="both", expand=True, padx=12, pady=8)

        scrollbar = tk.Scrollbar(frame_list)
        scrollbar.pack(side="right", fill="y")
        listbox = tk.Listbox(frame_list, font=("Arial", 10), yscrollcommand=scrollbar.set)
        listbox.pack(fill="both", expand=True)
        scrollbar.config(command=listbox.yview)

        def _refresh():
            listbox.delete(0, tk.END)
            for f in db_get_funciones():
                listbox.insert(tk.END, f)

        _refresh()

        frame_add = tk.Frame(top)
        frame_add.pack(fill="x", padx=12, pady=4)
        v_nueva = tk.StringVar()
        tk.Entry(frame_add, textvariable=v_nueva, width=30, font=("Arial", 10)).pack(side="left", padx=(0, 6))

        def _agregar():
            nombre = v_nueva.get().strip().upper()
            if not nombre:
                messagebox.showwarning("Aviso", "Ingrese un nombre de funcion.", parent=top)
                return
            if db_guardar_funcion(nombre):
                v_nueva.set("")
                _refresh()
            else:
                messagebox.showerror("Error", "No se pudo agregar la funcion.", parent=top)

        tk.Button(frame_add, text="Agregar", bg="#27AE60", fg="white",
                  font=("Arial", 9, "bold"), command=_agregar).pack(side="left")

        def _eliminar():
            sel = listbox.curselection()
            if not sel:
                messagebox.showwarning("Aviso", "Seleccione una funcion para eliminar.", parent=top)
                return
            nombre = listbox.get(sel[0])
            if messagebox.askyesno("Confirmar", f"Eliminar la funcion '{nombre}'?", parent=top):
                db_eliminar_funcion(nombre)
                _refresh()

        frame_btns = tk.Frame(top)
        frame_btns.pack(fill="x", padx=12, pady=(0, 10))
        tk.Button(frame_btns, text="Eliminar seleccionada", bg="#E74C3C", fg="white",
                  font=("Arial", 9, "bold"), command=_eliminar).pack(side="left")
        tk.Button(frame_btns, text="Cerrar", font=("Arial", 9),
                  command=top.destroy).pack(side="right")

    def abrir_gestor_aduanas_lop(self):
        """ABM completo de Aduanas y Lugares Operativos con SQLite."""
        try:
            top = tk.Toplevel(self.root)
            top.title("Gestión de Aduanas y Lugares Operativos")
            top.geometry("1100x680")
            top.resizable(True, True)

            # ── Header ──────────────────────────────────────────────────────────
            fh = tk.Frame(top, bg="#1B3A5C")
            fh.pack(fill="x")
            tk.Label(fh, text="GESTIÓN DE ADUANAS Y LUGARES OPERATIVOS",
                     bg="#1B3A5C", fg="white", font=("Arial", 11, "bold")).pack(side="left", padx=14, pady=10)
            tk.Label(fh, text="Administre las Aduanas habilitadas y sus Lugares Operativos asociados.",
                     bg="#1B3A5C", fg="#AED6F1", font=("Arial", 8)).pack(side="left", padx=4)

            # ── Notebook con dos pestañas ────────────────────────────────────────
            nb = ttk.Notebook(top)
            nb.pack(fill="both", expand=True, padx=6, pady=6)

            # ═══════════════════════════════════════════════════════════════════
            # PESTAÑA 1: ADUANAS
            # ═══════════════════════════════════════════════════════════════════
            tab_adu = tk.Frame(nb, bg="#F4F6F7")
            nb.add(tab_adu, text="  ADUANAS  ")

            # Toolbar aduanas
            ftb_a = tk.Frame(tab_adu, bg="#EBF5FB", bd=1, relief="flat")
            ftb_a.pack(fill="x", padx=0, pady=0)
            tk.Label(ftb_a, text="Filtrar:", bg="#EBF5FB", font=("Arial", 8, "bold")).pack(side="left", padx=8, pady=6)
            v_search_a = tk.StringVar()
            tk.Entry(ftb_a, textvariable=v_search_a, width=22, font=("Arial", 9)).pack(side="left", padx=4, pady=6)
            tk.Button(ftb_a, text="+ AGREGAR ADUANA", bg="#1A5276", fg="white",
                      font=("Arial", 8, "bold"), command=lambda: _abrir_form_aduana()).pack(side="left", padx=8, pady=6, ipadx=6, ipady=2)
            tk.Button(ftb_a, text="Editar", bg="#2980B9", fg="white",
                      font=("Arial", 8, "bold"), command=lambda: _editar_aduana()).pack(side="left", padx=4, pady=6, ipadx=6, ipady=2)
            tk.Button(ftb_a, text="Eliminar", bg="#E74C3C", fg="white",
                      font=("Arial", 8, "bold"), command=lambda: _del_aduana()).pack(side="left", padx=4, pady=6, ipadx=6, ipady=2)
            lbl_cnt_a = tk.Label(ftb_a, text="", bg="#EBF5FB", font=("Arial", 8), fg="#1B3A5C")
            lbl_cnt_a.pack(side="right", padx=12)

            # Treeview aduanas
            f_tr_a = tk.Frame(tab_adu)
            f_tr_a.pack(fill="both", expand=True, padx=8, pady=4)
            cols_a = ("Código", "Nombre")
            tr_a = ttk.Treeview(f_tr_a, columns=cols_a, show="headings", height=18, selectmode="extended")
            tr_a.heading("Código", text="Código", command=lambda: _sort_a("Código"))
            tr_a.heading("Nombre", text="Nombre de la Aduana", command=lambda: _sort_a("Nombre"))
            tr_a.column("Código", width=90); tr_a.column("Nombre", width=500)
            vsb_a = ttk.Scrollbar(f_tr_a, orient="vertical", command=tr_a.yview)
            tr_a.configure(yscrollcommand=vsb_a.set)
            tr_a.grid(row=0, column=0, sticky="nsew"); vsb_a.grid(row=0, column=1, sticky="ns")
            f_tr_a.grid_rowconfigure(0, weight=1); f_tr_a.grid_columnconfigure(0, weight=1)
            tr_a.bind("<Double-1>", lambda e: _editar_aduana())
            _sort_col_a = [None]; _sort_rev_a = [False]

            def _load_aduanas(q=""):
                tr_a.delete(*tr_a.get_children())
                rows = db_get_aduanas()
                q_lo = q.lower()
                if q_lo:
                    rows = [r for r in rows if q_lo in r["codigo"] or q_lo in r["nombre"].lower()]
                for r in rows:
                    tr_a.insert("", "end", values=(r["codigo"], r["nombre"]))
                lbl_cnt_a.config(text=f"{len(rows)} aduana(s)")

            def _sort_a(col):
                items = [(tr_a.set(iid, col), iid) for iid in tr_a.get_children("")]
                rev = _sort_col_a[0] == col and not _sort_rev_a[0]
                items.sort(reverse=rev)
                for i, (_, iid) in enumerate(items): tr_a.move(iid, "", i)
                _sort_col_a[0] = col; _sort_rev_a[0] = rev

            v_search_a.trace_add("write", lambda *a: _load_aduanas(v_search_a.get()))

            def _abrir_form_aduana(existing=None):
                frm = tk.Toplevel(top); frm.title("Nueva Aduana" if not existing else "Editar Aduana")
                frm.geometry("400x200"); frm.resizable(False, False); frm.grab_set()
                tk.Frame(frm, bg="#1B3A5C", height=40).pack(fill="x")
                tk.Label(frm, text="DATOS DE LA ADUANA", bg="#1B3A5C", fg="white",
                         font=("Arial", 9, "bold")).place(x=0, y=8, width=400)
                ff = tk.Frame(frm, pady=8); ff.pack(fill="both", expand=True, padx=16, pady=(50,0))
                v_cod  = tk.StringVar(value=existing["codigo"] if existing else "")
                v_nom  = tk.StringVar(value=existing["nombre"] if existing else "")
                tk.Label(ff, text="Código (3 dígitos):", font=("Arial", 8, "bold"), anchor="e", width=18).grid(row=0, column=0, sticky="e", pady=6)
                e_cod = tk.Entry(ff, textvariable=v_cod, width=10, font=("Arial", 9))
                e_cod.grid(row=0, column=1, sticky="w", padx=6)
                tk.Label(ff, text="Nombre:", font=("Arial", 8, "bold"), anchor="e", width=18).grid(row=1, column=0, sticky="e", pady=6)
                tk.Entry(ff, textvariable=v_nom, width=30, font=("Arial", 9)).grid(row=1, column=1, sticky="w", padx=6)
                def _guardar():
                    cod = v_cod.get().strip()
                    nom = v_nom.get().strip()
                    if not cod or not nom:
                        messagebox.showwarning("Error", "Código y Nombre son obligatorios.", parent=frm); return
                    if not cod.isdigit():
                        messagebox.showwarning("Error", "El código debe ser numérico.", parent=frm); return
                    new_cod = cod.zfill(3)
                    # If editing and codigo changed, delete the old one first
                    if existing and existing["codigo"] != new_cod:
                        db_eliminar_aduana(existing["codigo"])
                    db_guardar_aduana(new_cod, nom)
                    _load_aduanas(v_search_a.get())
                    _load_lop()   # actualizar también el combo de aduanas en LOP
                    frm.destroy()
                fb = tk.Frame(frm, bg="#2C3E50"); fb.pack(fill="x", side="bottom")
                tk.Button(fb, text="Guardar", bg="#27AE60", fg="white", font=("Arial", 8, "bold"),
                          command=_guardar).pack(side="left", padx=10, pady=6, ipadx=10, ipady=2)
                tk.Button(fb, text="Cancelar", bg="#7F8C8D", fg="white", font=("Arial", 8),
                          command=frm.destroy).pack(side="right", padx=10, pady=6)
                if not existing: e_cod.focus_set()

            def _editar_aduana():
                sel = tr_a.selection()
                if not sel: return
                vals = tr_a.item(sel[0], "values")
                _abrir_form_aduana({"codigo": vals[0], "nombre": vals[1]})

            def _del_aduana():
                sel = tr_a.selection()
                if not sel: return
                nombres = [tr_a.item(s, "values")[1] for s in sel]
                if not messagebox.askyesno("Confirmar",
                    f"Eliminar {len(sel)} aduana(s)?\n{', '.join(nombres)}\n\n"
                    "ATENCIÓN: También se eliminarán todos sus Lugares Operativos.", parent=top): return
                for s in sel:
                    vals = tr_a.item(s, "values")
                    db_eliminar_aduana(vals[0])
                _load_aduanas(v_search_a.get()); _load_lop()

            # ═══════════════════════════════════════════════════════════════════
            # PESTAÑA 2: LUGARES OPERATIVOS
            # ═══════════════════════════════════════════════════════════════════
            tab_lop = tk.Frame(nb, bg="#F4F6F7")
            nb.add(tab_lop, text="  LUGARES OPERATIVOS (LOT)  ")

            # Toolbar LOP
            ftb_l = tk.Frame(tab_lop, bg="#EBF5FB", bd=1, relief="flat")
            ftb_l.pack(fill="x", padx=0, pady=0)
            tk.Label(ftb_l, text="Filtrar por Aduana:", bg="#EBF5FB", font=("Arial", 8, "bold")).pack(side="left", padx=8, pady=6)
            v_adu_filter = tk.StringVar(value="TODAS")
            cb_adu_filter = ttk.Combobox(ftb_l, textvariable=v_adu_filter,
                                         values=["TODAS"], state="readonly", width=30, font=("Arial", 8))
            cb_adu_filter.pack(side="left", padx=4, pady=6)
            tk.Label(ftb_l, text="Buscar:", bg="#EBF5FB", font=("Arial", 8, "bold")).pack(side="left", padx=(10,2), pady=6)
            v_search_l = tk.StringVar()
            tk.Entry(ftb_l, textvariable=v_search_l, width=18, font=("Arial", 9)).pack(side="left", padx=4, pady=6)
            tk.Button(ftb_l, text="+ AGREGAR LOT", bg="#1D6A39", fg="white",
                      font=("Arial", 8, "bold"), command=lambda: _abrir_form_lop()).pack(side="left", padx=8, pady=6, ipadx=6, ipady=2)
            tk.Button(ftb_l, text="Editar", bg="#2980B9", fg="white",
                      font=("Arial", 8, "bold"), command=lambda: _editar_lop()).pack(side="left", padx=4, pady=6, ipadx=6, ipady=2)
            tk.Button(ftb_l, text="Eliminar", bg="#E74C3C", fg="white",
                      font=("Arial", 8, "bold"), command=lambda: _del_lop()).pack(side="left", padx=4, pady=6, ipadx=6, ipady=2)
            lbl_cnt_l = tk.Label(ftb_l, text="", bg="#EBF5FB", font=("Arial", 8), fg="#1B3A5C")
            lbl_cnt_l.pack(side="right", padx=12)

            # Treeview LOP
            f_tr_l = tk.Frame(tab_lop)
            f_tr_l.pack(fill="both", expand=True, padx=8, pady=4)
            cols_l = ("Código LOT", "Descripción", "Aduana")
            tr_l = ttk.Treeview(f_tr_l, columns=cols_l, show="headings", height=18, selectmode="extended")
            tr_l.heading("Código LOT",  text="Código LOT",   command=lambda: _sort_l("Código LOT"))
            tr_l.heading("Descripción", text="Descripción",  command=lambda: _sort_l("Descripción"))
            tr_l.heading("Aduana",      text="Aduana",       command=lambda: _sort_l("Aduana"))
            tr_l.column("Código LOT", width=100); tr_l.column("Descripción", width=380); tr_l.column("Aduana", width=200)
            vsb_l = ttk.Scrollbar(f_tr_l, orient="vertical", command=tr_l.yview)
            tr_l.configure(yscrollcommand=vsb_l.set)
            tr_l.grid(row=0, column=0, sticky="nsew"); vsb_l.grid(row=0, column=1, sticky="ns")
            f_tr_l.grid_rowconfigure(0, weight=1); f_tr_l.grid_columnconfigure(0, weight=1)
            tr_l.bind("<Double-1>", lambda e: _editar_lop())
            _sort_col_l = [None]; _sort_rev_l = [False]

            def _get_adu_opts():
                ads = db_get_aduanas()
                return ["TODAS"] + [f"{a['codigo']} - {a['nombre']}" for a in ads]

            def _load_lop(q="", adu_filter=None):
                tr_l.delete(*tr_l.get_children())
                # Actualizar combo aduanas
                opts = _get_adu_opts()
                cb_adu_filter["values"] = opts
                if v_adu_filter.get() not in opts: v_adu_filter.set("TODAS")
                # Filtro por aduana
                filt_adu = v_adu_filter.get()
                cod_adu = None
                if filt_adu and filt_adu != "TODAS" and " - " in filt_adu:
                    cod_adu = filt_adu.split(" - ", 1)[0]
                rows = db_get_lugares_operativos(cod_adu)
                q_lo = (q or v_search_l.get()).lower()
                if q_lo:
                    rows = [r for r in rows if q_lo in r["codigo"].lower() or q_lo in r["descripcion"].lower()]
                ads_map = {a["codigo"]: a["nombre"] for a in db_get_aduanas()}
                for r in rows:
                    adu_str = f"{r['aduana_codigo']} - {ads_map.get(r['aduana_codigo'], r['aduana_codigo'])}"
                    tr_l.insert("", "end", values=(r["codigo"], r["descripcion"], adu_str),
                                tags=(r["aduana_codigo"],))
                lbl_cnt_l.config(text=f"{len(rows)} lugar(es)")

            def _sort_l(col):
                items = [(tr_l.set(iid, col), iid) for iid in tr_l.get_children("")]
                rev = _sort_col_l[0] == col and not _sort_rev_l[0]
                items.sort(reverse=rev)
                for i, (_, iid) in enumerate(items): tr_l.move(iid, "", i)
                _sort_col_l[0] = col; _sort_rev_l[0] = rev

            v_search_l.trace_add("write", lambda *a: _load_lop())
            v_adu_filter.trace_add("write", lambda *a: _load_lop())

            def _abrir_form_lop(existing=None):
                frm = tk.Toplevel(top); frm.title("Nuevo LOT" if not existing else "Editar LOT")
                frm.geometry("460x240"); frm.resizable(False, False); frm.grab_set()
                tk.Frame(frm, bg="#1D6A39", height=40).pack(fill="x")
                tk.Label(frm, text="DATOS DEL LOT (Lugar Operativo de Tráfico)", bg="#1D6A39", fg="white",
                         font=("Arial", 9, "bold")).place(x=0, y=8, width=460)
                ff = tk.Frame(frm, pady=8); ff.pack(fill="both", expand=True, padx=16, pady=(50,0))
                v_cod_l = tk.StringVar(value=existing["codigo"] if existing else "")
                v_desc_l = tk.StringVar(value=existing["descripcion"] if existing else "")
                v_adu_l  = tk.StringVar()
                ads = db_get_aduanas()
                adu_opts = [f"{a['codigo']} - {a['nombre']}" for a in ads]
                # Pre-seleccionar aduana actual o la del existing
                if existing:
                    match = [o for o in adu_opts if o.startswith(existing["aduana_codigo"])]
                    v_adu_l.set(match[0] if match else (adu_opts[0] if adu_opts else ""))
                else:
                    cur_adu = self.get_var("car_lugar").get()
                    match = [o for o in adu_opts if cur_adu.split(" - ")[0].zfill(3) in o]
                    v_adu_l.set(match[0] if match else (adu_opts[0] if adu_opts else ""))
                tk.Label(ff, text="Código LOT:", font=("Arial", 8, "bold"), anchor="e", width=16).grid(row=0, column=0, sticky="e", pady=5)
                e_cod_l = tk.Entry(ff, textvariable=v_cod_l, width=12, font=("Arial", 9))
                e_cod_l.grid(row=0, column=1, sticky="w", padx=6)
                tk.Label(ff, text="Descripción:", font=("Arial", 8, "bold"), anchor="e", width=16).grid(row=1, column=0, sticky="e", pady=5)
                tk.Entry(ff, textvariable=v_desc_l, width=32, font=("Arial", 9)).grid(row=1, column=1, sticky="w", padx=6)
                tk.Label(ff, text="Aduana:", font=("Arial", 8, "bold"), anchor="e", width=16).grid(row=2, column=0, sticky="e", pady=5)
                ttk.Combobox(ff, textvariable=v_adu_l, values=adu_opts, state="readonly", width=30, font=("Arial", 8)).grid(row=2, column=1, sticky="w", padx=6)
                def _guardar_lop():
                    cod  = v_cod_l.get().strip()
                    desc = v_desc_l.get().strip()
                    adu  = v_adu_l.get().strip()
                    if not cod or not desc:
                        messagebox.showwarning("Error", "Código y Descripción son obligatorios.", parent=frm); return
                    if not adu or " - " not in adu:
                        messagebox.showwarning("Error", "Seleccione una Aduana.", parent=frm); return
                    adu_cod = adu.split(" - ", 1)[0]
                    # If editing and the key changed, delete the old record first
                    if existing:
                        old_cod = existing["codigo"]
                        old_adu = existing["aduana_codigo"]
                        if old_cod != cod or old_adu != adu_cod:
                            db_eliminar_lugar_operativo(old_cod, old_adu)
                    db_guardar_lugar_operativo(cod, desc, adu_cod)
                    _load_lop(); frm.destroy()
                fb = tk.Frame(frm, bg="#2C3E50"); fb.pack(fill="x", side="bottom")
                tk.Button(fb, text="Guardar", bg="#27AE60", fg="white", font=("Arial", 8, "bold"),
                          command=_guardar_lop).pack(side="left", padx=10, pady=6, ipadx=10, ipady=2)
                tk.Button(fb, text="Cancelar", bg="#7F8C8D", fg="white", font=("Arial", 8),
                          command=frm.destroy).pack(side="right", padx=10, pady=6)
                if not existing: e_cod_l.focus_set()

            def _editar_lop():
                sel = tr_l.selection()
                if not sel: return
                vals = tr_l.item(sel[0], "values")
                adu_cod = vals[2].split(" - ")[0] if " - " in vals[2] else vals[2]
                _abrir_form_lop({"codigo": vals[0], "descripcion": vals[1], "aduana_codigo": adu_cod})

            def _del_lop():
                sel = tr_l.selection()
                if not sel: return
                if not messagebox.askyesno("Confirmar", f"Eliminar {len(sel)} Lugar(es) Operativo(s)?", parent=top): return
                for s in sel:
                    vals = tr_l.item(s, "values")
                    adu_cod = vals[2].split(" - ")[0] if " - " in vals[2] else vals[2]
                    db_eliminar_lugar_operativo(vals[0], adu_cod)
                _load_lop()

            # ── Hint ────────────────────────────────────────────────────────────
            f_note = tk.Frame(tab_lop, bg="#FEF9E7"); f_note.pack(fill="x", padx=8, pady=2)
            tk.Label(f_note, text="Los Lugares Operativos están coordinados por la Aduana correspondiente. "
                     "Solo se pueden asignar Aduaneros a las Aduanas que figuran en la tabla de Aduanas.",
                     bg="#FEF9E7", fg="#6E4B00", font=("Arial", 7), anchor="w").pack(fill="x", padx=8, pady=3)

            # ── Footer ───────────────────────────────────────────────────────────
            f_bot = tk.Frame(top, bg="#2C3E50")
            f_bot.pack(fill="x", side="bottom")
            tk.Button(f_bot, text="Cerrar", bg="#7F8C8D", fg="white",
                      font=("Arial", 8, "bold"), command=top.destroy).pack(side="right", padx=10, pady=6, ipadx=8, ipady=2)
            tk.Button(f_bot, text="Aplicar aduana actual al LOT",
                      bg="#1D6A39", fg="white", font=("Arial", 8),
                      command=lambda: (nb.select(tab_lop), _load_lop())).pack(side="left", padx=10, pady=6, ipadx=6, ipady=2)

            # Carga inicial
            _load_aduanas(); _load_lop()

        except Exception as e:
            traceback.print_exc()
            messagebox.showerror("Error ABM", str(e))

    def borrar_funcionario(self, obj):
        if obj in self.funcionarios_data:
            self.funcionarios_data.remove(obj)
        if obj["frame"]:
            obj["frame"].destroy()

    def borrar_ddt(self, obj):
        if len(self.ddt_data) <= 1:
            messagebox.showwarning("Aviso", "Debe quedar al menos un Documento.")
            return
        if messagebox.askyesno("Confirmar", "¿Eliminar este Documento?"):
            obj["main_frame"].destroy()
            if obj in self.ddt_data: self.ddt_data.remove(obj)
            self.actualizar_combos_ddt()

    def agregar_salida(self, ddt_obj, data=None):
        salida = {
            "numero": tk.StringVar(value=data["numero"] if data else ""),
            "litros": tk.StringVar(value=data["litros"] if data else ""),
            "kilos": tk.StringVar(value=data["kilos"] if data else ""),
            "densidad": tk.StringVar(value=data["densidad"] if data else ""),
            "frame": None
        }
        salida["litros"].trace_add("write", lambda *args: self.auto_calc_densidad(salida["litros"], salida["kilos"], salida["densidad"]))
        salida["kilos"].trace_add("write", lambda *args: self.auto_calc_densidad(salida["litros"], salida["kilos"], salida["densidad"]))
        f_row = ttk.Frame(ddt_obj["salidas_frame"])
        f_row.pack(fill="x", pady=2)
        salida["frame"] = f_row
        ttk.Entry(f_row, textvariable=salida["numero"], width=20).pack(side="left", padx=2)
        ttk.Entry(f_row, textvariable=salida["litros"], width=12).pack(side="left", padx=2)
        ttk.Entry(f_row, textvariable=salida["kilos"], width=12).pack(side="left", padx=2)
        ttk.Entry(f_row, textvariable=salida["densidad"], width=12, state="readonly").pack(side="left", padx=2)
        tk.Button(f_row, text="x", bg="#ffcccc", width=2, command=lambda: self.borrar_salida(ddt_obj, salida)).pack(side="left", padx=5)
        ddt_obj["salidas"].append(salida)

    def borrar_salida(self, ddt_obj, salida):
        salida["frame"].destroy()
        if salida in ddt_obj["salidas"]: ddt_obj["salidas"].remove(salida)

    def actualizar_combos_ddt(self, event=None):
        lista = [d["numero"].get() for d in self.ddt_data if d["numero"].get()]
        lista.insert(0, "") 
        for cb in self.combos_ddt: 
            try: cb['values'] = lista
            except: pass

    def check_combos(self, event):
        self.actualizar_combos_ddt()

    # --- UI TABLA PRINCIPAL UNIFICADA ---
    def construir_pantalla_unificada(self, etapa, parent):
        main_frame = ttk.Frame(parent)
        main_frame.pack(fill="both", expand=True)

        canvas = tk.Canvas(main_frame, bg="white")
        vbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        hbar = ttk.Scrollbar(main_frame, orient="horizontal", command=canvas.xview)
        scroll_f = ttk.Frame(canvas)
        scroll_f.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scroll_f, anchor="nw")
        canvas.configure(yscrollcommand=vbar.set, xscrollcommand=hbar.set)
        canvas.grid(row=0, column=0, sticky="nsew")
        def _pu_mw(event): canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        def _pu_up(event): canvas.yview_scroll(-3, "units")
        def _pu_down(event): canvas.yview_scroll(3, "units")
        def _pu_bind(e):
            canvas.bind_all("<MouseWheel>", _pu_mw)
            canvas.bind_all("<Button-4>", _pu_up)
            canvas.bind_all("<Button-5>", _pu_down)
        def _pu_unbind(e):
            canvas.unbind_all("<MouseWheel>")
            canvas.unbind_all("<Button-4>")
            canvas.unbind_all("<Button-5>")
        canvas.bind("<Enter>", _pu_bind)
        canvas.bind("<Leave>", _pu_unbind)
        vbar.grid(row=0, column=1, sticky="ns")
        hbar.grid(row=1, column=0, sticky="ew")
        main_frame.grid_rowconfigure(0, weight=1)
        main_frame.grid_columnconfigure(0, weight=1)

        # Header Global
        f_top = ttk.Frame(scroll_f)
        f_top.pack(fill="x", padx=10, pady=10)
        font_style = ("Arial", 10)
        lbl_font = ("Arial", 8, "bold")

        # FILA 1: FECHA Y HORA
        f_row1 = tk.Frame(f_top)
        f_row1.pack(fill="x", pady=2)
        
        # Registrar validaciones
        vcmd_fecha = (f_row1.register(self.validar_fecha), '%P')
        vcmd_hora = (f_row1.register(self.validar_hora), '%P')
        
        ttk.Label(f_row1, text="Fecha:", font=lbl_font).pack(side="left", padx=5)
        var_fecha = self.get_var(f"{etapa}_Fecha")
        e_fecha = ttk.Entry(f_row1, textvariable=var_fecha, width=12, font=font_style, validate="key", validatecommand=vcmd_fecha)
        e_fecha.pack(side="left")
        e_fecha.bind("<FocusIn>", lambda ev, v=var_fecha: self.on_fecha_focus_in(ev, v))
        e_fecha.bind("<FocusOut>", lambda ev, v=var_fecha: self.on_fecha_focus_out(ev, v))
        if not var_fecha.get(): var_fecha.set("DD/MM/AAAA")

        ttk.Label(f_row1, text="Hora:", font=lbl_font).pack(side="left", padx=5)
        var_hora = self.get_var(f"{etapa}_Hora")
        e_hora = ttk.Entry(f_row1, textvariable=var_hora, width=8, font=font_style, validate="key", validatecommand=vcmd_hora)
        e_hora.pack(side="left")
        e_hora.bind("<FocusIn>", lambda ev, v=var_hora: self.on_hora_focus_in(ev, v))
        e_hora.bind("<FocusOut>", lambda ev, v=var_hora: self.on_hora_focus_out(ev, v))
        if not var_hora.get(): var_hora.set("00:00")

        # FILA 2: LONGITUDINAL (solo maritimos) / temperatura ambiente (tierra/camion)
        _tm_pu = self.get_tipo_medio()
        _mar_pu = self.es_maritimo()

        if _mar_pu:
            ttk.Label(f_row1, text="Proa:", font=lbl_font).pack(side="left", padx=5)
            ttk.Entry(f_row1, textvariable=self.get_var(f"{etapa}_Calados Proa"), width=8, font=font_style).pack(side="left")
            ttk.Label(f_row1, text="Popa:", font=lbl_font).pack(side="left", padx=5)
            ttk.Entry(f_row1, textvariable=self.get_var(f"{etapa}_Calados Popa"), width=8, font=font_style).pack(side="left")
            ttk.Label(f_row1, text="Trim:", font=lbl_font).pack(side="left", padx=5)
            tk.Label(f_row1, textvariable=self.get_var(f"{etapa}_Trimación"), width=8, bg="#e0e0e0", relief="sunken", font=font_style).pack(side="left")
            ttk.Label(f_row1, text="Babor:", font=lbl_font).pack(side="left", padx=5)
            ttk.Entry(f_row1, textvariable=self.get_var(f"{etapa}_Calados Babor"), width=8, font=font_style).pack(side="left")
            ttk.Label(f_row1, text="Estribor:", font=lbl_font).pack(side="left", padx=5)
            ttk.Entry(f_row1, textvariable=self.get_var(f"{etapa}_Calados Estribor"), width=8, font=font_style).pack(side="left")
            ttk.Label(f_row1, text="Escora:", font=lbl_font).pack(side="left", padx=5)
            tk.Label(f_row1, textvariable=self.get_var(f"{etapa}_Lista"), width=8, bg="#e0e0e0", relief="sunken", font=font_style).pack(side="left")
            self.get_var(f"{etapa}_Calados Popa").trace_add("write", lambda *args, e=etapa: self.calc_trim(e))
            self.get_var(f"{etapa}_Calados Proa").trace_add("write", lambda *args, e=etapa: self.calc_trim(e))
            self.get_var(f"{etapa}_Calados Babor").trace_add("write", lambda *args, e=etapa: self.calc_trim(e))
            self.get_var(f"{etapa}_Calados Estribor").trace_add("write", lambda *args, e=etapa: self.calc_trim(e))
        else:
            # Para tierra y camion: temperatura ambiente + modo sondaje
            ttk.Label(f_row1, text="Temp.Ambiente (°C):", font=lbl_font).pack(side="left", padx=5)
            ttk.Entry(f_row1, textvariable=self.get_var(f"{etapa}_Temp_Amb",""), width=8, font=font_style).pack(side="left")
            if self.es_tierra():
                ttk.Label(f_row1, text="Modo Sondaje:", font=lbl_font).pack(side="left", padx=5)
                _modo_var = self.get_var(f"{etapa}_Modo_Sondaje","INNAGE")
                cb_modo = ttk.Combobox(f_row1, textvariable=_modo_var,
                    values=["INNAGE (desde el fondo)","ULLAGE (espacio libre)"],
                    state="readonly", width=22, font=font_style)
                cb_modo.pack(side="left", padx=3)
                if "FLOTANTE" in _tm_pu:
                    ttk.Label(f_row1, text="Offset Techo (mm):", font=lbl_font).pack(side="left", padx=5)
                    ttk.Entry(f_row1, textvariable=self.get_var("car_tf_offset","0"),
                              width=8, font=font_style).pack(side="left")
            if self.es_ducto():
                ttk.Label(f_row1, text="P.Linea(kPa):", font=lbl_font).pack(side="left", padx=5)
                ttk.Entry(f_row1, textvariable=self.get_var(f"{etapa}_P_linea",""), width=8, font=font_style).pack(side="left")
                ttk.Label(f_row1, text="T.Linea(C):", font=lbl_font).pack(side="left", padx=5)
                ttk.Entry(f_row1, textvariable=self.get_var(f"{etapa}_T_linea",""), width=8, font=font_style).pack(side="left")
                ttk.Label(f_row1, text="Cond.base:", font=lbl_font).pack(side="left", padx=5)
                ttk.Combobox(f_row1, textvariable=self.get_var(f"{etapa}_Cond_base","15C/101.325kPa"),
                    values=["15C/101.325kPa","20C/101.325kPa","0C/101.325kPa"],
                    state="readonly", width=18, font=font_style).pack(side="left", padx=3)

        # --- TABLA EXTENDIDA --- (columnas adaptadas al tipo de medición)
        f_table = ttk.Frame(scroll_f)
        f_table.pack(fill="x", padx=10, pady=5)
        
        table_widgets = []
        
        # ── Determinar qué columnas mostrar según tipo de medio ──────────────
        _tm_tab   = self.get_tipo_medio()
        _is_mar_t = self.es_maritimo()
        _is_tie_t = self.es_tierra()
        _is_cam_t = self.es_camion()
        _is_gas_t = self.es_gasero()
        _is_esf_t = self.es_esfera()
        _is_duc_t = self.es_ducto()
        _is_el_t  = self.es_electrico()
        _is_cgb_t = self.es_camion_gas()
        _is_liq_t = _is_mar_t or _is_tie_t or (_is_cam_t and not _is_cgb_t)  # medición de líquidos clásica
        _is_gas_only = _is_gas_t or _is_esf_t or _is_cgb_t or ("METANERO" in _tm_tab)
        _show_dens_cols = not (_is_duc_t or _is_el_t)
        _show_prod_col  = not _is_el_t
        # Agua: solo para marítimos de líquidos + tierra + camión cisterna
        # Gasero/Metanero son marítimos pero NO usan UTI de agua
        _es_liq_mar_t   = _tm_tab in ("BUQUE", "BARCAZA", "BUQUE QUIMIQUERO", "DRAFT SURVEY")
        _show_agua_col  = _es_liq_mar_t or _is_tie_t or (_is_cam_t and not _is_cgb_t)

        # Encabezados fijos
        headers = [("TANQUE", "#ddd"), ("EDITAR", "#ddd")]
        headers.append(("DOCUMENTO", "#ddd"))
        if _show_prod_col:
            headers.append(("PRODUCTO", "#ddd"))

        if _is_el_t:
            # Electricidad: kWh activa, reactiva, cosfi
            headers += [
                ("kWh ACTIVA", "#d5f5e3"), ("kWh REACTIVA", "#d5f5e3"),
                ("COS FI", "#d5f5e3"), ("DEMANDA (kW)", "#d5f5e3"),
            ]
        elif _is_duc_t:
            # Ducto: volumen linea, volumen base, caudal
            headers += [
                ("VOL.LINEA (m3)", "#d6eaf8"), ("VOL.BASE (m3)", "#d6eaf8"),
                ("VOL.BASE (Km3)", "#d6eaf8"), ("CAUDAL (m3/h)", "#d6eaf8"),
            ]
        elif _show_dens_cols:
            # Líquidos y gas licuado: densidades y volúmenes
            headers += [
                ("DENS. LAB", "#ffffcc"), ("DENS. DOC", "#dcebf7"), ("DENS. SAL", "#dcf7dc"),
                ("VOL. LAB (L)", "#ffffcc"), ("KILOS LAB", "#ffffcc"),
                ("VOL. DOC (L)", "#dcebf7"), ("KILOS DOC", "#dcebf7"),
                ("VOL. SAL (L)", "#dcf7dc"), ("KILOS SAL", "#dcf7dc"),
            ]

        if _show_agua_col:
            headers.append(("AGUA (L)", "#ccffcc"))

        # Columna de volumen natural (siempre al final)
        headers.append(("VOL. NATURAL", "#e8e8e8"))

        for i, (h, bg_c) in enumerate(headers):
            lbl = tk.Label(f_table, text=h, font=("Arial", 8, "bold"), bg=bg_c, relief="raised", padx=5, anchor="center")
            lbl.grid(row=0, column=i, sticky="nsew")
            table_widgets.append(("header", lbl))
            f_table.grid_columnconfigure(i, weight=1)


        # Botón "Agregar" para tipos de tierra/esfera
        _es_tierra_esf = self.es_tierra() or self.es_esfera()
        if _es_tierra_esf:
            _btn_frame = tk.Frame(f_table, bg="#f0f0f0")
            _btn_frame.grid(row=0, column=14, rowspan=200, sticky="ns", padx=4)
            _tipo_str = "ESFERA" if self.es_esfera() else "TANQUE"
            _tank_bg = "#1D6A39" if not self.es_esfera() else "#6A1B9A"
            def _add_tank(tipo=_tipo_str, tab_frame=parent, tab_etapa=etapa, tab_bg=_tank_bg):
                n_existing = len([t for t in self.lista_tanques if tipo.split()[0] in t.upper()])
                self.lista_tanques.append(f"{tipo} {n_existing+1}")
                self.rebuild_all_tabs()
            def _del_tank():
                if len(self.lista_tanques) > 1:
                    self.lista_tanques.pop()
                    self.rebuild_all_tabs()
            tk.Button(_btn_frame, text=f"+ {_tipo_str}", bg=_tank_bg, fg="white",
                      font=("Arial", 7, "bold"), relief="flat", cursor="hand2",
                      command=_add_tank).pack(fill="x", padx=2, pady=(6,2), ipadx=3, ipady=4)
            tk.Button(_btn_frame, text="− Quitar", bg="#C0392B", fg="white",
                      font=("Arial", 7), relief="flat", cursor="hand2",
                      command=_del_tank).pack(fill="x", padx=2, pady=2, ipadx=3, ipady=3)

        # FILAS
        total_list = self.lista_tanques + self.lista_carbonera
        for idx, tk_name in enumerate(total_list):
            r = idx + 1
            bg_tk = "#fff9c4" if tk_name in self.lista_carbonera else "white"
            lbl_tk = tk.Label(f_table, text=tk_name, font=("Arial", 8, "bold"), bg=bg_tk, relief="solid")
            lbl_tk.grid(row=r, column=0, sticky="nsew")
            table_widgets.append(("name", lbl_tk))
            
            # Boton Lapiz (la interp × asiento se carga desde la ficha del tanque)
            f_btns = tk.Frame(f_table, bg="#eee")
            f_btns.grid(row=r, column=1, sticky="nsew", padx=1, pady=1)
            btn_edit = tk.Button(f_btns, text="Ed.", bg="#eee", font=("Arial", 10), command=lambda t=tk_name, e=etapa: self.abrir_popup_detalle(e, t))
            btn_edit.pack(side="left", fill="both", expand=True)
            table_widgets.append(("btn", f_btns))
            table_widgets.append(("btn", btn_edit))

            # Campos Resumen — DINÁMICOS según tipo (deben coincidir EXACTAMENTE con los headers)
            if _is_el_t:
                # Electricidad: documento, kWh activa, reactiva, cos fi, demanda, vol nat
                vars_resumen = [
                    ("ddt_assign",   30, "white"),
                    ("el_kwh_act",   12, "#d5f5e3"),
                    ("el_kwh_rea",   12, "#d5f5e3"),
                    ("el_fp",        10, "#d5f5e3"),
                    ("el_dem",       10, "#d5f5e3"),
                    ("vol_nat_prod", 12, "#e8e8e8"),
                ]
            elif _is_duc_t:
                # Ductos: documento, producto, vol linea, vol base, vol base km3, caudal, vol nat
                vars_resumen = [
                    ("ddt_assign",    30, "white"),
                    ("prod_name",     20, "white"),
                    ("vol_linea",     12, "#d6eaf8"),
                    ("vol_base",      12, "#d6eaf8"),
                    ("vol_base_km3",  12, "#d6eaf8"),
                    ("caudal_mh",     10, "#d6eaf8"),
                    ("vol_nat_prod",  12, "#e8e8e8"),
                ]
            else:
                # Líquidos y gas: documento, producto, densidades, volúmenes, [agua], vol nat
                vars_resumen = [
                    ("ddt_assign",  30, "white"),
                    ("prod_name",   24, "white"),
                    ("dens_lab",    10, "#ffffcc"), ("dens_doc", 10, "#dcebf7"), ("dens_salida", 10, "#dcf7dc"),
                    ("v15_lab",     10, "#ffffcc"), ("kv_lab",   10, "#ffffcc"),
                    ("v15_doc",     10, "#dcebf7"), ("kv_doc",   10, "#dcebf7"),
                    ("v15_sal",     10, "#dcf7dc"), ("kv_sal",   10, "#dcf7dc"),
                ]
                if _show_agua_col:
                    vars_resumen.append(("vol_nat_agua", 10, "#ccffcc"))
                vars_resumen.append(("vol_nat_prod", 12, "#e8e8e8"))
            
            for c_idx, (var_key, w_size, bg_c) in enumerate(vars_resumen):
                val_var = self.get_var(f"{etapa}_{tk_name}_{var_key}")
                e = tk.Entry(f_table, textvariable=val_var, state="readonly", justify="center", font=("Arial", 8), width=w_size, readonlybackground=bg_c)
                e.grid(row=r, column=c_idx+2, sticky="nsew", padx=1, pady=1)
                table_widgets.append(("cell", e))

        # --- ESCALAR FUENTES SEGÚN ANCHO DE VENTANA ---
        self._last_font_size = [8]
        def scale_fonts(event=None):
            try:
                w = self.root.winfo_width()
                # Escalar: 6 para 800px, 8 para 1200px, 10 para 1600px, max 11
                new_size = max(6, min(11, int(w / 160)))
                if new_size != self._last_font_size[0]:
                    self._last_font_size[0] = new_size
                    for wtype, widget in table_widgets:
                        try:
                            if wtype == "header":
                                widget.configure(font=("Arial", min(8, new_size), "bold"))
                            elif wtype == "name":
                                widget.configure(font=("Arial", min(8, new_size), "bold"))
                            elif wtype == "btn":
                                widget.configure(font=("Arial", new_size + 2))
                            else:
                                widget.configure(font=("Arial", new_size))
                        except: pass
            except: pass
        self._scale_fonts_pending = None
        def _scale_fonts_debounced(event=None):
            if self._scale_fonts_pending is not None:
                try: self.root.after_cancel(self._scale_fonts_pending)
                except: pass
            self._scale_fonts_pending = self.root.after(100, scale_fonts)
        self.root.bind("<Configure>", _scale_fonts_debounced)

    def construir_preview_tab(self):
        """Construye la solapa de Vista Previa con dibujo del buque y botones de reportes."""
        main_f = ttk.Frame(self.tab_preview)
        main_f.pack(fill="both", expand=True)

        # --- BARRA SUPERIOR ---
        top_bar = tk.Frame(main_f, bg="#1B3A5C", height=50)
        top_bar.pack(fill="x")
        top_bar.pack_propagate(False)
        tk.Label(top_bar, text="VISTA PREVIA DEL BUQUE Y REPORTES", font=("Arial", 8, "bold"), fg="white", bg="#1B3A5C").pack(side="left", padx=15)
        tk.Button(top_bar, text="  Actualizar Vista  ", bg="#2980B9", fg="white", font=("Arial", 8, "bold"),
                  command=lambda: self.refresh_preview_canvas(canvas_prev)).pack(side="left", padx=10, pady=8)
        tk.Button(top_bar, text="  GENERAR REPORTES PDF...  ", bg="#27AE60", fg="white", font=("Arial", 8, "bold"),
                  command=self.generar_con_seleccion, cursor="hand2").pack(side="right", padx=15, pady=8, ipadx=8)
        tk.Button(top_bar, text="  Vista Previa PDF  ", bg="#8E44AD", fg="white", font=("Arial", 8, "bold"),
                  command=self.preview_pdf_temp, cursor="hand2").pack(side="right", padx=5, pady=8)

        # --- CANVAS PRINCIPAL ---
        canvas_prev = tk.Canvas(main_f, bg="#E8F0FE")
        canvas_prev.pack(fill="both", expand=True)
        canvas_prev.bind("<Configure>", lambda e: self.refresh_preview_canvas(canvas_prev))

    def refresh_preview_canvas(self, cv):
        """Dibuja la vista del buque/unidad en el canvas de preview."""
        cv.delete("all")
        w = cv.winfo_width()
        h = cv.winfo_height()
        if w < 200 or h < 200: return

        buque_name = self.get_var("car_buque").get() or "Sin Nombre"
        tipo_nave = self.get_var("car_tipo_nave").get() or "BUQUE"
        cv.create_text(w/2, 18, text=f"{tipo_nave}: {buque_name}", font=("Arial", 8), fill="#1B3A5C")

        ship_w = w - 20
        if self.es_maritimo():
            # Buques: 2 vistas (babor arriba, estribor abajo)
            half_h = (h - 40) / 2
            self.dibujar_unidad_tk(cv, 10, 35, ship_w, half_h - 10, "BABOR", "inicial")
            self.dibujar_unidad_tk(cv, 10, 35 + half_h, ship_w, half_h - 10, "ESTRIBOR", "inicial")
        else:
            # Tierra/esfera/camion/ducto/electrico: canvas completo
            self.dibujar_unidad_tk(cv, 10, 35, ship_w, h - 50, "AMBOS", "inicial")

    def dibujar_unidad_tk(self, cv, x, y, width, height, side_label, etapa_key, tank_names=None, carb_names=None):
        """Dispatcher principal — redirige al dibujo correcto."""
        tm = self.get_tipo_medio()
        if "TANQUE FIJO" in tm:
            self._draw_vertical_tank(cv, x, y, width, height, etapa_key, tank_names, flotante=False)
        elif "TANQUE FLOTANTE" in tm:
            self._draw_vertical_tank(cv, x, y, width, height, etapa_key, tank_names, flotante=True)
        elif tm == "ESFERA DE GAS":
            self._draw_spheres(cv, x, y, width, height, etapa_key, tank_names)
        elif "CAMION GAS" in tm:
            self._draw_pressure_truck(cv, x, y, width, height, etapa_key, tank_names)
        elif "CAMION" in tm:
            self._draw_liquid_truck(cv, x, y, width, height, etapa_key, tank_names)
        elif tm in ("OLEODUCTO","POLIDUCTO","GASODUCTO"):
            self._draw_pipeline(cv, x, y, width, height, etapa_key, tm, tank_names)
        elif tm == "MEDICION ELECTRICA":
            self._draw_electric(cv, x, y, width, height, etapa_key, tank_names)
        elif "GASERO" in tm or ("GLP" in tm and "CAMION" not in tm):
            self._draw_moss_vessel(cv, x, y, width, height, side_label, etapa_key, tank_names, carb_names)
        elif "METANERO" in tm or "GNL" in tm:
            self._draw_membrane_vessel(cv, x, y, width, height, side_label, etapa_key, tank_names, carb_names)
        else:
            self.dibujar_buque_tk(cv, x, y, width, height, side_label, etapa_key, tank_names, carb_names)

    # ═══════════════════════════════════════════════════════════════════════════
    # SHARED DRAWING HELPERS
    # ═══════════════════════════════════════════════════════════════════════════

    def _sz(self, H, base=0.032, cap=8, floor=5):
        return ("Arial", min(cap, max(floor, int(H*base))))

    def _draw_bg(self, cv, x, y, W, H, color="#F4F6F7", border="#BDC3C7"):
        cv.create_rectangle(x, y, x+W, y+H, fill=color, outline=border, width=2)

    def _draw_fill_bar(self, cv, x, y_top, y_bot, x_left, x_right, pct, prod_fill, water_pct=0.0):
        """Draw product (bottom) and water layer (very bottom) with no blue for product."""
        if pct <= 0: return
        bar_h = y_bot - y_top
        # Water layer (blue)
        if water_pct > 0:
            wy = y_bot - int(bar_h * water_pct)
            cv.create_rectangle(x_left, wy, x_right, y_bot, fill="#3498DB", outline="")
        # Product layer above water
        prod_h = int(bar_h * max(0, pct - water_pct))
        if prod_h > 0:
            py = y_bot - int(bar_h*pct)
            # Gradient simulation: 3 shades
            shade1, shade2 = prod_fill
            cv.create_rectangle(x_left, py, x_right, y_bot - int(bar_h*water_pct), fill=shade1, outline="")
            # Highlight top 20% — sin stipple, dos capas de color
            hlt_h = max(2, prod_h//5)
            cv.create_rectangle(x_left, py, x_right, py+hlt_h, fill=shade2, outline="")

    def _draw_level_text(self, cv, cx, cy, pct, font, color="#1B3A5C"):
        if pct > 0.02:
            cv.create_text(cx, cy, text=f"{pct*100:.1f}%", font=font, fill=color)

    def _get_fill_pct(self, tk_name, etapa_key, alt_var="alt_ref", sond_var="s_corr"):
        """Return (vol_pct, water_pct) safely.
        vol_pct = fracción del tanque con producto total (incluye agua abajo).
        water_pct = fracción del tanque con agua (desde el fondo).
        Para dibujos: el producto ocupa [water_pct .. vol_pct], agua [0 .. water_pct].
        """
        try:
            if not etapa_key or not tk_name: return 0.0, 0.0
            ref = self.parse_float(self.get_var(f"{etapa_key}_{tk_name}_{alt_var}").get() or "0")
            if ref <= 0: return 0.0, 0.0
            s   = self.parse_float(self.get_var(f"{etapa_key}_{tk_name}_{sond_var}").get() or "0")
            # Intentar agua_s_real primero (marítimo), luego agua_mm (tierra), luego vol_nat_agua como fallback
            aw_str = (self.get_var(f"{etapa_key}_{tk_name}_agua_s_real").get()
                      or self.get_var(f"{etapa_key}_{tk_name}_agua_mm").get()
                      or "0")
            aw = self.parse_float(aw_str)
            vol_pct = min(max(s/ref, 0.0), 1.0)
            # agua como fracción de la altura total — limitada a 40% del volumen
            wat_pct = min(max(aw/ref, 0.0), vol_pct * 0.4)
            return vol_pct, wat_pct
        except: return 0.0, 0.0

    # ═══════════════════════════════════════════════════════════════════════════
    # TANQUE FIJO / FLOTANTE
    # ═══════════════════════════════════════════════════════════════════════════

    def _draw_vertical_tank(self, cv, x, y, W, H, etapa_key, tank_names, flotante=False):
        """Tanque vertical API 650 -- techo fijo conico 1:16 o techo flotante con ponton."""
        import math
        if tank_names is None: tank_names = self.lista_tanques
        if H < 40 or W < 60: return
        n = max(len(tank_names), 1)

        def bez(p0, p1, p2, p3, n_pts=16):
            pts = []
            for i in range(n_pts + 1):
                t = i / n_pts; u = 1 - t
                pts += [u**3*p0[0]+3*u**2*t*p1[0]+3*u*t**2*p2[0]+t**3*p3[0],
                        u**3*p0[1]+3*u**2*t*p1[1]+3*u*t**2*p2[1]+t**3*p3[1]]
            return pts

        fs  = min(10, max(5, int(H * 0.036))); fss = min(9, max(4, int(H * 0.029)))
        FT  = ("Arial", fs);  FTS = ("Arial", fss)

        TITLE_H  = max(16, int(H * 0.075))
        GROUND_Y = y + H - max(24, int(H * 0.12))
        BUND_H   = max(14, int(H * 0.09))

        # Vista lateral de tanque cilíndrico vertical: rectángulo más alto que ancho o cuadrado
        avail_h  = GROUND_Y - BUND_H - y - TITLE_H - 10
        pad      = max(10, int(W * 0.05))
        total_w  = W - 2 * pad
        raw_tw   = max(32, total_w // n)
        TK_W     = min(raw_tw, int(total_w * 0.90 / n))
        TK_H     = min(int(avail_h * 0.85), int(TK_W * 1.2))  # más alto que ancho o cuadrado
        TK_H     = max(40, TK_H)
        TK_BOT   = GROUND_Y - BUND_H - max(4, int(H * 0.02))
        TK_TOP   = TK_BOT - TK_H
        ELIPSE_RY = 0  # sin elipses, vista lateral pura

        gap = max(4, int(TK_W * 0.06))
        total_tanks_w = n * TK_W
        x_offset_cv = (W - total_tanks_w) // 2

        # ── Sky gradient background ───────────────────────────────────────
        sky_colors = ["#D6EAF8", "#DCE9F5", "#E2E9F1", "#E8EAEE", "#EFF0F2", "#F5F6F8", "#F7F9FC"]
        sky_zone_h = GROUND_Y - y
        for si, sc in enumerate(sky_colors):
            sy1 = y + int(sky_zone_h * si / len(sky_colors))
            sy2 = y + int(sky_zone_h * (si + 1) / len(sky_colors))
            cv.create_rectangle(x, sy1, x + W, sy2, fill=sc, outline="")
        cv.create_rectangle(x, y, x + W, y + H, fill="", outline="#85929E", width=1)

        tipo_lbl = "TANQUE TECHO FLOTANTE" if flotante else "TANQUE TECHO FIJO"
        instalacion = self.get_var("car_buque").get() or ""
        title_txt = f"{tipo_lbl}  --  {instalacion}" if instalacion else tipo_lbl
        cv.create_text(x + W // 2 + 1, y + TITLE_H // 2 + 1, text=title_txt, font=FT, fill="#B0B8C0")
        cv.create_text(x + W // 2, y + TITLE_H // 2, text=title_txt, font=FT, fill="#1B3A5C")

        # ── Ground (asphalt gradient) ─────────────────────────────────────
        ground_h = y + H - GROUND_Y
        gnd_cols = ["#6B7380", "#5F6672", "#545B66", "#4A5059", "#41474F"]
        for gi, gc in enumerate(gnd_cols):
            gy1 = GROUND_Y + int(ground_h * gi / len(gnd_cols))
            gy2 = GROUND_Y + int(ground_h * (gi + 1) / len(gnd_cols))
            cv.create_rectangle(x, gy1, x + W, gy2, fill=gc, outline="")
        cv.create_line(x, GROUND_Y, x + W, GROUND_Y, fill="#8A929C", width=1)

        # ── Bund (containment dike) ───────────────────────────────────────
        bund_x = x + pad // 2;  bund_w = W - pad
        bund_top = GROUND_Y - BUND_H
        cv.create_rectangle(bund_x + 3, bund_top + 3, bund_x + bund_w + 3, GROUND_Y + 6,
                            fill="#3A3F47", outline="")
        bund_grads = ["#D5D8DB", "#CDD0D3", "#C5C8CB", "#BDC0C3", "#B5B8BB", "#ADB0B3"]
        for bi, bc in enumerate(bund_grads):
            by1 = bund_top + int(BUND_H * bi / len(bund_grads))
            by2 = bund_top + int(BUND_H * (bi + 1) / len(bund_grads))
            cv.create_rectangle(bund_x, by1, bund_x + bund_w, by2, fill=bc, outline="")
        for li in range(0, bund_w, max(8, bund_w // 14)):
            cv.create_line(bund_x + li, bund_top + 4, bund_x + li + 5, GROUND_Y,
                           fill="#B0B5B8", width=1)
        for px2 in [bund_x, bund_x + bund_w]:
            cv.create_rectangle(px2 - 4, bund_top - 8, px2 + 4, GROUND_Y + 4,
                                fill="#909AA3", outline="#6D7880", width=1)
            cv.create_line(px2 - 4, bund_top - 8, px2 - 4, GROUND_Y + 4, fill="#C0C8CC", width=1)
            cv.create_line(px2 + 4, bund_top - 8, px2 + 4, GROUND_Y + 4, fill="#6D7880", width=1)
        cv.create_rectangle(bund_x, bund_top, bund_x + bund_w, GROUND_Y, fill="", outline="#7F8C8D", width=2)
        for ej in range(1, 4):
            ej_x = bund_x + int(bund_w * ej / 4)
            cv.create_line(ej_x, bund_top, ej_x, GROUND_Y, fill="#A0A5A8", width=1, dash=(6, 4))

        PROD_COL = [
            ("#C0392B", "#7B241C", "#F1948A"), ("#E67E22", "#A04000", "#FAD7A0"),
            ("#27AE60", "#1E8449", "#A9DFBF"), ("#D4AC0D", "#9A7D0A", "#F9E79F"),
            ("#784212", "#6E2C00", "#C39BD3"), ("#2C3E50", "#17202A", "#85929E"),
            ("#7D3C98", "#6C3483", "#D2B4DE"), ("#C2185B", "#AD1457", "#F48FB1"),
        ]

        for i, tn in enumerate(tank_names[:n]):
            tx = x + x_offset_cv + i * TK_W + gap // 2
            tw = TK_W - gap
            cx_t = tx + tw // 2
            inner_h = TK_BOT - TK_TOP
            sh_w = max(3, tw // 9)

            # ── Drop shadow ───────────────────────────────────────────────
            sh_off = max(3, int(tw * 0.04))
            cv.create_rectangle(tx + sh_off, TK_TOP + ELIPSE_RY + sh_off,
                                tx + tw + sh_off, TK_BOT + sh_off,
                                fill="#404850", outline="")
            cv.create_oval(tx + sh_off, TK_BOT - ELIPSE_RY + sh_off,
                           tx + tw + sh_off, TK_BOT + ELIPSE_RY + sh_off,
                           fill="#404850", outline="")

            # ── Concrete ringwall foundation ──────────────────────────────
            base_h = max(8, int(TK_H * 0.07))
            base_x = tx - max(6, tw // 8)
            base_w = tw + max(12, tw // 4)
            cv.create_rectangle(base_x + 2, TK_BOT + 2, base_x + base_w + 2, TK_BOT + base_h + 2,
                                fill="#6D7880", outline="")
            fnd_grads = ["#C8CDD1", "#BCC1C5", "#B0B5B9", "#A4A9AD", "#989DA1"]
            for fi, fc in enumerate(fnd_grads):
                fy1 = TK_BOT + int(base_h * fi / len(fnd_grads))
                fy2 = TK_BOT + int(base_h * (fi + 1) / len(fnd_grads))
                cv.create_rectangle(base_x, fy1, base_x + base_w, fy2, fill=fc, outline="")
            cv.create_rectangle(base_x, TK_BOT, base_x + base_w, TK_BOT + base_h,
                                fill="", outline="#8D9498", width=1)
            n_bolts = max(3, tw // 16)
            for bi in range(n_bolts):
                bx = base_x + int(base_w * (bi + 0.5) / n_bolts)
                cv.create_oval(bx - 2, TK_BOT + 1, bx + 2, TK_BOT + 5, fill="#6D7880", outline="#5D6D7E")

            # ── Fill percentages ──────────────────────────────────────────
            vp, wp = self._get_fill_pct(tn, etapa_key)
            pnm  = self.get_var(f"{etapa_key}_{tn}_prod_name").get() if etapa_key else ""
            vlit = self.get_var(f"{etapa_key}_{tn}_vol_nat_prod").get() if etapa_key else ""
            ci = abs(hash(pnm)) % len(PROD_COL) if pnm else i % len(PROD_COL)
            prod_c, prod_dk, prod_lt = PROD_COL[ci]

            # ── Cylinder body: 12-strip metallic gradient ─────────────────
            stripe_defs = [
                (0.00, 0.05, "#707D87"), (0.05, 0.10, "#808D97"),
                (0.10, 0.18, "#95A2AC"), (0.18, 0.28, "#A8B5BF"),
                (0.28, 0.40, "#BCC9D3"), (0.40, 0.52, "#D0DCE5"),
                (0.52, 0.60, "#E0E9EF"), (0.60, 0.68, "#E8F0F5"),
                (0.68, 0.76, "#DFE7EC"), (0.76, 0.84, "#CCD6DD"),
                (0.84, 0.92, "#B0BAC2"), (0.92, 1.00, "#8A9398"),
            ]
            for s_from, s_to, sc in stripe_defs:
                sx1 = int(tx + tw * s_from)
                sx2 = int(tx + tw * s_to)
                cv.create_rectangle(sx1, TK_TOP + ELIPSE_RY, sx2, TK_BOT, fill=sc, outline="")

            # ── Shell courses (5-7 horizontal weld seams) ─────────────────
            n_courses = min(7, max(5, TK_H // 25))
            for ci2 in range(1, n_courses):
                wy = TK_TOP + ELIPSE_RY + int((inner_h - ELIPSE_RY) * ci2 / n_courses)
                # Double weld line (realistic butt-weld seam)
                cv.create_line(tx + 2, wy, tx + tw - 2, wy, fill="#98A0A8", width=1)
                cv.create_line(tx + 2, wy + 1, tx + tw - 2, wy + 1, fill="#B8C0C8", width=1)

            # Vertical weld seam (staggered between courses)
            for ci2 in range(n_courses):
                wy1 = TK_TOP + ELIPSE_RY + int((inner_h - ELIPSE_RY) * ci2 / n_courses)
                wy2 = TK_TOP + ELIPSE_RY + int((inner_h - ELIPSE_RY) * (ci2 + 1) / n_courses)
                vwx = tx + int(tw * (0.45 + 0.10 * (ci2 % 2)))
                cv.create_line(vwx, wy1 + 2, vwx, wy2 - 2, fill="#B0B8C0", width=1, dash=(8, 12))

            # ── Water layer (blue, from bottom) ───────────────────────────
            if wp > 0:
                wy = int(TK_BOT - inner_h * wp)
                cv.create_rectangle(tx + sh_w, wy, tx + tw - sh_w, TK_BOT - ELIPSE_RY,
                                    fill="#5DADE2", outline="")
                cv.create_oval(tx + sh_w, wy - ELIPSE_RY // 2, tx + tw - sh_w, wy + ELIPSE_RY // 2,
                               fill="#2E86C1", outline="#2E86C1", width=1)
                for sw in range(3):
                    sw_y = wy + int((TK_BOT - ELIPSE_RY - wy) * (sw + 1) / 4)
                    cv.create_line(tx + sh_w + int(tw * 0.15), sw_y,
                                   tx + tw - sh_w - int(tw * 0.15), sw_y,
                                   fill="#85C1E9", width=1, dash=(4, 6))

            # ── Product fill above water ──────────────────────────────────
            if vp > wp + 0.01:
                py_top = int(TK_BOT - inner_h * vp)
                py_bot = int(TK_BOT - inner_h * wp) if wp > 0 else TK_BOT - ELIPSE_RY
                cv.create_rectangle(tx + sh_w, py_top, tx + tw - sh_w, py_bot, fill=prod_c, outline="")
                prod_zone_h = py_bot - py_top
                if prod_zone_h > 10:
                    refl_h = max(3, int(prod_zone_h * 0.12))
                    cv.create_rectangle(tx + sh_w, py_top, tx + tw - sh_w, py_top + refl_h,
                                        fill=prod_lt, outline="")
                    cv.create_rectangle(tx + sh_w, py_bot - max(2, refl_h // 2), tx + tw - sh_w, py_bot,
                                        fill=prod_dk, outline="")
                cv.create_oval(tx + sh_w, py_top - ELIPSE_RY // 2, tx + tw - sh_w, py_top + ELIPSE_RY // 2,
                               fill=prod_lt, outline=prod_dk, width=1)

            # ── Bottom ellipse ────────────────────────────────────────────
            cv.create_oval(tx, TK_BOT - ELIPSE_RY, tx + tw, TK_BOT + ELIPSE_RY,
                           fill="#8A9398", outline="#6D7880", width=1)
            if vp > 0.01:
                fill_col_bot = prod_c if vp > wp else "#5DADE2"
                cv.create_rectangle(tx + sh_w, TK_BOT - ELIPSE_RY, tx + tw - sh_w, TK_BOT,
                                    fill=fill_col_bot, outline="")
            else:
                cv.create_rectangle(tx, TK_BOT - ELIPSE_RY, tx + tw, TK_BOT, fill="#9BA7B2", outline="")

            # ── ROOF ──────────────────────────────────────────────────────
            if flotante:
                # ── Floating roof: pontón plano sobre el líquido ──────────
                PATAS_FRAC = 0.15
                ponton_h = max(6, int(TK_H * 0.04))
                liq_y   = int(TK_BOT - inner_h * vp)
                patas_y = int(TK_BOT - inner_h * PATAS_FRAC)
                fy = patas_y if vp <= PATAS_FRAC else liq_y

                # Support legs (patas de apoyo debajo del pontón)
                pata_base = fy + ponton_h
                for pi_frac in [0.20, 0.45, 0.70]:
                    px = int(tx + tw * pi_frac)
                    cv.create_line(px, pata_base, px, TK_BOT,
                                   fill="#7F8C8D", width=2)
                    cv.create_rectangle(px - 4, TK_BOT - 2, px + 4, TK_BOT + 1,
                                        fill="#808B96", outline="")

                # Pontón (rectángulo plano sobre el líquido)
                pont_margin = max(2, int(tw * 0.03))
                # Sombra del pontón
                cv.create_rectangle(tx + pont_margin + 2, fy + 2,
                                    tx + tw - pont_margin + 2, fy + ponton_h + 2,
                                    fill="#3A4550", outline="")
                # Cuerpo del pontón (gradiente)
                pont_grads = ["#6D8898", "#7A9CAC", "#88AAB8", "#7A9CAC", "#6D8898"]
                for pg_i, pg_c in enumerate(pont_grads):
                    pg_y1 = fy + int(ponton_h * pg_i / len(pont_grads))
                    pg_y2 = fy + int(ponton_h * (pg_i + 1) / len(pont_grads))
                    cv.create_rectangle(tx + pont_margin, pg_y1,
                                        tx + tw - pont_margin, pg_y2,
                                        fill=pg_c, outline="")
                cv.create_rectangle(tx + pont_margin, fy,
                                    tx + tw - pont_margin, fy + ponton_h,
                                    fill="", outline="#4A6878", width=2)

                # Rim seal (sello entre pontón y pared del tanque)
                seal_h = max(3, ponton_h // 2)
                # Sello izquierdo
                cv.create_rectangle(tx + 1, fy - seal_h, tx + pont_margin + 2, fy + ponton_h + seal_h,
                                    fill="#3D3D3D", outline="#2A2A2A", width=1)
                # Sello derecho
                cv.create_rectangle(tx + tw - pont_margin - 2, fy - seal_h,
                                    tx + tw - 1, fy + ponton_h + seal_h,
                                    fill="#3D3D3D", outline="#2A2A2A", width=1)

                # Drain pipe (articulado, del pontón al fondo)
                drain_x = tx + int(tw * 0.35)
                cv.create_line(drain_x, fy + ponton_h, drain_x - 4, TK_BOT - 6,
                               fill="#5D6D7E", width=2)
                cv.create_oval(drain_x - 6, fy + ponton_h - 3,
                               drain_x + 2, fy + ponton_h + 3,
                               fill="#4A5568", outline="")

                # Rolling ladder (escalera pivotante DENTRO del tanque, del borde superior al pontón)
                lad_x  = tx + int(tw * 0.82)
                lad_top_y = TK_TOP + 2
                lad_bot_y = fy
                lad_off = max(4, int(tw * 0.05))
                cv.create_line(lad_x, lad_top_y, lad_x, lad_bot_y, fill="#808B96", width=1)
                cv.create_line(lad_x + lad_off, lad_top_y, lad_x + lad_off, lad_bot_y, fill="#95A5A6", width=1)
                n_rungs = max(3, int((lad_bot_y - lad_top_y) / max(8, int(TK_H * 0.08))))
                for ri in range(n_rungs):
                    ry = int(lad_top_y + (lad_bot_y - lad_top_y) * (ri + 0.5) / n_rungs)
                    cv.create_line(lad_x, ry, lad_x + lad_off, ry, fill="#7F8C8D", width=1)

                # Guide poles (guías verticales)
                for cg_frac in [0.15, 0.85]:
                    cgx = int(tx + tw * cg_frac)
                    cv.create_line(cgx, fy, cgx, TK_TOP,
                                   fill="#95A5A6", width=1, dash=(4, 5))

                # Sounding reference (línea de medición)
                datum_x = tx + int(tw * 0.12)
                _sond_y_top = fy + 2
                _sond_y_bot = TK_BOT - 2
                if _sond_y_bot > _sond_y_top + 10:
                    cv.create_line(datum_x, _sond_y_top, datum_x, _sond_y_bot,
                                   fill="#F4D03F", width=1, dash=(3, 3))
                    cv.create_polygon(datum_x - 4, _sond_y_top + 6, datum_x, _sond_y_top,
                                      datum_x + 4, _sond_y_top + 6, fill="#F4D03F", outline="")
                    cv.create_polygon(datum_x - 4, _sond_y_bot - 6, datum_x, _sond_y_bot,
                                      datum_x + 4, _sond_y_bot - 6, fill="#F4D03F", outline="")
                    if fss >= 5:
                        cv.create_text(datum_x - 3, (_sond_y_top + _sond_y_bot) // 2,
                                       text="S", font=("Arial", max(4, fss - 2), "bold"),
                                       fill="#F4D03F", anchor="e")

                if vp > PATAS_FRAC and fss >= 5:
                    cv.create_text(cx_t, fy,
                                   text=f"Ponton {vp*100:.0f}%",
                                   font=("Arial", max(4, fss - 2)), fill="#D5D8DC")

            else:
                # ── Fixed cone roof (slope 1:16, almost flat) ─────────────
                # 1:16 slope = rise/run -> rise = tw / (2*16) -> very shallow
                roof_rise = max(4, int(tw / 32))  # 1:16 slope from edge to center
                apex_y = TK_TOP - roof_rise
                # Slight overhang beyond shell
                overhang = max(3, int(tw * 0.03))

                # Shadow behind roof
                cv.create_polygon(tx - overhang + 2, TK_TOP + ELIPSE_RY + 2,
                                  cx_t + 2, apex_y + 2,
                                  tx + tw + overhang + 2, TK_TOP + ELIPSE_RY + 2,
                                  fill="#707880", outline="")
                # Multi-face shading for very shallow cone
                faces = [
                    (0.00, 0.30, "#D8E0E5"), (0.30, 0.55, "#C8D2D8"),
                    (0.55, 0.75, "#B0BCC4"), (0.75, 1.00, "#96A8B0"),
                ]
                for f_from, f_to, f_col in faces:
                    fx1 = int(tx - overhang + (tw + 2 * overhang) * f_from)
                    fx2 = int(tx - overhang + (tw + 2 * overhang) * f_to)
                    cv.create_polygon(fx1, TK_TOP + ELIPSE_RY, cx_t, apex_y, fx2, TK_TOP + ELIPSE_RY,
                                      fill=f_col, outline="")
                # Specular band
                cv.create_polygon(tx + int(tw * 0.25), TK_TOP + ELIPSE_RY,
                                  cx_t, apex_y,
                                  tx + int(tw * 0.42), TK_TOP + ELIPSE_RY,
                                  fill="#E8EFF3", outline="")
                cv.create_polygon(tx - overhang, TK_TOP + ELIPSE_RY, cx_t, apex_y,
                                  tx + tw + overhang, TK_TOP + ELIPSE_RY,
                                  fill="", outline="#6D7D84", width=2)
                # Rafters
                for ri_a in [0.15, 0.30, 0.50, 0.70, 0.85]:
                    rib_x = int(tx + tw * ri_a)
                    cv.create_line(rib_x + 1, TK_TOP + ELIPSE_RY + 1, cx_t + 1, apex_y + 1,
                                   fill="#7A868F", width=1)
                    cv.create_line(rib_x, TK_TOP + ELIPSE_RY, cx_t, apex_y, fill="#9DABB5", width=1)
                # Center hub plate
                hub_r = max(3, tw // 16)
                cv.create_oval(cx_t - hub_r, apex_y - hub_r // 2, cx_t + hub_r, apex_y + hub_r // 2,
                               fill="#A0ABB4", outline="#7F8C8D")

                # PVRV (pressure-vacuum relief valve) - mushroom shape
                pvrv_h = max(8, int(TK_H * 0.08))
                pvrv_x = cx_t + int(tw * 0.15)
                # Nozzle stub
                cv.create_rectangle(pvrv_x - 3, apex_y - pvrv_h, pvrv_x + 3, apex_y - 2,
                                    fill="#808B96", outline="#5D6D7E", width=1)
                # Mushroom cap (wider disc on top)
                cap_w = max(8, tw // 8)
                cap_h = max(4, int(TK_H * 0.03))
                cv.create_oval(pvrv_x - cap_w // 2, apex_y - pvrv_h - cap_h,
                               pvrv_x + cap_w // 2, apex_y - pvrv_h + cap_h // 2,
                               fill="#95A5A6", outline="#636E72", width=1)
                cv.create_oval(pvrv_x - cap_w // 3, apex_y - pvrv_h - cap_h + 1,
                               pvrv_x + cap_w // 4, apex_y - pvrv_h,
                               fill="#B0BCC4", outline="")

                # Gauge hatch / thief hatch on roof
                gh_x = tx + int(tw * 0.30)
                gh_base = TK_TOP + ELIPSE_RY - 2
                # Interpolate roof height at gh_x
                ghf = abs(gh_x - cx_t) / max(1, tw // 2)
                gh_roof_y = int(apex_y + (TK_TOP + ELIPSE_RY - apex_y) * ghf)
                cv.create_rectangle(gh_x - 3, gh_roof_y - 8, gh_x + 3, gh_roof_y,
                                    fill="#5D6D7E", outline="#4A5568")
                cv.create_oval(gh_x - 5, gh_roof_y - 12, gh_x + 5, gh_roof_y - 7,
                               fill="#808B96", outline="#5D6D7E", width=1)

            # ── Top ellipse (cylinder cap) ────────────────────────────────
            if not flotante:
                cv.create_arc(tx, TK_TOP - ELIPSE_RY, tx + tw, TK_TOP + ELIPSE_RY,
                              start=180, extent=180,
                              fill="#CAD2D7", outline="#808B96", width=2, style="chord")
                cv.create_arc(tx + tw // 5, TK_TOP - ELIPSE_RY // 2,
                              tx + tw * 3 // 4, TK_TOP + ELIPSE_RY // 2,
                              start=195, extent=150,
                              fill="#E2E8EC", outline="", style="chord")

            # ── Wind girders (1-2 stiffening rings on shell) ──────────────
            wg_positions = [0.33, 0.66] if TK_H > 60 else [0.50]
            for ri_frac in wg_positions:
                ry = TK_TOP + ELIPSE_RY + int((TK_BOT - TK_TOP - ELIPSE_RY) * ri_frac)
                cv.create_oval(tx + 1, ry - 2, tx + tw + 1, ry + 4, fill="#808890", outline="")
                cv.create_oval(tx, ry - 3, tx + tw, ry + 3, fill="#A0ABB4", outline="#7F8C8D", width=1)
                cv.create_oval(tx + 2, ry - 2, tx + tw - 2, ry, fill="#C0CAD0", outline="")
                for lug_x in [tx - 5, tx + tw]:
                    cv.create_rectangle(lug_x, ry - 4, lug_x + 5, ry + 4, fill="#95A5A6", outline="#7F8C8D")

            # ── Nozzle stubs on shell ─────────────────────────────────────
            noz_fracs = [0.20, 0.50, 0.80]
            for nfrac in noz_fracs:
                noz_y = TK_TOP + ELIPSE_RY + int((inner_h - ELIPSE_RY) * nfrac)
                noz_len = max(4, int(tw * 0.06))
                # Nozzle pipe stub on right side of tank
                cv.create_rectangle(tx + tw, noz_y - 2, tx + tw + noz_len, noz_y + 2,
                                    fill="#6D7880", outline="#4A5568", width=1)
                # Blind flange cap
                cv.create_rectangle(tx + tw + noz_len, noz_y - 4, tx + tw + noz_len + 3, noz_y + 4,
                                    fill="#5D6D7E", outline="#4A5568", width=1)

            # ── Spiral staircase (diagonal line with landings) ────────────
            esc_x = tx + tw + max(6, int(tw * 0.08))
            esc_w = max(10, int(tw * 0.15))
            if esc_x + esc_w < x + W - 4:
                # Diagonal run (wrap simulation)
                n_landings = max(2, TK_H // 50)
                for li2 in range(n_landings):
                    lf1 = li2 / n_landings
                    lf2 = (li2 + 1) / n_landings
                    sy1 = int(TK_BOT - (inner_h - ELIPSE_RY) * lf1)
                    sy2 = int(TK_BOT - (inner_h - ELIPSE_RY) * lf2)
                    # Diagonal stair run
                    if li2 % 2 == 0:
                        cv.create_line(esc_x, sy1, esc_x + esc_w, sy2, fill="#808B96", width=2)
                        cv.create_line(esc_x + esc_w, sy1, esc_x + esc_w + max(4, esc_w // 3), sy2,
                                       fill="#95A5A6", width=1)
                    else:
                        cv.create_line(esc_x + esc_w, sy1, esc_x, sy2, fill="#808B96", width=2)
                    # Landing platform
                    cv.create_rectangle(esc_x - 2, sy2 - 2, esc_x + esc_w + 2, sy2 + 2,
                                        fill="#A0A8B0", outline="#808B96", width=1)
                # Handrail (continuous)
                cv.create_line(esc_x + esc_w + 3, TK_TOP + ELIPSE_RY, esc_x + esc_w + 3, TK_BOT,
                               fill="#BDC3C7", width=1)
                # Top platform
                plt_w = esc_w + 6
                cv.create_rectangle(esc_x - 3, TK_TOP - ELIPSE_RY - 6,
                                    esc_x + plt_w, TK_TOP - ELIPSE_RY,
                                    fill="#808B96", outline="#5D6D7E", width=1)
                rail_h = max(4, int(TK_H * 0.03))
                cv.create_line(esc_x - 3, TK_TOP - ELIPSE_RY - 6,
                               esc_x - 3, TK_TOP - ELIPSE_RY - 6 - rail_h, fill="#95A5A6", width=1)
                cv.create_line(esc_x + plt_w, TK_TOP - ELIPSE_RY - 6,
                               esc_x + plt_w, TK_TOP - ELIPSE_RY - 6 - rail_h, fill="#95A5A6", width=1)
                cv.create_line(esc_x - 3, TK_TOP - ELIPSE_RY - 6 - rail_h,
                               esc_x + plt_w, TK_TOP - ELIPSE_RY - 6 - rail_h, fill="#95A5A6", width=1)

            # ── Base piping and valves ────────────────────────────────────
            pipe_y_base = TK_BOT + base_h - 2
            asp_x = tx + tw // 4
            cv.create_rectangle(asp_x - 3, TK_BOT, asp_x + 3, pipe_y_base + 6, fill="#6D7880", outline="")
            cv.create_rectangle(asp_x - 2, TK_BOT, asp_x + 2, pipe_y_base + 6, fill="#808B96", outline="")
            cv.create_line(asp_x - 1, TK_BOT, asp_x - 1, pipe_y_base + 6, fill="#95A5A6", width=1)
            cv.create_rectangle(asp_x, pipe_y_base + 3, tx - 12, pipe_y_base + 7, fill="#6D7880", outline="")
            cv.create_rectangle(asp_x - 6, TK_BOT - 2, asp_x + 6, TK_BOT + 4,
                                fill="#4A5568", outline="#2C3E50", width=1)
            for fb in [asp_x - 5, asp_x + 4]:
                cv.create_oval(fb, TK_BOT - 1, fb + 2, TK_BOT + 1, fill="#808B96", outline="")
            vv_y = pipe_y_base + 5
            cv.create_polygon(asp_x - 7, vv_y - 6, asp_x + 7, vv_y + 6,
                              asp_x + 7, vv_y - 6, asp_x - 7, vv_y + 6,
                              fill="#2C3E50", outline="#1B2631", width=1)
            cv.create_line(asp_x, vv_y - 6, asp_x, vv_y - 14, fill="#5D6D7E", width=2)
            cv.create_oval(asp_x - 6, vv_y - 18, asp_x + 6, vv_y - 12, fill="", outline="#5D6D7E", width=2)

            # ── External level indicator (sight glass) ────────────────────
            lvl_x = tx - max(8, int(tw * 0.10))
            if lvl_x > x + 4:
                cv.create_rectangle(lvl_x - 3, TK_TOP + ELIPSE_RY, lvl_x + 3, TK_BOT - ELIPSE_RY,
                                    fill="#CDD4D8", outline="#95A5A6", width=1)
                cv.create_rectangle(lvl_x - 1, TK_TOP + ELIPSE_RY + 2, lvl_x + 1, TK_BOT - ELIPSE_RY - 2,
                                    fill="#E8F0F5", outline="")
                for mk_pct in [0.2, 0.4, 0.6, 0.8]:
                    mk_y = int(TK_BOT - inner_h * mk_pct)
                    cv.create_line(lvl_x - 7, mk_y, lvl_x + 7, mk_y, fill="#5D6D7E", width=1)
                    if fss >= 5:
                        cv.create_text(lvl_x - 9, mk_y, text=f"{int(mk_pct*100)}%",
                                       font=("Arial", max(4, fss - 2)), fill="#5D6D7E", anchor="e")
                if vp > 0:
                    lvl_y = int(TK_BOT - inner_h * vp)
                    cv.create_rectangle(lvl_x - 1, lvl_y, lvl_x + 1, TK_BOT - ELIPSE_RY - 2,
                                        fill=prod_c, outline="")
                    cv.create_rectangle(lvl_x - 7, lvl_y - 4, lvl_x + 7, lvl_y + 4,
                                        fill="#E74C3C", outline="#C0392B", width=2)

            # ── Final cylinder outline ────────────────────────────────────
            cv.create_rectangle(tx, TK_TOP + ELIPSE_RY, tx + tw, TK_BOT - ELIPSE_RY,
                                fill="", outline="#4A5568", width=2)
            cv.create_line(tx, TK_TOP + ELIPSE_RY, tx, TK_BOT - ELIPSE_RY, fill="#3D4B56", width=2)
            cv.create_line(tx + tw, TK_TOP + ELIPSE_RY, tx + tw, TK_BOT - ELIPSE_RY, fill="#3D4B56", width=2)

            # ── Bund drain ────────────────────────────────────────────────
            sump_x = bund_x + max(4, int(bund_w * 0.08))
            cv.create_rectangle(sump_x, bund_top + 1, sump_x + 10, bund_top + 5,
                                fill="#A9B2BC", outline="#7F8C8D")

            # ── Labels ────────────────────────────────────────────────────
            short = tn.replace("TANQUE ", "T.").replace("COMPARTIMENTO ", "C.").strip()[:8]
            plate_w = max(16, int(tw * 0.35))
            plate_h = max(10, int(inner_h * 0.06))
            plate_x_loc = cx_t - plate_w // 2
            plate_y_loc = TK_TOP + ELIPSE_RY + int(inner_h * 0.06)
            cv.create_rectangle(plate_x_loc, plate_y_loc, plate_x_loc + plate_w, plate_y_loc + plate_h,
                                fill="#F0F2F4", outline="#5D6D7E", width=1)
            txt_fill = self.contrast_text(prod_c) if vp > 0.2 else "#1B3A5C"
            cv.create_text(cx_t, plate_y_loc + plate_h // 2, text=short, font=FT, fill="#1B3A5C")
            cv.create_text(cx_t, plate_y_loc + plate_h + fs + 3, text=f"{vp*100:.0f}%", font=FTS,
                           fill=txt_fill if vp > 0.35 else "#1B3A5C")
            if pnm:
                cv.create_text(cx_t, TK_BOT - int(inner_h * 0.28), text=pnm[:12], font=FTS,
                               fill=txt_fill if vp > 0.4 else "#2C3E50")
            if vlit:
                cv.create_text(cx_t, TK_BOT - 12, text=f"{vlit} L", font=FTS,
                               fill=txt_fill if vp > 0.15 else "#1B3A5C")


    def _draw_spheres(self, cv, x, y, W, H, etapa_key, tank_names):
        """Horton spheres -- columns at EQUATOR splaying outward, beach-ball weld pattern, white/silver."""
        import math
        if tank_names is None: tank_names = self.lista_tanques
        if H < 60 or W < 80: return
        n = max(len(tank_names), 1)

        def bez(p0, p1, p2, p3, n_pts=16):
            pts = []
            for i in range(n_pts + 1):
                t = i / n_pts; u = 1 - t
                pts += [u**3*p0[0]+3*u**2*t*p1[0]+3*u*t**2*p2[0]+t**3*p3[0],
                        u**3*p0[1]+3*u**2*t*p1[1]+3*u*t**2*p2[1]+t**3*p3[1]]
            return pts

        fs  = min(10, max(5, int(H * 0.036)));  fss = min(9, max(4, int(H * 0.028)))
        FT  = ("Arial", fs);  FTS = ("Arial", fss)

        TITLE_H  = max(16, int(H * 0.075))
        GROUND_Y = y + H - max(18, int(H * 0.09))

        # ── Sky gradient background ───────────────────────────────────────
        sky_cols = ["#D6EAF8", "#DCEAF4", "#E2EBF0", "#E8ECEE", "#F0F2F4", "#F5F7FA"]
        sky_h = GROUND_Y - y
        for si, sc in enumerate(sky_cols):
            sy1 = y + int(sky_h * si / len(sky_cols))
            sy2 = y + int(sky_h * (si + 1) / len(sky_cols))
            cv.create_rectangle(x, sy1, x + W, sy2, fill=sc, outline="")
        cv.create_rectangle(x, y, x + W, y + H, fill="", outline="#5D6D7E", width=2)

        instalacion = self.get_var("car_buque").get() or "PLANTA"
        cv.create_text(x + W // 2 + 1, y + TITLE_H // 2 + 1,
                       text=f"ESFERAS DE GAS  --  {instalacion}", font=FT, fill="#B0B8C0")
        cv.create_text(x + W // 2, y + TITLE_H // 2,
                       text=f"ESFERAS DE GAS  --  {instalacion}", font=FT, fill="#1B3A5C")

        # ── Ground (concrete/gravel) ──────────────────────────────────────
        gnd_h = y + H - GROUND_Y
        gnd_cols = ["#8A929C", "#7D858F", "#707882", "#636B75", "#565E68"]
        for gi, gc in enumerate(gnd_cols):
            gy1 = GROUND_Y + int(gnd_h * gi / len(gnd_cols))
            gy2 = GROUND_Y + int(gnd_h * (gi + 1) / len(gnd_cols))
            cv.create_rectangle(x, gy1, x + W, gy2, fill=gc, outline="")
        cv.create_line(x + 1, GROUND_Y, x + W - 1, GROUND_Y, fill="#A0A8B0", width=1)

        zone_w = W - 30
        each_w = zone_w // n
        SPH_R  = min(int(H * 0.33), each_w // 2 - 14, int((GROUND_Y - y - TITLE_H - 22) * 0.46))
        SPH_R  = max(22, SPH_R)

        for i, tn in enumerate(tank_names[:n]):
            cx_s = x + 15 + i * each_w + each_w // 2
            # Columns attach at EQUATOR, splay outward to foundations
            LEG_H   = int(SPH_R * 0.85)
            EQ_Y    = GROUND_Y - LEG_H   # equator sits here
            cy_s    = EQ_Y                # sphere center = equator

            # ── Ground shadow ─────────────────────────────────────────────
            sh_r = int(SPH_R * 0.85)
            cv.create_oval(cx_s - sh_r, GROUND_Y - 3, cx_s + sh_r, GROUND_Y + 5,
                           fill="#404850", outline="")

            # ── 6-8 support columns from EQUATOR splaying outward ─────────
            n_cols = 8 if SPH_R > 40 else 6
            leg_w = max(3, SPH_R // 8)
            for ci2 in range(n_cols):
                ang_deg = ci2 * 360 / n_cols
                # Only draw columns visible in side view (front half roughly)
                ang_rad = math.radians(ang_deg)
                # Column attachment at equator (appears as ellipse projection)
                eq_x_off = int(SPH_R * 0.92 * math.cos(ang_rad))
                depth = math.sin(ang_rad)  # front/back (-1 to 1)
                # Skip columns on far side that are mostly hidden
                if abs(depth) > 0.85:
                    continue
                col_top_x = cx_s + eq_x_off
                col_top_y = cy_s
                # Splay outward: bottom is further from center than top
                splay = int(SPH_R * 0.35)
                col_bot_x = col_top_x + int(splay * math.cos(ang_rad) * 0.5)
                col_bot_y = GROUND_Y
                # Column shade based on depth
                col_fill = "#5D6D7E" if depth < 0 else "#808B96"
                # Tapered column
                tw_top = leg_w - 1
                tw_bot = leg_w + 2
                cv.create_polygon(col_top_x - tw_top, col_top_y,
                                  col_top_x + tw_top, col_top_y,
                                  col_bot_x + tw_bot, col_bot_y,
                                  col_bot_x - tw_bot, col_bot_y,
                                  fill=col_fill, outline="#4A5568", width=1)
                # Column highlight
                cv.create_line(col_top_x - tw_top + 1, col_top_y,
                               col_bot_x - tw_bot + 1, col_bot_y,
                               fill="#95A5A6", width=1)
                # Foundation pad
                bp_w = max(6, leg_w + 4)
                cv.create_rectangle(col_bot_x - bp_w, GROUND_Y - 3,
                                    col_bot_x + bp_w, GROUND_Y + 4,
                                    fill="#4A5568", outline="#2C3E50", width=1)
                for blt in [-bp_w + 2, bp_w - 2]:
                    cv.create_oval(col_bot_x + blt - 1, GROUND_Y - 1,
                                   col_bot_x + blt + 1, GROUND_Y + 1,
                                   fill="#3D4B56", outline="")

            # ── Diagonal cross-bracing between columns ────────────────────
            leg_spread = int(SPH_R * 0.92)
            mid_brace_y = cy_s + int(LEG_H * 0.50)
            cv.create_line(cx_s - leg_spread, cy_s, cx_s + leg_spread, GROUND_Y - 3,
                           fill="#7F8C8D", width=2, dash=(5, 4))
            cv.create_line(cx_s + leg_spread, cy_s, cx_s - leg_spread, GROUND_Y - 3,
                           fill="#7F8C8D", width=2, dash=(5, 4))
            # Horizontal bracing ring at mid-height
            cv.create_rectangle(cx_s - leg_spread, mid_brace_y - 2,
                                cx_s + leg_spread, mid_brace_y + 2,
                                fill="#6D7880", outline="#5D6D7E")

            # ── Equatorial walkway/platform with handrails ────────────────
            plt_w = int(SPH_R * 1.85)
            plt_h = max(5, int(SPH_R * 0.10))
            plt_y = cy_s
            # Platform shadow
            cv.create_rectangle(cx_s - plt_w // 2 + 2, plt_y + 2,
                                cx_s + plt_w // 2 + 2, plt_y + plt_h + 2,
                                fill="#4A5058", outline="")
            # Platform body with grating
            cv.create_rectangle(cx_s - plt_w // 2, plt_y, cx_s + plt_w // 2, plt_y + plt_h,
                                fill="#6D7880", outline="#5D6D7E", width=1)
            for gx in range(cx_s - plt_w // 2 + 4, cx_s + plt_w // 2 - 4, max(4, plt_w // 12)):
                cv.create_line(gx, plt_y + 1, gx, plt_y + plt_h - 1, fill="#5D6D7E", width=1)
            # Handrail posts and rails
            rail_h = max(8, int(SPH_R * 0.18))
            for rx in [-plt_w // 2 + 2, -plt_w // 4, 0, plt_w // 4, plt_w // 2 - 2]:
                bx2 = cx_s + rx
                cv.create_line(bx2, plt_y, bx2, plt_y - rail_h, fill="#808B96", width=1)
            cv.create_line(cx_s - plt_w // 2 + 2, plt_y - rail_h,
                           cx_s + plt_w // 2 - 2, plt_y - rail_h,
                           fill="#95A5A6", width=2)
            cv.create_line(cx_s - plt_w // 2 + 2, plt_y - rail_h // 2,
                           cx_s + plt_w // 2 - 2, plt_y - rail_h // 2,
                           fill="#808B96", width=1)

            # ── Access ladder from ground to equatorial platform ──────────
            lad_x = cx_s + int(SPH_R * 0.95)
            lad_w = max(6, int(SPH_R * 0.12))
            if lad_x + lad_w < x + W - 5:
                cv.create_line(lad_x, plt_y + plt_h, lad_x, GROUND_Y, fill="#808B96", width=2)
                cv.create_line(lad_x + lad_w, plt_y + plt_h, lad_x + lad_w, GROUND_Y, fill="#808B96", width=2)
                lad_h = GROUND_Y - plt_y - plt_h
                n_rungs = max(4, lad_h // max(6, int(SPH_R * 0.1)))
                for ri in range(n_rungs):
                    ry = plt_y + plt_h + int(lad_h * (ri + 0.5) / n_rungs)
                    cv.create_line(lad_x, ry, lad_x + lad_w, ry, fill="#7F8C8D", width=1)
                # Safety cage hoops
                for ci2 in range(0, n_rungs, max(1, n_rungs // 3)):
                    ry = plt_y + plt_h + int(lad_h * (ci2 + 0.5) / n_rungs)
                    cage_r = lad_w + 3
                    cv.create_arc(lad_x - 2, ry - cage_r, lad_x + lad_w + 2, ry + cage_r,
                                  start=0, extent=180, style="arc", outline="#95A5A6", width=1)

            # ── Sphere: white/silver 3D gradient (LPG reflective) ─────────
            # White/silver gradient for mandatory LPG reflective coating
            sphere_gradient = [
                "#A8B0B8", "#B0B8C0", "#B8C0C8", "#C0C8D0", "#C8D0D8",
                "#D0D8E0", "#D8E0E5", "#E0E5EA", "#E5EAEF", "#EAF0F4",
                "#EFF4F8", "#F2F6FA", "#F5F9FC", "#F8FBFD", "#FAFCFE",
                "#FCFDFE", "#FDFEFE", "#FEFEFE", "#FEFFFE", "#FFFFFF",
            ]
            n_layers = len(sphere_gradient)
            # Drop shadow
            cv.create_oval(cx_s - SPH_R + 4, cy_s - SPH_R + 4,
                           cx_s + SPH_R + 4, cy_s + SPH_R + 4,
                           fill="#2A3540", outline="")
            for k in range(n_layers - 1, -1, -1):
                frac  = k / (n_layers - 1)
                r_k   = int(SPH_R * (0.08 + frac * 0.92))
                if r_k < 1: continue
                max_off = int(SPH_R * 0.35)
                off_x   = int(max_off * (1 - frac) * 0.60)
                off_y   = int(max_off * (1 - frac) * 0.85)
                sc3 = sphere_gradient[k]
                cv.create_oval(cx_s - r_k - off_x // 2, cy_s - r_k - off_y // 2,
                               cx_s + r_k - off_x // 2, cy_s + r_k - off_y // 2,
                               fill=sc3, outline="")

            # ── Weld seams: beach-ball pattern (meridional + equatorial) ──
            # Equatorial weld
            cv.create_oval(cx_s - SPH_R + 2, cy_s - 1, cx_s + SPH_R - 2, cy_s + 1,
                           fill="", outline="#A0A8B0", width=1)
            # 4 meridional welds (vertical great circles seen as ellipses)
            for m_frac in [-0.55, -0.20, 0.20, 0.55]:
                m_off = int(SPH_R * m_frac)
                m_w = max(2, int(SPH_R * abs(1 - abs(m_frac)) * 0.30))
                cv.create_oval(cx_s + m_off - m_w, cy_s - SPH_R + 2,
                               cx_s + m_off + m_w, cy_s + SPH_R - 2,
                               fill="", outline="#A0A8B0", width=1)

            # Sphere outline
            cv.create_oval(cx_s - SPH_R, cy_s - SPH_R, cx_s + SPH_R, cy_s + SPH_R,
                           fill="", outline="#6D7880", width=3)

            # ── Gas fill level ────────────────────────────────────────────
            try:
                vol = self.parse_float(self.get_var(f"{etapa_key}_{tn}_vol_nat_prod").get() or "0") if etapa_key and tn else 0
                cap = self.parse_float(self.get_var(f"{etapa_key}_{tn}_alt_ref").get() or "1")
                vp  = min(max(vol / (cap if cap > 0 else 1), 0.0), 1.0)
            except: vp = 0.0

            pnm  = self.get_var(f"{etapa_key}_{tn}_prod_name").get() if etapa_key else ""
            pres = self.get_var(f"{etapa_key}_{tn}_temp").get() if etapa_key else ""

            if vp > 0.02:
                fill_h  = int(SPH_R * (2 * vp - 1))
                fill_top = cy_s - fill_h
                clip_top = max(cy_s - SPH_R + 3, fill_top)
                fill_c = self.get_prod_color(tn, etapa_key)[0] if etapa_key else "#F39C12"
                if fill_c in ("#3498DB", "#5DADE2", "#2E86C1", "#85C1E9", "#B8E0F7"):
                    fill_c = "#FF6B35"
                cv.create_rectangle(cx_s - SPH_R + 3, clip_top,
                                    cx_s + SPH_R - 3, cy_s + SPH_R - 3,
                                    fill=fill_c, outline="")
                # Re-draw sphere outline over fill
                cv.create_oval(cx_s - SPH_R, cy_s - SPH_R, cx_s + SPH_R, cy_s + SPH_R,
                               fill="", outline="#6D7880", width=3)
                # Level chord line
                lv_y = clip_top
                try:
                    half_chord = int((SPH_R**2 - (cy_s - lv_y)**2)**0.5)
                except: half_chord = 0
                if half_chord > 0:
                    cv.create_line(cx_s - half_chord, lv_y, cx_s + half_chord, lv_y,
                                   fill="#4A4A4A", width=2, dash=(6, 3))
                    cv.create_line(cx_s - half_chord + 3, lv_y + 2,
                                   cx_s + half_chord - 3, lv_y + 2,
                                   fill="#FFFFFF", width=1, dash=(3, 5))
                cv.create_text(cx_s, cy_s + int(SPH_R * 0.25), text=f"{vp*100:.0f}%",
                               font=FTS, fill=self.contrast_text(fill_c))

            # ── Small top platform for relief valve ───────────────────────
            top_plt_w = max(12, int(SPH_R * 0.40))
            top_plt_y = cy_s - SPH_R - 2
            cv.create_rectangle(cx_s - top_plt_w // 2, top_plt_y - 3,
                                cx_s + top_plt_w // 2, top_plt_y,
                                fill="#6D7880", outline="#5D6D7E", width=1)
            # Mini handrail on top platform
            cv.create_line(cx_s - top_plt_w // 2, top_plt_y - 3,
                           cx_s - top_plt_w // 2, top_plt_y - 3 - max(4, int(SPH_R * 0.08)),
                           fill="#808B96", width=1)
            cv.create_line(cx_s + top_plt_w // 2, top_plt_y - 3,
                           cx_s + top_plt_w // 2, top_plt_y - 3 - max(4, int(SPH_R * 0.08)),
                           fill="#808B96", width=1)
            cv.create_line(cx_s - top_plt_w // 2, top_plt_y - 3 - max(4, int(SPH_R * 0.08)),
                           cx_s + top_plt_w // 2, top_plt_y - 3 - max(4, int(SPH_R * 0.08)),
                           fill="#95A5A6", width=1)

            # PSV (pressure safety valve) on top platform
            psv_x = cx_s + int(SPH_R * 0.10)
            psv_top = top_plt_y - max(14, int(SPH_R * 0.25))
            cv.create_rectangle(psv_x - 4, psv_top, psv_x + 4, top_plt_y - 3,
                                fill="#E74C3C", outline="#922B21", width=1)
            cv.create_polygon(psv_x - 7, psv_top - 2, psv_x + 7, psv_top - 2,
                              psv_x + 4, psv_top + 5, psv_x - 4, psv_top + 5,
                              fill="#C0392B", outline="#7B241C")
            cv.create_rectangle(psv_x - 2, psv_top - 10, psv_x + 2, psv_top - 1,
                                fill="#5D6D7E", outline="")
            cv.create_text(psv_x, psv_top - 12, text="PSV", font=("Arial", max(3, fss - 2)), fill="#E74C3C")

            # ── Fire deluge piping ring near top ──────────────────────────
            ring_r = int(SPH_R * 1.05)
            ring_h = max(3, int(SPH_R * 0.06))
            # Upper deluge ring
            cv.create_oval(cx_s - int(ring_r * 0.85), cy_s - int(SPH_R * 0.55) - ring_h,
                           cx_s + int(ring_r * 0.85), cy_s - int(SPH_R * 0.55) + ring_h,
                           fill="", outline="#2E86C1", width=2, dash=(3, 4))
            # Spray nozzles on ring
            for dang in [30, 150, 210, 330]:
                dr_x = cx_s + int(ring_r * 0.85 * math.cos(math.radians(dang)))
                dr_y = cy_s - int(SPH_R * 0.55) + int(ring_h * math.sin(math.radians(dang)))
                cv.create_oval(dr_x - 2, dr_y - 2, dr_x + 2, dr_y + 2,
                               fill="#7F8C8D", outline="#5D6D7E", width=1)
                cv.create_line(dr_x, dr_y, dr_x, dr_y + max(4, int(SPH_R * 0.10)),
                               fill="#AAB2B9", width=1, dash=(2, 3))

            # ── Pressure gauge ────────────────────────────────────────────
            gx = cx_s + int(SPH_R * 0.58); gy = cy_s - int(SPH_R * 0.38)
            gauge_r = max(8, int(SPH_R * 0.15))
            cv.create_line(gx, gy + gauge_r, gx + int(SPH_R * 0.15), gy + gauge_r + 8,
                           fill="#808B96", width=2)
            cv.create_oval(gx - gauge_r - 2, gy - gauge_r - 2, gx + gauge_r + 2, gy + gauge_r + 2,
                           fill="#4A5568", outline="#2C3E50", width=1)
            cv.create_oval(gx - gauge_r, gy - gauge_r, gx + gauge_r, gy + gauge_r,
                           fill="#FDFEFE", outline="#E74C3C", width=2)
            cv.create_arc(gx - gauge_r + 3, gy - gauge_r + 3, gx + gauge_r - 3, gy + gauge_r - 3,
                          start=30, extent=240, style="arc", outline="#E8E8E8", width=1)
            ang_v = pres[:3] if pres else "???"
            try:
                p_frac = min(float(pres.replace(",", ".")) / 1000, 1.0) if pres else 0.3
            except: p_frac = 0.3
            ang = math.radians(-30 + 240 * p_frac)
            nx = gx + int(gauge_r * 0.7 * math.cos(ang))
            ny = gy - int(gauge_r * 0.7 * math.sin(ang))
            cv.create_line(gx, gy, nx, ny, fill="#C0392B", width=2)
            cv.create_oval(gx - 2, gy - 2, gx + 2, gy + 2, fill="#C0392B", outline="")
            cv.create_text(gx, gy + gauge_r // 2, text=ang_v[:3],
                           font=("Arial", max(4, fss - 1)), fill="#C0392B")

            # ── Nameplate ─────────────────────────────────────────────────
            lbl = tn[:12]
            plate_w = max(30, int(each_w * 0.5))
            cv.create_rectangle(cx_s - plate_w // 2, GROUND_Y + 3,
                                cx_s + plate_w // 2, GROUND_Y + 3 + fss + 6,
                                fill="#F0F2F4", outline="#5D6D7E", width=1)
            cv.create_text(cx_s, GROUND_Y + 6 + fss // 2, text=lbl, font=FTS, fill="#1B3A5C")
            if pnm:
                cv.create_text(cx_s, GROUND_Y + 8 + fss + fss // 2,
                               text=pnm[:12], font=FTS, fill="#5D6D7E")

        # ── Interconnecting piping manifold ───────────────────────────────
        if n > 1:
            manif_y = GROUND_Y - 10
            first_cx = x + 15 + each_w // 2
            last_cx  = x + 15 + (n - 1) * each_w + each_w // 2
            cv.create_rectangle(first_cx, manif_y - 5, last_cx, manif_y + 1, fill="#4A5568", outline="")
            cv.create_rectangle(first_cx, manif_y - 4, last_cx, manif_y, fill="#6D7880", outline="")
            cv.create_line(first_cx, manif_y - 3, last_cx, manif_y - 3, fill="#808B96", width=1)
            cv.create_rectangle(first_cx, manif_y - 5, last_cx, manif_y + 1, fill="", outline="#4A5568", width=1)
            for i2 in range(n):
                cx_m = x + 15 + i2 * each_w + each_w // 2
                cv.create_rectangle(cx_m - 2, manif_y - 16, cx_m + 2, manif_y - 4, fill="#5D6D7E", outline="")
                cv.create_polygon(cx_m - 4, manif_y - 16, cx_m + 4, manif_y - 10,
                                  cx_m + 4, manif_y - 16, cx_m - 4, manif_y - 10,
                                  fill="#E74C3C", outline="#922B21", width=1)
            for fx in [first_cx, last_cx]:
                cv.create_rectangle(fx - 2, manif_y - 7, fx + 2, manif_y + 3,
                                    fill="#4A5568", outline="#2C3E50")


    def _draw_liquid_truck(self, cv, x, y, W, H, etapa_key, tank_names):
        """DOT 406 fuel tanker -- elliptical barrel, multi-compartment, polished aluminum."""
        if tank_names is None: tank_names = self.lista_tanques
        if H < 80 or W < 180: return
        import math

        n   = max(len(tank_names), 1)
        fs  = min(9, max(5, int(H * 0.032))); fss = min(7, max(4, int(H * 0.026)))
        FT  = ("Arial", fs); FTS = ("Arial", fss)
        patente = self.get_var("car_patente").get() or "CISTERNA"

        # ── Proportions ───────────────────────────────────────────────────
        GROUND_Y    = y + H - max(14, int(H * 0.08))
        AXLE_R      = max(10, int(H * 0.12))
        CISTERN_BOT = GROUND_Y - AXLE_R * 2 - 4
        CISTERN_H   = max(30, int(H * 0.38))
        CISTERN_TOP = CISTERN_BOT - CISTERN_H
        CISTERN_MID = (CISTERN_TOP + CISTERN_BOT) // 2
        ELLIP_FLAT  = max(3, int(CISTERN_H * 0.08))  # elliptical top flatness

        PAD_L  = max(8, int(W * 0.03))
        CAB_W  = max(55, int(W * 0.17))
        CAB_X  = x + PAD_L
        CIS_X  = CAB_X + CAB_W + max(6, int(W * 0.01))
        CIS_W  = W - PAD_L - CAB_W - max(6, int(W * 0.01)) - max(8, int(W * 0.03))
        DOME_R = int(CISTERN_H * 0.50)  # end-cap dome radius

        PROD_COL = [
            ("#C0392B", "#922B21", "#F1948A"), ("#E67E22", "#A04000", "#FAD7A0"),
            ("#27AE60", "#1E8449", "#A9DFBF"), ("#D4AC0D", "#9A7D0A", "#F9E79F"),
            ("#784212", "#6E2C00", "#C39BD3"), ("#2C3E50", "#17202A", "#85929E"),
            ("#8E44AD", "#6C3483", "#D2B4DE"), ("#E74C3C", "#C0392B", "#F5B7B1"),
        ]

        # ── Background ────────────────────────────────────────────────────
        cv.create_rectangle(x, y, x + W, y + H, fill="#F5F7FA", outline="#5D6D7E", width=2)
        cv.create_rectangle(x, GROUND_Y, x + W, y + H, fill="#5A6370", outline="")
        cv.create_text(x + W // 2, y + max(10, int(H * 0.06)),
                       text=f"CAMION CISTERNA  --  {patente}", font=FT, fill="#1B3A5C")

        # ── Chassis frame rails ───────────────────────────────────────────
        ch_h = max(5, int(H * 0.04))
        ch_y = CISTERN_BOT
        cv.create_rectangle(CAB_X + int(CAB_W * 0.3), ch_y, CIS_X + CIS_W, ch_y + ch_h,
                            fill="#4A5568", outline="#2C3E50", width=1)
        cv.create_line(CAB_X + int(CAB_W * 0.3), ch_y + ch_h // 2,
                       CIS_X + CIS_W, ch_y + ch_h // 2, fill="#5D6D7E", width=1)
        for ci in range(5):
            cx_tr = CIS_X + int(CIS_W * ci / 4)
            cv.create_rectangle(cx_tr - 3, ch_y - max(4, int(H * 0.03)), cx_tr + 3, ch_y + ch_h,
                                fill="#3D4B56", outline="")

        # ── ELLIPTICAL BARREL with compartments ───────────────────────────
        COMP_W = CIS_W // n
        # Draw the overall barrel shape (elliptical top curve)
        # Top arc: slightly flatter than bottom = elliptical cross-section look
        cv.create_arc(CIS_X, CISTERN_TOP - ELLIP_FLAT, CIS_X + CIS_W, CISTERN_TOP + ELLIP_FLAT * 3,
                      start=0, extent=180, style="arc", outline="#B0B8C0", width=1)

        for i, tn in enumerate(tank_names[:n]):
            cx2 = CIS_X + i * COMP_W
            cw2 = COMP_W
            is_first = (i == 0)
            is_last  = (i == n - 1)

            body_x1 = cx2 + (DOME_R if is_first else 0)
            body_x2 = cx2 + cw2 - (DOME_R if is_last else 0)

            # Polished aluminum barrel gradient (8 strips, silver)
            barrel_cols = ["#B8C0C8", "#C4CCD4", "#D0D8E0", "#DCE4EC",
                           "#E4ECF0", "#DCE4EC", "#CCD4DC", "#BCC4CC"]
            for si, sc in enumerate(barrel_cols):
                sy = CISTERN_TOP + int(CISTERN_H * si / len(barrel_cols))
                sh = CISTERN_H // len(barrel_cols) + 1
                cv.create_rectangle(body_x1, sy, body_x2, sy + sh, fill=sc, outline="")
            # Specular highlight band near top
            cv.create_rectangle(body_x1 + 2, CISTERN_TOP + max(2, CISTERN_H // 8),
                                body_x2 - 2, CISTERN_TOP + max(4, CISTERN_H // 5),
                                fill="#F0F4F8", outline="")

            # Elliptical end caps (left=first, right=last)
            if is_first:
                for ki, kc in enumerate(["#A0A8B0", "#B0B8C0", "#C0C8D0", "#D0D8E0", "#E0E8F0"]):
                    kr = DOME_R - ki * 3
                    if kr > 0:
                        cv.create_oval(cx2, CISTERN_TOP + ki * 2, cx2 + 2 * DOME_R, CISTERN_BOT - ki * 2,
                                       fill=kc, outline="")
            if is_last:
                for ki, kc in enumerate(["#E0E8F0", "#D0D8E0", "#C0C8D0", "#B0B8C0", "#A0A8B0"]):
                    kr = DOME_R - ki * 3
                    if kr > 0:
                        cv.create_oval(cx2 + cw2 - 2 * DOME_R + ki * 2, CISTERN_TOP + ki * 2,
                                       cx2 + cw2, CISTERN_BOT - ki * 2, fill=kc, outline="")

            # ── Compartment rings (vertical baffle bands) ─────────────────
            if not is_last:
                sep_x = cx2 + cw2
                cv.create_rectangle(sep_x - 3, CISTERN_TOP - 1, sep_x + 3, CISTERN_BOT + 1,
                                    fill="#8A929C", outline="#5D6D7E", width=1)
                cv.create_line(sep_x, CISTERN_TOP, sep_x, CISTERN_BOT, fill="#A0A8B0", width=1)

            # ── Fill level ────────────────────────────────────────────────
            vp, wp = self._get_fill_pct(tn, etapa_key)
            pnm  = self.get_var(f"{etapa_key}_{tn}_prod_name").get() if etapa_key else ""
            vlit = self.get_var(f"{etapa_key}_{tn}_vol_nat_prod").get() if etapa_key else ""
            ci_c = abs(hash(pnm)) % len(PROD_COL) if pnm else i % len(PROD_COL)
            prod_c, prod_dk, prod_lt = PROD_COL[ci_c]

            fill_top = int(CISTERN_BOT - CISTERN_H * vp)
            _fx1 = (cx2 + 2) if is_first else (body_x1 + 2)
            _fx2 = (cx2 + cw2 - 2) if is_last else (body_x2 - 2)
            if vp > 0.02:
                if wp > 0:
                    wat_top = int(CISTERN_BOT - CISTERN_H * wp)
                    cv.create_rectangle(_fx1, wat_top, _fx2, CISTERN_BOT - 2,
                                        fill="#5DADE2", outline="")
                cv.create_rectangle(_fx1, fill_top, _fx2,
                                    int(CISTERN_BOT - CISTERN_H * wp) if wp > 0 else CISTERN_BOT - 2,
                                    fill=prod_c, outline="")
                cv.create_rectangle(_fx1, fill_top, _fx2,
                                    fill_top + max(2, int(CISTERN_H * 0.06)),
                                    fill=prod_lt, outline="")
                cv.create_line(body_x1, fill_top, body_x2, fill_top, fill=prod_dk, width=2)

            # Barrel outline
            cv.create_rectangle(body_x1, CISTERN_TOP, body_x2, CISTERN_BOT,
                                fill="", outline="#8A929C", width=2)

            # ── Manhole cover on top (one per compartment) ────────────────
            top_cx = cx2 + cw2 // 2
            mh_r = max(4, int(min(cw2, CISTERN_H) * 0.10))
            cv.create_oval(top_cx - mh_r, CISTERN_TOP - mh_r - 2,
                           top_cx + mh_r, CISTERN_TOP + mh_r - 2,
                           fill="#B8C0C8", outline="#636E72", width=2)
            # Hinge on manhole
            cv.create_rectangle(top_cx - mh_r - 3, CISTERN_TOP - 3,
                                top_cx - mh_r, CISTERN_TOP + 2,
                                fill="#808B96", outline="#5D6D7E")
            # Handle on manhole
            cv.create_line(top_cx - 2, CISTERN_TOP - mh_r - 2,
                           top_cx + 2, CISTERN_TOP - mh_r - 2, fill="#4A5568", width=2)

            # ── Bottom loading valve under each compartment ───────────────
            bv_x = top_cx
            bv_y = CISTERN_BOT + ch_h + 2
            cv.create_rectangle(bv_x - 3, CISTERN_BOT, bv_x + 3, bv_y + 4,
                                fill="#5D6D7E", outline="#4A5568", width=1)
            cv.create_polygon(bv_x - 5, bv_y, bv_x + 5, bv_y + 6,
                              bv_x + 5, bv_y, bv_x - 5, bv_y + 6,
                              fill="#E74C3C", outline="#922B21", width=1)

            # ── Labels ────────────────────────────────────────────────────
            short = tn.replace("COMPARTIMENTO ", "C.").replace("TK ", "").strip()[:6]
            txt_fill = "white" if vp > 0.35 else "#1B3A5C"
            cv.create_text(body_x1 + (body_x2 - body_x1) // 2,
                           CISTERN_TOP + max(10, int(CISTERN_H * 0.15)),
                           text=short, font=FT, fill=txt_fill)
            if vp > 0:
                cv.create_text(body_x1 + (body_x2 - body_x1) // 2,
                               CISTERN_TOP + max(22, int(CISTERN_H * 0.35)),
                               text=f"{vp*100:.0f}%", font=FTS, fill=txt_fill)
            if pnm:
                cv.create_text(body_x1 + (body_x2 - body_x1) // 2,
                               CISTERN_BOT - max(18, int(CISTERN_H * 0.25)),
                               text=pnm[:10], font=FTS, fill=txt_fill)
            if vlit:
                cv.create_text(body_x1 + (body_x2 - body_x1) // 2,
                               CISTERN_BOT - max(6, int(CISTERN_H * 0.10)),
                               text=f"{vlit}L", font=FTS, fill=txt_fill)

        # ── Hazmat diamond placard ────────────────────────────────────────
        plac_x = CIS_X + CIS_W - max(20, int(CIS_W * 0.08))
        plac_y = CISTERN_MID
        plac_s = max(8, int(CISTERN_H * 0.14))
        cv.create_polygon(plac_x, plac_y - plac_s, plac_x + plac_s, plac_y,
                          plac_x, plac_y + plac_s, plac_x - plac_s, plac_y,
                          fill="#E74C3C", outline="#922B21", width=1)
        cv.create_text(plac_x, plac_y, text="3", font=("Arial", max(3, fss - 1), "bold"), fill="white")

        # ── SEMI-TRUCK CAB ────────────────────────────────────────────────
        CAB_BOT = CISTERN_BOT
        CAB_TOP = CAB_BOT - CISTERN_H - max(8, int(CISTERN_H * 0.22))
        CAB_H   = CAB_BOT - CAB_TOP
        cv.create_rectangle(CAB_X + int(CAB_W * 0.1), CAB_BOT, CAB_X + CAB_W, CAB_BOT + ch_h,
                            fill="#4A5568", outline="#2C3E50", width=1)
        cab_slope = max(4, int(CAB_W * 0.12))
        cab_pts = [CAB_X, CAB_BOT, CAB_X + CAB_W, CAB_BOT,
                   CAB_X + CAB_W, CAB_TOP + int(CAB_H * 0.18),
                   CAB_X + CAB_W - cab_slope, CAB_TOP,
                   CAB_X + cab_slope, CAB_TOP,
                   CAB_X, CAB_TOP + int(CAB_H * 0.18)]
        cv.create_polygon(cab_pts, fill="#2E4053", outline="#1B2631", width=2)

        # Stripe and chrome
        stripe_y1 = CAB_TOP + int(CAB_H * 0.35)
        stripe_y2 = CAB_TOP + int(CAB_H * 0.52)
        cv.create_rectangle(CAB_X, stripe_y1, CAB_X + CAB_W, stripe_y2, fill="#C0392B", outline="")
        cv.create_rectangle(CAB_X, stripe_y2, CAB_X + CAB_W, stripe_y2 + max(2, int(CAB_H * 0.03)),
                            fill="#95A5A6", outline="")

        # Roof and spoiler
        cv.create_rectangle(CAB_X, CAB_TOP + 1, CAB_X + CAB_W, CAB_TOP + max(5, int(CAB_H * 0.09)),
                            fill="#3D5166", outline="")
        sp_h = max(6, int(CAB_H * 0.16)); sp_w = int(CAB_W * 0.75)
        cv.create_polygon(CAB_X + int(CAB_W * 0.12), CAB_TOP,
                          CAB_X + int(CAB_W * 0.12) + sp_w, CAB_TOP,
                          CAB_X + int(CAB_W * 0.12) + sp_w, CAB_TOP - sp_h + max(2, sp_h // 4),
                          CAB_X + int(CAB_W * 0.12) + int(sp_w * 0.8), CAB_TOP - sp_h,
                          CAB_X + int(CAB_W * 0.12), CAB_TOP - sp_h,
                          fill="#3D5166", outline="#2C3E50", width=1)

        # Window
        front_w = max(4, int(CAB_W * 0.12))
        cv.create_rectangle(CAB_X + CAB_W - front_w, CAB_TOP + int(CAB_H * 0.18),
                            CAB_X + CAB_W, CAB_TOP + int(CAB_H * 0.72),
                            fill="#1A5276", outline="#0E3850", width=1)
        cv.create_polygon(CAB_X + CAB_W - front_w, CAB_TOP + int(CAB_H * 0.18),
                          CAB_X + CAB_W, CAB_TOP + int(CAB_H * 0.18),
                          CAB_X + CAB_W, CAB_TOP + int(CAB_H * 0.35),
                          fill="#2E86C1", outline="")

        # Door
        d_ml = max(4, int(CAB_W * 0.08)); d_mr = front_w + max(4, int(CAB_W * 0.06))
        door_top = CAB_TOP + int(CAB_H * 0.18); door_bot = CAB_BOT - max(4, int(CAB_H * 0.05))
        dx1 = CAB_X + d_ml; dx2 = CAB_X + CAB_W - d_mr
        cv.create_rectangle(dx1, door_top, dx2, door_bot, fill="", outline="#1B2631", width=2)
        panel_div_y = door_top + int((door_bot - door_top) * 0.52)
        cv.create_line(dx1, panel_div_y, dx2, panel_div_y, fill="#1B2631", width=1)
        hx = dx1 + int((dx2 - dx1) * 0.72); hh = max(4, int(CAB_H * 0.07)); hw = max(6, int(CAB_W * 0.10))
        cv.create_rectangle(hx, panel_div_y - hh, hx + hw, panel_div_y, fill="#7F8C8D", outline="#5D6D7E")

        # Side window
        wm_l = d_ml + max(4, int(CAB_W * 0.05)); wm_r = d_mr + max(4, int(CAB_W * 0.05))
        wy = CAB_TOP + int(CAB_H * 0.08); wb = panel_div_y - max(3, int(CAB_H * 0.04))
        wx1 = CAB_X + wm_l; wx2 = CAB_X + CAB_W - wm_r
        cv.create_rectangle(wx1 - 2, wy - 2, wx2 + 2, wb + 2, fill="#1B2631", outline="")
        cv.create_rectangle(wx1, wy, wx2, wb, fill="#1A5276", outline="#0E3850", width=1)
        cv.create_polygon(wx1 + 2, wy + 2, wx1 + max(4, int((wx2 - wx1) * 0.45)), wy + 2,
                          wx1 + 2, wy + max(4, int((wb - wy) * 0.55)), fill="#2E86C1", outline="")
        if wx2 - wx1 > 20:
            cv.create_line(wx1 + int((wx2 - wx1) * 0.28), wy,
                           wx1 + int((wx2 - wx1) * 0.28), wb, fill="#0E3850", width=2)

        # Mirror
        mirror_x = CAB_X + CAB_W + max(3, int(CAB_W * 0.06))
        mirror_y = wy + max(2, int((wb - wy) * 0.15))
        mirror_w = max(6, int(CAB_W * 0.12)); mirror_h = max(4, int((wb - wy) * 0.3))
        cv.create_rectangle(mirror_x - 1, mirror_y - 1, mirror_x + mirror_w + 1, mirror_y + mirror_h + 1,
                            fill="#2C3E50", outline="")
        cv.create_rectangle(mirror_x, mirror_y, mirror_x + mirror_w, mirror_y + mirror_h,
                            fill="#7F8C8D", outline="")

        # Exhaust stack
        exh_x = CAB_X + CAB_W - max(6, int(CAB_W * 0.08))
        exh_h = max(16, int(CAB_H * 0.55))
        cv.create_rectangle(exh_x - 3, CAB_TOP - exh_h, exh_x + 3, CAB_TOP + 5,
                            fill="#4A5568", outline="#2C3E50", width=1)
        cv.create_oval(exh_x - 5, CAB_TOP - exh_h - 5, exh_x + 5, CAB_TOP - exh_h + 3,
                       fill="#3D4B56", outline="#2C3E50")
        for sm_j in range(3):
            sm_r2 = max(3, 4 * (sm_j + 1))
            sm_x2 = exh_x + sm_j * 2
            sm_y2 = CAB_TOP - exh_h - 6 - sm_j * max(4, exh_h // 7)
            cv.create_oval(sm_x2 - sm_r2, sm_y2 - sm_r2, sm_x2 + sm_r2, sm_y2 + sm_r2,
                           fill=["#808B96", "#AAB7B8", "#BDC3C7"][sm_j], outline="")

        # Grille and headlight
        grill_y = CAB_BOT - max(16, int(CAB_H * 0.26))
        grill_h = max(10, int(CAB_H * 0.24)); grill_w = max(8, int(CAB_W * 0.28))
        grill_x = CAB_X - grill_w
        cv.create_rectangle(grill_x, grill_y, CAB_X, grill_y + grill_h,
                            fill="#BDC3C7", outline="#7F8C8D", width=1)
        for gi in range(max(3, grill_h // 5)):
            gy2 = grill_y + 2 + gi * grill_h // max(3, grill_h // 5)
            cv.create_rectangle(grill_x + 2, gy2, CAB_X - 2, gy2 + max(1, grill_h // max(3, grill_h // 5) - 2),
                                fill="#2C3E50", outline="")
        hl_r = max(5, int(grill_h * 0.38))
        hl_cx = grill_x + grill_w // 2; hl_cy = grill_y + grill_h - hl_r - 3
        cv.create_oval(hl_cx - hl_r, hl_cy - hl_r, hl_cx + hl_r, hl_cy + hl_r,
                       fill="#F9E79F", outline="#BDC3C7")
        cv.create_oval(hl_cx - hl_r // 2, hl_cy - hl_r // 2, hl_cx + hl_r // 2, hl_cy + hl_r // 2,
                       fill="#FEFEFE", outline="")

        # ── Landing gear (near front of trailer) ─────────────────────────
        lg_x = CIS_X + max(8, int(CIS_W * 0.05))
        lg_h = GROUND_Y - CISTERN_BOT - ch_h
        cv.create_line(lg_x, CISTERN_BOT + ch_h, lg_x, GROUND_Y, fill="#4A5568", width=2)
        cv.create_line(lg_x + 6, CISTERN_BOT + ch_h, lg_x + 6, GROUND_Y, fill="#4A5568", width=2)
        cv.create_rectangle(lg_x - 2, GROUND_Y - 4, lg_x + 8, GROUND_Y, fill="#5D6D7E", outline="#4A5568")

        # ── ICC rear bumper bar ───────────────────────────────────────────
        rear_x = CIS_X + CIS_W + max(2, int(W * 0.005))
        cv.create_rectangle(rear_x, CISTERN_BOT + ch_h - 2, rear_x + max(4, int(W * 0.01)),
                            GROUND_Y - 2, fill="#4A5568", outline="#2C3E50")
        # Reflective tape
        cv.create_rectangle(rear_x, GROUND_Y - max(6, int(H * 0.04)),
                            rear_x + max(4, int(W * 0.01)), GROUND_Y - 2,
                            fill="#E74C3C", outline="")

        # ── WHEELS with hub detail ────────────────────────────────────────
        AXLE_Y = GROUND_Y - AXLE_R
        wheel_positions = [CAB_X + int(CAB_W * 0.62)]
        # Tandem rear axles
        rear_center = CIS_X + int(CIS_W * 0.70)
        wheel_positions.append(rear_center - max(5, AXLE_R // 4))
        wheel_positions.append(rear_center + max(5, AXLE_R // 4))
        if CIS_W > 120:
            wheel_positions.insert(1, CIS_X + int(CIS_W * 0.45))

        for wx in wheel_positions:
            cv.create_oval(wx - AXLE_R, AXLE_Y - AXLE_R, wx + AXLE_R, AXLE_Y + AXLE_R,
                           fill="#1B2631", outline="#0E1A27", width=2)
            rim_r = max(4, int(AXLE_R * 0.55))
            cv.create_oval(wx - rim_r, AXLE_Y - rim_r, wx + rim_r, AXLE_Y + rim_r,
                           fill="#7F8C8D", outline="#5D6D7E", width=1)
            for angle_r in [0, 60, 120, 180, 240, 300]:
                ar = math.radians(angle_r)
                cv.create_line(wx + int(rim_r * 0.25 * math.cos(ar)),
                               AXLE_Y + int(rim_r * 0.25 * math.sin(ar)),
                               wx + int(rim_r * 0.85 * math.cos(ar)),
                               AXLE_Y + int(rim_r * 0.85 * math.sin(ar)),
                               fill="#5D6D7E", width=1)
            cv.create_oval(wx - 3, AXLE_Y - 3, wx + 3, AXLE_Y + 3, fill="#4A5568", outline="")

        for wx in wheel_positions:
            cv.create_arc(wx - AXLE_R - 2, AXLE_Y - AXLE_R - 2, wx + AXLE_R + 2, AXLE_Y + AXLE_R + 2,
                          start=30, extent=120, style="arc", outline="#5D6D7E", width=3)


    def _draw_pressure_truck(self, cv, x, y, W, H, etapa_key, tank_names):
        """MC-331 LPG pressure truck -- circular cross-section, hemispherical heads, WHITE, single vessel."""
        import math
        if tank_names is None: tank_names = self.lista_tanques
        if H < 80 or W < 180: return

        n   = max(len(tank_names), 1)
        fs  = min(9, max(5, int(H * 0.032))); fss = min(7, max(4, int(H * 0.026)))
        FT  = ("Arial", fs); FTS = ("Arial", fss)
        patente = self.get_var("car_patente").get() or "GAS/GLP"

        GROUND_Y    = y + H - max(14, int(H * 0.08))
        AXLE_R      = max(10, int(H * 0.11))
        CISTERN_BOT = GROUND_Y - AXLE_R * 2 - 4
        # Circular cross-section: height = diameter -> perfectly round barrel in side view
        CISTERN_H   = max(30, int(H * 0.40))
        CISTERN_TOP = CISTERN_BOT - CISTERN_H
        DOME_R      = int(CISTERN_H * 0.52)  # hemispherical dished end caps
        PAD_L       = max(8, int(W * 0.03))
        CAB_W       = max(50, int(W * 0.16))
        CAB_X       = x + PAD_L
        CIS_X       = CAB_X + CAB_W + max(4, int(W * 0.01))
        CIS_W       = W - PAD_L - CAB_W - max(4, int(W * 0.01)) - max(8, int(W * 0.03))
        COMP_W      = CIS_W // n

        # ── Background ─────────────────────────────────────────────────────
        sky_cols = ["#D6EAF8", "#E0EDF6", "#EAF0F4", "#F5F7FA"]
        sky_h = GROUND_Y - y
        for si, sc in enumerate(sky_cols):
            sy1 = y + int(sky_h * si / len(sky_cols))
            sy2 = y + int(sky_h * (si + 1) / len(sky_cols))
            cv.create_rectangle(x, sy1, x + W, sy2, fill=sc, outline="")
        cv.create_rectangle(x, y, x + W, y + H, fill="", outline="#D4AC0D", width=2)
        gnd_cols = ["#808B96", "#747D88", "#686F7A", "#5C636C"]
        gnd_h = y + H - GROUND_Y
        for gi, gc in enumerate(gnd_cols):
            gy1 = GROUND_Y + int(gnd_h * gi / len(gnd_cols))
            gy2 = GROUND_Y + int(gnd_h * (gi + 1) / len(gnd_cols))
            cv.create_rectangle(x, gy1, x + W, gy2, fill=gc, outline="")
        cv.create_line(x, GROUND_Y, x + W, GROUND_Y, fill="#A0A8B0", width=1)

        cv.create_text(x + W // 2 + 1, y + max(10, int(H * 0.06)) + 1,
                       text=f"CAMION GAS/GLP (MC-331)  --  {patente}", font=FT, fill="#C0A050")
        cv.create_text(x + W // 2, y + max(10, int(H * 0.06)),
                       text=f"CAMION GAS/GLP (MC-331)  --  {patente}", font=FT, fill="#784212")

        cv.create_oval(CAB_X, GROUND_Y - 3, CIS_X + CIS_W, GROUND_Y + 5, fill="#404850", outline="")

        # ── Chassis frame ─────────────────────────────────────────────────
        ch_h = max(5, int(H * 0.04))
        cv.create_rectangle(CAB_X + int(CAB_W * 0.3), CISTERN_BOT, CIS_X + CIS_W, CISTERN_BOT + ch_h,
                            fill="#2C3E50", outline="#1B2631", width=1)
        for ci in range(6):
            cx_tr = CIS_X + int(CIS_W * ci / 5)
            cv.create_rectangle(cx_tr - 2, CISTERN_BOT - max(3, int(H * 0.02)), cx_tr + 2, CISTERN_BOT + ch_h,
                                fill="#1B2631", outline="")

        # ── SINGLE PRESSURE VESSEL (circular, WHITE, no compartment rings) ─
        # The entire barrel is one single vessel - iterate tanks for data only
        body_x1 = CIS_X + DOME_R
        body_x2 = CIS_X + CIS_W - DOME_R

        # WHITE vessel body (mandatory for LPG) with cylindrical gradient
        WHITE_GRAD = ["#E8E8E8", "#ECECEC", "#F0F0F0", "#F4F4F4", "#F8F8F8",
                      "#FCFCFC", "#FEFEFE", "#FCFCFC", "#F4F4F4", "#ECECEC"]
        for ki, kc in enumerate(WHITE_GRAD):
            ky = CISTERN_TOP + int(ki * CISTERN_H / len(WHITE_GRAD))
            kh = CISTERN_H // len(WHITE_GRAD) + 1
            cv.create_rectangle(body_x1, ky, body_x2, ky + kh, fill=kc, outline="")
        # Specular highlight
        cv.create_rectangle(body_x1 + 4, CISTERN_TOP + max(3, CISTERN_H // 5),
                            body_x2 - 4, CISTERN_TOP + max(5, CISTERN_H // 4),
                            fill="#FFFFFF", outline="")
        cv.create_rectangle(body_x1, CISTERN_TOP, body_x2, CISTERN_BOT,
                            fill="", outline="#C0C0C0", width=2)

        # Hemispherical dished end caps (clearly curved, NOT flat)
        for target, is_left in [(CIS_X, True), (CIS_X + CIS_W - 2 * DOME_R, False)]:
            hemi_colors = ["#D0D0D0", "#D8D8D8", "#E0E0E0", "#E8E8E8", "#F0F0F0",
                           "#F4F4F4", "#F8F8F8", "#FCFCFC", "#FEFEFE", "#FFFFFF"]
            if not is_left:
                hemi_colors = list(reversed(hemi_colors))
            for ki, kc in enumerate(hemi_colors):
                margin = ki * max(1, DOME_R // 12)
                r = DOME_R - margin
                if r > 2:
                    cv.create_oval(target + margin, CISTERN_TOP + margin,
                                   target + 2 * DOME_R - margin, CISTERN_BOT - margin,
                                   fill=kc, outline="")
            cv.create_oval(target, CISTERN_TOP, target + 2 * DOME_R, CISTERN_BOT,
                           fill="", outline="#C0C0C0", width=2)

        # ── Sunshield (half-cylinder shade over top half) ─────────────────
        ss_h = max(4, int(CISTERN_H * 0.08))
        ss_overhang = max(4, int(CIS_W * 0.02))
        cv.create_rectangle(CIS_X + DOME_R - ss_overhang, CISTERN_TOP - ss_h,
                            CIS_X + CIS_W - DOME_R + ss_overhang, CISTERN_TOP,
                            fill="#E0E0E0", outline="#B0B0B0", width=1)
        # Sunshield supports
        for si in range(3):
            sx2 = body_x1 + int((body_x2 - body_x1) * (si + 1) / 4)
            cv.create_line(sx2, CISTERN_TOP - ss_h, sx2, CISTERN_TOP - ss_h - 2,
                           fill="#B0B0B0", width=1)
            cv.create_line(sx2, CISTERN_TOP - ss_h - 2, sx2 + 4, CISTERN_TOP + 2,
                           fill="#B0B0B0", width=1)

        # ── Fill + labels per tank ────────────────────────────────────────
        for i, tn in enumerate(tank_names[:n]):
            # Compute fill zone within the single vessel
            seg_x1 = body_x1 + int((body_x2 - body_x1) * i / n) + 2
            seg_x2 = body_x1 + int((body_x2 - body_x1) * (i + 1) / n) - 2
            try:
                vol = self.parse_float(self.get_var(f"{etapa_key}_{tn}_vol_nat_prod").get() or "0")
                cap = self.parse_float(self.get_var(f"{etapa_key}_{tn}_alt_ref").get() or "1")
                vp  = min(max(vol / (cap if cap > 0 else 1), 0), 1)
            except: vp = 0.0
            pnm  = self.get_var(f"{etapa_key}_{tn}_prod_name").get() if etapa_key else ""
            pres = self.get_var(f"{etapa_key}_{tn}_esf_pres").get() if etapa_key else ""

            if vp > 0.02:
                fill_top = int(CISTERN_BOT - CISTERN_H * vp)
                cv.create_rectangle(seg_x1, fill_top, seg_x2, CISTERN_BOT - 2,
                                    fill="#FF6B35", outline="")
                cv.create_rectangle(seg_x1, fill_top, seg_x2,
                                    fill_top + max(3, int(CISTERN_H * 0.06)),
                                    fill="#FFB088", outline="")
                cv.create_line(seg_x1, fill_top, seg_x2, fill_top, fill="#CC4A1A", width=2)

            mid_y = (CISTERN_TOP + CISTERN_BOT) // 2
            seg_cx = (seg_x1 + seg_x2) // 2
            short = tn.replace("COMPARTIMENTO ", "C.").strip()[:6]
            txt_c = "#1B2631" if vp < 0.4 else "white"
            cv.create_text(seg_cx, CISTERN_TOP + max(10, int(CISTERN_H * 0.18)),
                           text=short, font=FT, fill="#4A5568")
            if vp > 0:
                cv.create_text(seg_cx, mid_y, text=f"{vp*100:.0f}%", font=FTS, fill=txt_c)
            if pnm:
                cv.create_text(seg_cx, CISTERN_BOT - max(8, int(CISTERN_H * 0.15)),
                               text=pnm[:10], font=FTS, fill=txt_c)

        # Redraw vessel outline over fill
        cv.create_rectangle(body_x1, CISTERN_TOP, body_x2, CISTERN_BOT,
                            fill="", outline="#C0C0C0", width=2)

        # ── Safety relief valves on top ───────────────────────────────────
        for vvi in range(2):
            vvx = body_x1 + int((body_x2 - body_x1) * (vvi + 1) / 3)
            cv.create_rectangle(vvx - 4, CISTERN_TOP - ss_h - 10, vvx + 4, CISTERN_TOP - ss_h,
                                fill="#E74C3C", outline="#C0392B", width=1)
            cv.create_rectangle(vvx - 6, CISTERN_TOP - ss_h - 14, vvx + 6, CISTERN_TOP - ss_h - 10,
                                fill="#C0392B", outline="")
            cv.create_rectangle(vvx - 1, CISTERN_TOP - ss_h - 20, vvx + 1, CISTERN_TOP - ss_h - 14,
                                fill="#5D6D7E", outline="")

        # ── Hazmat diamond: RED (flammable gas, Class 2.1) ────────────────
        plac_x = body_x2 - max(15, int(CIS_W * 0.06))
        plac_y = (CISTERN_TOP + CISTERN_BOT) // 2
        plac_s = max(8, int(CISTERN_H * 0.14))
        cv.create_polygon(plac_x, plac_y - plac_s, plac_x + plac_s, plac_y,
                          plac_x, plac_y + plac_s, plac_x - plac_s, plac_y,
                          fill="#E74C3C", outline="#922B21", width=1)
        cv.create_text(plac_x, plac_y - 2, text="2", font=("Arial", max(3, fss - 1), "bold"), fill="white")
        cv.create_text(plac_x, plac_y + plac_s - 4, text="GLP",
                       font=("Arial", max(2, fss - 2)), fill="white")

        # ── Heavy rear protective frame ───────────────────────────────────
        rear_x = CIS_X + CIS_W + max(2, int(W * 0.005))
        frame_w = max(6, int(W * 0.015))
        cv.create_rectangle(rear_x, CISTERN_TOP + int(CISTERN_H * 0.2),
                            rear_x + frame_w, CISTERN_BOT + ch_h,
                            fill="#2C3E50", outline="#1B2631", width=1)
        # Horizontal bars
        for bfi in range(3):
            bfy = CISTERN_TOP + int(CISTERN_H * (0.25 + 0.25 * bfi))
            cv.create_rectangle(rear_x, bfy - 2, rear_x + frame_w, bfy + 2,
                                fill="#4A5568", outline="")
        cv.create_rectangle(rear_x, CISTERN_BOT + ch_h - 2, rear_x + frame_w, GROUND_Y - 2,
                            fill="#4A5568", outline="#2C3E50")
        cv.create_rectangle(rear_x, CISTERN_BOT + ch_h, rear_x + 3, CISTERN_BOT + ch_h + 6,
                            fill="#E74C3C", outline="")

        # ── Cab (same chassis/cab as liquid truck with different color) ───
        CAB_BOT = CISTERN_BOT; CAB_H = CISTERN_H + max(8, int(CISTERN_H * 0.22))
        CAB_TOP = CAB_BOT - CAB_H
        cab_slope = max(4, int(CAB_W * 0.12))
        cab_pts = [CAB_X, CAB_BOT, CAB_X + CAB_W, CAB_BOT,
                   CAB_X + CAB_W, CAB_TOP + int(CAB_H * 0.18),
                   CAB_X + CAB_W - cab_slope, CAB_TOP,
                   CAB_X + cab_slope, CAB_TOP,
                   CAB_X, CAB_TOP + int(CAB_H * 0.18)]
        cv.create_polygon([p + 3 for p in cab_pts], fill="#404850", outline="")
        cv.create_polygon(cab_pts, fill="#2E4053", outline="#1B2631", width=2)
        cv.create_rectangle(CAB_X, CAB_TOP + 1, CAB_X + CAB_W, CAB_TOP + max(4, int(CAB_H * 0.08)),
                            fill="#3D5166", outline="")
        sp_h = max(6, int(CAB_H * 0.14))
        cv.create_polygon(CAB_X + int(CAB_W * 0.1), CAB_TOP,
                          CAB_X + int(CAB_W * 0.9), CAB_TOP,
                          CAB_X + int(CAB_W * 0.9), CAB_TOP - sp_h + 3,
                          CAB_X + int(CAB_W * 0.1), CAB_TOP - sp_h,
                          fill="#3D5166", outline="#2C3E50")
        # Window
        front_w = max(4, int(CAB_W * 0.12))
        cv.create_rectangle(CAB_X + CAB_W - front_w, CAB_TOP + int(CAB_H * 0.18),
                            CAB_X + CAB_W, CAB_TOP + int(CAB_H * 0.72),
                            fill="#AED6F1", outline="#2C3E50", width=1)
        d_ml = max(4, int(CAB_W * 0.08)); d_mr = front_w + max(4, int(CAB_W * 0.06))
        door_top = CAB_TOP + int(CAB_H * 0.18); door_bot = CAB_BOT - max(4, int(CAB_H * 0.05))
        dx1 = CAB_X + d_ml; dx2 = CAB_X + CAB_W - d_mr
        cv.create_rectangle(dx1, door_top, dx2, door_bot, fill="", outline="#1B2631", width=2)
        panel_y = door_top + int((door_bot - door_top) * 0.52)
        cv.create_line(dx1, panel_y, dx2, panel_y, fill="#1B2631", width=1)
        wm_l = d_ml + max(4, int(CAB_W * 0.05)); wm_r = d_mr + max(4, int(CAB_W * 0.05))
        wy = CAB_TOP + int(CAB_H * 0.08); wb = panel_y - max(3, int(CAB_H * 0.04))
        wx1 = CAB_X + wm_l; wx2 = CAB_X + CAB_W - wm_r
        cv.create_rectangle(wx1 - 2, wy - 2, wx2 + 2, wb + 2, fill="#1B2631", outline="")
        cv.create_rectangle(wx1, wy, wx2, wb, fill="#AED6F1", outline="#2C3E50", width=1)

        # Grille and headlight
        grill_h = max(10, int(CAB_H * 0.22)); grill_w = max(6, int(CAB_W * 0.25))
        grill_y = CAB_BOT - max(16, int(CAB_H * 0.24))
        cv.create_rectangle(CAB_X - grill_w, grill_y, CAB_X, grill_y + grill_h,
                            fill="#BDC3C7", outline="#7F8C8D", width=1)
        hl_r = max(4, int(grill_h * 0.35))
        hl_cx = CAB_X - grill_w // 2; hl_cy = grill_y + grill_h - hl_r - 3
        cv.create_oval(hl_cx - hl_r, hl_cy - hl_r, hl_cx + hl_r, hl_cy + hl_r,
                       fill="#F9E79F", outline="#D4AC0D")

        # Exhaust
        exh_x = CAB_X + CAB_W - max(6, int(CAB_W * 0.08))
        cv.create_rectangle(exh_x - 2, CAB_TOP - max(14, int(CAB_H * 0.45)), exh_x + 2, CAB_TOP + 3,
                            fill="#4A5568", outline="#2C3E50")

        # ── Wheels ────────────────────────────────────────────────────────
        AXLE_Y = GROUND_Y - AXLE_R
        wps = [CAB_X + int(CAB_W * 0.62)]
        rc = CIS_X + int(CIS_W * 0.72)
        wps += [rc - max(4, AXLE_R // 4), rc + max(4, AXLE_R // 4)]

        for wx in wps:
            cv.create_oval(wx - AXLE_R + 2, AXLE_Y - AXLE_R + 2,
                           wx + AXLE_R + 2, AXLE_Y + AXLE_R + 2,
                           fill="#0A0E12", outline="")
            cv.create_oval(wx - AXLE_R, AXLE_Y - AXLE_R, wx + AXLE_R, AXLE_Y + AXLE_R,
                           fill="#1B2631", outline="#0E1A27", width=2)
            rim_r = max(4, int(AXLE_R * 0.55))
            cv.create_oval(wx - rim_r, AXLE_Y - rim_r, wx + rim_r, AXLE_Y + rim_r,
                           fill="#808B96", outline="#5D6D7E", width=1)
            for angle_r in [0, 60, 120, 180, 240, 300]:
                ar = math.radians(angle_r)
                cv.create_line(wx + int(rim_r * 0.25 * math.cos(ar)),
                               AXLE_Y + int(rim_r * 0.25 * math.sin(ar)),
                               wx + int(rim_r * 0.85 * math.cos(ar)),
                               AXLE_Y + int(rim_r * 0.85 * math.sin(ar)),
                               fill="#5D6D7E", width=1)
            cv.create_oval(wx - 3, AXLE_Y - 3, wx + 3, AXLE_Y + 3, fill="#4A5568", outline="")


    def _draw_moss_vessel(self, cv, x, y, W, H, side_label, etapa_key, tank_names, carb_names):
        """MOSS type LNG carrier -- 4 large spheres DOMINATING profile above deck, stern superstructure."""
        import math
        if tank_names is None: tank_names = self.lista_tanques
        if H < 60 or W < 120: return

        def bez(p0, p1, p2, p3, n_pts=20):
            pts = []
            for i in range(n_pts + 1):
                t = i / n_pts; u = 1 - t
                pts += [u**3*p0[0]+3*u**2*t*p1[0]+3*u*t**2*p2[0]+t**3*p3[0],
                        u**3*p0[1]+3*u**2*t*p1[1]+3*u*t**2*p2[1]+t**3*p3[1]]
            return pts

        fs  = min(9, max(5, int(H*0.034)));  fss = min(7, max(4, int(H*0.026)))
        FT  = ("Arial", fs);  FTS = ("Arial", fss)

        TITLE_H = max(14, int(H*0.07))
        SPHERE_ZONE = int(H*0.55)   # Spheres DOMINATE the profile
        HULL_ZONE   = H - SPHERE_ZONE - TITLE_H - 4

        buque = self.get_var("car_buque").get() or "BUQUE GASERO"

        # ── Ocean/sky background ──────────────────────────────────────────
        # Sky gradient
        sky_cols = ["#87CEEB","#9DD5EE","#B3DCF1","#C9E4F4","#DFF0FA"]
        sky_h = int(H * 0.55)
        for si, sc in enumerate(sky_cols):
            sy1 = y + int(sky_h * si / len(sky_cols))
            sy2 = y + int(sky_h * (si + 1) / len(sky_cols))
            cv.create_rectangle(x, sy1, x + W, sy2, fill=sc, outline="")
        # Ocean water
        ocean_cols = ["#1A5276","#1F618D","#2471A3","#2980B9","#2E86C1","#3498DB"]
        ocean_top = y + sky_h
        ocean_h = y + H - ocean_top
        for oi, oc in enumerate(ocean_cols):
            oy1 = ocean_top + int(ocean_h * oi / len(ocean_cols))
            oy2 = ocean_top + int(ocean_h * (oi + 1) / len(ocean_cols))
            cv.create_rectangle(x, oy1, x + W, oy2, fill=oc, outline="")
        # Wave pattern on water surface
        for wi in range(0, W, max(20, W // 15)):
            wx = x + wi
            wave_pts = bez((wx, ocean_top), (wx + 5, ocean_top - 2), (wx + 10, ocean_top + 2), (wx + 15, ocean_top), 8)
            for wj in range(0, len(wave_pts) - 2, 2):
                cv.create_line(int(wave_pts[wj]), int(wave_pts[wj+1]),
                               int(wave_pts[wj+2]), int(wave_pts[wj+3]),
                               fill="#5DADE2", width=1)
        cv.create_rectangle(x, y, x + W, y + H, fill="", outline="#1B3A5C", width=2)

        # Title with shadow
        cv.create_text(x+W//2+1, y+TITLE_H//2+1, text=f"BUQUE GASERO/GLP  --  {buque}", font=FT, fill="#A0B0C0")
        cv.create_text(x+W//2, y+TITLE_H//2, text=f"BUQUE GASERO/GLP  --  {buque}", font=FT, fill="#1B3A5C")

        # ── Hull ──────────────────────────────────────────────────────────
        DECK_Y  = y + TITLE_H + SPHERE_ZONE + 4
        KEEL_Y  = y + H - 4
        HULL_H  = KEEL_Y - DECK_Y
        WL_Y    = DECK_Y + int(HULL_H * 0.55)
        BOW_X   = x + W - max(14, int(W * 0.03))
        STERN_X = x + max(14, int(W * 0.03))
        proa    = int(W * 0.045)
        CX      = x + W // 2

        # Freeboard (dark hull above waterline)
        bow_top = bez((BOW_X, DECK_Y), (BOW_X + proa, DECK_Y),
                      (BOW_X + proa, WL_Y - int(HULL_H * 0.1)), (BOW_X + proa // 2, WL_Y))
        top_poly = [STERN_X - 4, DECK_Y, BOW_X, DECK_Y] + bow_top + [STERN_X - 4, WL_Y]
        # Hull gradient (dark to slightly lighter)
        cv.create_polygon(top_poly, fill="#2C3E50", outline="")
        # Lighter band near deck
        for hi in range(3):
            hf = hi / 3
            hy = DECK_Y + int((WL_Y - DECK_Y) * hf * 0.3)
            hh = max(2, int(HULL_H * 0.04))
            hull_band_col = ["#34495E","#3D5166","#2C3E50"][hi]
            cv.create_rectangle(STERN_X, hy, BOW_X, hy + hh, fill=hull_band_col, outline="")
        cv.create_polygon(top_poly, fill="", outline="#1B2631", width=1)

        # Underwater hull (anti-fouling red)
        bow_bot = bez((BOW_X + proa // 2, WL_Y),
                      (BOW_X + proa // 2, KEEL_Y - int(HULL_H * 0.05)),
                      (BOW_X - int(W * 0.03), KEEL_Y), (CX, KEEL_Y + int(HULL_H * 0.03)))
        stern_bot = bez((CX, KEEL_Y + int(HULL_H * 0.03)),
                        (STERN_X + int(W * 0.04), KEEL_Y),
                        (STERN_X, KEEL_Y - int(HULL_H * 0.02)), (STERN_X - 4, WL_Y))
        bot_poly = [STERN_X - 4, WL_Y] + bow_bot + stern_bot
        cv.create_polygon(bot_poly, fill="#7B241C", outline="#5B1A14", width=1)
        # Boot-topping stripe
        cv.create_line(STERN_X - 10, WL_Y, BOW_X + proa // 2, WL_Y, fill="#1A1A1A", width=2)
        cv.create_line(STERN_X - 10, WL_Y + 2, BOW_X + proa // 2, WL_Y + 2, fill="#2E86C1", width=1, dash=(5, 3))
        cv.create_text(x + W - 25, WL_Y - 7, text="WL", font=FTS, fill="#2E86C1")

        # Deck with plating detail
        deck_x1 = STERN_X + int(W * 0.02)
        deck_x2 = BOW_X - int(W * 0.02)
        cv.create_rectangle(deck_x1, DECK_Y - 5, deck_x2, DECK_Y, fill="#3D4B56", outline="")
        cv.create_line(deck_x1, DECK_Y - 5, deck_x2, DECK_Y - 5, fill="#4A5568", width=1)
        # Deck plating lines
        for di in range(deck_x1, deck_x2, max(12, (deck_x2 - deck_x1) // 20)):
            cv.create_line(di, DECK_Y - 4, di, DECK_Y - 1, fill="#4A5568", width=1)

        # Bulwark (raised edge of deck)
        cv.create_rectangle(deck_x1, DECK_Y - 7, deck_x2, DECK_Y - 5, fill="#4A5568", outline="#3D4B56")

        # ── Superstructure (aft) ──────────────────────────────────────────
        SUP_W = max(50, int(W * 0.12))
        SUP_H = max(30, int(SPHERE_ZONE * 0.60))
        SUP_X = STERN_X + int(W * 0.02)

        # Multiple deck levels
        n_levels = 4
        for li in range(n_levels):
            lw = SUP_W - li * max(3, int(SUP_W * 0.05))
            lh = SUP_H // n_levels
            ly = DECK_Y - (li + 1) * lh
            lx = SUP_X + li * max(1, int(SUP_W * 0.025))
            col = ["#D5D8DC","#E0E3E7","#E5E8EC","#EBEDF0"][li]
            cv.create_rectangle(lx, ly, lx + lw, ly + lh, fill=col, outline="#ABB2B9", width=1)
            # Windows
            win_y = ly + max(2, lh // 4)
            win_h = max(3, lh // 3)
            for wi in range(max(2, lw // 10)):
                win_x = lx + 4 + wi * max(6, lw // max(2, lw // 10))
                if win_x + 4 < lx + lw - 4:
                    cv.create_rectangle(win_x, win_y, win_x + max(3, lw // 12), win_y + win_h,
                                        fill="#2E86C1", outline="#1A5276")
        # Bridge wing
        bridge_y = DECK_Y - SUP_H
        cv.create_rectangle(SUP_X - max(4, int(SUP_W * 0.08)), bridge_y,
                            SUP_X + SUP_W + max(4, int(SUP_W * 0.08)), bridge_y + max(4, SUP_H // n_levels),
                            fill="#E5E7E9", outline="#ABB2B9")

        # Funnel/chimney
        chim_w = max(8, int(SUP_W * 0.25))
        chim_h = max(14, int(SUP_H * 0.45))
        chim_x = SUP_X + SUP_W // 2 - chim_w // 2
        # Funnel body gradient
        chim_grads = ["#C0392B","#CD4535","#DA5040","#C0392B","#A82E22"]
        for ci2, cc in enumerate(chim_grads):
            cw1 = chim_x + int(chim_w * ci2 / len(chim_grads))
            cw2 = chim_x + int(chim_w * (ci2 + 1) / len(chim_grads))
            cv.create_rectangle(cw1, bridge_y - chim_h, cw2, bridge_y,
                                fill=cc, outline="")
        cv.create_rectangle(chim_x, bridge_y - chim_h, chim_x + chim_w, bridge_y,
                            fill="", outline="#922B21", width=1)
        # Funnel top (dark opening)
        cv.create_oval(chim_x, bridge_y - chim_h - 3, chim_x + chim_w, bridge_y - chim_h + 3,
                       fill="#4A1A14", outline="#922B21")
        # Company mark on funnel
        cv.create_rectangle(chim_x + 2, bridge_y - int(chim_h * 0.5),
                            chim_x + chim_w - 2, bridge_y - int(chim_h * 0.3),
                            fill="#F4D03F", outline="")

        # ── Spheres on deck ───────────────────────────────────────────────
        n = max(len(tank_names), 1)
        sph_start = SUP_X + SUP_W + max(8, int(W * 0.02))
        sph_end   = BOW_X - max(20, int(W * 0.05))
        sph_total = sph_end - sph_start
        sph_each  = sph_total // n
        SPH_R     = min(int(SPHERE_ZONE * 0.42), sph_each // 2 - 6)
        SPH_R     = max(18, SPH_R)
        SPH_BASE  = int(SPH_R * 0.22)

        for i, tn in enumerate(tank_names[:n]):
            cx_s = sph_start + i * sph_each + sph_each // 2
            cy_s = DECK_Y - SPH_BASE - SPH_R

            # Sphere shadow on deck
            sh_w = int(SPH_R * 0.7)
            cv.create_oval(cx_s - sh_w, DECK_Y - 5, cx_s + sh_w, DECK_Y + 1,
                           fill="#2A3540", outline="")

            # Skirt (cylindrical support)
            skirt_w = int(SPH_R * 0.5)
            skirt_top = DECK_Y - SPH_BASE - max(6, int(SPH_R * 0.25))
            # Skirt gradient
            cv.create_rectangle(cx_s - skirt_w, skirt_top, cx_s + skirt_w, DECK_Y,
                                fill="#4A5568", outline="")
            cv.create_line(cx_s - skirt_w + 2, skirt_top, cx_s - skirt_w + 2, DECK_Y,
                           fill="#5D6D7E", width=1)
            cv.create_rectangle(cx_s - skirt_w, skirt_top, cx_s + skirt_w, DECK_Y,
                                fill="", outline="#3D4456", width=1)
            # Skirt ventilation openings
            for vi in range(3):
                vx = cx_s - skirt_w + int(skirt_w * 2 * (vi + 1) / 4)
                cv.create_rectangle(vx - 2, DECK_Y - max(4, int(SPH_R * 0.06)),
                                    vx + 2, DECK_Y - 2, fill="#2C3E50", outline="")

            # Containment cover (dome housing visible above sphere equator)
            cover_r = int(SPH_R * 1.08)
            cover_h = int(SPH_R * 0.25)
            cv.create_arc(cx_s - cover_r, cy_s - cover_h, cx_s + cover_r, cy_s + cover_h,
                          start=0, extent=180, style="chord", fill="#D5D8DC", outline="#ABB2B9")

            # Sphere with 16-layer radial gradient
            sphere_gradient = [
                "#2D3E4C","#354A5A","#3D5668","#456276","#506E84","#5C7A92",
                "#6B88A0","#7E98AE","#94ABBC","#AAC0CC","#BFD4DC","#D4E4EC",
                "#E4EFF4","#EFF6F9","#F6FBFC","#FAFEFE",
            ]
            n_layers = len(sphere_gradient)
            for k in range(n_layers - 1, -1, -1):
                frac = k / (n_layers - 1)
                r_k = int(SPH_R * (0.10 + frac * 0.90))
                if r_k < 1: continue
                max_off = int(SPH_R * 0.30)
                off_x = int(max_off * (1 - frac) * 0.55)
                off_y = int(max_off * (1 - frac) * 0.80)
                sc3 = sphere_gradient[k]
                cv.create_oval(cx_s - r_k - off_x // 2, cy_s - r_k - off_y // 2,
                               cx_s + r_k - off_x // 2, cy_s + r_k - off_y // 2,
                               fill=sc3, outline="")

            # Weld seams on sphere
            cv.create_oval(cx_s - SPH_R + 2, cy_s - 1, cx_s + SPH_R - 2, cy_s + 1,
                           fill="", outline="#4A6070", width=1)
            for m_off in [-int(SPH_R * 0.40), int(SPH_R * 0.40)]:
                m_w = max(2, int(SPH_R * 0.12))
                cv.create_oval(cx_s + m_off - m_w, cy_s - SPH_R + 2,
                               cx_s + m_off + m_w, cy_s + SPH_R - 2,
                               fill="", outline="#4A6070", width=1)

            cv.create_oval(cx_s - SPH_R, cy_s - SPH_R, cx_s + SPH_R, cy_s + SPH_R,
                           fill="", outline="#2C3E50", width=3)

            # Nozzles at top
            for nxi in [-int(SPH_R * 0.3), 0, int(SPH_R * 0.3)]:
                nx2 = cx_s + nxi
                cv.create_rectangle(nx2 - 3, cy_s - SPH_R - 12, nx2 + 3, cy_s - SPH_R,
                                    fill="#4A5568", outline="#2C3E50", width=1)
                cv.create_rectangle(nx2 - 5, cy_s - SPH_R - 15, nx2 + 5, cy_s - SPH_R - 12,
                                    fill="#5D6D7E", outline="#2C3E50")

            # Fill level
            try:
                vol = self.parse_float(self.get_var(f"{etapa_key}_{tn}_vol_nat_prod").get() or "0") if etapa_key and tn else 0
                ref = self.parse_float(self.get_var(f"{etapa_key}_{tn}_alt_ref").get() or "1")
                vp  = min(max(vol / ref, 0), 1) if ref > 0 else 0
            except: vp = 0.0
            prod_fill = self.get_prod_color(tn, etapa_key)[0]

            if vp > 0.02:
                fh2 = int(SPH_R * (2 * vp - 1))
                ftop = cy_s - fh2
                clip_top = max(cy_s - SPH_R + 3, ftop)
                liq_fill = prod_fill if prod_fill not in ("#3498DB","#5DADE2","#2E86C1","#2471A3","#AED6F1","#85C1E9") else "#FF6B35"
                cv.create_rectangle(cx_s - SPH_R + 3, clip_top, cx_s + SPH_R - 3, cy_s + SPH_R - 3,
                                    fill=liq_fill, outline="")
                cv.create_oval(cx_s - SPH_R, cy_s - SPH_R, cx_s + SPH_R, cy_s + SPH_R,
                               fill="", outline="#2C3E50", width=3)
                lv_y2 = clip_top
                hc2 = int((SPH_R**2 - (cy_s - lv_y2)**2)**0.5) if abs(cy_s - lv_y2) < SPH_R else 0
                if hc2 > 0:
                    cv.create_line(cx_s - hc2, lv_y2, cx_s + hc2, lv_y2, fill="#4A4A4A", width=2, dash=(4, 3))
                txt_col_s = self.contrast_text(liq_fill) if vp > 0.5 else "white"
                cv.create_text(cx_s, cy_s, text=f"{vp*100:.0f}%", font=FTS, fill=txt_col_s)

            # Dome/trunk fitting above sphere
            cv.create_rectangle(cx_s - 7, cy_s - SPH_R - 12, cx_s + 7, cy_s - SPH_R,
                                fill="#5D6D7E", outline="#2C3E50")
            cv.create_oval(cx_s - 8, cy_s - SPH_R - 15, cx_s + 8, cy_s - SPH_R - 11,
                           fill="#808B96", outline="#5D6D7E")

            # Label
            tn_s = tn[:6]
            cv.create_text(cx_s, cy_s + SPH_R + SPH_BASE + 6, text=tn_s, font=FTS, fill="#1B3A5C")

        # ── Deck piping (manifold between spheres) ────────────────────────
        if n > 1:
            pipe_y = DECK_Y - 3
            first_cx = sph_start + sph_each // 2
            last_cx  = sph_start + (n - 1) * sph_each + sph_each // 2
            cv.create_rectangle(first_cx, pipe_y - 2, last_cx, pipe_y + 2, fill="#5D6D7E", outline="#4A5568")
            cv.create_line(first_cx, pipe_y - 1, last_cx, pipe_y - 1, fill="#808B96", width=1)

        # Bow mast
        mast_x = BOW_X - max(16, int(W * 0.04))
        mast_top = DECK_Y - max(25, int(SPHERE_ZONE * 0.60))
        cv.create_line(mast_x, DECK_Y - 5, mast_x, mast_top, fill="#626567", width=2)
        # Stays
        cv.create_line(mast_x, mast_top, BOW_X, DECK_Y - 5, fill="#95A5A6", width=1, dash=(3, 4))
        cv.create_line(mast_x, mast_top, mast_x - max(20, int(W * 0.05)), DECK_Y - 5,
                       fill="#95A5A6", width=1, dash=(3, 4))

        # Draft marks at bow
        for di in range(3):
            dm_y = WL_Y + int(di * HULL_H * 0.08)
            cv.create_text(BOW_X + proa // 2 - 4, dm_y, text=str(8 - di),
                           font=("Arial", max(3, fss - 2)), fill="#FDFEFE")

        # Side label
        cv.create_text(x + 12, y + TITLE_H + SPHERE_ZONE // 2, text=side_label,
                       font=FTS, fill="#1B3A5C", angle=90)

    def _draw_membrane_vessel(self, cv, x, y, W, H, side_label, etapa_key, tank_names, carb_names):
        """Membrane/GTT LNG carrier -- FLAT deck, NO spheres, low trunk coamings, clean sleek profile."""
        import math
        if tank_names is None: tank_names = self.lista_tanques
        if H < 60 or W < 120: return

        def bez(p0, p1, p2, p3, n_pts=20):
            pts = []
            for i in range(n_pts + 1):
                t = i / n_pts; u = 1 - t
                pts += [u**3*p0[0]+3*u**2*t*p1[0]+3*u*t**2*p2[0]+t**3*p3[0],
                        u**3*p0[1]+3*u**2*t*p1[1]+3*u*t**2*p2[1]+t**3*p3[1]]
            return pts

        fs  = min(9, max(5, int(H*0.034)));  fss = min(7, max(4, int(H*0.026)))
        FT  = ("Arial", fs);  FTS = ("Arial", fss)

        TITLE_H       = max(14, int(H*0.07))
        SUP_CLEARANCE = max(40, int(H*0.22))

        buque = self.get_var("car_buque").get() or "BUQUE METANERO"

        # ── Ocean/sky background ──────────────────────────────────────────
        sky_cols = ["#87CEEB","#9DD5EE","#B3DCF1","#C9E4F4","#E0F0FA"]
        sky_h = int(H * 0.45)
        for si, sc in enumerate(sky_cols):
            sy1 = y + int(sky_h * si / len(sky_cols))
            sy2 = y + int(sky_h * (si + 1) / len(sky_cols))
            cv.create_rectangle(x, sy1, x + W, sy2, fill=sc, outline="")
        ocean_cols = ["#154360","#1A5276","#1F618D","#2471A3","#2980B9","#3498DB"]
        ocean_top = y + sky_h
        ocean_h = y + H - ocean_top
        for oi, oc in enumerate(ocean_cols):
            oy1 = ocean_top + int(ocean_h * oi / len(ocean_cols))
            oy2 = ocean_top + int(ocean_h * (oi + 1) / len(ocean_cols))
            cv.create_rectangle(x, oy1, x + W, oy2, fill=oc, outline="")
        # Wave highlights
        for wi in range(0, W, max(18, W // 15)):
            wx = x + wi
            wave_pts = bez((wx, ocean_top), (wx + 5, ocean_top - 2), (wx + 10, ocean_top + 1), (wx + 14, ocean_top), 6)
            for wj in range(0, len(wave_pts) - 2, 2):
                cv.create_line(int(wave_pts[wj]), int(wave_pts[wj+1]),
                               int(wave_pts[wj+2]), int(wave_pts[wj+3]), fill="#5DADE2", width=1)
        cv.create_rectangle(x, y, x + W, y + H, fill="", outline="#1A5276", width=2)

        # Title
        cv.create_text(x+W//2+1, y+TITLE_H//2+1, text=f"BUQUE METANERO / GNL  --  {buque}", font=FT, fill="#A0B0C0")
        cv.create_text(x+W//2, y+TITLE_H//2, text=f"BUQUE METANERO / GNL  --  {buque}", font=FT, fill="#1A5276")

        # ── Hull ──────────────────────────────────────────────────────────
        DECK_Y  = y + TITLE_H + SUP_CLEARANCE
        KEEL_Y  = y + H - 4
        HULL_H  = KEEL_Y - DECK_Y
        WL_Y    = DECK_Y + int(HULL_H * 0.55)
        BOW_X   = x + W - max(14, int(W * 0.03))
        STERN_X = x + max(14, int(W * 0.03))
        proa    = int(W * 0.045)
        CX      = x + W // 2

        # Freeboard (solo contorno, interior se ve como cross-section con tanques)
        bow_top = bez((BOW_X, DECK_Y), (BOW_X + proa, DECK_Y),
                      (BOW_X + proa, WL_Y - int(HULL_H * 0.1)), (BOW_X + proa // 2, WL_Y))
        top_poly = [STERN_X - 4, DECK_Y, BOW_X, DECK_Y] + bow_top + [STERN_X - 4, WL_Y]
        # Fondo claro para cross-section del casco (los tanques se dibujan encima)
        cv.create_polygon(top_poly, fill="#2C3E50", outline="")
        cv.create_polygon(top_poly, fill="", outline="#0E1A27", width=1)

        # Underwater hull
        bow_bot = bez((BOW_X + proa // 2, WL_Y),
                      (BOW_X + proa // 2, KEEL_Y - int(HULL_H * 0.05)),
                      (BOW_X - int(W * 0.03), KEEL_Y), (CX, KEEL_Y + int(HULL_H * 0.03)))
        stern_bot = bez((CX, KEEL_Y + int(HULL_H * 0.03)),
                        (STERN_X + int(W * 0.04), KEEL_Y),
                        (STERN_X, KEEL_Y - int(HULL_H * 0.02)), (STERN_X - 4, WL_Y))
        bot_poly = [STERN_X - 4, WL_Y] + bow_bot + stern_bot
        cv.create_polygon(bot_poly, fill="#7B241C", outline="#5B1A14", width=1)
        # Boot-topping
        cv.create_line(STERN_X - 10, WL_Y, BOW_X + proa // 2, WL_Y, fill="#1A1A1A", width=2)
        cv.create_line(STERN_X - 10, WL_Y + 2, BOW_X + proa // 2, WL_Y + 2, fill="#2E86C1", width=1, dash=(5, 3))
        cv.create_text(x + W - 25, WL_Y - 7, text="WL", font=FTS, fill="#2E86C1")

        # Deck with plating detail
        deck_x1 = STERN_X + int(W * 0.02)
        deck_x2 = BOW_X - int(W * 0.02)
        cv.create_rectangle(deck_x1, DECK_Y - 5, deck_x2, DECK_Y, fill="#2C3E50", outline="")
        cv.create_line(deck_x1, DECK_Y - 5, deck_x2, DECK_Y - 5, fill="#3D4B56", width=1)
        # Bulwark
        cv.create_rectangle(deck_x1, DECK_Y - 7, deck_x2, DECK_Y - 5, fill="#3D4B56", outline="#2C3E50")

        # ── Superstructure ────────────────────────────────────────────────
        SUP_W = max(40, int(W * 0.08));  SUP_H = max(30, int(SUP_CLEARANCE * 0.85))
        SUP_X = STERN_X + int(W * 0.02)

        # Multi-level superstructure
        n_levels = 4
        for li in range(n_levels):
            lw = SUP_W - li * max(3, int(SUP_W * 0.05))
            lh = SUP_H // n_levels
            ly = DECK_Y - (li + 1) * lh
            lx = SUP_X + li * max(1, int(SUP_W * 0.02))
            col = ["#D5D8DC","#DFE2E6","#E5E8EC","#EBEDF0"][li]
            cv.create_rectangle(lx, ly, lx + lw, ly + lh, fill=col, outline="#ABB2B9", width=1)
            # Windows
            win_y = ly + max(2, lh // 4)
            win_h = max(3, lh // 3)
            for wi in range(max(2, lw // 10)):
                win_x = lx + 4 + wi * max(6, lw // max(2, lw // 10))
                if win_x + 4 < lx + lw - 4:
                    cv.create_rectangle(win_x, win_y, win_x + max(3, lw // 12), win_y + win_h,
                                        fill="#2E86C1", outline="#1A5276")
        # Bridge wing
        bridge_y = DECK_Y - SUP_H
        cv.create_rectangle(SUP_X - max(4, int(SUP_W * 0.08)), bridge_y,
                            SUP_X + SUP_W + max(4, int(SUP_W * 0.08)), bridge_y + max(4, SUP_H // n_levels),
                            fill="#E5E7E9", outline="#ABB2B9")

        # Chimney
        chim_w = max(8, int(SUP_W * 0.25))
        chim_h = max(12, int(SUP_H * 0.40))
        chim_x = SUP_X + SUP_W // 2 - chim_w // 2
        chim_grads = ["#C0392B","#CD4535","#DA5040","#C0392B","#A82E22"]
        for ci2, cc in enumerate(chim_grads):
            cw1 = chim_x + int(chim_w * ci2 / len(chim_grads))
            cw2 = chim_x + int(chim_w * (ci2 + 1) / len(chim_grads))
            cv.create_rectangle(cw1, bridge_y - chim_h, cw2, bridge_y, fill=cc, outline="")
        cv.create_rectangle(chim_x, bridge_y - chim_h, chim_x + chim_w, bridge_y,
                            fill="", outline="#922B21", width=1)
        cv.create_oval(chim_x, bridge_y - chim_h - 3, chim_x + chim_w, bridge_y - chim_h + 3,
                       fill="#4A1A14", outline="#922B21")

        # ── Membrane tanks (inside hull, visible through cross-section) ───
        # Un metanero real tiene mínimo 4 tanques
        n = max(len(tank_names), 4)
        # Tanques ocupan máximo espacio del casco
        tk_start = SUP_X + SUP_W + max(3, int(W * 0.005))
        tk_end   = BOW_X - max(8, int(W * 0.03))
        tk_total = max(40, tk_end - tk_start)
        cofferdam = max(2, int(tk_total * 0.008))
        tk_each  = max(10, (tk_total - cofferdam * (n - 1)) // n)
        TK_TOP   = DECK_Y + max(4, int(HULL_H * 0.04))
        TK_BOT   = KEEL_Y - max(6, int(HULL_H * 0.07))
        TK_H     = TK_BOT - TK_TOP
        CH       = max(5, int(min(tk_each * 0.9, TK_H) * 0.07))

        # Extender lista de tanques para que siempre haya n tanques visibles
        _display_tanks = list(tank_names[:n])
        while len(_display_tanks) < n:
            _display_tanks.append(f"TK {len(_display_tanks)+1}")


        for i, tn in enumerate(_display_tanks):
            tx = tk_start + i * (tk_each + cofferdam)
            tw = tk_each


            # Outer insulation layer
            insul_margin = max(2, int(tw * 0.03))
            cv.create_polygon(
                tx - insul_margin, TK_TOP + CH - insul_margin,
                tx + tw + insul_margin, TK_TOP + CH - insul_margin,
                tx + tw + insul_margin, TK_BOT + insul_margin,
                tx - insul_margin, TK_BOT + insul_margin,
                fill="#9EAAB2", outline="#7F8C8D", width=1
            )
            # Perlite/foam insulation pattern
            for pi in range(0, tw + 2 * insul_margin, max(6, tw // 8)):
                px = tx - insul_margin + pi
                cv.create_line(px, TK_TOP + CH - insul_margin, px, TK_BOT + insul_margin,
                               fill="#8A969E", width=1, dash=(1, 8))

            # Prismatic body with chamfers
            body = [
                tx + CH, TK_TOP,
                tx + tw - CH, TK_TOP,
                tx + tw, TK_TOP + CH,
                tx + tw, TK_BOT,
                tx, TK_BOT,
                tx, TK_TOP + CH,
            ]
            cv.create_polygon(body, fill="#2E4057", outline="", width=0)

            # Cryogenic stainless steel gradient (silver tones)
            cryo_colors = ["#F0F3F4","#E8ECF0","#E0E5EA","#D5DCE1","#CCD4DA","#C0CAD2",
                           "#B8C2CC","#B0BAC4","#A8B2BC","#A0AAB4"]
            for ki, cc in enumerate(cryo_colors):
                strip_y = TK_TOP + CH + int((TK_H - CH) * ki / len(cryo_colors))
                strip_h = int((TK_H - CH) / len(cryo_colors)) + 1
                clip_top2 = max(TK_TOP + CH, strip_y)
                clip_bot2 = min(TK_BOT, strip_y + strip_h)
                if clip_bot2 > clip_top2:
                    cv.create_polygon(
                        tx + CH, clip_top2, tx + tw - CH, clip_top2,
                        tx + tw, min(clip_top2 + CH, clip_bot2), tx + tw, clip_bot2,
                        tx, clip_bot2, tx, min(clip_top2 + CH, clip_bot2),
                        fill=cc, outline=""
                    )

            # Membrane corrugation pattern (GTT NO96 or Mark III style)
            corr_spacing = max(6, TK_H // 12)
            for ci2 in range(int(TK_TOP + CH + corr_spacing), int(TK_BOT), int(corr_spacing)):
                cv.create_line(tx + 3, ci2, tx + tw - 3, ci2, fill="#8A96A0", width=1, dash=(6, 4))
            # Vertical corrugations
            v_spacing = max(8, tw // 8)
            for vi in range(tx + v_spacing, tx + tw, v_spacing):
                cv.create_line(vi, TK_TOP + CH + 3, vi, TK_BOT - 3, fill="#8A96A0", width=1, dash=(6, 4))

            # Tank outline
            cv.create_polygon(body, fill="", outline="#1A4F6E", width=2)

            # Top chamfer highlight (metallic reflection)
            cv.create_polygon(tx + CH, TK_TOP, tx + tw - CH, TK_TOP, tx + tw, TK_TOP + CH,
                              tx + CH + 3, TK_TOP + 3, fill="#E8F0F5", outline="")
            # Left edge highlight
            cv.create_line(tx, TK_TOP + CH, tx, TK_BOT, fill="#D0D8E0", width=1)

            # Fill level
            try:
                s   = self.parse_float(self.get_var(f"{etapa_key}_{tn}_s_corr").get() or "0") if etapa_key and tn else 0
                ref = self.parse_float(self.get_var(f"{etapa_key}_{tn}_alt_ref").get() or "1")
                vp  = min(max(s / ref, 0), 1) if ref > 0 else 0
            except: vp = 0.0
            pnm = self.get_var(f"{etapa_key}_{tn}_prod_name").get() if etapa_key else ""

            if vp > 0.01:
                fill_h2 = int(TK_H * vp)
                fy = TK_BOT - fill_h2
                gnl_fill = self.get_prod_color(tn, etapa_key)[0] if etapa_key else "#E8E8F0"
                # LNG fill with gradient
                cv.create_rectangle(tx + 2, fy, tx + tw - 2, TK_BOT - 2, fill=gnl_fill, outline="")
                # Lighter band at top of liquid
                cv.create_rectangle(tx + 2, fy, tx + tw - 2, fy + max(2, fill_h2 // 8),
                                    fill="#F0F0F8", outline="")
                # Level line
                cv.create_line(tx + 2, fy, tx + tw - 2, fy, fill="#5D6D7E", width=2)
                # Boil-off gas bubbles
                for si in range(1, 5):
                    sx2 = tx + int(tw * si / 5)
                    bub_r = max(1, int(tw * 0.015))
                    cv.create_oval(sx2 - bub_r, fy - max(6, int(TK_H * 0.03)) - bub_r,
                                   sx2 + bub_r, fy - max(6, int(TK_H * 0.03)) + bub_r,
                                   fill="white", outline="#D0D8E0")
                    cv.create_line(sx2, fy, sx2, fy - max(6, int(TK_H * 0.03)),
                                   fill="#D0D8E0", width=1, dash=(2, 3))
                txt_gnl = self.contrast_text(gnl_fill)
                cv.create_text(tx + tw // 2, fy + fill_h2 // 2, text=f"{vp*100:.0f}%",
                               font=FTS, fill=txt_gnl)

            # Cryogenic dome (top of tank, above deck line)
            dm_cx = tx + tw // 2
            # Dome base
            cv.create_rectangle(dm_cx - max(12, int(tw * 0.12)), TK_TOP - max(12, int(TK_H * 0.04)),
                                dm_cx + max(12, int(tw * 0.12)), TK_TOP,
                                fill="#CCD1D9", outline="#7F8C8D", width=2)
            # Dome trunk
            cv.create_rectangle(dm_cx - max(6, int(tw * 0.06)), TK_TOP - max(18, int(TK_H * 0.06)),
                                dm_cx + max(6, int(tw * 0.06)), TK_TOP - max(12, int(TK_H * 0.04)),
                                fill="#5D6D7E", outline="#4A5568")
            # Dome cap
            dome_r = max(4, int(tw * 0.05))
            cv.create_oval(dm_cx - dome_r, TK_TOP - max(22, int(TK_H * 0.08)),
                           dm_cx + dome_r, TK_TOP - max(16, int(TK_H * 0.05)),
                           fill="#808B96", outline="#5D6D7E")
            # Cryo indicator
            cv.create_text(tx + 6, TK_TOP + 8, text="[C]", font=("Arial", max(4, fss - 1)), fill="#1ABC9C")

            # Tank number and labels
            tn_s = tn[:6]
            cv.create_text(tx + tw // 2, TK_BOT + max(8, int(HULL_H * 0.04)),
                           text=tn_s, font=FTS, fill="#2C3E50")
            if pnm:
                cv.create_text(tx + tw // 2, TK_TOP + TK_H // 2, text=pnm[:8], font=FTS, fill="#2C3E50")

        # Deck piping between tanks
        if n > 1:
            pipe_y = DECK_Y - 3
            cv.create_rectangle(tk_start, pipe_y - 1, tk_end, pipe_y + 1, fill="#5D6D7E", outline="#4A5568")

        # Bow mast with stays
        mast_x = BOW_X - max(16, int(W * 0.04))
        mast_top = DECK_Y - max(25, int(SUP_CLEARANCE * 0.60))
        cv.create_line(mast_x, DECK_Y - 5, mast_x, mast_top, fill="#626567", width=2)
        cv.create_line(mast_x, mast_top, BOW_X, DECK_Y - 5, fill="#95A5A6", width=1, dash=(3, 4))
        cv.create_line(mast_x, mast_top, mast_x - max(20, int(W * 0.05)), DECK_Y - 5,
                       fill="#95A5A6", width=1, dash=(3, 4))

        # Draft marks
        for di in range(3):
            dm_y = WL_Y + int(di * HULL_H * 0.08)
            cv.create_text(BOW_X + proa // 2 - 4, dm_y, text=str(10 - di),
                           font=("Arial", max(3, fss - 2)), fill="#FDFEFE")

        # Side label
        cv.create_text(x + 12, y + TITLE_H + HULL_H // 2, text=side_label,
                       font=FTS, fill="#1A5276", angle=90)

    def _draw_pipeline(self, cv, x, y, W, H, etapa_key, tipo, tank_names):
        """Underground pipeline cross-section -- earth layers, trench, FBE coating, warning tape, CP anode."""
        import math
        if H < 60 or W < 120: return
        if tank_names is None: tank_names = self.lista_tanques

        def bez(p0, p1, p2, p3, n_pts=16):
            pts = []
            for i in range(n_pts + 1):
                t = i / n_pts; u = 1 - t
                pts += [u**3*p0[0]+3*u**2*t*p1[0]+3*u*t**2*p2[0]+t**3*p3[0],
                        u**3*p0[1]+3*u**2*t*p1[1]+3*u*t**2*p2[1]+t**3*p3[1]]
            return pts

        fs  = min(10, max(5, int(H*0.036)));  fss = min(9, max(4, int(H*0.028)))
        FT  = ("Arial", fs);  FTS = ("Arial", fss)

        TITLE_H = max(16, int(H*0.08))
        instalacion = self.get_var("car_buque").get() or "DUCTO"
        tipo_color  = {"OLEODUCTO": "#C0392B", "POLIDUCTO": "#1A5276",
                       "GASODUCTO": "#1D6A39"}.get(tipo, "#5D4037")

        GROUND_Y = y + int(H * 0.42)

        # ── Sky background gradient ───────────────────────────────────────
        sky_cols = ["#87CEEB","#9DD5EE","#B3DCF1","#C9E4F4","#DFF0FA","#F0F7FC"]
        sky_h = GROUND_Y - y
        for si, sc in enumerate(sky_cols):
            sy1 = y + int(sky_h * si / len(sky_cols))
            sy2 = y + int(sky_h * (si + 1) / len(sky_cols))
            cv.create_rectangle(x, sy1, x + W, sy2, fill=sc, outline="")
        cv.create_rectangle(x, y, x + W, y + H, fill="", outline="#808B96", width=2)

        # Title with type color bar
        cv.create_rectangle(x + 1, y + 1, x + W - 1, y + TITLE_H + 2, fill="#F0F2F4", outline="")
        cv.create_rectangle(x + 1, y + TITLE_H, x + W - 1, y + TITLE_H + 4, fill=tipo_color, outline="")
        ttl = f"{tipo}  --  {instalacion}"
        cv.create_text(x + W // 2 + 1, y + TITLE_H // 2 + 1, text=ttl, font=FT, fill="#B0B8C0")
        cv.create_text(x + W // 2, y + TITLE_H // 2, text=ttl, font=FT, fill="#1B3A5C")

        # ── Ground surface with grass/vegetation ──────────────────────────
        # Grass line (irregular with Bezier)
        grass_pts = [x, GROUND_Y]
        seg_w = max(20, W // 12)
        for gi in range(0, W, seg_w):
            gx1 = x + gi
            gx2 = min(x + gi + seg_w, x + W)
            gh = max(2, int(H * 0.01)) * (1 + (gi % (seg_w * 2) == 0))
            grass_pts += bez((gx1, GROUND_Y), (gx1 + seg_w // 3, GROUND_Y - gh),
                             (gx2 - seg_w // 3, GROUND_Y - gh // 2), (gx2, GROUND_Y), 6)
        grass_pts += [x + W, GROUND_Y + 4, x, GROUND_Y + 4]
        cv.create_polygon(grass_pts, fill="#4A7C3F", outline="")
        # Grass tufts
        for gt in range(0, W, max(8, W // 30)):
            gx = x + gt + (hash(gt) % 5)
            gh = max(3, int(H * 0.015)) + (hash(gt * 7) % 3)
            cv.create_line(gx, GROUND_Y, gx - 2, GROUND_Y - gh, fill="#5D9B48", width=1)
            cv.create_line(gx, GROUND_Y, gx + 2, GROUND_Y - gh + 1, fill="#3E7A30", width=1)

        # ── Earth layers (cross section) ──────────────────────────────────
        UG_H = y + H - GROUND_Y
        earth_layers = [
            (0.00, 0.12, "#6B4F36", "#7D5E42"),  # topsoil (dark brown)
            (0.12, 0.30, "#8B6F47", "#9C7E55"),  # subsoil (lighter brown)
            (0.30, 0.55, "#C4A86C", "#D4B87C"),  # clay/sand
            (0.55, 0.75, "#A09080", "#B0A090"),  # gravel
            (0.75, 1.00, "#808080", "#909090"),  # bedrock
        ]
        for ef1, ef2, ec1, ec2 in earth_layers:
            ey1 = GROUND_Y + int(UG_H * ef1)
            ey2 = GROUND_Y + int(UG_H * ef2)
            # Two-tone gradient per layer
            mid = (ey1 + ey2) // 2
            cv.create_rectangle(x + 1, ey1, x + W - 1, mid, fill=ec1, outline="")
            cv.create_rectangle(x + 1, mid, x + W - 1, ey2, fill=ec2, outline="")
        # Layer boundary lines
        for ef1, ef2, ec1, ec2 in earth_layers[1:]:
            ey = GROUND_Y + int(UG_H * ef1)
            cv.create_line(x + 2, ey, x + W - 2, ey, fill="#5D5040", width=1, dash=(8, 4))
        # Rocks in gravel layer
        gravel_y1 = GROUND_Y + int(UG_H * 0.55)
        gravel_y2 = GROUND_Y + int(UG_H * 0.75)
        for ri in range(0, W, max(15, W // 20)):
            rx = x + ri + (hash(ri * 13) % 8)
            ry = gravel_y1 + (hash(ri * 7) % max(1, gravel_y2 - gravel_y1 - 6))
            rr = max(2, int(H * 0.008)) + hash(ri * 3) % 3
            cv.create_oval(rx, ry, rx + rr * 2, ry + rr, fill="#989088", outline="#807870")
        # Root-like lines in topsoil
        for ri in range(0, W, max(25, W // 10)):
            rx = x + ri + (hash(ri * 17) % 12)
            ry = GROUND_Y + int(UG_H * 0.02)
            cv.create_line(rx, ry, rx + max(6, W // 20), ry + int(UG_H * 0.08),
                           fill="#5A4030", width=1, dash=(3, 5))

        # ── Pipe parameters ───────────────────────────────────────────────
        PIPE_Y   = GROUND_Y + int(UG_H * 0.38)
        PIPE_R   = max(10, int(UG_H * 0.14))
        PAD_L    = max(40, int(W * 0.1))
        PAD_R    = PAD_L
        P_LEFT   = x + PAD_L
        P_RIGHT  = x + W - PAD_R
        PIPE_W   = P_RIGHT - P_LEFT

        # ── Yellow warning tape (300-500mm below surface, ABOVE pipe) ─────
        tape_y = GROUND_Y + int(UG_H * 0.15)  # about 300-500mm below surface
        tape_h = max(3, int(UG_H * 0.025))
        cv.create_rectangle(P_LEFT, tape_y, P_RIGHT, tape_y + tape_h,
                            fill="#F4D03F", outline="#D4AC0D", width=1)
        # "CAUTION BURIED PIPELINE" text on tape
        cv.create_text((P_LEFT + P_RIGHT) // 2, tape_y + tape_h // 2,
                       text="PELIGRO - DUCTO ENTERRADO", font=("Arial", max(3, fss - 2)), fill="#922B21")
        # Dashed edges for visibility
        cv.create_line(P_LEFT, tape_y, P_RIGHT, tape_y, fill="#D4AC0D", width=1, dash=(4, 4))
        cv.create_line(P_LEFT, tape_y + tape_h, P_RIGHT, tape_y + tape_h, fill="#D4AC0D", width=1, dash=(4, 4))

        # ── Trapezoidal trench with sloped walls ─────────────────────────
        trench_top_w = PIPE_R * 4
        trench_bot_w = PIPE_R * 3
        trench_top_y = PIPE_Y - PIPE_R - max(8, int(PIPE_R * 0.6))
        trench_bot_y = PIPE_Y + PIPE_R + max(8, int(PIPE_R * 0.6))
        trench_cx = (P_LEFT + P_RIGHT) // 2
        # Only draw trench cross-section at left end
        if True:
            tcx = P_LEFT
            cv.create_polygon(tcx - trench_top_w // 2, trench_top_y,
                              tcx + trench_top_w // 2, trench_top_y,
                              tcx + trench_bot_w // 2, trench_bot_y,
                              tcx - trench_bot_w // 2, trench_bot_y,
                              fill="", outline="#8B6F47", width=1, dash=(4, 3))

        # ── Cathodic protection anode ─────────────────────────────────────
        cp_x = P_LEFT + int(PIPE_W * 0.18)
        cp_y = PIPE_Y + PIPE_R + max(8, int(UG_H * 0.06))
        cp_h = max(12, int(UG_H * 0.12))
        # Anode body
        cv.create_rectangle(cp_x - 3, cp_y, cp_x + 3, cp_y + cp_h,
                            fill="#808080", outline="#606060")
        # Wire from anode to pipe
        cv.create_line(cp_x, cp_y, cp_x, PIPE_Y + PIPE_R + 2, fill="#F4D03F", width=1, dash=(2, 2))
        cv.create_text(cp_x, cp_y + cp_h + max(4, fss), text="CP", font=("Arial", max(3, fss - 2)), fill="#D4AC0D")

        # ── Buried pipe with coatings ─────────────────────────────────────
        n_pipes = 2 if tipo in ("POLIDUCTO",) else 1
        pipe_offsets = [0] if n_pipes == 1 else [-PIPE_R - 4, PIPE_R + 4]

        for pi, p_off in enumerate(pipe_offsets):
            py = PIPE_Y + p_off
            pc = tipo_color if pi == 0 else "#7D3C98"

            # Sand padding/bedding around pipe (light tan)
            cv.create_rectangle(P_LEFT - PIPE_R, py - PIPE_R - max(6, int(PIPE_R * 0.5)),
                                P_RIGHT + PIPE_R, py + PIPE_R + max(6, int(PIPE_R * 0.5)),
                                fill="#E8D8B0", outline="")

            # FBE coating (dark green layer around pipe - Fusion Bonded Epoxy)
            coat_t = max(2, int(PIPE_R * 0.15))
            cv.create_rectangle(P_LEFT, py - PIPE_R - coat_t, P_RIGHT, py + PIPE_R + coat_t,
                                fill="#1A3D1A", outline="")
            # FBE inner layer (dark green)
            cv.create_rectangle(P_LEFT, py - PIPE_R - coat_t + 1, P_RIGHT, py + PIPE_R + coat_t - 1,
                                fill="#2D5A2D" if pi == 0 else "#6A329F", outline="")

            # Steel pipe with cylindrical gradient
            cv.create_rectangle(P_LEFT, py - PIPE_R, P_RIGHT, py + PIPE_R, fill=pc, outline="")
            try:
                _r0 = int(pc[1:3], 16); _g0 = int(pc[3:5], 16); _b0 = int(pc[5:7], 16)
            except: _r0, _g0, _b0 = 100, 100, 100
            pipe_strips_grad = [
                (0.00, 0.08, 0.50), (0.08, 0.18, 0.68), (0.18, 0.32, 0.88),
                (0.32, 0.45, 1.15), (0.45, 0.55, 1.35), (0.55, 0.68, 1.15),
                (0.68, 0.82, 0.82), (0.82, 0.92, 0.62), (0.92, 1.00, 0.45),
            ]
            for pf1, pf2, bright in pipe_strips_grad:
                _rr = min(255, int(_r0 * bright)); _gg = min(255, int(_g0 * bright)); _bb = min(255, int(_b0 * bright))
                sc2 = f"#{_rr:02x}{_gg:02x}{_bb:02x}"
                sy1 = py - PIPE_R + int(pf1 * 2 * PIPE_R)
                sy2 = py - PIPE_R + int(pf2 * 2 * PIPE_R)
                cv.create_rectangle(P_LEFT + 2, sy1, P_RIGHT - 2, sy2 + 1, fill=sc2, outline="")
            # Specular highlight
            hl_y = py - PIPE_R + max(3, PIPE_R // 4)
            cv.create_rectangle(P_LEFT + PIPE_W // 8, hl_y, P_RIGHT - PIPE_W // 8,
                                hl_y + max(2, PIPE_R // 6), fill="#FFFFFF", outline="")
            # Pipe outline
            cv.create_rectangle(P_LEFT, py - PIPE_R, P_RIGHT, py + PIPE_R,
                                fill="", outline="#2C3E50", width=2)

            # End caps (cross-section view ellipses)
            # Left end - full cross section
            cap_r = PIPE_R + coat_t + 2
            cv.create_oval(P_LEFT - cap_r, py - cap_r, P_LEFT + cap_r, py + cap_r,
                           fill="#1A1A1A", outline="#0A0A0A", width=1)
            cv.create_oval(P_LEFT - PIPE_R, py - PIPE_R, P_LEFT + PIPE_R, py + PIPE_R,
                           fill=pc, outline="#2C3E50", width=2)
            # Inner bore
            inner_r = max(3, int(PIPE_R * 0.7))
            cv.create_oval(P_LEFT - inner_r, py - inner_r, P_LEFT + inner_r, py + inner_r,
                           fill="#2C2C2C", outline="#1A1A1A", width=1)
            # Product in bore
            if tipo == "OLEODUCTO":
                cv.create_oval(P_LEFT - inner_r + 1, py - inner_r + 1, P_LEFT + inner_r - 1, py + inner_r - 1,
                               fill="#7B241C", outline="")
            elif tipo == "GASODUCTO":
                cv.create_oval(P_LEFT - inner_r + 1, py - inner_r + 1, P_LEFT + inner_r - 1, py + inner_r - 1,
                               fill="#3D6B3D", outline="")

            # Right end
            cv.create_oval(P_RIGHT - cap_r, py - cap_r, P_RIGHT + cap_r, py + cap_r,
                           fill="#1A1A1A", outline="#0A0A0A", width=1)
            cv.create_oval(P_RIGHT - PIPE_R, py - PIPE_R, P_RIGHT + PIPE_R, py + PIPE_R,
                           fill="#4A5568", outline="#2C3E50", width=2)

            # Flanges (welded joints)
            n_flanges = max(3, min(7, PIPE_W // 60))
            flange_step = PIPE_W // (n_flanges + 1)
            for fi in range(n_flanges):
                fx = P_LEFT + (fi + 1) * flange_step
                # Flange ring
                cv.create_rectangle(fx - 2, py - PIPE_R - 4, fx + 2, py + PIPE_R + 4,
                                    fill="#5D6D7E", outline="#4A5568", width=1)
                # Bolts
                for by_off in [-PIPE_R - 3, PIPE_R + 2]:
                    cv.create_oval(fx - 1, py + by_off - 1, fx + 1, py + by_off + 1,
                                   fill="#3D4B56", outline="")

            # Weld seam indicators (circumferential)
            for wi in range(1, n_flanges + 2):
                wx = P_LEFT + int(PIPE_W * wi / (n_flanges + 2))
                cv.create_line(wx, py - PIPE_R + 2, wx, py + PIPE_R - 2,
                               fill="#808890", width=1, dash=(2, 6))

        # ── Flow direction arrows ─────────────────────────────────────────
        flow_val = self.get_var(f"{etapa_key}_{tank_names[0]}_vol_nat_prod").get() if etapa_key and tank_names else ""
        arrow_color = "#F4D03F" if flow_val else "#95A5A6"
        n_arrows = max(3, min(6, PIPE_W // 80))
        for ai in range(n_arrows):
            ax = P_LEFT + (ai + 1) * PIPE_W // (n_arrows + 1)
            aw = max(8, int(PIPE_R * 0.6))
            # Arrow with shadow
            cv.create_polygon(ax - aw // 2 + 1, PIPE_Y - aw // 3 + 1,
                              ax + aw // 2 + 1, PIPE_Y + 1,
                              ax - aw // 2 + 1, PIPE_Y + aw // 3 + 1,
                              fill="#80800A", outline="")
            cv.create_polygon(ax - aw // 2, PIPE_Y - aw // 3,
                              ax + aw // 2, PIPE_Y,
                              ax - aw // 2, PIPE_Y + aw // 3,
                              fill=arrow_color, outline=arrow_color)

        # ── Aboveground section: valve station ────────────────────────────
        # Riser pipe (pipe coming out of ground)
        riser_x = P_LEFT + int(PIPE_W * 0.15)
        riser_top = GROUND_Y - max(20, int(sky_h * 0.25))
        cv.create_rectangle(riser_x - 3, riser_top, riser_x + 3, GROUND_Y + 2,
                            fill=tipo_color, outline="#2C3E50", width=1)
        # Riser elbow
        cv.create_arc(riser_x - 8, riser_top - 4, riser_x + 8, riser_top + 12,
                       start=0, extent=180, style="arc", outline=tipo_color, width=4)
        # Block valve on riser
        vv_y = riser_top + max(6, int(sky_h * 0.06))
        vv_r = max(4, PIPE_R // 3)
        cv.create_polygon(riser_x - vv_r, vv_y - vv_r, riser_x + vv_r, vv_y + vv_r,
                          riser_x + vv_r, vv_y - vv_r, riser_x - vv_r, vv_y + vv_r,
                          fill="#E74C3C", outline="#C0392B", width=2)
        cv.create_line(riser_x, vv_y - vv_r, riser_x, vv_y - vv_r * 2, fill="#C0392B", width=2)
        cv.create_oval(riser_x - vv_r, vv_y - vv_r * 2 - 3, riser_x + vv_r, vv_y - vv_r * 2,
                       fill="", outline="#C0392B", width=1)

        # ── Pig launcher/receiver station ─────────────────────────────────
        pig_x = P_RIGHT - int(PIPE_W * 0.12)
        pig_top = GROUND_Y - max(18, int(sky_h * 0.22))
        pig_w = max(30, int(W * 0.08))
        pig_h = max(8, int(PIPE_R * 0.8))
        # Barrel (horizontal cylinder)
        cv.create_rectangle(pig_x, pig_top, pig_x + pig_w, pig_top + pig_h * 2,
                            fill=tipo_color, outline="#2C3E50", width=1)
        # Barrel gradient
        try:
            _r0 = int(tipo_color[1:3], 16); _g0 = int(tipo_color[3:5], 16); _b0 = int(tipo_color[5:7], 16)
        except: _r0, _g0, _b0 = 100, 100, 100
        for pgi in range(5):
            frac = pgi / 4
            bright = 0.6 + 0.8 * (1 - abs(frac - 0.35) * 2)
            _rr = min(255, int(_r0 * bright)); _gg = min(255, int(_g0 * bright)); _bb = min(255, int(_b0 * bright))
            pgc = f"#{_rr:02x}{_gg:02x}{_bb:02x}"
            pgy = pig_top + int(pig_h * 2 * pgi / 5)
            cv.create_rectangle(pig_x + 2, pgy, pig_x + pig_w - 2, pgy + pig_h * 2 // 5 + 1,
                                fill=pgc, outline="")
        cv.create_rectangle(pig_x, pig_top, pig_x + pig_w, pig_top + pig_h * 2,
                            fill="", outline="#2C3E50", width=1)
        # Closure door (end cap)
        cv.create_oval(pig_x + pig_w - 3, pig_top - 2, pig_x + pig_w + 5, pig_top + pig_h * 2 + 2,
                       fill="#4A5568", outline="#2C3E50", width=2)
        # Kicker line
        cv.create_rectangle(pig_x - 2, pig_top + pig_h, pig_x, pig_top + pig_h + 3,
                            fill="#5D6D7E", outline="")
        cv.create_line(pig_x, pig_top + pig_h + 1, pig_x - max(8, int(W * 0.02)), pig_top + pig_h + 1,
                       fill="#5D6D7E", width=2)
        # Label
        cv.create_text(pig_x + pig_w // 2, pig_top - max(4, fss),
                       text="PIG L/R", font=("Arial", max(3, fss - 2)), fill="#5D6D7E")
        # Support legs
        for sl in [pig_x + 4, pig_x + pig_w - 4]:
            cv.create_line(sl, pig_top + pig_h * 2, sl, GROUND_Y, fill="#5D6D7E", width=2)

        # ── Meter station (above ground) ──────────────────────────────────
        mtr_x = x + W // 2
        mtr_top = GROUND_Y - max(30, int(sky_h * 0.40))
        mtr_w = max(45, int(W * 0.13)); mtr_h = max(30, int(sky_h * 0.35))
        # Cabinet shadow
        cv.create_rectangle(mtr_x - mtr_w // 2 + 3, mtr_top + 3,
                            mtr_x + mtr_w // 2 + 3, mtr_top + mtr_h + 3,
                            fill="#404850", outline="")
        # Cabinet body with metallic gradient
        cab_grads = ["#D0D4D8","#C8CCD0","#C0C4C8","#B8BCC0","#B0B4B8","#A8ACB0"]
        for ci2, cc in enumerate(cab_grads):
            cy1 = mtr_top + int(mtr_h * ci2 / len(cab_grads))
            cy2 = mtr_top + int(mtr_h * (ci2 + 1) / len(cab_grads))
            cv.create_rectangle(mtr_x - mtr_w // 2, cy1, mtr_x + mtr_w // 2, cy2, fill=cc, outline="")
        cv.create_rectangle(mtr_x - mtr_w // 2, mtr_top, mtr_x + mtr_w // 2, mtr_top + mtr_h,
                            fill="", outline="#2C3E50", width=2)
        # Header bar
        cv.create_rectangle(mtr_x - mtr_w // 2, mtr_top, mtr_x + mtr_w // 2, mtr_top + max(5, mtr_h // 5),
                            fill="#2C3E50", outline="")
        cv.create_text(mtr_x, mtr_top + max(3, mtr_h // 10),
                       text="MEDIDOR", font=("Arial", max(4, fss - 1)), fill="white")
        # Digital display
        disp_bg = "#1A5276" if flow_val else "#4A5568"
        disp_y = mtr_top + mtr_h // 5 + 4
        disp_h = max(16, int(mtr_h * 0.35))
        cv.create_rectangle(mtr_x - mtr_w // 2 + 5, disp_y,
                            mtr_x + mtr_w // 2 - 5, disp_y + disp_h,
                            fill=disp_bg, outline="#1ABC9C", width=2)
        disp_val = flow_val[:8] if flow_val else "---"
        cv.create_text(mtr_x, disp_y + disp_h // 2,
                       text=disp_val, font=("Arial", max(5, fss)), fill="#1ABC9C")
        # Unit label
        unit_txt = {"OLEODUCTO": "m3", "GASODUCTO": "m3/h", "POLIDUCTO": "m3"}.get(tipo, "m3")
        cv.create_text(mtr_x, disp_y + disp_h + max(3, fss // 2),
                       text=unit_txt, font=("Arial", max(3, fss - 2)), fill="#85929E")
        # LED status lights
        for li, lc in enumerate(["#27AE60", "#F4D03F"]):
            cv.create_oval(mtr_x - mtr_w // 2 + 8 + li * 10, mtr_top + mtr_h - 10,
                           mtr_x - mtr_w // 2 + 14 + li * 10, mtr_top + mtr_h - 4,
                           fill=lc, outline="#2C3E50")
        # Connection pipe to underground
        cv.create_line(mtr_x, mtr_top + mtr_h, mtr_x, GROUND_Y + 2, fill="#5D6D7E", width=3)
        cv.create_line(mtr_x, GROUND_Y + 2, mtr_x, PIPE_Y - PIPE_R, fill="#5D6D7E", width=3)
        # Legs
        for sl in [mtr_x - mtr_w // 2 + 4, mtr_x + mtr_w // 2 - 4]:
            cv.create_line(sl, mtr_top + mtr_h, sl, GROUND_Y, fill="#4A5568", width=2)

        # ── Pressure gauges above ground ──────────────────────────────────
        for gpos in [P_LEFT + int(PIPE_W * 0.35), P_LEFT + int(PIPE_W * 0.70)]:
            # Gauge post
            g_top = GROUND_Y - max(12, int(sky_h * 0.15))
            cv.create_line(gpos, GROUND_Y, gpos, g_top, fill="#5D6D7E", width=2)
            g_r = max(6, int(sky_h * 0.06))
            cv.create_oval(gpos - g_r - 1, g_top - g_r - 1, gpos + g_r + 1, g_top + g_r + 1,
                           fill="#4A5568", outline="")
            cv.create_oval(gpos - g_r, g_top - g_r, gpos + g_r, g_top + g_r,
                           fill="white", outline="#E74C3C", width=2)
            # Scale
            cv.create_arc(gpos - g_r + 2, g_top - g_r + 2, gpos + g_r - 2, g_top + g_r - 2,
                          start=30, extent=240, style="arc", outline="#E0E0E0", width=1)
            # Needle
            cv.create_line(gpos, g_top, gpos + int(g_r * 0.6), g_top - int(g_r * 0.4),
                           fill="#C0392B", width=1)
            cv.create_oval(gpos - 1, g_top - 1, gpos + 1, g_top + 1, fill="#C0392B", outline="")

        # ── Warning marker posts ──────────────────────────────────────────
        for mp in [P_LEFT + int(PIPE_W * 0.05), P_LEFT + int(PIPE_W * 0.95)]:
            post_top = GROUND_Y - max(15, int(sky_h * 0.20))
            cv.create_rectangle(mp - 2, post_top, mp + 2, GROUND_Y, fill="#F4D03F", outline="#D4AC0D")
            cv.create_rectangle(mp - 5, post_top - 8, mp + 5, post_top,
                                fill="#F4D03F", outline="#D4AC0D", width=1)
            # Diamond hazmat sign
            cv.create_polygon(mp, post_top - 8, mp + 4, post_top - 4,
                              mp, post_top, mp - 4, post_top - 4,
                              fill=tipo_color, outline="#2C3E50")

        # ── Depth dimension annotations ────────────────────────────────────
        dim_x = x + max(10, int(W * 0.03))
        # Ground to pipe center
        cv.create_line(dim_x, GROUND_Y, dim_x, PIPE_Y, fill="#F4D03F", width=1)
        cv.create_line(dim_x - 3, GROUND_Y, dim_x + 3, GROUND_Y, fill="#F4D03F", width=1)
        cv.create_line(dim_x - 3, PIPE_Y, dim_x + 3, PIPE_Y, fill="#F4D03F", width=1)
        cv.create_text(dim_x + 6, (GROUND_Y + PIPE_Y) // 2,
                       text="Prof.", font=("Arial", max(3, fss - 2)), fill="#F4D03F", anchor="w")

        # ── Layer labels on right side ────────────────────────────────────
        lbl_x = x + W - max(30, int(W * 0.08))
        layer_lbls = [("Suelo", 0.06), ("Subsuelo", 0.21), ("Arcilla", 0.42), ("Grava", 0.65), ("Roca", 0.88)]
        for lt, lf in layer_lbls:
            ly = GROUND_Y + int(UG_H * lf)
            if ly < y + H - 4:
                cv.create_text(lbl_x, ly, text=lt, font=("Arial", max(3, fss - 2)), fill="#D5D0C8", anchor="e")

        # ── Flow direction label ──────────────────────────────────────────
        flow_dir = "-->  FLUJO"
        cv.create_text(x + W // 2, y + H - max(4, int(H * 0.02)),
                       text=flow_dir, font=FTS, fill="#5D6D7E")


    def _draw_electric(self, cv, x, y, W, H, etapa_key, tank_names):
        """Electrical metering room -- RAL 7035 grey cabinet, HV warning, CT symbol, anti-static floor."""
        import math
        if H < 60 or W < 100: return
        if tank_names is None: tank_names = self.lista_tanques

        fs  = min(10, max(5, int(H*0.036)));  fss = min(9, max(4, int(H*0.028)))
        FT  = ("Arial", fs);  FTS = ("Arial", fss)

        TITLE_H  = max(18, int(H*0.09))
        instalacion = self.get_var("car_buque").get() or "INSTALACION"

        # ── Room background (concrete wall) ───────────────────────────────
        wall_grads = ["#2D3436","#2C3335","#2A3133","#282F31","#262D2F","#242B2D"]
        for wi, wc in enumerate(wall_grads):
            wy1 = y + int(H * wi / len(wall_grads))
            wy2 = y + int(H * (wi + 1) / len(wall_grads))
            cv.create_rectangle(x, wy1, x + W, wy2, fill=wc, outline="")
        cv.create_rectangle(x, y, x + W, y + H, fill="", outline="#4A5568", width=3)

        # ── Conduit/cable tray at top ─────────────────────────────────────
        tray_h = max(6, int(H * 0.03))
        cv.create_rectangle(x + 4, y + 4, x + W - 4, y + 4 + tray_h, fill="#3D4B56", outline="#2C3E50")
        # Cables in tray
        cable_cols = ["#E74C3C","#3498DB","#2ECC71","#F4D03F","#95A5A6"]
        for ci in range(min(5, W // 20)):
            cx2 = x + 8 + ci * max(5, W // 20)
            cv.create_line(cx2, y + 5, cx2, y + 4 + tray_h - 1, fill=cable_cols[ci % 5], width=2)

        # ── Title bar with high voltage warning ───────────────────────────
        cv.create_rectangle(x + 1, y + 4 + tray_h + 2, x + W - 1, y + TITLE_H + tray_h + 2,
                            fill="#6A1B9A", outline="")
        # Lightning bolt symbols
        bolt_y = y + (TITLE_H + tray_h) // 2 + 3
        for bx in [x + max(12, int(W * 0.04)), x + W - max(12, int(W * 0.04))]:
            cv.create_polygon(bx - 3, bolt_y - 5, bx + 1, bolt_y - 1, bx - 1, bolt_y,
                              bx + 3, bolt_y + 5, bx - 1, bolt_y + 1, bx + 1, bolt_y,
                              fill="#F4D03F", outline="")
        cv.create_text(x + W // 2, bolt_y,
                       text=f"MEDICION ELECTRICA  --  {instalacion}",
                       font=FT, fill="white")

        # ── Main panel cabinet body ───────────────────────────────────────
        panel_x = x + int(W * 0.04); panel_y = y + TITLE_H + tray_h + max(8, int(H * 0.04))
        panel_w = W - int(W * 0.08); panel_h = H - TITLE_H - tray_h - max(28, int(H * 0.16))
        # Panel shadow
        cv.create_rectangle(panel_x + 3, panel_y + 3, panel_x + panel_w + 3, panel_y + panel_h + 3,
                            fill="#0A0E12", outline="")
        # Panel body: RAL 7035 light grey metallic gradient
        pan_grads = ["#C8CDD2","#C4C9CE","#C0C5CA","#BCC1C6","#B8BDC2","#B4B9BE"]
        for pi2, pc2 in enumerate(pan_grads):
            py1 = panel_y + int(panel_h * pi2 / len(pan_grads))
            py2 = panel_y + int(panel_h * (pi2 + 1) / len(pan_grads))
            cv.create_rectangle(panel_x, py1, panel_x + panel_w, py2, fill=pc2, outline="")
        cv.create_rectangle(panel_x, panel_y, panel_x + panel_w, panel_y + panel_h,
                            fill="", outline="#95A5A6", width=2)
        # Panel bevel highlight (RAL 7035 light grey)
        cv.create_line(panel_x, panel_y, panel_x + panel_w, panel_y, fill="#D0D5DA", width=1)
        cv.create_line(panel_x, panel_y, panel_x, panel_y + panel_h, fill="#A8ADB2", width=1)

        # ── Hinges on panel door ──────────────────────────────────────────
        for hy in [panel_y + int(panel_h * 0.2), panel_y + int(panel_h * 0.8)]:
            cv.create_rectangle(panel_x - 3, hy - 4, panel_x + 2, hy + 4,
                                fill="#5D6D7E", outline="#4A5568")
            cv.create_oval(panel_x - 2, hy - 1, panel_x + 1, hy + 1, fill="#808B96", outline="")

        # ── Door handle/lock ──────────────────────────────────────────────
        lock_x = panel_x + panel_w - max(8, int(panel_w * 0.03))
        lock_y = panel_y + panel_h // 2
        cv.create_rectangle(lock_x - 4, lock_y - 8, lock_x + 4, lock_y + 8,
                            fill="#7F8C8D", outline="#5D6D7E", width=1)
        cv.create_oval(lock_x - 2, lock_y - 2, lock_x + 2, lock_y + 2,
                       fill="#4A5568", outline="")

        # ── Status LEDs row ───────────────────────────────────────────────
        leds = [("#27AE60","ON"), ("#F4D03F","WARN"), ("#C0392B","ALM"), ("#3498DB","COM")]
        led_y = panel_y + max(6, int(panel_h * 0.03))
        for li, (lc, lt) in enumerate(leds):
            lx = panel_x + max(10, int(panel_w * 0.03)) + li * max(22, int(panel_w * 0.06))
            # LED housing
            cv.create_oval(lx - 1, led_y - 1, lx + max(10, int(panel_w * 0.025)) + 1,
                           led_y + max(10, int(panel_w * 0.025)) + 1,
                           fill="#1A252F", outline="#4A5568")
            led_r = max(4, int(panel_w * 0.01))
            cv.create_oval(lx + led_r // 2, led_y + led_r // 2,
                           lx + led_r * 3, led_y + led_r * 3,
                           fill=lc, outline="white", width=1)
            # LED glow effect
            cv.create_oval(lx + led_r, led_y + led_r, lx + led_r * 2, led_y + led_r * 2,
                           fill="white", outline="")
            cv.create_text(lx + led_r * 2, led_y + led_r * 3 + max(3, fss),
                           text=lt, font=("Arial", max(3, fss - 2)), fill="#BDC3C7")

        # ── High voltage warning sign ─────────────────────────────────────
        warn_x = panel_x + panel_w - max(30, int(panel_w * 0.10))
        warn_y = led_y
        warn_s = max(10, int(panel_h * 0.08))
        # Yellow triangle
        cv.create_polygon(warn_x + warn_s // 2, warn_y, warn_x, warn_y + warn_s,
                          warn_x + warn_s, warn_y + warn_s,
                          fill="#F4D03F", outline="#2C3E50", width=1)
        # Lightning bolt inside
        cv.create_polygon(warn_x + warn_s // 2 - 1, warn_y + 3,
                          warn_x + warn_s // 2 + 2, warn_y + warn_s // 2,
                          warn_x + warn_s // 2 - 1, warn_y + warn_s // 2,
                          warn_x + warn_s // 2 + 1, warn_y + warn_s - 3,
                          fill="#2C3E50", outline="")

        # ── Circuit breaker row ───────────────────────────────────────────
        cb_y = led_y + max(20, int(panel_h * 0.10))
        cb_h = max(12, int(panel_h * 0.08))
        cb_w_each = max(8, int(panel_w * 0.04))
        n_breakers = min(8, panel_w // (cb_w_each + 4))
        cb_start_x = panel_x + (panel_w - n_breakers * (cb_w_each + 3)) // 2
        for bi in range(n_breakers):
            bx = cb_start_x + bi * (cb_w_each + 3)
            # Breaker body
            cv.create_rectangle(bx, cb_y, bx + cb_w_each, cb_y + cb_h,
                                fill="#1A252F", outline="#4A5568", width=1)
            # Toggle switch (up = ON)
            toggle_h = max(4, cb_h // 3)
            cv.create_rectangle(bx + 2, cb_y + 2, bx + cb_w_each - 2, cb_y + 2 + toggle_h,
                                fill="#27AE60" if bi < n_breakers - 1 else "#E74C3C",
                                outline="#1A252F")
            # Label
            cv.create_text(bx + cb_w_each // 2, cb_y + cb_h + max(3, fss // 2),
                           text=f"{bi+1}", font=("Arial", max(3, fss - 2)), fill="#808B96")

        # ── Meters per measurement point ──────────────────────────────────
        n = max(len(tank_names), 1)
        mt_w = (panel_w - 20) // n
        mt_pad = 5
        meters_top = cb_y + cb_h + max(14, int(panel_h * 0.08))

        for i, tn in enumerate(tank_names[:n]):
            mx = panel_x + 10 + i * mt_w
            my = meters_top
            mw = mt_w - mt_pad * 2
            mh = panel_y + panel_h - meters_top - max(10, int(panel_h * 0.06))

            # Meter sub-panel with beveled edge
            cv.create_rectangle(mx + 1, my + 1, mx + mw + 1, my + mh + 1,
                                fill="#0A0E12", outline="")
            cv.create_rectangle(mx, my, mx + mw, my + mh, fill="#1A252F", outline="#5D6D7E", width=2)
            # Header with gradient
            hdr_h = max(12, int(mh * 0.12))
            cv.create_rectangle(mx, my, mx + mw, my + hdr_h, fill="#4A235A", outline="")
            cv.create_rectangle(mx, my + hdr_h - 2, mx + mw, my + hdr_h, fill="#6A1B9A", outline="")
            cv.create_text(mx + mw // 2, my + hdr_h // 2,
                           text=tn[:10], font=FTS, fill="white")

            # kWh readings
            kwh_val   = self.get_var(f"{etapa_key}_{tn}_el_ini_act").get() if etapa_key else ""
            kwh_final = self.get_var(f"{etapa_key}_{tn}_el_fin_act").get() if etapa_key else ""

            # ── Initial kWh display (LCD style) ───────────────────────────
            lcd_y = my + hdr_h + max(4, int(mh * 0.03))
            lcd_h = max(22, int(mh * 0.22))
            cv.create_text(mx + mw // 2, lcd_y - max(2, int(mh * 0.02)),
                           text="kWh INICIAL", font=("Arial", max(4, fss - 2)), fill="#85929E")
            # LCD bezel
            cv.create_rectangle(mx + 3, lcd_y - 1, mx + mw - 3, lcd_y + lcd_h + 1,
                                fill="#0A1520", outline="#1A252F")
            cv.create_rectangle(mx + 4, lcd_y, mx + mw - 4, lcd_y + lcd_h,
                                fill="#0D1B2A", outline="#1ABC9C", width=2)
            # LCD scanline effect
            for sl in range(lcd_y + 3, lcd_y + lcd_h - 2, 3):
                cv.create_line(mx + 6, sl, mx + mw - 6, sl, fill="#0A1822", width=1)
            kwh_disp = kwh_val[:10] if kwh_val else "---"
            cv.create_text(mx + mw // 2, lcd_y + lcd_h // 2, text=kwh_disp,
                           font=("Arial", max(6, min(10, mw // 8))), fill="#1ABC9C")

            # ── Final kWh display ─────────────────────────────────────────
            lcd2_y = lcd_y + lcd_h + max(6, int(mh * 0.04))
            lcd2_h = max(20, int(mh * 0.20))
            cv.create_text(mx + mw // 2, lcd2_y - max(2, int(mh * 0.02)),
                           text="kWh FINAL", font=("Arial", max(4, fss - 2)), fill="#85929E")
            cv.create_rectangle(mx + 3, lcd2_y - 1, mx + mw - 3, lcd2_y + lcd2_h + 1,
                                fill="#0A1520", outline="#1A252F")
            cv.create_rectangle(mx + 4, lcd2_y, mx + mw - 4, lcd2_y + lcd2_h,
                                fill="#0D1B2A", outline="#E74C3C", width=2)
            for sl in range(lcd2_y + 3, lcd2_y + lcd2_h - 2, 3):
                cv.create_line(mx + 6, sl, mx + mw - 6, sl, fill="#0A1218", width=1)
            kwh2_disp = kwh_final[:10] if kwh_final else "---"
            cv.create_text(mx + mw // 2, lcd2_y + lcd2_h // 2, text=kwh2_disp,
                           font=("Arial", max(6, min(10, mw // 8))), fill="#E74C3C")

            # ── Consumption bar with gradient ─────────────────────────────
            dif_y = lcd2_y + lcd2_h + max(6, int(mh * 0.04))
            dif_h = max(12, int(mh * 0.10))
            bar_w = mw - 12
            # Bar background with notches
            cv.create_rectangle(mx + 6, dif_y, mx + 6 + bar_w, dif_y + dif_h,
                                fill="#17202A", outline="#5D6D7E", width=1)
            # Notches on bar
            for ni in range(1, 5):
                nx = mx + 6 + int(bar_w * ni / 5)
                cv.create_line(nx, dif_y, nx, dif_y + dif_h, fill="#2C3E50", width=1)
            try:
                ini_f = float(kwh_val.replace(",", ".") if kwh_val else "0")
                fin_f = float(kwh_final.replace(",", ".") if kwh_final else "0")
                if fin_f > ini_f and ini_f >= 0:
                    pct = min((fin_f - ini_f) / max(ini_f, fin_f, 1), 1.0)
                    bar_filled = int(bar_w * pct)
                    bar_color = "#27AE60" if pct < 0.7 else ("#F39C12" if pct < 0.9 else "#E74C3C")
                    # Gradient fill on bar
                    cv.create_rectangle(mx + 6, dif_y, mx + 6 + bar_filled, dif_y + dif_h,
                                        fill=bar_color, outline="")
                    # Bar highlight
                    cv.create_rectangle(mx + 6, dif_y, mx + 6 + bar_filled, dif_y + max(2, dif_h // 4),
                                        fill="#FFFFFF", outline="", stipple="gray25")
                    cv.create_text(mx + mw // 2, dif_y + dif_h // 2,
                                   text=f"DkWh: {fin_f-ini_f:.0f}", font=FTS, fill="white")
            except: pass

            # ── Analog gauges (Voltmeter and Ammeter) ─────────────────────
            gauge_y = dif_y + dif_h + max(6, int(mh * 0.04))
            remaining_h = my + mh - gauge_y - 4
            for gi, (gl, gc, gv_key, gmax) in enumerate([
                ("V", "#F4D03F", "el_V", 500),
                ("A", "#E74C3C", "el_A", 500)
            ]):
                gx2 = mx + 4 + gi * (mw // 2)
                gy2 = gauge_y
                gr  = max(8, min(mw // 4 - 6, remaining_h // 2 - 4))
                gcx = gx2 + gr + 2
                gcy = gy2 + gr + 2

                # Gauge housing (dark bezel)
                cv.create_oval(gcx - gr - 3, gcy - gr - 3, gcx + gr + 3, gcy + gr + 3,
                               fill="#0A0E12", outline="#4A5568", width=1)
                # Gauge face
                cv.create_oval(gcx - gr, gcy - gr, gcx + gr, gcy + gr,
                               fill="#1A252F", outline="#5D6D7E", width=1)
                # Glass highlight
                cv.create_arc(gcx - gr + 2, gcy - gr + 2, gcx + gr - 2, gcy + gr - 2,
                              start=45, extent=90, style="chord", fill="#2A3540", outline="")
                # Scale arc
                cv.create_arc(gcx - gr + 4, gcy - gr + 4, gcx + gr - 4, gcy + gr - 4,
                              start=30, extent=240, style="arc", outline=gc, width=1)
                # Scale ticks
                for ti in range(11):
                    tick_ang = math.radians(210 - ti * 24)
                    tx1 = gcx + int((gr - 3) * math.cos(tick_ang))
                    ty1 = gcy - int((gr - 3) * math.sin(tick_ang))
                    tx2 = gcx + int((gr - max(4, gr // 4)) * math.cos(tick_ang))
                    ty2 = gcy - int((gr - max(4, gr // 4)) * math.sin(tick_ang))
                    tw = 2 if ti % 5 == 0 else 1
                    cv.create_line(tx1, ty1, tx2, ty2, fill=gc, width=tw)
                # Needle
                gval = self.get_var(f"{etapa_key}_{tn}_{gv_key}").get() if etapa_key else ""
                try:
                    gfrac = min(float(gval.replace(",", ".")) / gmax, 1.0) if gval else 0.0
                    ang = math.radians(210 - 240 * gfrac)
                    gnx = gcx + int(gr * 0.75 * math.cos(ang))
                    gny = gcy - int(gr * 0.75 * math.sin(ang))
                    cv.create_line(gcx, gcy, gnx, gny, fill=gc, width=2)
                except: pass
                # Center pivot
                cv.create_oval(gcx - 2, gcy - 2, gcx + 2, gcy + 2, fill=gc, outline="")
                # Label
                cv.create_text(gcx, gcy + gr + max(3, fss),
                               text=gl, font=FTS, fill=gc)

            # ── Conduit drops from cable tray to each meter ───────────────
            conduit_x = mx + mw // 2
            cv.create_line(conduit_x, y + 4 + tray_h, conduit_x, my, fill="#4A5568", width=2)
            # Conduit entry fitting
            cv.create_rectangle(conduit_x - 3, my - 2, conduit_x + 3, my + 2,
                                fill="#5D6D7E", outline="#4A5568")

        # ── Grounding bus bar (green-yellow stripe per IEC 60446) ─────────
        gnd_y = panel_y + panel_h + max(4, int(H * 0.02))
        gnd_h = max(8, int(H * 0.04))
        # Green-yellow alternating stripe pattern
        stripe_w = max(4, gnd_h)
        for gi in range(0, panel_w, stripe_w * 2):
            gx1 = panel_x + gi
            gx2 = min(panel_x + gi + stripe_w, panel_x + panel_w)
            cv.create_rectangle(gx1, gnd_y, gx2, gnd_y + gnd_h, fill="#2ECC71", outline="")
            gx3 = min(gx2 + stripe_w, panel_x + panel_w)
            if gx3 > gx2:
                cv.create_rectangle(gx2, gnd_y, gx3, gnd_y + gnd_h, fill="#F4D03F", outline="")
        cv.create_rectangle(panel_x, gnd_y, panel_x + panel_w, gnd_y + gnd_h,
                            fill="", outline="#27AE60", width=1)
        # Grounding connections
        for gi in range(0, panel_w, max(15, panel_w // 8)):
            gx = panel_x + gi + 6
            cv.create_oval(gx - 2, gnd_y + 1, gx + 2, gnd_y + gnd_h - 1,
                           fill="#1E8449", outline="")
        # Ground symbol
        gs_x = panel_x + panel_w // 2
        gs_y = gnd_y + gnd_h + 2
        cv.create_line(gs_x, gs_y, gs_x, gs_y + 6, fill="#2ECC71", width=2)
        for gsi, gsw in enumerate([8, 5, 2]):
            cv.create_line(gs_x - gsw, gs_y + 6 + gsi * 3, gs_x + gsw, gs_y + 6 + gsi * 3,
                           fill="#2ECC71", width=1)

        # ── Floor / cable channel ─────────────────────────────────────────
        floor_y = y + H - max(10, int(H * 0.05))
        cv.create_rectangle(x + 1, floor_y, x + W - 1, y + H - 1, fill="#1A1A1A", outline="")
        # Anti-static floor tiles pattern
        tile_w = max(12, W // 15)
        for ti in range(0, W, tile_w):
            cv.create_rectangle(x + ti + 1, floor_y + 1, x + ti + tile_w - 1, y + H - 2,
                                fill="#1E1E1E" if (ti // tile_w) % 2 == 0 else "#222222",
                                outline="#2A2A2A")


    def dibujar_tanque_tierra_tk(self, cv, x, y, W, H, flotante, etapa_key, tank_names=None):
        self._draw_vertical_tank(cv, x, y, W, H, etapa_key, tank_names, flotante=flotante)

    def dibujar_camion_tk(self, cv, x, y, W, H, etapa_key, tank_names=None):
        self._draw_liquid_truck(cv, x, y, W, H, etapa_key, tank_names)

    def dibujar_gasero_tk(self, cv, x, y, W, H, side_label, etapa_key, tank_names=None, carb_names=None):
        self._draw_moss_vessel(cv, x, y, W, H, side_label, etapa_key, tank_names, carb_names)

    def dibujar_metanero_tk(self, cv, x, y, W, H, side_label, etapa_key, tank_names=None, carb_names=None):
        self._draw_membrane_vessel(cv, x, y, W, H, side_label, etapa_key, tank_names, carb_names)

    def dibujar_camion_gas_tk(self, cv, x, y, W, H, etapa_key, tank_names=None):
        self._draw_pressure_truck(cv, x, y, W, H, etapa_key, tank_names)

    def dibujar_ducto_tk(self, cv, x, y, W, H, etapa_key, tipo="GASODUCTO", tank_names=None):
        self._draw_pipeline(cv, x, y, W, H, etapa_key, tipo, tank_names)

    def dibujar_electrico_tk(self, cv, x, y, W, H, etapa_key, tank_names=None):
        self._draw_electric(cv, x, y, W, H, etapa_key, tank_names)

    def dibujar_buque_tk(self, cv, x, y, width, height, side_label, etapa_key, tank_names=None, carb_names=None):
        """Dibujo proporcional al canvas. Sin valores absolutos: todo se escala con width/height."""
        tipo_nave = self.get_var("car_tipo_nave").get() or "BUQUE"
        if tank_names is None: tank_names = self.lista_tanques
        if carb_names is None: carb_names = self.lista_carbonera

        # Forzar aspect ratio apaisado mínimo 3:1 (buque es horizontal)
        min_ratio = 3.0
        if width / max(height, 1) < min_ratio:
            new_height = int(width / min_ratio)
            y = y + (height - new_height) // 2
            height = new_height

        # ── Bézier auxiliar ────────────────────────────────────────────────────
        def bez(p0, p1, p2, p3, n=20):
            pts = []
            for i in range(n + 1):
                t = i / n; u = 1 - t
                pts += [u**3*p0[0]+3*u**2*t*p1[0]+3*u*t**2*p2[0]+t**3*p3[0],
                        u**3*p0[1]+3*u**2*t*p1[1]+3*u*t**2*p2[1]+t**3*p3[1]]
            return pts

        # ── PROPORCIONES (todas relativas a width/height) ──────────────────────
        W = width;  H = height
        # zona de dibujo del casco dentro del rectángulo [x,y,x+W,y+H]
        # reservamos espacio arriba para título y superestructura
        TITLE_H  = max(14, int(H * 0.08))     # altura del título
        SUP_H    = max(40, int(H * 0.28))      # espacio sobre cubierta (superestructura+chimenea+mástil)
        DECK_Y   = y + TITLE_H + SUP_H         # línea de cubierta (Tk: menor Y = arriba)
        KEEL_Y   = y + H - max(6, int(H*0.04)) # quilla (Tk: mayor Y = abajo)
        HULL_H   = KEEL_Y - DECK_Y             # altura total del casco

        # ── Trim visual: leer calado popa/proa para inclinar la línea de flotación ──
        try:
            _trim_popa = self.parse_float(self.get_var(f"{etapa_key}_Calados Popa").get() or "0")
            _trim_proa = self.parse_float(self.get_var(f"{etapa_key}_Calados Proa").get() or "0")
            _trim_val  = _trim_popa - _trim_proa   # positivo = popa hundida
            _list_val  = self.parse_float(self.get_var(f"{etapa_key}_Lista").get() or "0")
        except:
            _trim_val = 0.0; _list_val = 0.0

        # Línea de flotación base: la obra viva ocupa ~35% del casco
        WL_Y_CENTER = DECK_Y + int(HULL_H * 0.65)  # waterline central
        # Ajuste de trim: max ±8% del casco en popa/proa visualmente
        _trim_px = max(-int(HULL_H*0.10), min(int(HULL_H*0.10), int(_trim_val * 0.5)))
        WL_Y_STERN = WL_Y_CENTER + _trim_px    # popa (izquierda): más baja si trim positivo
        WL_Y_BOW   = WL_Y_CENTER - _trim_px    # proa (derecha): más alta si trim positivo
        WL_Y       = WL_Y_CENTER               # waterline central para refs de tanques

        BOW_X    = x + W - max(14, int(W * 0.03))   # proa (derecha)
        STERN_X  = x + max(14, int(W * 0.03))        # popa (izquierda)
        CX       = x + W // 2

        # superestructura (escala proporcional)
        CAS_W    = max(60, int(W * 0.12))
        CAS_X    = STERN_X + max(8, int(W * 0.02))
        F1_H     = max(14, int(SUP_H * 0.30))    # nivel 1 (acomodaciones)
        F2_H     = max(11, int(SUP_H * 0.24))    # nivel 2 (puente)
        F3_H     = max(8,  int(SUP_H * 0.18))    # nivel 3 (ala)
        CHIM_H   = max(14, int(SUP_H * 0.28))    # chimenea
        CHIM_W   = max(10, int(CAS_W * 0.25))
        MAST_H   = max(20, int(SUP_H * 0.70))    # mástil de proa

        # fuentes proporcionales
        # Fuente fija (bitmap) — evita X11 RenderAddGlyphs | +25% sobre versión anterior
        _sz_tk   = min(10, max(6, int(H * 0.038)))   # tanque +25%
        _sz_pct  = min(10, max(6, int(H * 0.033)))   # porcentaje +25%
        _sz_prod = min(9,  max(5, int(H * 0.028)))   # producto +25%
        _sz_ttl  = min(9,  max(5, int(H * 0.035)))   # título +25%
        FONT_TITLE = ("Arial", _sz_ttl)
        FONT_TK    = ("Arial", _sz_tk)
        FONT_PCT   = ("Arial", _sz_pct)
        FONT_PROD  = ("Arial", _sz_prod)

        # ── Fondo neutro ──────────────────────────────────────────────────────────
        cv.create_rectangle(x, y, x+W, y+H, fill="#E8F0FE", outline="#5D6D7E", width=2)
        # ── TÍTULO ────────────────────────────────────────────────────────────
        cv.create_text(x + W//2, y + TITLE_H//2,
                       text=f"VISTA {side_label.upper()}  [{etapa_key.upper() if etapa_key else '-'}]",
                       font=FONT_TITLE, fill="#1B3A5C")

        if tipo_nave == "BARCAZA":
            cv.create_rectangle(STERN_X, DECK_Y, BOW_X, KEEL_Y,
                                fill="#2C3E50", outline="#1B2631", width=2)
            cv.create_rectangle(STERN_X, DECK_Y, BOW_X, DECK_Y + max(5, int(HULL_H*0.05)),
                                fill="#5D6D7E", outline="")
            cv.create_rectangle(STERN_X+1, WL_Y, BOW_X-1, KEEL_Y-1,
                                fill="#922B21", outline="")
            cv.create_line(STERN_X-8, WL_Y, BOW_X+8, WL_Y, fill="#2E86C1", width=2, dash=(5,3))

        else:
            # proa_offset = cómo se extiende la proa más allá de BOW_X
            proa = int(W * 0.045)

            # — Obra muerta con degradado vertical (más oscuro arriba, borda abajo) —
            cp_off = int(HULL_H * 0.12)
            bow_top = bez(
                (BOW_X,        DECK_Y),
                (BOW_X+proa,   DECK_Y),
                (BOW_X+proa,   WL_Y_BOW - cp_off),
                (BOW_X+proa//2, WL_Y_BOW)
            )
            top_poly = [STERN_X-4, DECK_Y, BOW_X, DECK_Y] + bow_top + [STERN_X-4, WL_Y_STERN]
            cv.create_polygon(top_poly, fill="#17202A", outline="")
            # Franja intermedia de la obra muerta (degradado manual)
            mid_wl = DECK_Y + int((WL_Y-DECK_Y)*0.5)
            cv.create_polygon([STERN_X-4, mid_wl, BOW_X, mid_wl] + bow_top + [STERN_X-4, WL_Y],
                              fill="#1B2631", outline="")
            # Borda (franja clara en el borde superior)
            borda_h = max(3, int(HULL_H*0.04))
            cv.create_polygon([STERN_X-4, DECK_Y, BOW_X, DECK_Y] + bez(
                (BOW_X, DECK_Y), (BOW_X+proa//2, DECK_Y), (BOW_X+proa//2, DECK_Y+borda_h), (BOW_X+proa//4, DECK_Y+borda_h)
            ) + [STERN_X-4, DECK_Y+borda_h], fill="#F0F3F4", outline="")
            # Riel de cubierta (línea de guarda)
            cv.create_line(STERN_X-4, DECK_Y+borda_h, BOW_X, DECK_Y+borda_h, fill="#808B96", width=2)
            # Línea de flotación activa (inclinada con trim)
            cv.create_line(STERN_X-4, WL_Y_STERN+1, BOW_X+proa//2, WL_Y_BOW+1,
                           fill="#F0F0F0", width=2)

            # Indicator de Trim visible si hay trim significativo
            if abs(_trim_val) > 0.1:
                trim_lbl = f"TRIM: {_trim_val:+.2f}m"
                trim_color = "#F4D03F" if abs(_trim_val) < 1.0 else "#E74C3C"
                cv.create_text(CX, WL_Y - max(5, int(HULL_H*0.08)),
                               text=trim_lbl, font=("Arial", max(5, int(H*0.028))),
                               fill=trim_color)

            # — Obra viva con degradado rojo-bordo —
            bow_bot = bez(
                (BOW_X+proa//2, WL_Y_BOW),
                (BOW_X+proa//2, KEEL_Y - int(HULL_H*0.05)),
                (BOW_X - int(W*0.04), KEEL_Y),
                (CX,             KEEL_Y + int(HULL_H*0.03))
            )
            stern_bot = bez(
                (CX,            KEEL_Y + int(HULL_H*0.03)),
                (STERN_X + int(W*0.05), KEEL_Y),
                (STERN_X,       KEEL_Y - int(HULL_H*0.02)),
                (STERN_X-4,     WL_Y_STERN)
            )
            bot_poly = [STERN_X-4, WL_Y_STERN] + bow_bot + stern_bot
            cv.create_polygon(bot_poly, fill="#922B21", outline="")
            # Franja más oscura en la quilla
            quilla_h = max(3, int(HULL_H*0.10))
            keel_poly = [STERN_X-4, WL_Y+int((KEEL_Y-WL_Y)*0.78)] + stern_bot[-4:] + bow_bot[-4:] + [BOW_X+proa//2, WL_Y+int((KEEL_Y-WL_Y)*0.78)]
            try: cv.create_polygon(keel_poly, fill="#641E16", outline="")
            except: pass
            # Línea de quilla
            cv.create_line(STERN_X+4, KEEL_Y, CX, KEEL_Y+int(HULL_H*0.035), fill="#1B2631", width=2)
            # Línea de flotación activa con olas (trim-aware)
            cv.create_line(STERN_X-10, WL_Y_STERN, BOW_X+proa//2, WL_Y_BOW,
                           fill="#5DADE2", width=2, dash=(5,3))

            # — Marca de Plimsoll (línea de carga) —
            plim_x = int(CX)
            plim_y = WL_Y_CENTER
            plim_r = max(6, int(HULL_H*0.09))
            # Círculo Plimsoll
            cv.create_oval(plim_x-plim_r, plim_y-plim_r, plim_x+plim_r, plim_y+plim_r,
                           fill="", outline="#F0F0F0", width=2)
            # Línea horizontal a través del círculo
            cv.create_line(plim_x-plim_r-8, plim_y, plim_x+plim_r+8, plim_y,
                           fill="#F0F0F0", width=2)
            # Letras de carga con líneas de nivel (TF, T, S, W, WNA)
            if plim_r > 7:
                fplim = max(4, plim_r//3)
                load_lines = [("S", -plim_r*1.8), ("T", -plim_r*3.2), ("F", -plim_r*4.6)]
                for label_pl, offset_pl in load_lines:
                    lx = plim_x - plim_r - 4
                    ly = WL_Y + int(offset_pl)
                    cv.create_line(lx-8, ly, lx+4, ly, fill="#F0F0F0", width=1)
                    if fplim >= 4:
                        cv.create_text(lx-10, ly, text=label_pl, font=("Arial", fplim),
                                       fill="#F0F0F0", anchor="e")
            # LR / Bureau Veritas letters inside Plimsoll circle
            if plim_r > 8:
                cv.create_text(plim_x, plim_y-plim_r//3, text="LR",
                               font=("Arial", max(4, plim_r//3), "bold"), fill="#F0F0F0")

            # ── MARCAS DE CALADO PROA/POPA (Draft Marks) ─────────────────
            # Estas marcas son características distintivas de los buques:
            # números pintados en el casco indicando el calado en esa sección
            _fdrft = ("Arial", max(4, int(HULL_H*0.065)), "bold")
            _fdsmall = ("Arial", max(3, int(HULL_H*0.050)))

            # Calados medidos: leer de las variables o estimar del trim
            try:
                _c_proa_val = self.parse_float(self.get_var(f"{etapa_key}_Calados Proa").get() or "0")
                _c_popa_val = self.parse_float(self.get_var(f"{etapa_key}_Calados Popa").get() or "0")
            except:
                _c_proa_val = 0.0; _c_popa_val = 0.0

            def _draw_draft_scale(cv2, dx, dy_wl, dy_keel, calado_val, anchor_side="w"):
                """Dibuja escala de calado en dx, desde keel hasta sobre WL."""
                _scale_h = dy_keel - dy_wl
                if _scale_h < 10: return
                # Rango de calado: asumir 0 en quilla → calado máximo en WL
                # Mostramos marcas cada 0.5m (aprox HULL_H/escala)
                _pix_per_m = _scale_h / max(calado_val, 4.0) if calado_val > 0 else _scale_h / 6.0
                # Línea vertical de escala
                _sx1 = dx - 3 if anchor_side == "e" else dx + 3
                cv2.create_line(dx, dy_wl - 6, dx, dy_keel + 2, fill="#D0D0D0", width=1)
                # Marcas numéricas (cada 1m)
                _num_marks = int(calado_val) + 1 if calado_val > 0 else 5
                for _mi in range(0, min(_num_marks + 2, 12)):
                    _my = int(dy_keel - _mi * _pix_per_m)
                    if _my < dy_wl - 8: break
                    if _my > dy_keel + 2: continue
                    _is_integer = True
                    # Marca corta para enteros
                    _llen = 6 if _is_integer else 3
                    _lx1 = dx - _llen if anchor_side == "e" else dx
                    _lx2 = dx if anchor_side == "e" else dx + _llen
                    cv2.create_line(_lx1, _my, _lx2, _my, fill="#D0D0D0", width=1)
                    if _mi > 0 and _fdsmall[1] >= 3:
                        _tx = dx - _llen - 2 if anchor_side == "e" else dx + _llen + 2
                        cv2.create_text(_tx, _my, text=str(_mi), font=_fdsmall,
                                        fill="#CCCCCC", anchor=anchor_side)
                # Línea de flotación actual (destacada con valor medido)
                if calado_val > 0:
                    _wl_mark_y = int(dy_keel - calado_val * _pix_per_m)
                    if dy_wl - 10 < _wl_mark_y < dy_keel:
                        # Flecha/indicador del calado medido
                        _arr_len = 10
                        _ax = dx + _arr_len if anchor_side == "e" else dx - _arr_len
                        cv2.create_line(dx - 3 if anchor_side=="e" else dx + 3,
                                        _wl_mark_y, _ax, _wl_mark_y,
                                        fill="#F4D03F", width=2)
                        cv2.create_text(_ax + (3 if anchor_side=="w" else -3),
                                        _wl_mark_y,
                                        text=f"{calado_val:.2f}m",
                                        font=("Arial", max(4, int(HULL_H*0.055)), "bold"),
                                        fill="#F4D03F",
                                        anchor="w" if anchor_side=="w" else "e")

            # Escala en PROA (derecha del dibujo)
            _bow_draft_x = BOW_X - max(6, int(W*0.015))
            _draw_draft_scale(cv, _bow_draft_x, WL_Y_BOW, KEEL_Y,
                              _c_proa_val if _c_proa_val > 0 else abs(_trim_val)+3.5, "w")

            # Escala en POPA (izquierda del dibujo)
            _stern_draft_x = STERN_X + max(6, int(W*0.015))
            _draw_draft_scale(cv, _stern_draft_x, WL_Y_STERN, KEEL_Y,
                              _c_popa_val if _c_popa_val > 0 else abs(_trim_val)+3.5, "e")

            # Label "PROA" y "POPA" sobre las escalas
            if _fdrft[1] >= 4:
                cv.create_text(_bow_draft_x, DECK_Y + int(HULL_H*0.04),
                               text="PROA", font=_fdsmall, fill="#A0A0A0", anchor="w")
                cv.create_text(_stern_draft_x, DECK_Y + int(HULL_H*0.04),
                               text="POPA", font=_fdsmall, fill="#A0A0A0", anchor="e")

            # — Hawse pipe / escobén (proa) —
            haw_x = int(BOW_X - max(8, int(W*0.025)))
            haw_y = int(DECK_Y + HULL_H*0.15)
            haw_r = max(4, int(HULL_H*0.05))
            cv.create_oval(haw_x-haw_r, haw_y-haw_r, haw_x+haw_r, haw_y+haw_r,
                           fill="#1B2631", outline="#808B96", width=2)
            # Cadena del ancla saliendo del escobén
            chain_y = haw_y + haw_r
            chain_pts = []
            import math as _math
            for ci_ch in range(10):
                cx_ch = haw_x + ci_ch*max(3, int(HULL_H*0.02))
                cy_ch = chain_y + int(_math.sin(ci_ch*0.8) * max(2, int(HULL_H*0.03))) + ci_ch*max(1, int(HULL_H*0.01))
                chain_pts.extend([cx_ch, cy_ch])
            if len(chain_pts) >= 4:
                try: cv.create_line(chain_pts, fill="#7F8C8D", width=2, smooth=True)
                except: pass
            # Ancla en el costado
            anc_x = int(BOW_X - max(6, int(W*0.03)))
            anc_y = int(DECK_Y - max(5, int(HULL_H*0.05)))
            anc_r = max(3, int(H*0.018))
            cv.create_oval(anc_x-anc_r, anc_y-anc_r, anc_x+anc_r, anc_y+anc_r,
                           fill="#7F8C8D", outline="#5D6D7E", width=1)
            cv.create_line(anc_x, anc_y, anc_x, anc_y+anc_r*2, fill="#7F8C8D", width=max(2,anc_r//2))
            cv.create_line(anc_x-anc_r, anc_y+anc_r, anc_x+anc_r, anc_y+anc_r,
                           fill="#7F8C8D", width=max(1,anc_r//3))


            # Se ubica ENCIMA de la cubierta (menor Y que DECK_Y)
            # Superestructura con pisos
            F0_H  = max(10, int(SUP_H * 0.22))
            f1_bot = DECK_Y;             f1_top = f1_bot - F1_H
            f2_bot = f1_top;             f2_top = f2_bot - F2_H
            f3_bot = f2_top;             f3_top = f3_bot - F3_H
            chim_bot = f3_top;           chim_top = chim_bot - CHIM_H
            chim_x = CAS_X + max(3, CHIM_W//3)

            # Nivel 0 — plataforma base (más ancha)
            cas_w0 = int(CAS_W * 1.10)
            cas_x0 = CAS_X - (cas_w0 - CAS_W)//2
            cv.create_rectangle(cas_x0+2, f1_bot-F0_H+2, cas_x0+cas_w0+2, f1_bot+2, fill="#8D99A4", outline="")
            cv.create_rectangle(cas_x0, f1_bot-F0_H, cas_x0+cas_w0, f1_bot, fill="#D0D3D4", outline="#ABB2B9", width=1)

            # Nivel 1 — acomodaciones
            cv.create_rectangle(CAS_X+2, f1_top+2, CAS_X+CAS_W+2, f1_bot+2, fill="#7D8C93", outline="")
            cv.create_rectangle(CAS_X, f1_top, CAS_X+CAS_W, f1_bot, fill="#E8EAEB", outline="#ABB2B9", width=1)
            # (ventanas acomodaciones removidas)

            # Nivel 2 — puente de mando (más estrecho)
            br_indent = max(2, CAS_W//10)
            cv.create_rectangle(CAS_X+br_indent+2, f2_top+2, CAS_X+CAS_W-br_indent+2, f2_bot+2, fill="#7D8C93", outline="")
            cv.create_rectangle(CAS_X+br_indent, f2_top, CAS_X+CAS_W-br_indent, f2_bot, fill="#E0E3E5", outline="#ABB2B9", width=1)
            # (ventana panorámica del puente removida)
            # Aletas del puente
            wing_ext = max(8, int(CAS_W*0.25))
            cv.create_rectangle(cas_x0, f2_top, CAS_X+br_indent, f2_bot, fill="#D0D3D4", outline="#95A5A6", width=1)
            cv.create_rectangle(CAS_X+CAS_W-br_indent, f2_top, CAS_X+CAS_W+wing_ext//2, f2_bot, fill="#D0D3D4", outline="#95A5A6", width=1)

            # Nivel 3 — techo con pasamanos
            cv.create_rectangle(CAS_X+br_indent//2, f3_top, CAS_X+CAS_W-br_indent//2, f3_bot, fill="#BDC3C7", outline="#95A5A6", width=1)
            rng_step = max(4, CAS_W//10)
            for pr_x in range(CAS_X+2, CAS_X+CAS_W-2, rng_step):
                cv.create_line(pr_x, f3_top-4, pr_x, f3_top, fill="#7F8C8D", width=1)
            cv.create_line(CAS_X+2, f3_top-4, CAS_X+CAS_W-2, f3_top-4, fill="#7F8C8D", width=1)

            # Radar/antena
            rad_cx = CAS_X + CAS_W//2
            rad_cy = f3_top - max(5, int(SUP_H*0.12))
            rad_r  = max(4, int(CAS_W*0.08))
            cv.create_line(rad_cx, f3_top, rad_cx, rad_cy, fill="#95A5A6", width=2)
            cv.create_arc(rad_cx-rad_r, rad_cy-rad_r, rad_cx+rad_r, rad_cy+rad_r,
                          start=0, extent=180, style="arc", outline="#7F8C8D", width=2)
            cv.create_line(rad_cx-rad_r, rad_cy, rad_cx+rad_r, rad_cy, fill="#95A5A6", width=1)

            # Chimenea
            cv.create_rectangle(chim_x+3, chim_top+4, chim_x+CHIM_W+3, chim_bot+3, fill="#641E16", outline="")
            cv.create_rectangle(chim_x, chim_top, chim_x+CHIM_W, chim_bot, fill="#B03A2E", outline="#7B241C", width=2)
            cv.create_rectangle(chim_x-2, chim_top, chim_x+CHIM_W+2, chim_top+max(3,CHIM_H//8),
                                fill="#808B96", outline="#5D6D7E", width=1)
            stripe_h = max(2, CHIM_H//5)
            cv.create_rectangle(chim_x+1, chim_top+stripe_h, chim_x+CHIM_W-1, chim_top+stripe_h*2,
                                fill="#212F3D", outline="")
            cv.create_rectangle(chim_x+1, chim_top+stripe_h*2+1, chim_x+CHIM_W-1, chim_top+stripe_h*3+1,
                                fill="#F4D03F", outline="")
            for sm_i in range(4):
                sm_r  = max(3, int(CHIM_H * (0.18 + sm_i*0.10)))
                sm_ox = int(sm_i * max(1, CHIM_W*0.10))
                sm_y  = chim_top - sm_i * max(3, CHIM_H//6) - sm_r
                smoke_shades = ["#95A5A6","#AAB7B8","#BDC3C7","#D5D8DC"]
                cv.create_oval(chim_x+CHIM_W//2-sm_r+sm_ox, sm_y,
                               chim_x+CHIM_W//2+sm_r+sm_ox, sm_y+sm_r*2,
                               fill=smoke_shades[sm_i], outline="")

                        # Mástil de proa con obenques (stays)
            mast_x = BOW_X - max(16, int(W*0.04))
            cv.create_line(mast_x, DECK_Y, mast_x, DECK_Y - MAST_H, fill="#626567", width=3)
            # Obenques
            stay_bot_1 = int(mast_x - MAST_H*0.35)
            stay_bot_2 = int(mast_x + MAST_H*0.25)
            cv.create_line(stay_bot_1, DECK_Y, mast_x, DECK_Y-int(MAST_H*0.8), fill="#7F8C8D", width=1, dash=(3,3))
            cv.create_line(stay_bot_2, DECK_Y, mast_x, DECK_Y-int(MAST_H*0.8), fill="#7F8C8D", width=1, dash=(3,3))
            # Verga horizontal
            cv.create_line(mast_x - max(7,MAST_H//4), DECK_Y-int(MAST_H*0.72),
                          mast_x + max(7,MAST_H//4), DECK_Y-int(MAST_H*0.72), fill="#626567", width=2)
            # Luces de navegación (tricolor: verde, rojo)
            cv.create_oval(mast_x-3, DECK_Y-int(MAST_H*0.72)-5, mast_x+3, DECK_Y-int(MAST_H*0.72)+5,
                          fill="#2ECC71", outline="")  # verde estribor
            cv.create_oval(mast_x-max(7,MAST_H//4)-5, DECK_Y-int(MAST_H*0.72)-3,
                           mast_x-max(7,MAST_H//4)+5, DECK_Y-int(MAST_H*0.72)+3,
                           fill="#E74C3C", outline="")  # rojo babor
            # Luz blanca en tope
            r_luz = max(3, MAST_H//8)
            cv.create_oval(mast_x-r_luz, DECK_Y-MAST_H-r_luz, mast_x+r_luz, DECK_Y-MAST_H+r_luz,
                          fill="#F8F9FA", outline="#BDC3C7")
            # Bandera en la cima (argentina)
            flag_w = max(10, MAST_H//3); flag_h = max(7, MAST_H//5)
            cv.create_rectangle(mast_x, DECK_Y-MAST_H,
                                mast_x+flag_w, DECK_Y-MAST_H+flag_h,
                                fill="#74B9FF", outline="#2980B9", width=1)
            cv.create_rectangle(mast_x, DECK_Y-MAST_H+flag_h//3,
                                mast_x+flag_w, DECK_Y-MAST_H+flag_h*2//3,
                                fill="#FFEAA7", outline="")
            cv.create_rectangle(mast_x, DECK_Y-MAST_H+flag_h*2//3,
                                mast_x+flag_w, DECK_Y-MAST_H+flag_h,
                                fill="#74B9FF", outline="")


            # ── Ojos de buey (portholes) ──────────────────────────────────
            n_ph = max(3, min(8, int((BOW_X - CAS_X - CAS_W - 30) // max(14, int(W*0.025)))))
            ph_y = WL_Y - int(HULL_H*0.28)
            ph_r = max(3, int(HULL_H*0.055))
            ph_step = (BOW_X - CAS_X - CAS_W - 20) // max(n_ph, 1)
            for phi in range(n_ph):
                ph_x = CAS_X + CAS_W + 10 + phi*ph_step
                cv.create_oval(ph_x-ph_r, ph_y-ph_r, ph_x+ph_r, ph_y+ph_r,
                               fill="#AED6F1", outline="#1B2631", width=2)
                cv.create_oval(ph_x-ph_r+2, ph_y-ph_r+2, ph_x, ph_y,
                               fill="white", outline="")  # reflejo

            # ── Nombre del buque en el casco ───────────────────────────────
            buq_txt = self.get_var("car_buque").get()[:18] if self.get_var("car_buque").get() else ""
            if buq_txt and tipo_nave != "BARCAZA":
                mid_hull_x = (CAS_X + CAS_W + BOW_X) // 2
                cv.create_text(mid_hull_x, ph_y + int(HULL_H*0.22),
                               text=buq_txt, font=("Arial", max(4, int(HULL_H*0.09))),
                               fill="#F0F0F0")

            # ── Grúa de cubierta (solo si el buque no es metanero/gasero) ──
            if tipo_nave in ("BUQUE",) and int(W*0.08) > 20:
                gc_x = int(CAS_X + CAS_W + (BOW_X - CAS_X - CAS_W)*0.55)
                gc_base_y = DECK_Y
                gc_h = int(SUP_H * 0.55)
                cv.create_rectangle(gc_x-4, gc_base_y-gc_h//3, gc_x+4, gc_base_y,
                                    fill="#5D6D7E", outline="#4A5568", width=1)
                # Pluma de la grúa
                cv.create_line(gc_x, gc_base_y-gc_h//3, gc_x+int(W*0.06), gc_base_y-gc_h,
                               fill="#4A5568", width=3)
                # Cable con gancho
                cable_x = gc_x+int(W*0.04); cable_top = gc_base_y-int(gc_h*0.75)
                cv.create_line(cable_x, cable_top, cable_x, cable_top+int(gc_h*0.35),
                               fill="#7F8C8D", width=1)
                cv.create_oval(cable_x-3, cable_top+int(gc_h*0.35)-3,
                               cable_x+3, cable_top+int(gc_h*0.35)+3,
                               fill="#E74C3C", outline="#C0392B")

            # ── Escotillas de carga (cargo hatches) en cubierta ──────────
            hatch_zone_start = CAS_X + CAS_W + max(8, int(W*0.015))
            hatch_zone_end   = BOW_X - max(20, int(W*0.04))
            hatch_zone_w     = hatch_zone_end - hatch_zone_start
            n_hatches = max(2, min(6, hatch_zone_w // max(20, int(W*0.07))))
            hatch_w   = hatch_zone_w // n_hatches
            hatch_h   = max(4, int(HULL_H*0.10))
            hatch_top = DECK_Y - hatch_h
            for hi in range(n_hatches):
                hx  = hatch_zone_start + hi*hatch_w + max(2, hatch_w//8)
                hw2 = hatch_w - max(4, hatch_w//4)
                # Tapa de la escotilla
                cv.create_rectangle(hx, hatch_top, hx+hw2, DECK_Y,
                                    fill="#2E4053", outline="#1A252F", width=1)
                # Panel central de la escotilla
                cv.create_rectangle(hx+2, hatch_top+2, hx+hw2-2, DECK_Y-1,
                                    fill="#2C3E50", outline="")
                # Resaltado superior (brillo metálico)
                cv.create_rectangle(hx+1, hatch_top, hx+hw2-1, hatch_top+2,
                                    fill="#5D6D7E", outline="")
                # Pestillo central
                pst_x = hx + hw2//2
                cv.create_rectangle(pst_x-2, hatch_top+hatch_h//3, pst_x+2, hatch_top+2*hatch_h//3,
                                    fill="#E74C3C", outline="")

            # ── Bollards (bitas de amarre) en cubierta ────────────────────
            bollard_positions = [int(STERN_X + (BOW_X-STERN_X)*f) for f in [0.08, 0.18, 0.85, 0.94]]
            boll_h = max(3, int(HULL_H*0.06))
            for bpx in bollard_positions:
                cv.create_rectangle(bpx-3, DECK_Y-boll_h, bpx+3, DECK_Y,
                                    fill="#808B96", outline="#5D6D7E", width=1)
                # Cabeza de la bita (más ancha)
                cv.create_rectangle(bpx-5, DECK_Y-boll_h, bpx+5, DECK_Y-boll_h+3,
                                    fill="#7F8C8D", outline="")

            # ── Timón y hélice en la popa ─────────────────────────────────
            rud_x = STERN_X - max(4, int(W*0.008))
            # Hélice (círculo con aspas)
            prop_cx = rud_x - max(5, int(W*0.012))
            prop_cy = WL_Y + int(HULL_H*0.55)
            prop_r  = max(5, int(HULL_H*0.14))
            # Aspas de la hélice
            for ang in [0, 120, 240]:
                pa_x = prop_cx + int(prop_r * math.cos(math.radians(ang)))
                pa_y = prop_cy + int(prop_r * math.sin(math.radians(ang)))
                cv.create_oval(pa_x-max(3,prop_r//4), pa_y-max(3,prop_r//4),
                               pa_x+max(3,prop_r//4), pa_y+max(3,prop_r//4),
                               fill="#5D6D7E", outline="")
            cv.create_oval(prop_cx-2, prop_cy-2, prop_cx+2, prop_cy+2,
                           fill="#95A5A6", outline="")
            # Bocina / eje de la hélice
            cv.create_line(prop_cx, prop_cy, rud_x+3, prop_cy, fill="#4A5568", width=3)
            # Timón
            rud_top = WL_Y + int(HULL_H*0.20)
            rud_bot = WL_Y + int(HULL_H*0.75)
            rud_w   = max(3, int(W*0.008))
            rud_h   = rud_bot - rud_top
            cv.create_polygon(rud_x, rud_top, rud_x+rud_w, rud_top+max(2,rud_h)//5,
                              rud_x+rud_w, rud_top+rud_h*4//5, rud_x, rud_bot,
                              fill="#4A5568", outline="#2C3E50", width=1)

        # ══════════ TANQUES ══════════════════════════════════════════════════
        target_side = side_label.upper()
        tanks_to_draw = [t for t in tank_names if target_side in t.upper()]
        if tanks_to_draw:
            # Zona: entre cubierta y línea de flotación
            t_top = DECK_Y + 2
            t_bot = WL_Y - 2
            t_h   = t_bot - t_top
            if t_h < 5: t_h = 5

            # Zona horizontal: entre superestructura y mástil
            sup_end = CAS_X + CAS_W + max(4, int(W*0.01))
            mast_x2 = BOW_X - max(16, int(W*0.04))
            t_start = sup_end
            t_end   = mast_x2 - max(4, int(W*0.01))
            total_tw = t_end - t_start
            if total_tw < 20: t_start = STERN_X + 4; total_tw = BOW_X - STERN_X - 8

            n_t = len(tanks_to_draw)
            t_w = total_tw / n_t
            gap = max(1.5, t_w * 0.04)

            PRODUCT_COLORS = ["#C0392B","#E67E22","#27AE60","#D4AC0D","#784212","#2C3E50","#7D3C98","#C2185B"]
            cur_tx = t_start
            for t_nm in tanks_to_draw:
                vliq = self.parse_float(self.get_var(f"{etapa_key}_{t_nm}_s_corr").get()) if etapa_key else 0
                vref = self.parse_float(self.get_var(f"{etapa_key}_{t_nm}_alt_ref").get()) if etapa_key else 0
                vwat = self.parse_float(self.get_var(f"{etapa_key}_{t_nm}_agua_s_real").get()) if etapa_key else 0
                vlit = self.get_var(f"{etapa_key}_{t_nm}_vol_nat_prod").get() if etapa_key else ""
                pnm  = self.get_var(f"{etapa_key}_{t_nm}_prod_name").get() if etapa_key else ""
                ref_h = vref if vref > 0 else 10000.0
                alt   = max(0.0, vref - vliq)
                pct_l = min(alt / ref_h, 1.0)
                pct_w = min(vwat / ref_h, 1.0) if vwat > 0 else 0.0
                px_f  = t_h * pct_l
                px_w  = t_h * pct_w

                tx = cur_tx + gap/2;  tw = t_w - gap
                # Fondo acero (con gradiente lateral)
                cv.create_rectangle(tx, t_top, tx+tw, t_bot, fill="#3D4B56", outline="")
                # Franjas de sombreado lateral (efecto 3D)
                sh_tk = max(2, tw//8)
                cv.create_rectangle(tx, t_top, tx+sh_tk, t_bot, fill="#2C3E50", outline="")
                cv.create_rectangle(tx+tw-sh_tk, t_top, tx+tw, t_bot, fill="#2A3742", outline="")
                cv.create_rectangle(tx, t_top, tx+tw, t_bot, fill="", outline="#2C3E50", width=1)
                # Agua (desde el fondo)
                if px_w > 0:
                    cv.create_rectangle(tx+1, t_bot-px_w, tx+tw-1, t_bot, fill="#5DADE2", outline="")
                # Producto — SIN stipple
                px_p = px_f - px_w
                ci_col = 0
                if px_p > 0:
                    dv = self.parse_float(self.get_var(f"{etapa_key}_{t_nm}_dens_lab").get()) if etapa_key else 0
                    ci_col = abs(int(round(dv,3)*1000)) % len(PRODUCT_COLORS) if dv else 0
                    cv.create_rectangle(tx+1, t_bot-px_f, tx+tw-1, t_bot-px_w,
                                        fill=PRODUCT_COLORS[ci_col], outline="")
                # Tapa
                cv.create_rectangle(tx, t_top, tx+tw, t_top+3, fill="#7F8C8D", outline="")
                # Textos — color legible según fondo del producto
                short = t_nm.replace("BABOR","B").replace("ESTRIBOR","E").strip()
                cyt_t = (t_top + t_bot) / 2
                # Calcular luminancia del color de producto para elegir letra legible
                _pc = PRODUCT_COLORS[ci_col] if px_p > 0 else "#3D4B56"
                try:
                    _r,_g,_b = int(_pc[1:3],16),int(_pc[3:5],16),int(_pc[5:7],16)
                    _lum = 0.299*_r + 0.587*_g + 0.114*_b
                    _txt_col = "#000000" if _lum > 140 else "white"
                    _pnm_col = "#000000" if _lum > 140 else "#E8F8F5"
                except: _txt_col = "white"; _pnm_col = "#E8F8F5"
                # Distribuir textos centrados evitando el borde inferior del tanque
                # 4 líneas: nombre, %, producto, litros — espaciado uniforme en zona segura
                _lines = [short[:14]]         # siempre
                _lines.append(f"{pct_l*100:.0f}%")
                if pnm:  _lines.append(pnm[:14])
                if vlit: _lines.append(f"{vlit} L")
                _n = len(_lines)
                # zona de texto: desde 15% desde arriba hasta 85% (evita tapa y línea de agua)
                _y_top  = t_top + t_h * 0.15
                _y_bot  = t_top + t_h * 0.82
                _step   = (_y_bot - _y_top) / max(_n - 1, 1) if _n > 1 else 0
                _fonts  = [FONT_TK, FONT_PCT] + [FONT_PROD] * max(0, _n - 2)
                _colors = [_txt_col, _txt_col] + [_pnm_col] * max(0, _n - 2)
                for _li, (_lt, _lf, _lc) in enumerate(zip(_lines, _fonts, _colors)):
                    _ly = _y_top + _li * _step if _n > 1 else (t_top + t_bot) / 2
                    cv.create_text(tx+tw/2, _ly, text=_lt, font=_lf, fill=_lc)
                cur_tx += t_w

        # ══════════ CARBONERAS ════════════════════════════════════════════════
        carbs = [c for c in carb_names
                 if target_side in c.upper() or ("BABOR" not in c.upper() and "ESTRIBOR" not in c.upper())]
        if carbs:
            czx = CAS_X + 2;  czw = CAS_W - 4
            cih = max(14, int((WL_Y - DECK_Y - 6) / max(len(carbs), 1)))
            for ci2, c_nm in enumerate(carbs):
                cyb = WL_Y - 2 - ci2*(cih+2)
                cyt2 = cyb - cih
                pc2 = 0.0; vc = ""
                if etapa_key:
                    sc = self.parse_float(self.get_var(f"{etapa_key}_{c_nm}_s_corr").get())
                    rc = self.parse_float(self.get_var(f"{etapa_key}_{c_nm}_alt_ref").get())
                    vc = self.get_var(f"{etapa_key}_{c_nm}_vol_nat_prod").get()
                    rh2 = rc if rc > 0 else 10000.0
                    pc2 = min(max(0.0, rh2-sc)/rh2, 1.0)
                px_cc = cih * pc2
                cv.create_rectangle(czx, cyt2, czx+czw, cyb, fill="#FEF9E7", outline="#D4AC0D", width=2)
                if px_cc > 0:
                    cv.create_rectangle(czx+1, cyb-px_cc, czx+czw-1, cyb-1, fill="#F0B429", outline="")
                ccy = (cyt2+cyb)/2
                sc2 = c_nm.replace("BABOR","B").replace("ESTRIBOR","E").strip()
                cv.create_text(czx+czw/2, ccy-3, text=sc2[:12], font=FONT_PROD, fill="#6E4B00")
                parts = ([vc] if vc else [])+[f"{pc2*100:.0f}%"]
                cv.create_text(czx+czw/2, ccy+5, text=" | ".join(parts), font=("Arial", max(4, FONT_PROD[1]-1)), fill="#6E4B00")

    def preview_pdf_temp(self):
        """Genera PDF temporal con TODOS los reportes y lo abre."""
        import tempfile
        try:
            tmpdir = tempfile.gettempdir()
            tmp_path = os.path.join(tmpdir, "_preview_medicion.pdf")
            c_pdf = canvas.Canvas(tmp_path, pagesize=landscape(A4))
            all_tanks = self.lista_tanques + self.lista_carbonera
            report_count = 0

            # 1. Reporte técnico global
            try:
                self.generar_reporte_tecnico_global("PREVIEW", "", shared_canvas=c_pdf)
                report_count += 1
            except: pass

            # 2. Reporte general
            try:
                self.generar_un_reporte("PREVIEW", all_tanks, is_partial=False, output_folder="", shared_canvas=c_pdf)
                report_count += 1
            except: pass

            # 3. Reportes por documento
            mapa = {}
            for tk_name in all_tanks:
                d_ini = self.get_var(f"inicial_{tk_name}_ddt_assign").get()
                d_fin = self.get_var(f"final_{tk_name}_ddt_assign").get()
                if d_ini:
                    if d_ini not in mapa: mapa[d_ini] = []
                    if tk_name not in mapa[d_ini]: mapa[d_ini].append(tk_name)
                if d_fin and d_fin != d_ini:
                    if d_fin not in mapa: mapa[d_fin] = []
                    if tk_name not in mapa[d_fin]: mapa[d_fin].append(tk_name)
            modes = [("SEGÚN_LAB", "dens_lab"), ("SEGÚN_DOC", "dens_doc"), ("SEGÚN_SAL", "dens_salida")]
            for ddt_num, tanks in mapa.items():
                ddt_obj = next((d for d in self.ddt_data if d["numero"].get() == ddt_num), None)
                if not ddt_obj: continue
                for suffix, mode_key in modes:
                    try:
                        self.generar_un_reporte(f"P_{suffix}", tanks, is_partial=True, ddt_obj=ddt_obj, output_folder="", density_mode_key=mode_key, shared_canvas=c_pdf)
                        report_count += 1
                    except: pass

            c_pdf.save()
            if report_count == 0:
                messagebox.showwarning("Vista Previa", "No se pudo generar ningún reporte.")
                return
            if platform.system() == 'Windows': os.startfile(tmp_path)
            elif platform.system() == 'Darwin': subprocess.call(('open', tmp_path))
            else: subprocess.call(('xdg-open', tmp_path))
        except Exception as e:
            traceback.print_exc()
            messagebox.showerror("Error", f"No se pudo generar la vista previa:\n{e}")

    def generar_con_seleccion(self):
        """Diálogo para seleccionar qué reportes generar."""
        dlg = tk.Toplevel(self.root)
        dlg.title("Seleccionar Reportes a Generar")
        dlg.transient(self.root)
        dlg.grab_set()
        dlg.update_idletasks()
        try:
            dlg.state("zoomed")
        except:
            try:
                dlg.attributes("-zoomed", True)
            except:
                _sw, _sh = dlg.winfo_screenwidth(), dlg.winfo_screenheight()
                dlg.geometry(f"{_sw}x{_sh}+0+0")

        tk.Label(dlg, text="Seleccione los reportes a incluir en el PDF:", font=("Arial", 8, "bold"), bg="#1B3A5C", fg="white").pack(fill="x", ipady=10)

        # --- BARRA INFERIOR PRIMERO ---
        f_bot = tk.Frame(dlg, bg="#2C3E50", height=60)
        f_bot.pack(side="bottom", fill="x")
        f_bot.pack_propagate(False)

        # --- ÁREA PRINCIPAL: 2 columnas izq/der ---
        f_main = tk.Frame(dlg)
        f_main.pack(fill="both", expand=True, padx=15, pady=8)
        f_main.columnconfigure(0, weight=1)
        f_main.columnconfigure(1, weight=1)

        # COLUMNA IZQUIERDA: Globales + Control
        f_izq = ttk.LabelFrame(f_main, text=" Opciones Generales ", padding=10)
        f_izq.grid(row=0, column=0, sticky="nsew", padx=(0,8), pady=4)

        # COLUMNA DERECHA: Documentos (se rellena abajo)
        f_der = ttk.LabelFrame(f_main, text=" Reportes por Documento ", padding=10)
        f_der.grid(row=0, column=1, sticky="nsew", padx=(8,0), pady=4)

        # Alias sf → f_izq para el código existente de checkboxes globales y control
        sf = f_izq

        var_global = tk.BooleanVar(value=True)
        var_general = tk.BooleanVar(value=True)

        ttk.Checkbutton(sf, text="Reporte Técnico Global (Memoria de Cálculo)", variable=var_global).pack(anchor="w", pady=6)
        ttk.Checkbutton(sf, text="Reporte General (Todos los tanques)", variable=var_general).pack(anchor="w", pady=6)

        # --- TIPO DE CONTROL ---
        ttk.Separator(sf, orient="horizontal").pack(fill="x", pady=6)
        f_ctrl = ttk.LabelFrame(sf, text="Tipo de Control (afecta planillas y reportes de cargo)")
        f_ctrl.pack(fill="x", padx=5, pady=4)
        tk.Label(f_ctrl, text="Seleccione según qué se realizó el control:", font=("Arial", 9)).pack(anchor="w", padx=10, pady=2)
        var_ctrl_doc = tk.BooleanVar(value=True)
        var_ctrl_sal = tk.BooleanVar(value=False)
        var_ctrl_lab = tk.BooleanVar(value=False)
        ttk.Checkbutton(f_ctrl, text="Por Documento (planillas y cargos segun documento declarado)", variable=var_ctrl_doc).pack(anchor="w", padx=20)
        ttk.Checkbutton(f_ctrl, text="Por Salida de Zona Primaria (planillas y cargos segun salida)", variable=var_ctrl_sal).pack(anchor="w", padx=20)
        ttk.Checkbutton(f_ctrl, text="Por Analisis de Laboratorio (planillas y cargos segun lab)", variable=var_ctrl_lab).pack(anchor="w", padx=20)

        # ── COLUMNA DERECHA: Documentos ──
        # redirigir a f_der
        sf_doc = f_der
        ttk.Label(sf_doc, text="Reportes por Documento:", font=("Arial", 8, "bold")).pack(anchor="w", pady=(0,4))
        all_tanks = self.lista_tanques + self.lista_carbonera
        mapa = {}
        for tk_name in all_tanks:
            d_ini = self.get_var(f"inicial_{tk_name}_ddt_assign").get()
            d_fin = self.get_var(f"final_{tk_name}_ddt_assign").get()
            if d_ini:
                if d_ini not in mapa: mapa[d_ini] = []
                if tk_name not in mapa[d_ini]: mapa[d_ini].append(tk_name)
            if d_fin and d_fin != d_ini:
                if d_fin not in mapa: mapa[d_fin] = []
                if tk_name not in mapa[d_fin]: mapa[d_fin].append(tk_name)

        doc_vars = {}
        mode_vars = {}
        for ddt_num in mapa:
            ddt_obj = next((d for d in self.ddt_data if d["numero"].get() == ddt_num), None)
            if not ddt_obj: continue
            tipo = ddt_obj["tipo_doc"].get() if "tipo_doc" in ddt_obj else "Detallada"
            prod = ddt_obj["producto"].get()
            n_tanks = len(mapa[ddt_num])

            f_doc = ttk.LabelFrame(sf_doc, text=f"Doc: {ddt_num} ({tipo}) - {prod} - {n_tanks} tanques")
            f_doc.pack(fill="x", pady=4, padx=5)

            var_lab = tk.BooleanVar(value=True)
            var_doc = tk.BooleanVar(value=True)
            var_sal = tk.BooleanVar(value=True)
            ttk.Checkbutton(f_doc, text="Según Laboratorio", variable=var_lab).pack(anchor="w", padx=10)
            ttk.Checkbutton(f_doc, text="Según Documento", variable=var_doc).pack(anchor="w", padx=10)
            ttk.Checkbutton(f_doc, text="Según Salida", variable=var_sal).pack(anchor="w", padx=10)

            doc_vars[ddt_num] = True  # siempre incluido si tiene modos
            mode_vars[ddt_num] = {"lab": var_lab, "doc": var_doc, "sal": var_sal}

        if not mapa:
            tk.Label(sf_doc, text="(No hay documentos asignados a tanques)", fg="gray", font=("Arial", 9, "italic")).pack(anchor="w", padx=10, pady=5)

        # Botones en f_bot (ya creado arriba)
        def seleccionar_todos():
            var_global.set(True); var_general.set(True)
            for mvs in mode_vars.values():
                for v in mvs.values(): v.set(True)
        def deseleccionar_todos():
            var_global.set(False); var_general.set(False)
            for mvs in mode_vars.values():
                for v in mvs.values(): v.set(False)

        tk.Button(f_bot, text="Seleccionar Todo", font=("Arial", 10), bg="#5D6D7E", fg="white", command=seleccionar_todos).pack(side="left", padx=10, pady=10)
        tk.Button(f_bot, text="Deseleccionar Todo", font=("Arial", 9), bg="#5D6D7E", fg="white", command=deseleccionar_todos).pack(side="left", padx=5, pady=10)

        def generar():
            # Detectar REMO que NO se pueden resolver automáticamente
            # (cuando el número no tiene el código de aduana embebido)
            remo_decisions = {}
            for ddt_num in mapa:
                ddt_obj_chk = next((d for d in self.ddt_data if d["numero"].get() == ddt_num), None)
                if ddt_obj_chk:
                    info = self.inferir_tipo_operacion(ddt_obj_chk)
                    if info["necesita_pregunta"]:
                        # Solo llega aquí si no pudo determinarse automáticamente
                        dlg_remo = tk.Toplevel(dlg)
                        dlg_remo.title("REMO - Confirmar direccion")
                        dlg_remo.transient(dlg)
                        dlg_remo.grab_set()
                        dlg_remo.update_idletasks()
                        sw2, sh2 = dlg_remo.winfo_screenwidth(), dlg_remo.winfo_screenheight()
                        dlg_remo.geometry(f"540x260+{(sw2-540)//2}+{(sh2-260)//2}")
                        tk.Label(dlg_remo, text=f"REMO: {ddt_num}", font=("Arial", 8, "bold"), bg="#1B3A5C", fg="white").pack(fill="x", ipady=6)
                        tk.Label(dlg_remo, text="No se pudo determinar la direccion automaticamente.\nIndique si la operacion REMO es a la CARGA o DESCARGA:",
                                 font=("Arial", 10), wraplength=500, justify="center").pack(pady=10)
                        var_remo_d = tk.StringVar(value="remo_descarga")
                        f_remo_d = ttk.Frame(dlg_remo); f_remo_d.pack(pady=5)
                        tk.Radiobutton(f_remo_d, text="A la DESCARGA (aduana origen distinta - Art. 954 C.A.)",
                                       variable=var_remo_d, value="remo_descarga", font=("Arial", 10)).pack(anchor="w", padx=20)
                        tk.Radiobutton(f_remo_d, text="A la CARGA (misma aduana - Art. 959 C.A.)",
                                       variable=var_remo_d, value="remo_carga", font=("Arial", 10)).pack(anchor="w", padx=20)
                        def _ok_remo_d(key=ddt_num):
                            remo_decisions[key] = var_remo_d.get()
                            dlg_remo.destroy()
                        tk.Button(dlg_remo, text="Confirmar", bg="#2E7D32", fg="white",
                                  font=("Arial", 8, "bold"), command=_ok_remo_d).pack(pady=8)
                        dlg_remo.wait_window()
                        if ddt_num not in remo_decisions:
                            remo_decisions[ddt_num] = "remo_descarga"
            dlg.destroy()
            self._generar_reportes_seleccionados(var_global.get(), var_general.get(), mapa, mode_vars,
                                                  ctrl_doc=var_ctrl_doc.get(), ctrl_sal=var_ctrl_sal.get(),
                                                  ctrl_lab=var_ctrl_lab.get(), remo_decisions=remo_decisions)

        def _generar_con_tributos():
            generar()   # first calls generar() which destroys dlg

        tk.Button(f_bot, text="  GENERAR PDF", bg="#27AE60", fg="white", font=("Arial", 8, "bold"),
                  command=generar, cursor="hand2").pack(side="right", padx=20, pady=8, ipadx=14, ipady=5)
        tk.Button(f_bot, text="⚖ TRIBUTOS", bg="#7B1FA2", fg="white", font=("Arial", 8, "bold"),
                  command=lambda: self.dialogo_tributos(dlg)).pack(side="right", padx=6, pady=8)
        tk.Button(f_bot, text="Cancelar", font=("Arial", 9), bg="#E74C3C", fg="white",
                  command=dlg.destroy).pack(side="right", padx=5, pady=10)

    def _generar_reportes_seleccionados(self, inc_global, inc_general, mapa, mode_vars, ctrl_doc=True, ctrl_sal=False, ctrl_lab=False, remo_decisions=None):
        """Genera solo los reportes seleccionados."""
        target_dir = filedialog.askdirectory(title="Seleccione carpeta destino")
        if not target_dir: return
        errors = []
        all_tanks = self.lista_tanques + self.lista_carbonera

        clean_buque = self.clean_filename(self.get_var('car_buque').get())
        if not clean_buque: clean_buque = "Reporte"
        _fecha_str = datetime.now().strftime("%Y%m%d")
        unified_path = os.path.join(target_dir, f"Reporte_Completo_{clean_buque}_{_fecha_str}.pdf")

        # === PASO 1: PRE-CALCULAR CARGOS/DENUNCIAS ANTES DE ABRIR EL CANVAS ===
        # Determinar modo_comp desde ctrl flags
        if ctrl_lab:
            modo_forzado = "laboratorio"
        elif ctrl_sal:
            modo_forzado = "salida"
        else:
            modo_forzado = "documento"

        # Pre-calcular diferencias para todos los documentos
        cargos_pendientes = []
        for ddt_num, tanks in mapa.items():
            ddt_obj = next((d for d in self.ddt_data if d["numero"].get() == ddt_num), None)
            if not ddt_obj: continue
            try:
                # Kilos vacío: para líquidos se calculan acá igual que la planilla
                # (NETO = bruto − agua, × VCF × densidad según modo) para que el
                # cargo/denuncia coincida exactamente con el PDF. Para otros tipos
                # se usan los kv_* pre-calculados de la UI.
                _tm_cp = self.get_tipo_medio()
                _es_liq_cp = _tm_cp in ("BUQUE","BARCAZA","BUQUE QUIMIQUERO","DRAFT SURVEY",
                                        "TANQUE FIJO","TANQUE FLOTANTE","CAMION CISTERNA")
                dens_key_cp = "dens_lab" if modo_forzado == "laboratorio" else ("dens_salida" if modo_forzado == "salida" else "dens_doc")
                kv_suffix = "kv_lab" if modo_forzado == "laboratorio" else ("kv_sal" if modo_forzado == "salida" else "kv_doc")
                sum_kv_i = 0; sum_kv_f = 0
                for tk_name in tanks:
                    for etapa in ["inicial", "final"]:
                        if _es_liq_cp:
                            d_raw = self.get_var(f"{etapa}_{tk_name}_{dens_key_cp}").get()
                            if not d_raw: d_raw = self.get_var(f"{etapa}_{tk_name}_dens_lab").get()
                            d_v = self.parse_float(d_raw)
                            tbl_v = self.get_var(f"{etapa}_{tk_name}_tabla_vcf").get() or "54B (Combustibles)"
                            t_v = self.parse_float(self.get_var(f"{etapa}_{tk_name}_temp").get())
                            neto_v = self.interpolar_prod(etapa, tk_name) - self.interpolar_agua(etapa, tk_name)
                            vcf_v = self.calc_vcf(d_v, t_v, tbl_v) if d_v > 0 else 1.0
                            _dg = d_v / 1000.0 if d_v > 2.0 else d_v
                            val_kv = neto_v * vcf_v * _dg if _dg > 0 else 0.0
                        else:
                            raw_kv = self.get_var(f"{etapa}_{tk_name}_{kv_suffix}").get()
                            try: val_kv = float(str(raw_kv).replace(",","")) if raw_kv and str(raw_kv).replace(".","").replace("-","").replace(",","").isdigit() else 0.0
                            except: val_kv = 0.0
                        if etapa == "inicial": sum_kv_i += val_kv
                        else:                  sum_kv_f += val_kv
                dif_kv = sum_kv_i - sum_kv_f
                doc_k = self.parse_float(ddt_obj["kilos"].get())
                sal_k = sum(self.parse_float(s["kilos"].get()) for s in ddt_obj["salidas"]) if ddt_obj.get("salidas") else 0
                # LAB: kv_lab vs salidas | SAL/DOC: kv vs documento
                target_k = sal_k if (modo_forzado == "laboratorio" and sal_k > 0) else (doc_k if doc_k > 0 else 1)
                diff_k = dif_kv - target_k
                permil_k_val = (diff_k / target_k * 1000) if target_k != 0 else 0
                if abs(permil_k_val) > 6.0:
                    cargos_pendientes.append({
                        "ddt_obj": ddt_obj, "dif_kv": dif_kv,
                        "target_k": target_k, "modo": modo_forzado
                    })
            except Exception as _precalc_err:
                traceback.print_exc()
                messagebox.showerror("Error pre-calculo cargo", f"Error calculando diferencias para {ddt_num}:\n{_precalc_err}")

        # Resolver tipo de operación para cargos (REMO automático o remo_decisions)
        if remo_decisions is None: remo_decisions = {}
        for cp in cargos_pendientes:
            info_op = self.inferir_tipo_operacion(cp["ddt_obj"])
            if info_op["necesita_pregunta"]:
                num_doc = cp["ddt_obj"]["numero"].get()
                remo_tipo = remo_decisions.get(num_doc, "remo_descarga")
                info_op["tipo"] = remo_tipo
                if remo_tipo == "remo_carga":
                    info_op["art_principal"] = "Art. 959 del Código Aduanero"
                    info_op["art_inc"] = "Art. 959 inc. c) C.A."
                    info_op["descripcion"] = "TRANSFERENCIA DE COMBUSTIBLE A LA CARGA (REMO)"
                else:
                    info_op["descripcion"] = "TRANSFERENCIA DE COMBUSTIBLE A LA DESCARGA (REMO)"
            cp["tipo_operacion_info"] = info_op

        # === PASO 2: SI HAY CARGOS, PREGUNTAR FORMATO ===
        modo_cargo_export = "pdf"   # default: embebido en PDF
        if cargos_pendientes:
            _dlg_fmt = tk.Toplevel(self.root)
            _dlg_fmt.title("Formato de Informes de Cargo/Denuncia")
            _dlg_fmt.transient(self.root)
            _dlg_fmt.grab_set()
            _dlg_fmt.update_idletasks()
            _sw, _sh = _dlg_fmt.winfo_screenwidth(), _dlg_fmt.winfo_screenheight()
            _dlg_fmt.geometry(f"500x220+{(_sw-500)//2}+{(_sh-220)//2}")
            _n_inf = len(cargos_pendientes)
            _tipos = sum(1 for cp in cargos_pendientes if cp["target_k"] and abs((cp["dif_kv"]-cp["target_k"])/cp["target_k"]*100) >= 2.0)
            _tipos_str = []
            for cp in cargos_pendientes:
                _dk = cp["dif_kv"] - cp["target_k"]
                _pct = abs(_dk/cp["target_k"]*100) if cp["target_k"] else 0
                _tipos_str.append("Denuncia" if _pct >= 2.0 else "Cargo")
            _resumen = ", ".join(f"{cp['ddt_obj']['numero'].get()} ({t})" for cp, t in zip(cargos_pendientes, _tipos_str))
            tk.Label(_dlg_fmt, text=f"Se generarán {_n_inf} informe(s):", font=("Arial", 8, "bold"), bg="#1B3A5C", fg="white").pack(fill="x", ipady=6)
            tk.Label(_dlg_fmt, text=_resumen, font=("Arial", 9), wraplength=480, justify="center", fg="#333").pack(pady=6)
            tk.Label(_dlg_fmt, text="¿Cómo desea los informes de cargo/denuncia?", font=("Arial", 10)).pack(pady=(4,2))
            _var_fmt = tk.StringVar(value="pdf")
            _f_opts = ttk.Frame(_dlg_fmt); _f_opts.pack(pady=4)
            tk.Radiobutton(_f_opts, text="[PDF]  Embebidos al inicio del PDF unificado",
                           variable=_var_fmt, value="pdf", font=("Arial", 10)).pack(anchor="w", padx=30)
            tk.Radiobutton(_f_opts, text="[WORD] Documentos Word/LibreOffice (.docx) separados",
                           variable=_var_fmt, value="docx", font=("Arial", 10)).pack(anchor="w", padx=30)
            _fmt_ok = [False]
            def _fmt_aceptar():
                _fmt_ok[0] = True; _dlg_fmt.destroy()
            tk.Button(_dlg_fmt, text="Continuar →", bg="#27AE60", fg="white", font=("Arial", 8, "bold"),
                      command=_fmt_aceptar).pack(pady=8)
            _dlg_fmt.wait_window()
            if not _fmt_ok[0]: return  # usuario cerró el diálogo
            modo_cargo_export = _var_fmt.get()
            # ── Paso 2b: Configurar tributos ────────────────────────────
            if not self.dialogo_tributos():
                return  # usuario canceló

        # === PASO 3: CREAR CANVAS ===
        try:
            c = canvas.Canvas(unified_path, pagesize=A4)  # portrait para cargos primero
        except Exception as e:
            traceback.print_exc()
            messagebox.showerror("Error", f"No se pudo crear el PDF:\n{e}")
            return

        report_count = 0
        cargo_count = 0

        # === PASO 4a: CARGOS EN PDF (si modo pdf, al principio) ===
        if modo_cargo_export == "pdf":
            for cp in cargos_pendientes:
                info_op = cp.get("tipo_operacion_info", {})
                try:
                    self._generar_cargo_en_canvas(
                        c,
                        cp["ddt_obj"], cp["dif_kv"],
                        cp["target_k"],
                        ctrl_doc=ctrl_doc, ctrl_sal=ctrl_sal, ctrl_lab=ctrl_lab,
                        modo_comp_forzado=cp["modo"],
                        tipo_operacion=info_op.get("tipo","importacion"),
                        tipo_operacion_info=info_op
                    )
                    cargo_count += 1
                except Exception as _cargo_err:
                    traceback.print_exc()
                    messagebox.showerror("Error en Cargo/Denuncia",
                        f"Error generando cargo para {cp['ddt_obj']['numero'].get()}:\n{_cargo_err}")

        # === PASO 4b: CAMBIAR A LANDSCAPE Y GENERAR REPORTES ===
        # Cambiar tamaño de página para los reportes (landscape A4)
        c.setPageSize(landscape(A4))

        modes_map = [
            ("lab", "SEGÚN_LABORATORIO", "dens_lab"),
            ("doc", "SEGÚN_DOCUMENTO", "dens_doc"),
            ("sal", "SEGÚN_SALIDA", "dens_salida")
        ]

        if inc_global:
            try:
                self.generar_reporte_tecnico_global("DETALLE_TECNICO_GLOBAL", target_dir, shared_canvas=c)
                report_count += 1
            except Exception as e:
                traceback.print_exc()
                errors.append(f"Reporte Técnico Global: {str(e)}")

        if inc_general:
            try:
                self.generar_un_reporte("GENERAL", all_tanks, is_partial=False, output_folder=target_dir, shared_canvas=c)
                report_count += 1
            except Exception as e:
                traceback.print_exc()
                errors.append(f"Reporte General: {str(e)}")

        for ddt_num, tanks in mapa.items():
            ddt_obj = next((d for d in self.ddt_data if d["numero"].get() == ddt_num), None)
            if not ddt_obj: continue
            safe_doc_num = self.clean_filename(ddt_num)
            if not safe_doc_num: safe_doc_num = "SinNombre"
            mvars = mode_vars.get(ddt_num, {})
            for mode_key, suffix, density_key in modes_map:
                if mode_key in mvars and mvars[mode_key].get():
                    try:
                        self.generar_un_reporte(f"DOC_{safe_doc_num}_{suffix}", tanks, is_partial=True, ddt_obj=ddt_obj, output_folder=target_dir, density_mode_key=density_key, shared_canvas=c)
                        report_count += 1
                    except Exception as e:
                        traceback.print_exc()
                        errors.append(f"Reporte {ddt_num} ({suffix}): {str(e)}")



        # === PASO 5: GUARDAR PDF ===
        c.save()

        if errors:
            msg = "Algunos reportes fallaron:\n" + "\n".join(errors[:5])
            if len(errors) > 5: msg += "\n..."
            messagebox.showerror("Errores de Generación", msg)

        # === PASO 6: SI DOCX, GENERAR WORD PARA CADA CARGO ===
        docx_paths = []
        if modo_cargo_export == "docx" and cargos_pendientes:
            for cp in cargos_pendientes:
                info_op = cp.get("tipo_operacion_info", {})
                try:
                    docx_path = self._generar_cargo_docx(
                        cp["ddt_obj"], cp["dif_kv"], cp["target_k"],
                        ctrl_doc=ctrl_doc, ctrl_sal=ctrl_sal, ctrl_lab=ctrl_lab,
                        modo_comp_forzado=cp["modo"],
                        tipo_operacion=info_op.get("tipo", "importacion"),
                        tipo_operacion_info=info_op,
                        output_folder=target_dir
                    )
                    if docx_path:
                        docx_paths.append(docx_path)
                        cargo_count += 1
                except Exception as _docx_err:
                    traceback.print_exc()
                    messagebox.showerror("Error Word", f"Error generando .docx para {cp['ddt_obj']['numero'].get()}:\n{_docx_err}")

        total = report_count + cargo_count
        if total == 0:
            messagebox.showwarning("Atención", "No se generaron reportes.")
        else:
            if modo_cargo_export == "docx" and docx_paths:
                extra = f" + {len(docx_paths)} cargo(s)/denuncia(s) en Word"
            elif cargo_count > 0:
                extra = f" ({cargo_count} cargo(s)/denuncia(s) al inicio del PDF)"
            else:
                extra = " (sin cargos/denuncias)"
            messagebox.showinfo("Listo", f"PDF: {report_count} reporte(s){extra}\n{os.path.basename(unified_path)}")
            try:
                if platform.system() == 'Windows': os.startfile(unified_path)
                elif platform.system() == 'Darwin': subprocess.call(('open', unified_path))
                else: subprocess.call(('xdg-open', unified_path))
            except: pass
            # Abrir los .docx también
            for dp in docx_paths:
                try:
                    if platform.system() == 'Windows': os.startfile(dp)
                    elif platform.system() == 'Darwin': subprocess.call(('open', dp))
                    else: subprocess.call(('xdg-open', dp))
                except: pass

   

    def abrir_tabla_calibrado(self, etapa, tk_name, label_col2="LITROS", es_buque=False):
        """Gestiona la tabla de calibrado (sondaje -> litros) con CSV import/export, ordenar.
        Para buques (es_buque=True) agrega columna de corrección de trim (L/m de asiento)."""
        import json, csv as csv_mod
        top = tk.Toplevel(self.root)
        top.title(f"Tabla de Calibrado — {tk_name} ({etapa.upper()})")
        top.geometry("980x640" if es_buque else "720x620")
        top.resizable(True, True)
        var_key = f"{etapa}_{tk_name}_tabla_cal_json"

        fh = tk.Frame(top, bg="#4A235A", height=50)
        fh.pack(fill="x"); fh.pack_propagate(False)
        tk.Label(fh, text=f"TABLA DE CALIBRADO  |  {tk_name}  |  {etapa.upper()}",
                 bg="#4A235A", fg="white", font=("Arial",10,"bold")).pack(side="left", padx=16, pady=12)
        col2_info = f"Col.2: {label_col2}" + ("  |  Col.3: CORR.TRIM (L/m asiento)" if es_buque else "")
        tk.Label(fh, text=col2_info, bg="#4A235A", fg="#CE93D8", font=("Arial",8)).pack(side="right", padx=12)

        if es_buque:
            f_info = tk.Frame(top, bg="#EBF5FB", bd=1, relief="flat")
            f_info.pack(fill="x")
            tk.Label(f_info,
                     text="BUQUE / BARCAZA — Col.3 CORR.TRIM: litros a sumar por cada metro de asiento (trim) positivo "
                          "(popa más hundida que proa). Negativo si el tanque está a proa del LCF. "
                          "Dejar en 0 si no aplica corrección de trim.",
                     bg="#EBF5FB", fg="#1A5276", font=("Arial", 7), anchor="w",
                     wraplength=940, justify="left").pack(fill="x", padx=10, pady=4)

        f_ctrl = tk.Frame(top, bg="#EDE7F6", pady=6)
        f_ctrl.pack(fill="x")
        tk.Button(f_ctrl, text="Cargar CSV", bg="#8E44AD", fg="white", font=("Arial",8,"bold"),
                  command=lambda: _cargar_csv()).pack(side="left", padx=10, pady=4, ipadx=8, ipady=3)
        tk.Button(f_ctrl, text="Exportar CSV", bg="#2980B9", fg="white", font=("Arial",8),
                  command=lambda: _exportar_csv()).pack(side="left", padx=4, pady=4, ipadx=6, ipady=3)
        tk.Button(f_ctrl, text="+ Agregar fila", bg="#16A085", fg="white", font=("Arial",8),
                  command=lambda: _add_row("","","")).pack(side="left", padx=4, pady=4, ipadx=6, ipady=3)
        tk.Button(f_ctrl, text="Ordenar", bg="#7F8C8D", fg="white", font=("Arial",8),
                  command=lambda: _ordenar()).pack(side="left", padx=4, pady=4, ipadx=6, ipady=3)
        tk.Button(f_ctrl, text="Limpiar todo", bg="#C0392B", fg="white", font=("Arial",8),
                  command=lambda: _limpiar()).pack(side="left", padx=4, pady=4, ipadx=6, ipady=3)
        lbl_pts = tk.Label(f_ctrl, text="0 puntos", bg="#EDE7F6", font=("Arial",8), fg="#6A1B9A")
        lbl_pts.pack(side="right", padx=12)
        csv_hint = "CSV: sondaje_mm, litros, corr_trim" if es_buque else "CSV: sondaje_mm, litros"
        tk.Label(f_ctrl, text=csv_hint, bg="#EDE7F6", font=("Arial",7), fg="#6A1B9A").pack(side="right", padx=4)

        f_tbl = tk.Frame(top, bg="white")
        f_tbl.pack(fill="both", expand=True, padx=10, pady=6)
        canvas_t = tk.Canvas(f_tbl, bg="white", highlightthickness=0)
        sb_t = ttk.Scrollbar(f_tbl, orient="vertical", command=canvas_t.yview)
        scroll_inner = tk.Frame(canvas_t, bg="white")
        scroll_inner.bind("<Configure>", lambda e: canvas_t.configure(scrollregion=canvas_t.bbox("all")))
        win_id = canvas_t.create_window((0,0), window=scroll_inner, anchor="nw")
        canvas_t.bind("<Configure>", lambda e: canvas_t.itemconfig(win_id, width=e.width))
        canvas_t.configure(yscrollcommand=sb_t.set)
        sb_t.pack(side="right", fill="y"); canvas_t.pack(fill="both", expand=True)
        canvas_t.bind("<MouseWheel>", lambda e: canvas_t.yview_scroll(int(-1*(e.delta/120)) if e.delta else (1 if e.num==5 else -1), "units"))
        canvas_t.bind("<Button-4>", lambda e: canvas_t.yview_scroll(-1,"units"))
        canvas_t.bind("<Button-5>", lambda e: canvas_t.yview_scroll(1,"units"))

        hdr = tk.Frame(scroll_inner, bg="#4A235A")
        hdr.pack(fill="x")
        tk.Label(hdr, text="#", width=5, bg="#4A235A", fg="white", font=("Arial",8,"bold")).pack(side="left", padx=2, pady=3)
        tk.Label(hdr, text="SONDAJE (mm)", width=18, bg="#4A235A", fg="white", font=("Arial",8,"bold")).pack(side="left", padx=2)
        tk.Label(hdr, text=label_col2, width=18, bg="#4A235A", fg="white", font=("Arial",8,"bold")).pack(side="left", padx=2)
        if es_buque:
            tk.Label(hdr, text="CORR.TRIM (L/m asiento)", width=22, bg="#1A5276",
                     fg="white", font=("Arial",8,"bold")).pack(side="left", padx=2)
        tk.Label(hdr, text="Quitar", width=6, bg="#4A235A", fg="white", font=("Arial",8,"bold")).pack(side="left", padx=2)

        rows_data = []

        def _update_count():
            n = len([r for r in rows_data if r["sond"].get().strip() and r["lits"].get().strip()])
            lbl_pts.config(text=f"{n} puntos")

        def _render_rows():
            for w in scroll_inner.winfo_children():
                if hasattr(w, "_cal_row"): w.destroy()
            for i, rd in enumerate(rows_data):
                rf = tk.Frame(scroll_inner, bg="#F3E5F5" if i%2==0 else "white")
                rf._cal_row = True
                rf.pack(fill="x")
                tk.Label(rf, text=str(i+1), width=5, bg=rf["bg"], font=("Arial",7)).pack(side="left", padx=2, pady=2)
                e1 = tk.Entry(rf, textvariable=rd["sond"], width=18, justify="center", font=("Arial",9), bg="#FDFEFE")
                e1.pack(side="left", padx=2)
                e2 = tk.Entry(rf, textvariable=rd["lits"], width=18, justify="center", font=("Arial",9), bg="#FDFEFE")
                e2.pack(side="left", padx=2)
                e1.bind("<KeyRelease>", lambda ev: _update_count())
                e2.bind("<KeyRelease>", lambda ev: _update_count())
                if es_buque:
                    e3 = tk.Entry(rf, textvariable=rd["trim"], width=22, justify="center",
                                  font=("Arial",9), bg="#EBF5FB")
                    e3.pack(side="left", padx=2)
                    e3.bind("<KeyRelease>", lambda ev: _update_count())
                tk.Button(rf, text="X", bg="#E74C3C", fg="white", width=4, font=("Arial",7),
                          command=lambda idx=i: _del_row(idx)).pack(side="left", padx=2)
            _update_count()

        def _add_row(s_val="", l_val="", t_val="0"):
            rows_data.append({"sond": tk.StringVar(value=str(s_val)),
                               "lits": tk.StringVar(value=str(l_val)),
                               "trim": tk.StringVar(value=str(t_val))})
            _render_rows()

        def _del_row(idx):
            if 0 <= idx < len(rows_data): del rows_data[idx]
            _render_rows()

        def _limpiar():
            rows_data.clear(); _render_rows()

        def _ordenar():
            pts = []
            for rd in rows_data:
                s = rd["sond"].get().strip().replace(",",".")
                l = rd["lits"].get().strip().replace(",",".")
                t = rd["trim"].get().strip().replace(",",".") if es_buque else "0"
                try:
                    pts.append((float(s), float(l), float(t) if t else 0.0))
                except: pass
            pts.sort(key=lambda x: x[0])
            rows_data.clear()
            for s, l, t in pts:
                rows_data.append({"sond": tk.StringVar(value=str(s)),
                                   "lits": tk.StringVar(value=str(l)),
                                   "trim": tk.StringVar(value=str(t))})
            _render_rows()

        def _parse_line(line):
            line = line.strip()
            if not line or line.startswith("#"): return None
            for sep in [",",";","\t","  "," "]:
                parts = line.split(sep)
                if len(parts) >= 2:
                    try:
                        s_v = parts[0].strip().replace(",",".")
                        l_v = parts[1].strip().replace(",",".")
                        t_v = parts[2].strip().replace(",",".") if len(parts) >= 3 else "0"
                        float(s_v); float(l_v); float(t_v)
                        return s_v, l_v, t_v
                    except: pass
            return None

        def _cargar_csv():
            f_path = filedialog.askopenfilename(
                title="Cargar tabla de calibrado",
                filetypes=[("CSV/TXT","*.csv *.txt"),("Todos","*.*")], parent=top)
            if not f_path: return
            try:
                nuevos = []
                with open(f_path, "r", encoding="utf-8-sig") as fcsv:
                    for line in fcsv:
                        parsed = _parse_line(line)
                        if parsed: nuevos.append(parsed)
                if not nuevos:
                    messagebox.showwarning("Sin datos", "No se encontraron pares numericos en el archivo.", parent=top); return
                if rows_data and any(r["sond"].get().strip() for r in rows_data):
                    if messagebox.askyesno("Reemplazar o agregar", "Reemplazar la tabla actual?\n(No = agregar a la existente)", parent=top):
                        rows_data.clear()
                for s_v, l_v, t_v in nuevos:
                    rows_data.append({"sond": tk.StringVar(value=s_v),
                                       "lits": tk.StringVar(value=l_v),
                                       "trim": tk.StringVar(value=t_v)})
                _render_rows()
                messagebox.showinfo("OK", f"Cargados {len(nuevos)} puntos.", parent=top)
            except Exception as ex:
                messagebox.showerror("Error", f"No se pudo leer el archivo:\n{ex}", parent=top)


        def _exportar_csv():
            if es_buque:
                pts = [(r["sond"].get().strip(), r["lits"].get().strip(), r["trim"].get().strip() or "0")
                       for r in rows_data if r["sond"].get().strip() and r["lits"].get().strip()]
            else:
                pts = [(r["sond"].get().strip(), r["lits"].get().strip())
                       for r in rows_data if r["sond"].get().strip() and r["lits"].get().strip()]
            if not pts:
                messagebox.showwarning("Sin datos", "No hay puntos para exportar.", parent=top); return
            f_path = filedialog.asksaveasfilename(
                title="Exportar tabla de calibrado", defaultextension=".csv",
                filetypes=[("CSV","*.csv"),("Texto","*.txt")],
                initialfile=f"calibrado_{tk_name}_{etapa}.csv", parent=top)
            if not f_path: return
            try:
                with open(f_path, "w", newline="", encoding="utf-8") as fc:
                    writer = csv_mod.writer(fc)
                    if es_buque:
                        writer.writerow(["sondaje_mm", label_col2.lower().replace(" ","_"), "corr_trim_L_m"])
                        for row in pts: writer.writerow(row)
                    else:
                        writer.writerow(["sondaje_mm", label_col2.lower().replace(" ","_")])
                        for s, l in pts: writer.writerow([s, l])
                messagebox.showinfo("Exportado", f"Exportados {len(pts)} puntos.\n{f_path}", parent=top)
            except Exception as ex:
                messagebox.showerror("Error", f"No se pudo exportar:\n{ex}", parent=top)

        existing = self.get_var(var_key).get()
        if existing:
            try:
                for pt in json.loads(existing):
                    _add_row(pt[0], pt[1], pt[2] if len(pt) > 2 else "0")
            except: pass
        if not rows_data:
            for _ in range(10): _add_row("", "", "0")
        _render_rows()

        f_foot = tk.Frame(top, bg="#2C3E50", height=50)
        f_foot.pack(fill="x", side="bottom"); f_foot.pack_propagate(False)

        def _guardar():
            pts = []
            for rd in rows_data:
                s = rd["sond"].get().strip().replace(",",".")
                l = rd["lits"].get().strip().replace(",",".")
                t = rd["trim"].get().strip().replace(",",".") if es_buque else "0"
                if not s or not l: continue
                try:
                    entry = [float(s), float(l)]
                    if es_buque: entry.append(float(t) if t else 0.0)
                    pts.append(entry)
                except: pass
            if not pts:
                self.get_var(var_key).set(""); top.destroy(); return
            pts.sort(key=lambda x: x[0])
            self.get_var(var_key).set(json.dumps(pts))
            top.destroy()
            messagebox.showinfo("Guardado", f"Tabla guardada: {len(pts)} puntos.\nRango: {pts[0][0]:.0f} mm -> {pts[-1][0]:.0f} mm")

        tk.Button(f_foot, text="  [OK] GUARDAR TABLA  ", bg="#27AE60", fg="white",
                  font=("Arial",9,"bold"), relief="flat", cursor="hand2",
                  command=_guardar).pack(side="left", padx=16, pady=8, ipadx=12, ipady=4)
        tk.Button(f_foot, text="Cancelar", bg="#7F8C8D", fg="white", font=("Arial",8), relief="flat",
                  command=top.destroy).pack(side="right", padx=8, pady=8, ipadx=6, ipady=3)
        tk.Label(f_foot, text="Ordenar antes de guardar para mejor interpolacion",
                 bg="#2C3E50", fg="#AED6F1", font=("Arial",7)).pack(side="right", padx=12)

    # ══════════════════════════════════════════════════════════════════════
    #  Interpolación bilineal por sondaje × asiento (trim) — buques/barcazas
    # ══════════════════════════════════════════════════════════════════════
    def _interp_1d(self, pts, x):
        """Interpola linealmente en una lista [(x0,y0),(x1,y1),...] (clamp en los extremos).
        Devuelve None si no hay puntos usables."""
        pts = [(float(a), float(b)) for a, b in pts if a is not None and b is not None]
        if not pts: return None
        pts.sort(key=lambda p: p[0])
        if len(pts) == 1: return pts[0][1]
        if x <= pts[0][0]:  return pts[0][1]
        if x >= pts[-1][0]: return pts[-1][1]
        for k in range(len(pts) - 1):
            x0, y0 = pts[k]; x1, y1 = pts[k+1]
            if x0 <= x <= x1:
                if x1 == x0: return y0
                return y0 + (x - x0) / (x1 - x0) * (y1 - y0)
        return pts[-1][1]

    def _interp_trim_table(self, obj, s, trim):
        """Interpolación bilineal en una tabla multi-asiento.
        obj = {"trims":[t0,t1,...], "rows":[[sondaje, v_t0, v_t1, ...], ...]}
        1) Para cada columna de asiento interpola volumen vs sondaje en 's'.
        2) Interpola esos volúmenes entre columnas según el 'trim' real (popa-proa).
        Devuelve (volumen, detalle_str) o (None, motivo)."""
        try:
            trims = [float(t) for t in obj.get("trims", [])]
            rows = [r for r in obj.get("rows", []) if r and r[0] is not None]
            if not trims or len(rows) < 1:
                return None, "tabla trim vacía"
            col_vols = []
            for c in range(len(trims)):
                pts = [(r[0], r[c+1]) for r in rows if len(r) > c+1 and r[c+1] is not None]
                v = self._interp_1d(pts, s)
                if v is not None:
                    col_vols.append((trims[c], v))
            if not col_vols:
                return None, "sin volúmenes por columna"
            vol = self._interp_1d(col_vols, trim)
            det = (f"Sondaje {s:.0f}mm, asiento {trim:+.2f}m → "
                   + " | ".join(f"t{tc:g}:{vc:.0f}" for tc, vc in col_vols)
                   + f" ⇒ {vol:.0f}")
            return vol, det
        except Exception as ex:
            return None, f"error: {ex}"

    def abrir_interp_trim_rapida(self, etapa, tk_name, agua=False):
        """Carga manual rápida para interpolar por asiento (trim) en buques.
        El usuario lee de la tabla de calibrado impresa los dos sondajes que
        rodean su medición y los litros en las dos páginas de asiento que
        rodean el asiento real (p.ej. 0 y 0.5). Interpola en 2D y guarda en
        {etapa}_{tanque}_tabla_trim[_agua]_json (mismo formato que el editor
        completo). Con agua=True opera sobre el sondaje de agua de fondo."""
        import json
        _que = "AGUA" if agua else "PRODUCTO"
        _hbg = "#117864" if agua else "#1A5276"
        top = tk.Toplevel(self.root)
        top.title(f"Interpolación × Asiento ({_que}) — {tk_name} ({etapa.upper()})")
        top.geometry("640x470")
        top.resizable(False, False)
        var_key = f"{etapa}_{tk_name}_tabla_trim{'_agua' if agua else ''}_json"

        _s_key = f"{etapa}_{tk_name}_agua_s_real" if agua else f"{etapa}_{tk_name}_s_corr"
        s_act = self.parse_float(self.get_var(_s_key).get())
        trim_act = self.parse_float(self.get_var(f"{etapa}_Trimación").get() or "0")

        fh = tk.Frame(top, bg=_hbg, height=46)
        fh.pack(fill="x"); fh.pack_propagate(False)
        tk.Label(fh, text=f"INTERPOLACIÓN × ASIENTO ({_que})  |  {tk_name}  |  {etapa.upper()}",
                 bg=_hbg, fg="white", font=("Arial",10,"bold")).pack(side="left", padx=14, pady=10)
        tk.Label(fh, text=f"Sondaje{' agua' if agua else ''}: {s_act:.0f} mm   Asiento: {trim_act:+.2f} m",
                 bg=_hbg, fg="#AED6F1", font=("Arial",9,"bold")).pack(side="right", padx=12)

        tk.Label(top, text="Cargá el sondaje de la tabla y sus litros en los 2 asientos que rodean el asiento real "
                           "del buque. La 2da fila es opcional: usala solo si tu sondaje medido cae entre dos "
                           "filas de la tabla (interpola también por sondaje).",
                 bg="#EBF5FB", fg="#1A5276", font=("Arial",8), anchor="w", wraplength=610,
                 justify="left").pack(fill="x", padx=0, pady=(0,4), ipadx=10, ipady=4)

        # Prefill desde tabla guardada (si es chica) o valores por defecto
        tA, tB = 0.0, 0.5
        pre = [["",""],["",""]]   # [fila][col] litros
        pre_s = ["",""]
        existing = self.get_var(var_key).get()
        if existing:
            try:
                obj0 = json.loads(existing)
                tr0 = [float(t) for t in obj0.get("trims", [])]
                rw0 = obj0.get("rows", [])
                if len(tr0) >= 2: tA, tB = tr0[0], tr0[1]
                for i in range(min(2, len(rw0))):
                    pre_s[i] = f"{rw0[i][0]:g}"
                    if len(rw0[i]) > 1 and rw0[i][1] is not None: pre[i][0] = f"{rw0[i][1]:g}"
                    if len(rw0[i]) > 2 and rw0[i][2] is not None: pre[i][1] = f"{rw0[i][2]:g}"
            except: pass

        f_g = tk.Frame(top, bg="white", padx=14, pady=8); f_g.pack(fill="both", expand=True)
        v_tA = tk.StringVar(value=f"{tA:g}"); v_tB = tk.StringVar(value=f"{tB:g}")
        v_s  = [tk.StringVar(value=pre_s[0]), tk.StringVar(value=pre_s[1])]
        v_l  = [[tk.StringVar(value=pre[0][0]), tk.StringVar(value=pre[0][1])],
                [tk.StringVar(value=pre[1][0]), tk.StringVar(value=pre[1][1])]]

        fbh = ("Arial", 9, "bold")
        tk.Label(f_g, text="", bg="white").grid(row=0, column=0)
        tk.Label(f_g, text="ASIENTO A (m)", bg="#D6EAF8", font=fbh, width=16).grid(row=0, column=1, padx=3, pady=3, sticky="ew")
        tk.Label(f_g, text="ASIENTO B (m)", bg="#D6EAF8", font=fbh, width=16).grid(row=0, column=2, padx=3, pady=3, sticky="ew")
        tk.Label(f_g, text="Asiento →", bg="white", font=fbh, anchor="e").grid(row=1, column=0, sticky="e", padx=3)
        e_tA = tk.Entry(f_g, textvariable=v_tA, justify="center", font=("Arial",11,"bold"), bg="#EBF5FB", width=14)
        e_tA.grid(row=1, column=1, padx=3, pady=2, ipady=3)
        e_tB = tk.Entry(f_g, textvariable=v_tB, justify="center", font=("Arial",11,"bold"), bg="#EBF5FB", width=14)
        e_tB.grid(row=1, column=2, padx=3, pady=2, ipady=3)

        tk.Label(f_g, text="SONDAJE (mm)", bg="#D5DBDB", font=fbh).grid(row=2, column=0, padx=3, pady=(10,3), sticky="ew")
        _lit = "LTS.AGUA" if agua else "LITROS"
        tk.Label(f_g, text=f"{_lit} @ A", bg="#D6EAF8", font=fbh).grid(row=2, column=1, padx=3, pady=(10,3), sticky="ew")
        tk.Label(f_g, text=f"{_lit} @ B", bg="#D6EAF8", font=fbh).grid(row=2, column=2, padx=3, pady=(10,3), sticky="ew")

        entries = []
        for i in range(2):
            e_s = tk.Entry(f_g, textvariable=v_s[i], justify="center", font=("Arial",12), bg="#FDFEFE", width=14)
            e_s.grid(row=3+i, column=0, padx=3, pady=3, ipady=4)
            e_a = tk.Entry(f_g, textvariable=v_l[i][0], justify="center", font=("Arial",12), bg="#F4FAFF", width=14)
            e_a.grid(row=3+i, column=1, padx=3, pady=3, ipady=4)
            e_b = tk.Entry(f_g, textvariable=v_l[i][1], justify="center", font=("Arial",12), bg="#F4FAFF", width=14)
            e_b.grid(row=3+i, column=2, padx=3, pady=3, ipady=4)
            entries += [e_s, e_a, e_b]

        lbl_res = tk.Label(f_g, text="—", bg="#FEF9E7", fg="#7D6608",
                           font=("Arial",11,"bold"), relief="groove", pady=8)
        lbl_res.grid(row=5, column=0, columnspan=3, sticky="ew", padx=3, pady=(14,4))
        lbl_det = tk.Label(f_g, text="", bg="white", fg="#616A6B", font=("Arial",8), wraplength=580, justify="left")
        lbl_det.grid(row=6, column=0, columnspan=3, sticky="ew", padx=3)

        def _build_obj():
            try:
                ta = float(v_tA.get().strip().replace(",","."))
                tb = float(v_tB.get().strip().replace(",","."))
            except: return None
            rows = []
            for i in range(2):
                s = v_s[i].get().strip().replace(",",".")
                la = v_l[i][0].get().strip().replace(",",".")
                lb = v_l[i][1].get().strip().replace(",",".")
                if not s: continue
                try:
                    rows.append([float(s),
                                 float(la) if la else None,
                                 float(lb) if lb else None])
                except: pass
            if not rows: return None
            return {"trims": [ta, tb], "rows": rows}

        def _preview(*_a):
            obj = _build_obj()
            if not obj:
                lbl_res.config(text="—"); lbl_det.config(text=""); return
            vol, det = self._interp_trim_table(obj, s_act, trim_act)
            if vol is None:
                lbl_res.config(text="Faltan datos"); lbl_det.config(text=str(det)); return
            lbl_res.config(text=f"VOLUMEN INTERPOLADO:  {vol:,.0f} L".replace(",","."))
            lbl_det.config(text=det)

        for vv in [v_tA, v_tB] + v_s + v_l[0] + v_l[1]:
            vv.trace_add("write", _preview)
        _preview()

        f_foot = tk.Frame(top, bg="#154360", height=48)
        f_foot.pack(fill="x", side="bottom"); f_foot.pack_propagate(False)

        def _guardar():
            obj = _build_obj()
            if not obj:
                self.get_var(var_key).set("")
            else:
                obj["rows"].sort(key=lambda r: r[0])
                self.get_var(var_key).set(json.dumps(obj))
            top.destroy()
            if agua: self.calc_volumen_agua_ui(etapa, tk_name)
            else: self.calc_volumen_prod_ui(etapa, tk_name)

        tk.Button(f_foot, text="  [OK] GUARDAR Y APLICAR  ", bg="#27AE60", fg="white",
                  font=("Arial",9,"bold"), relief="flat", cursor="hand2",
                  command=_guardar).pack(side="left", padx=14, pady=8, ipadx=10, ipady=4)
        tk.Button(f_foot, text="Tabla completa / CSV…", bg="#5D6D7E", fg="white", font=("Arial",8), relief="flat",
                  command=lambda: (top.destroy(), self.abrir_tabla_calibrado_trim(etapa, tk_name, agua=agua))
                  ).pack(side="left", padx=6, pady=8, ipadx=6, ipady=3)
        tk.Button(f_foot, text="Cancelar", bg="#7F8C8D", fg="white", font=("Arial",8), relief="flat",
                  command=top.destroy).pack(side="right", padx=10, pady=8, ipadx=6, ipady=3)

    def abrir_tabla_calibrado_trim(self, etapa, tk_name, label_col2="LITROS", agua=False):
        """Editor de tabla de calibrado con múltiples columnas de asiento (trim) para
        buques/barcazas. Cada sondaje tiene un volumen por cada asiento (p.ej. 0, 0.5,
        1.0, 1.5 m de diferencia popa-proa). El volumen final se interpola en 2D
        (sondaje × asiento real). Se guarda en {etapa}_{tanque}_tabla_trim[_agua]_json."""
        import json, csv as csv_mod
        if agua and label_col2 == "LITROS": label_col2 = "LTS.AGUA"
        top = tk.Toplevel(self.root)
        top.title(f"Calibrado × Asiento (Trim{', AGUA' if agua else ''}) — {tk_name} ({etapa.upper()})")
        top.geometry("1040x660")
        top.resizable(True, True)
        var_key = f"{etapa}_{tk_name}_tabla_trim{'_agua' if agua else ''}_json"

        fh = tk.Frame(top, bg="#1A5276", height=50)
        fh.pack(fill="x"); fh.pack_propagate(False)
        tk.Label(fh, text=f"CALIBRADO × ASIENTO  |  {tk_name}  |  {etapa.upper()}",
                 bg="#1A5276", fg="white", font=("Arial",10,"bold")).pack(side="left", padx=16, pady=12)
        tk.Label(fh, text=f"Col.2+: {label_col2} por asiento (m)", bg="#1A5276", fg="#AED6F1",
                 font=("Arial",8)).pack(side="right", padx=12)

        f_info = tk.Frame(top, bg="#EBF5FB"); f_info.pack(fill="x")
        tk.Label(f_info,
                 text="BUQUE / BARCAZA — La tabla de sondaje trae un volumen por cada 'asiento' (diferencia popa-proa, en metros). "
                      "Cargá una columna por cada asiento de tu tabla (típico: 0, 0.5, 1.0, 1.5). El programa interpola el volumen "
                      "según el sondaje corregido y el asiento (Trimación = Calado Popa − Calado Proa) real de la medición.",
                 bg="#EBF5FB", fg="#1A5276", font=("Arial",7), anchor="w",
                 wraplength=1000, justify="left").pack(fill="x", padx=10, pady=4)

        # ── Configuración de columnas de asiento ───────────────────────────
        f_cols = tk.Frame(top, bg="#D6EAF8", pady=4); f_cols.pack(fill="x")
        tk.Label(f_cols, text="Asientos (m), separados por coma:", bg="#D6EAF8",
                 font=("Arial",8,"bold"), fg="#1A5276").pack(side="left", padx=(12,4))
        var_trims = tk.StringVar(value="0, 0.5, 1.0, 1.5")
        e_trims = tk.Entry(f_cols, textvariable=var_trims, width=30, font=("Arial",9), justify="center")
        e_trims.pack(side="left", padx=4)
        tk.Button(f_cols, text="Aplicar columnas", bg="#2874A6", fg="white", font=("Arial",8,"bold"),
                  command=lambda: _aplicar_columnas()).pack(side="left", padx=8, ipadx=6, ipady=2)
        lbl_pts = tk.Label(f_cols, text="0 filas", bg="#D6EAF8", font=("Arial",8), fg="#1A5276")
        lbl_pts.pack(side="right", padx=12)

        f_ctrl = tk.Frame(top, bg="#EAF2F8", pady=6); f_ctrl.pack(fill="x")
        tk.Button(f_ctrl, text="Cargar CSV", bg="#1A5276", fg="white", font=("Arial",8,"bold"),
                  command=lambda: _cargar_csv()).pack(side="left", padx=10, pady=4, ipadx=8, ipady=3)
        tk.Button(f_ctrl, text="Exportar CSV", bg="#2980B9", fg="white", font=("Arial",8),
                  command=lambda: _exportar_csv()).pack(side="left", padx=4, pady=4, ipadx=6, ipady=3)
        tk.Button(f_ctrl, text="+ Agregar fila", bg="#16A085", fg="white", font=("Arial",8),
                  command=lambda: _add_row()).pack(side="left", padx=4, pady=4, ipadx=6, ipady=3)
        tk.Button(f_ctrl, text="Ordenar", bg="#7F8C8D", fg="white", font=("Arial",8),
                  command=lambda: _ordenar()).pack(side="left", padx=4, pady=4, ipadx=6, ipady=3)
        tk.Button(f_ctrl, text="Limpiar todo", bg="#C0392B", fg="white", font=("Arial",8),
                  command=lambda: _limpiar()).pack(side="left", padx=4, pady=4, ipadx=6, ipady=3)
        tk.Label(f_ctrl, text="CSV: sondaje_mm, v_asiento0, v_asiento1, ...  (1ra fila = encabezado con asientos)",
                 bg="#EAF2F8", font=("Arial",7), fg="#1A5276").pack(side="right", padx=8)

        f_tbl = tk.Frame(top, bg="white"); f_tbl.pack(fill="both", expand=True, padx=10, pady=6)
        canvas_t = tk.Canvas(f_tbl, bg="white", highlightthickness=0)
        sb_t = ttk.Scrollbar(f_tbl, orient="vertical", command=canvas_t.yview)
        scroll_inner = tk.Frame(canvas_t, bg="white")
        scroll_inner.bind("<Configure>", lambda e: canvas_t.configure(scrollregion=canvas_t.bbox("all")))
        win_id = canvas_t.create_window((0,0), window=scroll_inner, anchor="nw")
        canvas_t.bind("<Configure>", lambda e: canvas_t.itemconfig(win_id, width=e.width))
        canvas_t.configure(yscrollcommand=sb_t.set)
        sb_t.pack(side="right", fill="y"); canvas_t.pack(fill="both", expand=True)
        canvas_t.bind("<MouseWheel>", lambda e: canvas_t.yview_scroll(int(-1*(e.delta/120)) if e.delta else (1 if e.num==5 else -1), "units"))
        canvas_t.bind("<Button-4>", lambda e: canvas_t.yview_scroll(-1,"units"))
        canvas_t.bind("<Button-5>", lambda e: canvas_t.yview_scroll(1,"units"))

        state = {"trims": [0.0, 0.5, 1.0, 1.5]}
        rows_data = []   # cada fila: {"sond": StringVar, "vols": [StringVar, ...]}

        def _parse_trims():
            raw = var_trims.get().replace(";", ",")
            out = []
            for tok in raw.split(","):
                tok = tok.strip().replace(",", ".")
                if not tok: continue
                try: out.append(float(tok))
                except: pass
            return out or [0.0]

        def _update_count():
            n = len([r for r in rows_data if r["sond"].get().strip()])
            lbl_pts.config(text=f"{n} filas × {len(state['trims'])} asientos")

        def _render_rows():
            for w in scroll_inner.winfo_children(): w.destroy()
            hdr = tk.Frame(scroll_inner, bg="#1A5276"); hdr.pack(fill="x")
            tk.Label(hdr, text="#", width=4, bg="#1A5276", fg="white", font=("Arial",8,"bold")).pack(side="left", padx=2, pady=3)
            tk.Label(hdr, text="SONDAJE (mm)", width=14, bg="#1A5276", fg="white", font=("Arial",8,"bold")).pack(side="left", padx=2)
            for tval in state["trims"]:
                tk.Label(hdr, text=f"{label_col2} @ {tval:g}m", width=13, bg="#21618C",
                         fg="white", font=("Arial",8,"bold")).pack(side="left", padx=2)
            tk.Label(hdr, text="Quitar", width=6, bg="#1A5276", fg="white", font=("Arial",8,"bold")).pack(side="left", padx=2)
            for i, rd in enumerate(rows_data):
                rf = tk.Frame(scroll_inner, bg="#EBF5FB" if i%2==0 else "white"); rf.pack(fill="x")
                tk.Label(rf, text=str(i+1), width=4, bg=rf["bg"], font=("Arial",7)).pack(side="left", padx=2, pady=2)
                tk.Entry(rf, textvariable=rd["sond"], width=14, justify="center", font=("Arial",9),
                         bg="#FDFEFE").pack(side="left", padx=2)
                for vv in rd["vols"]:
                    tk.Entry(rf, textvariable=vv, width=13, justify="center", font=("Arial",9),
                             bg="#F4FAFF").pack(side="left", padx=2)
                tk.Button(rf, text="X", bg="#E74C3C", fg="white", width=4, font=("Arial",7),
                          command=lambda idx=i: _del_row(idx)).pack(side="left", padx=2)
            _update_count()

        def _add_row(s_val="", v_vals=None):
            v_vals = v_vals or []
            vols = [tk.StringVar(value=str(v_vals[c]) if c < len(v_vals) else "")
                    for c in range(len(state["trims"]))]
            rows_data.append({"sond": tk.StringVar(value=str(s_val)), "vols": vols})
            _render_rows()

        def _del_row(idx):
            if 0 <= idx < len(rows_data): del rows_data[idx]
            _render_rows()

        def _limpiar():
            rows_data.clear(); _render_rows()

        def _aplicar_columnas():
            new_trims = _parse_trims()
            old_n = len(state["trims"])
            state["trims"] = new_trims
            for rd in rows_data:
                cur = rd["vols"]
                if len(new_trims) > len(cur):
                    cur.extend(tk.StringVar(value="") for _ in range(len(new_trims) - len(cur)))
                elif len(new_trims) < len(cur):
                    del cur[len(new_trims):]
            var_trims.set(", ".join(f"{t:g}" for t in new_trims))
            _render_rows()

        def _ordenar():
            def _key(rd):
                try: return float(rd["sond"].get().strip().replace(",","."))
                except: return float("inf")
            rows_data.sort(key=_key)
            _render_rows()

        def _cargar_csv():
            f_path = filedialog.askopenfilename(
                title="Cargar tabla de calibrado × asiento",
                filetypes=[("CSV/TXT","*.csv *.txt"),("Todos","*.*")], parent=top)
            if not f_path: return
            try:
                raw_rows = []
                with open(f_path, "r", encoding="utf-8-sig") as fcsv:
                    for line in fcsv:
                        line = line.strip()
                        if not line or line.startswith("#"): continue
                        for sep in [",",";","\t"]:
                            if sep in line: parts = [p.strip() for p in line.split(sep)]; break
                        else: parts = line.split()
                        raw_rows.append(parts)
                if not raw_rows:
                    messagebox.showwarning("Sin datos", "Archivo vacío.", parent=top); return
                # ¿Primera fila es encabezado de asientos? (col0 no numérica)
                hdr_trims = None
                try: float(raw_rows[0][0].replace(",","."))
                except:
                    hdr_trims = []
                    for c in raw_rows[0][1:]:
                        cc = "".join(ch for ch in c.replace(",",".") if (ch.isdigit() or ch in ".-"))
                        try: hdr_trims.append(float(cc))
                        except: pass
                    raw_rows = raw_rows[1:]
                if hdr_trims:
                    state["trims"] = hdr_trims
                    var_trims.set(", ".join(f"{t:g}" for t in hdr_trims))
                rows_data.clear()
                for parts in raw_rows:
                    try:
                        s_v = parts[0].strip().replace(",",".")
                        float(s_v)
                    except: continue
                    v_vals = [p.strip().replace(",",".") for p in parts[1:]]
                    _add_row(s_v, v_vals)
                if not rows_data: _add_row()
                _render_rows()
                messagebox.showinfo("OK", f"Cargadas {len(raw_rows)} filas, {len(state['trims'])} asientos.", parent=top)
            except Exception as ex:
                messagebox.showerror("Error", f"No se pudo leer:\n{ex}", parent=top)

        def _exportar_csv():
            filas = []
            for rd in rows_data:
                s = rd["sond"].get().strip()
                if not s: continue
                filas.append([s] + [vv.get().strip() for vv in rd["vols"]])
            if not filas:
                messagebox.showwarning("Sin datos", "No hay filas para exportar.", parent=top); return
            f_path = filedialog.asksaveasfilename(
                title="Exportar calibrado × asiento", defaultextension=".csv",
                filetypes=[("CSV","*.csv"),("Texto","*.txt")],
                initialfile=f"calibrado_trim_{tk_name}_{etapa}.csv", parent=top)
            if not f_path: return
            try:
                with open(f_path, "w", newline="", encoding="utf-8") as fc:
                    writer = csv_mod.writer(fc)
                    writer.writerow(["sondaje_mm"] + [f"{t:g}" for t in state["trims"]])
                    for row in filas: writer.writerow(row)
                messagebox.showinfo("Exportado", f"Exportadas {len(filas)} filas.\n{f_path}", parent=top)
            except Exception as ex:
                messagebox.showerror("Error", f"No se pudo exportar:\n{ex}", parent=top)

        # Cargar datos existentes
        existing = self.get_var(var_key).get()
        if existing:
            try:
                obj = json.loads(existing)
                state["trims"] = [float(t) for t in obj.get("trims", [0.0])] or [0.0]
                var_trims.set(", ".join(f"{t:g}" for t in state["trims"]))
                for r in obj.get("rows", []):
                    if not r: continue
                    _add_row(r[0], r[1:])
            except: pass
        if not rows_data:
            for _ in range(8): _add_row()
        _render_rows()

        f_foot = tk.Frame(top, bg="#154360", height=50); f_foot.pack(fill="x", side="bottom")
        f_foot.pack_propagate(False)

        def _guardar():
            trims = _parse_trims()
            state["trims"] = trims
            rows = []
            for rd in rows_data:
                s = rd["sond"].get().strip().replace(",",".")
                if not s: continue
                try: s_f = float(s)
                except: continue
                vols = []
                any_v = False
                for c in range(len(trims)):
                    raw = rd["vols"][c].get().strip().replace(",",".") if c < len(rd["vols"]) else ""
                    if raw:
                        try: vols.append(float(raw)); any_v = True
                        except: vols.append(None)
                    else:
                        vols.append(None)
                if any_v: rows.append([s_f] + vols)
            _recalc = self.calc_volumen_agua_ui if agua else self.calc_volumen_prod_ui
            if not rows:
                self.get_var(var_key).set(""); top.destroy()
                _recalc(etapa, tk_name); return
            rows.sort(key=lambda r: r[0])
            self.get_var(var_key).set(json.dumps({"trims": trims, "rows": rows}))
            top.destroy()
            _recalc(etapa, tk_name)
            messagebox.showinfo("Guardado",
                f"Tabla × asiento guardada: {len(rows)} sondajes × {len(trims)} asientos.\n"
                f"Asientos: {', '.join(f'{t:g}' for t in trims)} m\n"
                f"Rango sondaje: {rows[0][0]:.0f} → {rows[-1][0]:.0f} mm")

        tk.Button(f_foot, text="  [OK] GUARDAR TABLA  ", bg="#27AE60", fg="white",
                  font=("Arial",9,"bold"), relief="flat", cursor="hand2",
                  command=_guardar).pack(side="left", padx=16, pady=8, ipadx=12, ipady=4)
        tk.Button(f_foot, text="Cancelar", bg="#7F8C8D", fg="white", font=("Arial",8), relief="flat",
                  command=top.destroy).pack(side="right", padx=8, pady=8, ipadx=6, ipady=3)
        tk.Label(f_foot, text="El volumen se interpola por sondaje y por asiento (Popa−Proa) real",
                 bg="#154360", fg="#AED6F1", font=("Arial",7)).pack(side="right", padx=12)



    def crear_interp_trim_inline(self, etapa, tk_name, parent, start_row, agua=False):
        """Sección embebida en la ficha del tanque para cargar la interpolación
        × asiento (trim) sin abrir el diálogo aparte. Mismo formato y variable
        que abrir_interp_trim_rapida ({etapa}_{tanque}_tabla_trim[_agua]_json).
        Aplica y recalcula en vivo al tipear. Devuelve las filas de grilla usadas."""
        import json
        var_key = f"{etapa}_{tk_name}_tabla_trim{'_agua' if agua else ''}_json"
        _hbg = "#117864" if agua else "#1A5276"
        _titulo = "AGUA" if agua else "PRODUCTO"
        _lit = "Lts.Agua" if agua else "Litros"
        _bgl = "#D1F2EB" if agua else "#D6EAF8"
        fb = ("Arial", 8, "bold")

        # Prefill desde la tabla guardada (si es chica, igual que el diálogo rápido)
        tA, tB = 0.0, 0.5
        pre = [["", ""], ["", ""]]
        pre_s = ["", ""]
        existing = self.get_var(var_key).get()
        if existing:
            try:
                obj0 = json.loads(existing)
                tr0 = [float(t) for t in obj0.get("trims", [])]
                rw0 = obj0.get("rows", [])
                if len(tr0) >= 2: tA, tB = tr0[0], tr0[1]
                for i in range(min(2, len(rw0))):
                    pre_s[i] = f"{rw0[i][0]:g}"
                    if len(rw0[i]) > 1 and rw0[i][1] is not None: pre[i][0] = f"{rw0[i][1]:g}"
                    if len(rw0[i]) > 2 and rw0[i][2] is not None: pre[i][1] = f"{rw0[i][2]:g}"
            except: pass

        v_tA = tk.StringVar(value=f"{tA:g}"); v_tB = tk.StringVar(value=f"{tB:g}")
        v_s = [tk.StringVar(value=pre_s[0]), tk.StringVar(value=pre_s[1])]
        v_l = [[tk.StringVar(value=pre[0][0]), tk.StringVar(value=pre[0][1])],
               [tk.StringVar(value=pre[1][0]), tk.StringVar(value=pre[1][1])]]
        v_res = tk.StringVar(value="—")

        r = start_row
        tk.Label(parent, text=f"── INTERPOLACIÓN × ASIENTO ({_titulo}) — 2da fila opcional ──",
                 font=fb, fg="white", bg=_hbg).grid(row=r, column=0, columnspan=8, pady=(8,2), sticky="ew")
        tk.Button(parent, text="Tabla completa/CSV…", bg="#5D6D7E", fg="white", font=("Arial", 7),
                  relief="flat", cursor="hand2",
                  command=lambda: self.abrir_tabla_calibrado_trim(etapa, tk_name, agua=agua)
                  ).grid(row=r, column=8, sticky="ew", padx=1, pady=(8, 2))
        r += 1
        _cols = [("Asiento A (m)", v_tA), ("Asiento B (m)", v_tB),
                 ("Sondaje 1 (mm)", v_s[0]), (f"{_lit} 1 @A", v_l[0][0]), (f"{_lit} 1 @B", v_l[0][1]),
                 ("Sondaje 2 (mm)", v_s[1]), (f"{_lit} 2 @A", v_l[1][0]), (f"{_lit} 2 @B", v_l[1][1])]
        for ci, (lbl_t, _v) in enumerate(_cols):
            tk.Label(parent, text=lbl_t, font=fb, bg=_bgl, anchor="w").grid(
                row=r, column=ci, sticky="ew", padx=1, pady=(2, 0))
            tk.Entry(parent, textvariable=_v, justify="center", bg="white").grid(
                row=r + 1, column=ci, sticky="ew", padx=1)
        tk.Label(parent, text="Vol. interp. (L)", font=fb, bg=_bgl, anchor="w").grid(
            row=r, column=8, sticky="ew", padx=1, pady=(2, 0))
        tk.Entry(parent, textvariable=v_res, justify="center", state="readonly",
                 readonlybackground="#FEF9E7").grid(row=r + 1, column=8, sticky="ew", padx=1)

        def _build_obj():
            try:
                ta = float(v_tA.get().strip().replace(",", "."))
                tb = float(v_tB.get().strip().replace(",", "."))
            except: return None
            rows = []
            for i in range(2):
                s = v_s[i].get().strip().replace(",", ".")
                la = v_l[i][0].get().strip().replace(",", ".")
                lb = v_l[i][1].get().strip().replace(",", ".")
                if not s: continue
                try:
                    rows.append([float(s),
                                 float(la) if la else None,
                                 float(lb) if lb else None])
                except: pass
            if not rows: return None
            return {"trims": [ta, tb], "rows": rows}

        def _recalc():
            if agua: self.calc_volumen_agua_ui(etapa, tk_name)
            else: self.calc_volumen_prod_ui(etapa, tk_name)

        def _preview(obj):
            _s_key = f"{etapa}_{tk_name}_agua_s_real" if agua else f"{etapa}_{tk_name}_s_corr"
            s_act = self.parse_float(self.get_var(_s_key).get())
            trim_act = self.parse_float(self.get_var(f"{etapa}_Trimación").get() or "0")
            vol, _det = self._interp_trim_table(obj, s_act, trim_act)
            v_res.set(f"{vol:,.0f}".replace(",", ".") if vol is not None else "—")

        def _aplicar(*_a):
            obj = _build_obj()
            if obj is None:
                v_res.set("—")
                # Solo borrar la tabla guardada si el usuario vació los campos
                _vacio = not any(vv.get().strip() for vv in v_s + v_l[0] + v_l[1])
                if _vacio and self.get_var(var_key).get():
                    self.get_var(var_key).set("")
                    _recalc()
                return
            obj["rows"].sort(key=lambda rr: rr[0])
            self.get_var(var_key).set(json.dumps(obj))
            _recalc()
            _preview(obj)

        for vv in [v_tA, v_tB] + v_s + v_l[0] + v_l[1]:
            vv.trace_add("write", _aplicar)
        # Vista previa inicial sin tocar lo guardado
        _obj_ini = _build_obj()
        if _obj_ini: _preview(_obj_ini)
        return 3

    def abrir_popup_detalle(self, etapa, tk_name):
        try:
            popup = tk.Toplevel(self.root)
            popup.title(f"Ficha de Carga: {tk_name} ({etapa.upper()})")
            popup.geometry("1120x680")
            popup.resizable(True, True)

            # Footer PRIMERO (fix pack order)
            f_bot_fixed = tk.Frame(popup, bg="#2C3E50", bd=0, height=50)
            f_bot_fixed.pack(side="bottom", fill="x")
            f_bot_fixed.pack_propagate(False)

            # ── Banner de tipo de medición (HEADER) ───────────────────────────
            _tm_now   = self.get_tipo_medio()
            _tipo_colores = {
                "BUQUE":              ("#1A3A5C","#5DADE2","⚓  BUQUE — Medición Marítima de Líquidos"),
                "BARCAZA":            ("#1A3A5C","#5DADE2","⚓  BARCAZA — Medición Marítima"),
                "BUQUE GASERO/GLP":   ("#4A235A","#A569BD","⚙  BUQUE GASERO/GLP — Presión/Temperatura/Factor Z"),
                "BUQUE QUIMIQUERO":   ("#1A3A5C","#76D7C4","⚗  BUQUE QUIMIQUERO — Productos Químicos"),
                "BUQUE METANERO/GNL": ("#1C2756","#85C1E9","❄  BUQUE METANERO/GNL — Gas Natural Licuado"),
                "DRAFT SURVEY":       ("#1A3A5C","#F4D03F","⚓  DRAFT SURVEY — Estimación por Calados"),
                "TANQUE FIJO":        ("#1A3B1A","#58D68D","🛢  TANQUE TECHO FIJO — Medición Terrestre"),
                "TANQUE FLOTANTE":    ("#1A3B1A","#58D68D","🛢  TANQUE TECHO FLOTANTE — Sondaje c/Pontón"),
                "ESFERA DE GAS":      ("#3A2010","#F0B429","⚙  ESFERA DE GAS — Almacenamiento Presurizado"),
                "CAMION CISTERNA":    ("#2C1810","#E59866","🚛  CAMION CISTERNA — Medición Automotor"),
                "CAMION GAS/GLP":     ("#2C1810","#E59866","🚛  CAMION GAS/GLP — Transporte de Gas"),
                "OLEODUCTO":          ("#1C1C2C","#A9CCE3","≋  OLEODUCTO — Línea de Líquidos"),
                "GASODUCTO":          ("#1C1C2C","#A9CCE3","≋  GASODUCTO — Línea de Gas"),
                "POLIDUCTO":          ("#1C1C2C","#A9CCE3","≋  POLIDUCTO — Línea Múltiple"),
                "ELECTRICO":          ("#1A1A1A","#F9E79F","⚡  ELÉCTRICO — Medición de Energía"),
            }
            _tbg, _tfg, _tlbl = _tipo_colores.get(_tm_now,
                                 ("#222222","#FFFFFF", f"■  {_tm_now}"))
            f_banner = tk.Frame(popup, bg=_tbg, height=32)
            f_banner.pack(side="top", fill="x")
            f_banner.pack_propagate(False)
            # Ícono+tipo a la izquierda
            tk.Label(f_banner, text=_tlbl,
                     bg=_tbg, fg=_tfg,
                     font=("Arial", 9, "bold"), pady=0).pack(side="left", padx=12)
            # Tanque y etapa a la derecha
            _etapa_colors = {"ini": "#27AE60", "fin": "#E74C3C"}
            _etapa_lbl = "INICIAL" if etapa.lower() in ("ini","inicial") else "FINAL" if etapa.lower() in ("fin","final") else etapa.upper()
            _etapa_col = _etapa_colors.get(etapa.lower()[:3], "#AED6F1")
            tk.Label(f_banner, text=f"[ {etapa.upper()} ]  {tk_name}",
                     bg=_tbg, fg=_etapa_col,
                     font=("Arial", 9, "bold")).pack(side="right", padx=12)

            # Canvas + scroll
            canvas = tk.Canvas(popup, bg="#f5f5f5", highlightthickness=0)
            v_scroll = ttk.Scrollbar(popup, orient="vertical", command=canvas.yview)
            canvas.configure(yscrollcommand=v_scroll.set)
            v_scroll.pack(side="right", fill="y")
            canvas.pack(side="left", fill="both", expand=True)

            frame_main = ttk.Frame(canvas)
            cv_win = canvas.create_window((0, 0), window=frame_main, anchor="nw")

            def _fc(e): canvas.configure(scrollregion=canvas.bbox("all"))
            def _cc(e): canvas.itemconfig(cv_win, width=e.width)
            frame_main.bind("<Configure>", _fc)
            canvas.bind("<Configure>", _cc)
            def _mw(e): canvas.yview_scroll(int(-1*(e.delta/120)) if e.delta else (1 if e.num==5 else -1), "units")
            canvas.bind("<MouseWheel>", _mw)
            canvas.bind("<Button-4>", _mw)
            canvas.bind("<Button-5>", _mw)

            for i in range(9): frame_main.grid_columnconfigure(i, weight=1)
            fb = ("Arial", 8, "bold")

            # ── Flags de tipo ─────────────────────────────────────────────────
            _tm   = self.get_tipo_medio()
            _mar  = self.es_maritimo()
            _tie  = self.es_tierra()
            _cam  = self.es_camion()
            _gas  = self.es_gasero() and not ("METANERO" in _tm or "GNL" in _tm)
            _met  = "METANERO" in _tm or "GNL" in _tm
            _esf  = self.es_esfera()
            _duc  = self.es_ducto()
            _el   = self.es_electrico()
            _cgb  = self.es_camion_gas()
            _flot = "FLOTANTE" in _tm
            # UTI (sondaje con UTI): SOLO para buques con carga LIQUIDA convencional
            # (BUQUE, BARCAZA, QUIMIQUERO). NO para gaseros, metaneros, esferas, gas.
            _es_maritimo_liq = _tm in ("BUQUE", "BARCAZA", "BUQUE QUIMIQUERO", "DRAFT SURVEY")
            _show_uti  = _es_maritimo_liq            # Sondaje UTI solo para líquidos marítimos
            _show_vcf  = not (_el or _duc or _cgb or _gas or _met or _esf)  # VCF para líquidos
            _show_prod = not _el                      # Producto no aplica a electricidad
            # Agua: solo para mediciones de líquidos (no gas, no eléctrico, no ducto)
            _show_agua_flag = _es_maritimo_liq or _tie or (_cam and not _cgb)

            # ══════════════════════════════════════════════════════════════════
            # SECCIÓN 1: Documento / Producto / Tabla VCF  (fila 0-1)
            # ══════════════════════════════════════════════════════════════════
            _s1_fields = [
                ("Documento Asignado", "ddt_assign", "combo", 3),
            ]
            if _show_prod:
                _s1_fields.append(("Producto", "prod_name", "entry", 3))
            if _show_vcf:
                _s1_fields.append(("Tabla VCF", "tabla_vcf", "combo_vcf", 3))
            self.crear_fila_popup(etapa, tk_name, frame_main, 0, _s1_fields, bg_row="#ddd")

            # ══════════════════════════════════════════════════════════════════
            # SECCIÓN 2: Sondaje UTI  (fila 2-3) — solo para tipos de sondaje
            # ══════════════════════════════════════════════════════════════════
            if _show_uti:
                self.crear_fila_popup(etapa, tk_name, frame_main, 2, [
                    ("Nro Uti",              "num_uti",   "entry",    2),
                    ("Medida Uti",           "alt_uti",   "entry",    2),
                    ("Desc.Tubo Sondaje",    "desc_tubo", "entry",    2),
                    ("Sondaje Corregido",    "s_corr",    "entry_ro", 2),
                    ("Temperatura UTI",      "temp",      "entry",    1),
                ], bg_row="#eee")

            # ══════════════════════════════════════════════════════════════════
            # SECCIÓN 3: Alturas / Sondajes 1 y 2  (fila 4-5) — solo sondaje
            # ══════════════════════════════════════════════════════════════════
            if _show_uti:
                self.crear_fila_popup(etapa, tk_name, frame_main, 4, [
                    ("Alt.Referencia",  "alt_ref",       "entry",    2),
                    ("Sondaje 1",       "prod_s1",       "entry",    1),
                    ("Litros 1",        "prod_l1",       "entry",    1),
                    ("Sondaje 2",       "prod_s2",       "entry",    1),
                    ("Litros 2",        "prod_l2",       "entry",    1),
                    ("Litros Naturales","vol_nat_prod",   "entry_ro", 3),
                ], bg_row="#eee")

            # ═══════════════════════════════════════════════════════════════
            # A partir de aquí usamos _cur_row dinámico
            # ═══════════════════════════════════════════════════════════════
            _cur_row = 6 if _show_uti else 2

            # ── Interp × asiento (PRODUCTO) embebida — buques líquidos ─────
            if _es_maritimo_liq:
                _cur_row += self.crear_interp_trim_inline(etapa, tk_name, frame_main, _cur_row, agua=False)

            # ── Tierra / Camión cisterna ──────────────────────────────────
            if _tie or (_cam and not _cgb):
                tk.Label(frame_main, text="── SONDAJE Y VOLUMEN ──",
                         font=fb, fg="#1B3A5C", bg="#D6EAF8").grid(
                         row=_cur_row, column=0, columnspan=9, pady=(8,2), sticky="ew")
                _cur_row += 1
                if _cam:
                    _tf = [
                        ("Varilla / Nivel (mm)", f"{etapa}_{tk_name}_s_tierra"),
                        ("Vol. Bruto (L)",        f"{etapa}_{tk_name}_vol_bruto"),
                        ("Ticket Planta (L)",     f"{etapa}_{tk_name}_ticket_l"),
                        ("Ticket Planta (Kg)",    f"{etapa}_{tk_name}_ticket_k"),
                    ]
                else:
                    _tf = [
                        ("Sondaje (mm)",             f"{etapa}_{tk_name}_s_tierra"),
                        ("Sondaje Corregido (mm)",   f"{etapa}_{tk_name}_s_corr"),
                        ("Vol. Bruto (L)",            f"{etapa}_{tk_name}_vol_bruto"),
                    ]
                    if _flot:
                        _tf.insert(1, ("Offset Techo (mm)", f"{etapa}_{tk_name}_tf_offset"))
                for ci, (lbl_t, var_t) in enumerate(_tf):
                    ro_t = "s_corr" in var_t or "vol_bruto" in var_t
                    tk.Label(frame_main, text=lbl_t, font=fb, bg="#D6EAF8", anchor="w").grid(
                        row=_cur_row, column=ci, sticky="ew", padx=1, pady=(2,0))
                    ent = tk.Entry(frame_main, textvariable=self.get_var(var_t),
                                   justify="center",
                                   state="readonly" if ro_t else "normal",
                                   bg="#dceeff" if ro_t else "white")
                    ent.grid(row=_cur_row+1, column=ci, sticky="ew", padx=1)
                    def _ct(*_a, _e=etapa, _t2=tk_name):
                        try:
                            s = self.parse_float(self.get_var(f"{_e}_{_t2}_s_tierra").get())
                            off = self.parse_float(self.get_var(f"{_e}_{_t2}_tf_offset").get() or self.get_var("car_tf_offset").get() or "0")
                            r  = self.parse_float(self.get_var("car_radio_m").get() or "0")
                            sc = s - off if "FLOTANTE" in self.get_tipo_medio() else s
                            self.get_var(f"{_e}_{_t2}_s_corr").set(str(int(sc)))
                            if r > 0:
                                vol = self.calc_volumen_geometrico_tierra(sc, r)
                                self.get_var(f"{_e}_{_t2}_vol_bruto").set(f"{vol:,.0f}")
                                self.get_var(f"{_e}_{_t2}_vol_nat_prod").set(f"{vol:,.0f}")
                        except: pass
                    if "s_tierra" in var_t or "tf_offset" in var_t:
                        ent.bind("<KeyRelease>", _ct)
                    if _cam:
                        def _cc2(*_a, _e=etapa, _t2=tk_name):
                            try:
                                s2  = self.parse_float(self.get_var(f"{_e}_{_t2}_s_tierra").get())
                                r2  = self.parse_float(self.get_var("car_radio_camion").get() or "0")
                                L2  = self.parse_float(self.get_var("car_largo_camion").get() or "0")
                                if r2 > 0 and L2 > 0:
                                    v2 = self.calc_volumen_cilindro_horizontal(s2, r2, L2)
                                    self.get_var(f"{_e}_{_t2}_vol_bruto").set(f"{v2:,.0f}")
                                    self.get_var(f"{_e}_{_t2}_vol_nat_prod").set(f"{v2:,.0f}")
                            except: pass
                        if "s_tierra" in var_t: ent.bind("<KeyRelease>", _cc2)
                _cur_row += 2

            # ── Gasero/GLP (presión, T, Z - cols 4-8, misma fila que tierra) ─
            if _gas:
                _gr = _cur_row - 2  # misma altura que sección tierra o section 3
                if not (_tie or _cam): _gr = _cur_row; _cur_row += 2
                tk.Label(frame_main, text="── DATOS GAS/PROCESO ──",
                         font=fb, fg="#721c24", bg="#f8d7da").grid(
                         row=_gr, column=4 if (_tie or _cam) else 0, columnspan=5 if (_tie or _cam) else 9,
                         pady=(8,2), sticky="ew")
                _gf = [
                    ("Presion (kPa)",        f"{etapa}_{tk_name}_presion"),
                    ("Temp.Liq (C)",         f"{etapa}_{tk_name}_temp_liq"),
                    ("Factor Z",             f"{etapa}_{tk_name}_factor_z"),
                    ("Dens.Vapor (kg/m3)",   f"{etapa}_{tk_name}_dens_vapor"),
                    ("Fase",                 f"{etapa}_{tk_name}_fase"),
                ]
                _col_off = 4 if (_tie or _cam) else 0
                for ci5, (lg, vg) in enumerate(_gf):
                    tk.Label(frame_main, text=lg, font=fb, bg="#f8d7da").grid(
                        row=_gr, column=_col_off+ci5, sticky="ew", padx=1, pady=(5,0))
                    tk.Entry(frame_main, textvariable=self.get_var(vg),
                             justify="center", bg="#fff0f0").grid(
                             row=_gr+1, column=_col_off+ci5, sticky="ew", padx=1)

            # ── GNL / Metanero: composición molar + fases ─────────────────
            if _met:
                tk.Label(frame_main, text="── COMPOSICION MOLAR GNL ──",
                         font=fb, fg="#1A5276", bg="#D6EAF8").grid(
                         row=_cur_row, column=0, columnspan=9, pady=(10,2), sticky="ew")
                _cur_row += 1
                GNL_C = [
                    ("Metano(CH4)%mol",    f"{etapa}_{tk_name}_gc_CH4"),
                    ("Etano(C2H6)%mol",    f"{etapa}_{tk_name}_gc_C2H6"),
                    ("Propano(C3H8)%mol",  f"{etapa}_{tk_name}_gc_C3H8"),
                    ("Butano(C4H10)%mol",  f"{etapa}_{tk_name}_gc_C4H10"),
                    ("N2 %mol",            f"{etapa}_{tk_name}_gc_N2"),
                    ("CO2 %mol",           f"{etapa}_{tk_name}_gc_CO2"),
                    ("i-C4 %mol",          f"{etapa}_{tk_name}_gc_iC4"),
                    ("n-C5 %mol",          f"{etapa}_{tk_name}_gc_nC5"),
                    ("H2S %mol",           f"{etapa}_{tk_name}_gc_H2S"),
                ]
                _COLS = 5
                for ci, (lc, vc) in enumerate(GNL_C):
                    col = ci % _COLS; roff = ci // _COLS
                    tk.Label(frame_main, text=lc, font=fb, bg="#EBF5FB", anchor="w").grid(
                        row=_cur_row+roff*2, column=col, sticky="ew", padx=1, pady=(4,0))
                    tk.Entry(frame_main, textvariable=self.get_var(vc),
                             justify="center", bg="#EAF4FD", width=10).grid(
                             row=_cur_row+roff*2+1, column=col, sticky="ew", padx=1)
                _n_comp_rows = ((len(GNL_C)-1)// _COLS + 1) * 2
                _cur_row += _n_comp_rows
                # Suma y PM
                v_sum = self.get_var(f"{etapa}_{tk_name}_gc_sum")
                v_PM  = self.get_var(f"{etapa}_{tk_name}_gc_PM")
                def _upd_gnl(*_a, _e=etapa, _t2=tk_name):
                    try:
                        MW = {"CH4":16.04,"C2H6":30.07,"C3H8":44.10,"C4H10":58.12,
                              "iC4":58.12,"nC5":72.15,"N2":28.01,"CO2":44.01,"H2S":34.08}
                        tot=0; PM=0
                        for k,mw in MW.items():
                            p = self.parse_float(self.get_var(f"{_e}_{_t2}_gc_{k}").get() or "0")
                            tot+=p; PM+=p/100*mw
                        self.get_var(f"{_e}_{_t2}_gc_sum").set(f"{tot:.2f}%")
                        self.get_var(f"{_e}_{_t2}_gc_PM").set(f"{PM:.4f}")
                    except: pass
                for _, vc in GNL_C: self.get_var(vc).trace_add("write", _upd_gnl)
                tk.Label(frame_main, text="Suma %mol", font=fb, bg="#EBF5FB").grid(
                    row=_cur_row, column=0, sticky="ew", padx=1, pady=(4,0))
                tk.Label(frame_main, textvariable=v_sum, bg="#dceeff", relief="sunken", font=fb).grid(
                    row=_cur_row+1, column=0, sticky="ew", padx=1)
                tk.Label(frame_main, text="PM calculado", font=fb, bg="#EBF5FB").grid(
                    row=_cur_row, column=1, sticky="ew", padx=1, pady=(4,0))
                tk.Label(frame_main, textvariable=v_PM, bg="#dceeff", relief="sunken").grid(
                    row=_cur_row+1, column=1, sticky="ew", padx=1)
                _cur_row += 2
                # Fases
                tk.Label(frame_main, text="── FASES ──", font=fb, fg="#1A5276", bg="#D5F5E3").grid(
                    row=_cur_row, column=0, columnspan=9, pady=(8,2), sticky="ew")
                _cur_row += 1
                _fase_f = [
                    ("Vol.Liq(m3)",      f"{etapa}_{tk_name}_vol_liq"),
                    ("Vol.Vap(m3)",      f"{etapa}_{tk_name}_vol_vap"),
                    ("Masa Liq(t)",      f"{etapa}_{tk_name}_masa_liq"),
                    ("Masa Vap(kg)",     f"{etapa}_{tk_name}_masa_vap"),
                    ("Dens.Liq(kg/m3)", f"{etapa}_{tk_name}_dens_liq"),
                    ("Dens.Vap(kg/m3)", f"{etapa}_{tk_name}_dens_vap"),
                    ("T Liq (C)",        f"{etapa}_{tk_name}_temp_liq"),
                    ("T Vap (C)",        f"{etapa}_{tk_name}_temp_vap"),
                    ("Presion (kPa)",    f"{etapa}_{tk_name}_pres_gnl"),
                ]
                for ci3, (lf, vf) in enumerate(_fase_f):
                    col3=ci3%_COLS; ro3=ci3//_COLS
                    ro_f = "masa_" in vf
                    tk.Label(frame_main, text=lf, font=fb, bg="#D5F5E3", anchor="w").grid(
                        row=_cur_row+ro3*2, column=col3, sticky="ew", padx=1, pady=(4,0))
                    tk.Entry(frame_main, textvariable=self.get_var(vf),
                             justify="center", bg="#EAFAF1" if not ro_f else "#dceeff",
                             state="readonly" if ro_f else "normal").grid(
                             row=_cur_row+ro3*2+1, column=col3, sticky="ew", padx=1)
                _n_fase_rows = ((len(_fase_f)-1)//_COLS+1)*2
                def _upd_fas(*_a, _e=etapa, _t2=tk_name):
                    try:
                        vl=self.parse_float(self.get_var(f"{_e}_{_t2}_vol_liq").get() or "0")
                        vv=self.parse_float(self.get_var(f"{_e}_{_t2}_vol_vap").get() or "0")
                        dl=self.parse_float(self.get_var(f"{_e}_{_t2}_dens_liq").get() or "0")
                        dv=self.parse_float(self.get_var(f"{_e}_{_t2}_dens_vap").get() or "0")
                        if dl>0: self.get_var(f"{_e}_{_t2}_masa_liq").set(f"{vl*dl/1000:.3f}")
                        if dv>0: self.get_var(f"{_e}_{_t2}_masa_vap").set(f"{vv*dv:.1f}")
                        if vl+vv>0: self.get_var(f"{_e}_{_t2}_vol_nat_prod").set(f"{(vl+vv)*1000:.0f}")
                    except: pass
                for vk in [f"{etapa}_{tk_name}_vol_liq",f"{etapa}_{tk_name}_vol_vap",
                            f"{etapa}_{tk_name}_dens_liq",f"{etapa}_{tk_name}_dens_vap"]:
                    self.get_var(vk).trace_add("write", _upd_fas)
                _cur_row += _n_fase_rows

            # ── Esfera de gas ──────────────────────────────────────────────
            if _esf:
                tk.Label(frame_main, text="── ESFERA DE GAS / PRESION ──",
                         font=fb, fg="#784212", bg="#FEF9E7").grid(
                         row=_cur_row, column=0, columnspan=9, pady=(10,2), sticky="ew")
                _cur_row += 1
                _esf_f = [
                    ("Presion (kPa)",       f"{etapa}_{tk_name}_esf_pres"),
                    ("Temperatura (C)",     f"{etapa}_{tk_name}_esf_temp"),
                    ("Dens.liq(kg/m3)",     f"{etapa}_{tk_name}_esf_dens"),
                    ("Vol.Liq(m3)",         f"{etapa}_{tk_name}_vol_liq"),
                    ("Vol.Gas Base(m3)",    f"{etapa}_{tk_name}_esf_vol_gas"),
                    ("Masa (t)",            f"{etapa}_{tk_name}_esf_masa"),
                    ("Fase",               f"{etapa}_{tk_name}_esf_fase"),
                    ("Producto",           f"{etapa}_{tk_name}_prod_name"),
                ]
                for ci4, (le, ve) in enumerate(_esf_f):
                    ro_e = "vol_gas" in ve or "esf_masa" in ve or "esf_fase" in ve
                    tk.Label(frame_main, text=le, font=fb, bg="#FEF9E7", anchor="w").grid(
                        row=_cur_row, column=ci4, sticky="ew", padx=1, pady=(4,0))
                    tk.Entry(frame_main, textvariable=self.get_var(ve),
                             justify="center",
                             state="readonly" if ro_e else "normal",
                             bg="#dceeff" if ro_e else "#FFFDE7").grid(
                             row=_cur_row+1, column=ci4, sticky="ew", padx=1)
                def _upd_esf(*_a, _e=etapa, _t2=tk_name):
                    try:
                        P  = self.parse_float(self.get_var(f"{_e}_{_t2}_esf_pres").get() or "101.325")
                        T  = self.parse_float(self.get_var(f"{_e}_{_t2}_esf_temp").get() or "15")
                        d  = self.parse_float(self.get_var(f"{_e}_{_t2}_esf_dens").get() or "500")
                        vl = self.parse_float(self.get_var(f"{_e}_{_t2}_vol_liq").get() or "0")
                        prod = self.get_var(f"{_e}_{_t2}_prod_name").get() or "GLP"
                        Z = self.calc_factor_z(P, T)
                        if d>0 and vl>0:
                            self.get_var(f"{_e}_{_t2}_esf_vol_gas").set(f"{self.calc_volumen_base_gas(vl,P,T,Z):.3f}")
                            self.get_var(f"{_e}_{_t2}_esf_masa").set(f"{vl*d/1000:.3f}")
                        self.get_var(f"{_e}_{_t2}_esf_fase").set(self.calc_gas_fase(P,T,prod))
                        self.get_var(f"{_e}_{_t2}_vol_nat_prod").set(f"{vl*1000:.0f}")
                    except: pass
                for fld in ["esf_pres","esf_temp","esf_dens","vol_liq"]:
                    self.get_var(f"{etapa}_{tk_name}_{fld}").trace_add("write", _upd_esf)
                _cur_row += 2

            # ── Ductos ────────────────────────────────────────────────────
            if _duc:
                tk.Label(frame_main, text="── CONTADORES Y CAUDAL ──",
                         font=fb, fg="#1B3A5C", bg="#D6EAF8").grid(
                         row=_cur_row, column=0, columnspan=9, pady=(8,2), sticky="ew")
                _cur_row += 1
                _duc1 = [
                    ("Contador Ini(m3)",  f"{etapa}_{tk_name}_cont_ini"),
                    ("Contador Fin(m3)",  f"{etapa}_{tk_name}_cont_fin"),
                    ("Vol.Linea(m3)",     f"{etapa}_{tk_name}_vol_linea"),
                    ("Caudal(m3/h)",      f"{etapa}_{tk_name}_caudal_mh"),
                    ("Coriolis(kg/h)",    f"{etapa}_{tk_name}_coriolis_kgh"),
                    ("Masa Coriolis(kg)", f"{etapa}_{tk_name}_masa_coriolis"),
                ]
                for ci, (ld, vd) in enumerate(_duc1[:9]):
                    ro_d = "vol_linea" in vd or "masa_coriolis" in vd
                    tk.Label(frame_main, text=ld, font=fb, bg="#D6EAF8", anchor="w").grid(
                        row=_cur_row, column=ci, sticky="ew", padx=1, pady=(4,0))
                    ed = tk.Entry(frame_main, textvariable=self.get_var(vd),
                                  justify="center",
                                  state="readonly" if ro_d else "normal",
                                  bg="#dceeff" if ro_d else "white")
                    ed.grid(row=_cur_row+1, column=ci, sticky="ew", padx=1)
                    def _udv(*a, _e=etapa, _t2=tk_name):
                        try:
                            i2=self.parse_float(self.get_var(f"{_e}_{_t2}_cont_ini").get())
                            f2=self.parse_float(self.get_var(f"{_e}_{_t2}_cont_fin").get())
                            self.get_var(f"{_e}_{_t2}_vol_linea").set(f"{f2-i2:,.3f}")
                        except: pass
                    if "cont_" in vd or "coriolis_kgh" in vd: ed.bind("<KeyRelease>", _udv)
                _cur_row += 2
                # P/T/Z
                tk.Label(frame_main, text="── CORRECCION P/T/Z ──",
                         font=fb, fg="#1B3A5C", bg="#EBF5FB").grid(
                         row=_cur_row, column=0, columnspan=9, pady=(8,2), sticky="ew")
                _cur_row += 1
                _duc2 = [
                    ("P linea(kPa)",   f"{etapa}_{tk_name}_P_lin"),
                    ("T linea(C)",     f"{etapa}_{tk_name}_T_lin"),
                    ("Factor Z",       f"{etapa}_{tk_name}_Z"),
                    ("Vol Base(m3)",   f"{etapa}_{tk_name}_vol_base"),
                    ("Vol Base(Km3)",  f"{etapa}_{tk_name}_vol_base_km3"),
                ]
                for ci2, (ld2, vd2) in enumerate(_duc2):
                    ro2 = "vol_base" in vd2 or vd2.endswith("_Z")
                    tk.Label(frame_main, text=ld2, font=fb, bg="#EBF5FB", anchor="w").grid(
                        row=_cur_row, column=ci2, sticky="ew", padx=1, pady=(4,0))
                    e2 = tk.Entry(frame_main, textvariable=self.get_var(vd2),
                                  justify="center",
                                  state="readonly" if ro2 else "normal",
                                  bg="#dceeff" if ro2 else "white")
                    e2.grid(row=_cur_row+1, column=ci2, sticky="ew", padx=1)
                    def _uz(*a, _e=etapa, _t2=tk_name):
                        try:
                            P2=self.parse_float(self.get_var(f"{_e}_{_t2}_P_lin").get() or self.get_var("car_presion_op").get() or "101.325")
                            T2=self.parse_float(self.get_var(f"{_e}_{_t2}_T_lin").get() or self.get_var("car_temp_op").get() or "15")
                            comp=self._get_cromatografia(f"{_e}_{_t2}")
                            Z2=self.calc_factor_z(P2,T2,composicion_pct=comp if comp else None)
                            self.get_var(f"{_e}_{_t2}_Z").set(f"{Z2:.5f}")
                            vl2=self.parse_float(self.get_var(f"{_e}_{_t2}_vol_linea").get() or "0")
                            vb=self.calc_volumen_base_gas(vl2,P2,T2,Z2)
                            self.get_var(f"{_e}_{_t2}_vol_base").set(f"{vb:,.3f}")
                            self.get_var(f"{_e}_{_t2}_vol_base_km3").set(f"{vb/1000:,.6f}")
                            self.get_var(f"{_e}_{_t2}_vol_nat_prod").set(f"{vb*1000:,.0f}")
                        except: pass
                    if "P_lin" in vd2 or "T_lin" in vd2: e2.bind("<KeyRelease>", _uz)
                _cur_row += 2
                # Cromatografía (gasoducto)
                if self.es_gasoducto():
                    tk.Label(frame_main, text="── CROMATOGRAFIA (% mol) ──",
                             font=fb, fg="#721c24", bg="#f8d7da").grid(
                             row=_cur_row, column=0, columnspan=9, pady=(8,2), sticky="ew")
                    _cur_row += 1
                    for ci_c, cn in enumerate(["CH4","C2H6","C3H8","iC4","nC4","iC5","nC5","C6+","N2"][:9]):
                        tk.Label(frame_main, text=cn, font=fb, bg="#f8d7da").grid(
                            row=_cur_row, column=ci_c, sticky="ew", padx=1, pady=(4,0))
                        ec = tk.Entry(frame_main, textvariable=self.get_var(f"{etapa}_{tk_name}_gc_{cn}",""),
                                      justify="center", bg="#fff0f0", width=7)
                        ec.grid(row=_cur_row+1, column=ci_c, sticky="ew", padx=1)
                        ec.bind("<KeyRelease>", _uz)
                    _cur_row += 2
                # PIG
                tk.Label(frame_main, text="── PIG CALIBRACION ──",
                         font=fb, fg="#6E4B00", bg="#FEF9E7").grid(
                         row=_cur_row, column=0, columnspan=9, pady=(8,2), sticky="ew")
                _cur_row += 1
                _pig = [
                    ("Diam.Nom(pulg)", f"{etapa}_{tk_name}_pig_diam"),
                    ("Long.Tramo(m)",  f"{etapa}_{tk_name}_pig_largo"),
                    ("Vol.Nom(m3)",    f"{etapa}_{tk_name}_pig_vol"),
                    ("Vel.Pig(m/s)",   f"{etapa}_{tk_name}_pig_vel"),
                    ("Obs/Resultado",  f"{etapa}_{tk_name}_pig_obs"),
                ]
                for ci_p, (lp, vp) in enumerate(_pig[:9]):
                    ro_p = "pig_vol" in vp
                    tk.Label(frame_main, text=lp, font=fb, bg="#FEF9E7").grid(
                        row=_cur_row, column=ci_p, sticky="ew", padx=1, pady=(4,0))
                    ep2 = tk.Entry(frame_main, textvariable=self.get_var(vp),
                                   justify="center",
                                   state="readonly" if ro_p else "normal",
                                   bg="#fff9e0" if ro_p else "white")
                    ep2.grid(row=_cur_row+1, column=ci_p, sticky="ew", padx=1)
                    def _upig(*a, _e=etapa, _t2=tk_name):
                        try:
                            d2=self.parse_float(self.get_var(f"{_e}_{_t2}_pig_diam").get() or "0")
                            l2=self.parse_float(self.get_var(f"{_e}_{_t2}_pig_largo").get() or "0")
                            self.get_var(f"{_e}_{_t2}_pig_vol").set(f"{self.calc_pig_volumen_m3(d2,l2):,.4f}")
                        except: pass
                    if "pig_diam" in vp or "pig_largo" in vp: ep2.bind("<KeyRelease>", _upig)
                _cur_row += 2

            # ── Medición eléctrica ─────────────────────────────────────────
            if _el:
                tk.Label(frame_main, text="── MEDICION ELECTRICA ──",
                         font=fb, fg="#1B3A5C", bg="#D5E8D4").grid(
                         row=_cur_row, column=0, columnspan=9, pady=(8,2), sticky="ew")
                _cur_row += 1
                _el1 = [
                    ("Lect.Ini kWh Act",  f"{etapa}_{tk_name}_el_ini_act"),
                    ("Lect.Fin kWh Act",  f"{etapa}_{tk_name}_el_fin_act"),
                    ("kWh Activa",        f"{etapa}_{tk_name}_el_kwh_act"),
                    ("Lect.Ini kWh Rea",  f"{etapa}_{tk_name}_el_ini_rea"),
                    ("Lect.Fin kWh Rea",  f"{etapa}_{tk_name}_el_fin_rea"),
                    ("kWh Reactiva",      f"{etapa}_{tk_name}_el_kwh_rea"),
                    ("Cte Medidor",       f"{etapa}_{tk_name}_el_const"),
                    ("cos fi calc",       f"{etapa}_{tk_name}_el_fp"),
                    ("Dem.Max.(kW)",      f"{etapa}_{tk_name}_el_dem"),
                ]
                for ci_e, (le2, ve2) in enumerate(_el1[:9]):
                    ro_e2 = "kwh_" in ve2 or "el_fp" in ve2
                    tk.Label(frame_main, text=le2, font=fb, bg="#D5E8D4", anchor="w").grid(
                        row=_cur_row, column=ci_e, sticky="ew", padx=1, pady=(4,0))
                    ee = tk.Entry(frame_main, textvariable=self.get_var(ve2),
                                  justify="center",
                                  state="readonly" if ro_e2 else "normal",
                                  bg="#d5f5e3" if ro_e2 else "white")
                    ee.grid(row=_cur_row+1, column=ci_e, sticky="ew", padx=1)
                    def _ue(*a, _e=etapa, _t2=tk_name):
                        try:
                            ct=self.parse_float(self.get_var(f"{_e}_{_t2}_el_const").get() or "1")
                            ia=self.parse_float(self.get_var(f"{_e}_{_t2}_el_ini_act").get() or "0")
                            fa=self.parse_float(self.get_var(f"{_e}_{_t2}_el_fin_act").get() or "0")
                            ir=self.parse_float(self.get_var(f"{_e}_{_t2}_el_ini_rea").get() or "0")
                            fr=self.parse_float(self.get_var(f"{_e}_{_t2}_el_fin_rea").get() or "0")
                            ka=self.calc_energia_electrica(ia,fa,ct)
                            kr=self.calc_energia_electrica(ir,fr,ct)
                            self.get_var(f"{_e}_{_t2}_el_kwh_act").set(f"{ka:,.3f}")
                            self.get_var(f"{_e}_{_t2}_el_kwh_rea").set(f"{kr:,.3f}")
                            self.get_var(f"{_e}_{_t2}_el_fp").set(f"{self.calc_fp_kwh(ka,kr):.4f}")
                            self.get_var(f"{_e}_{_t2}_vol_nat_prod").set(f"{ka:,.3f}")
                        except: pass
                    if any(x in ve2 for x in ["ini_","fin_","el_const"]): ee.bind("<KeyRelease>", _ue)
                _cur_row += 2
                tk.Label(frame_main, text="── TENSION / CORRIENTE ──",
                         font=fb, fg="#1B3A5C", bg="#D5E8D4").grid(
                         row=_cur_row, column=0, columnspan=9, pady=(5,2), sticky="ew")
                _cur_row += 1
                _el2 = [
                    ("Tension(V)",     f"{etapa}_{tk_name}_el_V"),
                    ("Corriente(A)",   f"{etapa}_{tk_name}_el_A"),
                    ("Pot.Apar.(VA)",  f"{etapa}_{tk_name}_el_VA"),
                    ("cos fi medido",  f"{etapa}_{tk_name}_el_fp_med"),
                    ("N fases",        f"{etapa}_{tk_name}_el_fases"),
                ]
                for ci_e2, (le3, ve3) in enumerate(_el2[:9]):
                    ro_e3 = "el_VA" in ve3
                    tk.Label(frame_main, text=le3, font=fb, bg="#D5E8D4").grid(
                        row=_cur_row, column=ci_e2, sticky="ew", padx=1, pady=(4,0))
                    ee3 = tk.Entry(frame_main, textvariable=self.get_var(ve3),
                                   justify="center",
                                   state="readonly" if ro_e3 else "normal",
                                   bg="#d5f5e3" if ro_e3 else "white")
                    ee3.grid(row=_cur_row+1, column=ci_e2, sticky="ew", padx=1)
                    def _uva(*a, _e=etapa, _t2=tk_name):
                        try:
                            V2=self.parse_float(self.get_var(f"{_e}_{_t2}_el_V").get() or "0")
                            A2=self.parse_float(self.get_var(f"{_e}_{_t2}_el_A").get() or "0")
                            self.get_var(f"{_e}_{_t2}_el_VA").set(f"{self.calc_potencia_aparente(V2,A2):,.1f}")
                        except: pass
                    if "el_V" in ve3 or "el_A" in ve3: ee3.bind("<KeyRelease>", _uva)
                _cur_row += 2

            # ── Camión gas/GLP ────────────────────────────────────────────
            if _cgb:
                tk.Label(frame_main, text="── CAMION GAS/GLP ──",
                         font=fb, fg="#721c24", bg="#f8d7da").grid(
                         row=_cur_row, column=0, columnspan=9, pady=(8,2), sticky="ew")
                _cur_row += 1
                _cg = [
                    ("Presion(kPa)",   f"{etapa}_{tk_name}_cg_pres"),
                    ("Temp(C)",        f"{etapa}_{tk_name}_cg_temp"),
                    ("Dens.liq(kg/m3)",f"{etapa}_{tk_name}_cg_dens"),
                    ("Masa(kg)",       f"{etapa}_{tk_name}_cg_masa"),
                    ("Vol.Liq(L)",     f"{etapa}_{tk_name}_cg_vol"),
                    ("Vol.Gas 15C(m3)",f"{etapa}_{tk_name}_cg_vol_gas"),
                ]
                for ci_cg, (lc2, vc2) in enumerate(_cg[:9]):
                    ro_cg = "cg_vol" in vc2
                    tk.Label(frame_main, text=lc2, font=fb, bg="#f8d7da").grid(
                        row=_cur_row, column=ci_cg, sticky="ew", padx=1, pady=(4,0))
                    ecg = tk.Entry(frame_main, textvariable=self.get_var(vc2),
                                   justify="center",
                                   state="readonly" if ro_cg else "normal",
                                   bg="#fff0f0" if ro_cg else "white")
                    ecg.grid(row=_cur_row+1, column=ci_cg, sticky="ew", padx=1)
                    def _ucg(*a, _e=etapa, _t2=tk_name):
                        try:
                            masa=self.parse_float(self.get_var(f"{_e}_{_t2}_cg_masa").get() or "0")
                            dens=self.parse_float(self.get_var(f"{_e}_{_t2}_cg_dens").get() or "500")
                            P3=self.parse_float(self.get_var(f"{_e}_{_t2}_cg_pres").get() or "101.325")
                            T3=self.parse_float(self.get_var(f"{_e}_{_t2}_cg_temp").get() or "15")
                            prod=self.get_var(f"{_e}_{_t2}_prod_name").get() or "GLP"
                            if masa>0 and dens>0:
                                vl3=self.calc_volumen_camion_gas(masa,dens)
                            else:
                                var=self.parse_float(self.get_var(f"{_e}_{_t2}_s_tierra").get() or "0")
                                r3=self.parse_float(self.get_var("car_radio_camion").get() or "0")
                                L3=self.parse_float(self.get_var("car_largo_camion").get() or "0")
                                vl3=self.calc_volumen_cilindro_horizontal(var,r3,L3) if r3>0 and L3>0 and var>0 else 0
                            if vl3>0:
                                Z3=self.calc_factor_z(P3,T3)
                                self.get_var(f"{_e}_{_t2}_cg_vol").set(f"{vl3:,.1f}")
                                self.get_var(f"{_e}_{_t2}_cg_vol_gas").set(f"{self.calc_volumen_base_gas(vl3/1000,P3,T3,Z3):,.3f}")
                                self.get_var(f"{_e}_{_t2}_vol_nat_prod").set(f"{vl3:,.0f}")
                                self.get_var(f"{_e}_{_t2}_fase").set(self.calc_gas_fase(P3,T3,prod))
                        except: pass
                    if any(x in vc2 for x in ["cg_masa","cg_dens","cg_pres","cg_temp"]): ecg.bind("<KeyRelease>", _ucg)
                _cur_row += 2

            # ══════════════════════════════════════════════════════════════
            # AGUA  (solo para mediciones de líquidos — NO para gas/eléctrico/ducto)
            # ══════════════════════════════════════════════════════════════
            _show_agua = _show_agua_flag
            if _show_agua:
                tk.Label(frame_main, text="── CALCULO DE AGUA ──",
                         font=fb, fg="blue", bg="#ccffcc").grid(
                         row=_cur_row, column=0, columnspan=9, pady=(14,4), sticky="ew")
                _cur_row += 1
                _hdrs_a = ["Sondaje 1(Agua)","Litros 1","Sondaje 2(Agua)","Litros 2",
                           "Lectura Agua","Desc.Tubo","Sond.Real","Litros Agua Nat."]
                _keys_a = ["agua_s1","agua_l1","agua_s2","agua_l2",
                           "agua_lectura","agua_desc","agua_s_real","vol_nat_agua"]
                for i2, h2 in enumerate(_hdrs_a):
                    tk.Label(frame_main, text=h2, font=fb, bg="#ccffcc").grid(
                        row=_cur_row, column=i2, sticky="ew", padx=1)
                for i2, k2 in enumerate(_keys_a):
                    st_a = "readonly" if k2 in ("agua_s_real","vol_nat_agua") else "normal"
                    va = self.get_var(f"{etapa}_{tk_name}_{k2}")
                    ea = tk.Entry(frame_main, textvariable=va,
                                  justify="center", state=st_a, bg="#eeffee")
                    ea.grid(row=_cur_row+1, column=i2, sticky="ew", padx=1, pady=(0,8))
                    if k2 in ("agua_lectura","agua_desc"):
                        ea.bind("<KeyRelease>", lambda ev,t=tk_name,ep=etapa: self.calc_sondaje_agua(ep,t))
                    if k2 in ("agua_s1","agua_l1","agua_s2","agua_l2"):
                        ea.bind("<KeyRelease>", lambda ev,t=tk_name,ep=etapa: self.calc_volumen_agua_ui(ep,t))
                _cur_row += 2

                # ── Interp × asiento (AGUA) embebida — buques líquidos ─────
                if _es_maritimo_liq:
                    _cur_row += self.crear_interp_trim_inline(etapa, tk_name, frame_main, _cur_row, agua=True)

            # ══════════════════════════════════════════════════════════════
            # DENSIDADES / VOLÚMENES (solo para mediciones de producto líquido)
            # ══════════════════════════════════════════════════════════════
            _show_dens = not (_el or _duc)
            if _show_dens:
                self.crear_fila_popup(etapa, tk_name, frame_main, _cur_row, [
                    ("Densidad Lab",          "dens_lab",  "entry",    2),
                    ("Litros a 15 Lab",       "v15_lab",   "entry_ro", 2),
                    ("Kilos Vacio Lab",       "kv_lab",    "entry_ro", 2),
                    ("Kilos al Aire Lab",     "ka_lab",    "entry_ro", 3),
                ], bg_row="#ffffcc")
                self.crear_fila_popup(etapa, tk_name, frame_main, _cur_row+2, [
                    ("Densidad Doc",          "dens_doc",  "entry_ro", 2),
                    ("Litros a 15 Doc",       "v15_doc",   "entry_ro", 2),
                    ("Kilos Vacio Doc",       "kv_doc",    "entry_ro", 2),
                    ("Kilos al Aire Doc",     "ka_doc",    "entry_ro", 3),
                ], bg_row="#dcebf7")
                self.crear_fila_popup(etapa, tk_name, frame_main, _cur_row+4, [
                    ("Densidad Salida",       "dens_salida","entry_ro", 2),
                    ("Litros a 15 Salida",    "v15_sal",   "entry_ro", 2),
                    ("Kilos Vacio Salida",    "kv_sal",    "entry_ro", 2),
                    ("Kilos al Aire Salida",  "ka_sal",    "entry_ro", 3),
                ], bg_row="#dcf7dc")

            # ══════════════════════════════════════════════════════════════
            # Botón Guardar y Cerrar
            # ══════════════════════════════════════════════════════════════
            tk.Button(f_bot_fixed, text="  Guardar y Cerrar  ", bg="#27AE60", fg="white",
                      font=fb, command=popup.destroy,
                      relief="flat", cursor="hand2"
                      ).pack(side="left", padx=16, pady=8, ipadx=10, ipady=4)
            if not (_duc or _el):
                tk.Button(f_bot_fixed, text="  Tabla Calibrado  ", bg="#8E44AD", fg="white",
                          font=fb, relief="flat", cursor="hand2",
                          command=lambda: self.abrir_tabla_calibrado(etapa, tk_name, es_buque=_mar)
                          ).pack(side="left", padx=8, pady=8, ipadx=8, ipady=4)
            if _mar:
                tk.Button(f_bot_fixed, text="  Interp × Asiento  ", bg="#1A5276", fg="white",
                          font=fb, relief="flat", cursor="hand2",
                          command=lambda: self.abrir_interp_trim_rapida(etapa, tk_name)
                          ).pack(side="left", padx=8, pady=8, ipadx=8, ipady=4)
                if _tm in ("BUQUE", "BARCAZA", "BUQUE QUIMIQUERO", "DRAFT SURVEY"):
                    tk.Button(f_bot_fixed, text="  Interp Agua  ", bg="#117864", fg="white",
                              font=fb, relief="flat", cursor="hand2",
                              command=lambda: self.abrir_interp_trim_rapida(etapa, tk_name, agua=True)
                              ).pack(side="left", padx=8, pady=8, ipadx=8, ipady=4)
            tk.Label(f_bot_fixed, text=f"{tk_name}  |  {etapa.upper()}",
                     bg="#2C3E50", fg="#AED6F1", font=("Arial",9)).pack(side="right", padx=12)

            popup.update_idletasks()

        except Exception as e:
            import traceback
            traceback.print_exc()
            import tkinter.messagebox as messagebox
            messagebox.showerror("Error", f"Error al abrir ficha:\n{e}")

    def crear_fila_popup(self, etapa, tk_name, parent, start_row, fields, bg_row="#eee"):
        col_cursor = 0
        font_bold = ("Arial", 8, "bold")
        for item in fields:
            lbl, key, w_type, span = item
            tk.Label(parent, text=lbl, font=font_bold, bg=bg_row, anchor="w").grid(row=start_row, column=col_cursor, columnspan=span, sticky="ew", padx=1, pady=(5,0))
            col_cursor += span
        col_cursor = 0
        for item in fields:
            lbl, key, w_type, span = item
            var_name = f"{etapa}_{tk_name}_{key}"
            widget = None
            if w_type == "combo":
                # 1. Obtenemos la lista de documentos que tienen número
                lista_docs = [d["numero"].get() for d in self.ddt_data if d["numero"].get()]
                
                # 2. AGREGA ESTA LINEA: Insertamos una opción vacía al principio
                values_docs = [""] + lista_docs 
                
                widget = ttk.Combobox(parent, textvariable=self.get_var(var_name), values=values_docs, state="readonly")
                widget.bind("<<ComboboxSelected>>", lambda event, t=tk_name, e=etapa: self.on_ddt_selected(e, t))
                        
            elif w_type == "combo_vcf":
                vcf_v = self.get_var(var_name)
                if not vcf_v.get(): vcf_v.set("54B (Combustibles)")
                widget = ttk.Combobox(parent, textvariable=vcf_v, values=[
                "54B (Combustibles)", "54A (Crudos)", "54D (Lubricantes)",
                "Químico (Lineal)", "GLP/Propano (tabla 54E)", "GNL/Criogénico",
                "Amoniaco/Refrigerante", "Sin corrección (VCF=1)"
            ], state="readonly")
                widget.bind("<<ComboboxSelected>>", lambda event, t=tk_name, e=etapa: self.calc_volumen_prod_ui(e, t))
            else:
                st = "readonly" if "ro" in w_type else "normal"
                widget = tk.Entry(parent, textvariable=self.get_var(var_name), justify="center", state=st)
                if key in ["alt_uti", "desc_tubo", "alt_ref"]:
                    widget.bind("<KeyRelease>", lambda event, t=tk_name, e=etapa: self.calc_sondaje_prod(e, t))
                elif "prod_" in key or key == "temp" or "dens_lab" in key:
                    widget.bind("<KeyRelease>", lambda event, t=tk_name, e=etapa: self.calc_volumen_prod_ui(e, t))
            if widget:
                widget.grid(row=start_row+1, column=col_cursor, columnspan=span, sticky="ew", padx=1, pady=(0,5))
            col_cursor += span

    # --- CALCULO LOGICA ---
    def on_ddt_selected(self, etapa, tanque):
        if self.is_loading_data: return
        ddt_num = self.get_var(f"{etapa}_{tanque}_ddt_assign").get()
        if not ddt_num:
            self.get_var(f"{etapa}_{tanque}_prod_name").set("")
            self.get_var(f"{etapa}_{tanque}_dens_doc").set("")
            self.get_var(f"{etapa}_{tanque}_dens_salida").set("")
            self.calc_volumen_prod_ui(etapa, tanque)
            return
        found = next((d for d in self.ddt_data if d["numero"].get() == ddt_num), None)
        if found:
            prod_val = found["producto"].get()
            dens_doc = found["densidad"].get()
            self.get_var(f"{etapa}_{tanque}_prod_name").set(prod_val)
            self.get_var(f"{etapa}_{tanque}_dens_doc").set(dens_doc)
            if not self.get_var(f"{etapa}_{tanque}_dens_lab").get():
                self.get_var(f"{etapa}_{tanque}_dens_lab").set(dens_doc)
            salidas = found["salidas"]
            avg_sal = 0.0
            if salidas:
                total_mass = 0.0
                total_liters = 0.0
                for s in salidas:
                    l = self.parse_float(s["litros"].get())
                    d = self.parse_float(s["densidad"].get())
                    total_liters += l
                    total_mass += (l * d)
                if total_liters > 0:
                    avg_sal = total_mass / total_liters
                else:
                    avg_sal = 0.0
            if avg_sal > 0:
                self.get_var(f"{etapa}_{tanque}_dens_salida").set(f"{avg_sal:.5f}")
            else:
                self.get_var(f"{etapa}_{tanque}_dens_salida").set(dens_doc)
            self.calc_volumen_prod_ui(etapa, tanque)

    def calc_trim(self, etapa):
        if self.is_loading_data: return
        try:
            p1 = self.parse_float(self.get_var(f"{etapa}_Calados Popa").get())
            p2 = self.parse_float(self.get_var(f"{etapa}_Calados Proa").get())
            self.get_var(f"{etapa}_Trimación").set(f"{p1-p2:.2f}")
            b = self.parse_float(self.get_var(f"{etapa}_Calados Babor").get())
            e = self.parse_float(self.get_var(f"{etapa}_Calados Estribor").get())
            self.get_var(f"{etapa}_Lista").set(f"{b-e:.2f}")
        except: pass

    def calc_sondaje_prod(self, etapa, tanque):
        if self.is_loading_data: return
        try:
            uti = self.parse_float(self.get_var(f"{etapa}_{tanque}_alt_uti").get())
            desc = self.parse_float(self.get_var(f"{etapa}_{tanque}_desc_tubo").get())
            res = uti - desc
            self.get_var(f"{etapa}_{tanque}_s_corr").set(f"{res:.0f}")
            self.calc_volumen_prod_ui(etapa, tanque)
        except: pass

    def calc_sondaje_agua(self, etapa, tanque):
        if self.is_loading_data: return
        try:
            lec = self.parse_float(self.get_var(f"{etapa}_{tanque}_agua_lectura").get())
            desc = self.parse_float(self.get_var(f"{etapa}_{tanque}_agua_desc").get())
            res = lec - desc
            self.get_var(f"{etapa}_{tanque}_agua_s_real").set(f"{res:.0f}")
            self.calc_volumen_agua_ui(etapa, tanque)
        except: pass

    def calc_volumen_prod_ui(self, etapa, tanque):
        if self.is_loading_data: return
        import json
        try:
            s = self.parse_float(self.get_var(f"{etapa}_{tanque}_s_corr").get())
            # ── Tabla de calibrado × asiento (buque/barcaza) — prioridad ────
            trim_json = self.get_var(f"{etapa}_{tanque}_tabla_trim_json").get()
            handled_trim = False
            if trim_json:
                obj = json.loads(trim_json)
                trim_m = self.parse_float(self.get_var(f"{etapa}_Trimación").get() or "0")
                vol_t, det_t = self._interp_trim_table(obj, s, trim_m)
                if vol_t is not None:
                    self.get_var(f"{etapa}_{tanque}_vol_nat_prod").set(f"{vol_t:.0f}")
                    val = vol_t
                    handled_trim = True
            # ── Tabla de calibrado multi-punto (si existe) ──────────────────
            tabla_json = "" if handled_trim else self.get_var(f"{etapa}_{tanque}_tabla_cal_json").get()
            if tabla_json:
                pts = json.loads(tabla_json)
                if len(pts) >= 2:
                    pts.sort(key=lambda x: x[0])
                    sonds = [p[0] for p in pts]; lits = [p[1] for p in pts]
                    trims = [p[2] if len(p) > 2 else 0.0 for p in pts]
                    if s <= sonds[0]: val = lits[0]; trim_corr = trims[0]
                    elif s >= sonds[-1]: val = lits[-1]; trim_corr = trims[-1]
                    else:
                        val = lits[0]; trim_corr = 0.0
                        for k in range(len(pts)-1):
                            if sonds[k] <= s <= sonds[k+1]:
                                ds = sonds[k+1] - sonds[k]
                                if ds != 0:
                                    frac = (s - sonds[k]) / ds
                                    val = lits[k] + frac * (lits[k+1] - lits[k])
                                    trim_corr = trims[k] + frac * (trims[k+1] - trims[k])
                                else:
                                    val = lits[k]; trim_corr = trims[k]
                                break
                    # Aplicar corrección de trim si hay datos de trimación
                    try:
                        trim_m = self.parse_float(self.get_var(f"{etapa}_Trimación").get() or "0")
                        if trim_corr != 0.0 and trim_m != 0.0:
                            val += trim_corr * trim_m
                    except: pass
                    self.get_var(f"{etapa}_{tanque}_vol_nat_prod").set(f"{val:.0f}")
            elif not handled_trim:
                # ── Interpolación 2 puntos (modo original) ──────────────────
                # Ignorar marcadores tipo "[5pts]" que versiones previas
                # escribían en prod_s1 (parse_float los convertía en número).
                def _pfm(k):
                    v = self.get_var(k).get()
                    return 0.0 if "[" in v else self.parse_float(v)
                s1 = _pfm(f"{etapa}_{tanque}_prod_s1")
                l1 = _pfm(f"{etapa}_{tanque}_prod_l1")
                s2 = _pfm(f"{etapa}_{tanque}_prod_s2")
                l2 = _pfm(f"{etapa}_{tanque}_prod_l2")
                if s2 == s1: val = l1
                else: val = l1 + ((s - s1) / (s2 - s1)) * (l2 - l1)
                self.get_var(f"{etapa}_{tanque}_vol_nat_prod").set(f"{val:.0f}")
            temp_str = self.get_var(f"{etapa}_{tanque}_temp").get()
            if not temp_str: return
            temp = self.parse_float(temp_str)
            table_type = self.get_var(f"{etapa}_{tanque}_tabla_vcf").get()
            if not table_type: table_type = "54B (Combustibles)"
            # Litros a 15° y kilos sobre el volumen NETO (bruto − agua de fondo),
            # igual que la planilla PDF (generar_un_reporte) y el cálculo de cargos.
            agua_lts = self.parse_float(self.get_var(f"{etapa}_{tanque}_vol_nat_agua").get() or "0")
            val_neto = val - agua_lts
            def calc_set(dens_key_suffix, out_v15, out_kv, out_ka):
                d_str = self.get_var(f"{etapa}_{tanque}_{dens_key_suffix}").get()
                if d_str:
                    dens = self.parse_float(d_str)
                    vcf = self.calc_vcf(dens, temp, table_type)
                    vol_15 = val_neto * vcf
                    d_calc = dens
                    if d_calc > 2.0: d_calc = d_calc / 1000.0
                    kg_vacio = vol_15 * d_calc
                    kg_aire = vol_15 * (d_calc - 0.0011)
                    self.get_var(f"{etapa}_{tanque}_{out_v15}").set(f"{vol_15:.0f}")
                    self.get_var(f"{etapa}_{tanque}_{out_kv}").set(f"{kg_vacio:.0f}")
                    self.get_var(f"{etapa}_{tanque}_{out_ka}").set(f"{kg_aire:.0f}")
                else:
                    self.get_var(f"{etapa}_{tanque}_{out_v15}").set("---")
                    self.get_var(f"{etapa}_{tanque}_{out_kv}").set("---")
                    self.get_var(f"{etapa}_{tanque}_{out_ka}").set("---")
            calc_set("dens_lab", "v15_lab", "kv_lab", "ka_lab")
            calc_set("dens_doc", "v15_doc", "kv_doc", "ka_doc")
            calc_set("dens_salida", "v15_sal", "kv_sal", "ka_sal")
        except: pass 

    def calc_volumen_agua_ui(self, etapa, tanque):
        if self.is_loading_data: return
        import json
        try:
            s = self.parse_float(self.get_var(f"{etapa}_{tanque}_agua_s_real").get())
            # ── Tabla × asiento para agua (buque/barcaza) — prioridad ──────
            trim_json = self.get_var(f"{etapa}_{tanque}_tabla_trim_agua_json").get()
            handled = False
            if trim_json:
                obj = json.loads(trim_json)
                trim_m = self.parse_float(self.get_var(f"{etapa}_Trimación").get() or "0")
                vol_t, _det = self._interp_trim_table(obj, s, trim_m)
                if vol_t is not None:
                    self.get_var(f"{etapa}_{tanque}_vol_nat_agua").set(f"{vol_t:.0f}")
                    handled = True
            if not handled:
                s1 = self.parse_float(self.get_var(f"{etapa}_{tanque}_agua_s1").get())
                l1 = self.parse_float(self.get_var(f"{etapa}_{tanque}_agua_l1").get())
                s2 = self.parse_float(self.get_var(f"{etapa}_{tanque}_agua_s2").get())
                l2 = self.parse_float(self.get_var(f"{etapa}_{tanque}_agua_l2").get())
                if s2 == s1: val = l1
                else: val = l1 + ((s - s1) / (s2 - s1)) * (l2 - l1)
                self.get_var(f"{etapa}_{tanque}_vol_nat_agua").set(f"{val:.0f}")
        except: self.get_var(f"{etapa}_{tanque}_vol_nat_agua").set("0")
        # El agua afecta v15/kv/ka del producto → refrescar
        self.calc_volumen_prod_ui(etapa, tanque)

    def interpolar_prod(self, etapa, tanque):
        self.calc_volumen_prod_ui(etapa, tanque)
        try: return self.parse_float(self.get_var(f"{etapa}_{tanque}_vol_nat_prod").get())
        except: return 0.0

    def interpolar_agua(self, etapa, tanque):
        self.calc_volumen_agua_ui(etapa, tanque)
        try: return self.parse_float(self.get_var(f"{etapa}_{tanque}_vol_nat_agua").get())
        except: return 0.0

    def get_interpolation_details(self, etapa, tanque):
        import json
        try:
            s = self.get_var(f"{etapa}_{tanque}_s_corr").get()
            # Tabla × asiento (trim): detalle real de la interpolación 2D
            tj = self.get_var(f"{etapa}_{tanque}_tabla_trim_json").get()
            if tj:
                trim_m = self.parse_float(self.get_var(f"{etapa}_Trimación").get() or "0")
                vol, det = self._interp_trim_table(json.loads(tj), self.parse_float(s), trim_m)
                if vol is not None: return det
            # Tabla de calibrado multi-punto
            cj = self.get_var(f"{etapa}_{tanque}_tabla_cal_json").get()
            if cj:
                pts = json.loads(cj)
                return f"Tabla de calibrado ({len(pts)} pts, {pts[0][0]:.0f}→{pts[-1][0]:.0f} mm), sondaje {s} mm"
            s1 = self.get_var(f"{etapa}_{tanque}_prod_s1").get()
            l1 = self.get_var(f"{etapa}_{tanque}_prod_l1").get()
            s2 = self.get_var(f"{etapa}_{tanque}_prod_s2").get()
            l2 = self.get_var(f"{etapa}_{tanque}_prod_l2").get()
            return f"{l1} + (({s} - {s1}) / ({s2} - {s1})) * ({l2} - {l1})"
        except: return "Error en datos"

    def get_water_interp_details(self, etapa, tanque):
        import json
        try:
            s = self.get_var(f"{etapa}_{tanque}_agua_s_real").get()
            tj = self.get_var(f"{etapa}_{tanque}_tabla_trim_agua_json").get()
            if tj:
                trim_m = self.parse_float(self.get_var(f"{etapa}_Trimación").get() or "0")
                vol, det = self._interp_trim_table(json.loads(tj), self.parse_float(s), trim_m)
                if vol is not None: return det
            s1 = self.get_var(f"{etapa}_{tanque}_agua_s1").get()
            l1 = self.get_var(f"{etapa}_{tanque}_agua_l1").get()
            s2 = self.get_var(f"{etapa}_{tanque}_agua_s2").get()
            l2 = self.get_var(f"{etapa}_{tanque}_agua_l2").get()
            return f"{l1} + (({s} - {s1}) / ({s2} - {s1})) * ({l2} - {l1})"
        except: return ""

    def get_vcf_details(self, dens, temp, table):
        """Detalle de cálculo VCF para incluir en reportes PDF."""
        try:
            rho = float(dens)
            t   = float(temp)
            if rho < 2.0: rho = rho * 1000.0
            if rho <= 0: return "S/D"

            norma = getattr(self, "norma_astm", None)
            usar_1980 = (norma is None) or (norma.get() == "1980")
            norma_str = "ASTM D1250-80" if usar_1980 else "API MPMS 11.1-2004"

            if "54A" in table:
                k0, k1, alpha_ov = 613.9723, 0.0, None
            elif "54D" in table:
                k0, k1, alpha_ov = 1489.0672, 0.0, None
            elif "54B" in table or True:
                if usar_1980:
                    k0, k1, alpha_ov = self._vcf_k0k1_1980_54B(rho)
                    if alpha_ov is not None:
                        # Zona transición
                        alpha = alpha_ov
                        zona = "Zona transición (770<ρ<778)"
                        coef_str = f"α directo={alpha:.7f}"
                    else:
                        alpha = (k0/rho**2) + (k1/rho)
                        if rho <= 770:   zona = "Zona 1 (ρ≤770)"
                        elif rho < 839:  zona = "Zona 3 (778≤ρ<839)"
                        else:            zona = "Zona 4 (ρ≥839)"
                        coef_str = f"K0={k0}, K1={k1}"
                else:
                    k0, k1, alpha_ov = 346.4228, 0.4033, None
                    alpha = (k0/rho**2) + (k1/rho)
                    zona = "—"
                    coef_str = f"K0={k0}, K1={k1}"

            if alpha_ov is None:
                alpha = (k0/rho**2) + (k1/rho) if "54B" not in table or not usar_1980 else alpha
                coef_str = coef_str if "coef_str" in dir() else f"K0={k0}, K1={k1}"
                zona = zona if "zona" in dir() else "—"

            dt  = t - 15.0
            vcf = round(math.exp(-alpha * dt * (1.0 + 0.8 * alpha * dt)), 5)
            return f"{norma_str} | {coef_str} | Alpha={alpha:.7f}, ΔT={dt:.1f}°C, VCF={vcf:.5f}"
        except:
            return "S/D"


    # ─────────────────────────────────────────────────────────────────────────
    # CÁLCULOS ESPECÍFICOS: GAS, DUCTOS, ELECTRICIDAD
    # ─────────────────────────────────────────────────────────────────────────

    COMPONENTES_GAS = [
        ("CH4",  16.043, 190.56, 4599.0),
        ("C2H6", 30.069, 305.32, 4872.0),
        ("C3H8", 44.096, 369.83, 4248.0),
        ("iC4",  58.122, 407.82, 3640.0),
        ("nC4",  58.122, 425.12, 3796.0),
        ("iC5",  72.149, 460.35, 3381.0),
        ("nC5",  72.149, 469.70, 3370.0),
        ("C6+",  86.175, 507.60, 3025.0),
        ("N2",   28.014, 126.19, 3396.0),
        ("CO2",  44.010, 304.13, 7375.0),
        ("H2S",  34.082, 373.10, 8937.0),
    ]

    def calc_gas_propiedades(self, composicion_pct):
        """MW, Tc(K), Pc(kPa) de mezcla por Kay's rule."""
        try:
            total = sum(composicion_pct.values())
            if total == 0: return {"MW":16.0,"Tc":190.56,"Pc":4599.0}
            MW=Tc=Pc=0.0
            for nombre, mw, tc, pc in self.COMPONENTES_GAS:
                yi = composicion_pct.get(nombre,0.0)/total
                MW+=yi*mw; Tc+=yi*tc; Pc+=yi*pc
            return {"MW":MW,"Tc":Tc,"Pc":Pc}
        except: return {"MW":16.0,"Tc":190.56,"Pc":4599.0}

    def calc_factor_z(self, P_kPa, T_C, Tc_K=190.56, Pc_kPa=4599.0, composicion_pct=None):
        """Factor Z por ecuacion de Papay (AGA simplificado, prec. +-0.3%)."""
        try:
            if composicion_pct:
                props = self.calc_gas_propiedades(composicion_pct)
                Tc_K=props["Tc"]; Pc_kPa=props["Pc"]
            Tr = (T_C+273.15)/Tc_K
            Pr = P_kPa/Pc_kPa
            Z = 1-(3.52*Pr)/(10**(0.9813*Tr))+(0.274*Pr**2)/(10**(0.8157*Tr))
            return max(0.50,min(1.05,Z))
        except: return 1.0

    def calc_volumen_base_gas(self, vol_m3, P_kPa, T_C, Z, P_base=101.325, T_base_C=15.0):
        """V condiciones de línea → V condiciones base (IRAM/AGA)."""
        try:
            return vol_m3*(P_kPa/P_base)*((T_base_C+273.15)/(T_C+273.15))*(1.0/Z)
        except: return 0.0

    def calc_pig_volumen_m3(self, diametro_pulg, longitud_m):
        """Volumen nominal del tramo (calibración pig)."""
        try:
            D=float(diametro_pulg)*0.0254; L=float(longitud_m)
            return math.pi/4*D**2*L
        except: return 0.0

    def calc_energia_electrica(self, lect_ini, lect_fin, constante=1.0):
        """kWh = (fin - ini) * constante de transformacion."""
        try: return (float(lect_fin)-float(lect_ini))*float(constante)
        except: return 0.0

    def calc_fp_kwh(self, kwh_act, kwh_react):
        """cos φ desde kWh activa/reactiva."""
        try:
            a,r=float(kwh_act),float(kwh_react)
            return a/math.sqrt(a**2+r**2) if (a**2+r**2)>0 else 1.0
        except: return 1.0

    def calc_volumen_camion_gas(self, masa_kg, dens_kgm3):
        """Vol líquido GLP/GNC en cisterna a presión (L)."""
        try: return (float(masa_kg)/float(dens_kgm3))*1000.0
        except: return 0.0

    def calc_masa_gas_cil_horizontal(self, varilla_mm, radio_m, largo_m, densidad_kgm3):
        """Masa de líquido GLP en cisterna horizontal (por varilla)."""
        try:
            vol_l = self.calc_volumen_cilindro_horizontal(varilla_mm, radio_m, largo_m)
            return (vol_l / 1000.0) * float(densidad_kgm3)
        except: return 0.0

    def calc_gas_fase(self, presion_kpa, temp_c, producto="GLP"):
        """Determina fase (líquido/gas/bifásico) según P/T.
        GLP propano: P_sat a 20°C ≈ 836 kPa, a -42°C ≈ 101 kPa.
        GLP butano:  P_sat a 20°C ≈ 210 kPa, a -0.5°C ≈ 101 kPa.
        GNC: siempre gas (alta presión)."""
        try:
            if "GNC" in producto.upper() or "GNV" in producto.upper():
                return "GAS COMPRIMIDO"
            # Approximación lineal de P_sat para propano/butano mixtura
            # P_sat(T) ≈ 101.325 * exp(a + b/T_K) — usamos tabla simplificada
            t_k = temp_c + 273.15
            # Propano: log(P_sat/kPa) = 6.803 - 803.8/T_K
            psat_prop = 10 ** (6.803 - 803.8/t_k)
            if presion_kpa > psat_prop * 1.05: return "LÍQUIDO"
            elif presion_kpa < psat_prop * 0.95: return "GAS"
            else: return "BIFÁSICO"
        except: return "DESCONOCIDO"

    # ═══════════════════════════════════════════════════════════════════════════
    # TRIBUTOS DINÁMICOS
    # ═══════════════════════════════════════════════════════════════════════════

    # Tributos disponibles: (nombre, alicuota_default, activo_default, es_configurable)
    TRIBUTOS_CATALOGO = [
        ("Derechos de Importación",        8.0,  True,  True),
        ("Derechos de Exportación",        8.0,  True,  True),
        ("IVA",                           21.0,  False, True),
        ("IVA reducido",                  10.5,  False, True),
        ("Estadística",                    0.5,  True,  True),
        ("Anticipo Ganancias",             0.5,  True,  True),
        ("Tasa Comprobación Destino (TCD)",0.5,  False, True),
        ("Antidumping / Compensatorio",    0.0,  False, True),
        ("Ingresos Brutos (prov.)",        3.0,  False, True),
        ("Tasa Aduanera de Servicios",     0.5,  False, True),
    ]

    def get_tributos_activos(self):
        """Devuelve lista de (nombre, alicuota) de tributos activos."""
        if not hasattr(self, "_tributos_sesion") or not self._tributos_sesion:
            # Default: Derechos + Estadística + Ganancias
            self._tributos_sesion = [
                {"nombre": t[0], "alicuota": t[1], "activo": t[2]}
                for t in self.TRIBUTOS_CATALOGO
            ]
        return [(t["nombre"], t["alicuota"]) for t in self._tributos_sesion if t["activo"]]

    def dialogo_tributos(self, parent=None):
        """Abre diálogo interactivo para seleccionar y configurar tributos."""
        if not hasattr(self, "_tributos_sesion") or not self._tributos_sesion:
            self._tributos_sesion = [
                {"nombre": t[0], "alicuota": t[1], "activo": t[2]}
                for t in self.TRIBUTOS_CATALOGO
            ]
        resultado = {"ok": False}
        dlg = tk.Toplevel(parent or self.root)
        dlg.title("Configurar Tributos para Cargo/Denuncia")
        dlg.grab_set()
        dlg.transient(parent or self.root)
        dlg.update_idletasks()
        sw, sh = dlg.winfo_screenwidth(), dlg.winfo_screenheight()
        dlg.geometry(f"620x480+{(sw-620)//2}+{(sh-480)//2}")

        tk.Label(dlg, text="TRIBUTOS A INCLUIR EN EL CÁLCULO", bg="#1B3A5C", fg="white",
                 font=("Arial", 9, "bold")).pack(fill="x", ipady=7)
        tk.Label(dlg, text="Seleccioná los tributos y ajustá los porcentajes según el caso.",
                 font=("Arial", 9), fg="#555").pack(pady=(6,2))

        # Headers
        fhdr = tk.Frame(dlg, bg="#DEE2E6")
        fhdr.pack(fill="x", padx=15, pady=(4,0))
        for txt, w in [("Incluir", 6), ("Concepto", 34), ("Alícuota (%)", 14), ("", 6)]:
            tk.Label(fhdr, text=txt, bg="#DEE2E6", font=("Arial", 8, "bold"),
                     width=w, anchor="w").pack(side="left", padx=3, pady=3)

        # Scrollable rows
        cf = tk.Frame(dlg); cf.pack(fill="both", expand=True, padx=15, pady=2)
        cv_t = tk.Canvas(cf, bg="white"); sb_t = ttk.Scrollbar(cf, orient="v", command=cv_t.yview)
        sf_t = tk.Frame(cv_t, bg="white")
        cv_t.create_window((0,0), window=sf_t, anchor="nw")
        sf_t.bind("<Configure>", lambda e: cv_t.configure(scrollregion=cv_t.bbox("all")))
        cv_t.configure(yscrollcommand=sb_t.set)
        cv_t.pack(side="left", fill="both", expand=True)
        sb_t.pack(side="right", fill="y")

        row_vars = []  # (activo_var, nombre_var, alicuota_var)
        for trib in self._tributos_sesion:
            v_act  = tk.BooleanVar(value=trib["activo"])
            v_nom  = tk.StringVar(value=trib["nombre"])
            v_alic = tk.StringVar(value=str(trib["alicuota"]))
            row_vars.append((v_act, v_nom, v_alic))
            row = tk.Frame(sf_t, bg="white")
            row.pack(fill="x", pady=1)
            tk.Checkbutton(row, variable=v_act, bg="white").pack(side="left", padx=4)
            tk.Label(row, textvariable=v_nom, bg="white", font=("Arial", 9),
                     width=32, anchor="w").pack(side="left")
            e_alic = tk.Entry(row, textvariable=v_alic, width=10, font=("Arial", 9),
                              justify="center")
            e_alic.pack(side="left", padx=6)
            tk.Label(row, text="%", bg="white", font=("Arial", 9)).pack(side="left")

        # Custom tributo row
        sep = ttk.Separator(dlg, orient="horizontal")
        sep.pack(fill="x", padx=15, pady=4)
        f_custom = tk.Frame(dlg); f_custom.pack(fill="x", padx=15)
        tk.Label(f_custom, text="+ Agregar concepto:", font=("Arial", 8, "bold")).pack(side="left")
        v_custom_nom  = tk.StringVar()
        v_custom_alic = tk.StringVar(value="0.0")
        tk.Entry(f_custom, textvariable=v_custom_nom, width=28,
                 font=("Arial", 9)).pack(side="left", padx=6)
        tk.Entry(f_custom, textvariable=v_custom_alic, width=8,
                 font=("Arial", 9)).pack(side="left")
        tk.Label(f_custom, text="%", font=("Arial", 9)).pack(side="left", padx=2)

        def _add_custom():
            nom = v_custom_nom.get().strip()
            if not nom: return
            try: alic = float(v_custom_alic.get())
            except: alic = 0.0
            new_t = {"nombre": nom, "alicuota": alic, "activo": True}
            self._tributos_sesion.append(new_t)
            v_act  = tk.BooleanVar(value=True)
            v_nom  = tk.StringVar(value=nom)
            v_alic = tk.StringVar(value=str(alic))
            row_vars.append((v_act, v_nom, v_alic))
            row = tk.Frame(sf_t, bg="white"); row.pack(fill="x", pady=1)
            tk.Checkbutton(row, variable=v_act, bg="white").pack(side="left", padx=4)
            tk.Label(row, textvariable=v_nom, bg="white", font=("Arial",9), width=32, anchor="w").pack(side="left")
            tk.Entry(row, textvariable=v_alic, width=10, font=("Arial",9), justify="center").pack(side="left", padx=6)
            tk.Label(row, text="%", bg="white", font=("Arial",9)).pack(side="left")
            v_custom_nom.set(""); v_custom_alic.set("0.0")

        tk.Button(f_custom, text="Agregar", bg="#2196F3", fg="white",
                  font=("Arial", 8, "bold"), command=_add_custom).pack(side="left", padx=8)

        def _confirmar():
            for i, (v_act, v_nom, v_alic) in enumerate(row_vars):
                if i < len(self._tributos_sesion):
                    self._tributos_sesion[i]["activo"]   = v_act.get()
                    self._tributos_sesion[i]["alicuota"] = float(v_alic.get() or "0")
            resultado["ok"] = True
            dlg.destroy()

        fbot_t = tk.Frame(dlg); fbot_t.pack(fill="x", padx=15, pady=8)
        tk.Button(fbot_t, text="Confirmar y Continuar", bg="#27AE60", fg="white",
                  font=("Arial", 9, "bold"), command=_confirmar).pack(side="right", padx=6)
        tk.Button(fbot_t, text="Cancelar", bg="#E74C3C", fg="white",
                  font=("Arial", 9), command=dlg.destroy).pack(side="right", padx=4)

        dlg.wait_window()
        return resultado["ok"]

    def _get_cromatografia(self, prefix):
        """Devuelve dict {nombre: %mol} desde variables de cromatografia."""
        comp_names = ["CH4","C2H6","C3H8","iC4","nC4","iC5","nC5","C6+","N2","CO2","H2S"]
        result = {}
        for cn in comp_names:
            val = self.parse_float(self.get_var(f"{prefix}_gc_{cn}").get() or "0")
            if val > 0: result[cn] = val
        return result if result else None

    def calc_potencia_aparente(self, V, I):
        try: return float(V)*float(I)
        except: return 0.0


    # ═══════════════════════════════════════════════════════════════════════════
    # DRAFT SURVEY — ESTIMACIÓN DE PESO POR CALADOS
    # Métodos de cálculo + ventana completa
    # ═══════════════════════════════════════════════════════════════════════════

    def calc_draft_trim(self, proa_b, proa_e, popa_b, popa_e, medio_b, medio_e):
        """Calados medios en proa, medio y popa."""
        proa  = (proa_b  + proa_e)  / 2
        popa  = (popa_b  + popa_e)  / 2
        medio = (medio_b + medio_e) / 2
        trim  = popa - proa  # positivo = trimado a popa (stern trim)
        return proa, popa, medio, trim

    def calc_draft_corr_trim(self, proa, popa, medio, lbp):
        """Corrección de calado medio por trimado (método estándar Lloyd's/RINA).
        Dc = (Dproa + 6*Dmedio + Dpopa) / 8  → calado hidrostático (mean of means)
        Corrección por diferencia de centro de carena: δ = trim * (LCF - LBP/2) / LBP
        """
        try:
            d_medio_corr = (proa + 6*medio + popa) / 8
            return d_medio_corr
        except: return (proa + 6*medio + popa) / 8

    def calc_draft_hog_sag(self, proa, popa, medio):
        """Hog (camello) o Sag (pandeo).
        Positivo = Hog (buque arqueado hacia arriba, medio más alto que extremos).
        Negativo = Sag (buque pandeado, medio más bajo)."""
        teorico_medio = (proa + popa) / 2
        diff = medio - teorico_medio
        if diff > 0.002:   return ("HOG", round(diff*1000, 1))
        elif diff < -0.002: return ("SAG", round(abs(diff)*1000, 1))
        else:               return ("RECTO", 0.0)

    def calc_desplazamiento_interpolado(self, calado_m, tabla_hidro):
        """Interpola desplazamiento desde tabla hidrostática.
        tabla_hidro: lista de (calado_m, desplaz_t, tpc, lcf_m) ordenada por calado."""
        if not tabla_hidro: return 0.0, 0.0, 0.0
        tabla = sorted(tabla_hidro, key=lambda x: x[0])
        if calado_m <= tabla[0][0]:  return tabla[0][1], tabla[0][2], tabla[0][3]
        if calado_m >= tabla[-1][0]: return tabla[-1][1], tabla[-1][2], tabla[-1][3]
        for i in range(len(tabla)-1):
            c0,d0,t0,l0 = tabla[i]
            c1,d1,t1,l1 = tabla[i+1]
            if c0 <= calado_m <= c1:
                f = (calado_m - c0) / (c1 - c0)
                return d0+f*(d1-d0), t0+f*(t1-t0), l0+f*(l1-l0)
        return 0.0, 0.0, 0.0

    def calc_draft_desp_trimado(self, desp_medio, trim_m, tpc, lcf, lbp):
        """Corrección del desplazamiento por trimado (1ª corrección de trim).
        δD = TPC * trim * (LCF - LBP/2) * 100 / LBP
        """
        try:
            if lbp <= 0: return desp_medio
            corr = tpc * trim_m * (lcf - lbp/2) * 100.0 / lbp
            return desp_medio + corr
        except: return desp_medio

    def calc_draft_peso_carga(self, desp_llegada, desp_salida, constante_buque,
                              peso_combustible_ini, peso_combustible_fin,
                              peso_agua_ini, peso_agua_fin,
                              otros_ini=0.0, otros_fin=0.0):
        """Peso de la carga por diferencia de desplazamiento.
        Carga = (Desp_fin - Desp_ini)
                - [(Constante_fin - Constante_ini)]
                - [(Comb_fin - Comb_ini) + (Agua_fin - Agua_ini) + (Otros_fin - Otros_ini)]
        """
        delta_desp    = desp_llegada - desp_salida
        delta_const   = constante_buque
        delta_deducts = (peso_combustible_fin - peso_combustible_ini) + \
                        (peso_agua_fin - peso_agua_ini) + \
                        (otros_fin - otros_ini)
        peso_carga = delta_desp - delta_const - delta_deducts
        return peso_carga, delta_desp, delta_deducts

    def abrir_draft_survey(self):
        """Ventana completa de Draft Survey."""
        try:
            top = tk.Toplevel(self.root)
            top.title("Draft Survey — Estimación de Peso por Calados")
            top.geometry("1020x820")
            top.resizable(True, True)
            top.transient(self.root)

            fnt_b = ("Arial", 8, "bold")
            fnt_n = ("Arial", 8)

            # ── Canvas scrollable ─────────────────────────────────────────────
            main_frame = tk.Frame(top)
            main_frame.pack(fill="both", expand=True)
            canvas_ds = tk.Canvas(main_frame)
            sb_ds = ttk.Scrollbar(main_frame, orient="vertical", command=canvas_ds.yview)
            sf = ttk.Frame(canvas_ds)
            sf.bind("<Configure>", lambda e: canvas_ds.configure(scrollregion=canvas_ds.bbox("all")))
            canvas_ds.create_window((0,0), window=sf, anchor="nw")
            canvas_ds.configure(yscrollcommand=sb_ds.set)
            canvas_ds.pack(side="left", fill="both", expand=True)
            sb_ds.pack(side="right", fill="y")
            canvas_ds.bind("<MouseWheel>", lambda e: canvas_ds.yview_scroll(int(-1*(e.delta/120)),"units"))

            def make_section(parent, title, color="#1B3A5C"):
                lf = tk.LabelFrame(parent, text=f"  {title}  ", bg="#F8F9FA",
                                   font=fnt_b, fg=color, relief="groove", bd=2)
                lf.pack(fill="x", padx=15, pady=6)
                return lf

            def fv(default=""):
                return tk.StringVar(value=str(default))

            # ── SECCIÓN 1: Datos del buque ─────────────────────────────────────
            s1 = make_section(sf, "DATOS DEL BUQUE")
            buque_fields = [
                ("Nombre del Buque:", self.get_var("car_buque")),
                ("LBP (m):",          fv("0")),
                ("LCF desde popa (m):", fv("0")),
                ("Constante Buque (t):", fv("0")),
            ]
            _lbp_v  = buque_fields[1][1]
            _lcf_v  = buque_fields[2][1]
            for ci, (lbl, var) in enumerate(buque_fields):
                tk.Label(s1, text=lbl, bg="#F8F9FA", font=fnt_b).grid(row=0, column=ci*2, sticky="e", padx=6, pady=6)
                tk.Entry(s1, textvariable=var, width=14, font=fnt_n).grid(row=0, column=ci*2+1, sticky="w", padx=4)

            # ── SECCIÓN 2: Calados ─────────────────────────────────────────────
            def make_calados_section(parent, etiqueta, prefix, color):
                s = make_section(parent, etiqueta, color)
                vs = {}
                for ri, pos in enumerate(["proa","medio","popa"]):
                    tk.Label(s, text=f"{pos.upper()}:", bg="#F8F9FA", font=fnt_b, width=8
                             ).grid(row=ri, column=0, sticky="e", padx=6, pady=3)
                    for ci_s, (side, lbl_s) in enumerate([("b","Babor (m)"),("e","Estribor (m)")]):
                        k = f"{prefix}_{pos}_{side}"
                        v = fv("0"); vs[k] = v
                        tk.Label(s, text=lbl_s, bg="#F8F9FA", font=fnt_n).grid(row=ri, column=ci_s*2+1, sticky="e", padx=3)
                        tk.Entry(s, textvariable=v, width=9, font=fnt_n).grid(row=ri, column=ci_s*2+2, sticky="w", padx=3)
                        v.trace_add("write", lambda *a, px=prefix: _recalc(px))
                tk.Label(s, text="Cal. corr.:", bg="#F8F9FA", font=fnt_b).grid(row=3, column=0, sticky="e", padx=6, pady=3)
                vs[f"{prefix}_calado_corr"] = fv()
                tk.Label(s, textvariable=vs[f"{prefix}_calado_corr"], bg="#dceeff", relief="sunken",
                         font=fnt_n, width=20, anchor="w").grid(row=3, column=1, columnspan=3, sticky="w", padx=3)
                tk.Label(s, text="Hog/Sag:", bg="#F8F9FA", font=fnt_b).grid(row=3, column=4, sticky="e", padx=6)
                vs[f"{prefix}_hog_sag"] = fv()
                tk.Label(s, textvariable=vs[f"{prefix}_hog_sag"], bg="#fff3cd", relief="sunken",
                         font=fnt_n, width=18, anchor="w").grid(row=3, column=5, sticky="w", padx=3)
                tk.Label(s, text="Desplaz. (t):", bg="#F8F9FA", font=fnt_b).grid(row=4, column=0, sticky="e", padx=6, pady=3)
                vs[f"{prefix}_desp"] = fv()
                tk.Label(s, textvariable=vs[f"{prefix}_desp"], bg="#d4edda", relief="sunken",
                         font=(fnt_b[0], fnt_b[1]+1, "bold"), width=14, anchor="center"
                         ).grid(row=4, column=1, columnspan=2, sticky="w", padx=3)
                tk.Label(s, text="TPC:", bg="#F8F9FA", font=fnt_b).grid(row=4, column=4, sticky="e", padx=6)
                vs[f"{prefix}_tpc"] = fv()
                tk.Label(s, textvariable=vs[f"{prefix}_tpc"], bg="#dceeff", relief="sunken",
                         font=fnt_n, width=10, anchor="w").grid(row=4, column=5, sticky="w", padx=3)
                return s, vs

            s2, vs_ini = make_calados_section(sf, "CALADOS INICIALES (ANTES DE CARGA/DESCARGA)", "ini", "#1B3A5C")
            s3, vs_fin = make_calados_section(sf, "CALADOS FINALES (DESPUÉS DE CARGA/DESCARGA)", "fin", "#1D6A39")

            # ── SECCIÓN 3: Tabla hidrostática ──────────────────────────────────
            s4 = make_section(sf, "TABLA HIDROSTÁTICA (Calado → Desplazamiento/TPC/LCF)")
            f_hidro_btns = tk.Frame(s4, bg="#F8F9FA")
            f_hidro_btns.pack(anchor="w", padx=8, pady=4)
            hidro_data = []  # lista de [calado, desp, tpc, lcf]

            def _parse_hidro():
                rows = []
                for row in hidro_data:
                    try:
                        rows.append((float(row[0].get()), float(row[1].get()),
                                     float(row[2].get()), float(row[3].get())))
                    except: pass
                return sorted(rows, key=lambda r: r[0])

            hidro_frame = tk.Frame(s4, bg="#F8F9FA")
            hidro_frame.pack(fill="x", padx=8)
            for ci_h, htxt in enumerate(["Calado (m)","Desplaz. (t)","TPC (t/cm)","LCF desde popa (m)"]):
                tk.Label(hidro_frame, text=htxt, bg="#E8EAF6", font=fnt_b, width=18, relief="groove"
                         ).grid(row=0, column=ci_h, padx=2, pady=2)

            def _add_hidro_row(vals=None):
                ri = len(hidro_data)+1
                row_vars = []
                for ci in range(4):
                    v = fv(vals[ci] if vals else "0")
                    row_vars.append(v)
                    tk.Entry(hidro_frame, textvariable=v, width=18, font=fnt_n
                             ).grid(row=ri, column=ci, padx=2, pady=1)
                hidro_data.append(row_vars)

            for _ in range(6): _add_hidro_row()

            def _exportar_hidro_csv():
                import csv, tkinter.filedialog as fd2
                path2 = fd2.asksaveasfilename(defaultextension=".csv",
                    filetypes=[("CSV","*.csv")], initialfile="hidrostatica.csv", parent=top)
                if not path2: return
                with open(path2, "w", newline="", encoding="utf-8") as f2:
                    w2c = csv.writer(f2)
                    w2c.writerow(["Calado(m)","Desplaz(t)","TPC","LCF"])
                    for row in hidro_data:
                        try: w2c.writerow([r.get() for r in row])
                        except: pass
                messagebox.showinfo("Exportado", f"Tabla exportada: {path2}", parent=top)

            def _importar_hidro_csv():
                import csv, tkinter.filedialog as fd2
                path2 = fd2.askopenfilename(filetypes=[("CSV","*.csv")], parent=top)
                if not path2: return
                with open(path2, newline="", encoding="utf-8") as f2:
                    reader = csv.reader(f2)
                    rows_csv = [r for r in reader if r and r[0] != "Calado(m)"]
                while len(hidro_data) < len(rows_csv): _add_hidro_row()
                for ri2, row_vals in enumerate(rows_csv):
                    if ri2 >= len(hidro_data): break
                    for ci2, v2 in enumerate(row_vals[:4]):
                        hidro_data[ri2][ci2].set(v2)

            tk.Button(f_hidro_btns, text="+ Fila", bg="#1B3A5C", fg="white", font=("Arial",7,"bold"),
                      command=_add_hidro_row).pack(side="left", padx=4, ipadx=4, ipady=1)
            tk.Button(f_hidro_btns, text="Exportar CSV", bg="#455A64", fg="white", font=("Arial",7,"bold"),
                      command=_exportar_hidro_csv).pack(side="right", padx=4, ipadx=4, ipady=1)
            tk.Button(f_hidro_btns, text="Importar CSV", bg="#546E7A", fg="white", font=("Arial",7,"bold"),
                      command=_importar_hidro_csv).pack(side="right", padx=4, ipadx=4, ipady=1)

            # ═══════════════════════════════════════════════════════════════════
            # BALLAST WATERS — dinámica (agregar / quitar tanques)
            # ═══════════════════════════════════════════════════════════════════
            s_bw = make_section(sf, "BALLAST WATERS (Lastre)", "#1D6A39")
            f_bw_btns = tk.Frame(s_bw, bg="#F8F9FA")
            f_bw_btns.pack(anchor="w", padx=8, pady=4)

            bw_frame = tk.Frame(s_bw, bg="#F8F9FA")
            bw_frame.pack(fill="x", padx=8, pady=2)
            bw_rows = []   # lista de dicts: {name_var, bb_ini, bb_fin, es_ini, es_fin, widgets}

            # Cabecera fija
            for ci_h, htxt in enumerate(["TANQUE","BABOR ini(t)","BABOR fin(t)","ESTRIBOR ini(t)","ESTRIBOR fin(t)",""]):
                tk.Label(bw_frame, text=htxt, bg="#C8E6C9", font=fnt_b,
                         width=14 if ci_h==0 else (3 if ci_h==5 else 12), relief="groove"
                         ).grid(row=0, column=ci_h, padx=2, pady=2)

            def _bw_rebuild_totals():
                _recalc_final()

            def _bw_remove_row(row_dict):
                for w in row_dict["widgets"]:
                    try: w.grid_forget(); w.destroy()
                    except: pass
                bw_rows.remove(row_dict)
                # Re-grid remaining rows
                for ri3, rd in enumerate(bw_rows, 1):
                    for w in rd["widgets"]:
                        info = w.grid_info()
                        if info:
                            w.grid(row=ri3, column=info["column"])
                _bw_rebuild_totals()

            def _bw_add_row(name=""):
                ri = len(bw_rows) + 1
                row_dict = {}
                row_widgets = []
                name_v = fv(name or f"TK {ri} BALLAST")
                row_dict["name_var"] = name_v
                e_name = tk.Entry(bw_frame, textvariable=name_v, width=14, font=fnt_n, bg="#EAF4FC")
                e_name.grid(row=ri, column=0, padx=2, pady=2)
                row_widgets.append(e_name)
                for ci_s, key_s in enumerate(["bb_ini","bb_fin","es_ini","es_fin"]):
                    v = fv("0")
                    row_dict[key_s] = v
                    bg_e = "#E8F8EF" if "fin" in key_s else "#FFFDE7"
                    ew = tk.Entry(bw_frame, textvariable=v, width=12, font=fnt_n, bg=bg_e)
                    ew.grid(row=ri, column=ci_s+1, padx=2, pady=2)
                    row_widgets.append(ew)
                    v.trace_add("write", lambda *a: _recalc_final())
                btn_del = tk.Button(bw_frame, text="✕", font=("Arial",7), bg="#E53935", fg="white",
                                    width=2, cursor="hand2",
                                    command=lambda rd=row_dict: _bw_remove_row(rd))
                btn_del.grid(row=ri, column=5, padx=2, pady=2)
                row_widgets.append(btn_del)
                row_dict["widgets"] = row_widgets
                bw_rows.append(row_dict)

            # 8 tanques por defecto
            for i in range(1, 9):
                _bw_add_row(f"TK {i} BALLAST")

            tk.Button(f_bw_btns, text="+ Agregar tanque", bg="#1D6A39", fg="white", font=("Arial",7,"bold"),
                      command=_bw_add_row).pack(side="left", padx=4, ipadx=6, ipady=2)
            tk.Label(f_bw_btns, text="  (podés agregar o quitar tanques con ✕)", bg="#F8F9FA",
                     font=("Arial",7), fg="#555").pack(side="left")

            # ═══════════════════════════════════════════════════════════════════
            # DEDUCCIONES
            # ═══════════════════════════════════════════════════════════════════
            s5 = make_section(sf, "DEDUCCIONES ADICIONALES (Combustible, Agua Dulce, Otros)")
            ded_fields = [
                ("Constante Buque (t)",  "ded_cte_buque_ini"),
                ("FO inicio (t)",  "ded_fo_ini"),
                ("FO final (t)",   "ded_fo_fin"),
                ("DO inicio (t)",  "ded_do_ini"),
                ("DO final (t)",   "ded_do_fin"),
                ("FW inicio (t)",  "ded_fw_ini"),
                ("FW final (t)",   "ded_fw_fin"),
                ("Otros ini (t)",  "ded_ot_ini"),
                ("Otros fin (t)",  "ded_ot_fin"),
            ]
            ded_vs = {}
            for ci, (lbl, key) in enumerate(ded_fields):
                r, c = divmod(ci, 4)
                tk.Label(s5, text=lbl, bg="#F8F9FA", font=fnt_b).grid(row=r, column=c*2, sticky="e", padx=6, pady=3)
                v = fv("0"); ded_vs[key] = v
                tk.Entry(s5, textvariable=v, width=10, font=fnt_n).grid(row=r, column=c*2+1, sticky="w", padx=3)
                v.trace_add("write", lambda *a: _recalc_final())

            # ── Resultados ─────────────────────────────────────────────────────
            s6 = make_section(sf, "RESULTADO FINAL", "#7B1FA2")
            res_labels = [
                ("Desplaz. Inicial Corregido (t)", "res_desp_ini"),
                ("Desplaz. Final Corregido (t)",   "res_desp_fin"),
                ("Δ Desplazamiento (t)",            "res_delta_desp"),
                ("Total Deducciones (t)",           "res_deducs"),
                ("BW Total Inicio (t)",             "res_bw_ini"),
                ("BW Total Final (t)",              "res_bw_fin"),
                ("Constante Buque (t)",             "res_cte"),
                ("PESO CARGA ESTIMADO (t)",         "res_peso"),
                ("Obs.",                             "res_obs"),
            ]
            res_vs = {}
            for ri, (lbl, key) in enumerate(res_labels):
                r, c = divmod(ri, 2)
                bold = key == "res_peso"
                bg = "#D5E8D4" if bold else "#dceeff"
                tk.Label(s6, text=lbl, bg="#F8F9FA",
                         font=fnt_b if bold else ("Arial",8),
                         fg="#7B1FA2" if bold else "#222"
                         ).grid(row=r, column=c*2, sticky="e", padx=8, pady=4)
                v = fv("—"); res_vs[key] = v
                tk.Label(s6, textvariable=v, bg=bg, relief="sunken",
                         font=("Arial",11,"bold") if bold else fnt_n,
                         width=16, anchor="center"
                         ).grid(row=r, column=c*2+1, sticky="w", padx=4)

            # ── Cálculo ────────────────────────────────────────────────────────
            def _recalc(prefix):
                try:
                    vs = vs_ini if prefix == "ini" else vs_fin
                    pb = self.parse_float(vs.get(f"{prefix}_proa_b", fv()).get())
                    pe = self.parse_float(vs.get(f"{prefix}_proa_e", fv()).get())
                    mb = self.parse_float(vs.get(f"{prefix}_medio_b", fv()).get())
                    me = self.parse_float(vs.get(f"{prefix}_medio_e", fv()).get())
                    ab = self.parse_float(vs.get(f"{prefix}_popa_b", fv()).get())
                    ae = self.parse_float(vs.get(f"{prefix}_popa_e", fv()).get())
                    proa, popa, medio, trim = self.calc_draft_trim(pb,pe,ab,ae,mb,me)
                    _lbp_val = self.parse_float(_lbp_v.get())
                    if _lbp_val <= 0:
                        _lbp_v.set(""); return
                    calado_corr = self.calc_draft_corr_trim(proa, popa, medio, _lbp_val)
                    tipo_hs, mm_hs = self.calc_draft_hog_sag(proa, popa, medio)
                    tabla = _parse_hidro()
                    desp, tpc, lcf = self.calc_desplazamiento_interpolado(calado_corr, tabla)
                    lbp = _lbp_val
                    lcf_v = self.parse_float(_lcf_v.get()) or lcf
                    desp_corr = self.calc_draft_desp_trimado(desp, trim, tpc, lcf_v, lbp)
                    vs[f"{prefix}_calado_corr"].set(f"{calado_corr:.3f} m  (trim: {trim:+.3f})")
                    vs[f"{prefix}_hog_sag"].set(f"{tipo_hs} {mm_hs} mm")
                    vs[f"{prefix}_desp"].set(f"{desp_corr:,.1f}")
                    vs[f"{prefix}_tpc"].set(f"{tpc:.2f}")
                except Exception:
                    pass
                _recalc_final()

            def _recalc_final():
                try:
                    d_ini = self.parse_float(vs_ini.get("ini_desp", fv()).get().replace(",","") or "0")
                    d_fin = self.parse_float(vs_fin.get("fin_desp", fv()).get().replace(",","") or "0")
                    fo_i = self.parse_float(ded_vs["ded_fo_ini"].get())
                    fo_f = self.parse_float(ded_vs["ded_fo_fin"].get())
                    do_i = self.parse_float(ded_vs["ded_do_ini"].get())
                    do_f = self.parse_float(ded_vs["ded_do_fin"].get())
                    fw_i = self.parse_float(ded_vs["ded_fw_ini"].get())
                    fw_f = self.parse_float(ded_vs["ded_fw_fin"].get())
                    ot_i = self.parse_float(ded_vs["ded_ot_ini"].get())
                    ot_f = self.parse_float(ded_vs["ded_ot_fin"].get())
                    cte  = self.parse_float(ded_vs.get("ded_cte_buque_ini", fv("0")).get())
                    bw_i = 0.0; bw_f = 0.0
                    for rd in bw_rows:
                        bw_i += self.parse_float(rd["bb_ini"].get()) + self.parse_float(rd["es_ini"].get())
                        bw_f += self.parse_float(rd["bb_fin"].get()) + self.parse_float(rd["es_fin"].get())
                    comb_ini = fo_i + do_i
                    comb_fin = fo_f + do_f
                    agua_ini = fw_i + bw_i
                    agua_fin = fw_f + bw_f
                    peso, delta_d, delta_ded = self.calc_draft_peso_carga(
                        d_fin, d_ini, cte, comb_ini, comb_fin, agua_ini, agua_fin, ot_i, ot_f)
                    res_vs["res_desp_ini"].set(f"{d_ini:,.1f}")
                    res_vs["res_desp_fin"].set(f"{d_fin:,.1f}")
                    res_vs["res_delta_desp"].set(f"{delta_d:+,.1f}")
                    res_vs["res_deducs"].set(f"{delta_ded:+,.1f}")
                    res_vs["res_peso"].set(f"{peso:,.2f} t")
                    res_vs["res_obs"].set("CARGANDO" if peso > 0 else "DESCARGANDO")
                    res_vs["res_bw_ini"].set(f"{bw_i:,.1f}")
                    res_vs["res_bw_fin"].set(f"{bw_f:,.1f}")
                    res_vs["res_cte"].set(f"{cte:,.1f}")
                except Exception:
                    pass

            # ── PDF export ─────────────────────────────────────────────────────
            def _export_pdf():
                try:
                    from reportlab.lib.pagesizes import A4
                    from reportlab.pdfgen import canvas as rlc
                    import tkinter.filedialog as fd
                    from datetime import datetime

                    buque_nombre = self.get_var("car_buque").get() or "buque"
                    path = fd.asksaveasfilename(
                        defaultextension=".pdf",
                        filetypes=[("PDF","*.pdf")],
                        initialfile=f"DraftSurvey_{buque_nombre}.pdf",
                        parent=top)
                    if not path: return

                    cw, ch = A4
                    M = 45
                    c2 = rlc.Canvas(path, pagesize=A4)

                    def new_page():
                        nonlocal y2
                        c2.showPage()
                        y2 = ch - 50

                    def check_page(needed=25):
                        nonlocal y2
                        if y2 < needed:
                            new_page()

                    # Header
                    c2.setFillColorRGB(0.106, 0.227, 0.361)
                    c2.rect(M, ch-72, cw-2*M, 50, fill=1, stroke=0)
                    c2.setFillColorRGB(1,1,1)
                    c2.setFont("Helvetica-Bold", 15)
                    c2.drawCentredString(cw/2, ch-50, "DRAFT SURVEY — ESTIMACIÓN DE PESO DE CARGA")
                    c2.setFont("Helvetica", 9)
                    c2.drawCentredString(cw/2, ch-64,
                        f"Buque: {buque_nombre}   |   Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
                    y2 = ch - 90

                    def draw_section_title(text, color=(0.11,0.23,0.36)):
                        nonlocal y2
                        check_page(30)
                        c2.setFillColorRGB(*color)
                        c2.rect(M, y2-12, cw-2*M, 16, fill=1, stroke=0)
                        c2.setFillColorRGB(1,1,1)
                        c2.setFont("Helvetica-Bold", 9)
                        c2.drawString(M+6, y2-6, text)
                        y2 -= 20
                        c2.setFillColorRGB(0,0,0)
                        c2.setFont("Helvetica", 9)

                    def two_fields(l1, v1, l2="", v2=""):
                        nonlocal y2
                        check_page()
                        c2.setFont("Helvetica-Bold", 8); c2.setFillColorRGB(0.2,0.2,0.2)
                        c2.drawString(M, y2, str(l1))
                        c2.setFont("Helvetica", 8); c2.setFillColorRGB(0,0,0)
                        c2.drawString(M+130, y2, str(v1))
                        if l2:
                            c2.setFont("Helvetica-Bold", 8); c2.setFillColorRGB(0.2,0.2,0.2)
                            c2.drawString(cw/2+10, y2, str(l2))
                            c2.setFont("Helvetica", 8); c2.setFillColorRGB(0,0,0)
                            c2.drawString(cw/2+140, y2, str(v2))
                        y2 -= 12

                    def draw_field(label, val):
                        nonlocal y2
                        check_page()
                        c2.setFont("Helvetica-Bold", 8); c2.setFillColorRGB(0.2,0.2,0.2)
                        c2.drawString(M, y2, str(label))
                        c2.setFont("Helvetica", 8); c2.setFillColorRGB(0,0,0)
                        c2.drawString(M+130, y2, str(val))
                        y2 -= 12

                    # Datos buque
                    draw_section_title("DATOS DEL BUQUE")
                    two_fields("Nombre:", buque_nombre, "LBP (m):", _lbp_v.get())
                    two_fields("LCF desde popa (m):", _lcf_v.get(),
                               "Constante Buque (t):", ded_vs.get("ded_cte_buque_ini", fv("0")).get())
                    y2 -= 4

                    # Calados
                    for pref, label_c, vs_d in [("ini","CALADOS INICIALES", vs_ini),
                                                ("fin","CALADOS FINALES",   vs_fin)]:
                        draw_section_title(label_c)
                        for pos in ["proa","medio","popa"]:
                            pb2 = vs_d.get(f"{pref}_{pos}_b", fv()).get()
                            pe2 = vs_d.get(f"{pref}_{pos}_e", fv()).get()
                            two_fields(f"  {pos.upper()} babor:", pb2, f"  {pos.upper()} estribor:", pe2)
                        draw_field("  Calado corregido:", vs_d.get(f"{pref}_calado_corr", fv()).get())
                        draw_field("  Hog/Sag:", vs_d.get(f"{pref}_hog_sag", fv()).get())
                        draw_field("  Desplazamiento (t):", vs_d.get(f"{pref}_desp", fv()).get())
                        draw_field("  TPC:", vs_d.get(f"{pref}_tpc", fv()).get())
                        y2 -= 4

                    # Ballast Waters dinámico
                    draw_section_title("BALLAST WATERS", color=(0.11,0.42,0.24))
                    bw_col_x = [M, M+110, M+185, M+260, M+335]
                    c2.setFont("Helvetica-Bold", 7); c2.setFillColorRGB(0,0,0)
                    for ci_h2, htx in enumerate(["TANQUE","BABOR ini(t)","BABOR fin(t)","ESTRIBOR ini(t)","ESTRIBOR fin(t)"]):
                        c2.drawString(bw_col_x[ci_h2], y2, htx)
                    y2 -= 11
                    bw_total_ini = 0.0; bw_total_fin = 0.0
                    for rd in bw_rows:
                        check_page()
                        c2.setFont("Helvetica", 7)
                        c2.drawString(bw_col_x[0], y2, rd["name_var"].get())
                        for ci_s2, key_s2 in enumerate(["bb_ini","bb_fin","es_ini","es_fin"]):
                            val_str = rd[key_s2].get()
                            c2.drawString(bw_col_x[ci_s2+1], y2, val_str)
                            try:
                                fval = float(val_str)
                                if "ini" in key_s2: bw_total_ini += fval
                                else:               bw_total_fin += fval
                            except: pass
                        y2 -= 10
                    c2.setFont("Helvetica-Bold", 8)
                    c2.drawString(M, y2, f"TOTAL ini: {bw_total_ini:,.1f} t   |   TOTAL fin: {bw_total_fin:,.1f} t")
                    y2 -= 16

                    # Deducciones
                    draw_section_title("DEDUCCIONES ADICIONALES")
                    ded_show = [
                        ("FO inicio (t)", ded_vs["ded_fo_ini"].get()),
                        ("FO final (t)",  ded_vs["ded_fo_fin"].get()),
                        ("DO inicio (t)", ded_vs["ded_do_ini"].get()),
                        ("DO final (t)",  ded_vs["ded_do_fin"].get()),
                        ("FW inicio (t)", ded_vs["ded_fw_ini"].get()),
                        ("FW final (t)",  ded_vs["ded_fw_fin"].get()),
                        ("Otros ini (t)", ded_vs["ded_ot_ini"].get()),
                        ("Otros fin (t)", ded_vs["ded_ot_fin"].get()),
                    ]
                    for ii_d in range(0, len(ded_show), 2):
                        l1d, v1d = ded_show[ii_d]
                        l2d, v2d = ded_show[ii_d+1] if ii_d+1 < len(ded_show) else ("","")
                        two_fields(f"  {l1d}", v1d, f"  {l2d}", v2d)
                    y2 -= 6

                    # Resultado
                    draw_section_title("RESULTADO FINAL", color=(0.48,0.11,0.63))
                    for lbl_r, key_r in [
                        ("Desplaz. Inicial Corregido (t):", "res_desp_ini"),
                        ("Desplaz. Final Corregido (t):",   "res_desp_fin"),
                        ("Δ Desplazamiento (t):",           "res_delta_desp"),
                        ("Total Deducciones (t):",          "res_deducs"),
                        ("BW Total Inicial (t):",           "res_bw_ini"),
                        ("BW Total Final (t):",             "res_bw_fin"),
                        ("Constante Buque (t):",            "res_cte"),
                    ]:
                        draw_field(f"  {lbl_r}", res_vs.get(key_r, fv("—")).get())
                    y2 -= 4
                    check_page(30)
                    c2.setFillColorRGB(0.11,0.42,0.24)
                    c2.setFont("Helvetica-Bold", 13)
                    peso_txt = res_vs.get("res_peso", fv("—")).get()
                    obs_txt  = res_vs.get("res_obs",  fv("—")).get()
                    c2.drawString(M, y2, f"PESO DE CARGA ESTIMADO: {peso_txt}  —  {obs_txt}")
                    c2.setFillColorRGB(0,0,0)
                    y2 -= 20
                    
                    c2.save()
                    messagebox.showinfo("Draft Survey", f"PDF guardado:\n{path}", parent=top)
                except Exception as ex:
                    import traceback as tb2; tb2.print_exc()
                    messagebox.showerror("Error al exportar", str(ex), parent=top)

            # ── Botones ────────────────────────────────────────────────────────
            fbot = tk.Frame(top, bg="#1B3A5C")
            fbot.pack(fill="x", side="bottom")
            tk.Button(fbot, text="RECALCULAR TODO", bg="#2196F3", fg="white",
                      font=("Arial", 8, "bold"),
                      command=lambda: [_recalc("ini"), _recalc("fin")]
                      ).pack(side="left", padx=10, pady=6)
            tk.Button(fbot, text="EXPORTAR PDF", bg="#27AE60", fg="white",
                      font=("Arial", 8, "bold"),
                      command=_export_pdf
                      ).pack(side="left", padx=6, pady=6)
            tk.Button(fbot, text="CERRAR", font=("Arial", 8, "bold"),
                      command=top.destroy).pack(side="right", padx=10, pady=6)

        except Exception as e:
            traceback.print_exc()
            messagebox.showerror("Error Draft Survey", str(e))

    def calc_vcf(self, dens_input, temp_input, table_type):
        """
        Calcula VCF (Volume Correction Factor) según la norma seleccionada.

        NORMA 1980 — ASTM D1250-80 (tablas impresas, uso habitual aduana AR):
          54B: 4 zonas de densidad con coeficientes distintos
               ρ ≤ 770        K0=346.42278  K1=0.43884
               770 < ρ < 778  transición:   α = -0.0033612 + 2680.32/ρ²
               778 ≤ ρ < 839  K0=594.5418   K1=0
               ρ ≥ 839        K0=186.9696   K1=0.48618
          54A: K0=613.9723  K1=0  (igual en ambas normas)
          54D: K0=1489.0672 K1=0  (igual en ambas normas)

        NORMA 2004 — API MPMS 11.1 / ASTM D1250-04 (norma digital):
          54B: K0=346.4228  K1=0.4033  (fórmula única para todo el rango)
          54A: K0=613.9723  K1=0
          54D: K0=1489.0672 K1=0

        Resultado siempre redondeado a 5 decimales (ASTM D1250).
        """
        try:
            if not dens_input or not temp_input: return 1.0
            rho = float(dens_input)
            t   = float(temp_input)
            if rho <= 0: return 1.0
            if rho < 2.0: rho = rho * 1000.0     # g/cm³ → kg/m³

            # ── Tablas con fórmula lineal ─────────────────────────────────────
            if "Químico" in table_type:
                return round(max(0.50, min(1.50, 1.0 - 0.0011 * (t - 15.0))), 5)

            elif "GLP" in table_type or "Propano" in table_type:
                # API MPMS 11.2.4 / GPA 8217 tabla 54E — alpha depende de rho
                alpha_glp = 757.0 / (rho ** 2)
                return round(max(0.80, min(1.20, 1.0 - alpha_glp * (t - 15.0))), 5)

            elif "GNL" in table_type or "Criog" in table_type:
                # Metano líquido, base -162°C, alpha=0.00468/°C (NIST)
                return round(max(0.50, min(1.50, 1.0 - 0.00468 * (t - (-162.0)))), 5)

            elif "Amoniaco" in table_type or "Refrigerante" in table_type:
                # NH3, alpha=0.00226/°C (Measurement Canada / Haar & Gallagher)
                return round(max(0.70, min(1.30, 1.0 - 0.00226 * (t - 15.0))), 5)

            elif "Sin corrección" in table_type or "VCF=1" in table_type:
                return 1.00000

            # ── Tablas ASTM 54 con fórmula exponencial ────────────────────────
            norma = getattr(self, "norma_astm", None)
            usar_1980 = (norma is None) or (norma.get() == "1980")

            if "54A" in table_type:
                # Igual en 1980 y 2004
                return self._calc_vcf_exponencial(rho, t, 613.9723, 0.0)

            elif "54D" in table_type:
                # Igual en 1980 y 2004
                return self._calc_vcf_exponencial(rho, t, 1489.0672, 0.0)

            elif "54B" in table_type or True:   # fallback = 54B
                if usar_1980:
                    k0, k1, alpha_ov = self._vcf_k0k1_1980_54B(rho)
                    return self._calc_vcf_exponencial(rho, t, k0, k1, alpha_ov)
                else:
                    # 2004: fórmula única
                    return self._calc_vcf_exponencial(rho, t, 346.4228, 0.4033)

        except:
            return 1.0

    def generar_todos_reportes(self):
        target_dir = filedialog.askdirectory(title="Seleccione carpeta destino")
        if not target_dir: return
        errors = []
        
        all_tanks = self.lista_tanques + self.lista_carbonera
        
        # --- CREAR UN SOLO PDF UNIFICADO ---
        clean_buque = self.clean_filename(self.get_var('car_buque').get())
        if not clean_buque: clean_buque = "Reporte"
        unified_path = os.path.join(target_dir, f"Reporte_Completo_{clean_buque}.pdf")
        
        try:
            c = canvas.Canvas(unified_path, pagesize=landscape(A4))
        except Exception as e:
            traceback.print_exc()
            messagebox.showerror("Error", f"No se pudo crear el PDF:\n{unified_path}\n\nError: {str(e)}")
            return
        
        report_count = 0
        
        # 1. REPORTE TECNICO GLOBAL
        try:
            self.generar_reporte_tecnico_global("DETALLE_TECNICO_GLOBAL", target_dir, shared_canvas=c)
            report_count += 1
        except Exception as e:
            traceback.print_exc()
            errors.append(f"Reporte Técnico Global: {str(e)}")
        
        # 2. REPORTE GENERAL (incluye carbonera)
        try:
            self.generar_un_reporte("GENERAL", all_tanks, is_partial=False, output_folder=target_dir, shared_canvas=c)
            report_count += 1
        except Exception as e:
            traceback.print_exc()
            errors.append(f"Reporte General: {str(e)}")
        
        # 3. REPORTES POR DOCUMENTO (incluye carbonera en el mapa)
        mapa = {} 
        for tk_name in all_tanks:
            d_ini = self.get_var(f"inicial_{tk_name}_ddt_assign").get()
            d_fin = self.get_var(f"final_{tk_name}_ddt_assign").get()
            if d_ini:
                if d_ini not in mapa: mapa[d_ini] = []
                if tk_name not in mapa[d_ini]: mapa[d_ini].append(tk_name)
            if d_fin and d_fin != d_ini:
                if d_fin not in mapa: mapa[d_fin] = []
                if tk_name not in mapa[d_fin]: mapa[d_fin].append(tk_name)
        modes = [
            ("SEGÚN_LABORATORIO", "dens_lab"),
            ("SEGÚN_DOCUMENTO", "dens_doc"),
            ("SEGÚN_SALIDA", "dens_salida")
        ]
        for ddt_num, tanks in mapa.items():
            ddt_obj = next((d for d in self.ddt_data if d["numero"].get() == ddt_num), None)
            if not ddt_obj: continue
            safe_doc_num = self.clean_filename(ddt_num)
            if not safe_doc_num: safe_doc_num = "SinNombre"
            for suffix, mode_key in modes:
                try:
                    self.generar_un_reporte(f"DOC_{safe_doc_num}_{suffix}", tanks, is_partial=True, ddt_obj=ddt_obj, output_folder=target_dir, density_mode_key=mode_key, shared_canvas=c)
                    report_count += 1
                except Exception as e: 
                    traceback.print_exc()
                    errors.append(f"Reporte Parcial {ddt_num} ({suffix}): {str(e)}")
        
        # --- GUARDAR EL PDF UNIFICADO ---
        c.save()
        
        if errors:
            msg = "Algunos reportes fallaron:\n" + "\n".join(errors[:5]) 
            if len(errors) > 5: msg += "\n..."
            messagebox.showerror("Errores de Generación", msg)
        
        if report_count == 0:
            messagebox.showwarning("Atención", "No se generaron reportes. Verifique que no tenga los archivos abiertos.")
        else:
            messagebox.showinfo("Listo", f"Se generaron {report_count} reportes en un solo PDF:\n{os.path.basename(unified_path)}")
            try:
                if platform.system() == 'Windows': os.startfile(unified_path)
                elif platform.system() == 'Darwin': subprocess.call(('open', unified_path))
                else: subprocess.call(('xdg-open', unified_path))
            except: pass
    

    # ═══════════════════════════════════════════════════════════════════════════
    # PDF DRAWING METHODS (ReportLab) — mirror of TK drawings
    # ═══════════════════════════════════════════════════════════════════════════

    @staticmethod
    def _pdf_polygon(c, pts, fill=1, stroke=0):
        """Draw a polygon on a ReportLab canvas from a flat list of coordinates [x0,y0,x1,y1,...]."""
        if len(pts) < 4:
            return
        p = c.beginPath()
        p.moveTo(pts[0], pts[1])
        for i in range(2, len(pts), 2):
            p.lineTo(pts[i], pts[i+1])
        p.close()
        c.drawPath(p, fill=fill, stroke=stroke)

    def _pdf_prod_color(self, tk_name, etapa_key):
        """Return ReportLab HexColor for product (same palette as TK, never blue)."""
        from reportlab.lib import colors
        rgb_fill, rgb_out = self.get_prod_color(tk_name, etapa_key)
        try:
            return colors.HexColor(rgb_fill), colors.HexColor(rgb_out)
        except:
            return colors.HexColor("#F0B429"), colors.HexColor("#D4880A")

    def _pdf_draw_vertical_tank(self, c, x, y, W, H, etapa_key, title):
        """ReportLab vertical tank — versión ultra-realista con perspectiva 3D cilíndrica y elipses."""
        from reportlab.lib import colors
        tank_names = self.lista_tanques or ["TK 1"]
        flotante = "FLOTANTE" in self.get_tipo_medio()
        n = max(len(tank_names), 1)
        pad = 10
        TK_W = max(24, (W - 2*pad) // n)
        TK_H = H * 0.58
        # Real storage tanks: diameter ≈ 80-100% of height
        # For single tank, use wider proportions; for many tanks, narrower
        if n == 1:
            MAX_TK_W = TK_H * 1.0  # 1:1 ratio for single tank
        elif n <= 3:
            MAX_TK_W = TK_H * 0.75
        else:
            MAX_TK_W = TK_H * 0.55
        TK_W = min(TK_W, int(MAX_TK_W))
        # Center tanks horizontally
        total_tanks_w = n * TK_W
        x_offset = (W - total_tanks_w) / 2
        y_base = y + H * 0.14
        y_top  = y_base + TK_H
        # Radio vertical de la elipse de perspectiva del cilindro
        EL_RY = max(3, TK_H * 0.052)
        BUND_H = TK_H * 0.12

        c.saveState()
        # ── Fondo neutro ─────────────────────────────────────────────────
        c.setFillColor(colors.HexColor("#F7F9FC")); c.setLineWidth(0)
        c.rect(x, y, W, H, fill=1, stroke=0)
        # Suelo
        c.setFillColor(colors.HexColor("#5A6370")); c.setLineWidth(0)
        c.rect(x, y, W, H*0.12, fill=1, stroke=0)

        # Dique de contención
        c.setFillColor(colors.HexColor("#C2C6C9")); c.setStrokeColor(colors.HexColor("#7F8C8D"))
        c.setLineWidth(1.2); c.rect(x+pad//2, y_base-BUND_H, W-pad, BUND_H, fill=1, stroke=1)
        c.setFillColor(colors.HexColor("#9EA5A8")); c.setLineWidth(0)
        c.rect(x+pad//2, y_base-BUND_H, W-pad, 2.5, fill=1, stroke=0)
        c.setStrokeColor(colors.HexColor("#A8ADB0")); c.setLineWidth(0.4)
        for li in range(0, int(W-pad), max(5, int(W-pad)//14)):
            c.line(x+pad//2+li, y_base-BUND_H+2, x+pad//2+li+4, y_base)

        # Título
        # Título dentro del área de dibujo con fondo para legibilidad
        c.setFont("Helvetica-Bold", 6.5)
        label = "TANQUE TECHO FLOTANTE" if flotante else "TANQUE TECHO FIJO"
        instalacion = self.get_var("car_buque").get() or ""
        title_txt = f"{label}  —  {title}" + (f"  ({instalacion})" if instalacion else "")
        _tw = c.stringWidth(title_txt, "Helvetica-Bold", 6.5)
        c.setFillColor(colors.HexColor("#F7F9FC")); c.setLineWidth(0)
        c.rect(x+W/2 - _tw/2 - 4, y+H-14-4, _tw+8, 12, fill=1, stroke=0)
        c.setFillColor(colors.HexColor("#1B3A5C"))
        c.drawCentredString(x+W/2, y+H-14, title_txt)

        for i, tn in enumerate(tank_names[:n]):
            tx = x + x_offset + i*TK_W; tw = TK_W - 5
            mid_x = tx + tw/2

            # Base de hormigón
            base_h = max(3, TK_H*0.04)
            base_x = tx - max(2, tw*0.09)
            base_w = tw + max(4, tw*0.18)
            c.setFillColor(colors.HexColor("#AAB2B5")); c.setStrokeColor(colors.HexColor("#8D9498"))
            c.setLineWidth(0.5); c.rect(base_x, y_base-base_h, base_w, base_h, fill=1, stroke=1)
            # Detalle base
            c.setStrokeColor(colors.HexColor("#9EA8B0")); c.setLineWidth(0.3); c.setDash([2,3])
            c.rect(base_x+3, y_base-base_h+2, base_w-6, base_h-4, fill=0, stroke=1)
            c.setDash()

            # ── Cuerpo cilíndrico: franjas verticales detalladas ──────────
            stripe_defs_pdf = [
                (0.00, 0.08, "#7E8A94"),
                (0.08, 0.20, "#96A2AC"),
                (0.20, 0.38, "#B2BEC8"),
                (0.38, 0.55, "#D2DAE0"),
                (0.55, 0.66, "#E0E8EC"),  # brillo máximo
                (0.66, 0.78, "#D8DFE3"),
                (0.78, 0.90, "#C4CDD4"),
                (0.90, 1.00, "#9EA8B2"),
            ]
            for s_from, s_to, sc_p in stripe_defs_pdf:
                c.setFillColor(colors.HexColor(sc_p)); c.setLineWidth(0)
                c.rect(tx + tw*s_from, y_base+EL_RY, tw*(s_to-s_from)+0.5, TK_H-EL_RY*2, fill=1, stroke=0)

            # ── Juntas de soldadura entre virolas (líneas horizontales) ───
            c.setStrokeColor(colors.HexColor("#6B7880")); c.setLineWidth(0.4)
            c.setDash([3, 2])
            for sc_frac in [0.15, 0.30, 0.45, 0.60, 0.75]:
                sc_y = y_base + TK_H * sc_frac
                c.line(tx, sc_y, tx + tw, sc_y)
            c.setDash()

            # ── Nivel de llenado ─────────────────────────────────────────
            vp, wp = self._get_fill_pct(tn, etapa_key)
            pnm = self.get_var(f"{etapa_key}_{tn}_prod_name").get() if etapa_key else ""
            vlit = self.get_var(f"{etapa_key}_{tn}_vol_nat_prod").get() if etapa_key else ""
            if vp > 0.02:
                fc, oc = self._pdf_prod_color(tn, etapa_key)
                fill_y = y_base + TK_H*wp
                fill_h = TK_H*(vp-wp)
                # Agua (si corresponde)
                if wp > 0:
                    c.setFillColor(colors.HexColor("#5DADE2")); c.setLineWidth(0)
                    c.rect(tx+2, y_base+EL_RY, tw-4, TK_H*wp-EL_RY, fill=1, stroke=0)
                    # Elipse del nivel de agua
                    c.setFillColor(colors.HexColor("#2E86C1"))
                    c.ellipse(tx+2, y_base+TK_H*wp-EL_RY, tx+tw-2, y_base+TK_H*wp+EL_RY, fill=1, stroke=0)
                # Producto
                c.setFillColor(fc); c.setLineWidth(0)
                c.rect(tx+2, fill_y+EL_RY, tw-4, fill_h-EL_RY*2, fill=1, stroke=0)
                # Reflejo superior del producto
                c.setFillColor(colors.HexColor("#FFFFFF")); c.setFillAlpha(0.30)
                c.rect(tx+2, fill_y+fill_h-fill_h*0.12, tw-4, fill_h*0.10, fill=1, stroke=0)
                c.setFillAlpha(1.0)
                # Elipse de la SUPERFICIE del producto (perspectiva)
                c.setFillColor(fc)
                c.ellipse(tx+2, fill_y+fill_h-EL_RY, tx+tw-2, fill_y+fill_h+EL_RY, fill=1, stroke=0)
                c.setStrokeColor(oc); c.setLineWidth(0.8)
                c.ellipse(tx+2, fill_y+fill_h-EL_RY, tx+tw-2, fill_y+fill_h+EL_RY, fill=0, stroke=1)
                # Porcentaje — color con contraste correcto
                fc_hex = self.get_prod_color(tn, etapa_key)[0]
                txt_col_pdf = colors.HexColor(self.contrast_text(fc_hex))
                c.setFillColor(txt_col_pdf)
                c.setFont("Helvetica-Bold", 6)
                c.drawCentredString(mid_x, fill_y + fill_h*0.45, f"{vp*100:.1f}%")
                if pnm:
                    c.setFont("Helvetica", 5); c.setFillColor(txt_col_pdf)
                    c.drawCentredString(mid_x, fill_y + fill_h*0.22, pnm[:12])

            # ── Elipse BASE del cilindro (perspectiva inferior) ───────────
            # La parte inferior de la elipse simula el fondo del cilindro visible
            c.setFillColor(colors.HexColor("#808D97")); c.setLineWidth(0)
            c.ellipse(tx, y_base-EL_RY, tx+tw, y_base+EL_RY, fill=1, stroke=0)
            # Línea de contorno de la elipse base
            c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(1)
            c.ellipse(tx, y_base-EL_RY, tx+tw, y_base+EL_RY, fill=0, stroke=1)

            # ── Anillos de rigidización — virolas delgadas ────────────────
            for ri in [0.18, 0.36, 0.54, 0.72]:
                ry = y_base + TK_H * ri
                c.setFillColor(colors.HexColor("#9EA9B3")); c.setLineWidth(0)
                c.ellipse(tx, ry-1.2, tx+tw, ry+1.2, fill=1, stroke=0)
                c.setStrokeColor(colors.HexColor("#7F8C8D")); c.setLineWidth(0.5)
                c.ellipse(tx, ry-1.2, tx+tw, ry+1.2, fill=0, stroke=1)

            # ── Viento girder — anillo prominente techo fijo a 82% ────────
            if not flotante:
                wg_y = y_base + TK_H * 0.82
                c.setFillColor(colors.HexColor("#8C9BA5")); c.setLineWidth(0)
                c.ellipse(tx-2, wg_y-3.5, tx+tw+2, wg_y+3.5, fill=1, stroke=0)
                c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(1.0)
                c.ellipse(tx-2, wg_y-3.5, tx+tw+2, wg_y+3.5, fill=0, stroke=1)
                # Orejetas del wind girder
                c.setFillColor(colors.HexColor("#8C9BA5")); c.setLineWidth(0)
                c.rect(tx-5, wg_y-2.5, 5, 5, fill=1, stroke=0)
                c.rect(tx+tw, wg_y-2.5, 5, 5, fill=1, stroke=0)

            # ── Techo con tapa elíptica superior ─────────────────────────
            if flotante:
                # ── FÍSICA CORRECTA DEL TECHO FLOTANTE ───────────────────────
                # Patas estándar ≈ 1.8m → float level ≈ 15% del tanque
                # Por debajo: techo apoyado en patas (posición fija)
                # Por encima: techo flota exactamente al nivel del líquido
                PATAS_FRAC = 0.15
                ponton_h = max(3, EL_RY*0.95)
                inner_h = TK_H - 2*EL_RY

                liq_y   = y_base + TK_H * vp
                patas_y = y_base + inner_h * PATAS_FRAC + EL_RY

                if vp <= PATAS_FRAC:
                    deck_y = patas_y   # apoyado en patas
                else:
                    deck_y = liq_y     # flotando en el líquido

                # ── Patas (siempre visibles, desde base del pontón hasta fondo) ──
                pata_base_y = deck_y - ponton_h
                c.setStrokeColor(colors.HexColor("#7F8C8D")); c.setLineWidth(1.2)
                for pi_frac in [0.15, 0.38, 0.62, 0.85]:
                    px = tx + tw * pi_frac
                    c.line(px, pata_base_y, px, y_base + EL_RY)
                    # Zapata (foot plate)
                    c.setFillColor(colors.HexColor("#95A5A6")); c.setLineWidth(0)
                    c.rect(px-2.5, y_base+EL_RY-1.5, 5, 3, fill=1, stroke=0)

                # ── Pontón perimetral ──────────────────────────────────────
                c.setFillColor(colors.HexColor("#5B7B8A"))
                c.setStrokeColor(colors.HexColor("#4A6270")); c.setLineWidth(1)
                c.ellipse(tx, deck_y-ponton_h, tx+tw, deck_y+ponton_h, fill=1, stroke=1)
                c.setFillColor(colors.HexColor("#6D92A1")); c.setLineWidth(0)
                c.ellipse(tx+ponton_h+2, deck_y-ponton_h+1.5, tx+tw-ponton_h-2, deck_y+ponton_h-1.5, fill=1, stroke=0)
                # Brillo del pontón
                c.setFillColor(colors.HexColor("#A8C4CF")); c.setLineWidth(0)
                c.ellipse(tx+ponton_h+4, deck_y-ponton_h+2, tx+tw*0.45, deck_y-ponton_h*0.2, fill=1, stroke=0)
                # Rim seal (sello perimetral) — elemento clave del techo flotante
                c.setStrokeColor(colors.HexColor("#4A5568")); c.setLineWidth(1.5); c.setDash([2,2])
                c.ellipse(tx-2, deck_y-ponton_h-3, tx+tw+2, deck_y+ponton_h-1, fill=0, stroke=1)
                c.setDash()
                # Cables guía
                c.setStrokeColor(colors.HexColor("#95A5A6")); c.setLineWidth(0.5); c.setDash([3,4])
                for cg_f in [0.22, 0.5, 0.78]:
                    cgx = tx + tw*cg_f
                    c.line(cgx, deck_y+ponton_h, cgx, y_base+TK_H)
                c.setDash()
                # ── Flecha de SONDAJE (datum→pontón) ────────────────────────
                _pdf_s_x = tx + tw*0.12
                _pdf_s_bot = y_base + EL_RY + 2
                _pdf_s_top = deck_y - ponton_h - 1
                if _pdf_s_top > _pdf_s_bot + 8:
                    c.setStrokeColor(colors.HexColor("#F4D03F"))
                    c.setFillColor(colors.HexColor("#F4D03F"))
                    c.setLineWidth(0.8)
                    c.setDash([2,3])
                    c.line(_pdf_s_x, _pdf_s_bot, _pdf_s_x, _pdf_s_top)
                    c.setDash()
                    # Flechas en los extremos
                    _arr = c.beginPath()
                    _arr.moveTo(_pdf_s_x-2, _pdf_s_top-4)
                    _arr.lineTo(_pdf_s_x, _pdf_s_top)
                    _arr.lineTo(_pdf_s_x+2, _pdf_s_top-4)
                    c.drawPath(_arr, fill=1, stroke=0)
                    _arr2 = c.beginPath()
                    _arr2.moveTo(_pdf_s_x-2, _pdf_s_bot+4)
                    _arr2.lineTo(_pdf_s_x, _pdf_s_bot)
                    _arr2.lineTo(_pdf_s_x+2, _pdf_s_bot+4)
                    c.drawPath(_arr2, fill=1, stroke=0)
                    # Placa datum en la base
                    c.setFillColor(colors.HexColor("#F4D03F"))
                    c.rect(_pdf_s_x-3, _pdf_s_bot-2, 6, 2, fill=1, stroke=0)
                    # Label "S" centrado
                    _mid_s = (_pdf_s_bot + _pdf_s_top) / 2
                    c.setFont("Helvetica-Bold", max(3, int(TK_H*0.04)))
                    c.setFillColor(colors.HexColor("#D4AC0D"))
                    c.drawCentredString(_pdf_s_x, _mid_s, "⌇")
                # Indicador de nivel lateral
                lv_x_pdf = tx + tw*0.88
                c.setStrokeColor(colors.HexColor("#E74C3C")); c.setLineWidth(1.2); c.setDash([3,3])
                c.line(lv_x_pdf, deck_y-ponton_h, lv_x_pdf, y_base)
                c.setDash()
                c.setFillColor(colors.HexColor("#E74C3C")); c.setLineWidth(0)
                c.rect(lv_x_pdf-2.5, deck_y-ponton_h-4, 5, 4, fill=1, stroke=0)
            else:
                # Techo cónico
                apex_y = y_base + TK_H + EL_RY + max(5, TK_H*0.09)
                # Cara iluminada del cono
                p = c.beginPath()
                p.moveTo(tx, y_base+TK_H); p.lineTo(mid_x, apex_y); p.lineTo(tx+tw*0.55, y_base+TK_H); p.close()
                c.setFillColor(colors.HexColor("#D2D9DE")); c.drawPath(p, fill=1, stroke=0)
                # Cara sombreada
                p2 = c.beginPath()
                p2.moveTo(tx+tw*0.45, y_base+TK_H); p2.lineTo(mid_x, apex_y); p2.lineTo(tx+tw, y_base+TK_H); p2.close()
                c.setFillColor(colors.HexColor("#A8B5B8")); c.drawPath(p2, fill=1, stroke=0)
                # Contorno cono
                c.setStrokeColor(colors.HexColor("#7A8896")); c.setLineWidth(1.2)
                c.line(tx, y_base+TK_H, mid_x, apex_y)
                c.line(tx+tw, y_base+TK_H, mid_x, apex_y)
                # Nervaduras
                c.setStrokeColor(colors.HexColor("#95A5A6")); c.setLineWidth(0.6)
                for ri2 in [0.3, 0.7]:
                    c.line(tx+tw*ri2, y_base+TK_H, mid_x, apex_y)
                # Tapa ELÍPTICA superior del cilindro (collar del techo — efecto 3D clave)
                c.setFillColor(colors.HexColor("#CAD2D7"))
                c.ellipse(tx, y_base+TK_H-EL_RY, tx+tw, y_base+TK_H+EL_RY, fill=1, stroke=0)
                c.setStrokeColor(colors.HexColor("#808B96")); c.setLineWidth(1.2)
                c.ellipse(tx, y_base+TK_H-EL_RY, tx+tw, y_base+TK_H+EL_RY, fill=0, stroke=1)
                # Brillo en la tapa
                c.setFillColor(colors.HexColor("#E2E8EC")); c.setLineWidth(0)
                c.ellipse(tx+tw*0.15, y_base+TK_H-EL_RY*0.5, tx+tw*0.55, y_base+TK_H+EL_RY*0.3, fill=1, stroke=0)

            # ── Tobera de venteo ──────────────────────────────────────────
            if not flotante:
                vnt_x = tx + tw*0.70
                vnt_base = y_base + TK_H + EL_RY
                c.setFillColor(colors.HexColor("#808B96")); c.setLineWidth(0)
                c.rect(vnt_x-2.5, vnt_base, 5, max(6, TK_H*0.07), fill=1, stroke=0)
                c.setFillColor(colors.HexColor("#4A5568"))
                c.rect(vnt_x-3, vnt_base+TK_H*0.07, 6, 2, fill=1, stroke=0)
                p3 = c.beginPath()
                p3.moveTo(vnt_x-5, vnt_base+TK_H*0.07); p3.lineTo(vnt_x, vnt_base+TK_H*0.10+5); p3.lineTo(vnt_x+5, vnt_base+TK_H*0.07); p3.close()
                c.setFillColor(colors.HexColor("#95A5A6")); c.drawPath(p3, fill=1, stroke=0)

            # ── Escalera cat ladder ───────────────────────────────────────
            esc_x = tx + tw + 3.5
            if esc_x < x + W - 4:
                c.setStrokeColor(colors.HexColor("#95A5A6")); c.setLineWidth(1)
                c.line(esc_x-1.5, y_base, esc_x-1.5, y_base+TK_H)
                c.line(esc_x+2.5, y_base, esc_x+2.5, y_base+TK_H)
                c.setLineWidth(0.5)
                sy2 = y_base
                step_h = max(4, TK_H/10)
                while sy2 < y_base+TK_H:
                    c.line(esc_x-2.5, sy2, esc_x+3.5, sy2)
                    sy2 += step_h
                # Plataforma en la cima
                c.setFillColor(colors.HexColor("#808B96")); c.setLineWidth(0)
                c.rect(esc_x-4, y_base+TK_H+EL_RY, 9, 2.5, fill=1, stroke=0)

            # ── Tubería base con válvula ──────────────────────────────────
            pipe_y = y_base - base_h - 2
            pipe_x = tx + tw*0.28
            c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(2.5)
            c.line(pipe_x, y_base, pipe_x, pipe_y)
            c.line(pipe_x, pipe_y, tx - 9, pipe_y)
            c.setFillColor(colors.HexColor("#4A5568")); c.setLineWidth(0)
            c.rect(pipe_x-3, y_base-2, 6, 3, fill=1, stroke=0)
            # Válvula (mariposa)
            vv_y = pipe_y
            c.setFillColor(colors.HexColor("#2C3E50")); c.setStrokeColor(colors.HexColor("#1B2631")); c.setLineWidth(0.8)
            p4 = c.beginPath()
            p4.moveTo(pipe_x-4, vv_y-3.5); p4.lineTo(pipe_x+4, vv_y+3.5); p4.lineTo(pipe_x+4, vv_y-3.5); p4.lineTo(pipe_x-4, vv_y+3.5); p4.close()
            c.drawPath(p4, fill=1, stroke=1)
            c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(0.8)
            c.line(pipe_x, vv_y-3.5, pipe_x, vv_y-8)
            c.circle(pipe_x, vv_y-9.5, 2.5, stroke=1, fill=0)
            # Tubería de retorno
            ret_x2 = tx + tw*0.72
            c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(2)
            c.line(ret_x2, y_base, ret_x2, pipe_y-1)
            c.setFillColor(colors.HexColor("#4A5568")); c.setLineWidth(0)
            c.rect(ret_x2-3, y_base-2, 6, 3, fill=1, stroke=0)

            # ── Indicador de nivel lateral ────────────────────────────────
            c.setStrokeColor(colors.HexColor("#BDC3C7")); c.setLineWidth(0.8); c.setDash([1.5,2])
            c.line(tx-5, y_base, tx-5, y_base+TK_H)
            c.setDash()
            if vp > 0:
                lv_y = y_base + TK_H*vp
                c.setFillColor(colors.HexColor("#E74C3C")); c.setLineWidth(0)
                c.rect(tx-8, lv_y-2, 6, 4, fill=1, stroke=0)
            c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(0.5)
            for mk in [0.2, 0.4, 0.6, 0.8]:
                my = y_base + TK_H*mk
                c.line(tx-7, my, tx-3, my)

            # ── Contorno final ────────────────────────────────────────────
            c.setFillColor(colors.white); c.setStrokeColor(colors.HexColor("#4A5568")); c.setLineWidth(1.8)
            c.rect(tx, y_base+EL_RY, tw, TK_H-EL_RY*2, fill=0, stroke=1)

            # ── Etiqueta ──────────────────────────────────────────────────
            c.setFillColor(colors.HexColor("#2C3E50")); c.setFont("Helvetica-Bold", 5.5)
            c.drawCentredString(mid_x, y_base - 9, tn[:10])
            if vlit:
                c.setFont("Helvetica", 4.5); c.setFillColor(colors.HexColor("#5D6D7E"))
                c.drawCentredString(mid_x, y_base - 15, f"{vlit} L")

        c.restoreState()

    def _pdf_draw_spheres(self, c, x, y, W, H, etapa_key, title):
        """ReportLab sphere drawing — versión premium con gradiente 3D."""
        from reportlab.lib import colors
        tank_names = self.lista_tanques or ["ESFERA 1"]
        n = max(len(tank_names), 1)
        instalacion = self.get_var("car_buque").get() or ""

        c.saveState()
        c.setFillColor(colors.HexColor("#F5F7FA")); c.setStrokeColor(colors.HexColor("#1B3A5C")); c.setLineWidth(1.5)
        c.rect(x, y, W, H, fill=1, stroke=1)
        c.setFillColor(colors.HexColor("#1B3A5C")); c.setFont("Helvetica-Bold", min(8, max(6, H*0.04)))
        title_txt = f"ESFERAS DE GAS - {title}" + (f" ({instalacion})" if instalacion else "")
        max_chars = int(W / 4.5)
        if len(title_txt) > max_chars:
            title_txt = title_txt[:max_chars-3] + "..."
        c.drawCentredString(x+W/2, y+H-10, title_txt)

        # Suelo
        c.setFillColor(colors.HexColor("#CCD1D9")); c.setLineWidth(0)
        c.rect(x+1, y+2, W-2, H*0.07, fill=1, stroke=0)

        sph_zone = W / n; ground_y = y + H*0.12
        for i, tn in enumerate(tank_names[:n]):
            cx_s = x + i*sph_zone + sph_zone/2
            # Cap sphere radius: max 22% of H, max 35% of zone, hard max 65pt
            sph_r = min(H*0.22, sph_zone*0.35, 65)
            # Center sphere vertically between ground and top
            cy_s = ground_y + (y + H - 18 - ground_y) * 0.55
            lsp = sph_r*0.68; lh = sph_r*0.65

            # Patas (2 delanteras, 2 traseras)
            c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(max(1.5, sph_r*0.07))
            c.line(cx_s-lsp*0.65, cy_s+sph_r*0.78, cx_s-lsp*0.65, ground_y)
            c.line(cx_s+lsp*0.65, cy_s+sph_r*0.78, cx_s+lsp*0.65, ground_y)
            c.setLineWidth(max(2, sph_r*0.09))
            c.line(cx_s-lsp, cy_s+sph_r*0.82, cx_s-lsp, ground_y)
            c.line(cx_s+lsp, cy_s+sph_r*0.82, cx_s+lsp, ground_y)
            # Anillo base
            c.setFillColor(colors.HexColor("#4A5568")); c.setLineWidth(0)
            c.ellipse(cx_s-lsp-4, ground_y-3, cx_s+lsp+4, ground_y+2, fill=1, stroke=0)
            # Arriostramiento
            c.setStrokeColor(colors.HexColor("#7F8C8D")); c.setLineWidth(0.8)
            c.setDash([4,3])
            c.line(cx_s-lsp, cy_s+sph_r*0.82, cx_s+lsp, ground_y)
            c.line(cx_s+lsp, cy_s+sph_r*0.82, cx_s-lsp, ground_y)
            c.setDash()
            # Anillo horizontal
            mid_leg_y = cy_s + sph_r*0.60
            c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(1.5)
            c.line(cx_s-lsp-2, mid_leg_y, cx_s+lsp+2, mid_leg_y)

            # Gradiente esfera (capas concéntricas — gris-plata industrial, sin azules)
            sphere_gradient = [
                ("#2C3E50", 1.00), ("#3D5166", 0.92), ("#4A6174", 0.83),
                ("#5D7585", 0.74), ("#7B8D9A", 0.64), ("#9EADB7", 0.53),
                ("#C4CDD3", 0.42), ("#E0E6E9", 0.30),
            ]
            for sg_col, sg_frac in sphere_gradient:
                sr = sph_r * sg_frac
                offset_x = int(sph_r * (1-sg_frac) * 0.4)
                offset_y = int(sph_r * (1-sg_frac) * 0.5)
                c.setFillColor(colors.HexColor(sg_col)); c.setLineWidth(0)
                c.circle(cx_s + offset_x*0.5, cy_s - offset_y*0.5, sr, fill=1, stroke=0)
            c.setFillColor(colors.HexColor("#1A4F6E")); c.setLineWidth(0); c.setStrokeColor(colors.HexColor("#1A4F6E")); c.setLineWidth(2)
            c.circle(cx_s, cy_s, sph_r, fill=0, stroke=1)

            # Llenado
            try:
                vol = self.parse_float(self.get_var(f"{etapa_key}_{tn}_vol_nat_prod").get() or
                                       self.get_var(f"{etapa_key}_{tn}_vol_liq").get() or "0")
                cap = self.parse_float(self.get_var(f"{etapa_key}_{tn}_alt_ref").get() or "1")
                vp = min(max(vol/(cap if cap>0 else 1), 0), 1)
            except: vp = 0.0
            if vp > 0.02:
                fc, oc = self._pdf_prod_color(tn, etapa_key)
                fill_h = sph_r * 2 * vp
                fill_bot = cy_s - sph_r
                fill_top = fill_bot + fill_h
                # Clip al contorno circular de la esfera
                c.saveState()
                clip_path = c.beginPath()
                clip_path.circle(cx_s, cy_s, sph_r)
                c.clipPath(clip_path, stroke=0)
                c.setFillColor(fc); c.setLineWidth(0)
                c.rect(cx_s-sph_r, fill_bot, sph_r*2, fill_h, fill=1, stroke=0)
                c.restoreState()
                c.setStrokeColor(colors.HexColor("#1A4F6E")); c.setLineWidth(2)
                c.circle(cx_s, cy_s, sph_r, fill=0, stroke=1)
                c.setStrokeColor(oc); c.setLineWidth(1); c.setDash([4,3])
                c.line(cx_s-sph_r+2, fill_top, cx_s+sph_r-2, fill_top); c.setDash()
                c.setFillColor(colors.white); c.setFont("Helvetica-Bold", max(5, int(sph_r*0.25)))
                c.drawCentredString(cx_s, cy_s, f"{vp*100:.0f}%")

            # Toberas cima (3 nozzle pipes on top)
            c.setFillColor(colors.HexColor("#4A5568")); c.setLineWidth(0)
            for nxi in [-sph_r*0.3, 0, sph_r*0.3]:
                c.rect(cx_s+nxi-2, cy_s+sph_r, 4, sph_r*0.18, fill=1, stroke=0)

            # Top access platform
            plat_y = cy_s + sph_r + sph_r*0.18 + 2
            plat_hw = sph_r * 0.5
            c.setFillColor(colors.HexColor("#4A5568")); c.setLineWidth(0)
            c.rect(cx_s - plat_hw, plat_y, plat_hw*2, 2, fill=1, stroke=0)

            # Handrail posts (2 thin vertical lines rising from platform ends)
            c.setStrokeColor(colors.HexColor("#4A5568")); c.setLineWidth(0.75)
            c.line(cx_s - plat_hw + 1, plat_y + 2, cx_s - plat_hw + 1, plat_y + 5)
            c.line(cx_s + plat_hw - 1, plat_y + 2, cx_s + plat_hw - 1, plat_y + 5)

            # Central access trunk (thin dashed line from platform down through sphere center to bottom)
            c.setStrokeColor(colors.HexColor("#4A5568")); c.setLineWidth(0.5); c.setDash([3, 2])
            c.line(cx_s, plat_y, cx_s, cy_s - sph_r)
            c.setDash()

            # Etiqueta
            c.setFillColor(colors.HexColor("#1B3A5C")); c.setFont("Helvetica-Bold", max(4, min(7, int(sph_r*0.15))))
            c.drawCentredString(cx_s, ground_y-6, tn[:12])
        c.restoreState()

    def _pdf_draw_liquid_truck(self, c, x, y, W, H, etapa_key, title):
        """ReportLab cistern truck — versión premium con cabina, ruedas y detalles industriales."""
        from reportlab.lib import colors
        import math
        tank_names = self.lista_tanques or ["TK 1"]
        n = max(len(tank_names), 1)
        pat = self.get_var("car_patente").get() or ""

        c.saveState()
        # Fondo neutro
        c.setFillColor(colors.HexColor("#F5F7FA")); c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(1.5)
        c.rect(x, y, W, H, fill=1, stroke=1)
        # Suelo simple
        c.setFillColor(colors.HexColor("#808B96")); c.setLineWidth(0)
        c.rect(x+1, y+2, W-2, H*0.10, fill=1, stroke=0)

        c.setFillColor(colors.HexColor("#1B3A5C")); c.setFont("Helvetica-Bold", min(8, max(6, H*0.04)))
        hdr = f"CAMION CISTERNA - {title}" + (f" | Patente: {pat}" if pat else "")
        max_chars = int(W / 4.5)
        if len(hdr) > max_chars:
            hdr = hdr[:max_chars-3] + "..."
        c.drawCentredString(x+W/2, y+H-10, hdr)

        # Proporciones
        AXLE_R   = max(7, H*0.085)
        GND_Y    = y + H*0.12 + AXLE_R*2   # nivel de suelo (Y en ReportLab crece hacia arriba)
        CIS_BOT  = GND_Y + AXLE_R*0.3
        CIS_H    = H * 0.40
        CIS_TOP  = CIS_BOT + CIS_H
        PAD_L    = W * 0.03
        CAB_W    = W * 0.17
        CAB_X    = x + PAD_L
        CIS_X    = CAB_X + CAB_W + W*0.01
        CIS_W    = W - PAD_L - CAB_W - W*0.01 - W*0.03
        COMP_W   = CIS_W / n
        # Dome radius capped: no wider than 18% of compartment, and max 40% of cistern height
        DOME_R   = min(CIS_H * 0.40, COMP_W * 0.18)
        ch_h     = H * 0.04

        # Chasis
        c.setFillColor(colors.HexColor("#4A5568")); c.setLineWidth(0)
        c.rect(CAB_X+CAB_W*0.3, CIS_BOT-ch_h, CIS_W+CIS_X-CAB_X-CAB_W*0.3, ch_h, fill=1, stroke=0)

        # Compartimentos
        PROD_COLS = [
            "#C0392B","#E67E22","#27AE60","#D4AC0D","#784212","#2471A3","#8E44AD","#E74C3C",
        ]
        for i, tn in enumerate(tank_names[:n]):
            cx2 = CIS_X + i*COMP_W; cw2 = COMP_W
            is_first = (i==0); is_last = (i==n-1)
            bx1 = cx2 + (DOME_R if is_first else 0)
            bx2 = cx2 + cw2 - (DOME_R if is_last else 0)

            # Gradiente acero
            STRIPS = ["#D5D8DC","#CDD4D8","#C5CDD2","#BFC8CC","#BAC3C8","#D2D8DC"]
            sh = CIS_H / len(STRIPS)
            for si, sc in enumerate(STRIPS):
                c.setFillColor(colors.HexColor(sc)); c.setLineWidth(0)
                c.rect(bx1, CIS_BOT+si*sh, bx2-bx1, sh+1, fill=1, stroke=0)
            c.setFillColor(colors.white); c.setStrokeColor(colors.HexColor("#7F8C8D")); c.setLineWidth(1.5)
            c.rect(bx1, CIS_BOT, bx2-bx1, CIS_H, fill=0, stroke=1)

            # Cabezas esféricas
            if is_first:
                for ki, kc in enumerate(["#A9B2BC","#B8C4CC","#CDD4D8","#D5DCE0"]):
                    mg = ki*1.2
                    c.setFillColor(colors.HexColor(kc)); c.setLineWidth(0)
                    c.ellipse(cx2+mg, CIS_BOT+mg, cx2+2*DOME_R-mg, CIS_BOT+CIS_H-mg, fill=1, stroke=0)
                c.setFillColor(colors.white); c.setStrokeColor(colors.HexColor("#7F8C8D")); c.setLineWidth(1.5)
                c.ellipse(cx2, CIS_BOT, cx2+2*DOME_R, CIS_BOT+CIS_H, fill=0, stroke=1)
            if is_last:
                for ki, kc in enumerate(["#D5DCE0","#CDD4D8","#B8C4CC","#A9B2BC"]):
                    mg = ki*1.2
                    c.setFillColor(colors.HexColor(kc)); c.setLineWidth(0)
                    c.ellipse(cx2+cw2-2*DOME_R+mg, CIS_BOT+mg, cx2+cw2-mg, CIS_BOT+CIS_H-mg, fill=1, stroke=0)
                c.setFillColor(colors.white); c.setStrokeColor(colors.HexColor("#7F8C8D")); c.setLineWidth(1.5)
                c.ellipse(cx2+cw2-2*DOME_R, CIS_BOT, cx2+cw2, CIS_BOT+CIS_H, fill=0, stroke=1)

            # Llenado
            vp, wp = self._get_fill_pct(tn, etapa_key)
            pnm = self.get_var(f"{etapa_key}_{tn}_prod_name").get() if etapa_key else ""
            ci_c = abs(hash(pnm)) % len(PROD_COLS) if pnm else i % len(PROD_COLS)
            prod_col = colors.HexColor(PROD_COLS[ci_c])
            if vp > 0.02:
                fill_y = CIS_BOT; fill_h2 = CIS_H * vp
                # Fill ONLY the body rect area (between domes) to avoid clip bleeding
                # Agua
                if wp > 0:
                    c.setFillColor(colors.HexColor("#5DADE2")); c.setLineWidth(0)
                    c.rect(bx1+1, CIS_BOT+1, max(1, bx2-bx1-2), CIS_H*wp-1, fill=1, stroke=0)
                    fill_y = CIS_BOT + CIS_H*wp; fill_h2 = CIS_H*(vp-wp)
                # Producto (body area only)
                c.setFillColor(prod_col); c.setLineWidth(0)
                c.rect(bx1+1, fill_y, max(1, bx2-bx1-2), fill_h2, fill=1, stroke=0)
                # Reflejo
                c.setFillColor(colors.Color(1,1,1,alpha=0.35))
                c.rect(bx1+1, fill_y+fill_h2*0.85, max(1, bx2-bx1-2), fill_h2*0.10, fill=1, stroke=0)
                # Recontorno body
                c.setFillColor(colors.white); c.setStrokeColor(colors.HexColor("#7F8C8D")); c.setLineWidth(1.5)
                c.rect(bx1, CIS_BOT, bx2-bx1, CIS_H, fill=0, stroke=1)
                # Recontorno domes
                if is_first:
                    c.setFillColor(colors.white); c.setStrokeColor(colors.HexColor("#7F8C8D")); c.setLineWidth(1.5)
                    c.ellipse(cx2, CIS_BOT, cx2+2*DOME_R, CIS_BOT+CIS_H, fill=0, stroke=1)
                if is_last:
                    c.setFillColor(colors.white); c.setStrokeColor(colors.HexColor("#7F8C8D")); c.setLineWidth(1.5)
                    c.ellipse(cx2+cw2-2*DOME_R, CIS_BOT, cx2+cw2, CIS_BOT+CIS_H, fill=0, stroke=1)
                c.setFillColor(colors.white if vp>0.35 else colors.HexColor("#1B3A5C"))
                c.setFont("Helvetica-Bold", max(4.5, min(6.5, COMP_W*0.12)))
                c.drawCentredString((bx1+bx2)/2, CIS_BOT+CIS_H*vp*0.45, f"{vp*100:.0f}%")

            # Divisor entre compartimentos
            if not is_last:
                c.setFillColor(colors.HexColor("#5D6D7E")); c.setLineWidth(0)
                c.rect(bx2-1.5, CIS_BOT, 3, CIS_H, fill=1, stroke=0)

            # Boca de hombre
            mh_r = max(4, COMP_W*0.10); mh_x = (bx1+bx2)/2
            c.setFillColor(colors.HexColor("#B2BEC3")); c.setStrokeColor(colors.HexColor("#636E72")); c.setLineWidth(1)
            c.circle(mh_x, CIS_TOP, mh_r, fill=1, stroke=1)
            c.setFillColor(colors.HexColor("#C0C8CC")); c.setLineWidth(0)
            c.circle(mh_x, CIS_TOP, mh_r*0.6, fill=1, stroke=0)
            # PRV
            c.setFillColor(colors.HexColor("#E74C3C")); c.setLineWidth(0)
            c.rect(mh_x+mh_r+2, CIS_TOP-3, 4, max(6, H*0.05), fill=1, stroke=0)

            # Etiqueta compartimento
            short = tn.replace("COMPARTIMENTO ","C.").replace("TK ","").strip()[:5]
            c.setFillColor(colors.HexColor("#2C3E50")); c.setFont("Helvetica-Bold", max(4, min(5.5, COMP_W*0.10)))
            c.drawCentredString((bx1+bx2)/2, CIS_TOP+4, short)
            if pnm:
                c.setFont("Helvetica", 4); c.setFillColor(colors.HexColor("#5D6D7E"))
                c.drawCentredString((bx1+bx2)/2, CIS_BOT+CIS_H*0.08, pnm[:10])

        # Tubería inferior (manifold)
        c.setStrokeColor(colors.HexColor("#4A5568")); c.setLineWidth(3)
        c.line(CIS_X+DOME_R, CIS_BOT-ch_h*0.5, CIS_X+CIS_W-DOME_R, CIS_BOT-ch_h*0.5)

        # ── CABINA ──────────────────────────────────────────────────────────
        CAB_BOT  = CIS_BOT; CAB_H = CIS_H + H*0.15; CAB_TOP = CAB_BOT + CAB_H
        slope    = CAB_W * 0.14

        # Cuerpo cabina
        c.setFillColor(colors.HexColor("#4A5568")); c.setLineWidth(0)
        self._pdf_polygon(c,[CAB_X, CAB_BOT, CAB_X+CAB_W, CAB_BOT, CAB_X+CAB_W, CAB_BOT+CAB_H*0.82,
                   CAB_X+CAB_W-slope, CAB_TOP, CAB_X+slope, CAB_TOP, CAB_X, CAB_BOT+CAB_H*0.82],
                  fill=1)
        c.setFillColor(colors.white); c.setStrokeColor(colors.HexColor("#2C3E50")); c.setLineWidth(1.5)
        self._pdf_polygon(c,[CAB_X, CAB_BOT, CAB_X+CAB_W, CAB_BOT, CAB_X+CAB_W, CAB_BOT+CAB_H*0.82,
                   CAB_X+CAB_W-slope, CAB_TOP, CAB_X+slope, CAB_TOP, CAB_X, CAB_BOT+CAB_H*0.82],
                  fill=0)

        # Parabrisas
        wm = CAB_W*0.13; wh = CAB_H*0.32; wy = CAB_BOT+CAB_H*0.60
        c.setFillColor(colors.HexColor("#AED6F1")); c.setStrokeColor(colors.HexColor("#2C3E50")); c.setLineWidth(1)
        c.rect(CAB_X+wm, wy, CAB_W-2*wm, wh, fill=1, stroke=1)
        # División
        c.setLineWidth(1); c.line(CAB_X+CAB_W/2, wy, CAB_X+CAB_W/2, wy+wh)
        # Reflejo parabrisas
        c.setFillColor(colors.Color(1,1,1,alpha=0.5)); c.setLineWidth(0)
        _pt2 = c.beginPath(); _pt2.moveTo(CAB_X+wm+1, wy+wh-1); _pt2.lineTo(CAB_X+wm+CAB_W*0.28, wy+wh-1); _pt2.lineTo(CAB_X+wm+1, wy+wh*0.45); _pt2.close()
        c.drawPath(_pt2, fill=1, stroke=0)

        # Parrilla
        gr_h = CAB_H*0.22; gr_x = CAB_X-CAB_W*0.22; gr_y = CAB_BOT+gr_h*0.15
        c.setFillColor(colors.HexColor("#2C3E50")); c.setLineWidth(0)
        c.rect(gr_x, gr_y, CAB_W*0.22, gr_h, fill=1, stroke=0)
        for li in range(4):
            gy2 = gr_y + li*gr_h/4
            c.setFillColor(colors.HexColor("#4A5568")); c.rect(gr_x+1, gy2+1, CAB_W*0.22-2, gr_h/4-2, fill=1, stroke=0)
        # Faro
        c.setFillColor(colors.HexColor("#F9E79F")); c.setStrokeColor(colors.HexColor("#D4AC0D")); c.setLineWidth(0.8)
        hl_r = gr_h*0.3
        c.circle(gr_x+CAB_W*0.11, gr_y+gr_h-hl_r-2, hl_r, fill=1, stroke=1)
        # Espejo lateral
        c.setFillColor(colors.HexColor("#5D6D7E")); c.setLineWidth(0)
        c.rect(CAB_X+CAB_W+2, wy+wh*0.2, CAB_W*0.09, wh*0.33, fill=1, stroke=0)
        # Caño de escape (clamped to stay within drawing area)
        ex_x = CAB_X+CAB_W-CAB_W*0.09; ex_h = CAB_H*0.45
        ex_top = min(CAB_TOP + ex_h, y+H-15)
        c.setFillColor(colors.HexColor("#4A5568")); c.rect(ex_x-3, CAB_TOP, 6, ex_top-CAB_TOP, fill=1, stroke=0)
        c.setFillColor(colors.HexColor("#3D4B56")); c.circle(ex_x, ex_top, 5, fill=1, stroke=0)

        # ── RUEDAS ──────────────────────────────────────────────────────────
        AXLE_Y   = GND_Y + AXLE_R
        wheel_xs = [CAB_X+CAB_W*0.62]
        rc2 = CIS_X+CIS_W*0.72
        wheel_xs += [rc2-AXLE_R*0.4, rc2+AXLE_R*0.4]
        if CIS_W > W*0.4:
            wheel_xs.insert(1, CIS_X+CIS_W*0.44)
        for wx in wheel_xs:
            # Neumático
            c.setFillColor(colors.HexColor("#1B2631")); c.setStrokeColor(colors.HexColor("#0E1A27")); c.setLineWidth(1)
            c.circle(wx, AXLE_Y, AXLE_R, fill=1, stroke=1)
            # Llanta
            rim_r = AXLE_R*0.58
            c.setFillColor(colors.HexColor("#7F8C8D")); c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(0.8)
            c.circle(wx, AXLE_Y, rim_r, fill=1, stroke=1)
            # Rayos
            c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(0.8)
            for ang_d in [0, 60, 120, 180, 240, 300]:
                ar = math.radians(ang_d)
                c.line(wx+math.cos(ar)*rim_r*0.25, AXLE_Y+math.sin(ar)*rim_r*0.25,
                       wx+math.cos(ar)*rim_r*0.9,  AXLE_Y+math.sin(ar)*rim_r*0.9)
            # Centro
            c.setFillColor(colors.HexColor("#4A5568")); c.setLineWidth(0)
            c.circle(wx, AXLE_Y, AXLE_R*0.12, fill=1, stroke=0)
            # Guardabarro
            c.setFillColor(colors.white); c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(1)
            c.arc(wx-AXLE_R-2, AXLE_Y-AXLE_R-2, wx+AXLE_R+2, AXLE_Y+AXLE_R+2, startAng=30, extent=120)

        # ── HAZMAT PLACARD (diamond) ─────────────────────────────────────────
        plaq_cx = CIS_X + CIS_W - DOME_R - 10  # near right end of cistern body
        plaq_cy = (CIS_BOT + CIS_TOP) / 2      # vertically centred on tank
        plaq_s  = 8.5                           # half-diagonal (~12×12 pt visible)
        c.saveState()
        c.setFillColor(colors.HexColor("#E74C3C"))
        c.setStrokeColor(colors.white)
        c.setLineWidth(1.2)
        _ph = c.beginPath()
        _ph.moveTo(plaq_cx,          plaq_cy + plaq_s)  # top
        _ph.lineTo(plaq_cx + plaq_s, plaq_cy)            # right
        _ph.lineTo(plaq_cx,          plaq_cy - plaq_s)  # bottom
        _ph.lineTo(plaq_cx - plaq_s, plaq_cy)            # left
        _ph.close()
        c.drawPath(_ph, fill=1, stroke=1)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 7)
        c.drawCentredString(plaq_cx, plaq_cy - 2.5, "3")
        c.restoreState()
        # ─────────────────────────────────────────────────────────────────────

        c.restoreState()

    def _pdf_draw_pressure_truck(self, c, x, y, W, H, etapa_key, title):
        """ReportLab pressure vessel truck GLP — recipiente toriesférico premium."""
        from reportlab.lib import colors
        import math
        tank_names = self.lista_tanques or ["TK 1"]
        n = max(len(tank_names), 1)
        pat = self.get_var("car_patente").get() or ""

        c.saveState()
        c.setFillColor(colors.HexColor("#F5F7FA")); c.setStrokeColor(colors.HexColor("#D4AC0D")); c.setLineWidth(2)
        c.rect(x, y, W, H, fill=1, stroke=1)
        # Suelo simple
        c.setFillColor(colors.HexColor("#808B96")); c.setLineWidth(0)
        c.rect(x+1, y+2, W-2, H*0.10, fill=1, stroke=0)
        c.setFillColor(colors.HexColor("#784212")); c.setFont("Helvetica-Bold", min(8, max(6, H*0.04)))
        hdr = f"CAMION GAS/GLP - {title}" + (f" | Patente: {pat}" if pat else "")
        max_chars = int(W / 4.5)
        if len(hdr) > max_chars:
            hdr = hdr[:max_chars-3] + "..."
        c.drawCentredString(x+W/2, y+H-10, hdr)

        AXLE_R  = max(7, H*0.085); GND_Y = y+H*0.12+AXLE_R*2
        CIS_BOT = GND_Y+AXLE_R*0.3; CIS_H = H*0.40; CIS_TOP = CIS_BOT+CIS_H
        PAD_L   = W*0.03; CAB_W = W*0.17; CAB_X = x+PAD_L
        CIS_X   = CAB_X+CAB_W+W*0.01; CIS_W = W-PAD_L-CAB_W-W*0.01-W*0.03
        COMP_W  = CIS_W/n
        # Cap dome radius: max 22% of cistern height and 15% of compartment width, hard max 25pt
        DOME_R  = min(CIS_H*0.22, COMP_W*0.15, 25)
        ch_h = H*0.04

        c.setFillColor(colors.HexColor("#6E4B00")); c.setLineWidth(0)
        c.rect(CAB_X+CAB_W*0.3, CIS_BOT-ch_h, CIS_X+CIS_W-CAB_X-CAB_W*0.3, ch_h, fill=1, stroke=0)

        for i, tn in enumerate(tank_names[:n]):
            cx2 = CIS_X+i*COMP_W; cw2 = COMP_W
            is_first=(i==0); is_last=(i==n-1)
            bx1 = cx2+(DOME_R if is_first else 0); bx2 = cx2+cw2-(DOME_R if is_last else 0)

            # Gradiente dorado (acero de alta presión)
            STRIPS = ["#FEF9E7","#FDEBD0","#FDEBD0","#FCEBD0","#FAE5C0","#F5DBA8"]
            sh = CIS_H/len(STRIPS)
            for si, sc in enumerate(STRIPS):
                c.setFillColor(colors.HexColor(sc)); c.setLineWidth(0)
                c.rect(bx1, CIS_BOT+si*sh, bx2-bx1, sh+1, fill=1, stroke=0)
            c.setFillColor(colors.white); c.setStrokeColor(colors.HexColor("#D4AC0D")); c.setLineWidth(2)
            c.rect(bx1, CIS_BOT, bx2-bx1, CIS_H, fill=0, stroke=1)

            # Cabezas hemisféricas doradas
            for is_side, start_x in [(is_first, cx2), (is_last, cx2+cw2-2*DOME_R)]:
                if not is_side: continue
                hemi_cols = ["#B8860B","#C9960C","#D4AC0D","#E8C44B","#F0D060","#FEF9E7"]
                if not is_first: hemi_cols = list(reversed(hemi_cols))
                for ki, kc in enumerate(hemi_cols):
                    mg = ki*1.2
                    c.setFillColor(colors.HexColor(kc)); c.setLineWidth(0)
                    c.ellipse(start_x+mg, CIS_BOT+mg, start_x+2*DOME_R-mg, CIS_BOT+CIS_H-mg, fill=1, stroke=0)
                c.setFillColor(colors.white); c.setStrokeColor(colors.HexColor("#D4AC0D")); c.setLineWidth(2)
                c.ellipse(start_x, CIS_BOT, start_x+2*DOME_R, CIS_BOT+CIS_H, fill=0, stroke=1)

            # Anillos de refuerzo
            n_rings = max(2, int(cw2/40))
            c.setStrokeColor(colors.HexColor("#B8860B")); c.setFillColor(colors.HexColor("#C9960C"))
            for ri in range(1, n_rings+1):
                rx = bx1+(bx2-bx1)*ri/(n_rings+1)
                c.setLineWidth(0); c.rect(rx-2, CIS_BOT, 4, CIS_H, fill=1, stroke=0)

            # Llenado GLP
            try:
                vol = self.parse_float(self.get_var(f"{etapa_key}_{tn}_vol_nat_prod").get() or "0")
                cap = self.parse_float(self.get_var(f"{etapa_key}_{tn}_alt_ref").get() or "1")
                vp  = min(max(vol/(cap if cap>0 else 1),0),1)
            except: vp = 0.0
            if vp > 0.02:
                fill_h = CIS_H * vp
                # Fill ONLY the body rect area (between domes) — no bleeding
                c.setFillColor(colors.HexColor("#FF6B35")); c.setLineWidth(0)
                c.rect(bx1+1, CIS_BOT+1, max(1, bx2-bx1-2), fill_h-2, fill=1, stroke=0)
                # Reflejo
                c.setFillColor(colors.Color(1,1,1,alpha=0.35))
                c.rect(bx1+1, CIS_BOT+fill_h*0.85, max(1, bx2-bx1-2), fill_h*0.10, fill=1, stroke=0)
                # Redraw body contour over fill
                c.setFillColor(colors.white); c.setStrokeColor(colors.HexColor("#D4AC0D")); c.setLineWidth(2)
                c.rect(bx1, CIS_BOT, bx2-bx1, CIS_H, fill=0, stroke=1)
                for is_side_r, start_x_r in [(is_first, cx2), (is_last, cx2+cw2-2*DOME_R)]:
                    if is_side_r:
                        c.ellipse(start_x_r, CIS_BOT, start_x_r+2*DOME_R, CIS_BOT+CIS_H, fill=0, stroke=1)
                c.setFillColor(colors.white); c.setFont("Helvetica-Bold", max(4,min(6, COMP_W*0.12)))
                c.drawCentredString((bx1+bx2)/2, CIS_BOT+CIS_H*vp*0.45, f"{vp*100:.0f}%")

            # PRV (válvulas de seguridad) en la cima
            pnm = self.get_var(f"{etapa_key}_{tn}_prod_name").get() if etapa_key else ""
            pres = self.get_var(f"{etapa_key}_{tn}_esf_pres").get() if etapa_key else ""
            for vvi in range(2):
                vvx = bx1+(bx2-bx1)*(vvi+1)/3
                c.setFillColor(colors.HexColor("#E74C3C")); c.setLineWidth(0)
                c.rect(vvx-2.5, CIS_TOP, 5, max(5,H*0.04), fill=1, stroke=0)
                c.rect(vvx-4, CIS_TOP+max(5,H*0.04)-1, 8, 3, fill=1, stroke=0)

            # Manómetro
            g_x = (bx1+bx2)/2; g_r = max(7, H*0.065); g_y = CIS_TOP+g_r+max(3,H*0.02)
            c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(1)
            c.line(g_x, CIS_TOP, g_x, g_y-g_r)
            c.setFillColor(colors.white); c.setStrokeColor(colors.HexColor("#E74C3C")); c.setLineWidth(1.5)
            c.circle(g_x, g_y, g_r, fill=1, stroke=1)
            try: p_frac = min(float((pres or "0").replace(",","."))/1000, 1.0)
            except: p_frac = 0.3
            ang = math.radians(-30+240*p_frac)
            c.setStrokeColor(colors.HexColor("#C0392B")); c.setLineWidth(1.2)
            c.line(g_x, g_y, g_x+math.cos(ang)*g_r*0.75, g_y+math.sin(ang)*g_r*0.75)
            if pres:
                c.setFillColor(colors.HexColor("#784212")); c.setFont("Helvetica", 4)
                c.drawCentredString(g_x, g_y-g_r-4, f"{pres[:5]}kPa")

            # Etiqueta
            short = tn.replace("COMPARTIMENTO ","C.").strip()[:5]
            c.setFillColor(colors.HexColor("#784212")); c.setFont("Helvetica-Bold", max(4,min(5.5,COMP_W*0.10)))
            c.drawCentredString((bx1+bx2)/2, CIS_TOP+g_r*2+6, short)
            if pnm:
                c.setFont("Helvetica", 4); c.setFillColor(colors.HexColor("#8E6915"))
                c.drawCentredString((bx1+bx2)/2, CIS_BOT+CIS_H*0.08, pnm[:10])

        # ── CABINA ─────────────────────────────────────────────────────────
        CAB_BOT=CIS_BOT; CAB_H=CIS_H+H*0.15; CAB_TOP=CAB_BOT+CAB_H; slope=CAB_W*0.14
        c.setFillColor(colors.HexColor("#6E4B00")); c.setLineWidth(0)
        self._pdf_polygon(c,[CAB_X,CAB_BOT,CAB_X+CAB_W,CAB_BOT,CAB_X+CAB_W,CAB_BOT+CAB_H*0.82,
                   CAB_X+CAB_W-slope,CAB_TOP,CAB_X+slope,CAB_TOP,CAB_X,CAB_BOT+CAB_H*0.82],fill=1)
        c.setFillColor(colors.white); c.setStrokeColor(colors.HexColor("#5D3D00")); c.setLineWidth(1.5)
        self._pdf_polygon(c,[CAB_X,CAB_BOT,CAB_X+CAB_W,CAB_BOT,CAB_X+CAB_W,CAB_BOT+CAB_H*0.82,
                   CAB_X+CAB_W-slope,CAB_TOP,CAB_X+slope,CAB_TOP,CAB_X,CAB_BOT+CAB_H*0.82],fill=0)
        wm=CAB_W*0.13; wh=CAB_H*0.32; wy=CAB_BOT+CAB_H*0.60
        c.setFillColor(colors.HexColor("#AED6F1")); c.setStrokeColor(colors.HexColor("#2C3E50")); c.setLineWidth(1)
        c.rect(CAB_X+wm,wy,CAB_W-2*wm,wh,fill=1,stroke=1)
        c.setLineWidth(1); c.line(CAB_X+CAB_W/2,wy,CAB_X+CAB_W/2,wy+wh)
        gr_h=CAB_H*0.22; gr_x=CAB_X-CAB_W*0.22; gr_y=CAB_BOT+gr_h*0.15
        c.setFillColor(colors.HexColor("#5D3D00")); c.setLineWidth(0)
        c.rect(gr_x,gr_y,CAB_W*0.22,gr_h,fill=1,stroke=0)
        c.setFillColor(colors.HexColor("#F9E79F")); c.setStrokeColor(colors.HexColor("#D4AC0D")); c.setLineWidth(0.8)
        c.circle(gr_x+CAB_W*0.11, gr_y+gr_h-gr_h*0.35-2, gr_h*0.3, fill=1, stroke=1)

        # Ruedas
        AXLE_Y=GND_Y+AXLE_R
        for wx in [CAB_X+CAB_W*0.62, CIS_X+CIS_W*0.72-AXLE_R*0.4, CIS_X+CIS_W*0.72+AXLE_R*0.4]:
            c.setFillColor(colors.HexColor("#1B2631")); c.setStrokeColor(colors.HexColor("#0E1A27")); c.setLineWidth(1)
            c.circle(wx,AXLE_Y,AXLE_R,fill=1,stroke=1)
            rim_r=AXLE_R*0.58
            c.setFillColor(colors.HexColor("#7F8C8D")); c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(0.8)
            c.circle(wx,AXLE_Y,rim_r,fill=1,stroke=1)
            c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(0.8)
            for ang_d in [0,60,120,180,240,300]:
                ar=math.radians(ang_d)
                c.line(wx+math.cos(ar)*rim_r*0.25,AXLE_Y+math.sin(ar)*rim_r*0.25,
                       wx+math.cos(ar)*rim_r*0.9,AXLE_Y+math.sin(ar)*rim_r*0.9)
            c.setFillColor(colors.HexColor("#4A5568")); c.setLineWidth(0); c.circle(wx,AXLE_Y,AXLE_R*0.12,fill=1,stroke=0)
        c.restoreState()

    def _pdf_draw_pipeline(self, c, x, y, W, H, etapa_key, title):
        """ReportLab pipeline drawing — versión industrial premium con gradientes y detalles."""
        from reportlab.lib import colors
        import math as _pm
        tm = self.get_tipo_medio()
        instalacion = self.get_var("car_buque").get() or ""
        STYLES = {
            "GASODUCTO": ("#D6EAF8","#1A5276","#AED6F1","#2874A6","#1A5276"),
            "OLEODUCTO": ("#FEF9E7","#784212","#F0B429","#D4880A","#784212"),
            "POLIDUCTO": ("#EAFAF1","#1D6A39","#A9DFBF","#27AE60","#1D6A39"),
        }
        bg, tc, pf, po, dark = STYLES.get(tm, STYLES["GASODUCTO"])

        c.saveState()
        # ── Fondo neutro ──────────────────────────────────────────────────
        c.setFillColor(colors.HexColor("#F5F7FA")); c.setLineWidth(0)
        c.rect(x, y, W, H, fill=1, stroke=0)
        # Border
        c.setStrokeColor(colors.HexColor(dark)); c.setLineWidth(1.5)
        c.rect(x, y, W, H, fill=0, stroke=1)
        # Header bar with title
        hdr_h = 14
        c.setFillColor(colors.HexColor(tc)); c.setLineWidth(0)
        c.rect(x+1, y+H-hdr_h, W-2, hdr_h-1, fill=1, stroke=0)

        # ── Terreno (simple) ──────────────────────────────────────────────
        GROUND_Y = y + H * 0.22
        c.setFillColor(colors.HexColor("#D5D8DC")); c.setLineWidth(0)
        c.rect(x+1, y+1, W-2, GROUND_Y-y-2, fill=1, stroke=0)

        # ── Parámetros de la tubería ──────────────────────────────────────────
        PIPE_Y  = y + H*0.55
        PIPE_R  = max(9, H*0.13)
        PAD_L   = max(28, W*0.08); PAD_R = PAD_L
        P_LEFT  = x + PAD_L; P_RIGHT = x + W - PAD_R
        PIPE_W  = P_RIGHT - P_LEFT

        n_pipes = 2 if "POLI" in tm else 1
        pipe_offsets = [0] if n_pipes == 1 else [-PIPE_R-3, PIPE_R+3]

        for pi, p_off in enumerate(pipe_offsets):
            py = PIPE_Y + p_off
            _rc = int(pf[1:3],16); _gc = int(pf[3:5],16); _bc = int(pf[5:7],16)
            pc_use = pf if pi == 0 else "#7D3C98"
            if pi == 1: _rc,_gc,_bc = 125,60,152

            # Cuerpo principal
            c.setFillColor(colors.HexColor(pc_use)); c.setLineWidth(0)
            c.rect(P_LEFT, py-PIPE_R, PIPE_W, PIPE_R*2, fill=1, stroke=0)
            # Gradiente cilíndrico
            strips = [(0.00,0.10,0.60),(0.10,0.22,0.75),(0.22,0.38,0.95),(0.38,0.52,1.25),(0.52,0.65,1.08),(0.65,0.80,0.80),(0.80,1.00,0.53)]
            for s1,s2,bright in strips:
                rr2=min(255,int(_rc*bright)); gg2=min(255,int(_gc*bright)); bb2=min(255,int(_bc*bright))
                sc2=f"#{rr2:02x}{gg2:02x}{bb2:02x}"
                c.setFillColor(colors.HexColor(sc2)); c.setLineWidth(0)
                c.rect(P_LEFT+2, py-PIPE_R+int(s1*2*PIPE_R), PIPE_W-4, max(1,int((s2-s1)*2*PIPE_R)), fill=1, stroke=0)
            # Brillo especular
            hl_y = py-PIPE_R+max(2,PIPE_R//5)
            c.setFillColor(colors.white); c.setLineWidth(0)
            c.rect(P_LEFT+PIPE_W//8, hl_y, PIPE_W*3//4, max(1,PIPE_R//6), fill=1, stroke=0)
            # Contorno
            c.setStrokeColor(colors.HexColor(dark)); c.setLineWidth(1.5)
            c.rect(P_LEFT, py-PIPE_R, PIPE_W, PIPE_R*2, fill=0, stroke=1)
            # Tapas elípticas
            c.setFillColor(colors.HexColor("#4A5568")); c.setStrokeColor(colors.HexColor(dark)); c.setLineWidth(1.2)
            c.ellipse(P_LEFT-PIPE_R//2, py-PIPE_R, P_LEFT+PIPE_R//2, py+PIPE_R, fill=1, stroke=1)
            c.ellipse(P_RIGHT-PIPE_R//2, py-PIPE_R, P_RIGHT+PIPE_R//2, py+PIPE_R, fill=1, stroke=1)
            # Soportes al suelo
            for sx in [P_LEFT+PIPE_W//4, P_LEFT+PIPE_W//2, P_LEFT+3*PIPE_W//4]:
                c.setFillColor(colors.HexColor("#7F8C8D")); c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(0.8)
                c.rect(sx-3, GROUND_Y, 6, max(1, (py - PIPE_R) - GROUND_Y), fill=1, stroke=1)
                c.setFillColor(colors.HexColor("#5D6D7E")); c.setLineWidth(0)
                c.rect(sx-7, GROUND_Y-3, 14, 4, fill=1, stroke=0)
            # Bridas
            n_fl = max(2, min(5, int(PIPE_W//60)))
            fl_step = PIPE_W//(n_fl+1)
            for fi in range(n_fl):
                fx = P_LEFT + (fi+1)*fl_step
                c.setFillColor(colors.HexColor("#5D6D7E")); c.setStrokeColor(colors.HexColor("#4A5568")); c.setLineWidth(0.8)
                c.rect(fx-2, py-PIPE_R-3, 4, PIPE_R*2+6, fill=1, stroke=1)

        # ── Tramo subterráneo ────────────────────────────────────────────────
        sg_w = max(20, PIPE_W//8)
        sg_x = x + W - PAD_R - sg_w - 8
        ug_depth = max(8, PIPE_R*2)
        c.setFillColor(colors.HexColor("#9E8A6E")); c.setLineWidth(0)
        c.rect(sg_x-4, y+1, sg_w+14, GROUND_Y-y, fill=1, stroke=0)
        c.setFillColor(colors.HexColor("#7A6350")); c.setLineWidth(0)
        c.rect(sg_x-2, y+2, sg_w+10, ug_depth*0.4, fill=1, stroke=0)
        py0 = PIPE_Y + pipe_offsets[0]
        c.setFillColor(colors.HexColor(pf)); c.setStrokeColor(colors.HexColor(dark)); c.setLineWidth(1.5)
        c.rect(sg_x, py0-PIPE_R, sg_w, PIPE_R*2, fill=1, stroke=1)
        # Tapa izquierda
        c.setFillColor(colors.HexColor("#4A5568"))
        c.ellipse(sg_x-PIPE_R//2, py0-PIPE_R, sg_x+PIPE_R//2, py0+PIPE_R, fill=1, stroke=1)
        c.setFillColor(colors.HexColor("#2C3E50")); c.setFont("Helvetica", max(4,int(PIPE_R//3)))
        c.drawCentredString(sg_x+sg_w//2, (GROUND_Y + y+2) / 2, "─ ─ ─")

        # ── Flechas de dirección de flujo ────────────────────────────────────
        arr_y = PIPE_Y + pipe_offsets[0]
        c.setFillColor(colors.HexColor(tc)); c.setStrokeColor(colors.HexColor(tc)); c.setLineWidth(0)
        for afx in [P_LEFT+PIPE_W//5, P_LEFT+PIPE_W*2//5, P_LEFT+PIPE_W*3//5]:
            arr = c.beginPath()
            arr.moveTo(afx-4, arr_y-2); arr.lineTo(afx+5, arr_y); arr.lineTo(afx-4, arr_y+2); arr.close()
            c.drawPath(arr, fill=1, stroke=0)

        # ── Instrumento de presión ────────────────────────────────────────────
        inst_x = P_LEFT + PIPE_W//2; inst_y = PIPE_Y + pipe_offsets[0] - PIPE_R
        inst_r = max(6, PIPE_R*0.65)
        # Línea de conexión
        c.setStrokeColor(colors.HexColor(dark)); c.setLineWidth(1.2)
        c.line(inst_x, inst_y, inst_x, inst_y-inst_r*0.6)
        # Cuerpo del manómetro
        c.setFillColor(colors.HexColor("#D5DBDB")); c.setStrokeColor(colors.HexColor(dark)); c.setLineWidth(1)
        c.circle(inst_x, inst_y-inst_r*0.6-inst_r, inst_r, fill=1, stroke=1)
        c.setFillColor(colors.white); c.setLineWidth(0)
        c.circle(inst_x, inst_y-inst_r*0.6-inst_r, inst_r*0.7, fill=1, stroke=0)
        # Aguja
        c.setStrokeColor(colors.HexColor("#E74C3C")); c.setLineWidth(0.8)
        ang_p = _pm.radians(45)
        c.line(inst_x, inst_y-inst_r*0.6-inst_r,
               inst_x+inst_r*0.55*_pm.cos(ang_p), inst_y-inst_r*0.6-inst_r+inst_r*0.55*_pm.sin(ang_p))
        c.setFillColor(colors.HexColor("#1B3A5C")); c.setFont("Helvetica-Bold", max(4,int(inst_r*0.55)))
        c.drawCentredString(inst_x, inst_y-inst_r*0.6-inst_r*2-3, "P")

        # ── Válvula en el extremo ─────────────────────────────────────────────
        vx = P_LEFT + PIPE_W//6; vy = PIPE_Y + pipe_offsets[0]
        vr = max(4, PIPE_R*0.75)
        c.setFillColor(colors.HexColor("#4A5568")); c.setStrokeColor(colors.HexColor(dark)); c.setLineWidth(1)
        c.rect(vx-2, vy-vr, 4, vr*2, fill=1, stroke=1)
        # Forma de mariposa
        v_pts = [vx-vr, vy-1, vx-2, vy-vr+1, vx+2, vy-vr+1, vx+vr, vy-1,
                 vx+vr, vy+1, vx+2, vy+vr-1, vx-2, vy+vr-1, vx-vr, vy+1]
        c.setFillColor(colors.HexColor("#7F8C8D")); c.setLineWidth(0)
        self._pdf_polygon(c,v_pts, fill=1)
        c.setStrokeColor(colors.HexColor(dark)); c.setLineWidth(0.8)
        self._pdf_polygon(c,v_pts, fill=0)
        # Tope de la válvula
        c.setFillColor(colors.HexColor("#2C3E50")); c.setLineWidth(0)
        c.rect(vx-3, vy-vr-5, 6, 5, fill=1, stroke=0)
        c.rect(vx-6, vy-vr-7, 12, 2, fill=1, stroke=0)

        # ── Cross-section inset (upper-left) ───────────────────────────────
        cs_r = max(12, min(22, H*0.07, W*0.035))
        cs_cx = x + PAD_L + cs_r + 8; cs_cy = y + H - hdr_h - cs_r - 8
        # Background circle
        c.setFillColor(colors.HexColor("#E8EBF0")); c.setStrokeColor(colors.HexColor(dark)); c.setLineWidth(1)
        c.circle(cs_cx, cs_cy, cs_r+4, fill=1, stroke=1)
        # Outer coating (FBE - dark green)
        c.setFillColor(colors.HexColor("#1D6A39")); c.setLineWidth(0)
        c.circle(cs_cx, cs_cy, cs_r, fill=1, stroke=0)
        # Steel wall
        c.setFillColor(colors.HexColor("#7F8C8D")); c.setLineWidth(0)
        c.circle(cs_cx, cs_cy, cs_r*0.85, fill=1, stroke=0)
        # Internal lining
        c.setFillColor(colors.HexColor("#F0E68C")); c.setLineWidth(0)
        c.circle(cs_cx, cs_cy, cs_r*0.70, fill=1, stroke=0)
        # Product flow (inside)
        flow_col = "#5D4037" if "OLEO" in tm else ("#AED6F1" if "GAS" in tm else "#82E0AA")
        c.setFillColor(colors.HexColor(flow_col)); c.setLineWidth(0)
        c.circle(cs_cx, cs_cy, cs_r*0.62, fill=1, stroke=0)
        # Flow direction arrow
        c.setFillColor(colors.white); c.setLineWidth(0)
        arr_s = cs_r * 0.3
        arr = c.beginPath()
        arr.moveTo(cs_cx-arr_s*0.6, cs_cy-arr_s*0.4); arr.lineTo(cs_cx+arr_s*0.8, cs_cy)
        arr.lineTo(cs_cx-arr_s*0.6, cs_cy+arr_s*0.4); arr.close()
        c.drawPath(arr, fill=1, stroke=0)
        # Label
        c.setFillColor(colors.HexColor(dark)); c.setFont("Helvetica-Bold", max(3.5, cs_r*0.22))
        c.drawCentredString(cs_cx, cs_cy-cs_r-7, "SECCION")

        # ── Título (centrado en la barra de header) ─────────────────────────
        c.setFillColor(colors.white); c.setFont("Helvetica-Bold", 7)
        title_txt = f"{tm}" + (f" - {instalacion}" if instalacion else "") + f" - {title}"
        max_chars = int(W / 4.2)
        if len(title_txt) > max_chars:
            title_txt = title_txt[:max_chars-3] + "..."
        c.drawCentredString(x+W/2, y+H-11, title_txt)

        c.restoreState()

    def _pdf_draw_electric(self, c, x, y, W, H, etapa_key, title):
        """ReportLab electrical metering drawing — versión industrial premium con panel de control."""
        from reportlab.lib import colors
        import math as _pm
        tank_names = self.lista_tanques or ["MED 1"]
        n = max(len(tank_names), 1)
        instalacion = self.get_var("car_buque").get() or ""

        c.saveState()
        # ── Fondo claro (sala eléctrica) ──────────────────────────────────
        c.setFillColor(colors.HexColor("#F0F2F5")); c.setLineWidth(0)
        c.rect(x, y, W, H, fill=1, stroke=0)
        # Borde
        c.setStrokeColor(colors.HexColor("#4A235A")); c.setLineWidth(1.5)
        c.rect(x, y, W, H, fill=0, stroke=1)
        # Header con color solido
        hdr_h = 14
        c.setFillColor(colors.HexColor("#4A235A")); c.setLineWidth(0)
        c.rect(x+1, y+H-hdr_h, W-2, hdr_h-1, fill=1, stroke=0)
        # Título
        c.setFillColor(colors.white); c.setFont("Helvetica-Bold", min(8, max(5, H*0.04)))
        title_txt = f"MEDICION ELECTRICA" + (f" | {instalacion}" if instalacion else "") + f" - {title}"
        max_chars = int(W / 4.5)
        if len(title_txt) > max_chars:
            title_txt = title_txt[:max_chars-3] + "..."
        c.drawCentredString(x+W/2, y+H-11, title_txt)

        # ── Panel de control principal ─────────────────────────────────────
        px = x+8; pw = W-16; py_p = y+8; ph = H-hdr_h-18
        c.setFillColor(colors.HexColor("#E8EBF0")); c.setStrokeColor(colors.HexColor("#4A235A")); c.setLineWidth(1.5)
        c.roundRect(px, py_p, pw, ph, 4, fill=1, stroke=1)
        # Tornillos de panel (4 esquinas)
        c.setFillColor(colors.HexColor("#4A5568")); c.setLineWidth(0)
        for sx2,sy2 in [(px+4,py_p+4),(px+pw-6,py_p+4),(px+4,py_p+ph-6),(px+pw-6,py_p+ph-6)]:
            c.circle(sx2, sy2, 2.5, fill=1, stroke=0)
            c.setStrokeColor(colors.HexColor("#2C3E50")); c.setLineWidth(0.6)
            c.line(sx2-1.5,sy2,sx2+1.5,sy2)

        # ── LEDs de estado ────────────────────────────────────────────────
        leds = [("#27AE60","ON"),("#F4D03F","WARN"),("#E74C3C","ALM")]
        for li2, (lc2,lt2) in enumerate(leds):
            lx2 = px+10+li2*22; ly2 = py_p+ph-10
            c.setFillColor(colors.HexColor(lc2)); c.setLineWidth(0)
            c.circle(lx2, ly2, 4, fill=1, stroke=0)
            # Halo del LED (glow effect)
            c.setFillColor(colors.HexColor(lc2))
            c.circle(lx2, ly2, 5, fill=1, stroke=0)
            c.setFillColor(colors.HexColor(lc2)); c.circle(lx2, ly2, 3.5, fill=1, stroke=0)
            c.setFillColor(colors.HexColor("#BDC3C7")); c.setFont("Helvetica", max(3,H*0.027))
            c.drawCentredString(lx2, ly2-8, lt2)

        # ── Medidores individuales ────────────────────────────────────────
        mt_w = max(24, (pw-20)//n)
        mt_pad = 4
        for i, tn in enumerate(tank_names[:n]):
            mx = px+10+i*mt_w; my_bot = py_p+8; mw = mt_w-mt_pad*2; mh = ph-22

            # Marco del medidor
            c.setFillColor(colors.HexColor("#D5D8DC")); c.setStrokeColor(colors.HexColor("#4A235A")); c.setLineWidth(1)
            c.roundRect(mx, my_bot, mw, mh, 3, fill=1, stroke=1)

            # === Proportional layout: divide available height into zones ===
            mhdr = min(12, max(8, mh*0.04))
            # Header del medidor (top of meter)
            c.setFillColor(colors.HexColor("#4A235A")); c.setLineWidth(0)
            c.roundRect(mx, my_bot+mh-mhdr, mw, mhdr, 3, fill=1, stroke=0)
            c.setFillColor(colors.white); c.setFont("Helvetica-Bold", min(6, max(4.5, mhdr*0.55)))
            c.drawCentredString(mx+mw//2, my_bot+mh-mhdr+mhdr*0.25, tn[:10])

            # Usable content area (below header, above bottom margin)
            content_bot = my_bot + 4
            content_top_y = my_bot + mh - mhdr - 3
            ch = content_top_y - content_bot  # total content height

            # Proportional zones: LCD_INI(25%) + LCD_FIN(25%) + BAR(10%) + GAUGES(30%) + gaps(10%)
            gap = max(3, ch * 0.025)
            lcd_h = max(14, ch * 0.20)
            bar_h = max(6, ch * 0.06)
            gauge_zone = max(20, ch * 0.30)
            gr = min(gauge_zone * 0.42, mw * 0.18, 22)
            label_fs = min(5, max(3.5, ch * 0.018))
            lcd_fs = max(5, min(9, mw * 0.10, lcd_h * 0.45))

            # Content data
            kwh_ini = self.get_var(f"{etapa_key}_{tn}_el_ini_act").get() if etapa_key else ""
            kwh_fin = self.get_var(f"{etapa_key}_{tn}_el_fin_act").get() if etapa_key else ""

            # ── kWh INICIAL: label + LCD ──
            cur_y = content_top_y
            c.setFillColor(colors.HexColor("#2C3E50")); c.setFont("Helvetica-Bold", label_fs)
            c.drawCentredString(mx+mw//2, cur_y - label_fs - 1, "kWh INICIAL")
            cur_y -= label_fs + 3
            lcd_y = cur_y - lcd_h
            c.setFillColor(colors.HexColor("#0D4D3E")); c.setStrokeColor(colors.HexColor("#1ABC9C")); c.setLineWidth(1)
            c.roundRect(mx+4, lcd_y, mw-8, lcd_h, 2, fill=1, stroke=1)
            c.setFillColor(colors.HexColor("#1ABC9C")); c.setFont("Courier-Bold", lcd_fs)
            c.drawCentredString(mx+mw//2, lcd_y+lcd_h*0.35, kwh_ini[:9] if kwh_ini else "---")
            cur_y = lcd_y - gap

            # ── kWh FINAL: label + LCD ──
            c.setFillColor(colors.HexColor("#2C3E50")); c.setFont("Helvetica-Bold", label_fs)
            c.drawCentredString(mx+mw//2, cur_y - label_fs - 1, "kWh FINAL")
            cur_y -= label_fs + 3
            lcd2_y = cur_y - lcd_h
            c.setFillColor(colors.HexColor("#2C0808")); c.setStrokeColor(colors.HexColor("#E74C3C")); c.setLineWidth(1)
            c.roundRect(mx+4, lcd2_y, mw-8, lcd_h, 2, fill=1, stroke=1)
            c.setFillColor(colors.HexColor("#E74C3C")); c.setFont("Courier-Bold", lcd_fs)
            c.drawCentredString(mx+mw//2, lcd2_y+lcd_h*0.35, kwh_fin[:9] if kwh_fin else "---")
            cur_y = lcd2_y - gap

            # ── Consumption bar ──
            bar_w = mw - 12
            bar_y = cur_y - bar_h
            c.setFillColor(colors.HexColor("#BDC3C7")); c.setStrokeColor(colors.HexColor("#4A235A")); c.setLineWidth(0.8)
            c.rect(mx+6, bar_y, bar_w, bar_h, fill=1, stroke=1)
            c.setFillColor(colors.HexColor("#2C3E50")); c.setFont("Helvetica-Bold", min(4, label_fs))
            c.drawCentredString(mx+mw//2, bar_y + bar_h + 1, "CONSUMO")
            try:
                ini_v = float(kwh_ini.replace(",",".")) if kwh_ini else 0.0
                fin_v = float(kwh_fin.replace(",",".")) if kwh_fin else 0.0
                if fin_v > ini_v > 0:
                    pct_b = min((fin_v-ini_v)/max(fin_v,1), 1.0)
                    bar_c = "#27AE60" if pct_b < 0.7 else ("#F39C12" if pct_b < 0.9 else "#E74C3C")
                    c.setFillColor(colors.HexColor(bar_c)); c.setLineWidth(0)
                    c.rect(mx+6, bar_y, int(bar_w*pct_b), bar_h, fill=1, stroke=0)
                    c.setFillColor(colors.white); c.setFont("Helvetica-Bold", min(5, bar_h*0.7))
                    c.drawCentredString(mx+mw//2, bar_y+bar_h*0.2, f"D{fin_v-ini_v:.0f}")
            except: pass
            cur_y = bar_y - gap

            # ── Voltímetro y Amperímetro analógicos ──
            gauge_y = cur_y - gr - 2
            # Center gauges vertically in remaining space
            remaining = gauge_y - gr - content_bot
            if remaining > gr:
                gauge_y = content_bot + (cur_y - content_bot) * 0.45 + gr * 0.5
            for gi2,(gl2,gc2,gvk) in enumerate([("V","#F4D03F","el_V"),("A","#E74C3C","el_A")]):
                gx3 = mx + mw//4 + gi2*(mw//2); gy3 = gauge_y
                c.setFillColor(colors.HexColor("#1A252F")); c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(0.8)
                c.circle(gx3, gy3, gr, fill=1, stroke=1)
                # Marcas de la esfera
                c.setStrokeColor(colors.HexColor("#85929E")); c.setLineWidth(0.5)
                for am in range(0,180,30):
                    ar2 = _pm.radians(180+am)
                    c.line(gx3+_pm.cos(ar2)*gr*0.7,gy3+_pm.sin(ar2)*gr*0.7,gx3+_pm.cos(ar2)*gr*0.9,gy3+_pm.sin(ar2)*gr*0.9)
                # Aguja
                gval = self.get_var(f"{etapa_key}_{tn}_{gvk}").get() if etapa_key else ""
                try:
                    gfrac = min(float(gval.replace(",","."))/500.0,1.0) if gval else 0.0
                except: gfrac = 0.0
                ang3 = _pm.radians(180+gfrac*180)
                c.setStrokeColor(colors.HexColor(gc2)); c.setLineWidth(1)
                c.line(gx3,gy3,gx3+_pm.cos(ang3)*gr*0.75,gy3+_pm.sin(ang3)*gr*0.75)
                c.setFillColor(colors.HexColor("#2C3E50")); c.setLineWidth(0)
                c.circle(gx3,gy3,2,fill=1,stroke=0)
                c.setFillColor(colors.HexColor(gc2)); c.setFont("Helvetica-Bold", max(4, min(6, gr*0.4)))
                c.drawCentredString(gx3, gy3-gr-max(4, min(6, gr*0.4))-1, gl2)

        # ── Barras de cables en la parte inferior ─────────────────────────
        cab_h = max(6, min(10, H*0.03))
        c.setFillColor(colors.HexColor("#17202A")); c.setLineWidth(0)
        c.rect(x+1, y+1, W-2, cab_h, fill=1, stroke=0)
        cable_cols = ["#E74C3C","#F4D03F","#3498DB","#ECF0F1","#27AE60"]
        n_cables = int(max(6, W//15))
        cw = (W-8)//n_cables
        for ci2 in range(n_cables):
            cc2 = cable_cols[ci2 % len(cable_cols)]
            c.setFillColor(colors.HexColor(cc2)); c.setLineWidth(0)
            c.roundRect(x+4+ci2*cw, y+2, max(3,cw-2), cab_h-3, 1, fill=1, stroke=0)
        c.setFillColor(colors.HexColor("#BDC3C7")); c.setFont("Helvetica", max(3, H*0.027))
        c.drawCentredString(x+W//2, y+cab_h//2, "BARRAS Y CABLEADO")

        c.restoreState()

    def _pdf_draw_gasero_ship(self, c, x, y, W, H, etapa_key, title):
        """Buque gasero/GLP — casco con hasta 4 tanques esféricos MOSS sobre cubierta."""
        from reportlab.lib import colors
        import math
        c.saveState()
        tank_names = [t for t in self.lista_tanques if not t.startswith("CARBONERA")]
        n = max(len(tank_names), 1)

        # Fondo neutro
        c.setFillColor(colors.HexColor("#E8F0FE")); c.setLineWidth(0)
        c.rect(x, y, W, H, fill=1, stroke=0)

        # Título
        c.setFont("Helvetica-Bold", min(8, max(6, H*0.04)))
        c.setFillColor(colors.HexColor("#1D6A39"))
        _gtxt = f"BUQUE GASERO/GLP - {title}"
        max_chars = int(W / 4.5)
        if len(_gtxt) > max_chars:
            _gtxt = _gtxt[:max_chars-3] + "..."
        c.drawCentredString(x+W/2, y+H-10, _gtxt)

        # Dimensiones casco
        M = 20
        hull_bot = y + H*0.10
        hull_top = y + H*0.52
        hull_h   = hull_top - hull_bot
        bow_x    = x + W - M
        stern_x  = x + M
        mid_y    = hull_bot + hull_h*0.45

        # Casco (proa puntiaguda, popa cuadrada)
        p = c.beginPath()
        p.moveTo(stern_x, hull_top)
        p.lineTo(bow_x - 20, hull_top)
        p.curveTo(bow_x+10, hull_top, bow_x+15, mid_y+8, bow_x+5, mid_y)
        p.lineTo(stern_x, mid_y)
        p.close()
        c.setFillColor(colors.HexColor("#BDC3C7")); c.setLineWidth(1.5)
        c.setStrokeColor(colors.HexColor("#7F8C8D"))
        c.drawPath(p, fill=1, stroke=1)

        # Obra viva
        p2 = c.beginPath()
        p2.moveTo(stern_x, mid_y)
        p2.lineTo(bow_x+5, mid_y)
        p2.curveTo(bow_x+15, mid_y, bow_x+10, hull_bot+5, bow_x-10, hull_bot)
        p2.lineTo(stern_x+5, hull_bot)
        p2.curveTo(stern_x-5, hull_bot, stern_x-5, mid_y-5, stern_x, mid_y)
        p2.close()
        c.setFillColor(colors.HexColor("#E74C3C")); c.drawPath(p2, fill=1, stroke=0)
        # Franja azul sobre línea de flotación
        c.setFillColor(colors.HexColor("#1B3A5C")); c.setLineWidth(0)
        c.rect(stern_x, mid_y-2, bow_x-stern_x, 4, fill=1, stroke=0)

        # Cubierta (plataforma entre caseta y proa)
        c.setFillColor(colors.HexColor("#95A5A6")); c.setLineWidth(0)
        c.rect(stern_x+50, hull_top-3, bow_x-stern_x-65, 4, fill=1, stroke=0)

        # Caseta de popa
        c.setFillColor(colors.HexColor("#ECF0F1")); c.setStrokeColor(colors.HexColor("#BDC3C7")); c.setLineWidth(0.8)
        c.roundRect(stern_x, hull_top, 52, 22, 2, fill=1, stroke=1)
        c.setFillColor(colors.HexColor("#D5D8DC"))
        c.roundRect(stern_x+4, hull_top+22, 44, 12, 2, fill=1, stroke=1)
        c.roundRect(stern_x-3, hull_top+30, 58, 6, 1, fill=1, stroke=1)
        # Chimenea corta
        c.setFillColor(colors.HexColor("#C0392B")); c.setLineWidth(0)
        c.roundRect(stern_x+18, hull_top+36, 12, 14, 2, fill=1, stroke=0)

        # ── TANQUES ESFÉRICOS MOSS ─────────────────────────────────────────────
        sphere_zone_x1 = stern_x + 55
        sphere_zone_x2 = bow_x - 15
        zone_w = sphere_zone_x2 - sphere_zone_x1
        sphere_spacing = zone_w / n
        sphere_r = min(hull_h * 0.70, sphere_spacing * 0.42)
        sphere_cy = hull_top + sphere_r * 0.80   # centro esferas sobre cubierta

        for i, tn in enumerate(tank_names[:n]):
            cx_s = sphere_zone_x1 + i * sphere_spacing + sphere_spacing / 2

            # Falda (skirt)
            skirt_w = sphere_r * 0.55
            skirt_h = sphere_r * 0.55
            c.setFillColor(colors.HexColor("#7F8C8D")); c.setLineWidth(0)
            c.rect(cx_s - skirt_w/2, hull_top, skirt_w, skirt_h, fill=1, stroke=0)
            # Plataforma cubierta
            c.setFillColor(colors.HexColor("#5D6D7E"))
            c.rect(cx_s - sphere_r*0.65, hull_top-2, sphere_r*1.3, 4, fill=1, stroke=0)

            # Nivel de llenado
            try:
                vn_val = self.parse_float(self.get_var(f"{etapa_key}_{tn}_vol_nat_prod").get() or "0")
                ar_val = self.parse_float(self.get_var(f"{etapa_key}_{tn}_alt_ref").get() or "1")
                pct = min(vn_val / (ar_val if ar_val > 0 else 1), 1.0)
            except: pct = 0.0
            prod_name = self.get_var(f"{etapa_key}_{tn}_prod_name").get() if etapa_key else ""

            # Esfera gradiente (capas de blanco-aluminio)
            sphere_layers = [
                ("#1A4F6E",1.00),("#2471A3",0.92),("#2E86C1",0.82),
                ("#5DADE2",0.70),("#AED6F1",0.55),("#D6EAF8",0.38),
            ]
            for sc, sf in sphere_layers:
                c.setFillColor(colors.HexColor(sc)); c.setLineWidth(0)
                c.circle(cx_s, sphere_cy, sphere_r*sf, fill=1, stroke=0)

            # Color del producto — nunca azul (azul=agua)
            glp_col = self.get_prod_color(tn, etapa_key)[0] if etapa_key else "#FF6B35"
            if glp_col.upper() in ("#3498DB","#5DADE2","#2E86C1","#2471A3","#AED6F1","#85C1E9","#E8E8F0"):
                glp_col = "#FF6B35"

            # Llenado (producto dentro de la esfera)
            if pct > 0.02:
                fill_h = sphere_r * 2 * pct
                fill_y = sphere_cy - sphere_r
                clip_h = min(fill_h, sphere_r * 2 - 1)
                c.setFillColor(colors.HexColor(glp_col)); c.setLineWidth(0)
                c.saveState()
                p_clip = c.beginPath()
                p_clip.circle(cx_s, sphere_cy, sphere_r-1)
                c.clipPath(p_clip, stroke=0)
                c.rect(cx_s-sphere_r, fill_y, sphere_r*2, clip_h, fill=1, stroke=0)
                c.restoreState()

            # Contorno esfera
            c.setFillColor(colors.white); c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(1.5)
            c.circle(cx_s, sphere_cy, sphere_r, fill=0, stroke=1)
            # Anillo ecuatorial
            c.setStrokeColor(colors.HexColor("#2C3E50")); c.setLineWidth(0.8)
            c.ellipse(cx_s-sphere_r, sphere_cy-sphere_r*0.1, cx_s+sphere_r, sphere_cy+sphere_r*0.1, fill=0, stroke=1)

            # Tuberías y válvulas en la cima
            c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(1.2)
            c.line(cx_s, sphere_cy+sphere_r, cx_s, sphere_cy+sphere_r+8)
            c.setFillColor(colors.HexColor("#E74C3C")); c.setLineWidth(0)
            c.rect(cx_s-3, sphere_cy+sphere_r+6, 6, 4, fill=1, stroke=0)  # PRV

            # Etiqueta
            lbl_bg = glp_col if pct > 0.02 else "#7B8D9A"
            lbl_txt = self.contrast_text(lbl_bg)
            c.setFillColor(colors.HexColor(lbl_txt)); c.setFont("Helvetica-Bold", max(4.5, sphere_r*0.12))
            short = tn.replace("TK ","T").replace(" BABOR","B").replace(" ESTRIBOR","E")
            c.drawCentredString(cx_s, sphere_cy, short)
            if pct > 0:
                c.setFont("Helvetica", max(4, sphere_r*0.10))
                c.drawCentredString(cx_s, sphere_cy-sphere_r*0.18, f"{pct*100:.0f}%")
            if prod_name:
                c.setFillColor(colors.HexColor("#2C3E50"))
                c.setFont("Helvetica", max(3.5, sphere_r*0.09))
                c.drawCentredString(cx_s, hull_top - 8, prod_name[:14])

        # Mastil de proa
        c.setStrokeColor(colors.HexColor("#626567")); c.setLineWidth(1.2)
        c.line(bow_x-18, hull_top, bow_x-18, hull_top+28)
        c.line(bow_x-26, hull_top+20, bow_x-10, hull_top+20)
        c.setFillColor(colors.HexColor("#F4D03F")); c.setLineWidth(0)
        c.circle(bow_x-18, hull_top+28, 2, fill=1, stroke=0)

        # Línea de flotación
        c.setStrokeColor(colors.HexColor("#2E86C1")); c.setLineWidth(1)
        c.setDash(5, 2); c.line(stern_x-8, mid_y, bow_x+15, mid_y); c.setDash()

        c.restoreState()

    def _pdf_draw_metanero_ship(self, c, x, y, W, H, etapa_key, title):
        """Buque metanero/GNL — casco con hasta 4 tanques Kvaerner Moss o prisma."""
        from reportlab.lib import colors
        import math
        c.saveState()
        tank_names = [t for t in self.lista_tanques if not t.startswith("CARBONERA")]
        n = max(len(tank_names), 1)

        # Fondo neutro
        c.setFillColor(colors.HexColor("#E8F0FE")); c.setLineWidth(0)
        c.rect(x, y, W, H, fill=1, stroke=0)

        # Título
        c.setFont("Helvetica-Bold", min(8, max(6, H*0.04)))
        c.setFillColor(colors.HexColor("#1B3A5C"))
        _mtxt = f"BUQUE METANERO/GNL - {title}"
        max_chars = int(W / 4.5)
        if len(_mtxt) > max_chars:
            _mtxt = _mtxt[:max_chars-3] + "..."
        c.drawCentredString(x+W/2, y+H-10, _mtxt)

        M = 20
        hull_bot = y + H*0.10
        hull_top = y + H*0.50
        hull_h   = hull_top - hull_bot
        bow_x    = x + W - M
        stern_x  = x + M
        mid_y    = hull_bot + hull_h * 0.45

        # Casco obra muerta + obra viva (forma de LNG con popa recta)
        p = c.beginPath()
        p.moveTo(stern_x, hull_top)
        p.lineTo(bow_x-15, hull_top)
        p.curveTo(bow_x+12, hull_top, bow_x+15, mid_y+10, bow_x+5, mid_y)
        p.lineTo(stern_x, mid_y); p.close()
        c.setFillColor(colors.HexColor("#D5D8DC")); c.setStrokeColor(colors.HexColor("#7F8C8D")); c.setLineWidth(1.5)
        c.drawPath(p, fill=1, stroke=1)

        p2 = c.beginPath()
        p2.moveTo(stern_x, mid_y)
        p2.lineTo(bow_x+5, mid_y)
        p2.curveTo(bow_x+14, mid_y, bow_x+8, hull_bot+5, bow_x-12, hull_bot)
        p2.lineTo(stern_x+4, hull_bot)
        p2.curveTo(stern_x-4, hull_bot, stern_x-4, mid_y-4, stern_x, mid_y); p2.close()
        c.setFillColor(colors.HexColor("#D35400")); c.drawPath(p2, fill=1, stroke=0)
        c.setFillColor(colors.HexColor("#2C3E50")); c.setLineWidth(0)
        c.rect(stern_x, mid_y-2, bow_x-stern_x, 4, fill=1, stroke=0)

        # Cubierta (aislada = blanca)
        c.setFillColor(colors.HexColor("#F0F3F4")); c.setLineWidth(0)
        c.rect(stern_x+48, hull_top-2, bow_x-stern_x-60, 4, fill=1, stroke=0)

        # Caseta popa
        c.setFillColor(colors.HexColor("#ECF0F1")); c.setStrokeColor(colors.HexColor("#BDC3C7")); c.setLineWidth(0.8)
        c.roundRect(stern_x, hull_top, 48, 20, 2, fill=1, stroke=1)
        c.setFillColor(colors.HexColor("#D5D8DC"))
        c.roundRect(stern_x+4, hull_top+20, 40, 12, 2, fill=1, stroke=1)
        c.roundRect(stern_x-2, hull_top+28, 52, 6, 1, fill=1, stroke=1)
        c.setFillColor(colors.HexColor("#154360")); c.setLineWidth(0)
        c.roundRect(stern_x+14, hull_top+34, 10, 12, 2, fill=1, stroke=0)

        # ── TANQUES PRISMÁTICOS GNL (tipo Membrane) ────────────────────────────
        # Forma: trapezoide con esquinas biseladas
        tz_x1 = stern_x + 52
        tz_x2 = bow_x - 12
        tz_w   = tz_x2 - tz_x1
        tk_w_each = tz_w / n
        gap = 4
        tank_h_gnl = hull_h * 1.30   # tanques sobresalen sobre cubierta
        tank_bot_y = hull_top - hull_h * 0.80

        for i, tn in enumerate(tank_names[:n]):
            tx = tz_x1 + i * tk_w_each + gap/2
            tw = tk_w_each - gap
            ty_bot = hull_top - hull_h * 0.75
            ty_top = hull_top + hull_h * 0.95

            try:
                vn_val = self.parse_float(self.get_var(f"{etapa_key}_{tn}_vol_nat_prod").get() or "0")
                ar_val = self.parse_float(self.get_var(f"{etapa_key}_{tn}_alt_ref").get() or "1")
                pct = min(vn_val / (ar_val if ar_val > 0 else 1), 1.0)
            except: pct = 0.0
            prod_name = self.get_var(f"{etapa_key}_{tn}_prod_name").get() if etapa_key else ""
            tank_inner_h = ty_top - ty_bot

            # Cuerpo aislado (blanco-azulado)
            c.setFillColor(colors.HexColor("#EBF5FB")); c.setStrokeColor(colors.HexColor("#2E86C1")); c.setLineWidth(1.2)
            c.roundRect(tx, ty_bot, tw, tank_inner_h, 5, fill=1, stroke=1)

            # Color del producto — nunca azul (azul es SOLO para agua)
            gnl_col = self.get_prod_color(tn, etapa_key)[0] if etapa_key else "#E8E8F0"
            if gnl_col.upper() in ("#3498DB","#5DADE2","#2E86C1","#2471A3","#AED6F1","#85C1E9"):
                gnl_col = "#D5D8DC"

            # Llenado GNL (plata criogénica)
            if pct > 0.01:
                fill_px = tank_inner_h * pct
                c.setFillColor(colors.HexColor(gnl_col)); c.setLineWidth(0)
                c.saveState()
                p_clip2 = c.beginPath()
                p_clip2.roundRect(tx+1, ty_bot+1, tw-2, tank_inner_h-2, 4)
                c.clipPath(p_clip2, stroke=0)
                c.rect(tx+1, ty_bot+1, tw-2, fill_px-2, fill=1, stroke=0)
                # Brillo criogénico (franja clara en la superficie)
                c.setFillColor(colors.HexColor("#F0F3F4")); c.setLineWidth(0)
                c.rect(tx+1, ty_bot+fill_px-4, tw-2, 3, fill=1, stroke=0)
                c.restoreState()

            # Borde
            c.setFillColor(colors.white); c.setStrokeColor(colors.HexColor("#7F8C8D")); c.setLineWidth(1.2)
            c.roundRect(tx, ty_bot, tw, tank_inner_h, 5, fill=0, stroke=1)

            # Aislamiento (líneas diagonales)
            c.setStrokeColor(colors.HexColor("#BDC3C7")); c.setLineWidth(0.4)
            for il in range(0, int(tw), 8):
                c.line(tx+il, ty_bot, tx+il+6, ty_bot+6)
            # Dom (cima redondeada = cúpula de aislamiento)
            c.setFillColor(colors.HexColor("#CCD1D9")); c.setStrokeColor(colors.HexColor("#7F8C8D")); c.setLineWidth(1)
            c.roundRect(tx+tw*0.2, ty_top-4, tw*0.6, 10, 3, fill=1, stroke=1)
            # PRV + tuberías
            c.setStrokeColor(colors.HexColor("#7F8C8D")); c.setLineWidth(1)
            c.line(tx+tw/2, ty_top+6, tx+tw/2, ty_top+12)
            c.setFillColor(colors.HexColor("#C0392B")); c.setLineWidth(0)
            c.rect(tx+tw/2-3, ty_top+10, 6, 4, fill=1, stroke=0)

            # Etiquetas
            c.setFillColor(colors.HexColor("#2C3E50")); c.setFont("Helvetica-Bold", max(4.5, tw*0.095))
            short = tn.replace("TANQUE ","T").replace("TK ","T").replace(" BABOR","B").replace(" ESTRIBOR","E")
            c.drawCentredString(tx+tw/2, ty_bot + tank_inner_h*0.55, short)
            if pct > 0:
                txt_gnl = self.contrast_text(gnl_col) if pct > 0.01 else "#2C3E50"
                c.setFillColor(colors.HexColor(txt_gnl))
                c.setFont("Helvetica", max(4, tw*0.08))
                c.drawCentredString(tx+tw/2, ty_bot + tank_inner_h*0.35, f"{pct*100:.0f}%")
            if prod_name:
                c.setFillColor(colors.HexColor("#2C3E50")); c.setFont("Helvetica", max(3.5, tw*0.07))
                c.drawCentredString(tx+tw/2, hull_top-8, prod_name[:14])

        # Línea de flotación
        c.setStrokeColor(colors.HexColor("#2E86C1")); c.setLineWidth(1)
        c.setDash(5, 2); c.line(stern_x-8, mid_y, bow_x+10, mid_y); c.setDash()
        # Mastil
        c.setStrokeColor(colors.HexColor("#626567")); c.setLineWidth(1.2)
        c.line(bow_x-16, hull_top, bow_x-16, hull_top+26)
        c.line(bow_x-24, hull_top+18, bow_x-8, hull_top+18)
        c.setFillColor(colors.HexColor("#F4D03F")); c.setLineWidth(0)
        c.circle(bow_x-16, hull_top+26, 2, fill=1, stroke=0)

        c.restoreState()

    def dibujar_perfil_buque(self, c, x, y, width, height, title, tanks_data, ref_heights, water_levels, trim_sign):
        """PDF ship/tank profile drawing. Routes to specific renderers by tipo_medio."""
        tipo_nave = self.get_var("car_tipo_nave").get()
        etapa_key = 'inicial' if 'INICIAL' in title else 'final'
        tm = self.get_tipo_medio()

        # ── Non-maritime types: use ReportLab equivalent of TK drawings ──────
        if "TANQUE" in tm:
            self._pdf_draw_vertical_tank(c, x, y, width, height, etapa_key, title)
            return
        if tm == "ESFERA DE GAS":
            self._pdf_draw_spheres(c, x, y, width, height, etapa_key, title)
            return
        if "CAMION GAS" in tm:
            self._pdf_draw_pressure_truck(c, x, y, width, height, etapa_key, title)
            return
        if "CAMION" in tm:
            self._pdf_draw_liquid_truck(c, x, y, width, height, etapa_key, title)
            return
        if self.es_ducto():
            self._pdf_draw_pipeline(c, x, y, width, height, etapa_key, title)
            return
        if self.es_electrico():
            self._pdf_draw_electric(c, x, y, width, height, etapa_key, title)
            return
        if "GASERO" in tm or ("GLP" in tm and "CAMION" not in tm):
            self._pdf_draw_gasero_ship(c, x, y, width, height, etapa_key, title)
            return
        if "METANERO" in tm or "GNL" in tm:
            self._pdf_draw_metanero_ship(c, x, y, width, height, etapa_key, title)
            return
        c.saveState()
        cx, cy = x + width/2, y + height/2
        angle = 0
        if trim_sign > 0: angle = 1.5
        elif trim_sign < 0: angle = -1.5
        c.translate(cx, cy)
        c.rotate(angle)
        c.translate(-cx, -cy)

        # --- DIMENSIONES BASE ---
        hull_top_y = y + height - 25
        hull_bot_y = y + 15
        bow_x = x + width - 15
        stern_x = x + 15
        mid_y = (hull_top_y + hull_bot_y) / 2
        hull_h = hull_top_y - hull_bot_y

        # --- TITULO (inside drawing area, near top) ---
        c.setFont("Helvetica-Bold", 8)
        c.setFillColor(colors.HexColor("#1B3A5C"))
        c.drawCentredString(x + width/2, y + height - 10, title)

        if tipo_nave == "BARCAZA":
            # === BARCAZA ===
            r = 6
            # Casco
            c.setFillColor(colors.HexColor("#2C3E50"))
            c.setStrokeColor(colors.HexColor("#1B2631"))
            c.setLineWidth(1.5)
            c.roundRect(stern_x, hull_bot_y, bow_x - stern_x, hull_top_y - hull_bot_y, r, fill=1, stroke=1)
            # Cubierta
            c.setFillColor(colors.HexColor("#5D6D7E"))
            c.roundRect(stern_x, hull_top_y - 6, bow_x - stern_x, 6, r, fill=1, stroke=0)
            # Hatches en cubierta
            hatch_count = max(2, int((bow_x - stern_x - 30) // 30))
            hatch_step = (bow_x - stern_x - 30) / (hatch_count + 1)
            for hi2 in range(hatch_count):
                hx2 = stern_x + 15 + (hi2+1)*hatch_step - 10
                c.setFillColor(colors.HexColor("#485564")); c.setStrokeColor(colors.HexColor("#2C3E50")); c.setLineWidth(0.5)
                c.roundRect(hx2, hull_top_y - 5, 20, 4, 1, fill=1, stroke=1)
            # Franja roja de flotacion
            c.setFillColor(colors.HexColor("#922B21"))
            c.rect(stern_x + 1, hull_bot_y + 1, bow_x - stern_x - 2, (mid_y - hull_bot_y) - 1, fill=1, stroke=0)
            # Linea de flotacion
            c.setStrokeColor(colors.HexColor("#2E86C1"))
            c.setLineWidth(1.2)
            c.setDash(4, 2)
            c.line(stern_x - 8, mid_y, bow_x + 8, mid_y)
            c.setDash()

        else:
            # === BUQUE ===
            # --- Casco superior (azul marino oscuro - obra muerta) ---
            p_top = c.beginPath()
            p_top.moveTo(stern_x - 5, hull_top_y)
            p_top.lineTo(bow_x + 5, hull_top_y)
            p_top.curveTo(bow_x + 30, hull_top_y, bow_x + 35, mid_y + 10, bow_x + 15, mid_y)
            p_top.lineTo(stern_x - 5, mid_y)
            p_top.close()
            c.setFillColor(colors.HexColor("#1B2631"))
            c.setStrokeColor(colors.HexColor("#0E1A27"))
            c.setLineWidth(1)
            c.drawPath(p_top, fill=1, stroke=1)
            
            # Franja decorativa (linea blanca entre obra muerta y viva)
            c.setStrokeColor(colors.HexColor("#F0F0F0"))
            c.setLineWidth(1.2)
            c.line(stern_x - 5, mid_y + 1, bow_x + 15, mid_y + 1)

            # Ojos de buey (portholes) en la obra muerta
            port_y = mid_y + (hull_top_y - mid_y) * 0.45
            port_spacing = max(14, (bow_x - stern_x - 80) / 8)
            for pxi in range(7):
                px2 = stern_x + 60 + pxi * port_spacing
                if px2 < bow_x - 10:
                    c.setFillColor(colors.HexColor("#AED6F1")); c.setStrokeColor(colors.HexColor("#85C1E9")); c.setLineWidth(0.6)
                    c.circle(px2, port_y, 3.5, fill=1, stroke=1)
                    c.setFillColor(colors.HexColor("#FDFEFE")); c.setLineWidth(0)
                    c.circle(px2-1, port_y+1, 1, fill=1, stroke=0)

            # --- Casco inferior (rojo burdeos - obra viva) ---
            p_bot = c.beginPath()
            p_bot.moveTo(stern_x - 5, mid_y)
            p_bot.lineTo(bow_x + 15, mid_y)
            p_bot.curveTo(bow_x + 10, hull_bot_y + 10, bow_x - 10, hull_bot_y, bow_x - 25, hull_bot_y)
            p_bot.curveTo(cx, hull_bot_y - 6, stern_x + 20, hull_bot_y - 2, stern_x - 5, hull_bot_y + 5)
            p_bot.close()
            c.setFillColor(colors.HexColor("#7B241C"))
            c.setStrokeColor(colors.HexColor("#5B1A14"))
            c.drawPath(p_bot, fill=1, stroke=1)

            # Bulbo de proa (sugerido como elipse oscura)
            c.setFillColor(colors.HexColor("#2C3E50")); c.setStrokeColor(colors.HexColor("#1B2631")); c.setLineWidth(0.8)
            c.ellipse(bow_x-15, hull_bot_y-3, bow_x+8, hull_bot_y+8, fill=1, stroke=1)

            # Propulsión en la popa (hélice sugerida)
            prop_x = stern_x - 5
            prop_y = hull_bot_y + (mid_y - hull_bot_y) * 0.4
            c.setFillColor(colors.HexColor("#4A5568")); c.setStrokeColor(colors.HexColor("#2C3E50")); c.setLineWidth(0.8)
            c.circle(prop_x, prop_y, 5, fill=1, stroke=1)
            c.setStrokeColor(colors.HexColor("#7F8C8D")); c.setLineWidth(1.2)
            import math as _shm2
            for ang_p in [0, 60, 120, 180, 240, 300]:
                ar_p = _shm2.radians(ang_p)
                c.line(prop_x, prop_y, prop_x+_shm2.cos(ar_p)*5, prop_y+_shm2.sin(ar_p)*5)

            # Ancla en proa
            anc_x = bow_x - 5; anc_y = hull_top_y
            c.setFillColor(colors.HexColor("#7F8C8D")); c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(1)
            c.line(anc_x, anc_y, anc_x, anc_y-6)  # Chain
            c.circle(anc_x, anc_y-6, 2, fill=1, stroke=1)  # Anchor ring

            # --- Linea de flotacion ---
            c.setStrokeColor(colors.HexColor("#2E86C1"))
            c.setLineWidth(1.2)
            c.setDash(4, 2)
            c.line(stern_x - 12, mid_y, bow_x + 20, mid_y)
            c.setDash()

            # --- Marca de Plimsoll (Load Line) en cuaderna maestra ─────────
            plim_cx = (stern_x + bow_x) / 2
            plim_cy = mid_y
            plim_r  = max(4, hull_h * 0.08)
            c.setStrokeColor(colors.HexColor("#F0F0F0"))
            c.setLineWidth(1.2)
            c.circle(plim_cx, plim_cy, plim_r, fill=0, stroke=1)
            c.line(plim_cx - plim_r - 5, plim_cy, plim_cx + plim_r + 5, plim_cy)
            # LR en el centro
            c.setFillColor(colors.HexColor("#F0F0F0"))
            c.setFont("Helvetica-Bold", max(3, int(plim_r*0.55)))
            c.drawCentredString(plim_cx, plim_cy - plim_r * 0.28, "LR")
            # Líneas de carga (S, T, F)
            c.setStrokeColor(colors.HexColor("#D0D0D0"))
            c.setLineWidth(0.7)
            c.setFont("Helvetica", max(3, int(plim_r*0.45)))
            for _lll, _loff in [("S", -plim_r*1.5), ("T", -plim_r*2.7), ("F", -plim_r*3.9)]:
                _ly = mid_y + _loff
                c.line(plim_cx - plim_r - 4, _ly, plim_cx - plim_r, _ly)
                c.setFillColor(colors.HexColor("#C0C0C0"))
                c.drawRightString(plim_cx - plim_r - 5, _ly - 1.5, _lll)

            # --- Marcas de Calado (Draft Marks) en PROA y POPA ──────────────
            import math as _shmd
            try:
                _pdf_c_proa = self.parse_float(self.get_var(f"{etapa_key}_Calados Proa").get() or "0")
                _pdf_c_popa = self.parse_float(self.get_var(f"{etapa_key}_Calados Popa").get() or "0")
            except:
                _pdf_c_proa = 0.0; _pdf_c_popa = 0.0

            def _pdf_draft_scale(c2, dx, y_wl, y_keel, calado, side="right"):
                _h = y_wl - y_keel  # en PDF, y aumenta hacia arriba
                if _h < 8: return
                _ppm = _h / max(calado, 4.0) if calado > 0 else _h / 6.0
                _nmarks = int(calado) + 1 if calado > 0 else 5
                c2.setStrokeColor(colors.HexColor("#C0C0C0"))
                c2.setLineWidth(0.5)
                c2.line(dx, y_keel - 2, dx, y_wl + 4)
                c2.setFont("Helvetica", max(3, int(_h/12)))
                for _mi in range(0, min(_nmarks + 2, 12)):
                    _my = y_keel + _mi * _ppm
                    if _my > y_wl + 5: break
                    _lx1 = dx if side == "right" else dx - 4
                    _lx2 = dx + 4 if side == "right" else dx
                    c2.line(_lx1, _my, _lx2, _my)
                    if _mi > 0:
                        c2.setFillColor(colors.HexColor("#B0B0B0"))
                        if side == "right":
                            c2.drawString(_lx2 + 1, _my - 1.5, str(_mi))
                        else:
                            c2.drawRightString(_lx1 - 1, _my - 1.5, str(_mi))
                # Calado medido (destacado)
                if calado > 0:
                    _wl_mark = y_keel + calado * _ppm
                    if y_keel < _wl_mark < y_wl + 4:
                        c2.setFillColor(colors.HexColor("#F4D03F"))
                        c2.setStrokeColor(colors.HexColor("#F4D03F"))
                        c2.setLineWidth(1.0)
                        _ax = dx + 8 if side == "right" else dx - 8
                        c2.line(dx, _wl_mark, _ax, _wl_mark)
                        c2.setFont("Helvetica-Bold", max(3, int(_h/10)))
                        if side == "right":
                            c2.drawString(_ax + 1, _wl_mark - 1.5, f"{calado:.2f}m")
                        else:
                            c2.drawRightString(_ax - 1, _wl_mark - 1.5, f"{calado:.2f}m")
                        c2.setLineWidth(0.5)

            # Escala en PROA (derecha)
            _bow_dx = bow_x - 8
            _pdf_draft_scale(c, _bow_dx, mid_y, hull_bot_y,
                             _pdf_c_proa if _pdf_c_proa > 0 else 4.0, "left")
            c.setFillColor(colors.HexColor("#A0A0A0"))
            c.setFont("Helvetica", max(3, int(hull_h/14)))
            c.drawCentredString(_bow_dx, hull_top_y + 3, "PROA")

            # Escala en POPA (izquierda)
            _stern_dx = stern_x + 8
            _pdf_draft_scale(c, _stern_dx, mid_y, hull_bot_y,
                             _pdf_c_popa if _pdf_c_popa > 0 else 4.0, "right")
            c.drawCentredString(_stern_dx, hull_top_y + 3, "POPA")

            # Nombre del buque en el casco (si disponible)
            _buq_nm = self.get_var("car_buque").get()[:22]
            if _buq_nm and tipo_nave != "BARCAZA":
                c.setFillColor(colors.HexColor("#E8E8E8"))
                c.setFont("Helvetica-Bold", max(4, int(hull_h * 0.12)))
                _hull_mid_x = (stern_x + 65 + bow_x - 25) / 2
                _hull_name_y = mid_y + hull_h * 0.35
                c.drawCentredString(_hull_mid_x, _hull_name_y, _buq_nm)

            # --- Superestructura (casillaje) mejorada ---
            cas_x = stern_x
            cas_w = 60
            # Nivel 1 - acomodaciones con ventanas
            c.setFillColor(colors.HexColor("#F2F3F4"))
            c.setStrokeColor(colors.HexColor("#ABB2B9"))
            c.setLineWidth(0.8)
            c.roundRect(cas_x, hull_top_y, cas_w, 20, 2, fill=1, stroke=1)
            # Ventanas nivel 1
            for wni in range(3):
                wnx = cas_x + 6 + wni*16
                c.setFillColor(colors.HexColor("#AED6F1")); c.setStrokeColor(colors.HexColor("#5D6D7E")); c.setLineWidth(0.5)
                c.roundRect(wnx, hull_top_y+5, 10, 8, 1, fill=1, stroke=1)
                c.setFillColor(colors.HexColor("#FDFEFF")); c.setLineWidth(0)
                c.rect(wnx+1, hull_top_y+10, 4, 2, fill=1, stroke=0)

            # Nivel 2 - puente de mando con ventanas
            c.setFillColor(colors.HexColor("#E5E7E9"))
            c.roundRect(cas_x + 4, hull_top_y + 20, cas_w - 8, 14, 2, fill=1, stroke=1)
            # Ventanas puente
            for bni in range(4):
                bnx = cas_x + 6 + bni*12
                c.setFillColor(colors.HexColor("#D6EAF8")); c.setStrokeColor(colors.HexColor("#85C1E9")); c.setLineWidth(0.5)
                c.roundRect(bnx, hull_top_y+23, 9, 7, 1, fill=1, stroke=1)

            # Nivel 3 - alas del puente
            c.setFillColor(colors.HexColor("#D5D8DC")); c.setStrokeColor(colors.HexColor("#ABB2B9")); c.setLineWidth(0.6)
            c.roundRect(cas_x - 5, hull_top_y + 30, cas_w + 10, 7, 1, fill=1, stroke=1)

            # Chimenea
            chim_x = cas_x + cas_w/2 - 7
            chim_y = hull_top_y + 37
            c.setFillColor(colors.HexColor("#C0392B"))
            c.setStrokeColor(colors.HexColor("#922B21"))
            c.roundRect(chim_x, chim_y, 14, 18, 2, fill=1, stroke=1)
            # Franjas chimenea
            c.setFillColor(colors.HexColor("#1B2631"))
            c.rect(chim_x + 1, chim_y + 13, 12, 3, fill=1, stroke=0)
            c.setFillColor(colors.HexColor("#F4D03F"))
            c.rect(chim_x + 1, chim_y + 10, 12, 2, fill=1, stroke=0)
            # Humo de chimenea (elipses difusas)
            for sm_i, (sm_r, sm_a) in enumerate([(3,0),(4,2),(5,4)]):
                c.setFillColor(colors.HexColor("#C8D6DF")); c.setLineWidth(0)
                c.ellipse(chim_x+7+sm_a-sm_r, chim_y+18+sm_i*4,
                          chim_x+7+sm_a+sm_r, chim_y+18+sm_i*4+sm_r*1.3, fill=1, stroke=0)
            # --- Mastil de proa con antenas ---
            mast_x = bow_x - 20
            c.setStrokeColor(colors.HexColor("#626567"))
            c.setLineWidth(1.5)
            c.line(mast_x, hull_top_y, mast_x, hull_top_y + 32)
            c.setLineWidth(0.7)
            c.line(mast_x - 8, hull_top_y + 24, mast_x + 8, hull_top_y + 24)  # verga
            # Antena de radar (círculo pequeño rotado)
            c.setStrokeColor(colors.HexColor("#E74C3C")); c.setLineWidth(0.6)
            c.line(mast_x-6, hull_top_y+30, mast_x+6, hull_top_y+28)
            c.line(mast_x-6, hull_top_y+26, mast_x+6, hull_top_y+28)
            # Luz de tope (blanca)
            c.setFillColor(colors.HexColor("#F4D03F")); c.setLineWidth(0)
            c.circle(mast_x, hull_top_y + 32, 2, fill=1, stroke=0)
            # Luces laterales (verde estribor / rojo babor)
            c.setFillColor(colors.HexColor("#27AE60"))
            c.circle(mast_x + 8, hull_top_y + 24, 1.5, fill=1, stroke=0)
            c.setFillColor(colors.HexColor("#E74C3C"))
            c.circle(mast_x - 8, hull_top_y + 24, 1.5, fill=1, stroke=0)

        # === TANQUES DE CARGA ===
        target_side = "BABOR" if "BABOR" in title else "ESTRIBOR"
        tanks_to_draw = [tk_name for tk_name in self.lista_tanques if target_side in tk_name]

        if tanks_to_draw:
            start_x = stern_x + 62
            end_x = bow_x - 30
            total_w = end_x - start_x
            t_w = total_w / len(tanks_to_draw)
            gap = 1.5
            current_tx = start_x

            tank_zone_top = hull_top_y - 2
            tank_zone_bot = mid_y + 2
            tank_h = tank_zone_top - tank_zone_bot

            c.setLineWidth(0.5)

            for t_name in tanks_to_draw:
                val_liq = self.parse_float(self.get_var(f"{etapa_key}_{t_name}_s_corr").get())
                val_ref = self.parse_float(self.get_var(f"{etapa_key}_{t_name}_alt_ref").get())
                val_wat = self.parse_float(self.get_var(f"{etapa_key}_{t_name}_agua_s_real").get())
                val_lit = self.get_var(f"{etapa_key}_{t_name}_vol_nat_prod").get()
                prod_name = self.get_var(f"{etapa_key}_{t_name}_prod_name").get()

                ref_h = val_ref if val_ref > 0 else 10000.0
                altura_liquido = val_ref - val_liq
                if altura_liquido < 0: altura_liquido = 0
                pct_liq = min(altura_liquido / ref_h, 1.0)
                pct_wat = min(val_wat / ref_h, 1.0) if val_wat > 0 else 0

                px_fill = tank_h * pct_liq
                px_wat_fill = tank_h * pct_wat

                tx = current_tx + gap/2
                tw = t_w - gap

                # Fondo tanque (gris acero)
                c.setFillColor(colors.HexColor("#4A545E"))
                c.setStrokeColor(colors.HexColor("#2C3E50"))
                c.roundRect(tx, tank_zone_bot, tw, tank_h, 2, fill=1, stroke=1)

                # Agua (azul claro)
                if px_wat_fill > 0:
                    c.setFillColor(colors.HexColor("#5DADE2"))
                    c.rect(tx + 0.5, tank_zone_bot + 0.5, tw - 1, px_wat_fill, fill=1, stroke=0)

                # Producto — usar paleta sin azules (azul es exclusivo para AGUA)
                px_prod = px_fill - px_wat_fill
                if px_prod > 0:
                    _col, _ = self.get_prod_color(t_name, etapa_key)
                    c.setFillColor(colors.HexColor(_col))
                    c.rect(tx + 0.5, tank_zone_bot + px_wat_fill, tw - 1, px_prod, fill=1, stroke=0)
                else:
                    _col = "#4A545E"

                # Tapa tanque
                c.setFillColor(colors.HexColor("#7F8C8D"))
                c.rect(tx, tank_zone_top - 2, tw, 2, fill=1, stroke=0)

                # Textos — contraste correcto vs color de producto
                txt_col_ship = colors.HexColor(self.contrast_text(_col)) if pct_liq > 0.15 else colors.white
                c.setFillColor(txt_col_ship)
                c.setFont("Helvetica-Bold", 5.5)
                short_name = t_name.replace("BABOR","B").replace("ESTRIBOR","E").strip()
                center_y_tk = tank_zone_bot + (tank_h / 2)
                c.drawCentredString(tx + tw/2, center_y_tk + 10, short_name)
                c.setFont("Helvetica", 4.5)
                c.drawCentredString(tx + tw/2, center_y_tk + 3, f"{pct_liq*100:.0f}%")
                if prod_name:
                    c.setFont("Helvetica", 4)
                    c.drawCentredString(tx + tw/2, center_y_tk - 4, prod_name[:17])
                if val_lit:
                    c.setFont("Helvetica", 4)
                    c.drawCentredString(tx + tw/2, center_y_tk - 10, f"{val_lit} L")

                current_tx += t_w

        # === CARBONERA (filtrada por lado, igual que los tanques) ===
        if self.lista_carbonera:
            # Filtrar carboneras para esta vista
            carbs_to_draw = []
            for c_name in self.lista_carbonera:
                c_upper = c_name.upper()
                if target_side in c_upper:
                    carbs_to_draw.append(c_name)
                elif "BABOR" not in c_upper and "ESTRIBOR" not in c_upper:
                    # Sin lado = AMBOS, aparece en las dos vistas
                    carbs_to_draw.append(c_name)
            
            if carbs_to_draw:
                carb_zone_x = stern_x + 2
                carb_zone_w = 55
                carb_individual_h = min(22, (hull_h * 0.35) / max(len(carbs_to_draw), 1))
                # Posicionar entre mitad del casco y la cubierta (misma zona vertical que los tanques)
                carb_y_start = mid_y + 2

                for ci, c_name in enumerate(carbs_to_draw):
                    cy_pos = carb_y_start + ci * (carb_individual_h + 2)

                    val_liq = self.parse_float(self.get_var(f"{etapa_key}_{c_name}_s_corr").get())
                    val_ref = self.parse_float(self.get_var(f"{etapa_key}_{c_name}_alt_ref").get())
                    val_lit = self.get_var(f"{etapa_key}_{c_name}_vol_nat_prod").get()

                    ref_h = val_ref if val_ref > 0 else 10000.0
                    altura_liquido = val_ref - val_liq
                    if altura_liquido < 0: altura_liquido = 0
                    pct_liq = min(altura_liquido / ref_h, 1.0)
                    px_fill = carb_individual_h * pct_liq

                    # Fondo (crema cálido)
                    c.setFillColor(colors.HexColor("#FEF9E7"))
                    c.setStrokeColor(colors.HexColor("#D4AC0D"))
                    c.setLineWidth(1)
                    c.roundRect(carb_zone_x, cy_pos, carb_zone_w, carb_individual_h, 3, fill=1, stroke=1)

                    # Producto (dorado)
                    if px_fill > 0:
                        c.setFillColor(colors.HexColor("#F0B429"))
                        c.rect(carb_zone_x + 1, cy_pos + 1, carb_zone_w - 2, min(px_fill, carb_individual_h - 2), fill=1, stroke=0)

                    # Texto
                    c.setFillColor(colors.HexColor("#6E4B00"))
                    c.setFont("Helvetica-Bold", 5)
                    ccy = cy_pos + carb_individual_h / 2
                    short_carb = c_name.replace("BABOR","B").replace("ESTRIBOR","E").strip()
                    c.drawCentredString(carb_zone_x + carb_zone_w/2, ccy + 3, short_carb[:14])
                    c.setFont("Helvetica", 4)
                    info_parts = []
                    if val_lit: info_parts.append(f"{val_lit} L")
                    info_parts.append(f"{pct_liq*100:.0f}%")
                    c.drawCentredString(carb_zone_x + carb_zone_w/2, ccy - 4, " | ".join(info_parts))

        c.restoreState()

        # === VISTA DE POPA (ESCORA) ===
        x_popa = x + width + 20
        y_popa = y + 30
        size_popa = 70

        c.setFont("Helvetica-Bold", 7)
        c.setFillColor(colors.HexColor("#1B3A5C"))
        c.drawCentredString(x_popa + size_popa/2, y_popa + size_popa + 15, "VISTA POPA")

        list_val = self.parse_float(self.get_var(f"{etapa_key}_Lista").get())
        angle_list = list_val * 2.5
        if angle_list > 15: angle_list = 15
        if angle_list < -15: angle_list = -15

        # Agua de fondo popa
        c.setFillColor(colors.HexColor("#D6EAF8"))
        wl_y = y_popa + size_popa * 0.45
        c.rect(x_popa - 10, y_popa - 5, size_popa + 20, wl_y - y_popa + 8, fill=1, stroke=0)

        c.saveState()
        c.translate(x_popa + size_popa/2, y_popa + size_popa/2)
        c.rotate(angle_list)
        c.translate(-(x_popa + size_popa/2), -(y_popa + size_popa/2))

        c.setLineWidth(1)
        c.setStrokeColor(colors.HexColor("#1B2631"))

        if tipo_nave == "BARCAZA":
            c.setFillColor(colors.HexColor("#2C3E50"))
            c.roundRect(x_popa, y_popa + 15, size_popa, size_popa * 0.5, 4, fill=1, stroke=1)
            # Franja roja
            c.setFillColor(colors.HexColor("#922B21"))
            c.rect(x_popa + 1, y_popa + 16, size_popa - 2, size_popa * 0.2, fill=1, stroke=0)
        else:
            # Casco superior
            p = c.beginPath()
            p.moveTo(x_popa + 3, y_popa + size_popa)
            p.lineTo(x_popa + size_popa - 3, y_popa + size_popa)
            p.lineTo(x_popa + size_popa - 3, y_popa + 15)
            p.curveTo(x_popa + size_popa * 0.7, y_popa, x_popa + size_popa * 0.3, y_popa, x_popa + 3, y_popa + 15)
            p.close()
            c.setFillColor(colors.HexColor("#1B2631"))
            c.drawPath(p, fill=1, stroke=1)
            # Quilla roja
            mid_popa = y_popa + size_popa * 0.45
            p2 = c.beginPath()
            p2.moveTo(x_popa + 3, mid_popa)
            p2.lineTo(x_popa + size_popa - 3, mid_popa)
            p2.lineTo(x_popa + size_popa - 3, y_popa + 15)
            p2.curveTo(x_popa + size_popa * 0.7, y_popa, x_popa + size_popa * 0.3, y_popa, x_popa + 3, y_popa + 15)
            p2.close()
            c.setFillColor(colors.HexColor("#7B241C"))
            c.drawPath(p2, fill=1, stroke=0)
            # Franja blanca
            c.setStrokeColor(colors.HexColor("#F0F0F0"))
            c.setLineWidth(0.8)
            c.line(x_popa + 4, mid_popa, x_popa + size_popa - 4, mid_popa)

        c.restoreState()

        # Linea de agua (horizontal real)
        c.setStrokeColor(colors.HexColor("#2E86C1"))
        c.setLineWidth(1.2)
        c.setDash(4, 2)
        c.line(x_popa - 8, wl_y, x_popa + size_popa + 8, wl_y)
        c.setDash()

        # Etiquetas B / E con info de lista
        c.setFont("Helvetica-Bold", 6)
        c.setFillColor(colors.HexColor("#1B3A5C"))
        c.drawCentredString(x_popa - 5, y_popa + size_popa * 0.7, "B")
        c.drawCentredString(x_popa + size_popa + 5, y_popa + size_popa * 0.7, "E")
        # Ángulo de lista si hay escora
        if abs(list_val) > 0.05:
            _list_col = "#E74C3C" if abs(list_val) > 1.5 else "#F4D03F"
            c.setFillColor(colors.HexColor(_list_col))
            c.setFont("Helvetica-Bold", 5)
            _list_side = "B" if list_val > 0 else "E"
            c.drawCentredString(x_popa + size_popa/2, y_popa - 3,
                                f"Lista: {abs(list_val):.2f}° → {_list_side}")
        # Calados de babor/estribor si están disponibles
        try:
            _c_bab = self.parse_float(self.get_var(f"{etapa_key}_Calados Babor").get() or "0")
            _c_est = self.parse_float(self.get_var(f"{etapa_key}_Calados Estribor").get() or "0")
            if _c_bab > 0 or _c_est > 0:
                c.setFillColor(colors.HexColor("#4A90D9"))
                c.setFont("Helvetica", 4.5)
                if _c_bab > 0:
                    c.drawCentredString(x_popa - 5, y_popa + size_popa * 0.55, f"{_c_bab:.2f}m")
                if _c_est > 0:
                    c.drawCentredString(x_popa + size_popa + 5, y_popa + size_popa * 0.55, f"{_c_est:.2f}m")
        except: pass


    def generar_reporte_tecnico_global(self, suffix, output_folder, shared_canvas=None):
        clean_name = self.clean_filename(self.get_var('car_buque').get())
        filename = f"Reporte_Tecnico_Global_{clean_name}_{datetime.now().strftime('%Y%m%d')}.pdf"
        full_path = os.path.join(output_folder, filename)
        
        if shared_canvas:
            c = shared_canvas
            w, h = landscape(A4)
        else:
            try: c = canvas.Canvas(full_path, pagesize=landscape(A4))
            except: return None
            w, h = landscape(A4)
        
        def draw_logo_header(title_suffix, etapa=None):
            try:
                raw_b64 = ICON_REPORT_B64.strip().replace("\n", "").replace("\r", "")
                icon_data = base64.b64decode(raw_b64)
                img = ImageReader(BytesIO(icon_data))
                c.drawImage(img, 30, h-70, width=50, height=50, preserveAspectRatio=True, mask=None)
            except: 
                c.saveState()
                c.rect(30, h-70, 50, 50)
                c.drawString(40, h-45, "ARCA")
                c.restoreState()

            c.setFont("Helvetica-Bold", 14)
            c.drawCentredString(w/2, h-40, "ARCA - DGA")
            
            # TITULO DINAMICO
            _tm_pl = self.get_tipo_medio()
            if "BARCAZA" in _tm_pl: base_title = "PLANILLA DE SONDAJES - BARCAZA"
            elif "CAMION GAS" in _tm_pl: base_title = "PLANILLA DE MEDICION - CAMION GAS/GLP"
            elif "CAMION" in _tm_pl: base_title = "PLANILLA DE MEDICION - CAMION CISTERNA"
            elif "TANQUE FIJO" in _tm_pl: base_title = "PLANILLA DE SONDAJES - TANQUE FIJO"
            elif "TANQUE FLOTANTE" in _tm_pl: base_title = "PLANILLA DE SONDAJES - TANQUE FLOTANTE"
            elif _tm_pl == "OLEODUCTO": base_title = "ACTA DE MEDICION - OLEODUCTO"
            elif _tm_pl == "POLIDUCTO": base_title = "ACTA DE MEDICION - POLIDUCTO"
            elif _tm_pl == "GASODUCTO": base_title = "ACTA DE MEDICION - GASODUCTO"
            elif _tm_pl == "MEDICION ELECTRICA": base_title = "ACTA DE MEDICION ELECTRICA"
            elif "GASERO" in _tm_pl or ("GLP" in _tm_pl and "CAMION" not in _tm_pl): base_title = "PLANILLA DE SONDAJES - BUQUE GASERO/GLP"
            elif "METANERO" in _tm_pl or "GNL" in _tm_pl: base_title = "PLANILLA DE SONDAJES - BUQUE METANERO/GNL"
            elif "QUIMIQUERO" in _tm_pl: base_title = "PLANILLA DE SONDAJES - BUQUE QUIMIQUERO"
            elif "ESFERA" in _tm_pl: base_title = "PLANILLA DE MEDICIÓN - ESFERA DE GAS A PRESIÓN"
            elif "DRAFT SURVEY" in _tm_pl: base_title = "PLANILLA DE SONDAJES - DRAFT SURVEY"
            else: base_title = "PLANILLA DE SONDAJES DE TANQUES DE BUQUES"
            
            c.setFont("Helvetica-Bold", 10)
            c.drawCentredString(w/2, h-55, f"{base_title} - {title_suffix}")
            
            if etapa:
                fecha = self.get_var(f"{etapa}_Fecha").get()
                hora  = self.get_var(f"{etapa}_Hora").get()
                c.setFont("Helvetica", 9)
                # ── Mostrar calados SOLO para tipos marítimos con calados reales ──
                _tmh = self.get_tipo_medio()
                _es_mar_h = _tmh in ("BUQUE","BARCAZA","BUQUE QUIMIQUERO","DRAFT SURVEY",
                                     "BUQUE GASERO/GLP","BUQUE METANERO/GNL")
                if _es_mar_h:
                    c_proa    = self.get_var(f"{etapa}_Calados Proa").get()
                    c_popa    = self.get_var(f"{etapa}_Calados Popa").get()
                    trim      = self.get_var(f"{etapa}_Trimación").get()
                    c_babor   = self.get_var(f"{etapa}_Calados Babor").get()
                    c_estribor= self.get_var(f"{etapa}_Calados Estribor").get()
                    asiento   = self.get_var(f"{etapa}_Lista").get()
                    header_info = (f"FECHA: {fecha}  HORA: {hora}  |  "
                                   f"CALADO PROA: {c_proa}  |  CALADO POPA: {c_popa}  |  "
                                   f"TRIM: {trim}  |  BABOR: {c_babor}  |  "
                                   f"ESTRIBOR: {c_estribor}  |  ESCORA: {asiento}")
                else:
                    header_info = f"FECHA: {fecha}  |  HORA: {hora}  |  {_tmh}"
                c.drawCentredString(w/2, h-75, header_info)
        
        def draw_signatures_2_lines():
            y_sig = 40
            c.setLineWidth(1)
            c.setFillColor(colors.black)
            c.line(100, y_sig, 300, y_sig)
            c.drawCentredString(200, y_sig-12, "Aduana")
            c.line(500, y_sig, 700, y_sig)
            c.drawCentredString(600, y_sig-12, "Interesado")

        all_tanks = self.lista_tanques + self.lista_carbonera
        
        # 1. TABLA INICIAL
        _tm_pdf = self.get_tipo_medio()
        _es_tie_pdf = self.es_tierra()
        _es_cam_pdf = self.es_camion()
        _es_gas_pdf = self.es_gasero()
        _es_duc_pdf = self.es_ducto()
        _es_elec_pdf = self.es_electrico()
        _es_cgb_pdf = self.es_camion_gas()
        _es_esf_pdf = self.es_esfera()
        if _es_duc_pdf:
            draw_logo_header("ACTA DE MEDICION - DUCTO", "inicial")
            headers = ["TRAMO/PUNTO", "PRODUCTO", "CONT.INI", "CONT.FIN", "VOL.LIN(m3)", "P(kPa)", "T(C)", "Z", "VOL.BASE(m3)", "CORIOLIS(kg)"]
            x_pos = [50, 120, 195, 255, 315, 370, 415, 455, 505, 595]
        elif _es_elec_pdf:
            draw_logo_header("ACTA DE MEDICION ELECTRICA", "inicial")
            headers = ["PUNTO", "INI.kWh Act", "FIN.kWh Act", "kWh Act", "INI.kWh Rea", "FIN.kWh Rea", "kWh Rea", "cos fi", "V", "A"]
            x_pos = [50, 115, 180, 245, 295, 355, 415, 465, 510, 550]
        elif _es_esf_pdf:
            draw_logo_header("DETALLE TÉCNICO - ESFERA DE GAS - INICIAL", "inicial")
            headers = ["ESFERA", "PRODUCTO", "P(kPa)", "T(°C)", "DENS", "VOL.LIQ", "FASE", "VOL.NAT", "Z", "MASA(kg)"]
            x_pos = [50, 130, 210, 270, 330, 390, 450, 510, 575, 640]
        elif _es_gas_pdf:
            draw_logo_header("DETALLE TÉCNICO - BUQUE GASERO - INICIAL", "inicial")
            headers = ["TANQUE", "PRODUCTO", "P(kPa)", "T(°C)", "DENS.LIQ", "VOL.LIQ", "FASE", "VOL.NAT", "Z", "MASA(kg)"]
            x_pos = [50, 130, 210, 270, 330, 390, 450, 510, 575, 640]
        elif _es_tie_pdf or _es_cam_pdf or _es_cgb_pdf:
            draw_logo_header("DETALLE TECNICO - INICIAL (TIERRA/CAMION)", "inicial")
            headers = ["TANQUE", "PRODUCTO", "N° UTI", "ALT REF", "SONDAJE", "VOL.BRUTO", "AGUA", "S.CORR", "TEMP", "LTS NAT"]
            x_pos = [50, 120, 210, 265, 310, 365, 420, 480, 545, 605]
        else:
            draw_logo_header("DETALLE TÉCNICO - INICIAL PRODUCTO", "inicial")
            headers = ["TANQUE", "PRODUCTO", "N° UTI", "ALT REF", "SOND 1", "LTS 1", "SOND 2", "LTS 2", "DESC", "CORREGIDO", "TEMP", "LTS NAT"]
            x_pos = [50, 130, 230, 280, 330, 380, 430, 480, 530, 580, 650, 715]
        y = h - 110
        c.setFont("Helvetica-Bold", 8)
        for i, txt in enumerate(headers): c.drawString(x_pos[i], y, txt)
        y -= 15
        c.setFont("Helvetica", 8)
        
        for tk_name in all_tanks:
            if _es_duc_pdf:
                vals = [tk_name,
                    self.get_var(f"inicial_{tk_name}_prod_name").get(),
                    self.get_var(f"inicial_{tk_name}_cont_ini").get(),
                    self.get_var(f"inicial_{tk_name}_cont_fin").get(),
                    self.get_var(f"inicial_{tk_name}_vol_linea").get(),
                    self.get_var(f"inicial_{tk_name}_P_lin").get() or self.get_var("car_presion_op").get(),
                    self.get_var(f"inicial_{tk_name}_T_lin").get() or self.get_var("car_temp_op").get(),
                    self.get_var(f"inicial_{tk_name}_Z").get(),
                    self.get_var(f"inicial_{tk_name}_vol_base").get(),
                    self.get_var(f"inicial_{tk_name}_masa_coriolis").get()]
            elif _es_elec_pdf:
                vals = [tk_name,
                    self.get_var(f"inicial_{tk_name}_el_ini_act").get(),
                    self.get_var(f"inicial_{tk_name}_el_fin_act").get(),
                    self.get_var(f"inicial_{tk_name}_el_kwh_act").get(),
                    self.get_var(f"inicial_{tk_name}_el_ini_rea").get(),
                    self.get_var(f"inicial_{tk_name}_el_fin_rea").get(),
                    self.get_var(f"inicial_{tk_name}_el_kwh_rea").get(),
                    self.get_var(f"inicial_{tk_name}_el_fp").get(),
                    self.get_var(f"inicial_{tk_name}_el_V").get(),
                    self.get_var(f"inicial_{tk_name}_el_A").get()]
            elif _es_esf_pdf or _es_gas_pdf:
                vals = [tk_name,
                    self.get_var(f"inicial_{tk_name}_prod_name").get(),
                    self.get_var(f"inicial_{tk_name}_P_lin").get() or self.get_var("car_presion_op").get(),
                    self.get_var(f"inicial_{tk_name}_temp").get(),
                    self.get_var(f"inicial_{tk_name}_dens_lab").get(),
                    self.get_var(f"inicial_{tk_name}_vol_liq").get() or self.get_var(f"inicial_{tk_name}_vol_bruto").get(),
                    self.get_var(f"inicial_{tk_name}_fase").get(),
                    self.get_var(f"inicial_{tk_name}_vol_nat_prod").get(),
                    self.get_var(f"inicial_{tk_name}_Z").get(),
                    self.get_var(f"inicial_{tk_name}_masa_kg").get()]
            elif _es_tie_pdf or _es_cam_pdf or _es_cgb_pdf:
                vals = [tk_name,
                    self.get_var(f"inicial_{tk_name}_prod_name").get(),
                    self.get_var(f"inicial_{tk_name}_num_uti").get(),
                    self.get_var(f"inicial_{tk_name}_alt_ref").get(),
                    self.get_var(f"inicial_{tk_name}_s_tierra").get(),
                    self.get_var(f"inicial_{tk_name}_vol_bruto").get(),
                    self.get_var(f"inicial_{tk_name}_vol_nat_agua").get(),
                    self.get_var(f"inicial_{tk_name}_s_corr").get(),
                    self.get_var(f"inicial_{tk_name}_temp").get(),
                    self.get_var(f"inicial_{tk_name}_vol_nat_prod").get()]
            else:
                vals = [tk_name, self.get_var(f"inicial_{tk_name}_prod_name").get(), self.get_var(f"inicial_{tk_name}_num_uti").get(), self.get_var(f"inicial_{tk_name}_alt_ref").get(), self.get_var(f"inicial_{tk_name}_prod_s1").get(), self.get_var(f"inicial_{tk_name}_prod_l1").get(), self.get_var(f"inicial_{tk_name}_prod_s2").get(), self.get_var(f"inicial_{tk_name}_prod_l2").get(), self.get_var(f"inicial_{tk_name}_desc_tubo").get(), self.get_var(f"inicial_{tk_name}_s_corr").get(), self.get_var(f"inicial_{tk_name}_temp").get(), self.get_var(f"inicial_{tk_name}_vol_nat_prod").get()]
            for i, v in enumerate(vals):
                if i < len(x_pos): c.drawString(x_pos[i], y, str(v))
            y -= 12
            if y < 80:
                draw_signatures_2_lines()
                c.showPage()
                if _es_duc_pdf: _cont_lbl = "ACTA DE MEDICION - DUCTO (CONT.)"
                elif _es_elec_pdf: _cont_lbl = "ACTA DE MEDICION ELECTRICA (CONT.)"
                elif _es_esf_pdf: _cont_lbl = "DETALLE TÉCNICO - ESFERA DE GAS (CONT.)"
                elif _es_gas_pdf: _cont_lbl = "DETALLE TÉCNICO - BUQUE GASERO (CONT.)"
                elif _es_tie_pdf or _es_cam_pdf or _es_cgb_pdf: _cont_lbl = "DETALLE TECNICO - INICIAL (CONT.)"
                else: _cont_lbl = "DETALLE TÉCNICO - INICIAL PRODUCTO (CONT.)"
                draw_logo_header(_cont_lbl, "inicial")
                y = h - 110
        
        draw_signatures_2_lines()
        c.showPage()
        
        # 2. TABLA FINAL
        if _es_duc_pdf: _lbl_final = "ACTA DE MEDICION - DUCTO (FINAL)"
        elif _es_elec_pdf: _lbl_final = "ACTA DE MEDICION ELECTRICA (FINAL)"
        elif _es_esf_pdf: _lbl_final = "DETALLE TÉCNICO - ESFERA DE GAS (FINAL)"
        elif _es_gas_pdf: _lbl_final = "DETALLE TÉCNICO - BUQUE GASERO (FINAL)"
        elif _es_tie_pdf or _es_cam_pdf or _es_cgb_pdf: _lbl_final = "DETALLE TECNICO - FINAL (TIERRA/CAMION)"
        else: _lbl_final = "DETALLE TÉCNICO - FINAL PRODUCTO"
        draw_logo_header(_lbl_final, "final")
        y = h - 110
        c.setFont("Helvetica-Bold", 8)
        for i, txt in enumerate(headers): c.drawString(x_pos[i], y, txt)
        y -= 15
        c.setFont("Helvetica", 8)
        for tk_name in all_tanks:
            if _es_duc_pdf:
                vals = [tk_name,
                    self.get_var(f"final_{tk_name}_prod_name").get(),
                    self.get_var(f"final_{tk_name}_cont_ini").get(),
                    self.get_var(f"final_{tk_name}_cont_fin").get(),
                    self.get_var(f"final_{tk_name}_vol_linea").get(),
                    self.get_var(f"final_{tk_name}_P_lin").get() or self.get_var("car_presion_op").get(),
                    self.get_var(f"final_{tk_name}_T_lin").get() or self.get_var("car_temp_op").get(),
                    self.get_var(f"final_{tk_name}_Z").get(),
                    self.get_var(f"final_{tk_name}_vol_base").get(),
                    self.get_var(f"final_{tk_name}_masa_coriolis").get()]
            elif _es_elec_pdf:
                vals = [tk_name,
                    self.get_var(f"final_{tk_name}_el_ini_act").get(),
                    self.get_var(f"final_{tk_name}_el_fin_act").get(),
                    self.get_var(f"final_{tk_name}_el_kwh_act").get(),
                    self.get_var(f"final_{tk_name}_el_ini_rea").get(),
                    self.get_var(f"final_{tk_name}_el_fin_rea").get(),
                    self.get_var(f"final_{tk_name}_el_kwh_rea").get(),
                    self.get_var(f"final_{tk_name}_el_fp").get(),
                    self.get_var(f"final_{tk_name}_el_V").get(),
                    self.get_var(f"final_{tk_name}_el_A").get()]
            elif _es_esf_pdf or _es_gas_pdf:
                vals = [tk_name,
                    self.get_var(f"final_{tk_name}_prod_name").get(),
                    self.get_var(f"final_{tk_name}_P_lin").get() or self.get_var("car_presion_op").get(),
                    self.get_var(f"final_{tk_name}_temp").get(),
                    self.get_var(f"final_{tk_name}_dens_lab").get(),
                    self.get_var(f"final_{tk_name}_vol_liq").get() or self.get_var(f"final_{tk_name}_vol_bruto").get(),
                    self.get_var(f"final_{tk_name}_fase").get(),
                    self.get_var(f"final_{tk_name}_vol_nat_prod").get(),
                    self.get_var(f"final_{tk_name}_Z").get(),
                    self.get_var(f"final_{tk_name}_masa_kg").get()]
            elif _es_tie_pdf or _es_cam_pdf or _es_cgb_pdf:
                vals = [tk_name,
                    self.get_var(f"final_{tk_name}_prod_name").get(),
                    self.get_var(f"final_{tk_name}_num_uti").get(),
                    self.get_var(f"final_{tk_name}_alt_ref").get(),
                    self.get_var(f"final_{tk_name}_s_tierra").get(),
                    self.get_var(f"final_{tk_name}_vol_bruto").get(),
                    self.get_var(f"final_{tk_name}_vol_nat_agua").get(),
                    self.get_var(f"final_{tk_name}_s_corr").get(),
                    self.get_var(f"final_{tk_name}_temp").get(),
                    self.get_var(f"final_{tk_name}_vol_nat_prod").get()]
            else:
                vals = [tk_name, self.get_var(f"final_{tk_name}_prod_name").get(), self.get_var(f"final_{tk_name}_num_uti").get(), self.get_var(f"final_{tk_name}_alt_ref").get(), self.get_var(f"final_{tk_name}_prod_s1").get(), self.get_var(f"final_{tk_name}_prod_l1").get(), self.get_var(f"final_{tk_name}_prod_s2").get(), self.get_var(f"final_{tk_name}_prod_l2").get(), self.get_var(f"final_{tk_name}_desc_tubo").get(), self.get_var(f"final_{tk_name}_s_corr").get(), self.get_var(f"final_{tk_name}_temp").get(), self.get_var(f"final_{tk_name}_vol_nat_prod").get()]
            for i, v in enumerate(vals):
                if i < len(x_pos): c.drawString(x_pos[i], y, str(v))
            y -= 12
            if y < 80:
                draw_signatures_2_lines()
                c.showPage()
                if _es_duc_pdf: _cont_lbl2 = "ACTA DUCTO - FINAL (CONT.)"
                elif _es_elec_pdf: _cont_lbl2 = "ACTA ELECTRICA - FINAL (CONT.)"
                elif _es_esf_pdf: _cont_lbl2 = "ESFERA DE GAS - FINAL (CONT.)"
                elif _es_gas_pdf: _cont_lbl2 = "BUQUE GASERO - FINAL (CONT.)"
                elif _es_tie_pdf or _es_cam_pdf or _es_cgb_pdf: _cont_lbl2 = "DETALLE TECNICO - FINAL (CONT.)"
                else: _cont_lbl2 = "DETALLE TÉCNICO - FINAL PRODUCTO (CONT.)"
                draw_logo_header(_cont_lbl2, "final")
                y = h - 110
                
        draw_signatures_2_lines()
        c.showPage()

        # 3. MEMORIA DE CALCULO — adaptada al tipo de medición
        _tm_rep = self.get_tipo_medio()
        _es_met_rep = "METANERO" in _tm_rep or "GNL" in _tm_rep
        _es_gas_rep = ("GASERO" in _tm_rep or "GLP" in _tm_rep) and not _es_met_rep
        _es_liq_mar = _tm_rep in ("BUQUE","BARCAZA","BUQUE QUIMIQUERO","DRAFT SURVEY")
        _es_tie_rep = self.es_tierra()
        _es_cam_rep = self.es_camion() and not self.es_camion_gas()
        _es_duc_rep = self.es_ducto()
        _es_el_rep  = self.es_electrico()
        _es_esf_rep = self.es_esfera()

        for etapa in ['inicial', 'final']:
            draw_logo_header(f"MEMORIA DE CÁLCULO ({etapa.upper()})", etapa)
            y = h - 100
            tanks_on_page = 0

            # Cabecera de sección según tipo
            c.setFont("Helvetica-Bold", 9)
            c.setFillColor(colors.HexColor("#1B3A5C"))
            tipo_label = {
                True: "BUQUE METANERO/GNL — Composición Molar + Fases",
                False: None,
            }
            if _es_met_rep:
                c.drawString(40, y, "TIPO: BUQUE METANERO/GNL — Composición Molar + Fases Líquida/Gaseosa")
            elif _es_gas_rep:
                c.drawString(40, y, "TIPO: BUQUE GASERO/GLP — Presión, Temperatura, Factor Z")
            elif _es_liq_mar:
                c.drawString(40, y, "TIPO: BUQUE/QUIMIQUERO — Sondaje UTI, Tablas de Calibrado, VCF")
            elif _es_esf_rep:
                c.drawString(40, y, "TIPO: ESFERA DE GAS — Presión/Temperatura/Densidad Líquida")
            elif _es_duc_rep:
                c.drawString(40, y, "TIPO: DUCTO — Contadores volumétricos, Factor Z, Volumen Base")
            elif _es_el_rep:
                c.drawString(40, y, "TIPO: MEDICIÓN ELÉCTRICA — Lectura kWh activa/reactiva")
            else:
                c.drawString(40, y, f"TIPO: {_tm_rep} — Sondaje / Varilla / Calibrado")
            c.setFont("Helvetica", 8)
            c.setFillColor(colors.black)
            y -= 16

            for tk_name in all_tanks:
                if not self.get_var(f"{etapa}_{tk_name}_prod_name").get(): continue

                if tanks_on_page >= 5 or y < 80:
                    draw_signatures_2_lines()
                    c.showPage()
                    draw_logo_header(f"MEMORIA DE CÁLCULO ({etapa.upper()} - CONT.)", etapa)
                    y = h - 100; tanks_on_page = 0

                prod_name = self.get_var(f"{etapa}_{tk_name}_prod_name").get()
                vol_nat = self.get_var(f"{etapa}_{tk_name}_vol_nat_prod").get() or "0"

                # Encabezado de tanque
                c.setFillColor(colors.HexColor("#ECF0F1"))
                c.setLineWidth(0)
                c.rect(38, y-3, w-76, 14, fill=1, stroke=0)
                c.setStrokeColor(colors.HexColor("#AEB6BF")); c.setLineWidth(0.5)
                c.rect(38, y-3, w-76, 14, fill=0, stroke=1)
                c.setFont("Helvetica-Bold", 8)
                c.setFillColor(colors.HexColor("#1B3A5C"))
                c.drawString(42, y+2, f"TANQUE: {tk_name}  —  PRODUCTO: {prod_name}")
                c.setFont("Helvetica", 7.5)
                c.setFillColor(colors.black)
                y -= 14

                # ── Tipo MARITIMO LIQUIDO (UTI) ─────────────────────────────
                if _es_liq_mar:
                    s_uti = self.get_var(f"{etapa}_{tk_name}_alt_uti").get() or "—"
                    desc  = self.get_var(f"{etapa}_{tk_name}_desc_tubo").get() or "0"
                    s_corr = self.get_var(f"{etapa}_{tk_name}_s_corr").get() or "—"
                    alt_ref = self.get_var(f"{etapa}_{tk_name}_alt_ref").get() or "—"
                    num_uti = self.get_var(f"{etapa}_{tk_name}_num_uti").get() or "—"
                    interp_str = self.get_interpolation_details(etapa, tk_name)
                    dens = self.parse_float(self.get_var(f"{etapa}_{tk_name}_dens_lab").get() or "0") or \
                           self.parse_float(self.get_var(f"{etapa}_{tk_name}_dens_doc").get() or "0")
                    temp = self.parse_float(self.get_var(f"{etapa}_{tk_name}_temp").get() or "0")
                    tbl  = self.get_var(f"{etapa}_{tk_name}_tabla_vcf").get() or "54B"
                    vcf_str = self.get_vcf_details(dens, temp, tbl)
                    agua_real = self.get_var(f"{etapa}_{tk_name}_agua_s_real").get() or "0"
                    agua_lts  = self.get_var(f"{etapa}_{tk_name}_vol_nat_agua").get() or "0"
                    _neto_v = self.parse_float(vol_nat) - self.parse_float(agua_lts)
                    try:
                        # Vol.15°C y peso sobre el NETO (bruto − agua), igual que la planilla
                        v15_val = _neto_v * self.calc_vcf(dens, temp, tbl)
                        _d_air = dens / 1000.0 if dens > 2.0 else dens
                        kg_air_val = v15_val * (_d_air - 0.0011) if _d_air > 0 else 0
                    except: v15_val = 0; kg_air_val = 0
                    lines_m = [
                        f"  UTI Nro: {num_uti}  |  Alt.Referencia: {alt_ref} mm",
                        f"  1. Sondaje UTI: {s_uti} mm  —  Desc.tubo: {desc} mm  →  Sondaje Corr.: {s_corr} mm",
                        f"  2. Interpolación tabla: {interp_str}  =  {vol_nat} Lts (bruto)",
                        f"  3. Agua de fondo: {agua_real} mm  →  {agua_lts} Lts  |  Vol.Neto: {_neto_v:,.0f} Lts",
                        f"  4. VCF ({tbl}): {vcf_str}  |  Vol.15°C: {_neto_v:,.0f}×VCF = {v15_val:,.0f} Lts",
                        f"  5. Dens.lab: {dens}  |  Temp: {temp}°C  |  Peso en aire: {v15_val:,.0f}×({dens}−0.0011) = {kg_air_val:,.0f} Kg",
                    ]
                    if self.get_var(f"{etapa}_{tk_name}_tabla_trim_agua_json").get():
                        lines_m.insert(4, f"     Interp. agua: {self.get_water_interp_details(etapa, tk_name)}")

                # ── Tipo METANERO (GNL) ───────────────────────────────────────
                elif _es_met_rep:
                    MW = {"CH4":16.04,"C2H6":30.07,"C3H8":44.10,"C4H10":58.12,
                          "iC4":58.12,"nC5":72.15,"N2":28.01,"CO2":44.01,"H2S":34.08}
                    comp_parts = []
                    for k in MW.keys():
                        v = self.get_var(f"{etapa}_{tk_name}_gc_{k}").get()
                        if v and float(v or "0") > 0:
                            comp_parts.append(f"{k}:{v}%")
                    gc_str = "  ".join(comp_parts) or "—"
                    gc_pm  = self.get_var(f"{etapa}_{tk_name}_gc_PM").get() or "—"
                    gc_sum = self.get_var(f"{etapa}_{tk_name}_gc_sum").get() or "—"
                    vl = self.get_var(f"{etapa}_{tk_name}_vol_liq").get() or "—"
                    vv = self.get_var(f"{etapa}_{tk_name}_vol_vap").get() or "—"
                    dl = self.get_var(f"{etapa}_{tk_name}_dens_liq").get() or "—"
                    dv = self.get_var(f"{etapa}_{tk_name}_dens_vap").get() or "—"
                    tl = self.get_var(f"{etapa}_{tk_name}_temp_liq").get() or "—"
                    tv = self.get_var(f"{etapa}_{tk_name}_temp_vap").get() or "—"
                    pres = self.get_var(f"{etapa}_{tk_name}_pres_gnl").get() or "—"
                    masa_l = self.get_var(f"{etapa}_{tk_name}_masa_liq").get() or "—"
                    masa_v = self.get_var(f"{etapa}_{tk_name}_masa_vap").get() or "—"
                    lines_m = [
                        f"  Composición molar: {gc_str}",
                        f"  Suma %mol: {gc_sum}  |  PM calculado: {gc_pm} g/mol",
                        f"  FASE LÍQUIDA: Vol={vl} m³  Dens={dl} kg/m³  Temp={tl}°C  Masa={masa_l} t",
                        f"  FASE GASEOSA: Vol={vv} m³  Dens={dv} kg/m³  Temp={tv}°C  Masa={masa_v} kg",
                        f"  Presión: {pres} kPa",
                    ]

                # ── Tipo GASERO (GLP/Propano/Butano) ──────────────────────────
                elif _es_gas_rep:
                    pres = self.get_var(f"{etapa}_{tk_name}_presion").get() or "—"
                    tl   = self.get_var(f"{etapa}_{tk_name}_temp_liq").get() or "—"
                    z    = self.get_var(f"{etapa}_{tk_name}_factor_z").get() or "—"
                    dv   = self.get_var(f"{etapa}_{tk_name}_dens_vapor").get() or "—"
                    fase = self.get_var(f"{etapa}_{tk_name}_fase").get() or "—"
                    # vol_nat en m³ o litros según disponible
                    lines_m = [
                        f"  Presión: {pres} kPa  |  Temperatura: {tl}°C  |  Factor Z: {z}",
                        f"  Densidad vapor: {dv} kg/m³  |  Fase: {fase}",
                        f"  Volumen liquido natural: {vol_nat}",
                    ]

                # ── Tipo ESFERA ───────────────────────────────────────────────
                elif _es_esf_rep:
                    pres  = self.get_var(f"{etapa}_{tk_name}_esf_pres").get() or "—"
                    temp  = self.get_var(f"{etapa}_{tk_name}_esf_temp").get() or "—"
                    dens  = self.get_var(f"{etapa}_{tk_name}_esf_dens").get() or "—"
                    vl    = self.get_var(f"{etapa}_{tk_name}_vol_liq").get() or "—"
                    vgas  = self.get_var(f"{etapa}_{tk_name}_esf_vol_gas").get() or "—"
                    masa  = self.get_var(f"{etapa}_{tk_name}_esf_masa").get() or "—"
                    fase  = self.get_var(f"{etapa}_{tk_name}_esf_fase").get() or "—"
                    lines_m = [
                        f"  Presión: {pres} kPa  |  Temperatura: {temp}°C",
                        f"  Densidad líq.: {dens} kg/m³  |  Vol.líquido: {vl} m³",
                        f"  Vol.gas base (15°C): {vgas} m³  |  Masa: {masa} t  |  Fase: {fase}",
                    ]

                # ── Tipo TIERRA o CAMION ──────────────────────────────────────
                elif _es_tie_rep or _es_cam_rep:
                    s_tierra = self.get_var(f"{etapa}_{tk_name}_s_tierra").get() or "—"
                    s_corr   = self.get_var(f"{etapa}_{tk_name}_s_corr").get() or "—"
                    vol_bruto = self.get_var(f"{etapa}_{tk_name}_vol_bruto").get() or "—"
                    interp_str = self.get_interpolation_details(etapa, tk_name)
                    dens = self.parse_float(self.get_var(f"{etapa}_{tk_name}_dens_lab").get() or "0")
                    temp = self.parse_float(self.get_var(f"{etapa}_{tk_name}_temp").get() or "0")
                    tbl  = self.get_var(f"{etapa}_{tk_name}_tabla_vcf").get() or "54B"
                    vcf_str = self.get_vcf_details(dens, temp, tbl)
                    _agua_t = self.parse_float(self.get_var(f"{etapa}_{tk_name}_vol_nat_agua").get() or "0")
                    _neto_t = self.parse_float(vol_nat) - _agua_t
                    try:
                        # parse_float (no float()): vol_nat de tierra viene con coma de miles
                        v15_val = _neto_t * self.calc_vcf(dens, temp, tbl)
                        _d_air = dens / 1000.0 if dens > 2.0 else dens
                        kg_air_val = v15_val * (_d_air - 0.0011) if _d_air > 0 else 0
                    except: v15_val = 0; kg_air_val = 0
                    lines_m = [
                        f"  1. Sondaje/Varilla: {s_tierra} mm  →  Corregido: {s_corr} mm",
                        f"  2. Interpolación: {interp_str}  =  {vol_nat} Lts",
                        f"  3. Agua: {_agua_t:,.0f} Lts  |  Vol.Neto: {_neto_t:,.0f} Lts",
                        f"  4. VCF ({tbl}): {vcf_str}  |  Vol.15°C: {v15_val:,.0f} Lts",
                        f"  5. Dens: {dens}  |  Temp: {temp}°C  |  Peso en aire: {kg_air_val:,.0f} Kg",
                    ]

                # ── Tipo DUCTO ────────────────────────────────────────────────
                elif _es_duc_rep:
                    c_ini = self.get_var(f"{etapa}_{tk_name}_cont_ini").get() or "—"
                    c_fin = self.get_var(f"{etapa}_{tk_name}_cont_fin").get() or "—"
                    vl    = self.get_var(f"{etapa}_{tk_name}_vol_linea").get() or "—"
                    P_lin = self.get_var(f"{etapa}_{tk_name}_P_lin").get() or "—"
                    T_lin = self.get_var(f"{etapa}_{tk_name}_T_lin").get() or "—"
                    Z     = self.get_var(f"{etapa}_{tk_name}_Z").get() or "—"
                    vb    = self.get_var(f"{etapa}_{tk_name}_vol_base").get() or "—"
                    lines_m = [
                        f"  Contador ini: {c_ini} m³  |  Contador fin: {c_fin} m³  |  Vol.línea: {vl} m³",
                        f"  P línea: {P_lin} kPa  |  T línea: {T_lin}°C  |  Factor Z: {Z}",
                        f"  Vol.base (ref 101.325 kPa / 15°C): {vb} m³",
                    ]

                # ── Tipo ELECTRICO ─────────────────────────────────────────────
                elif _es_el_rep:
                    ia = self.get_var(f"{etapa}_{tk_name}_el_ini_act").get() or "—"
                    fa = self.get_var(f"{etapa}_{tk_name}_el_fin_act").get() or "—"
                    ka = self.get_var(f"{etapa}_{tk_name}_el_kwh_act").get() or "—"
                    ir = self.get_var(f"{etapa}_{tk_name}_el_ini_rea").get() or "—"
                    fr = self.get_var(f"{etapa}_{tk_name}_el_fin_rea").get() or "—"
                    kr = self.get_var(f"{etapa}_{tk_name}_el_kwh_rea").get() or "—"
                    ct = self.get_var(f"{etapa}_{tk_name}_el_const").get() or "1"
                    fp = self.get_var(f"{etapa}_{tk_name}_el_fp").get() or "—"
                    lines_m = [
                        f"  kWh Activa:  Ini={ia}  Fin={fa}  Cte={ct}  →  {ka} kWh",
                        f"  kWh Reactiva: Ini={ir}  Fin={fr}  Cte={ct}  →  {kr} kWh",
                        f"  cos fi calculado: {fp}",
                    ]

                else:
                    lines_m = [f"  Vol.Natural: {vol_nat}"]

                # Imprimir líneas de la memoria
                for line_m in lines_m:
                    if y < 70:
                        draw_signatures_2_lines(); c.showPage()
                        draw_logo_header(f"MEMORIA DE CÁLCULO ({etapa.upper()} - CONT.)", etapa)
                        y = h - 100; tanks_on_page = 0
                    c.setFont("Helvetica", 7.5)
                    c.setFillColor(colors.black)
                    c.drawString(40, y, line_m)
                    y -= 11
                y -= 4
                tanks_on_page += 1

            draw_signatures_2_lines()
            c.showPage()
        
        # 4. GRAFICOS — tipos marítimos: vista lateral; otros: dibujo técnico específico
        _tm_gr = self.get_tipo_medio()
        _es_mar_gr = _tm_gr in ("BUQUE","BARCAZA","BUQUE QUIMIQUERO","DRAFT SURVEY",
                                "BUQUE GASERO/GLP","BUQUE METANERO/GNL")
        if _es_mar_gr:
            draw_logo_header("DETALLE TÉCNICO - GRÁFICO VISUAL - INICIAL", "inicial")
            trim_i = self.parse_float(self.get_var("inicial_Trimación").get())
            self.dibujar_perfil_buque(c, 85, 350, 630, 135, "VISTA LATERAL - TANQUES BABOR (INICIAL)", [], [], [], trim_i)
            self.dibujar_perfil_buque(c, 85, 100, 630, 135, "VISTA LATERAL - TANQUES ESTRIBOR (INICIAL)", [], [], [], trim_i)
            draw_signatures_2_lines()
            c.showPage()

            draw_logo_header("DETALLE TÉCNICO - GRÁFICO VISUAL - FINAL", "final")
            trim_f = self.parse_float(self.get_var("final_Trimación").get())
            self.dibujar_perfil_buque(c, 85, 350, 630, 135, "VISTA LATERAL - TANQUES BABOR (FINAL)", [], [], [], trim_f)
            self.dibujar_perfil_buque(c, 85, 100, 630, 135, "VISTA LATERAL - TANQUES ESTRIBOR (FINAL)", [], [], [], trim_f)
            draw_signatures_2_lines()
            c.showPage()
        else:
            # Página de dibujo para tipos terrestres/ductos/eléctricos
            draw_logo_header("DETALLE TÉCNICO - GRÁFICO VISUAL - INICIAL", "inicial")
            self.dibujar_perfil_buque(c, 70, 70, w-140, h-220, "VISTA TÉCNICA (INICIAL)", [], [], [], 0)
            draw_signatures_2_lines()
            c.showPage()

            draw_logo_header("DETALLE TÉCNICO - GRÁFICO VISUAL - FINAL", "final")
            self.dibujar_perfil_buque(c, 70, 70, w-140, h-220, "VISTA TÉCNICA (FINAL)", [], [], [], 0)
            draw_signatures_2_lines()
            c.showPage()
        
        if not shared_canvas:
            c.save()
            return full_path
        return None


    def generar_un_reporte(self, suffix, tank_list, is_partial=False, ddt_obj=None, output_folder="", density_mode_key="dens_lab", shared_canvas=None):
        clean_buque = self.clean_filename(self.get_var('car_buque').get())
        clean_suffix = self.clean_filename(suffix)
        filename = f"Planilla_{clean_buque}_{clean_suffix}_{datetime.now().strftime('%Y%m%d')}.pdf"
        full_path = os.path.join(output_folder, filename)
        if shared_canvas:
            c = shared_canvas
            w, h = landscape(A4)
        else:
            print(f"Intentando guardar en: {full_path}")
            try: 
                c = canvas.Canvas(full_path, pagesize=landscape(A4))
            except Exception as e:
                traceback.print_exc() 
                messagebox.showerror("Error de PDF", f"No se pudo crear el archivo:\n{full_path}\n\nError: {str(e)}")
                return None
            w, h = landscape(A4)
        sum_ini = {"neto":0, "15":0, "kv": 0, "ka": 0}
        sum_fin = {"neto":0, "15":0, "kv": 0, "ka": 0}
        
        # Para tipos sin VCF/densidad clásica, usar vol_nat_prod directamente
        _sum_tm = self.get_tipo_medio()
        _sum_es_liq = _sum_tm in ("BUQUE","BARCAZA","BUQUE QUIMIQUERO","DRAFT SURVEY",
                                   "TANQUE FIJO","TANQUE FLOTANTE","CAMION CISTERNA")
        _sum_es_elec = self.es_electrico()
        _sum_es_duc  = self.es_ducto()

        for tk_name in tank_list:
            if _sum_es_elec:
                # kWh activa como "neto"
                i_v = self.parse_float(self.get_var(f"inicial_{tk_name}_el_kwh_act").get() or "0")
                f_v = self.parse_float(self.get_var(f"final_{tk_name}_el_kwh_act").get() or "0")
                sum_ini["neto"] += i_v; sum_ini["15"] += i_v
                sum_fin["neto"] += f_v; sum_fin["15"] += f_v
                continue
            elif _sum_es_duc:
                # vol_base como "neto"
                i_v = self.parse_float(self.get_var(f"inicial_{tk_name}_vol_base").get() or "0")
                f_v = self.parse_float(self.get_var(f"final_{tk_name}_vol_base").get() or "0")
                sum_ini["neto"] += i_v; sum_ini["15"] += i_v
                sum_fin["neto"] += f_v; sum_fin["15"] += f_v
                continue
            elif not _sum_es_liq:
                # Gas/esfera/camion gas: vol_nat_prod como "neto"
                i_v = self.parse_float(self.get_var(f"inicial_{tk_name}_vol_nat_prod").get() or "0")
                f_v = self.parse_float(self.get_var(f"final_{tk_name}_vol_nat_prod").get() or "0")
                sum_ini["neto"] += i_v; sum_ini["15"] += i_v
                sum_fin["neto"] += f_v; sum_fin["15"] += f_v
                continue
            
            # Líquidos clásicos: cálculo completo con VCF y densidades
            d_i_raw = self.get_var(f"inicial_{tk_name}_{density_mode_key}").get()
            if not d_i_raw: d_i_raw = self.get_var(f"inicial_{tk_name}_dens_lab").get()
            d_i = self.parse_float(d_i_raw)
            table_type = self.get_var(f"inicial_{tk_name}_tabla_vcf").get()
            if not table_type: table_type = "54B (Combustibles)"
            i_temp = self.get_var(f"inicial_{tk_name}_temp").get()
            i_temp = self.parse_float(i_temp) 
            i_bruto = self.interpolar_prod("inicial", tk_name)
            i_agua = self.interpolar_agua("inicial", tk_name)
            i_neto = i_bruto - i_agua
            i_vcf = self.calc_vcf(d_i, i_temp, table_type) if d_i > 0 else 1.0
            i_15 = i_neto * i_vcf
            _di_g = d_i / 1000.0 if d_i > 2.0 else d_i
            i_kv = i_15 * _di_g if _di_g > 0 else 0
            i_ka = i_15 * (_di_g - 0.0011) if _di_g > 0 else 0
            sum_ini["neto"] += i_neto
            sum_ini["15"] += i_15
            sum_ini["kv"] += i_kv
            sum_ini["ka"] += i_ka
            d_f_raw = self.get_var(f"final_{tk_name}_{density_mode_key}").get()
            if not d_f_raw: d_f_raw = self.get_var(f"final_{tk_name}_dens_lab").get()
            d_f = self.parse_float(d_f_raw)
            table_type_f = self.get_var(f"final_{tk_name}_tabla_vcf").get()
            if not table_type_f: table_type_f = "54B (Combustibles)"
            f_temp = self.get_var(f"final_{tk_name}_temp").get()
            f_temp = self.parse_float(f_temp) 
            f_bruto = self.interpolar_prod("final", tk_name)
            f_agua = self.interpolar_agua("final", tk_name)
            f_neto = f_bruto - f_agua
            f_vcf = self.calc_vcf(d_f, f_temp, table_type_f) if d_f > 0 else 1.0
            f_15 = f_neto * f_vcf
            _df_g = d_f / 1000.0 if d_f > 2.0 else d_f
            f_kv = f_15 * _df_g if _df_g > 0 else 0
            f_ka = f_15 * (_df_g - 0.0011) if _df_g > 0 else 0
            sum_fin["neto"] += f_neto
            sum_fin["15"] += f_15
            sum_fin["kv"] += f_kv
            sum_fin["ka"] += f_ka
        
        avg_d_i = sum_ini["kv"] / sum_ini["15"] if sum_ini["15"] > 0 else 0
        avg_d_f = sum_fin["kv"] / sum_fin["15"] if sum_fin["15"] > 0 else 0
        kv_i, ka_i = sum_ini["kv"], sum_ini["ka"]
        kv_f, ka_f = sum_fin["kv"], sum_fin["ka"]
        dif_net = sum_ini['neto'] - sum_fin['neto']
        dif_15 = sum_ini['15'] - sum_fin['15']
        dif_ka = ka_i - ka_f
        dif_kv = kv_i - kv_f

        try:
            raw_b64 = ICON_REPORT_B64.strip().replace("\n", "").replace("\r", "")
            icon_data = base64.b64decode(raw_b64)
            img = ImageReader(BytesIO(icon_data))
            c.drawImage(img, 30, h-70, width=100, height=50, preserveAspectRatio=True, mask=None)
        except: pass
        
        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(w/2, h-40, "ARCA - DGA")
        num_print = self.get_var('car_num_planilla_gen').get()
        if is_partial and ddt_obj:
            val = ddt_obj['num_planilla'].get()
            num_print = val if val else num_print

        # TITULO DINAMICO (aplica a todos los reportes, parciales y generales)
        _tm_pl = self.get_tipo_medio()
        if "BARCAZA" in _tm_pl: base_title = "PLANILLA DE SONDAJES - BARCAZA"
        elif "CAMION GAS" in _tm_pl: base_title = "PLANILLA DE MEDICION - CAMION GAS/GLP"
        elif "CAMION" in _tm_pl: base_title = "PLANILLA DE MEDICION - CAMION CISTERNA"
        elif "TANQUE FIJO" in _tm_pl: base_title = "PLANILLA DE SONDAJES - TANQUE FIJO"
        elif "TANQUE FLOTANTE" in _tm_pl: base_title = "PLANILLA DE SONDAJES - TANQUE FLOTANTE"
        elif _tm_pl == "OLEODUCTO": base_title = "ACTA DE MEDICION - OLEODUCTO"
        elif _tm_pl == "POLIDUCTO": base_title = "ACTA DE MEDICION - POLIDUCTO"
        elif _tm_pl == "GASODUCTO": base_title = "ACTA DE MEDICION - GASODUCTO"
        elif _tm_pl == "MEDICION ELECTRICA": base_title = "ACTA DE MEDICION ELECTRICA"
        elif _tm_pl == "ESFERA DE GAS": base_title = "PLANILLA DE MEDICIÓN - ESFERA DE GAS"
        elif "GASERO" in _tm_pl or ("GLP" in _tm_pl and "CAMION" not in _tm_pl): base_title = "PLANILLA DE SONDAJES - BUQUE GASERO/GLP"
        elif "METANERO" in _tm_pl or "GNL" in _tm_pl: base_title = "PLANILLA DE SONDAJES - BUQUE METANERO/GNL"
        elif "QUIMIQUERO" in _tm_pl: base_title = "PLANILLA DE SONDAJES - BUQUE QUIMIQUERO"
        elif "DRAFT SURVEY" in _tm_pl: base_title = "PLANILLA DE SONDAJES - DRAFT SURVEY"
        else: base_title = "PLANILLA DE SONDAJES DE TANQUES DE BUQUES"
            

        c.drawRightString(w-40, h-55, f"NUMERO: {num_print}")
        
        doc_num_safe = ""
        if ddt_obj: doc_num_safe = ddt_obj['numero'].get()
        
        calc_title = suffix.replace(f"DOC_{self.clean_filename(doc_num_safe)}_", "").replace("_", " ") if is_partial else "GENERAL"
        c.setFont("Helvetica-Bold", 9)
        c.drawCentredString(w/2, h-55, base_title)
        c.setFont("Helvetica", 8)
        c.drawCentredString(w/2, h-68, f"AGENCIA DE RECAUDACIÓN ADUANERA - {calc_title}")
        
        all_ddts_assigned = set()
        all_prods_assigned = set()
        for tk_name in tank_list:
            d = self.get_var(f"inicial_{tk_name}_ddt_assign").get()
            p = self.get_var(f"inicial_{tk_name}_prod_name").get()
            if d: all_ddts_assigned.add(d)
            if p: all_prods_assigned.add(p)
        txt_ddt = ", ".join(sorted(list(all_ddts_assigned)))
        label_ddt_title = "Documentos" if len(all_ddts_assigned) > 1 else "Documento"
        txt_prod = ", ".join(sorted(list(all_prods_assigned)))
        y_header_start = h - 90
        col1 = 53; col2 = 335; col3 = 595 
        c.setFont("Helvetica", 8)
        # Actores: del documento si es reporte parcial; si es general, los de
        # todos los documentos (distintos, unidos con ' / '), fallback carátula
        _act = self._actores_pdf(ddt_obj if is_partial else None)
        c.drawString(col1, y_header_start,    f"Buque: {self.get_var('car_buque').get()} (IMO: {self.get_var('car_imo').get()})")
        c.drawString(col1, y_header_start-12, f"Despachante: {_act['despachante']} ({_act['cuit_desp']})")
        c.drawString(col1, y_header_start-24, f"Importador/Exportador: {_act['impexp']} ({_act['cuit_impexp']})")
        c.drawString(col1, y_header_start-36, f"MANI: {self.get_var('car_mani').get()}")
        aduana_cod  = self.aduana_codigo()
        aduana_nom  = self.aduana_nombre()
        lugar_op_val = (self.get_var("car_lop_codigo").get() + " " + self.get_var("car_lop_desc").get()).strip()
        aduana_str  = f"Aduana: {aduana_cod} - {aduana_nom}" if aduana_cod else f"Aduana: {aduana_nom}"
        c.drawString(col1, y_header_start-48, aduana_str)
        c.drawString(col1, y_header_start-60, f"{label_ddt_title}: {txt_ddt}")
        if lugar_op_val:
            c.drawString(col1, y_header_start-72, f"Lugar Op.: {lugar_op_val}")
        c.drawString(col2, y_header_start,    aduana_str)
        c.drawString(col2, y_header_start-12, f"ATA: {_act['ata']} ({_act['cuit_ata']})")
        c.drawString(col2, y_header_start-24, f"Viaje: {self.get_var('car_conocimientos').get()}")
        c.drawString(col2, y_header_start-36, f"Producto: {txt_prod}")
        current_y = y_header_start - 84  # shifted down for Lugar Op line
        if is_partial and ddt_obj:
            tot_sal_l = 0; tot_sal_k = 0
            for s in ddt_obj["salidas"]:
                try: tot_sal_l += self.parse_float(s['litros'].get()); tot_sal_k += self.parse_float(s['kilos'].get())
                except: pass
            doc_l = self.parse_float(ddt_obj['litros'].get())
            doc_k = self.parse_float(ddt_obj['kilos'].get())
            target_l = tot_sal_l if tot_sal_l > 0 else doc_l
            target_k = tot_sal_k if tot_sal_k > 0 else doc_k
            label_diff = "con Salidas" if tot_sal_l > 0 else "con Documento"
            diff_l = dif_15 - target_l
            diff_k = dif_kv - target_k 
            # Por mil: diferencia / valor declarado * 1000
            permil_l = (diff_l / target_l * 1000) if target_l != 0 else 0
            permil_k = (diff_k / target_k * 1000) if target_k != 0 else 0
            is_red_l = abs(permil_l) > 6.0
            is_red_k = abs(permil_k) > 6.0
            c.setFont("Helvetica", 8)
            dens_ddt = self.parse_float(ddt_obj['densidad'].get())
            tipo_doc_str = ddt_obj['tipo_doc'].get() if 'tipo_doc' in ddt_obj else "Detallada"
            c.drawString(col1, current_y, f"Documento {tipo_doc_str}: Litros {doc_l:,.0f}, Kilos {doc_k:,.0f}, Densidad {dens_ddt}")
            y_diff_top = y_header_start - 48 
            col_font_l = colors.red if is_red_l else colors.blue
            txt_l = f"Diferencias Litros (15°C) {label_diff}: {diff_l:,.0f} ({permil_l:.2f} ‰)"
            c.setFillColor(col_font_l)
            c.setFont("Helvetica", 8)
            c.drawString(col2, y_diff_top, txt_l)
            col_font_k = colors.red if is_red_k else colors.blue
            txt_k = f"Diferencias Kilos (Vacío) {label_diff}: {diff_k:,.0f} ({permil_k:.2f} ‰)"
            c.setFillColor(col_font_k)
            c.drawString(col2, y_diff_top - 12, txt_k)
            c.setFillColor(colors.black) 
            y_dens = y_header_start - 12
            c.setFont("Helvetica", 7)
            if lugar_op_val:
                c.drawString(col3, y_dens, f"Lugar Op.: {lugar_op_val}")
                y_dens -= 10
            c.setFont("Helvetica-Bold", 8)
            c.drawString(col3, y_dens, "COMPARATIVAS DENSIDAD:")
            y_dens -= 12
            c.setFont("Helvetica", 7)
            dens_salida_avg = 0
            if ddt_obj["salidas"]:
                tm=0; tv=0
                for s in ddt_obj["salidas"]:
                    dv=self.parse_float(s['densidad'].get()); lv=self.parse_float(s['litros'].get())
                    tm+=(dv*lv); tv+=lv
                if tv>0: dens_salida_avg = tm/tv
            def fmt_diff(v1, v2):
                d = v1 - v2
                pct = (d / v1) * 1000 if v1 != 0 else 0
                return f"{d:.5f} ({pct:.2f} ‰)"
            c.drawString(col3, y_dens, f"Diferencia Doc vs Salida: {fmt_diff(dens_ddt, dens_salida_avg)}")
            y_dens -= 10
            c.drawString(col3, y_dens, f"Diferencia Salida vs Inicial: {fmt_diff(dens_salida_avg, avg_d_i)}")
            y_dens -= 10
            c.drawString(col3, y_dens, f"Diferencia Salida vs Final: {fmt_diff(dens_salida_avg, avg_d_f)}")
            y_dens -= 10
            c.drawString(col3, y_dens, f"Diferencia Doc vs Inicial: {fmt_diff(dens_ddt, avg_d_i)}")
            y_dens -= 10
            c.drawString(col3, y_dens, f"Diferencia Doc vs Final: {fmt_diff(dens_ddt, avg_d_f)}")
            y_dens -= 10
            c.drawString(col3, y_dens, f"Diferencia Ponderada Inicial vs Final: {fmt_diff(avg_d_i, avg_d_f)}")
            current_y -= 10
            c.setFont("Helvetica", 7)
            for s in ddt_obj["salidas"]:
                line_salida = f"Salida {s['numero'].get()}: Litros {s['litros'].get()}, Kilos {s['kilos'].get()}, Densidad {s['densidad'].get()}"
                c.drawString(col1, current_y, line_salida) 
                current_y -= 10
            current_y -= 5
        y_med = current_y - 10
        col_final_header = 600
        # Determinar si este reporte es de tipo marítimo para mostrar calados
        _tmr = self.get_tipo_medio()
        _es_mar_r = _tmr in ("BUQUE","BARCAZA","BUQUE QUIMIQUERO","DRAFT SURVEY",
                             "BUQUE GASERO/GLP","BUQUE METANERO/GNL")

        c.setFont("Helvetica-Bold", 8)
        c.drawString(col1, y_med, "MEDICION INICIAL")
        c.setFont("Helvetica", 8)
        
        f_ini = self.get_var('inicial_Fecha').get()
        h_ini = self.get_var('inicial_Hora').get()
        
        c.drawString(col1 + 140, y_med, f"Fecha: {f_ini}")
        c.drawString(col1 + 220, y_med, f"Hora: {h_ini}")
        if _es_mar_r:
            c_proa_i  = self.get_var('inicial_Calados Proa').get()
            c_popa_i  = self.get_var('inicial_Calados Popa').get()
            trim_i    = self.get_var('inicial_Trimación').get()
            bab_i     = self.get_var('inicial_Calados Babor').get()
            est_i     = self.get_var('inicial_Calados Estribor').get()
            asiento_i = self.get_var('inicial_Lista').get()
            c.drawString(col1, y_med - 10,
                         f"Calados: Proa {c_proa_i} | Popa {c_popa_i} | Trim {trim_i} | Babor {bab_i} | Estribor {est_i} | Escora {asiento_i}")

        c.setFont("Helvetica-Bold", 8)
        c.drawString(col_final_header - 180, y_med, "MEDICION FINAL")
        c.setFont("Helvetica", 8)
        
        f_fin = self.get_var('final_Fecha').get()
        h_fin = self.get_var('final_Hora').get()

        c.drawString(col_final_header + 90, y_med, f"Fecha: {f_fin}")
        c.drawString(col_final_header + 170, y_med, f"Hora: {h_fin}")
        if _es_mar_r:
            c_proa_f  = self.get_var('final_Calados Proa').get()
            c_popa_f  = self.get_var('final_Calados Popa').get()
            trim_f    = self.get_var('final_Trimación').get()
            bab_f     = self.get_var('final_Calados Babor').get()
            est_f     = self.get_var('final_Calados Estribor').get()
            asiento_f = self.get_var('final_Lista').get()
            c.drawString(col_final_header - 180, y_med - 10,
                         f"Calados: Proa {c_proa_f} | Popa {c_popa_f} | Trim {trim_f} | Babor {bab_f} | Estribor {est_f} | Escora {asiento_f}")

        
  
 
        
        y_table = y_med - 25
        c.setLineWidth(2)
        c.line(35, y_table, w-35, y_table)
        
        # ── Headers adaptativos según tipo de medición ──────────────────
        _tm_rep_t = self.get_tipo_medio()
        _es_liq_rep = _tm_rep_t in ("BUQUE","BARCAZA","BUQUE QUIMIQUERO","DRAFT SURVEY")
        _es_gas_rep_t = self.es_gasero()
        _es_duc_rep_t = self.es_ducto()
        _es_el_rep_t  = self.es_electrico()
        _es_tie_rep_t = self.es_tierra()
        _es_cam_rep_t = self.es_camion() and not self.es_camion_gas()
        _es_cgb_rep_t = self.es_camion_gas()
        _es_esf_rep_t = self.es_esfera()
        
        if _es_liq_rep:
            h_col = ["Tanque N° B/E", "Alt.Ref", "Sondeo\nUti", "Sondeo\nAgua", "Temp\nºC", "Sondeo\nCorr", "Litros\nc/agua", "Litros\nAgua", "Litros\nNetos", "F.C.V.", "Litros\n15ºC"]
        elif _es_el_rep_t:
            h_col = ["Punto Medición", "kWh\nActiva", "kWh\nReactiva", "cos fi", "Demanda\nkW", "Tensión\nV", "Corriente\nA", "V·A", "Fases", "—", "—"]
        elif _es_duc_rep_t:
            h_col = ["Tramo/Punto", "Cont.\nInicial", "Cont.\nFinal", "Vol.Línea\nm3", "P línea\nkPa", "T línea\nºC", "Factor Z", "Vol.Base\nm3", "Vol.Base\nKm3", "Caudal\nm3/h", "Masa\nCoriolis"]
        elif _es_gas_rep_t:
            h_col = ["Tanque", "Producto", "Presión\nkPa", "Temp\nºC", "Dens.Liq\nkg/m3", "Vol.Liq\nm3", "Fase", "Vol.Nat\nLts", "Factor Z", "—", "—"]
        elif _es_esf_rep_t:
            h_col = ["Esfera", "Presión\nkPa", "Temp\nºC", "Dens.Liq\nkg/m3", "Vol.Liq\nm3", "Vol.Gas\nBase m3", "Masa\nt", "Fase", "—", "—", "—"]
        elif _es_cgb_rep_t:
            h_col = ["Compartim.", "Presión\nkPa", "Temp\nºC", "Dens.Liq\nkg/m3", "Masa\nkg", "Vol.Líq\nL", "Vol.Gas\nm3", "Fase", "—", "—", "—"]
        elif _es_tie_rep_t or _es_cam_rep_t:
            h_col = ["Tanque/Comp.", "N°Util.", "Sondaje\nmm", "Sondaje\nCorr mm", "Temp\nºC", "Vol.Bruto\nLts", "Litros\nAgua", "Litros\nNetos", "F.C.V.", "Litros\n15ºC", "—"]
        else:
            h_col = ["Tanque N° B/E", "Altura Ref", "Sondeo\nUti", "Sondeo\nAgua", "Temp\nºC", "Sondeo\nCorregido", "Litros\nc/agua", "Litros\nAgua", "Litros\nNetos", "F.C.V.", "Litros\n15ºC"]
        
        h_left = h_col
        h_right = h_col
        x_start = 53
        col_w_name = 50 
        col_w_num = 31 
        c.setFont("Helvetica-Bold", 6)
        def draw_header_row(start_x, headers):
            curr_x = start_x
            for i, h in enumerate(headers):
                w_curr = col_w_name if i == 0 else col_w_num
                y_off = 0 if "\n" not in h else 6
                for l in h.split("\n"): 
                    c.drawString(curr_x, y_table - 8 - y_off, l)
                    y_off -= 6
                curr_x += w_curr
            return curr_x
        end_x_left = draw_header_row(x_start, h_left)
        x_mid = x_start + 360 + 15  # = 428 centrado
        draw_header_row(x_mid, h_right)
        c.line(35, y_table - 25, w-35, y_table - 25)
        y_row = y_table - 35
        c.setFont("Helvetica", 6)
        def val_or_zero(v):
            """Return '0' if value is empty/placeholder."""
            s = str(v).strip()
            if not s or s in ("DD/MM/AAAA", "00:00"): return "0"
            return s
        for tk_name in tank_list:
            i_ref = val_or_zero(self.get_var(f"inicial_{tk_name}_alt_ref").get())
            i_uti = val_or_zero(self.get_var(f"inicial_{tk_name}_alt_uti").get())
            i_temp = self.get_var(f"inicial_{tk_name}_temp").get()
            i_temp = self.parse_float(i_temp) 
            i_corr = val_or_zero(self.get_var(f"inicial_{tk_name}_s_corr").get())
            i_agua_real = val_or_zero(self.get_var(f"inicial_{tk_name}_agua_s_real").get())
            d_i_raw = self.get_var(f"inicial_{tk_name}_{density_mode_key}").get()
            if not d_i_raw: d_i_raw = self.get_var(f"inicial_{tk_name}_dens_lab").get()
            d_i = self.parse_float(d_i_raw)
            table_type = self.get_var(f"inicial_{tk_name}_tabla_vcf").get()
            if not table_type: table_type = "54B (Combustibles)"
            i_bruto = self.interpolar_prod("inicial", tk_name)
            i_agua = self.interpolar_agua("inicial", tk_name)
            i_neto = i_bruto - i_agua
            i_vcf = self.calc_vcf(d_i, i_temp, table_type) if d_i > 0 else 1.0
            i_15 = i_neto * i_vcf
            f_ref = val_or_zero(self.get_var(f"final_{tk_name}_alt_ref").get())
            f_uti = val_or_zero(self.get_var(f"final_{tk_name}_alt_uti").get())
            f_temp = self.get_var(f"final_{tk_name}_temp").get()
            f_temp = self.parse_float(f_temp) 
            f_corr = val_or_zero(self.get_var(f"final_{tk_name}_s_corr").get())
            f_agua_real = val_or_zero(self.get_var(f"final_{tk_name}_agua_s_real").get())
            d_f_raw = self.get_var(f"final_{tk_name}_{density_mode_key}").get()
            if not d_f_raw: d_f_raw = self.get_var(f"final_{tk_name}_dens_lab").get()
            d_f = self.parse_float(d_f_raw)
            table_type_f = self.get_var(f"final_{tk_name}_tabla_vcf").get()
            if not table_type_f: table_type_f = "54B (Combustibles)"
            f_bruto = self.interpolar_prod("final", tk_name)
            f_agua = self.interpolar_agua("final", tk_name)
            f_neto = f_bruto - f_agua
            f_vcf = self.calc_vcf(d_f, f_temp, table_type_f) if d_f > 0 else 1.0
            f_15 = f_neto * f_vcf
            
            # Construir vals_i y vals_f según tipo
            def draw_data_row(start_x, vals):
                curr_x = start_x
                for i, v in enumerate(vals):
                    w_curr = col_w_name if i == 0 else col_w_num
                    c.drawString(curr_x, y_row, str(v))
                    curr_x += w_curr
            
            if _es_liq_rep:
                vals_i = [tk_name, i_ref, i_uti, i_agua_real, f"{i_temp:.1f}" if i_temp else "0", i_corr, f"{i_bruto:.0f}", f"{i_agua:.0f}", f"{i_neto:.0f}", f"{i_vcf:.4f}", f"{i_15:.0f}"]
                vals_f = [tk_name, f_ref, f_uti, f_agua_real, f"{f_temp:.1f}" if f_temp else "0", f_corr, f"{f_bruto:.0f}", f"{f_agua:.0f}", f"{f_neto:.0f}", f"{f_vcf:.4f}", f"{f_15:.0f}"]
            elif _es_el_rep_t:
                def _eg(k): return val_or_zero(self.get_var(k).get())
                vals_i = [tk_name, _eg(f"inicial_{tk_name}_el_kwh_act"), _eg(f"inicial_{tk_name}_el_kwh_rea"),
                          _eg(f"inicial_{tk_name}_el_fp"), _eg(f"inicial_{tk_name}_el_dem"),
                          _eg(f"inicial_{tk_name}_el_V"), _eg(f"inicial_{tk_name}_el_A"),
                          _eg(f"inicial_{tk_name}_el_VA"), _eg(f"inicial_{tk_name}_el_fases"), "—", "—"]
                vals_f = [tk_name, _eg(f"final_{tk_name}_el_kwh_act"), _eg(f"final_{tk_name}_el_kwh_rea"),
                          _eg(f"final_{tk_name}_el_fp"), _eg(f"final_{tk_name}_el_dem"),
                          _eg(f"final_{tk_name}_el_V"), _eg(f"final_{tk_name}_el_A"),
                          _eg(f"final_{tk_name}_el_VA"), _eg(f"final_{tk_name}_el_fases"), "—", "—"]
            elif _es_duc_rep_t:
                def _dg(k): return val_or_zero(self.get_var(k).get())
                vals_i = [tk_name, _dg(f"inicial_{tk_name}_cont_ini"), _dg(f"inicial_{tk_name}_cont_fin"),
                          _dg(f"inicial_{tk_name}_vol_linea"), _dg(f"inicial_{tk_name}_P_lin"),
                          _dg(f"inicial_{tk_name}_T_lin"), _dg(f"inicial_{tk_name}_Z"),
                          _dg(f"inicial_{tk_name}_vol_base"), _dg(f"inicial_{tk_name}_vol_base_km3"),
                          _dg(f"inicial_{tk_name}_caudal_mh"), _dg(f"inicial_{tk_name}_masa_coriolis")]
                vals_f = [tk_name, _dg(f"final_{tk_name}_cont_ini"), _dg(f"final_{tk_name}_cont_fin"),
                          _dg(f"final_{tk_name}_vol_linea"), _dg(f"final_{tk_name}_P_lin"),
                          _dg(f"final_{tk_name}_T_lin"), _dg(f"final_{tk_name}_Z"),
                          _dg(f"final_{tk_name}_vol_base"), _dg(f"final_{tk_name}_vol_base_km3"),
                          _dg(f"final_{tk_name}_caudal_mh"), _dg(f"final_{tk_name}_masa_coriolis")]
            elif _es_gas_rep_t:
                def _gg(k): return val_or_zero(self.get_var(k).get())
                vals_i = [tk_name, _gg(f"inicial_{tk_name}_prod_name"), _gg(f"inicial_{tk_name}_presion"),
                          _gg(f"inicial_{tk_name}_temp_liq"), _gg(f"inicial_{tk_name}_dens_liq"),
                          _gg(f"inicial_{tk_name}_vol_liq"), _gg(f"inicial_{tk_name}_fase"),
                          _gg(f"inicial_{tk_name}_vol_nat_prod"), _gg(f"inicial_{tk_name}_factor_z"), "—", "—"]
                vals_f = [tk_name, _gg(f"final_{tk_name}_prod_name"), _gg(f"final_{tk_name}_presion"),
                          _gg(f"final_{tk_name}_temp_liq"), _gg(f"final_{tk_name}_dens_liq"),
                          _gg(f"final_{tk_name}_vol_liq"), _gg(f"final_{tk_name}_fase"),
                          _gg(f"final_{tk_name}_vol_nat_prod"), _gg(f"final_{tk_name}_factor_z"), "—", "—"]
            elif _es_esf_rep_t:
                def _sfg(k): return val_or_zero(self.get_var(k).get())
                vals_i = [tk_name, _sfg(f"inicial_{tk_name}_esf_pres"), _sfg(f"inicial_{tk_name}_esf_temp"),
                          _sfg(f"inicial_{tk_name}_esf_dens"), _sfg(f"inicial_{tk_name}_vol_liq"),
                          _sfg(f"inicial_{tk_name}_esf_vol_gas"), _sfg(f"inicial_{tk_name}_esf_masa"),
                          _sfg(f"inicial_{tk_name}_esf_fase"), "—", "—", "—"]
                vals_f = [tk_name, _sfg(f"final_{tk_name}_esf_pres"), _sfg(f"final_{tk_name}_esf_temp"),
                          _sfg(f"final_{tk_name}_esf_dens"), _sfg(f"final_{tk_name}_vol_liq"),
                          _sfg(f"final_{tk_name}_esf_vol_gas"), _sfg(f"final_{tk_name}_esf_masa"),
                          _sfg(f"final_{tk_name}_esf_fase"), "—", "—", "—"]
            elif _es_cgb_rep_t:
                def _cgg(k): return val_or_zero(self.get_var(k).get())
                vals_i = [tk_name, _cgg(f"inicial_{tk_name}_cg_pres"), _cgg(f"inicial_{tk_name}_cg_temp"),
                          _cgg(f"inicial_{tk_name}_cg_dens"), _cgg(f"inicial_{tk_name}_cg_masa"),
                          _cgg(f"inicial_{tk_name}_cg_vol"), _cgg(f"inicial_{tk_name}_cg_vol_gas"),
                          _cgg(f"inicial_{tk_name}_fase"), "—", "—", "—"]
                vals_f = [tk_name, _cgg(f"final_{tk_name}_cg_pres"), _cgg(f"final_{tk_name}_cg_temp"),
                          _cgg(f"final_{tk_name}_cg_dens"), _cgg(f"final_{tk_name}_cg_masa"),
                          _cgg(f"final_{tk_name}_cg_vol"), _cgg(f"final_{tk_name}_cg_vol_gas"),
                          _cgg(f"final_{tk_name}_fase"), "—", "—", "—"]
            else:
                # Tierra y camión cisterna
                def _tg(k): return val_or_zero(self.get_var(k).get())
                vals_i = [tk_name, _tg(f"inicial_{tk_name}_num_uti"), _tg(f"inicial_{tk_name}_s_tierra"),
                          _tg(f"inicial_{tk_name}_s_corr"), f"{i_temp:.1f}" if i_temp else "0",
                          _tg(f"inicial_{tk_name}_vol_bruto"), f"{i_agua:.0f}",
                          f"{i_neto:.0f}", f"{i_vcf:.4f}", f"{i_15:.0f}", "—"]
                vals_f = [tk_name, _tg(f"final_{tk_name}_num_uti"), _tg(f"final_{tk_name}_s_tierra"),
                          _tg(f"final_{tk_name}_s_corr"), f"{f_temp:.1f}" if f_temp else "0",
                          _tg(f"final_{tk_name}_vol_bruto"), f"{f_agua:.0f}",
                          f"{f_neto:.0f}", f"{f_vcf:.4f}", f"{f_15:.0f}", "—"]
            draw_data_row(x_start, vals_i)
            draw_data_row(x_mid, vals_f)
            y_row -= 10
            if y_row < 50: c.showPage(); y_row = h - 50
        c.line(35, y_row+5, w-35, y_row+5); y_row -= 5 
        c.setFont("Helvetica-Bold", 8)
        c.setFillColor(colors.blue)
        texto_inicial = f"El total de litros naturales es {sum_ini['neto']:,.0f} y de litros a 15 es {sum_ini['15']:,.0f}"
        c.drawString(x_start, y_row, texto_inicial)
        texto_final = f"El total de litros naturales es {sum_fin['neto']:,.0f} y de litros a 15 es {sum_fin['15']:,.0f}"
        c.drawString(x_mid, y_row, texto_final)
        y_row -= 10
        c.setFillColor(colors.black)
        y_row -= 15

        if is_partial:
            x_transf = x_mid + 245 
            def draw_stat_line(lbl, val_i, val_f, transf_label=None, transf_val=None, color_transf=colors.black):
                c.setFillColor(colors.black)
                c.drawString(x_start, y_row, f"{lbl}: {val_i}")
                c.drawString(x_mid, y_row, f"{lbl}: {val_f}")
                if transf_label:
                    c.setFillColor(colors.black)
                    c.setFont("Helvetica-Bold", 8)
                    c.drawString(x_transf, y_row, f"{transf_label}: ")
                    if transf_val is not None:
                        c.setFillColor(color_transf)
                        c.setFont("Helvetica", 8)
                        width_lbl = c.stringWidth(f"{transf_label}: ", "Helvetica-Bold", 8)
                        c.drawString(x_transf + width_lbl, y_row, str(transf_val))
                c.setFillColor(colors.black) 
            
            c.setFont("Helvetica", 8)
            draw_stat_line("Total Litros Naturales Netos", f"{sum_ini['neto']:,.0f}", f"{sum_fin['neto']:,.0f}", "Diferencia Litros Naturales", f"{dif_net:,.0f}")
            y_row -= 12
            draw_stat_line("Total Litros a 15º C", f"{sum_ini['15']:,.0f}", f"{sum_fin['15']:,.0f}", "Diferencia Litros a 15 C", f"{dif_15:,.0f}")
            y_row -= 12
            draw_stat_line("Densidad del producto", f"{avg_d_i:.4f}", f"{avg_d_f:.4f}", "Diferencia Kilos Aire", f"{dif_ka:,.0f}")
            y_row -= 12
            draw_stat_line("Total Kilos Aire", f"{ka_i:,.0f}", f"{ka_f:,.0f}", "Diferencia Kilos Vacío", f"{dif_kv:,.0f}")
            y_row -= 12
            c.setFont("Helvetica-Bold", 8)
            c.drawString(x_start, y_row, f"Total Kilogramos al Vacío: {kv_i:,.0f}")
            c.drawString(x_mid, y_row, f"Total Kilogramos al Vacío: {kv_f:,.0f}")
            dens_transf = avg_d_f if avg_d_f > 0 else avg_d_i
            c.drawString(x_transf, y_row, "Densidad Doc: ")
            c.setFont("Helvetica", 8)
            c.drawString(x_transf + c.stringWidth("Densidad Doc: ", "Helvetica-Bold", 8), y_row, f"{dens_transf:.4f}")
            y_row -= 12
            y_row -= 25
        else:
            c.setFont("Helvetica-Bold", 10)
            c.drawString(53, y_row, "Producto por Tanque:")
            y_row -= 15
            mapa_ddt = {}
            for tk_name in tank_list:
                d_assign = self.get_var(f"inicial_{tk_name}_ddt_assign").get()
                if d_assign:
                    if d_assign not in mapa_ddt: mapa_ddt[d_assign] = []
                    mapa_ddt[d_assign].append(tk_name)
            c.setFont("Helvetica", 8)
            for d_num, tks in mapa_ddt.items():
                found = next((d for d in self.ddt_data if d["numero"].get() == d_num), None)
                if found:
                    prod = found["producto"].get()
                    pos = found["pos_arancel"].get()
                    tipo = found["tipo_doc"].get() if "tipo_doc" in found else "Detallada"
                    chunks = [tks[i:i + 4] for i in range(0, len(tks), 4)]
                    first_chunk = ", ".join(chunks[0])
                    linea = f"Documento {tipo}: {d_num} - {prod} - Pos. Arancelaria: {pos} - Tanques: {first_chunk}"
                    c.drawString(53, y_row, linea)
                    y_row -= 12
                    for chunk in chunks[1:]:
                        linea_cont = f"       (Cont. Tanques): {', '.join(chunk)}"
                        c.drawString(53, y_row, linea_cont)
                        y_row -= 12
        y_sig_line = 40
        y_sig_text = 30
        c.setLineWidth(1)
        c.setFillColor(colors.black)
        c.line(40, y_sig_line, 190, y_sig_line)
        c.drawCentredString(115, y_sig_text, "Aduana Inicial")
        c.line(240, y_sig_line, 390, y_sig_line)
        c.drawCentredString(315, y_sig_text, "Interesado Inicial")
        c.line(440, y_sig_line, 590, y_sig_line)
        c.drawCentredString(515, y_sig_text, "Aduana Final")
        c.line(640, y_sig_line, 790, y_sig_line)
        c.drawCentredString(715, y_sig_text, "Interesado Final")
        c.showPage()
        if not shared_canvas:
            c.save()
            return full_path
        return None

    def _generar_cargo_docx(self, ddt_obj, dif_kv, target_k,
                            ctrl_doc=True, ctrl_sal=False, ctrl_lab=False,
                            modo_comp_forzado=None, tipo_operacion="importacion",
                            tipo_operacion_info=None, output_folder=None):
        """Genera el informe de cargo/denuncia como documento Word (.docx)."""
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_ALIGN_VERTICAL
        except ImportError:
            messagebox.showerror("Error", "python-docx no está instalado.\nInstale con: pip install python-docx")
            return None

        if target_k == 0: return None
        diff_k = dif_kv - target_k
        permil_k = (diff_k / target_k * 1000) if target_k != 0 else 0
        if abs(permil_k) <= 6.0: return None
        diff_abs_k = abs(diff_k)
        pct_k = (diff_abs_k / target_k * 100) if target_k != 0 else 0
        es_faltante = diff_k < 0
        tipo_diff = "FALTANTE" if es_faltante else "SOBRANTE"
        es_denuncia = pct_k >= 2.0
        modo_comp = modo_comp_forzado or ("laboratorio" if ctrl_lab else ("salida" if ctrl_sal else "documento"))

        # Info operación
        if tipo_operacion_info:
            art_principal  = tipo_operacion_info.get("art_principal", "Art. 954 del Código Aduanero")
            art_infraccion = tipo_operacion_info.get("art_inc", "Art. 954 inc. c) C.A.")
            desc_operacion = tipo_operacion_info.get("descripcion", "OPERACIÓN ADUANERA")
            cod_subregimen = tipo_operacion_info.get("codigo", "")
        elif tipo_operacion in ("exportacion", "remo_carga"):
            art_principal = "Art. 959 del Código Aduanero"; art_infraccion = "Art. 959 inc. c) C.A."
            desc_operacion = "EXPORTACIÓN"; cod_subregimen = ""
        else:
            art_principal = "Art. 954 del Código Aduanero"; art_infraccion = "Art. 954 inc. c) C.A."
            desc_operacion = "IMPORTACIÓN"; cod_subregimen = ""

        aduana    = self.aduana_nombre()
        buque     = self.get_var('car_buque').get()
        operador  = self._ddt_actor(ddt_obj, "impexp")
        dest_num  = ddt_obj["numero"].get()
        producto  = ddt_obj["producto"].get()
        doc_k     = self.parse_float(ddt_obj["kilos"].get())
        fecha_op  = self.get_var('car_fecha').get()
        fecha_hoy = datetime.now().strftime("%d/%m/%Y")
        tipo_doc_str = "DENUNCIA" if es_denuncia else "CARGO TRIBUTARIO"
        modo_labels   = {"documento": "Documento declarado", "salida": "Salida de Zona Primaria", "laboratorio": "Análisis de Laboratorio"}
        _vl_var = ddt_obj.get("valor_litro")
        valor_kg = self.parse_float(_vl_var.get() if isinstance(_vl_var, tk.StringVar) else str(_vl_var or "0"))
        divisa   = ddt_obj.get("divisa", tk.StringVar(value="ARS"))
        divisa   = divisa.get() if isinstance(divisa, tk.StringVar) else str(divisa)
        divisa_sym = {"ARS": "$", "USD": "U$S", "EUR": "€", "BRL": "R$", "GBP": "£"}.get(divisa, "$")
        _tc_var = ddt_obj.get("tipo_cambio")
        tc_ddt  = self.parse_float(_tc_var.get() if isinstance(_tc_var, tk.StringVar) else str(_tc_var or "0"))
        # T/C global del documento (carátula) para la línea de equivalencia en USD
        tc_global = self.parse_float(self.get_var("car_tipo_cambio", "").get())
        # Regla: para divisa!=ARS usamos el T/C del DDT (convierte precio a ARS)
        #        para la línea "Equiv. en USD" usamos SIEMPRE el T/C global del documento
        tipo_cambio = tc_ddt if (divisa != "ARS" and tc_ddt > 0) else tc_global

        doc = Document()

        # ── Márgenes ──
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        def add_heading(text, level=1, color=None):
            p = doc.add_heading(text, level=level)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
            if color:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(*color)
            return p

        def add_para(text, bold=False, italic=False, size=10, color=None, align=None):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = bold; run.italic = italic; run.font.size = Pt(size)
            if color: run.font.color.rgb = RGBColor(*color)
            if align: p.alignment = align
            return p

        def add_table_row(table, cells_data):
            row = table.add_row()
            for i, (text, bold, bg) in enumerate(cells_data):
                cell = row.cells[i]
                cell.text = text
                if bold:
                    for run in cell.paragraphs[0].runs: run.bold = True
                cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            return row

        # ── ENCABEZADO ──
        color_warn = (192, 57, 43) if es_denuncia else (230, 126, 34)
        add_heading(f"INFORME DE {tipo_doc_str}", level=1, color=color_warn)
        add_para(f"Generado: {fecha_hoy}", italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

        # ── DATOS ──
        add_heading("I. DATOS DE LA OPERACIÓN", level=2)
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        tbl.columns[0].width = Cm(6); tbl.columns[1].width = Cm(11)
        rows_data = [
            ("Aduana", aduana), ("Buque / Nave", buque), ("Operador / Imp-Exp", operador),
            ("Documento", dest_num), ("Subregimen", f"{desc_operacion} {cod_subregimen}".strip()),
            ("Producto", producto), ("Fecha operación", fecha_op),
            ("Base de comparación", modo_labels.get(modo_comp, modo_comp)),
        ]
        for lbl, val in rows_data:
            row = tbl.add_row()
            row.cells[0].text = lbl; row.cells[1].text = str(val)
            for run in row.cells[0].paragraphs[0].runs: run.bold = True
        doc.add_paragraph()

        # ── DIFERENCIA ──
        add_heading("II. DIFERENCIA DETECTADA", level=2)
        tbl2 = doc.add_table(rows=0, cols=2); tbl2.style = "Table Grid"
        tbl2.columns[0].width = Cm(8); tbl2.columns[1].width = Cm(9)
        diff_rows = [
            ("Tipo de diferencia", tipo_diff),
            ("Diferencia en Kilos", f"{diff_abs_k:,.0f} kg"),
            ("Por mil (‰)", f"{abs(permil_k):.2f}‰"),
            ("Porcentaje (%)", f"{pct_k:.2f}%"),
            ("Total declarado (kg)", f"{doc_k:,.0f} kg"),
            ("Clasificación", "DENUNCIA (>=2%)" if es_denuncia else "CARGO TRIBUTARIO (entre 6 o/oo y 2%)"),
        ]
        for lbl, val in diff_rows:
            row = tbl2.add_row()
            row.cells[0].text = lbl; row.cells[1].text = val
            for run in row.cells[0].paragraphs[0].runs: run.bold = True
        doc.add_paragraph()

        # ── ENCUADRE TÉCNICO ──
        add_heading("III. ENCUADRE TÉCNICO Y NORMATIVO", level=2)
        if es_denuncia:
            add_para(
                f"La diferencia detectada ({diff_abs_k:,.0f} kg) representa el {pct_k:.2f}% del total declarado, "
                f"SUPERANDO el umbral del 2%, lo que configura la presunción de infracción al {art_principal} "
                f"por declaración inexacta de cantidad de mercadería.", size=10)
        else:
            add_para(
                f"La diferencia detectada ({diff_abs_k:,.0f} kg) representa el {pct_k:.2f}% del total declarado, "
                f"superando la franquicia técnica del 6‰ pero resultando INFERIOR al 2%, "
                f"por lo cual es de aplicación la excusa absolutoria del {art_infraccion}. "
                f"No obstante, en materia tributaria existe un hecho imponible que genera el cargo correspondiente.", size=10)
        doc.add_paragraph()
        add_para("Marco Normativo:", bold=True, size=10)
        normas = [
            f"• {art_principal}: Infracción por declaración inexacta de cantidad.",
            "• Arts. 637/638 C.A.: Determinación y nacimiento de la obligación tributaria.",
            "• Arts. 790/791 C.A.: Liquidación de tributos y actuación aduanera.",
            "• Res. ex-ANA 2220/90: Tolerancia técnica 0,6% para graneles líquidos.",
        ]
        for n in normas:
            doc.add_paragraph(n, style="List Bullet")
        doc.add_paragraph()

        # ── LIQUIDACIÓN ──
        # ── Conversión de divisa ──────────────────────────────────────────────
        # Si el precio está en moneda extranjera, convertir a ARS usando el
        # tipo de cambio del documento para la base imponible final.
        usa_tc = (divisa != "ARS") and (tipo_cambio > 0)
        if usa_tc:
            # valor_kg_ars: precio por kg en pesos (para la liquidación ARS)
            valor_kg_ars = valor_kg * tipo_cambio
            tc_nota = (f"T/C: {divisa_sym} 1 = $ {tipo_cambio:,.2f}  →  "
                       f"Precio en ARS: {divisa_sym} {valor_kg:,.4f} × {tipo_cambio:,.2f} = $ {valor_kg_ars:,.4f}/Kg")
        else:
            valor_kg_ars = valor_kg   # ya en ARS, sin conversión
            tc_nota = ""

        if es_denuncia:
            add_heading("IV. BASE IMPONIBLE ESTIMADA — DENUNCIA", level=2)
            base_k = doc_k if doc_k > 0 else diff_abs_k
            add_para(f"La base imponible correspondería a la totalidad de la carga: {base_k:,.0f} kg", size=10)
            if valor_kg > 0:
                base_imp_orig = base_k * valor_kg          # en divisa original
                base_imp      = base_k * valor_kg_ars      # en ARS (= base_imp_orig si ya es ARS)
                if usa_tc:
                    add_para(tc_nota, size=9, italic=True)
                    add_para(
                        f"Total: {base_k:,.0f} Kg × {divisa_sym} {valor_kg:,.4f}/Kg"
                        f" = {divisa_sym} {base_imp_orig:,.2f}"
                        f"  →  $ {base_imp:,.2f} (tipo de cambio $ {tipo_cambio:,.2f})",
                        bold=True, size=10)
                else:
                    add_para(f"Total: {base_k:,.0f} Kg × $ {valor_kg:,.4f}/Kg = $ {base_imp:,.2f}", bold=True, size=10)
                tbl3 = doc.add_table(rows=0, cols=4); tbl3.style = "Table Grid"
                tbl3.add_row().cells[0].text = "CONCEPTO"
                tbl3.rows[0].cells[1].text = "ALÍC."; tbl3.rows[0].cells[2].text = "BASE (ARS)"; tbl3.rows[0].cells[3].text = "IMPORTE EST. (ARS)"
                for run in tbl3.rows[0].cells[0].paragraphs[0].runs: run.bold = True
                total_t = 0
                for conc, alic in (self.get_tributos_activos() or [("Derechos de Imp/Exp (ref.)", 8.0)]):
                    imp = base_imp * alic / 100; total_t += imp
                    r = tbl3.add_row()
                    r.cells[0].text = conc; r.cells[1].text = f"{alic}%"
                    r.cells[2].text = f"$ {base_imp:,.2f}"; r.cells[3].text = f"$ {imp:,.2f}"
                r = tbl3.add_row()
                r.cells[0].text = "TOTAL ESTIMADO"; r.cells[3].text = f"$ {total_t:,.2f}"
                for run in r.cells[0].paragraphs[0].runs: run.bold = True
        else:
            add_heading("IV. LIQUIDACIÓN ESTIMADA DE CARGO", level=2)
            add_para(f"El cargo tributario corresponde al {tipo_diff.lower()} de {diff_abs_k:,.0f} kg.", size=10)
            if valor_kg > 0:
                base_imp_orig = diff_abs_k * valor_kg      # en divisa original
                base_imp      = diff_abs_k * valor_kg_ars  # en ARS
                if usa_tc:
                    add_para(tc_nota, size=9, italic=True)
                    add_para(
                        f"{tipo_diff}: {diff_abs_k:,.0f} Kg × {divisa_sym} {valor_kg:,.4f}/Kg"
                        f" = {divisa_sym} {base_imp_orig:,.2f}"
                        f"  →  $ {base_imp:,.2f} (tipo de cambio $ {tipo_cambio:,.2f})",
                        bold=True, size=10)
                else:
                    add_para(f"{tipo_diff}: {diff_abs_k:,.0f} Kg × $ {valor_kg:,.4f}/Kg = $ {base_imp:,.2f}", bold=True, size=10)
                tbl3 = doc.add_table(rows=0, cols=4); tbl3.style = "Table Grid"
                r0 = tbl3.add_row()
                r0.cells[0].text = "CONCEPTO"; r0.cells[1].text = "ALÍC."
                r0.cells[2].text = "BASE (ARS)"; r0.cells[3].text = "IMPORTE EST. (ARS)"
                for run in r0.cells[0].paragraphs[0].runs: run.bold = True
                total_t = 0
                for conc, alic in (self.get_tributos_activos() or [("Derechos de Imp/Exp (ref.)", 8.0)]):
                    imp = base_imp * alic / 100; total_t += imp
                    r = tbl3.add_row()
                    r.cells[0].text = conc; r.cells[1].text = f"{alic}%"
                    r.cells[2].text = f"$ {base_imp:,.2f}"; r.cells[3].text = f"$ {imp:,.2f}"
                r = tbl3.add_row()
                r.cells[0].text = f"TOTAL ESTIMADO {tipo_diff}"; r.cells[3].text = f"$ {total_t:,.2f}"
                for run in r.cells[0].paragraphs[0].runs: run.bold = True
        doc.add_paragraph()

        # ── CONCLUSIÓN ──
        add_heading("V. CONCLUSIÓN", level=2)
        if es_denuncia:
            add_para(
                f"Se elevarían las presentes actuaciones a fin de que la superioridad evaluara la posible "
                f"instrucción de sumario al amparo del {art_principal}, en virtud del {tipo_diff.lower()} "
                f"detectado de {diff_abs_k:,.0f} kilogramos ({pct_k:.2f}% del total declarado), "
                f"sin perjuicio del cargo tributario que eventualmente correspondiere. Es todo cuanto se informaría.", size=10)
        else:
            add_para(
                f"Se elevarían las presentes actuaciones a fin de que la superioridad evaluara la procedencia "
                f"del cargo tributario en los términos del {art_principal}, en virtud del {tipo_diff.lower()} "
                f"detectado de {diff_abs_k:,.0f} kilogramos ({pct_k:.2f}% del total declarado). Es todo cuanto se informaría.", size=10)
        doc.add_paragraph()
        add_para(f"Fecha: {fecha_hoy}", italic=True, size=9)

        # ── GUARDAR ──
        clean_name  = self.clean_filename(buque)
        fecha_arch  = datetime.now().strftime("%Y%m%d")
        tipo_arch   = "Denuncia" if es_denuncia else "Cargo"
        docx_file   = f"{tipo_arch}_Tributario_{clean_name}_{self.clean_filename(dest_num)}_{fecha_arch}.docx"
        docx_path   = os.path.join(output_folder or ".", docx_file)
        doc.save(docx_path)
        return docx_path

    def detectar_tolerancia_y_cargo(self, ddt_obj, dif_kv, target_k, output_folder,
                                   ctrl_doc=True, ctrl_sal=False, ctrl_lab=False,
                                   modo_comp_forzado=None, shared_canvas=None,
                                   tipo_operacion_override=None, tipo_operacion_info=None):
        """Detecta si excede tolerancia (en KILOS) y genera cargo o denuncia."""
        if target_k == 0: return
        diff_k = dif_kv - target_k
        permil_k = (diff_k / target_k * 1000) if target_k != 0 else 0
        if abs(permil_k) <= 6.0: return
        diff_abs_k = abs(diff_k)
        pct_k = (diff_abs_k / target_k * 100) if target_k != 0 else 0
        es_faltante = diff_k < 0
        tipo_diff = "FALTANTE" if es_faltante else "SOBRANTE"
        es_denuncia = pct_k >= 2.0
        modo_comp = modo_comp_forzado or ("laboratorio" if ctrl_lab else ("salida" if ctrl_sal else "documento"))

        if tipo_operacion_override:
            return self._generar_cargo_pdf(ddt_obj, diff_abs_k, permil_k,
                                           pct_k, tipo_diff, es_faltante, output_folder, modo_comp,
                                           es_denuncia=es_denuncia, tipo_operacion=tipo_operacion_override,
                                           shared_canvas=shared_canvas, tipo_operacion_info=tipo_operacion_info)

        # Diálogo: solo importación/exportación
        dlg = tk.Toplevel(self.root)
        dlg.title("Tolerancia Excedida - Confirmar Cargo / Denuncia")
        dlg.transient(self.root)
        dlg.grab_set()
        dlg.update_idletasks()
        sw, sh = dlg.winfo_screenwidth(), dlg.winfo_screenheight()
        w_dlg, h_dlg = 560, 280
        dlg.geometry(f"{w_dlg}x{h_dlg}+{(sw-w_dlg)//2}+{(sh-h_dlg)//2}")
        color_warn = "#C0392B" if es_denuncia else "#E67E22"
        titulo = "DENUNCIA - SUPERA 2% EN KILOS" if es_denuncia else "CARGO TRIBUTARIO - Entre 6 o/oo y 2% en Kilos"
        tk.Label(dlg, text=titulo, font=("Arial", 8, "bold"), fg=color_warn, bg="#FDF2F2").pack(fill="x", ipady=8)
        comp_labels = {"documento": "Documento declarado", "salida": "Salida de Zona Primaria", "laboratorio": "Análisis de Laboratorio"}
        dest_num_val = ddt_obj["numero"].get()
        lbl_info  = f"Destino: {dest_num_val}\n"
        lbl_info += f"Diferencia: {diff_abs_k:,.0f} kg {tipo_diff}  ({abs(permil_k):.2f} o/oo  |  {pct_k:.2f}%)\n"
        lbl_info += f"Comparación: {comp_labels.get(modo_comp, modo_comp)}"
        tk.Label(dlg, text=lbl_info, font=("Arial", 10), justify="left").pack(padx=20, pady=8)
        f_op = ttk.LabelFrame(dlg, text="Tipo de operación")
        f_op.pack(fill="x", padx=20, pady=4)
        var_op = tk.StringVar(value="importacion")
        tk.Radiobutton(f_op, text="Importación  (Art. 954 C.A.)", variable=var_op, value="importacion", font=("Arial", 10)).pack(side="left", padx=20)
        tk.Radiobutton(f_op, text="Exportación  (Art. 959 C.A.)", variable=var_op, value="exportacion", font=("Arial", 10)).pack(side="left", padx=20)
        resultado = {"generar": False, "operacion": "importacion"}
        def aceptar():
            resultado["generar"] = True; resultado["operacion"] = var_op.get(); dlg.destroy()
        def cancelar(): dlg.destroy()
        lbl_btn = "GENERAR DENUNCIA" if es_denuncia else "GENERAR CARGO"
        f_bot = tk.Frame(dlg); f_bot.pack(pady=10)
        tk.Button(f_bot, text=lbl_btn, bg="#C0392B", fg="white", font=("Arial", 8, "bold"), command=aceptar).pack(side="left", padx=10)
        tk.Button(f_bot, text="Cancelar", bg="gray", fg="white", font=("Arial", 10), command=cancelar).pack(side="left", padx=10)
        dlg.wait_window()
        if not resultado["generar"]: return
        return self._generar_cargo_pdf(ddt_obj, diff_abs_k, permil_k,
                                pct_k, tipo_diff, es_faltante, output_folder, modo_comp,
                                es_denuncia=es_denuncia, tipo_operacion=resultado["operacion"],
                                shared_canvas=shared_canvas)

    def _generar_cargo_en_canvas(self, c, ddt_obj, dif_kv, target_k,
                                  ctrl_doc=True, ctrl_sal=False, ctrl_lab=False,
                                  modo_comp_forzado=None, tipo_operacion="importacion", tipo_operacion_info=None):
        """Calcula tolerancia en KILOS y si supera dibuja cargo/denuncia en el canvas c."""
        if target_k == 0: return
        diff_k = dif_kv - target_k
        permil_k = (diff_k / target_k * 1000) if target_k != 0 else 0
        if abs(permil_k) <= 6.0: return
        diff_abs_k = abs(diff_k)
        pct_k = (diff_abs_k / target_k * 100) if target_k != 0 else 0
        es_faltante = diff_k < 0
        tipo_diff = "FALTANTE" if es_faltante else "SOBRANTE"
        es_denuncia = pct_k >= 2.0
        modo_comp = modo_comp_forzado or ("laboratorio" if ctrl_lab else ("salida" if ctrl_sal else "documento"))
        self._generar_cargo_pdf(
            ddt_obj, diff_abs_k, permil_k,
            pct_k, tipo_diff, es_faltante, output_folder=None, modo_comp=modo_comp,
            es_denuncia=es_denuncia, tipo_operacion=tipo_operacion,
            shared_canvas=c, tipo_operacion_info=tipo_operacion_info
        )


    def _generar_cargo_pdf(self, ddt_obj, diff_abs_k, permil_k,
                           pct_max, tipo_diff, es_faltante, output_folder, modo_comp,
                           es_denuncia=False, tipo_operacion="importacion", shared_canvas=None,
                           tipo_operacion_info=None, _buf_override=None):
        """Genera el informe de cargo/denuncia — todo en KILOS, sin litros.
        Si shared_canvas: dibuja en él y retorna None.
        Si output_folder: guarda a archivo y retorna la ruta.
        """
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.utils import simpleSplit
        import io
        
        clean_name = self.clean_filename(self.get_var('car_buque').get())
        dest_num = ddt_obj["numero"].get()
        tipo_doc_str = "Denuncia" if es_denuncia else "Cargo"
        filename = f"{tipo_doc_str}_Tributario_{clean_name}_{self.clean_filename(dest_num)}_{datetime.now().strftime('%Y%m%d')}.pdf"
        full_path = os.path.join(output_folder, filename) if output_folder else None
        
        # Artículos y descripción de operación
        if tipo_operacion_info:
            art_principal  = tipo_operacion_info.get("art_principal", "Art. 954 del Código Aduanero")
            art_infraccion = tipo_operacion_info.get("art_inc",        "Art. 954 inc. c) C.A.")
            art_denuncia   = tipo_operacion_info.get("art_principal",  "Art. 954 del Código Aduanero")
            desc_operacion = tipo_operacion_info.get("descripcion",    "OPERACIÓN ADUANERA")
            cod_subregimen = tipo_operacion_info.get("codigo",         "")
        elif tipo_operacion in ("exportacion", "remo_carga"):
            art_principal  = "Art. 959 del Código Aduanero"
            art_infraccion = "Art. 959 inc. c) C.A."
            art_denuncia   = "Art. 959 del Código Aduanero"
            desc_operacion = "EXPORTACIÓN"
            cod_subregimen = ""
        else:
            art_principal  = "Art. 954 del Código Aduanero"
            art_infraccion = "Art. 954 inc. c) C.A."
            art_denuncia   = "Art. 954 del Código Aduanero"
            desc_operacion = "IMPORTACIÓN"
            cod_subregimen = ""
        
        standalone = shared_canvas is None
        if shared_canvas is not None:
            # Dibujar directamente en el canvas compartido (ya en portrait A4)
            c = shared_canvas
            c.setPageSize(A4)  # asegurar portrait
        else:
            _buf = _buf_override if _buf_override is not None else io.BytesIO()
            try:
                c = canvas.Canvas(_buf, pagesize=A4)
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo crear buffer de cargo: {e}")
                return
        
        w, h = A4  # 595 x 842
        margin = 50
        usable_w = w - 2 * margin
        
        aduana = self.aduana_nombre()
        buque = self.get_var('car_buque').get() or self.get_var('car_patente').get() or "S/I"
        operador = self._ddt_actor(ddt_obj, "impexp")
        cuit_op = self._ddt_actor(ddt_obj, "cuit_impexp")
        ata = self._ddt_actor(ddt_obj, "ata")
        cuit_ata = self._ddt_actor(ddt_obj, "cuit_ata")
        fecha_med = self.get_var('final_Fecha').get() or self.get_var('inicial_Fecha').get()
        
        _vl_var2 = ddt_obj.get("valor_litro")
        valor_litro = self.parse_float(_vl_var2.get() if isinstance(_vl_var2, tk.StringVar) else str(_vl_var2 or "0"))
        divisa = ddt_obj.get("divisa", tk.StringVar(value="ARS")).get() if isinstance(ddt_obj.get("divisa"), tk.StringVar) else ddt_obj.get("divisa", "ARS")
        divisa_desc = ddt_obj.get("divisa_desc", tk.StringVar()).get() if isinstance(ddt_obj.get("divisa_desc"), tk.StringVar) else ddt_obj.get("divisa_desc", "")
        if divisa == "Otra" and divisa_desc: divisa = divisa_desc
        tipo_cambio_ddt = self.parse_float(ddt_obj.get("tipo_cambio", tk.StringVar()).get() if isinstance(ddt_obj.get("tipo_cambio"), tk.StringVar) else ddt_obj.get("tipo_cambio", "0"))
        tc_global_pdf   = self.parse_float(self.get_var("car_tipo_cambio", "").get())
        # Para divisa!=ARS: T/C del DDT convierte precio a ARS; para equiv USD: T/C global
        tipo_cambio = tipo_cambio_ddt if (divisa != "ARS" and tipo_cambio_ddt > 0) else tc_global_pdf
        
        doc_l = self.parse_float(ddt_obj['litros'].get())
        doc_k = self.parse_float(ddt_obj['kilos'].get())
        dens_doc = self.parse_float(ddt_obj['densidad'].get())
        producto_desc = ddt_obj.get('producto', tk.StringVar()).get() if isinstance(ddt_obj.get('producto'), tk.StringVar) else str(ddt_obj.get('producto', ''))
        pos_arancel = ddt_obj.get('pos_arancel', tk.StringVar()).get() if isinstance(ddt_obj.get('pos_arancel'), tk.StringVar) else str(ddt_obj.get('pos_arancel', ''))
        # Salidas declaradas en la destinación
        salidas_list = []
        for s in (ddt_obj.get('salidas') or []):
            s_num = s['numero'].get() if isinstance(s.get('numero'), tk.StringVar) else str(s.get('numero',''))
            s_k   = self.parse_float(s['kilos'].get()  if isinstance(s.get('kilos'),   tk.StringVar) else s.get('kilos',  '0'))
            s_l   = self.parse_float(s['litros'].get() if isinstance(s.get('litros'),  tk.StringVar) else s.get('litros', '0'))
            s_d   = self.parse_float(s['densidad'].get() if isinstance(s.get('densidad'), tk.StringVar) else s.get('densidad','0'))
            if s_num or s_k or s_l:
                salidas_list.append({'numero': s_num, 'kilos': s_k, 'litros': s_l, 'densidad': s_d})
        sal_k_total = sum(s['kilos']  for s in salidas_list)
        sal_l_total = sum(s['litros'] for s in salidas_list)
        
        def draw_text(txt, y_pos, font_name="Helvetica", font_size=10, bold=False, centered=False):
            fn = f"{font_name}-Bold" if bold else font_name
            c.setFont(fn, font_size)
            if centered:
                c.drawCentredString(w/2, y_pos, txt)
            else:
                c.drawString(margin, y_pos, txt)
            return y_pos - font_size - 4
        
        def draw_wrapped(txt, y_pos, font_name="Helvetica", font_size=10, indent=0):
            c.setFont(font_name, font_size)
            lines = simpleSplit(txt, font_name, font_size, usable_w - indent)
            for line in lines:
                c.drawString(margin + indent, y_pos, line)
                y_pos -= font_size + 3
            return y_pos
        
        def new_page_if_needed(y_pos, needed=80):
            if y_pos < needed:
                c.showPage()
                return h - 50
            return y_pos
        
        # === I. TÍTULO Y ENCABEZADO ===
        # ======================================================
        # === I. ENCABEZADO
        # ======================================================
        y = h - 50
        # Datos de lugar operativo
        lop_cod  = self.get_var("car_lop_codigo").get().strip()
        lop_desc = self.get_var("car_lop_desc").get().strip()
        lugar_str = aduana
        if lop_cod or lop_desc:
            lugar_str = f"{aduana} — Lugar Operativo: {lop_cod} {lop_desc}".strip(" —")

        tipo_doc_titulo = "INFORME DE DENUNCIA POR POSIBLE INFRACCIÓN" if es_denuncia else "INFORME DE CARGO TRIBUTARIO"
        aduana_full = self.aduana_completa()  # "067 - USHUAIA"
        aduana_cod_h = self.aduana_codigo()   # "067"
        aduana_nom_h = self.aduana_nombre()   # "USHUAIA"
        # Primer renglón: solo aduana (código + nombre)
        y = draw_text(f"ADUANA {aduana_cod_h} - {aduana_nom_h.upper()} — DIRECCIÓN GENERAL DE ADUANAS — ARCA", y, font_size=11, bold=True, centered=True)
        # Segundo renglón: Lugar Operativo (si existe), separado
        if lop_cod or lop_desc:
            lop_str = f"Lugar Operativo: {lop_cod}  {lop_desc}".strip()
            y = draw_text(lop_str, y, font_size=9, centered=True)
        y = draw_text(f"{aduana_nom_h}, {fecha_med}", y, font_size=9, centered=True)
        y -= 8
        c.setFont("Helvetica-Bold", 13)
        c.drawCentredString(w/2, y, tipo_doc_titulo)
        y -= 16
        # Tipo de operación resaltado
        _desc_font = 11 if len(desc_operacion) <= 60 else (9 if len(desc_operacion) > 80 else 10)
        c.setFont("Helvetica-Bold", _desc_font)
        c.setFillColor(colors.HexColor("#1B3A5C"))
        c.drawCentredString(w/2, y, f"[ {desc_operacion} ]")
        c.setFillColor(colors.black)
        y -= (_desc_font + 7)
        asunto = (f"Posible infracción al {art_principal} — {tipo_diff.capitalize()} detectado en destinación."
                  if es_denuncia else
                  f"Posible procedencia de Cargo Tributario por {tipo_diff.lower()} en destinación.")
        y = draw_text(f"Ref.: {asunto}", y, font_size=10, bold=True)
        y = draw_text(f"Destinación N°: {dest_num}", y, font_size=10)
        y = draw_text(f"Operador: {operador}  —  CUIT: {cuit_op}", y, font_size=10)
        y = draw_text(f"Agente de Transporte Aduanero: {ata}  —  CUIT: {cuit_ata}", y, font_size=10)
        tipo_nave = self.label_unidad()
        tm_cargo = self.get_tipo_medio()
        lugar_body = aduana_nom_h
        if lop_cod or lop_desc:
            lugar_body = f"{aduana_nom_h} / Lugar Op.: {lop_cod} {lop_desc}".strip(" /")
        y = draw_text(f"Lugar: {lugar_body}  -  {tipo_nave}: {buque}", y, font_size=10)
        y -= 12

        # I.a FUNCIONARIOS INTERVINIENTES
        if self.funcionarios_data:
            y = new_page_if_needed(y, 120)
            y = draw_text("I. FUNCIONARIOS INTERVINIENTES", y, font_size=11, bold=True)
            y -= 4
            y = draw_wrapped(
                "Los abajo firmantes, en calidad de funcionarios de la Dirección General de Aduanas "
                "— ARCA, certificarían haber intervenido en la operación de medición de referencia "
                "ejerciendo las funciones que se detallan:",
                y, font_size=10)
            y -= 6
            for f_obj in self.funcionarios_data:
                apell = f_obj.get("apellido", tk.StringVar()).get() if isinstance(f_obj.get("apellido"), tk.StringVar) else ""
                nom = f_obj.get("nombre", tk.StringVar()).get() if isinstance(f_obj.get("nombre"), tk.StringVar) else ""
                nombre_completo = f"{apell.upper()}, {nom}".strip(", ")
                legajo = f_obj["legajo"].get() if isinstance(f_obj.get("legajo"), tk.StringVar) else str(f_obj.get("legajo", ""))
                cuil = f_obj["cuil"].get() if isinstance(f_obj.get("cuil"), tk.StringVar) else str(f_obj.get("cuil", ""))
                funcion = f_obj["funcion"].get() if isinstance(f_obj.get("funcion"), tk.StringVar) else str(f_obj.get("funcion", ""))
                linea = f"• {nombre_completo}  —  Leg. {legajo}  —  CUIL {cuil}  —  Función: {funcion}"
                y = draw_wrapped(linea, y, font_size=9, indent=10)
                y -= 2
            y -= 8

        # ======================================================
        # === II. RELACIÓN DE LOS HECHOS
        # ======================================================
        y = new_page_if_needed(y, 150)
        sec_num = "II" if self.funcionarios_data else "I"
        y = draw_text(f"{sec_num}. RELACIÓN DE LOS HECHOS Y CONSTANCIAS TÉCNICAS", y, font_size=11, bold=True)
        y -= 5
        comp_label = {"documento": "Documento declarado", "salida": "Salida de Zona Primaria", "laboratorio": "Análisis de Laboratorio"}
        comp_str = comp_label.get(modo_comp, modo_comp)

        # Párrafo 1: introducción — adaptado por tipo de medio
        _tm_legal = self.get_tipo_medio()
        if _tm_legal in ("OLEODUCTO","POLIDUCTO","GASODUCTO"):
            intro_p1 = (
                f"1. El presente informe se elevaría a los fines de que la superioridad determine el "
                f"tratamiento tributario{'e infraccional' if es_denuncia else ''} que pudiere corresponder "
                f"en relación a las diferencias detectadas en el volumen (y/o masa) de "
                f"{'gas natural / hidrocarburos gaseosos' if 'GASODUCTO' in _tm_legal else 'hidrocarburo líquido'} "
                f"transportado por el {tipo_nave.lower()} de referencia, entre lo declarado y lo medido "
                f"según las constancias instrumentales del período."
            )
        elif _tm_legal == "MEDICION ELECTRICA":
            intro_p1 = (
                f"1. El presente informe se elevaría a los fines de que la superioridad determine el "
                f"tratamiento tributario{'e infraccional' if es_denuncia else ''} que pudiere corresponder "
                f"en relación a las diferencias constatadas en la medición de energía eléctrica "
                f"en la instalación de referencia, entre lo declarado y lo registrado en el medidor "
                f"durante el período de control."
            )
        elif "CAMION GAS" in _tm_legal:
            intro_p1 = (
                f"1. El presente informe se elevaría a los fines de que la superioridad determine el "
                f"tratamiento tributario{'e infraccional' if es_denuncia else ''} que pudiere corresponder "
                f"en relación a las diferencias de masa detectadas en la carga de gas licuado / GLP "
                f"transportada en el vehículo cisterna de referencia."
            )
        else:
            intro_p1 = (
                f"1. El presente informe se elevaría a los fines de que la superioridad determine el "
                f"tratamiento tributario{'e infraccional' if es_denuncia else ''} que pudiere corresponder, "
                f"a efectos de salvaguardar la renta fiscal en relación a las diferencias de peso constatadas "
                f"a la finalización de la operación en el {tipo_nave.lower()} de referencia."
            )
        y = draw_wrapped(intro_p1, y, font_size=10, indent=15)
        y -= 5

        # Párrafo 2: producto
        prod_str = producto_desc if producto_desc else "sin especificar"
        pos_str  = f" (Pos. Arancel: {pos_arancel})" if pos_arancel else ""
        y = draw_wrapped(
            f"2. La mercadería objeto de la destinación {dest_num} sería: {prod_str}{pos_str}. "
            f"La documentación declaró {doc_l:,.0f} Litros y {doc_k:,.0f} Kilogramos"
            f"{f', con densidad {dens_doc:.4f} Kg/L' if dens_doc > 0 else ''}.",
            y, font_size=10, indent=15)
        y -= 5

        # Párrafo 3: salidas (si aplica modo salida o laboratorio y hay salidas)
        nro_item = 3
        if modo_comp in ("salida", "laboratorio") and salidas_list:
            y = new_page_if_needed(y, 80)
            y = draw_wrapped(
                f"{nro_item}. Las salidas de Zona Primaria declaradas en la destinación son las siguientes:",
                y, font_size=10, indent=15)
            y -= 3
            # Tabla de salidas
            c.setFont("Helvetica-Bold", 9)
            sx = margin + 25
            c.drawString(sx,       y, "N° Salida")
            c.drawString(sx+90,    y, "Kilos")
            c.drawString(sx+170,   y, "Litros")
            c.drawString(sx+250,   y, "Densidad (Kg/L)")
            y -= 3
            c.line(sx, y, sx+350, y)
            y -= 10
            c.setFont("Helvetica", 9)
            for s in salidas_list:
                c.drawString(sx,    y, str(s['numero']) if s['numero'] else "-")
                c.drawString(sx+90, y, f"{s['kilos']:,.0f}" if s['kilos'] else "-")
                c.drawString(sx+170,y, f"{s['litros']:,.0f}" if s['litros'] else "-")
                c.drawString(sx+250,y, f"{s['densidad']:.4f}" if s['densidad'] else "-")
                y -= 11
            y -= 4
            c.line(sx, y, sx+350, y)
            y -= 13          # bajar suficiente para que el texto quede bajo la línea
            # Total salidas
            c.setFont("Helvetica-Bold", 9)
            c.drawString(sx,     y, "TOTAL")
            c.drawString(sx+90,  y, f"{sal_k_total:,.0f}")
            c.drawString(sx+170, y, f"{sal_l_total:,.0f}")
            y -= 14
            nro_item += 1
            y -= 4

        # Párrafo: sondaje / medición (adaptado por tipo de medio)
        y = new_page_if_needed(y, 60)
        _tm_sond = self.get_tipo_medio()
        if _tm_sond in ("OLEODUCTO","POLIDUCTO","GASODUCTO"):
            sond_txt = (
                f"{nro_item}. Se habría procedido a realizar las lecturas de los contadores de caudal (inicio/fin del período), "
                f"la verificación de los medidores Coriolis (masa directa), el registro de condiciones de línea (T/P), "
                f"el cálculo del factor de compresibilidad Z y la corrección a condiciones base. "
                f"{'La cromatografía de gas (composición % mol) fue registrada según lo indicado en planilla adjunta.' if _tm_sond=='GASODUCTO' else ''} "
                f"Se adjuntan las actas de medición debidamente rubricadas."
            )
        elif _tm_sond == "MEDICION ELECTRICA":
            sond_txt = (
                f"{nro_item}. Se habría procedido a realizar la lectura inicial y final del medidor de energía, "
                f"el registro de tensión, corriente, factor de potencia y demanda máxima, aplicando la constante "
                f"de transformación correspondiente. Se adjuntan las actas de medición debidamente rubricadas."
            )
        else:
            sond_txt = (
                f"{nro_item}. Se habría procedido a realizar el correspondiente sondaje de tanques del "
                f"{tipo_nave.lower()} y la toma de muestras, adjuntándose las planillas de medición debidamente rubricadas."
            )
        y = draw_wrapped(sond_txt, y, font_size=10, indent=15)
        y -= 5
        nro_item += 1

        # Párrafo: explicación modo de control y diferencia
        y = new_page_if_needed(y, 100)
        if modo_comp == "laboratorio":
            if salidas_list:
                sal_desc = f"{sal_k_total:,.0f} Kg ({sal_l_total:,.0f} Litros)" if sal_l_total else f"{sal_k_total:,.0f} Kg"
                txt_modo = (
                    f"{nro_item}. El control se habría efectuado mediante análisis de laboratorio. "
                    f"Las salidas de la Declaración Detallada acumulan {sal_desc}. "
                    f"Al comparar las mediciones realizadas a bordo con la densidad determinada por análisis de laboratorio, "
                    f"se habría constatado una diferencia de {diff_abs_k:,.0f} Kg ({tipo_diff.lower()}) "
                    f"equivalente al {abs(permil_k):.2f}‰ ({pct_max:.2f}%) respecto de las salidas declaradas."
                )
            else:
                txt_modo = (
                    f"{nro_item}. El control se habría efectuado mediante análisis de laboratorio. "
                    f"Al aplicar la densidad determinada en laboratorio sobre las mediciones realizadas a bordo, "
                    f"se habría constatado una diferencia de {diff_abs_k:,.0f} Kg ({tipo_diff.lower()}) "
                    f"equivalente al {abs(permil_k):.2f}‰ ({pct_max:.2f}%) respecto de la carga declarada."
                )
        elif modo_comp == "salida":
            if salidas_list:
                sal_desc = f"{sal_k_total:,.0f} Kg ({sal_l_total:,.0f} Litros)" if sal_l_total else f"{sal_k_total:,.0f} Kg"
                txt_modo = (
                    f"{nro_item}. El control se habría efectuado cotejando las mediciones con las Salidas de Zona Primaria "
                    f"que acumulan {sal_desc}. La diferencia entre las salidas declaradas y las mediciones "
                    f"realizadas a bordo resultaría de {diff_abs_k:,.0f} Kg ({tipo_diff.lower()}), "
                    f"equivalente al {abs(permil_k):.2f}‰ ({pct_max:.2f}%)."
                )
            else:
                txt_modo = (
                    f"{nro_item}. El control se habría efectuado cotejando las mediciones con las Salidas de Zona Primaria. "
                    f"La diferencia entre las salidas declaradas y las mediciones realizadas a bordo resultaría de "
                    f"{diff_abs_k:,.0f} Kg ({tipo_diff.lower()}), equivalente al {abs(permil_k):.2f}‰ ({pct_max:.2f}%)."
                )
        else:  # documento
            txt_modo = (
                f"{nro_item}. El control se habría efectuado confrontando las mediciones realizadas a bordo con la "
                f"cantidad declarada en la documentación ({doc_k:,.0f} Kg / {doc_l:,.0f} Litros). "
                f"La diferencia entre la documentación y las mediciones resultaría de {diff_abs_k:,.0f} Kg ({tipo_diff.lower()}), "
                f"equivalente al {abs(permil_k):.2f}‰ ({pct_max:.2f}%) del total declarado."
            )
        y = draw_wrapped(txt_modo, y, font_size=10, indent=15)
        y -= 5
        nro_item += 1

        # Diferencia resumida (bold)
        y = new_page_if_needed(y, 60)
        y = draw_wrapped(f"{nro_item}. La medición habría arrojado la siguiente diferencia:", y, font_size=10, indent=15)
        c.setFont("Helvetica-Bold", 10)
        y -= 2
        c.drawString(margin + 25, y, f"• Diferencia en Kilos: {diff_abs_k:,.0f} Kg {tipo_diff}S.")
        y -= 13
        c.setFont("Helvetica", 10)
        c.drawString(margin + 25, y, f"• Por mil: {abs(permil_k):.2f} o/oo  |  Porcentaje: {pct_max:.2f}%")
        y -= 18

        y = new_page_if_needed(y, 200)
        y = draw_text("III. ENCUADRE TÉCNICO Y NORMATIVO", y, font_size=11, bold=True)
        y -= 5
        
        if es_denuncia:
            y = draw_wrapped(
                f"La diferencia detectada ({diff_abs_k:,.0f} kg) representa el {pct_max:.2f}% del total de la carga, "
                f"SUPERANDO el umbral del 2%, lo que configura la presunción de infracción al {art_principal}, "
                f"por declaración inexacta de cantidad de mercadería. La diferencia excede la franquicia técnica "
                f"de medición del 6‰ (0,6%) y el umbral de excusa absolutoria del 2%, habilitando la instrucción "
                f"de sumario por infracción.",
                y, font_size=10)
            y -= 8
            y = new_page_if_needed(y, 120)
            c.setFont("Helvetica-Bold", 10)
            c.drawString(margin, y, f"A. Instrucción de Sumario ({art_principal}):")
            y -= 14
            y = draw_wrapped(
                f"Dado que la diferencia supera el 2%, corresponde la instrucción de sumario por declaración inexacta "
                f"al amparo del {art_principal}, sin perjuicio del cargo tributario que proceda por el tributo omitido. "
                f"La excusa absolutoria del 2% prevista en el {art_infraccion} es inaplicable al superarse dicho umbral.",
                y, font_size=10, indent=15)
            y -= 8
        else:
            y = draw_wrapped(
                f"La diferencia detectada ({diff_abs_k:,.0f} kg) representa el {pct_max:.2f}% del total de la carga, "
                f"superando la franquicia técnica del 6‰ (0,6%) pero resultando INFERIOR al 2%, "
                f"por lo cual es de aplicación la excusa absolutoria del {art_infraccion} (no corresponde multa). "
                f"No obstante, en materia TRIBUTARIA rige el principio de integridad: existe un hecho imponible "
                f"que genera el cargo tributario correspondiente.",
                y, font_size=10)
            y -= 8
            y = new_page_if_needed(y, 120)
            c.setFont("Helvetica-Bold", 10)
            c.drawString(margin, y, f"A. Improcedencia de Multa — Procedencia de Cargo ({art_infraccion}):")
            y -= 14
            y = draw_wrapped(
                f"Al ser la diferencia inferior al 2%, resulta aplicable la excusa absolutoria del {art_infraccion}, "
                f"no correspondiendo sanción por infracción de declaración inexacta. Sin embargo, la diferencia "
                f"supera la tolerancia técnica del 0,6%, por lo que existe mercadería que físicamente falta "
                f"y no puede atribuirse a error instrumental, generando el hecho imponible.",
                y, font_size=10, indent=15)
            y -= 8
        
        y = new_page_if_needed(y, 150)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(margin, y, "B. Procedencia del Cargo Tributario (Art. 790 y 791 C.A.):")
        y -= 14
        y = draw_wrapped(
            f"El {tipo_diff.lower()} supera la tolerancia técnica del 0,6% (seis por mil) admitida para mermas "
            f"en transporte acuático de graneles líquidos (Res. ex-ANA 2220/90 y concordantes). "
            f"Dicha mercadería no arribó físicamente. La tolerancia del 2% es una ficción jurídica penal "
            f"(Art. {art_infraccion}). La tolerancia del 0,6% es técnica. Todo lo que supere el 0,6% y no llegue "
            f"al 2% es mercadería que físicamente falta y no puede atribuirse a error de instrumento. "
            f"Por ende, existe un Hecho Imponible (Arts. 637/638 C.A.).",
            y, font_size=10, indent=15)
        y -= 8
        
        y = new_page_if_needed(y, 100)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(margin, y, "C. Marco Normativo Aplicable:")
        y -= 14
        normas = [
            f"• {art_principal}: Infracción por declaración inexacta de cantidad.",
            "• Arts. 637/638 C.A.: Determinación y nacimiento de la obligación tributaria.",
            "• Arts. 790/791 C.A.: Liquidación de tributos y actuación aduanera.",
            "• Arts. 164/390 C.A.: Presunción de destino de mercadería faltante.",
            "• Res. ex-ANA 2220/90: Tolerancia técnica 0,6% para graneles líquidos.",
        ]
        for n in normas:
            y = draw_wrapped(n, y, font_size=9, indent=15)
            y -= 2
        y -= 6
        
        # === IV. JURISPRUDENCIA ===
        y = new_page_if_needed(y, 200)
        y = draw_text("IV. JURISPRUDENCIA Y ANTECEDENTES", y, font_size=11, bold=True)
        y -= 5
        
        juris = [
            ('CASO "ESSO": Autonomía del Tributo respecto de la Multa.',
             'Autos: "Esso S.A.P.A. c/ DGA s/ Recurso de Apelación" — Tribunal Fiscal de la Nación, Sala E.',
             'El TFN confirmó que las franquicias del Art. 959 del C.A. (2%) operan exclusivamente en el ámbito '
             'infraccional, pero no liberan del pago de tributos por mercadería faltante que exceda la merma técnica.'),
            ('CASO "SHELL": La Inexistencia de Tolerancia Fiscal.',
             'Autos: "Shell C.A.P.S.A. c/ DGA" — Tribunal Fiscal de la Nación, Sala F.',
             'Se estableció que la tolerancia del 2% del Art. 959 es una excusa absolutoria penal. '
             'En materia tributaria rige el principio de integridad. Si la merma supera lo técnicamente '
             'aceptable (0,6%), la diferencia está sujeta a tributación.'),
            ('FALLO CORTE SUPREMA (CSJN): "I.B.M. World Trade Corp."',
             'Referencia: Fallos 303:1363.',
             'La obligación tributaria nace del hecho imponible objetivo y es independiente de la intención '
             'dolosa o culposa del operador.')
        ]
        for titulo, ref, doctrina in juris:
            y = new_page_if_needed(y, 80)
            c.setFont("Helvetica-Bold", 9)
            c.drawString(margin + 10, y, titulo)
            y -= 12
            c.setFont("Helvetica-Oblique", 8)
            c.drawString(margin + 20, y, ref)
            y -= 12
            y = draw_wrapped(doctrina, y, font_size=9, indent=20)
            y -= 8
        
        # === V. LIQUIDACIÓN ESTIMADA ===
        y = new_page_if_needed(y, 200)
        if es_denuncia:
            sec_title = "V. BASE IMPONIBLE ESTIMADA — DENUNCIA (total de la carga)"
        else:
            sec_title = "V. LIQUIDACIÓN ESTIMADA DE CARGO (sobre el faltante/sobrante)"
        y = draw_text(sec_title, y, font_size=11, bold=True)
        y -= 5

        moneda_sym = {"ARS": "$", "USD": "U$S", "EUR": "€", "BRL": "R$", "GBP": "£"}.get(divisa, "$")

        lbl_tipo = "FALTANTE" if es_faltante else "SOBRANTE"
        # valor_litro en el DDT se usa como $/kg para el cargo (se trabaja todo en kilos)
        _vl_var = ddt_obj.get("valor_litro")
        valor_kg = self.parse_float(_vl_var.get() if isinstance(_vl_var, tk.StringVar) else str(_vl_var or "0"))
        divisa_sym = {"ARS": "$", "USD": "U$S", "EUR": "€", "BRL": "R$", "GBP": "£"}.get(divisa, "$")

        if es_denuncia:
            base_kilos = doc_k if doc_k > 0 else diff_abs_k
            y = draw_wrapped(
                f"En el presente caso se habría constatado un {lbl_tipo} de {diff_abs_k:,.0f} kilogramos, "
                f"representando el {pct_max:.2f}% del total declarado ({base_kilos:,.0f} kg). "
                f"Por tratarse de una denuncia, la base imponible estimada correspondería a la TOTALIDAD "
                f"de la mercadería objeto de la destinación ({base_kilos:,.0f} kg), "
                f"sin perjuicio de la liquidación que practicara el área de valoración.",
                y, font_size=10)
            y -= 8
            if valor_kg > 0:
                usa_tc_pdf = (divisa != "ARS") and (tipo_cambio > 0)
                base_imp_orig = base_kilos * valor_kg          # en divisa original
                base_imp_ars  = base_imp_orig * tipo_cambio if usa_tc_pdf else base_imp_orig
                if usa_tc_pdf:
                    y = draw_text(f"T/C: {divisa_sym} 1 = $ {tipo_cambio:,.2f}  →  precio ARS: {divisa_sym} {valor_kg:,.4f} × {tipo_cambio:,.2f} = $ {base_imp_orig * tipo_cambio / base_kilos:,.4f}/Kg", y, font_size=9)
                    y -= 4
                    y = draw_text(f"Total declarado: {base_kilos:,.0f} Kg × {divisa_sym} {valor_kg:,.4f}/Kg = {divisa_sym} {base_imp_orig:,.2f}  →  $ {base_imp_ars:,.2f}", y, font_size=10, bold=True)
                else:
                    y = draw_text(f"Total declarado: {base_kilos:,.0f} Kg  ×  $ {valor_kg:,.4f}/Kg = $ {base_imp_ars:,.2f}", y, font_size=10, bold=True)
                y -= 6
                y = draw_wrapped("Los valores a continuación serían estimativos. La liquidación definitiva la practicaría el área de valoración.", y, font_size=9)
                y -= 8
                tributos = self.get_tributos_activos() or [("Derechos de Importación/Exportación (ref.)", 8.0), ("Anticipo Ganancias (ref.)", 0.5)]
                c.setFont("Helvetica-Bold", 9)
                c.drawString(margin, y, "CONCEPTO"); c.drawString(270, y, "ALÍC."); c.drawString(345, y, "BASE (ARS)"); c.drawString(475, y, "IMPORTE EST. (ARS)")
                y -= 3; c.line(margin, y, w - margin, y); y -= 14
                c.setFont("Helvetica", 9)
                total_trib = 0
                for conc, alic in tributos:
                    importe = base_imp_ars * alic / 100; total_trib += importe
                    c.drawString(margin, y, conc); c.drawString(270, y, f"{alic}%")
                    c.drawString(345, y, f"$ {base_imp_ars:,.2f}"); c.drawString(475, y, f"$ {importe:,.2f}")
                    y -= 14
                c.setFont("Helvetica-Bold", 10)
                c.drawString(margin, y, "TOTAL ESTIMADO SOBRE TOTAL CARGA (referencia):")
                c.drawString(475, y, f"$ {total_trib:,.2f}")
                y -= 20; c.line(margin, y, w - margin, y); y -= 10
                # Equivalencia en divisa extranjera (solo si hay T/C cargado)
                if tipo_cambio > 0 and divisa != "ARS":
                    # ya mostramos en ARS arriba; agregamos nota en divisa original
                    c.setFont("Helvetica", 9)
                    c.drawString(margin, y, f"Monto en {divisa}: {divisa_sym} {total_trib / tipo_cambio:,.2f}  (T/C: 1 {divisa} = $ {tipo_cambio:,.2f})")
                    y -= 14
                elif tipo_cambio > 0 and divisa == "ARS":
                    # Precio en ARS, mostramos equiv en USD usando el T/C del documento
                    c.setFont("Helvetica", 9)
                    c.drawString(margin, y, f"Equiv. est. en USD: U$S {total_trib / tipo_cambio:,.2f}  (T/C ref.: 1 U$S = $ {tipo_cambio:,.2f})")
                    y -= 14
            else:
                y = draw_text("(Sin valor unitario — liquidación pendiente de valoración oficial)", y, font_size=10)
        else:
            # CARGO: base = solo el faltante/sobrante en KILOS
            y = draw_wrapped(
                f"En el presente caso se habría constatado un {lbl_tipo} de {diff_abs_k:,.0f} kilogramos "
                f"({pct_max:.2f}% del total declarado). El cargo tributario correspondería ÚNICAMENTE "
                f"a la diferencia detectada de {diff_abs_k:,.0f} kg, conforme Arts. 790/791 C.A.",
                y, font_size=10)
            y -= 8
            if valor_kg > 0:
                usa_tc_pdf = (divisa != "ARS") and (tipo_cambio > 0)
                base_imp_orig = diff_abs_k * valor_kg
                base_imp_ars  = base_imp_orig * tipo_cambio if usa_tc_pdf else base_imp_orig
                if usa_tc_pdf:
                    y = draw_text(f"T/C: {divisa_sym} 1 = $ {tipo_cambio:,.2f}  →  precio ARS: $ {valor_kg * tipo_cambio:,.4f}/Kg", y, font_size=9)
                    y -= 4
                    y = draw_text(f"{lbl_tipo}: {diff_abs_k:,.0f} Kg × {divisa_sym} {valor_kg:,.4f}/Kg = {divisa_sym} {base_imp_orig:,.2f}  →  $ {base_imp_ars:,.2f}", y, font_size=10, bold=True)
                else:
                    y = draw_text(f"{lbl_tipo}: {diff_abs_k:,.0f} Kg  ×  $ {valor_kg:,.4f}/Kg = $ {base_imp_ars:,.2f}", y, font_size=10, bold=True)
                y -= 6
                y = draw_wrapped("Los valores a continuación serían estimativos. La liquidación definitiva la practicaría el área de valoración.", y, font_size=9)
                y -= 8
                tributos = self.get_tributos_activos() or [("Derechos de Importación/Exportación (ref.)", 8.0), ("Anticipo Ganancias (ref.)", 0.5)]
                c.setFont("Helvetica-Bold", 9)
                c.drawString(margin, y, "CONCEPTO"); c.drawString(270, y, "ALÍC."); c.drawString(345, y, "BASE (ARS)"); c.drawString(475, y, "IMPORTE EST. (ARS)")
                y -= 3; c.line(margin, y, w - margin, y); y -= 14
                c.setFont("Helvetica", 9)
                total_trib = 0
                for conc, alic in tributos:
                    importe = base_imp_ars * alic / 100; total_trib += importe
                    c.drawString(margin, y, conc); c.drawString(270, y, f"{alic}%")
                    c.drawString(345, y, f"$ {base_imp_ars:,.2f}"); c.drawString(475, y, f"$ {importe:,.2f}")
                    y -= 14
                c.setFont("Helvetica-Bold", 10)
                c.drawString(margin, y, f"TOTAL ESTIMADO CARGO SOBRE {lbl_tipo} (referencia):")
                c.drawString(475, y, f"$ {total_trib:,.2f}")
                y -= 20; c.line(margin, y, w - margin, y); y -= 10
                if tipo_cambio > 0 and divisa != "ARS":
                    c.setFont("Helvetica", 9)
                    c.drawString(margin, y, f"Monto en {divisa}: {divisa_sym} {total_trib / tipo_cambio:,.2f}  (T/C: 1 {divisa} = $ {tipo_cambio:,.2f})")
                    y -= 14
                elif tipo_cambio > 0 and divisa == "ARS":
                    c.setFont("Helvetica", 9)
                    c.drawString(margin, y, f"Equiv. est. en USD: U$S {total_trib / tipo_cambio:,.2f}  (T/C ref.: 1 U$S = $ {tipo_cambio:,.2f})")
                    y -= 14
            else:
                y = draw_text("(Sin valor unitario — liquidación pendiente de valoración oficial)", y, font_size=10)

        # === VI. CONCLUSIÓN ===
        y = new_page_if_needed(y, 120)
        y -= 10
        y = draw_text("VI. CONCLUSIÓN", y, font_size=11, bold=True)
        y -= 5
        lbl_tipo_final = "FALTANTE" if es_faltante else "SOBRANTE"
        if es_denuncia:
            y = draw_wrapped(
                f"En mérito a lo expuesto, se elevarían las presentes actuaciones a fin de que la superioridad "
                f"evaluara la posible instrucción de sumario al amparo del {art_principal}, en virtud del "
                f"{lbl_tipo_final} detectado de {diff_abs_k:,.0f} kilogramos ({pct_max:.2f}% del total declarado), "
                f"sin perjuicio del cargo tributario que eventualmente correspondiere. "
                f"Es todo cuanto se informaría.",
                y, font_size=10)
        else:
            y = draw_wrapped(
                f"En mérito a lo expuesto, se elevarían las presentes actuaciones a fin de que la superioridad "
                f"evaluara la procedencia del cargo tributario en los términos del {art_principal}, "
                f"en virtud del {lbl_tipo_final} detectado de {diff_abs_k:,.0f} kilogramos "
                f"({pct_max:.2f}% del total declarado). Es todo cuanto se informaría.",
                y, font_size=10)
        y -= 15

        try:
            c.showPage()
            if shared_canvas is None:
                c.save()
        except Exception as _pdf_err:
            messagebox.showerror("Error PDF Cargo", f"Error al finalizar PDF de cargo:\n{_pdf_err}")
            return None
        # Si se usó canvas compartido, las páginas ya están en él, no guardar archivo
        if shared_canvas is not None:
            return None
        # Si es un buffer override (embed en PDF unificado), no guardar archivo
        if _buf_override is not None:
            return None
        # Escribir archivo del cargo (standalone)
        _buf.seek(0)
        _data = _buf.read()
        if full_path:
            try:
                with open(full_path, "wb") as f_out:
                    f_out.write(_data)
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo guardar {filename}: {e}")
                return None
            return full_path
        return None

    def auto_save_loop(self):
        clean_ddts = []
        for d in self.ddt_data:
            c_d = {
                "numero": d["numero"].get(),
                "num_planilla": d["num_planilla"].get(),
                "tipo_doc": d["tipo_doc"].get(),
                "producto": d["producto"].get(),
                "pos_arancel": d["pos_arancel"].get(),
                "densidad": d["densidad"].get(),
                "litros": d["litros"].get(),
                "kilos": d["kilos"].get(),
                "valor_litro": d["valor_litro"].get() if "valor_litro" in d else "",
                "divisa": d["divisa"].get() if "divisa" in d else "ARS",
                "divisa_desc": d["divisa_desc"].get() if "divisa_desc" in d else "",
                "tipo_cambio": d["tipo_cambio"].get() if "tipo_cambio" in d else "",
                "salidas": []
            }
            for _ak in self.DDT_ACTOR_KEYS:
                c_d[_ak] = d[_ak].get() if _ak in d else ""
            for s in d["salidas"]:
                c_d["salidas"].append({
                    "numero": s["numero"].get(),
                    "litros": s["litros"].get(),
                    "kilos": s["kilos"].get(),
                    "densidad": s["densidad"].get()
                })
            clean_ddts.append(c_d)
        data = {k: v.get() for k, v in self.vars.items() if "tkinter" not in k and "<module" not in k}
        data["_ddt_struct"] = clean_ddts
        data["_tank_list"] = self.lista_tanques
        data["_carb_list"] = self.lista_carbonera
        data["_funcionarios"] = [{"cuil": f["cuil"].get(), "legajo": f["legajo"].get(), "apellido": f["apellido"].get(), "nombre": f["nombre"].get(), "funcion": f["funcion"].get()} for f in self.funcionarios_data]
        with open("autosave.json", 'w') as f: json.dump(data, f)
        self.root.after(30000, self.auto_save_loop)
    
    def guardar_datos_manual(self):
        _buque_meg = self.clean_filename(self.get_var('car_buque').get()) or 'medicion'
        _fecha_meg = datetime.now().strftime('%Y%m%d')
        f = filedialog.asksaveasfilename(defaultextension=".meg", initialfile=f"{_buque_meg}_{_fecha_meg}", filetypes=[("Archivos MEG", "*.meg"), ("Todos", "*.*")])
        if f: 
            clean_ddts = []
            for d in self.ddt_data:
                c_d = {
                    "numero": d["numero"].get(),
                    "num_planilla": d["num_planilla"].get(),
                    "tipo_doc": d["tipo_doc"].get(),
                    "producto": d["producto"].get(),
                    "pos_arancel": d["pos_arancel"].get(),
                    "densidad": d["densidad"].get(),
                    "litros": d["litros"].get(),
                    "kilos": d["kilos"].get(),
                    "valor_litro": d["valor_litro"].get() if "valor_litro" in d else "",
                    "divisa": d["divisa"].get() if "divisa" in d else "ARS",
                    "divisa_desc": d["divisa_desc"].get() if "divisa_desc" in d else "",
                    "tipo_cambio": d["tipo_cambio"].get() if "tipo_cambio" in d else "",
                    "salidas": []
                }
                for _ak in self.DDT_ACTOR_KEYS:
                    c_d[_ak] = d[_ak].get() if _ak in d else ""
                for s in d["salidas"]:
                    c_d["salidas"].append({
                        "numero": s["numero"].get(),
                        "litros": s["litros"].get(),
                        "kilos": s["kilos"].get(),
                        "densidad": s["densidad"].get()
                    })
                clean_ddts.append(c_d)
            data = {k: v.get() for k, v in self.vars.items() if "tkinter" not in k and "<module" not in k}
            data["_ddt_struct"] = clean_ddts
            data["_tipo_medio"] = self.get_tipo_medio()
            data["_tank_list"] = self.lista_tanques
            data["_carb_list"] = self.lista_carbonera
            data["_funcionarios"] = [{"cuil": f["cuil"].get(), "legajo": f["legajo"].get(), "apellido": f["apellido"].get(), "nombre": f["nombre"].get(), "funcion": f["funcion"].get()} for f in self.funcionarios_data]
            with open(f, 'w') as file: json.dump(data, file)
        
    def cargar_datos(self):
        f = filedialog.askopenfilename(filetypes=[("Archivos MEG", "*.meg"), ("Todos", "*.*")])
        if not f: return
        try:
            with open(f, 'r') as file:
                data = json.load(file)
            # Limpiar vars antiguas
            for key in list(self.vars.keys()): self.vars[key].set("")
            self.vars.clear()
            for k,v in data.items():
                if k.startswith("_"): continue
                if "tkinter" in k or "<module" in k: continue
                self.get_var(k).set(v)
            if "_tipo_medio" in data:
                tm = data["_tipo_medio"]
                self.get_var("car_tipo_medio").set(tm)
                self.get_var("car_tipo_nave").set(tm)
                # LOCK: guardar tipo original para impedir cambio de categoría
                self._archivo_tipo_medio = tm
                self._archivo_bloqueado  = True
            # Tank lists
            if "_tank_list" in data:  self.lista_tanques  = data["_tank_list"]
            if "_carb_list"  in data: self.lista_carbonera = data["_carb_list"]
            # Rebuild UI PRIMERO - construir_caratula crea func_stack y ddt_stack frescos
            # Rebuild UI
            for w in self.tab_caratula.winfo_children(): w.destroy()
            self.combos_ddt = []
            self.construir_caratula()
            self.rebuild_all_tabs()
            self._apply_tipo_lock()  # apply lock after rebuild
            # Funcionarios - DESPUES de construir_caratula (func_stack ya existe)
            # Funcionarios
            if "_funcionarios" in data:
                self.funcionarios_data = []
                self.func_counter = 0
                for fc in data["_funcionarios"]:
                    self.agregar_funcionario_row(data=fc)
            # DDT - DESPUES de construir_caratula (ddt_stack ya existe)
            # DDT data
            for d in self.ddt_data[:]: d["main_frame"].destroy()
            self.ddt_data = []
            self.ddt_counter = 0
            if "_ddt_struct" in data:
                for ddt_d in data["_ddt_struct"]:
                    self.agregar_ddt_row(data=ddt_d)
            elif "_ddt_list" in data:  # compatibilidad con archivos viejos
                for ddt_d in data["_ddt_list"]:
                    self.agregar_ddt_row(data=ddt_d)
            elif not self.ddt_data:
                self.agregar_ddt_row(def_prod="GASOIL")
        except Exception as e:
            traceback.print_exc()
            messagebox.showerror("Error", f"No se pudo cargar el archivo:\n{e}")

    def _apply_tipo_lock(self):
        """Bloquea el combobox de tipo_medio a la categoría de la medición actual.
        Ej: BUQUE solo puede cambiar entre tipos marítimos; TANQUE solo entre tierra, etc."""
        if not getattr(self, "_archivo_bloqueado", False): return
        tm_locked = getattr(self, "_archivo_tipo_medio", "")
        allowed = self.CATEGORIA_TIPOS.get(tm_locked, [tm_locked])
        try:
            for w in self.tab_caratula.winfo_children():
                for w2 in self._find_all_comboboxes(w):
                    try:
                        curr_var = str(self.get_var("car_tipo_medio"))
                        if hasattr(w2, "cget") and w2.cget("textvariable") == curr_var:
                            w2.config(values=allowed)
                            if len(allowed) <= 1:
                                w2.config(state="disabled")
                            else:
                                w2.config(state="readonly")
                    except: pass
        except: pass

    def _find_all_comboboxes(self, widget):
        result = []
        if isinstance(widget, ttk.Combobox): result.append(widget)
        try:
            for child in widget.winfo_children():
                result.extend(self._find_all_comboboxes(child))
        except: pass
        return result

